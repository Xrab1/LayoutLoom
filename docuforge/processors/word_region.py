"""Rebuild a LayoutLoom fixed-layout DOCX with page-relative editable regions.

The PDF-to-Word converter remains the source of truth.  This module is a
separate Word-to-Word postprocessor: it consumes the converter's positioned
line paragraphs, groups them into paragraph/column regions, and writes a new
document whose page contains only a small number of modern DrawingML text
boxes and page-relative visual crops.  Text inside every box is made from
ordinary Word paragraphs, so copying a paragraph no longer depends on a chain
of independently positioned ``w:framePr`` lines.
"""

from __future__ import annotations

import io
import re
import shutil
import statistics
import tempfile
import unicodedata
import weakref
from collections import Counter
from dataclasses import dataclass, field
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any, Callable, Iterable, Mapping

from docuforge.models import MissingEngineError, ValidationError
from docuforge.utils import atomic_output, unique_path

from .word_flow import (
    _FrameLine,
    _MergedLine,
    _VisualRegion,
    _assign_column,
    _background_color,
    _can_merge_lines,
    _connected_boxes,
    _heading_line,
    _line_break_hyphen_is_soft,
    _merge_frame_rows,
    _remove_last_hyphen,
    _source_pages,
    _text_joiner,
    _two_column_layout,
    _visual_regions,
    _word_counter,
)


_WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
_WPS_GRAPHIC_URI = _WPS_NS
_PAGE_NAME = re.compile(r"^(?:LayoutLoom|DocuForge) region page (\d+) ")
_VISUAL_NAME = re.compile(
    r"^(?:LayoutLoom|DocuForge) region page (\d+) visual (\d+)(?: ([a-z_]+))?$"
)
_JOINED_ENGLISH = re.compile(r"[A-Za-z]{7,}")
_INTEGRITY_ENGLISH_WORD = re.compile(r"[A-Za-z]+(?:[-'][A-Za-z]+)*")
_CJK = r"\u3400-\u9fff"
_MIN_CHARACTER_RECALL = 0.995
_MIN_CHARACTER_PRECISION = 0.995
_MIN_SEQUENCE_COVERAGE = 0.985
_MIN_ENGLISH_WORD_RECALL = 0.985
_MIN_ADJACENT_WORD_RECALL = 0.970
_MIN_PRECISE_FRAME_COVERAGE = 0.80
_MIN_PRECISE_FRAME_CONFLICT = 0.08
_BACKGROUND_SAFE_VISUAL_PREFIX = "background_safe_"
_DENSE_BACKGROUND_HINT_THRESHOLD = 24

_PreciseVisualHint = tuple[str, int, int, int, int]


class _WorkspaceCleanupToken:
    """Weak-reference owner for exception-safe temporary workspace cleanup."""


def _progress_reporter(
    callback: Callable[[float, str], None] | None,
) -> Callable[[float, str], None]:
    if callback is not None:
        return callback
    from docuforge.runner import report_progress

    return report_progress


@dataclass
class _TextBlock:
    lines: list[_MergedLine]
    x: int
    y: int
    width: int
    height: int
    column: str
    font_size: float
    heading: bool
    precision: bool = False

    @property
    def x1(self) -> int:
        return self.x + self.width

    @property
    def y1(self) -> int:
        return self.y + self.height

    @property
    def center_y(self) -> float:
        return self.y + self.height / 2


@dataclass
class _TextRegion:
    blocks: list[_TextBlock]
    x: int
    y: int
    width: int
    height: int
    column: str
    precision: bool = False

    @property
    def x1(self) -> int:
        return self.x + self.width

    @property
    def y1(self) -> int:
        return self.y + self.height

    @property
    def center_y(self) -> float:
        return self.y + self.height / 2


@dataclass
class _RegionPage:
    index: int
    width: int
    height: int
    text_regions: list[_TextRegion] = field(default_factory=list)
    visual_regions: list[_VisualRegion] = field(default_factory=list)
    precise_audit_regions: list[tuple[int, int, int, int]] = field(
        default_factory=list
    )


@dataclass(frozen=True)
class RegionWordInspection:
    source_pages: int
    source_frames: int
    output_text_regions: int
    output_visual_regions: int
    output_paragraphs: int
    source_words: int
    retained_words: int
    word_recall: float
    structural_passed: bool
    rendered_passed: bool
    used_correction_pass: bool
    character_recall: float = 1.0
    character_precision: float = 1.0
    sequence_coverage: float = 1.0
    english_word_recall: float = 1.0
    adjacent_word_recall: float = 1.0


@dataclass(frozen=True)
class _TextIntegrityMetrics:
    source_tokens: int
    retained_tokens: int
    token_recall: float
    source_characters: int
    retained_characters: int
    character_recall: float
    character_precision: float
    sequence_coverage: float
    english_word_recall: float
    adjacent_word_recall: float


@dataclass(frozen=True)
class _RenderAudit:
    passed: bool
    reason: str = ""
    page_count: int = 0
    text_recall: float = 0.0
    worst_mean_difference: float = 0.0
    worst_changed_fraction: float = 0.0


def _restore_english_word_boundaries(text: str) -> str:
    """Restore spaces removed by PDF text extraction without changing real words."""

    try:
        import wordninja
    except ImportError:
        return text

    def replace(match: re.Match[str]) -> str:
        value = match.group(0)
        if value.isupper():
            return value
        parts = wordninja.split(value)
        if len(parts) < 2:
            return value
        if "".join(parts).casefold() != value.casefold():
            return value
        return " ".join(parts)

    return _JOINED_ENGLISH.sub(replace, text)


def _normalize_mixed_script_text(text: str) -> str:
    """Normalize visible Chinese/Latin boundaries without touching identifiers."""

    normalized = re.sub(
        rf"(?<=[{_CJK}])\s*:\s*(?=[A-Za-z0-9])",
        "：",
        text,
    )
    normalized = re.sub(
        rf"(?<=[A-Za-z0-9])(?=[{_CJK}])",
        " ",
        normalized,
    )
    normalized = re.sub(
        rf"(?<=[{_CJK}])(?=[A-Za-z0-9])",
        " ",
        normalized,
    )
    return normalized


def _normalized_source_text(text: str) -> str:
    return _normalize_mixed_script_text(_restore_english_word_boundaries(text))


def _normalize_text_boundary(previous: str, current: str) -> str:
    """Apply the same mixed-script rule when a boundary crosses Word runs."""

    if not previous or not current or previous[-1].isspace() or current[0].isspace():
        return current
    combined = _normalize_mixed_script_text(previous[-1] + current)
    if combined.startswith(previous[-1]):
        return combined[1:]
    return current


def _normalized_frame_runs(frame: _FrameLine) -> list[tuple[Any, str]]:
    """Normalize one visual frame while retaining its original run formatting."""

    runs = list(frame.paragraph.runs)
    if not runs:
        return []
    original_parts = [str(run.text or "") for run in runs]
    original = "".join(original_parts)
    normalized = _normalized_source_text(original)
    if normalized == original:
        return list(zip(runs, original_parts))
    if not original:
        return list(zip(runs, original_parts))

    owners: list[int] = []
    for index, part in enumerate(original_parts):
        owners.extend([index] * len(part))
    output: list[list[str]] = [[] for _run in runs]
    matcher = SequenceMatcher(None, original, normalized, autojunk=False)
    for tag, source_start, source_end, target_start, target_end in matcher.get_opcodes():
        replacement = normalized[target_start:target_end]
        if tag == "delete":
            continue
        if tag == "equal":
            for offset, character in enumerate(replacement):
                output[owners[source_start + offset]].append(character)
            continue
        if tag == "insert":
            if source_start > 0:
                owner = owners[source_start - 1]
            elif source_start < len(owners):
                owner = owners[source_start]
            else:
                owner = 0
            output[owner].append(replacement)
            continue
        source_length = max(1, source_end - source_start)
        target_length = max(1, len(replacement))
        for offset, character in enumerate(replacement):
            source_offset = min(
                source_length - 1,
                int(offset * source_length / target_length),
            )
            output[owners[source_start + source_offset]].append(character)
    return [(run, "".join(parts)) for run, parts in zip(runs, output)]


def _overlap_fraction(first: Any, second: Any) -> float:
    x0 = max(first.x, second.x)
    y0 = max(first.y, second.y)
    x1 = min(first.x1, second.x1)
    y1 = min(first.y1, second.y1)
    if x1 <= x0 or y1 <= y0:
        return 0.0
    intersection = (x1 - x0) * (y1 - y0)
    minimum = min(first.width * first.height, second.width * second.height)
    return intersection / max(1, minimum)


def _hint_intersection_fraction(
    item: Any,
    hint: tuple[int, int, int, int],
) -> float:
    x0 = max(item.x, hint[0])
    y0 = max(item.y, hint[1])
    x1 = min(item.x1, hint[2])
    y1 = min(item.y1, hint[3])
    if x1 <= x0 or y1 <= y0:
        return 0.0
    intersection = (x1 - x0) * (y1 - y0)
    hint_area = max(1, (hint[2] - hint[0]) * (hint[3] - hint[1]))
    item_area = max(1, item.width * item.height)
    return intersection / min(hint_area, item_area)


def _center_inside_hint(item: Any, hint: tuple[int, int, int, int]) -> bool:
    center_x = item.x + item.width / 2.0
    center_y = item.y + item.height / 2.0
    return hint[0] <= center_x <= hint[2] and hint[1] <= center_y <= hint[3]


def _block_from_lines(lines: list[_MergedLine], median_size: float) -> _TextBlock:
    x0 = min(line.x for line in lines)
    y0 = min(line.y for line in lines)
    x1 = max(line.x1 for line in lines)
    y1 = max(line.y1 for line in lines)
    dominant = max(lines, key=lambda line: (len(line.text.strip()), line.font_size))
    return _TextBlock(
        list(lines),
        x0,
        y0,
        max(20, x1 - x0),
        max(20, y1 - y0),
        dominant.column,
        dominant.font_size,
        len(lines) == 1 and _heading_line(dominant, median_size),
        any(bool(getattr(line, "precision", False)) for line in lines),
    )


def _ordered_lines(lines: list[_MergedLine], page_width: int) -> list[_MergedLine]:
    two_columns = _two_column_layout(lines, page_width)
    for line in lines:
        line.column = _assign_column(line, page_width, two_columns)
    if not two_columns:
        return sorted(lines, key=lambda item: (item.y, item.x))

    full = sorted(
        (line for line in lines if line.column == "full"),
        key=lambda item: (item.y, item.x),
    )
    remaining = [line for line in lines if line.column != "full"]
    ordered: list[_MergedLine] = []
    for separator in full:
        before = [line for line in remaining if line.center_y < separator.center_y]
        for column in ("left", "right"):
            ordered.extend(
                sorted(
                    (line for line in before if line.column == column),
                    key=lambda item: (item.y, item.x),
                )
            )
        before_ids = {id(line) for line in before}
        remaining = [line for line in remaining if id(line) not in before_ids]
        ordered.append(separator)
    for column in ("left", "right"):
        ordered.extend(
            sorted(
                (line for line in remaining if line.column == column),
                key=lambda item: (item.y, item.x),
            )
        )
    return ordered


def _coalesce_same_row_fragments(
    lines: list[_MergedLine], page_width: int
) -> list[_MergedLine]:
    """Join horizontally separated fragments that belong to one visual row.

    Forms and resumes often store a label and its value as separate frame
    clusters with the same baseline.  Treating those clusters as independent
    regions creates overlapping boxes.  A single row with tab-aligned runs is
    both more faithful and substantially easier to copy.
    """

    if not lines:
        return []
    two_columns = _two_column_layout(lines, page_width)
    for line in lines:
        line.column = _assign_column(line, page_width, two_columns)
    rows: list[list[_MergedLine]] = []

    def row_match_score(
        row: list[_MergedLine],
        line: _MergedLine,
    ) -> tuple[float, float] | None:
        first = row[0]
        maximum_height = max(line.height, *(item.height for item in row))
        center_delta = abs(
            statistics.median(item.center_y for item in row) - line.center_y
        )
        tolerance = max(45, round(maximum_height * 0.42))
        first_precision = bool(getattr(first, "precision", False))
        line_precision = bool(getattr(line, "precision", False))
        row_x0 = min(item.x for item in row)
        row_x1 = max(item.x1 for item in row)
        row_y0 = min(item.y for item in row)
        row_y1 = max(item.y1 for item in row)
        vertical_overlap = min(row_y1, line.y1) - max(row_y0, line.y)
        minimum_height = max(1, min(row_y1 - row_y0, line.height))
        horizontal_separation = max(
            line.x - row_x1,
            row_x0 - line.x1,
        )
        same_column = first.column == line.column
        compatible_precision_columns = bool(
            same_column
            or first.column in {"full", "main"}
            or line.column in {"full", "main"}
        )
        ordinary_same_row = bool(
            same_column
            and not first_precision
            and not line_precision
            and center_delta <= tolerance
        )
        precision_same_row = bool(
            first_precision
            and line_precision
            and compatible_precision_columns
            and -20
            <= horizontal_separation
            <= max(900, round(page_width * 0.14))
            and (
                vertical_overlap / minimum_height >= 0.22
                or center_delta <= max(55, round(maximum_height * 0.75))
            )
            and max(row_y1, line.y1) - min(row_y0, line.y)
            <= max(80, round(maximum_height * 1.72))
        )
        if not (ordinary_same_row or precision_same_row):
            return None
        return (
            vertical_overlap / minimum_height,
            -center_delta,
        )

    for line in sorted(lines, key=lambda item: (item.y, item.x)):
        matches = [
            (score, index)
            for index, row in enumerate(rows)
            if (score := row_match_score(row, line)) is not None
        ]
        if matches:
            _score, row_index = max(matches, key=lambda item: item[0])
            rows[row_index].append(line)
        else:
            rows.append([line])

    output: list[_MergedLine] = []
    for row in sorted(
        rows,
        key=lambda items: (
            min(item.y for item in items),
            min(item.x for item in items),
        ),
    ):
        if len(row) == 1:
            output.append(row[0])
            continue
        frames = sorted(
            (frame for line in row for frame in line.frames),
            key=lambda item: item.x,
        )
        x0 = min(line.x for line in row)
        y0 = min(line.y for line in row)
        x1 = max(line.x1 for line in row)
        y1 = max(line.y1 for line in row)
        text = ""
        previous: _FrameLine | None = None
        for frame in frames:
            if previous is not None:
                text += _text_joiner(
                    previous.text,
                    frame.text,
                    frame.x - previous.x1,
                    max(previous.font_size, frame.font_size),
                )
            text += frame.text
            previous = frame
        dominant = max(
            row,
            key=lambda item: (len(item.text.strip()), item.font_size),
        )
        largest_font_size = max(
            [item.font_size for item in row]
            + [frame.font_size for frame in frames]
        )
        merged_column = dominant.column
        if any(item.column == "full" for item in row):
            merged_column = "full"
        elif any(item.column == "main" for item in row):
            merged_column = "main"
        merged = _MergedLine(
            frames,
            x0,
            y0,
            max(20, x1 - x0),
            max(20, y1 - y0),
            text,
            largest_font_size,
            dominant.bold,
            merged_column,
        )
        if any(bool(getattr(item, "precision", False)) for item in row):
            setattr(merged, "precision", True)
        output.append(merged)
    return output


def _paragraph_blocks(lines: list[_MergedLine], page_width: int) -> list[_TextBlock]:
    if not lines:
        return []
    lines = _coalesce_same_row_fragments(lines, page_width)
    median_size = statistics.median(line.font_size for line in lines)
    ordered = _ordered_lines(lines, page_width)
    blocks: list[_TextBlock] = []
    active: list[_MergedLine] = []
    for line in ordered:
        if bool(getattr(line, "precision", False)):
            if active:
                blocks.append(_block_from_lines(active, median_size))
                active = []
            blocks.append(_block_from_lines([line], median_size))
            continue
        missing_visual_line = bool(
            active
            and line.y - active[-1].y
            > max(active[-1].height, line.height) * 1.62
        )
        if active and (
            missing_visual_line
            or not _can_merge_lines(active[-1], line, median_size=median_size)
        ):
            blocks.append(_block_from_lines(active, median_size))
            active = []
        active.append(line)
    if active:
        blocks.append(_block_from_lines(active, median_size))
    return blocks


def _visual_ink_mask(image: Any) -> tuple[Any, float]:
    from PIL import Image, ImageChops

    rgb = image.convert("RGB")
    background = _background_color(rgb)
    difference = ImageChops.difference(
        rgb,
        Image.new("RGB", rgb.size, background),
    ).convert("L")
    mask = difference.point(lambda value: 255 if value >= 11 else 0)
    histogram = mask.histogram()
    ink_ratio = histogram[255] / max(1, rgb.width * rgb.height)
    return mask, float(ink_ratio)


def _merge_sparse_formula_boxes(
    boxes: list[tuple[int, int, int, int]],
    width: int,
    height: int,
) -> list[tuple[int, int, int, int]]:
    rows: list[list[tuple[int, int, int, int]]] = []
    for box in sorted(boxes, key=lambda item: (item[1], item[0])):
        center = (box[1] + box[3]) / 2
        if rows:
            first = rows[-1][0]
            first_center = (first[1] + first[3]) / 2
            tolerance = max(5, max(first[3] - first[1], box[3] - box[1]) * 0.75)
            if abs(center - first_center) <= tolerance:
                rows[-1].append(box)
                continue
        rows.append([box])
    merged_rows: list[tuple[int, int, int, int]] = []
    maximum_gap = max(10, round(width * 0.055))
    for row in rows:
        active: tuple[int, int, int, int] | None = None
        for box in sorted(row, key=lambda item: item[0]):
            if active is None:
                active = box
                continue
            if box[0] - active[2] <= maximum_gap:
                active = (
                    min(active[0], box[0]),
                    min(active[1], box[1]),
                    max(active[2], box[2]),
                    max(active[3], box[3]),
                )
            else:
                merged_rows.append(active)
                active = box
        if active is not None:
            merged_rows.append(active)

    output: list[tuple[int, int, int, int]] = []
    for box in merged_rows:
        if output:
            previous = output[-1]
            vertical_gap = box[1] - previous[3]
            horizontal_overlap = min(box[2], previous[2]) - max(box[0], previous[0])
            combined_height = max(box[3], previous[3]) - min(box[1], previous[1])
            if (
                vertical_gap <= max(6, round(height * 0.025))
                and horizontal_overlap >= min(box[2] - box[0], previous[2] - previous[0]) * 0.35
                and combined_height <= height * 0.24
            ):
                output[-1] = (
                    min(previous[0], box[0]),
                    min(previous[1], box[1]),
                    max(previous[2], box[2]),
                    max(previous[3], box[3]),
                )
                continue
        output.append(box)
    return output


def _refine_visual_regions(
    regions: list[_VisualRegion],
    page_width: int,
    page_height: int,
    *,
    lines: Iterable[_MergedLine] = (),
    formula_hints: Iterable[tuple[int, int, int, int]] | None = None,
    table_hints: Iterable[tuple[int, int, int, int]] = (),
    visual_hints: Iterable[_PreciseVisualHint] = (),
) -> list[_VisualRegion]:
    """Split sparse, over-merged formula islands while retaining real figures."""

    try:
        from PIL import Image, ImageFilter
    except ImportError as exc:
        raise MissingEngineError("区域级 Word 视觉区域分析需要 Pillow") from exc

    source_lines = list(lines)
    formula_boxes = tuple(formula_hints or ())
    table_boxes = tuple(table_hints)
    typed_visual_boxes = tuple(visual_hints)
    typed_hints = bool(formula_boxes or typed_visual_boxes)
    median_line_height = (
        statistics.median(line.height for line in source_lines)
        if source_lines
        else max(120, round(page_height * 0.015))
    )

    def typed_kind(region: _VisualRegion, ink_ratio: float) -> str:
        if any(
            _center_inside_hint(region, hint)
            or _hint_intersection_fraction(region, hint) >= 0.55
            for hint in table_boxes
        ):
            return "table"
        if any(
            _center_inside_hint(region, hint)
            or _hint_intersection_fraction(region, hint) >= 0.55
            for hint in formula_boxes
        ):
            return "formula"
        for kind, x0, y0, x1, y1 in typed_visual_boxes:
            hint = (x0, y0, x1, y1)
            if (
                _center_inside_hint(region, hint)
                or _hint_intersection_fraction(region, hint) >= 0.55
            ):
                return kind
        heuristic_formula = bool(
            ink_ratio < (0.075 if typed_hints else 0.10)
            and region.height <= page_height * (0.12 if typed_hints else 0.15)
            and region.width <= page_width * (0.62 if typed_hints else 0.72)
        )
        if heuristic_formula:
            return "formula_unconfirmed" if typed_hints else "formula"
        return "figure"

    refined: list[_VisualRegion] = []
    for region in regions:
        try:
            image = Image.open(io.BytesIO(region.blob)).convert("RGB")
        except Exception:
            setattr(region, "kind", "figure")
            refined.append(region)
            continue
        mask, ink_ratio = _visual_ink_mask(image)
        overlapping_lines = [
            line
            for line in source_lines
            if min(line.x1, region.x1) > max(line.x, region.x)
            and min(line.y1, region.y1) > max(line.y, region.y)
        ]
        intersects_formula_hint = any(
            _hint_intersection_fraction(region, hint) >= 0.08
            for hint in formula_boxes
        )
        intersects_table_hint = any(
            _hint_intersection_fraction(region, hint) >= 0.08
            for hint in table_boxes
        )
        intersects_precise_visual_hint = any(
            _hint_intersection_fraction(region, (x0, y0, x1, y1)) >= 0.08
            for _kind, x0, y0, x1, y1 in typed_visual_boxes
        )
        page_area_ratio = (region.width * region.height) / max(1, page_width * page_height)
        sparse_and_large = bool(
            ink_ratio < 0.075
            and page_area_ratio >= 0.045
            and (region.height >= page_height * 0.12 or region.width >= page_width * 0.62)
        )
        sparse_multi_line = bool(
            ink_ratio < 0.075
            and len(overlapping_lines) >= 2
            and region.height
            >= max(median_line_height * 1.35, page_height * 0.018)
        )
        split_candidate = bool(
            sparse_and_large
            or (typed_hints and intersects_formula_hint and ink_ratio < 0.12)
            or (
                typed_hints
                and intersects_precise_visual_hint
                and ink_ratio < 0.12
            )
            or (
                sparse_multi_line
                and not intersects_table_hint
                and not intersects_formula_hint
                and not intersects_precise_visual_hint
            )
        )
        has_long_rule = False
        if split_candidate:
            has_long_rule = any(
                x1 - x0 >= image.width * 0.62
                and y1 - y0 <= max(8, image.height * 0.08)
                for x0, y0, x1, y1, _count in _connected_boxes(mask)
            )
        should_split = bool(
            split_candidate
            and (not has_long_rule or intersects_precise_visual_hint)
        )
        if not should_split:
            setattr(region, "kind", typed_kind(region, ink_ratio))
            refined.append(region)
            continue

        compact = mask.filter(ImageFilter.MaxFilter(5))
        raw = [
            (x0, y0, x1, y1)
            for x0, y0, x1, y1, count in _connected_boxes(compact)
            if count >= 8 and x1 - x0 >= 3 and y1 - y0 >= 3
        ]
        boxes = _merge_sparse_formula_boxes(raw, image.width, image.height)
        split_regions: list[_VisualRegion] = []
        for x0, y0, x1, y1 in boxes:
            margin = max(3, round(min(image.width, image.height) * 0.008))
            x0 = max(0, x0 - margin)
            y0 = max(0, y0 - margin)
            x1 = min(image.width, x1 + margin)
            y1 = min(image.height, y1 + margin)
            width = x1 - x0
            height = y1 - y0
            if width < 4 or height < 4:
                continue
            crop = image.crop((x0, y0, x1, y1))
            payload = io.BytesIO()
            crop.save(payload, format="PNG", optimize=True, compress_level=7)
            child = _VisualRegion(
                payload.getvalue(),
                region.x + round(x0 * region.width / image.width),
                region.y + round(y0 * region.height / image.height),
                max(20, round(width * region.width / image.width)),
                max(20, round(height * region.height / image.height)),
                crop.width,
                crop.height,
            )
            _child_mask, child_ink_ratio = _visual_ink_mask(crop)
            setattr(child, "kind", typed_kind(child, child_ink_ratio))
            split_regions.append(child)
        if 1 <= len(split_regions) <= 40:
            refined.extend(split_regions)
        else:
            setattr(region, "kind", typed_kind(region, ink_ratio))
            refined.append(region)
    return sorted(refined, key=lambda item: (item.y, item.x))


def _line_is_replaced_by_formula(
    line: _MergedLine,
    visuals: list[_VisualRegion],
) -> bool:
    for visual in visuals:
        visual_kind = str(getattr(visual, "kind", "figure"))
        if visual_kind not in {
            "formula",
            "formula_unconfirmed",
        }:
            continue
        if _overlap_fraction(line, visual) < 0.42:
            continue
        if _unconfirmed_formula_line_text_candidate(line):
            return True
    return False


def _formula_line_text_candidate(line: _MergedLine) -> bool:
    """Reject prose lines before converting a formula-bearing row to pixels."""

    text = line.text or ""
    compact = re.sub(r"\s+", "", text)
    if not compact:
        return True
    # Keep the original word boundaries for prose detection.  Running the
    # expression over ``compact`` turns a normal English sentence into one or
    # two very long tokens, so a version number, slash, or inline ``x = y`` can
    # incorrectly make the whole line look like a display formula.
    words = re.findall(r"[A-Za-z]{3,}|[\u3400-\u9fff]{2,}", text)
    natural_words = re.findall(
        r"(?<![A-Za-z_])[A-Za-z]{2,}(?![A-Za-z_])",
        text,
    )
    long_natural_words = [word for word in natural_words if len(word) >= 3]
    latin_runs = re.findall(r"[A-Za-z]+", compact)
    letters = sum(character.isalpha() for character in compact)
    digits = sum(character.isdigit() for character in compact)
    formula_operators = sum(
        character in "=+−-*/^_∑∫√≈≠≤≥±→⇒()[]{}"
        for character in compact
    )
    symbols = sum(not character.isalnum() for character in compact)
    symbol_ratio = symbols / max(1, len(compact))
    letter_ratio = letters / max(1, len(compact))
    cjk_characters = sum("\u3400" <= character <= "\u9fff" for character in text)
    cjk_body_punctuation = sum(character in "，。！？；：" for character in text)
    prose_connectors = {
        "a",
        "an",
        "and",
        "are",
        "as",
        "at",
        "be",
        "been",
        "by",
        "for",
        "from",
        "has",
        "have",
        "in",
        "into",
        "is",
        "it",
        "of",
        "on",
        "or",
        "that",
        "the",
        "then",
        "this",
        "to",
        "was",
        "we",
        "were",
        "while",
        "with",
    }
    connector_count = sum(
        word.casefold() in prose_connectors for word in natural_words
    )
    multiword_prose = bool(
        len(natural_words) >= 5
        and len(long_natural_words) >= 3
        and letter_ratio >= 0.45
        and symbol_ratio <= 0.34
    )
    grammatical_prose = bool(
        len(natural_words) >= 3
        and len(long_natural_words) >= 2
        and connector_count >= 1
        and letter_ratio >= 0.42
        and symbol_ratio <= 0.34
    )
    academic_reference_prose = bool(
        len(natural_words) >= 2
        and re.search(
            r"\b(?:fig(?:ure)?|table|algorithm|eq(?:uation)?)\.?\s*\d",
            text,
            flags=re.IGNORECASE,
        )
        and letter_ratio >= 0.30
        and symbol_ratio <= 0.48
    )
    cjk_prose = bool(
        (
            cjk_characters >= 8
            and cjk_body_punctuation >= 1
            and letter_ratio >= 0.40
            and symbol_ratio <= 0.34
        )
        or (
            cjk_characters >= 5
            and cjk_body_punctuation >= 1
            and re.search(r"(?:图|表|算法|公式)\s*\d", text)
            and letter_ratio >= 0.35
            and symbol_ratio <= 0.38
        )
    )
    long_unspaced_prose = bool(
        len(compact) >= 32
        and letter_ratio >= 0.65
        and max((len(run) for run in latin_runs), default=0) >= 20
        and symbol_ratio <= 0.22
    )
    prose_like = bool(
        multiword_prose
        or grammatical_prose
        or academic_reference_prose
        or cjk_prose
        or long_unspaced_prose
    )
    if prose_like:
        return False
    return bool(
        formula_operators >= 1
        or (digits >= 2 and len(words) <= 4)
        or symbol_ratio >= 0.16
        or (len(compact) <= 10 and len(words) <= 1)
    )


def _unconfirmed_formula_line_text_candidate(line: _MergedLine) -> bool:
    """Require positive math evidence before rasterizing an unconfirmed line."""

    text = str(line.text or "")
    compact = re.sub(r"\s+", "", text)
    if not compact:
        return False
    if re.fullmatch(r"[A-Za-z]+|[A-Za-z]+\d+|\d+", compact):
        return False
    if re.fullmatch(r"[•●◦▪■□☐☑✓✔→←↔—–-]+", compact):
        return False
    if re.fullmatch(r"[A-Za-z]+(?:-[A-Za-z]+)+", compact):
        return False
    if re.fullmatch(r"[A-Za-z]+-\d+", compact, flags=re.IGNORECASE):
        return False
    if re.fullmatch(r"\([A-Za-z0-9]\)", compact):
        return False
    has_strong_operator = any(
        character in "=^_∑∫√≈≠≤≥±⇒" for character in compact
    )
    if re.search(r"\b[A-Za-z]{3,}\b", text) and not has_strong_operator:
        return False
    cjk_characters = sum("\u3400" <= character <= "\u9fff" for character in compact)
    if cjk_characters >= 2:
        return False
    strong_operators = sum(
        character in "=^_∑∫√≈≠≤≥±⇒"
        for character in compact
    )
    mathematical_unicode = any(
        unicodedata.category(character) == "Sm"
        or "\u0370" <= character <= "\u03ff"
        or character in "²³¹⁰⁴⁵⁶⁷⁸⁹₀₁₂₃₄₅₆₇₈₉"
        for character in compact
    )
    digits = sum(character.isdigit() for character in compact)
    weak_operators = sum(character in "+−-*/()[]{}" for character in compact)
    variables = re.findall(r"(?<![A-Za-z])[A-Za-z](?![A-Za-z])", text)
    math_terms = re.findall(r"(?<![A-Za-z])[A-Za-z](?![A-Za-z])|\d+(?:\.\d+)?", text)
    return bool(
        mathematical_unicode
        or strong_operators >= 1
        or (
            weak_operators >= 1
            and len(math_terms) >= 2
            and (len(variables) >= 2 or (variables and digits) or digits >= 2)
        )
    )


def _line_has_clear_math_evidence(line: _MergedLine) -> bool:
    """Return true only when a line has positive, non-prose math evidence."""

    if (
        _formula_line_text_candidate(line)
        and _unconfirmed_formula_line_text_candidate(line)
    ):
        return True
    compact = re.sub(r"\s+", "", str(line.text or ""))
    if not compact or len(compact) > 4:
        return False
    if re.fullmatch(r"[•●◦▪■□☐☑✓✔→←↔—–-]+", compact):
        return False
    return all(
        unicodedata.category(character) == "Sm"
        or "\u0370" <= character <= "\u03ff"
        or character in "ˆ˜ˉ˘˙²³¹⁰⁴⁵⁶⁷⁸⁹₀₁₂₃₄₅₆₇₈₉=+−*/^_∑∫√≈≠≤≥±⇒()[]{}"
        for character in compact
    )


def _background_safe_visual_kind(kind: str) -> str:
    normalized = re.sub(r"[^a-z_]", "_", str(kind or "text_visual").casefold())
    if normalized.startswith(_BACKGROUND_SAFE_VISUAL_PREFIX):
        return normalized
    return f"{_BACKGROUND_SAFE_VISUAL_PREFIX}{normalized}"


def _visual_separates(
    first: _TextBlock,
    second: _TextBlock,
    visuals: Iterable[_VisualRegion],
) -> bool:
    top = min(first.y1, second.y)
    bottom = max(first.y1, second.y)
    for visual in visuals:
        if str(getattr(visual, "kind", "figure")) == "background_safe_page":
            continue
        if visual.y1 <= top or visual.y >= bottom:
            continue
        horizontal = min(first.x1, second.x1, visual.x1) - max(
            first.x, second.x, visual.x
        )
        if horizontal > min(first.width, second.width) * 0.18:
            return True
    return False


def _may_share_region(
    active: list[_TextBlock],
    candidate: _TextBlock,
    visuals: list[_VisualRegion],
    page_width: int,
    page_height: int,
) -> bool:
    previous = active[-1]
    if previous.precision or candidate.precision:
        if not (previous.precision and candidate.precision):
            return False
        if previous.column != candidate.column or len(active) >= 4:
            return False
        gap = candidate.y - previous.y1
        if gap < -max(previous.height, candidate.height) * 0.12:
            return False
        maximum_gap = max(
            180,
            round(max(previous.font_size, candidate.font_size) * 24),
        )
        if gap > maximum_gap:
            return False
        if candidate.y1 - active[0].y > page_height * 0.12:
            return False
        overlap = min(previous.x1, candidate.x1) - max(
            previous.x, candidate.x
        )
        aligned = bool(
            overlap >= min(previous.width, candidate.width) * 0.35
            or abs(previous.x - candidate.x) <= page_width * 0.025
        )
        if not aligned or _visual_separates(previous, candidate, visuals):
            return False
        return True
    if previous.column != candidate.column:
        return False
    if len(active) >= 4:
        return False
    gap = candidate.y - previous.y1
    if gap < -max(previous.height, candidate.height) * 0.20:
        return False
    maximum_gap = max(260, round(max(previous.font_size, candidate.font_size) * 28))
    if gap > maximum_gap:
        return False
    if candidate.y1 - active[0].y > page_height * 0.36:
        return False
    overlap = min(previous.x1, candidate.x1) - max(previous.x, candidate.x)
    aligned = bool(
        overlap >= min(previous.width, candidate.width) * 0.42
        or abs(previous.x - candidate.x) <= page_width * 0.035
    )
    if not aligned:
        return False
    if previous.heading and candidate.heading:
        return False
    if _visual_separates(previous, candidate, visuals):
        return False
    return True


def _region_from_blocks(blocks: list[_TextBlock], page_width: int, page_height: int) -> _TextRegion:
    precision = any(block.precision for block in blocks)
    if precision:
        x0 = max(0, min(block.x for block in blocks))
        y0 = max(0, min(block.y for block in blocks))
        x1 = min(page_width, max(block.x1 for block in blocks) + 8)
        y1 = min(page_height, max(block.y1 for block in blocks) + 10)
    else:
        x0 = max(0, min(block.x for block in blocks) - 18)
        y0 = max(0, min(block.y for block in blocks) - 12)
        x1 = min(page_width, max(block.x1 for block in blocks) + 28)
        y1 = min(page_height, max(block.y1 for block in blocks) + 28)
    return _TextRegion(
        list(blocks),
        x0,
        y0,
        max(40, x1 - x0),
        max(40, y1 - y0),
        blocks[0].column,
        precision,
    )


def _coalesce_overlapping_text_regions(
    regions: list[_TextRegion],
    page_width: int,
    page_height: int,
) -> list[_TextRegion]:
    """Merge compatible region boxes whose padding or table rows overlap."""

    pending = list(regions)
    changed = True
    while changed:
        changed = False
        output: list[_TextRegion] = []
        while pending:
            current = pending.pop(0)
            index = 0
            while index < len(pending):
                candidate = pending[index]
                compatible_columns = bool(
                    current.column == candidate.column
                    or current.column in {"full", "main"}
                    or candidate.column in {"full", "main"}
                )
                if (
                    current.precision
                    or candidate.precision
                    or not compatible_columns
                    or _overlap_fraction(current, candidate) <= 0.12
                ):
                    index += 1
                    continue
                blocks = sorted(
                    [*current.blocks, *candidate.blocks],
                    key=lambda block: (block.y, block.x),
                )
                current = _region_from_blocks(blocks, page_width, page_height)
                if any(block.column in {"full", "main"} for block in blocks):
                    current.column = "full"
                pending.pop(index)
                changed = True
                index = 0
            output.append(current)
        pending = output
    return sorted(pending, key=lambda region: (region.y, region.x))


def _merge_blocks_to_regions(
    blocks: list[_TextBlock],
    visuals: list[_VisualRegion],
    page_width: int,
    page_height: int,
) -> list[_TextRegion]:
    regions: list[_TextRegion] = []
    active: list[_TextBlock] = []
    for block in blocks:
        if active and not _may_share_region(
            active, block, visuals, page_width, page_height
        ):
            regions.append(_region_from_blocks(active, page_width, page_height))
            active = []
        active.append(block)
    if active:
        regions.append(_region_from_blocks(active, page_width, page_height))

    # Do not let region padding create avoidable overlaps.  Shrinking the lower
    # edge is safer than allowing one text box to paint over the next one.
    by_column: dict[str, list[_TextRegion]] = {}
    for region in regions:
        by_column.setdefault(region.column, []).append(region)
    for column_regions in by_column.values():
        column_regions.sort(key=lambda item: (item.y, item.x))
        for current, following in zip(column_regions, column_regions[1:]):
            if _overlap_fraction(current, following) <= 0:
                continue
            new_bottom = max(current.y + 40, following.y - 8)
            current.height = max(40, new_bottom - current.y)
    return _coalesce_overlapping_text_regions(regions, page_width, page_height)


def _formula_line_replacements(
    lines: list[_MergedLine],
    visuals: list[_VisualRegion],
    rendered_page: bytes,
    page_width: int,
    page_height: int,
) -> tuple[list[_MergedLine], list[_VisualRegion]]:
    """Replace formula-bearing visual lines with exact crops of the source render."""

    try:
        from PIL import Image
    except ImportError as exc:
        raise MissingEngineError("区域级 Word 公式行重建需要 Pillow") from exc

    try:
        page_image = Image.open(io.BytesIO(rendered_page)).convert("RGB")
    except Exception:
        return lines, visuals
    formula_indices = [
        index
        for index, visual in enumerate(visuals)
        if getattr(visual, "kind", "figure")
        in {"formula", "formula_unconfirmed"}
    ]
    assigned: dict[int, list[int]] = {}
    for visual_index in formula_indices:
        visual = visuals[visual_index]
        best_line: int | None = None
        best_score = 0.0
        for line_index, line in enumerate(lines):
            horizontal = min(line.x1, visual.x1) - max(line.x, visual.x)
            vertical = min(line.y1, visual.y1) - max(line.y, visual.y)
            if horizontal <= 0:
                continue
            center_distance = abs(line.center_y - visual.center_y)
            if vertical <= 0 and center_distance > max(line.height, visual.height) * 0.85:
                continue
            vertical_score = max(0, vertical) / max(1, min(line.height, visual.height))
            horizontal_score = horizontal / max(1, min(line.width, visual.width))
            distance_score = max(
                0.0,
                1.0 - center_distance / max(1, line.height + visual.height),
            )
            score = vertical_score * 1.8 + horizontal_score * 0.5 + distance_score
            if score > best_score:
                best_score = score
                best_line = line_index
        if best_line is not None and best_score >= 0.55:
            assigned.setdefault(best_line, []).append(visual_index)

    candidate_boxes: list[tuple[int, int, int, int]] = []
    consumed: set[int] = set()
    for line_index, line in enumerate(lines):
        related = assigned.get(line_index, [])
        if not related or not _line_is_replaced_by_formula(
            line,
            [visuals[index] for index in related],
        ):
            continue
        consumed.update(related)
        related_visuals = [visuals[index] for index in related]
        x0 = max(0, min([line.x, *(visual.x for visual in related_visuals)]) - 8)
        y0 = max(0, min([line.y, *(visual.y for visual in related_visuals)]) - 4)
        x1 = min(page_width, max([line.x1, *(visual.x1 for visual in related_visuals)]) + 8)
        y1 = min(page_height, max([line.y1, *(visual.y1 for visual in related_visuals)]) + 4)
        candidate_boxes.append((x0, y0, x1, y1))

    merged_boxes = list(candidate_boxes)
    changed = True
    while changed:
        changed = False
        output_boxes: list[tuple[int, int, int, int]] = []
        while merged_boxes:
            current = merged_boxes.pop(0)
            index = 0
            while index < len(merged_boxes):
                candidate = merged_boxes[index]
                horizontal_overlap = min(current[2], candidate[2]) - max(
                    current[0], candidate[0]
                )
                vertical_overlap = min(current[3], candidate[3]) - max(
                    current[1], candidate[1]
                )
                vertical_gap = max(
                    0,
                    max(current[1], candidate[1]) - min(current[3], candidate[3]),
                )
                minimum_width = min(
                    current[2] - current[0], candidate[2] - candidate[0]
                )
                should_merge = bool(
                    (horizontal_overlap > 0 and vertical_overlap > 0)
                    or (
                        horizontal_overlap >= minimum_width * 0.30
                        and vertical_gap <= 36
                    )
                )
                if should_merge:
                    current = (
                        min(current[0], candidate[0]),
                        min(current[1], candidate[1]),
                        max(current[2], candidate[2]),
                        max(current[3], candidate[3]),
                    )
                    merged_boxes.pop(index)
                    changed = True
                    index = 0
                    continue
                index += 1
            output_boxes.append(current)
        merged_boxes = output_boxes

    # Expand every crop to the complete source lines it intersects.  Without
    # this closure step a formula island can cause half of a line to remain
    # editable while the other half is present in the crop, producing a dark
    # double print.  Two bounded passes are sufficient for multi-line formulas
    # without allowing one formula to absorb an entire column.
    for _pass in range(2):
        expanded: list[tuple[int, int, int, int]] = []
        for box in merged_boxes:
            x0, y0, x1, y1 = box
            for line in lines:
                horizontal = min(x1, line.x1) - max(x0, line.x)
                vertical = min(y1, line.y1) - max(y0, line.y)
                if horizontal <= 0 or vertical <= 0:
                    continue
                if (
                    horizontal / max(1, line.width) >= 0.12
                    and vertical / max(1, line.height) >= 0.18
                    and _unconfirmed_formula_line_text_candidate(line)
                ):
                    x0 = min(x0, line.x - 4)
                    y0 = min(y0, line.y - 2)
                    x1 = max(x1, line.x1 + 4)
                    y1 = max(y1, line.y1 + 2)
            expanded.append(
                (
                    max(0, x0),
                    max(0, y0),
                    min(page_width, x1),
                    min(page_height, y1),
                )
            )
        merged_boxes = expanded

    replacements: list[_VisualRegion] = []
    for x0, y0, x1, y1 in merged_boxes:
        pixel_box = (
            max(0, round(x0 * page_image.width / page_width)),
            max(0, round(y0 * page_image.height / page_height)),
            min(page_image.width, round(x1 * page_image.width / page_width)),
            min(page_image.height, round(y1 * page_image.height / page_height)),
        )
        if pixel_box[2] <= pixel_box[0] or pixel_box[3] <= pixel_box[1]:
            continue
        crop = page_image.crop(pixel_box)
        payload = io.BytesIO()
        crop.save(payload, format="PNG", optimize=True, compress_level=7)
        replacement = _VisualRegion(
            payload.getvalue(),
            x0,
            y0,
            max(20, x1 - x0),
            max(20, y1 - y0),
            crop.width,
            crop.height,
        )
        setattr(replacement, "kind", "formula_line")
        replacements.append(replacement)
    output_visuals = [
        visual for index, visual in enumerate(visuals) if index not in consumed
    ]
    for visual in output_visuals:
        if str(getattr(visual, "kind", "figure")) in {
            "formula",
            "formula_unconfirmed",
        }:
            # Residual heuristic components that were not backed by positive
            # mathematical text evidence are ordinary visual content.  Treat
            # them as figures so broad background bboxes cannot suppress valid
            # headings, bullets, page numbers, or form labels.
            setattr(visual, "kind", "figure")
    output_visuals.extend(replacements)
    output_visuals.sort(key=lambda item: (item.y, item.x))
    kept_lines = [
        line
        for line in lines
        if not any(
            (
                min(line.x1, replacement.x1) - max(line.x, replacement.x)
                > line.width * 0.12
                and min(line.y1, replacement.y1) - max(line.y, replacement.y)
                > line.height * 0.18
                and _unconfirmed_formula_line_text_candidate(line)
            )
            for replacement in replacements
        )
    ]
    return kept_lines, output_visuals


def _merged_line_from_frames(
    frames: list[_FrameLine],
    template: _MergedLine,
    *,
    precision: bool,
) -> _MergedLine:
    ordered = sorted(frames, key=lambda item: item.x)
    text = ""
    previous: _FrameLine | None = None
    for frame in ordered:
        if previous is not None:
            text += _text_joiner(
                previous.text,
                frame.text,
                frame.x - previous.x1,
                max(previous.font_size, frame.font_size),
            )
        text += frame.text
        previous = frame
    x0 = min(frame.x for frame in ordered)
    y0 = min(frame.y for frame in ordered)
    x1 = max(frame.x1 for frame in ordered)
    y1 = max(frame.y1 for frame in ordered)
    dominant = max(
        ordered,
        key=lambda frame: (len(frame.text.strip()), frame.font_size),
    )
    largest_font_size = max(frame.font_size for frame in ordered)
    line = _MergedLine(
        ordered,
        x0,
        y0,
        max(20, x1 - x0),
        max(20, y1 - y0),
        text,
        largest_font_size,
        dominant.bold,
        template.column,
    )
    setattr(line, "precision", bool(precision))
    return line


def _strict_position_frame_lines(
    frames: Iterable[_FrameLine],
    template: _MergedLine,
) -> list[_MergedLine]:
    output: list[_MergedLine] = []
    for frame in sorted(frames, key=lambda item: (item.y, item.x)):
        line = _MergedLine(
            [frame],
            frame.x,
            frame.y,
            frame.width,
            frame.height,
            frame.text,
            frame.font_size,
            frame.bold,
            template.column,
        )
        setattr(line, "precision", True)
        setattr(line, "strict_position", True)
        output.append(line)
    return output


def _strict_position_lines_near_visual_hints(
    lines: Iterable[_MergedLine],
    hints: Iterable[tuple[int, int, int, int]],
    page_width: int,
) -> list[_MergedLine]:
    """Preserve source frame geometry only beside dense inline mathematics."""

    boxes = tuple(hints)
    if not boxes:
        return list(lines)
    side_gap_limit = max(240, round(page_width * 0.04))
    output: list[_MergedLine] = []
    for line in lines:
        touches_visual = any(
            _same_hint_row(line, hint)
            and (
                (hint[0] < line.x1 and hint[2] > line.x)
                or -20 <= hint[0] - line.x1 <= side_gap_limit
                or -20 <= line.x - hint[2] <= side_gap_limit
            )
            for hint in boxes
        )
        if touches_visual:
            output.extend(_strict_position_frame_lines(line.frames, line))
        else:
            output.append(line)
    return sorted(output, key=lambda item: (item.y, item.x))


def _frame_hint_overlap_fraction(
    frame: _FrameLine,
    hint: tuple[int, int, int, int],
) -> float:
    x0 = max(frame.x, hint[0])
    y0 = max(frame.y, hint[1])
    x1 = min(frame.x1, hint[2])
    y1 = min(frame.y1, hint[3])
    if x1 <= x0 or y1 <= y0:
        return 0.0
    return ((x1 - x0) * (y1 - y0)) / max(1, frame.width * frame.height)


def _validate_pending_visual_text_clearance(
    pending_hints: Iterable[_PreciseVisualHint],
    lines: Iterable[_MergedLine],
    removed_frames: set[int],
) -> None:
    """Reject an expanded rendered crop that newly absorbs editable text."""

    frames = {
        id(frame): frame
        for line in lines
        for frame in line.frames
        if id(frame) not in removed_frames
    }.values()
    for kind, x0, y0, x1, y1 in pending_hints:
        box = (x0, y0, x1, y1)
        for frame in frames:
            overlap = _frame_hint_overlap_fraction(frame, box)
            center_inside = bool(
                x0 <= frame.x + frame.width / 2.0 <= x1
                and y0 <= frame.y + frame.height / 2.0 <= y1
            )
            if overlap >= _MIN_PRECISE_FRAME_CONFLICT or center_inside:
                raise ValidationError(
                    f"{kind} 精确视觉裁图扩展后会覆盖可编辑文字，"
                    "已保留原固定坐标结果"
                )


def _line_intersects_hint(
    line: _MergedLine,
    hint: tuple[int, int, int, int],
) -> bool:
    return bool(
        min(line.x1, hint[2]) > max(line.x, hint[0])
        and min(line.y1, hint[3]) > max(line.y, hint[1])
    )


def _partition_remaining_formula_line(
    line: _MergedLine,
    retained_frames: list[_FrameLine],
    hints: Iterable[tuple[int, int, int, int]],
) -> list[_MergedLine]:
    ordered = sorted(retained_frames, key=lambda frame: frame.x)
    if not ordered:
        return []
    relevant = [hint for hint in hints if _line_intersects_hint(line, hint)]
    groups: list[list[_FrameLine]] = []
    for frame in ordered:
        if not groups:
            groups.append([frame])
            continue
        previous = groups[-1][-1]
        crosses_formula = any(
            hint[0] < frame.x
            and hint[2] > previous.x1
            and min(line.y1, hint[3]) > max(line.y, hint[1])
            for hint in relevant
        )
        if crosses_formula:
            groups.append([frame])
        else:
            groups[-1].append(frame)
    return [
        _merged_line_from_frames(group, line, precision=True) for group in groups
    ]


def _precise_formula_replacements(
    lines: list[_MergedLine],
    visuals: list[_VisualRegion],
    rendered_page: bytes,
    page_width: int,
    page_height: int,
    hints: Iterable[tuple[int, int, int, int]],
    *,
    audit_regions: list[tuple[int, int, int, int]] | None = None,
    background_page: bytes | None = None,
) -> tuple[list[_MergedLine], list[_VisualRegion]]:
    """Replace only confirmed formula rectangles from the WPS source render."""

    try:
        from PIL import Image
    except ImportError as exc:
        raise MissingEngineError("区域级 Word 精确公式重建需要 Pillow") from exc

    formula_boxes = [
        (x0, y0, x1, y1)
        for _kind, x0, y0, x1, y1 in _normalized_precise_visual_hints(
            (("inline_math", x0, y0, x1, y1) for x0, y0, x1, y1 in hints),
            page_width,
            page_height,
        )
    ]
    if not formula_boxes:
        return lines, visuals
    try:
        page_image = Image.open(io.BytesIO(rendered_page)).convert("RGB")
    except Exception as exc:
        raise ValidationError("无法读取 WPS 公式基准页面") from exc

    background_image = None
    if background_page:
        try:
            background_image = Image.open(io.BytesIO(background_page)).convert("RGB")
        except Exception:
            background_image = None

    removed_frames: set[int] = set()
    affected_lines: set[int] = set()
    if background_image is None:
        for line_index, line in enumerate(lines):
            for frame in line.frames:
                for hint in formula_boxes:
                    overlap = _frame_hint_overlap_fraction(frame, hint)
                    center_x = frame.x + frame.width / 2.0
                    center_y = frame.y + frame.height / 2.0
                    center_inside = (
                        hint[0] <= center_x <= hint[2]
                        and hint[1] <= center_y <= hint[3]
                    )
                    if overlap >= _MIN_PRECISE_FRAME_COVERAGE:
                        if not _line_has_clear_math_evidence(line):
                            raise ValidationError(
                                "公式坐标覆盖了非数学可编辑文字，"
                                "已拒绝不确定的区域候选"
                            )
                        removed_frames.add(id(frame))
                        affected_lines.add(line_index)
                        break
                    if center_inside or overlap >= _MIN_PRECISE_FRAME_CONFLICT:
                        raise ValidationError(
                            "公式坐标未高覆盖可编辑文字，已拒绝不确定的区域候选"
                        )

    kept_visuals, pending_formula_hints = _reconcile_precise_visual_hints(
        visuals,
        [("formula_exact", *box) for box in formula_boxes],
        audit_regions=audit_regions,
    )
    if background_image is None:
        _validate_pending_visual_text_clearance(
            pending_formula_hints,
            lines,
            removed_frames,
        )
    replacements: list[_VisualRegion] = []
    for _kind, x0, y0, x1, y1 in pending_formula_hints:
        crop_source = background_image or page_image
        pixel_box = (
            max(0, round(x0 * crop_source.width / page_width)),
            max(0, round(y0 * crop_source.height / page_height)),
            min(crop_source.width, round(x1 * crop_source.width / page_width)),
            min(crop_source.height, round(y1 * crop_source.height / page_height)),
        )
        if pixel_box[2] <= pixel_box[0] or pixel_box[3] <= pixel_box[1]:
            raise ValidationError("公式坐标无法生成有效裁图")
        crop = crop_source.crop(pixel_box)
        payload = io.BytesIO()
        crop.save(payload, format="PNG", optimize=True, compress_level=7)
        replacement = _VisualRegion(
            payload.getvalue(),
            x0,
            y0,
            max(1, x1 - x0),
            max(1, y1 - y0),
            crop.width,
            crop.height,
        )
        if background_image is not None:
            setattr(replacement, "background_safe", True)
            setattr(
                replacement,
                "kind",
                _background_safe_visual_kind("formula_exact"),
            )
        else:
            setattr(replacement, "kind", "formula_exact")
        replacements.append(replacement)

    for visual in kept_visuals:
        visual_kind = str(getattr(visual, "kind", "figure"))
        if bool(getattr(visual, "background_safe", False)) and visual_kind in {
            "formula",
            "formula_unconfirmed",
            "formula_line",
            "formula_exact",
            "inline_math",
        }:
            setattr(visual, "kind", _background_safe_visual_kind(visual_kind))
    kept_visuals.extend(replacements)
    kept_visuals.sort(key=lambda item: (item.y, item.x))

    output_lines: list[_MergedLine] = []
    for line_index, line in enumerate(lines):
        retained = [frame for frame in line.frames if id(frame) not in removed_frames]
        if not retained:
            continue
        intersects_formula = any(
            _line_intersects_hint(line, hint) for hint in formula_boxes
        )
        if line_index in affected_lines or intersects_formula:
            output_lines.append(
                _merged_line_from_frames(retained, line, precision=True)
            )
        else:
            output_lines.append(line)
    return output_lines, kept_visuals


def _normalized_precise_visual_hints(
    hints: Iterable[_PreciseVisualHint],
    page_width: int,
    page_height: int,
) -> list[_PreciseVisualHint]:
    normalized: set[_PreciseVisualHint] = set()
    for raw_kind, raw_x0, raw_y0, raw_x1, raw_y1 in hints:
        kind = re.sub(r"[^a-z_]", "_", str(raw_kind or "text_visual").casefold())
        if kind not in {"inline_math", "text_visual"}:
            kind = "text_visual"
        x0 = max(0, min(page_width, int(raw_x0)))
        y0 = max(0, min(page_height, int(raw_y0)))
        x1 = max(0, min(page_width, int(raw_x1)))
        y1 = max(0, min(page_height, int(raw_y1)))
        if x1 <= x0 or y1 <= y0:
            continue
        normalized.add((kind, x0, y0, x1, y1))
    pending = sorted(normalized, key=lambda item: (item[2], item[1], item[0]))
    coalesced: list[_PreciseVisualHint] = []
    while pending:
        current = pending.pop(0)
        changed = True
        while changed:
            changed = False
            current_kind, current_x0, current_y0, current_x1, current_y1 = current
            current_area = max(
                1, (current_x1 - current_x0) * (current_y1 - current_y0)
            )
            for index, candidate in enumerate(pending):
                candidate_kind, x0, y0, x1, y1 = candidate
                intersection_width = min(current_x1, x1) - max(current_x0, x0)
                intersection_height = min(current_y1, y1) - max(current_y0, y0)
                if intersection_width <= 0 or intersection_height <= 0:
                    continue
                candidate_area = max(1, (x1 - x0) * (y1 - y0))
                overlap = intersection_width * intersection_height
                if overlap / min(current_area, candidate_area) < 0.80:
                    continue
                current = (
                    "inline_math"
                    if "inline_math" in {current_kind, candidate_kind}
                    else "text_visual",
                    min(current_x0, x0),
                    min(current_y0, y0),
                    max(current_x1, x1),
                    max(current_y1, y1),
                )
                pending.pop(index)
                changed = True
                break
        coalesced.append(current)
    return sorted(coalesced, key=lambda item: (item[2], item[1], item[0]))


def _materialize_page_hints(
    hints_by_page: Mapping[int, Iterable[Any]] | None,
) -> dict[int, tuple[Any, ...]] | None:
    """Freeze public Iterable values so validation never consumes generators."""

    if hints_by_page is None:
        return None
    materialized: dict[int, list[Any]] = {}
    for raw_page_index, raw_values in hints_by_page.items():
        try:
            page_index = int(raw_page_index)
        except (TypeError, ValueError):
            continue
        materialized.setdefault(page_index, []).extend(tuple(raw_values))
    return {
        page_index: tuple(values)
        for page_index, values in materialized.items()
    }


def _same_hint_row(line: _MergedLine, hint: tuple[int, int, int, int]) -> bool:
    vertical = min(line.y1, hint[3]) - max(line.y, hint[1])
    hint_height = max(1, hint[3] - hint[1])
    minimum_height = max(1, min(line.height, hint_height))
    if vertical / minimum_height >= 0.30:
        return True
    hint_center = (hint[1] + hint[3]) / 2.0
    tolerance = max(28, round(minimum_height * 0.34))
    return abs(line.center_y - hint_center) <= tolerance


def _visual_hint_coverage(
    visual: _VisualRegion,
    hint: tuple[int, int, int, int],
) -> tuple[float, float, float, float]:
    """Return area and per-axis coverage for an old visual/exact hint pair."""

    x0 = max(visual.x, hint[0])
    y0 = max(visual.y, hint[1])
    x1 = min(visual.x1, hint[2])
    y1 = min(visual.y1, hint[3])
    if x1 <= x0 or y1 <= y0:
        return 0.0, 0.0, 0.0, 0.0
    intersection_width = x1 - x0
    intersection_height = y1 - y0
    intersection = intersection_width * intersection_height
    visual_area = max(1, visual.width * visual.height)
    hint_area = max(1, (hint[2] - hint[0]) * (hint[3] - hint[1]))
    return (
        intersection / visual_area,
        intersection / hint_area,
        intersection_width / max(1, min(visual.width, hint[2] - hint[0])),
        intersection_height / max(1, min(visual.height, hint[3] - hint[1])),
    )


def _reconcile_precise_visual_hints(
    visuals: list[_VisualRegion],
    hints: Iterable[_PreciseVisualHint],
    *,
    audit_regions: list[tuple[int, int, int, int]] | None = None,
) -> tuple[list[_VisualRegion], list[_PreciseVisualHint]]:
    """Resolve exact crops against visual components without deleting content.

    A small old component that is almost fully covered can be replaced by the
    exact WPS crop.  When a larger existing figure or table already contains
    the hint, keeping it is both more faithful and lighter than adding a second
    overlapping image.  Partial, ambiguous overlap rejects the candidate so
    the caller can retain the legacy fixed-coordinate document.
    """

    removed: set[int] = set()
    pending: list[_PreciseVisualHint] = []
    formula_family = {
        "formula",
        "formula_unconfirmed",
        "formula_line",
        "formula_exact",
        "inline_math",
    }

    def base_kind(value: str) -> str:
        normalized = str(value or "figure")
        if normalized.startswith(_BACKGROUND_SAFE_VISUAL_PREFIX):
            return normalized[len(_BACKGROUND_SAFE_VISUAL_PREFIX) :]
        return normalized

    def compatible(existing_kind: str, hint_kind: str) -> bool:
        existing_base = base_kind(existing_kind)
        hint_base = base_kind(hint_kind)
        if existing_base == hint_base:
            return True
        return existing_base in formula_family and hint_base in formula_family

    for typed_hint in hints:
        kind, x0, y0, x1, y1 = typed_hint
        box = (x0, y0, x1, y1)
        meaningful: list[tuple[_VisualRegion, float, float, float, float]] = []
        for visual in visuals:
            if id(visual) in removed:
                continue
            (
                visual_coverage,
                hint_coverage,
                horizontal_coverage,
                vertical_coverage,
            ) = _visual_hint_coverage(visual, box)
            if (
                max(visual_coverage, hint_coverage) < 0.08
                or horizontal_coverage < 0.12
                or vertical_coverage < 0.12
            ):
                continue
            meaningful.append(
                (
                    visual,
                    visual_coverage,
                    hint_coverage,
                    horizontal_coverage,
                    vertical_coverage,
                )
            )

        already_covered = any(
            (
                hint_coverage >= 0.85
                and (
                    visual_coverage < 0.85
                    or base_kind(
                        str(getattr(visual, "kind", "figure"))
                    ).startswith("table")
                )
            )
            or (
                compatible(str(getattr(visual, "kind", "figure")), kind)
                and hint_coverage >= 0.60
                and visual_coverage < 0.20
                and max(horizontal_coverage, vertical_coverage) >= 0.95
            )
            for (
                visual,
                visual_coverage,
                hint_coverage,
                horizontal_coverage,
                vertical_coverage,
            ) in meaningful
        )
        if already_covered:
            continue

        removable: list[_VisualRegion] = []
        for (
            visual,
            visual_coverage,
            hint_coverage,
            _horizontal_coverage,
            _vertical_coverage,
        ) in meaningful:
            visual_kind = str(getattr(visual, "kind", "figure"))
            visual_base_kind = base_kind(visual_kind)
            if visual_base_kind.startswith("table"):
                raise ValidationError(
                    "精确视觉提示与表格视觉层部分重叠，已保留原固定坐标结果"
                )
            if visual_coverage >= 0.80:
                removable.append(visual)
                continue
            visual_center_x = visual.x + visual.width / 2.0
            visual_center_y = visual.y + visual.height / 2.0
            hint_center_x = (x0 + x1) / 2.0
            hint_center_y = (y0 + y1) / 2.0
            component_coverage = (
                0.30
                if (
                    visual_base_kind == base_kind(kind)
                    or (
                        base_kind(kind) == "formula_exact"
                        and visual_base_kind in formula_family
                    )
                )
                else 0.55
            )
            same_component = bool(
                compatible(visual_kind, kind)
                and min(visual_coverage, hint_coverage) >= component_coverage
                and abs(visual_center_x - hint_center_x)
                <= max(visual.width, x1 - x0) * 0.35
                and abs(visual_center_y - hint_center_y)
                <= max(visual.height, y1 - y0) * 0.35
            )
            if same_component:
                removable.append(visual)
                continue
            raise ValidationError(
                f"{kind} 精确视觉提示与既有 {visual_kind} 视觉层部分重叠，"
                f"覆盖 {visual_coverage:.1%}/{hint_coverage:.1%}，"
                "已保留原固定坐标结果"
            )
        removed.update(id(visual) for visual in removable)
        expanded_box = (
            min([x0, *(visual.x for visual in removable)]),
            min([y0, *(visual.y for visual in removable)]),
            max([x1, *(visual.x1 for visual in removable)]),
            max([y1, *(visual.y1 for visual in removable)]),
        )
        expanded_hint = (kind, *expanded_box)
        pending.append(expanded_hint)
        if audit_regions is not None:
            audit_regions.append(expanded_box)

    return [visual for visual in visuals if id(visual) not in removed], pending


def _precise_visual_replacements(
    lines: list[_MergedLine],
    visuals: list[_VisualRegion],
    rendered_page: bytes,
    page_width: int,
    page_height: int,
    hints: Iterable[_PreciseVisualHint],
    *,
    audit_regions: list[tuple[int, int, int, int]] | None = None,
    background_page: bytes | None = None,
) -> tuple[list[_MergedLine], list[_VisualRegion]]:
    """Keep visual-only PDF spans as exact crops and split text around them.

    The fixed-layout builder deliberately leaves unportable math/font spans in
    the page background.  When neighboring editable frames are later grouped
    into one large text box, that box can paint across the retained glyphs.
    Exact typed hints let us preserve the visual island while keeping the text
    on either side independently selectable and copyable.
    """

    normalized = _normalized_precise_visual_hints(hints, page_width, page_height)
    if not normalized:
        return lines, visuals
    try:
        from PIL import Image
    except ImportError as exc:
        raise MissingEngineError("区域级 Word 精确视觉文字重建需要 Pillow") from exc
    try:
        page_image = Image.open(io.BytesIO(rendered_page)).convert("RGB")
    except Exception as exc:
        raise ValidationError("无法读取 WPS 行内视觉文字基准页面") from exc

    background_image = None
    if background_page:
        try:
            background_image = Image.open(io.BytesIO(background_page)).convert("RGB")
        except Exception:
            background_image = None

    typed_boxes = [
        (kind, (x0, y0, x1, y1))
        for kind, x0, y0, x1, y1 in normalized
    ]
    boxes = [box for _kind, box in typed_boxes]
    removed_frames: set[int] = set()
    # The page background was generated after editable spans were redacted.
    # Cropping it therefore cannot duplicate neighboring editable words; keep
    # every frame and let the unchanged local pixel audit verify the result.
    if background_image is None:
        for line in lines:
            for frame in line.frames:
                for kind, hint in typed_boxes:
                    if not _same_hint_row(line, hint):
                        continue
                    overlap = _frame_hint_overlap_fraction(frame, hint)
                    center_x = frame.x + frame.width / 2.0
                    center_y = frame.y + frame.height / 2.0
                    center_inside = bool(
                        hint[0] <= center_x <= hint[2]
                        and hint[1] <= center_y <= hint[3]
                    )
                    if (
                        overlap < _MIN_PRECISE_FRAME_CONFLICT
                        and not center_inside
                    ):
                        continue
                    compact = re.sub(r"\s+", "", frame.text or "")
                    visual_only_symbol = bool(
                        compact
                        and len(compact) <= 4
                        and not re.search(r"[\u3400-\u9fff]", compact)
                    )
                    removable = bool(
                        overlap >= _MIN_PRECISE_FRAME_COVERAGE
                        and visual_only_symbol
                        and (
                            kind == "text_visual"
                            or _line_has_clear_math_evidence(line)
                        )
                    )
                    if removable:
                        removed_frames.add(id(frame))
                        break
                    raise ValidationError(
                        "视觉文字坐标未以高覆盖命中明确视觉/数学字形，"
                        "已拒绝不确定的区域候选"
                    )

    kept_visuals, pending_hints = _reconcile_precise_visual_hints(
        visuals,
        normalized,
        audit_regions=audit_regions,
    )
    if background_image is None:
        _validate_pending_visual_text_clearance(
            pending_hints,
            lines,
            removed_frames,
        )
    replacements: list[_VisualRegion] = []
    for kind, x0, y0, x1, y1 in pending_hints:
        crop_source = background_image or page_image
        pixel_box = (
            max(0, round(x0 * crop_source.width / page_width)),
            max(0, round(y0 * crop_source.height / page_height)),
            min(crop_source.width, round(x1 * crop_source.width / page_width)),
            min(crop_source.height, round(y1 * crop_source.height / page_height)),
        )
        if pixel_box[2] <= pixel_box[0] or pixel_box[3] <= pixel_box[1]:
            raise ValidationError("视觉文字坐标无法生成有效裁图")
        crop = crop_source.crop(pixel_box)
        payload = io.BytesIO()
        crop.save(payload, format="PNG", optimize=True, compress_level=7)
        replacement = _VisualRegion(
            payload.getvalue(),
            x0,
            y0,
            max(1, x1 - x0),
            max(1, y1 - y0),
            crop.width,
            crop.height,
        )
        if background_image is not None:
            setattr(replacement, "background_safe", True)
            setattr(replacement, "kind", _background_safe_visual_kind(kind))
        else:
            setattr(replacement, "kind", kind)
        replacements.append(replacement)

    for visual in kept_visuals:
        visual_kind = str(getattr(visual, "kind", "figure"))
        if bool(getattr(visual, "background_safe", False)) and visual_kind in {
            "inline_math",
            "text_visual",
        }:
            setattr(visual, "kind", _background_safe_visual_kind(visual_kind))
    kept_visuals.extend(replacements)
    kept_visuals.sort(key=lambda item: (item.y, item.x))

    working_lines: list[_MergedLine] = []
    for line in lines:
        retained_frames = [
            frame for frame in line.frames if id(frame) not in removed_frames
        ]
        if not retained_frames:
            continue
        working_lines.append(
            _merged_line_from_frames(retained_frames, line, precision=True)
            if len(retained_frames) != len(line.frames)
            else line
        )

    # Rejoin the editable fragments on both sides into one precision row.  A
    # left-aligned paragraph with explicit tab stops keeps the visual gap while
    # remaining much easier to copy and far lighter than one text box per
    # fragment.
    side_gap_limit = max(240, round(page_width * 0.04))
    for hint in boxes:
        row_lines = [line for line in working_lines if _same_hint_row(line, hint)]
        left = [
            line
            for line in row_lines
            if -20 <= hint[0] - line.x1 <= side_gap_limit
        ]
        right = [
            line
            for line in row_lines
            if -20 <= line.x - hint[2] <= side_gap_limit
        ]
        intersecting = [
            line
            for line in row_lines
            if hint[0] < line.x1 and hint[2] > line.x
        ]
        selected: list[_MergedLine] = list(intersecting)
        if left and right:
            selected.extend(
                (
                    max(left, key=lambda item: item.x1),
                    min(right, key=lambda item: item.x),
                )
            )
        elif left:
            selected.append(max(left, key=lambda item: item.x1))
        elif right:
            selected.append(min(right, key=lambda item: item.x))
        selected_ids = {id(line) for line in selected}
        selected = [
            line for line in working_lines if id(line) in selected_ids
        ]
        if not selected:
            continue
        if len(selected) == 1:
            setattr(selected[0], "precision", True)
            continue
        frames = sorted(
            (frame for line in selected for frame in line.frames),
            key=lambda item: item.x,
        )
        template = max(
            selected,
            key=lambda item: (len(item.text.strip()), item.font_size),
        )
        merged = _merged_line_from_frames(frames, template, precision=True)
        working_lines = [
            line for line in working_lines if id(line) not in selected_ids
        ]
        working_lines.append(merged)
    return sorted(working_lines, key=lambda item: (item.y, item.x)), kept_visuals


def _deduplicate_contained_background_visuals(
    visuals: list[_VisualRegion],
    background_page: bytes | None,
    page_width: int,
    page_height: int,
) -> list[_VisualRegion]:
    """Remove a smaller exact crop already present in a same-source crop.

    Precise formula and visual-span hints can independently describe the same
    glyph island.  When both replacements were cropped from the redacted page
    background, retaining a fully contained second anchor only makes WPS do
    more layout and image-decoding work.  The smaller anchor is removed only
    after the outer crop is proven, pixel for pixel, to contain the matching
    source-background pixels at that position.  Any ambiguity keeps both
    visuals so the existing structural gate can reject the candidate safely.
    """

    if background_page is None or len(visuals) < 2:
        return visuals
    try:
        from PIL import Image, ImageChops

        background = Image.open(io.BytesIO(background_page)).convert("RGB")
    except Exception:
        return visuals
    if page_width <= 0 or page_height <= 0:
        return visuals

    def source_box(visual: _VisualRegion) -> tuple[int, int, int, int]:
        return (
            max(0, round(visual.x * background.width / page_width)),
            max(0, round(visual.y * background.height / page_height)),
            min(background.width, round(visual.x1 * background.width / page_width)),
            min(background.height, round(visual.y1 * background.height / page_height)),
        )

    def safe_kind(visual: _VisualRegion) -> str | None:
        kind = str(getattr(visual, "kind", "figure"))
        if not bool(getattr(visual, "background_safe", False)):
            return None
        if not kind.startswith(_BACKGROUND_SAFE_VISUAL_PREFIX):
            return None
        return kind[len(_BACKGROUND_SAFE_VISUAL_PREFIX) :]

    image_cache: dict[int, Any | None] = {}

    def outer_contains_source_pixels(
        outer: _VisualRegion,
        inner: _VisualRegion,
    ) -> bool:
        outer_box = source_box(outer)
        inner_box = source_box(inner)
        if (
            inner_box[2] <= inner_box[0]
            or inner_box[3] <= inner_box[1]
            or outer_box[0] > inner_box[0]
            or outer_box[1] > inner_box[1]
            or outer_box[2] < inner_box[2]
            or outer_box[3] < inner_box[3]
        ):
            return False
        cache_key = id(outer)
        if cache_key not in image_cache:
            try:
                image_cache[cache_key] = Image.open(
                    io.BytesIO(outer.blob)
                ).convert("RGB")
            except Exception:
                image_cache[cache_key] = None
        outer_image = image_cache[cache_key]
        if outer_image is None or outer_image.size != (
            outer_box[2] - outer_box[0],
            outer_box[3] - outer_box[1],
        ):
            return False
        relative_box = (
            inner_box[0] - outer_box[0],
            inner_box[1] - outer_box[1],
            inner_box[2] - outer_box[0],
            inner_box[3] - outer_box[1],
        )
        retained_pixels = outer_image.crop(relative_box)
        source_pixels = background.crop(inner_box)
        return ImageChops.difference(retained_pixels, source_pixels).getbbox() is None

    ordered = sorted(
        enumerate(visuals),
        key=lambda item: (
            -(item[1].width * item[1].height),
            item[1].y,
            item[1].x,
            item[0],
        ),
    )
    retained: list[_VisualRegion] = []
    removed: set[int] = set()
    for _index, candidate in ordered:
        candidate_kind = safe_kind(candidate)
        if candidate_kind is not None:
            for outer in retained:
                if safe_kind(outer) != candidate_kind:
                    continue
                if (
                    outer.x <= candidate.x
                    and outer.y <= candidate.y
                    and outer.x1 >= candidate.x1
                    and outer.y1 >= candidate.y1
                    and outer_contains_source_pixels(outer, candidate)
                ):
                    removed.add(id(candidate))
                    break
        if id(candidate) not in removed:
            retained.append(candidate)
    return sorted(
        (visual for visual in visuals if id(visual) not in removed),
        key=lambda item: (item.y, item.x),
    )


def _align_background_visuals_to_source_pixels(
    visuals: list[_VisualRegion],
    background_page: bytes | None,
    page_width: int,
    page_height: int,
) -> list[_VisualRegion]:
    """Recrop safe visuals on exact source-pixel/twip boundaries.

    A crop selected in twips is first rounded to source pixels.  Reusing the
    pre-rounding twip rectangle when inserting that bitmap makes WPS rescale a
    56-pixel glyph island as, for example, 56.4 pixels, which changes its
    antialiasing phase relative to the original full-page background.  Anchor
    the crop to the inverse-mapped pixel edges instead so both render paths use
    the same pixels-per-twip ratio.
    """

    if background_page is None or not visuals or page_width <= 0 or page_height <= 0:
        return visuals
    try:
        from PIL import Image

        background = Image.open(io.BytesIO(background_page)).convert("RGB")
    except Exception:
        return visuals
    for visual in visuals:
        if not bool(getattr(visual, "background_safe", False)):
            continue
        pixel_box = (
            max(0, round(visual.x * background.width / page_width)),
            max(0, round(visual.y * background.height / page_height)),
            min(background.width, round(visual.x1 * background.width / page_width)),
            min(background.height, round(visual.y1 * background.height / page_height)),
        )
        if pixel_box[2] <= pixel_box[0] or pixel_box[3] <= pixel_box[1]:
            continue
        aligned_box = (
            round(pixel_box[0] * page_width / background.width),
            round(pixel_box[1] * page_height / background.height),
            round(pixel_box[2] * page_width / background.width),
            round(pixel_box[3] * page_height / background.height),
        )
        if aligned_box[2] <= aligned_box[0] or aligned_box[3] <= aligned_box[1]:
            continue
        crop = background.crop(pixel_box)
        payload = io.BytesIO()
        crop.save(payload, format="PNG", optimize=True, compress_level=7)
        visual.blob = payload.getvalue()
        visual.x = aligned_box[0]
        visual.y = aligned_box[1]
        visual.width = aligned_box[2] - aligned_box[0]
        visual.height = aligned_box[3] - aligned_box[1]
        visual.pixel_width = crop.width
        visual.pixel_height = crop.height
    return sorted(visuals, key=lambda item: (item.y, item.x))


def _full_page_background_visual(
    background_page: bytes,
    page_width: int,
    page_height: int,
) -> _VisualRegion:
    """Return one exact redacted-background anchor for a math-dense page."""

    try:
        from PIL import Image

        with Image.open(io.BytesIO(background_page)) as image:
            pixel_width, pixel_height = image.size
    except Exception as exc:
        raise ValidationError("无法读取高密度公式页的安全背景图") from exc
    visual = _VisualRegion(
        background_page,
        0,
        0,
        page_width,
        page_height,
        pixel_width,
        pixel_height,
    )
    setattr(visual, "background_safe", True)
    setattr(visual, "kind", "background_safe_page")
    return visual


def _mark_table_precision_lines(
    lines: list[_MergedLine],
    table_hints: Iterable[tuple[int, int, int, int]],
) -> list[_MergedLine]:
    hints = tuple(table_hints)
    if not hints:
        return lines
    for line in lines:
        if any(_line_intersects_hint(line, hint) for hint in hints):
            setattr(line, "precision", True)
    return lines


def _plan_pages(
    source_document: Any,
    rendered_pages: list[bytes] | None = None,
    *,
    preserve_editable_text: bool = False,
    formula_hints_by_page: Mapping[
        int, Iterable[tuple[int, int, int, int]]
    ]
    | None = None,
    table_hints_by_page: Mapping[
        int, Iterable[tuple[int, int, int, int]]
    ]
    | None = None,
    visual_hints_by_page: Mapping[int, Iterable[_PreciseVisualHint]] | None = None,
) -> list[_RegionPage]:
    source_pages = _source_pages(source_document)
    planned: list[_RegionPage] = []
    for page in source_pages:
        lines = _merge_frame_rows(page.frames, page.width)
        raw_formula_hints = (
            None
            if formula_hints_by_page is None
            else tuple(formula_hints_by_page.get(page.index, ()))
        )
        formula_hints = (
            None
            if raw_formula_hints is None
            else tuple(
                (x0, y0, x1, y1)
                for _kind, x0, y0, x1, y1 in _normalized_precise_visual_hints(
                    (
                        ("inline_math", x0, y0, x1, y1)
                        for x0, y0, x1, y1 in raw_formula_hints
                    ),
                    page.width,
                    page.height,
                )
            )
        )
        table_hints = (
            ()
            if table_hints_by_page is None
            else tuple(table_hints_by_page.get(page.index, ()))
        )
        visual_hints = _normalized_precise_visual_hints(
            (
                ()
                if visual_hints_by_page is None
                else tuple(visual_hints_by_page.get(page.index, ()))
            ),
            page.width,
            page.height,
        )
        precise_audit_regions = list(formula_hints or ())
        precise_audit_regions.extend(
            (x0, y0, x1, y1) for _kind, x0, y0, x1, y1 in visual_hints
        )
        dense_page_background = bool(
            page.background is not None
            and len(visual_hints) >= _DENSE_BACKGROUND_HINT_THRESHOLD
        )
        if dense_page_background:
            visuals = [
                _full_page_background_visual(
                    page.background,
                    page.width,
                    page.height,
                )
            ]
        else:
            visuals = (
                _visual_regions(page.background, page.width, page.height)
                if page.background is not None
                else []
            )
            visuals = _refine_visual_regions(
                visuals,
                page.width,
                page.height,
                lines=lines,
                formula_hints=formula_hints,
                table_hints=table_hints,
                visual_hints=visual_hints,
            )
        if page.background is not None and not dense_page_background:
            for visual in visuals:
                setattr(visual, "background_safe", True)
        if formula_hints:
            if rendered_pages is None or page.index >= len(rendered_pages):
                raise ValidationError("精确公式区域缺少 WPS 渲染基准")
            lines, visuals = _precise_formula_replacements(
                lines,
                visuals,
                rendered_pages[page.index],
                page.width,
                page.height,
                formula_hints,
                audit_regions=precise_audit_regions,
                background_page=page.background,
            )
        if visual_hints:
            if rendered_pages is None or page.index >= len(rendered_pages):
                raise ValidationError("精确视觉文字区域缺少 WPS 渲染基准")
            lines, visuals = _precise_visual_replacements(
                lines,
                visuals,
                rendered_pages[page.index],
                page.width,
                page.height,
                visual_hints,
                audit_regions=precise_audit_regions,
                background_page=page.background,
            )
        # Confirmed PDF hints are authoritative only for their own rectangles.
        # Always run the conservative residual pass: it only consumes regions
        # classified as formula/formula_unconfirmed after table and typed-hint
        # exclusion, so a mixed page can still preserve an unnumbered formula.
        if not preserve_editable_text and not dense_page_background:
            if rendered_pages is not None and page.index < len(rendered_pages):
                lines, visuals = _formula_line_replacements(
                    lines,
                    visuals,
                    rendered_pages[page.index],
                    page.width,
                    page.height,
                )
            else:
                lines = [
                    line
                    for line in lines
                    if not _line_is_replaced_by_formula(line, visuals)
                ]
        visuals = _align_background_visuals_to_source_pixels(
            visuals,
            page.background,
            page.width,
            page.height,
        )
        visuals = _deduplicate_contained_background_visuals(
            visuals,
            page.background,
            page.width,
            page.height,
        )
        lines = _mark_table_precision_lines(lines, table_hints)
        blocks = _paragraph_blocks(lines, page.width)
        regions = _merge_blocks_to_regions(
            blocks, visuals, page.width, page.height
        )
        precise_audit_regions = sorted(
            set(precise_audit_regions),
            key=lambda box: (box[1], box[0], box[3], box[2]),
        )
        planned.append(
            _RegionPage(
                page.index,
                page.width,
                page.height,
                regions,
                visuals,
                precise_audit_regions,
            )
        )
    return planned


def _configure_section(section: Any, width: int, height: int) -> None:
    from docx.enum.section import WD_ORIENT
    from docx.shared import Pt, Twips

    section.page_width = Twips(width)
    section.page_height = Twips(height)
    section.orientation = WD_ORIENT.LANDSCAPE if width > height else WD_ORIENT.PORTRAIT
    margin = Pt(1)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin
    section.header_distance = Pt(0)
    section.footer_distance = Pt(0)


def _new_page_carrier(document: Any, *, page_break: bool) -> Any:
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(1)
    if page_break:
        paragraph.paragraph_format.page_break_before = True
    paragraph.paragraph_format.keep_together = True
    return paragraph


def _frame_tab_threshold(
    previous: _FrameLine,
    current: _FrameLine,
    *,
    precision: bool,
) -> int:
    return max(
        18 if precision else 220,
        round(
            max(previous.font_size, current.font_size)
            * (3 if precision else 52)
        ),
    )


def _line_uses_tab_alignment(line: _MergedLine, *, precision: bool) -> bool:
    ordered = sorted(line.frames, key=lambda item: item.x)
    return any(
        current.x - previous.x1
        >= _frame_tab_threshold(previous, current, precision=precision)
        for previous, current in zip(ordered, ordered[1:])
    )


def _emission_blocks(block: _TextBlock) -> list[_TextBlock]:
    """Split multi-line tab layouts so each line owns its tab-stop map."""

    if len(block.lines) <= 1 or not any(
        _line_uses_tab_alignment(
            line,
            precision=bool(block.precision or getattr(line, "precision", False)),
        )
        for line in block.lines
    ):
        return [block]

    output: list[_TextBlock] = []
    for index, line in enumerate(block.lines):
        line_font_size = max(
            [line.font_size]
            + [frame.font_size for frame in line.frames]
        )
        output.append(
            _TextBlock(
                [line],
                line.x,
                line.y,
                line.width,
                line.height,
                line.column,
                line_font_size,
                bool(block.heading and index == 0),
                bool(block.precision or getattr(line, "precision", False)),
            )
        )
    return output


def _validate_source_paragraph_for_region(paragraph: Any) -> None:
    """Reject source semantics that a fresh region document cannot copy safely."""

    from docx.oxml.ns import qn

    source = paragraph._p
    nested_runs = [
        run
        for run in source.xpath(".//w:r")
        if run.getparent() is not source
    ]
    if nested_runs:
        raise ValidationError(
            "固定坐标源 Word 包含超链接、内容控件或修订层中的嵌套文字；"
            "区域候选已拒绝并保留原高精度结果"
        )

    properties = source.pPr
    if properties is not None and any(
        properties.find(qn(name)) is not None
        for name in ("w:pStyle", "w:numPr", "w:bidi", "w:textDirection")
    ):
        raise ValidationError(
            "固定坐标源 Word 的段落依赖样式、编号或双向排版；"
            "区域候选已拒绝并保留原高精度结果"
        )

    for run in paragraph.runs:
        run_properties = run._r.rPr
        if run_properties is None:
            continue
        if run_properties.find(qn("w:rStyle")) is not None:
            raise ValidationError(
                "固定坐标源 Word 的文字依赖字符样式；"
                "区域候选已拒绝并保留原高精度结果"
            )
        for element in run_properties.iter():
            for attribute in element.attrib:
                local_name = str(attribute).rsplit("}", 1)[-1]
                if local_name in {
                    "asciiTheme",
                    "hAnsiTheme",
                    "eastAsiaTheme",
                    "cstheme",
                    "themeColor",
                    "themeTint",
                    "themeShade",
                    "themeFill",
                    "themeFillTint",
                    "themeFillShade",
                }:
                    raise ValidationError(
                        "固定坐标源 Word 的文字依赖主题字体或主题颜色；"
                        "区域候选已拒绝并保留原高精度结果"
                    )


def _run_can_carry_soft_break(run: Any) -> bool:
    from docx.oxml.ns import qn

    properties = run._r.rPr
    if properties is None:
        return True
    return not any(
        properties.find(qn(name)) is not None
        for name in (
            "w:vanish",
            "w:specVanish",
            "w:webHidden",
            "w:rtl",
            "w:vertAlign",
        )
    )


def _line_spacing(lines: list[_MergedLine], font_size: float) -> float:
    largest_font_size = max(
        [font_size]
        + [line.font_size for line in lines]
        + [frame.font_size for line in lines for frame in line.frames]
    )
    gaps = [
        current.y - previous.y
        for previous, current in zip(lines, lines[1:])
        if current.y > previous.y
    ]
    measured = (
        statistics.median(gaps) / 20
        if gaps
        else largest_font_size * 1.08
    )
    return max(
        largest_font_size,
        min(largest_font_size * 1.38, measured),
    )


def _append_region_paragraph(
    document: Any,
    block: _TextBlock,
    region: _TextRegion,
    *,
    previous_block: _TextBlock | None,
    word_lexicon: set[str],
    normalize_text: bool,
) -> Any:
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.shared import Pt, Twips
    from lxml import etree

    def append_text(run: Any, value: str) -> None:
        """Append text without creating another run or losing tabs/breaks."""

        value = value.replace("\r\n", "\n")
        pieces = re.split(r"([\t\r\n])", value)
        for piece in pieces:
            if not piece:
                continue
            if piece == "\t":
                run.add_tab()
            elif piece in {"\r", "\n"}:
                run.add_break()
            else:
                run.add_text(piece)

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    before_twips = 0 if previous_block is None else max(0, block.y - previous_block.y1)
    paragraph.paragraph_format.space_before = Twips(before_twips)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(
        _line_spacing(block.lines, block.font_size)
    )
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = bool(block.heading)

    continuation_x = (
        statistics.median(line.x for line in block.lines[1:])
        if len(block.lines) > 1
        else block.lines[0].x
    )
    left_indent = max(0, round(continuation_x - region.x))
    first_indent = round(block.lines[0].x - continuation_x)
    paragraph.paragraph_format.left_indent = Twips(left_indent)
    if abs(first_indent) >= 20:
        paragraph.paragraph_format.first_line_indent = Twips(first_indent)

    first = block.lines[0]
    centered = bool(
        block.heading
        and block.column == "full"
        and abs((first.x + first.x1) / 2 - (region.x + region.x1) / 2)
        <= region.width * 0.14
    )
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.LEFT
        if block.precision
        else WD_ALIGN_PARAGRAPH.CENTER
        if centered
        else WD_ALIGN_PARAGRAPH.LEFT
        if block.heading
        else WD_ALIGN_PARAGRAPH.JUSTIFY
    )

    previous_line: _MergedLine | None = None
    emitted_text = ""
    previous_target_run: Any | None = None
    previous_run_properties: bytes | None = None
    pending_line_break = False

    def emit_run(
        text: str,
        properties_key: bytes,
        copied_properties: Any | None,
    ) -> Any | None:
        nonlocal pending_line_break
        nonlocal previous_target_run
        nonlocal previous_run_properties

        if not text:
            return previous_target_run
        if pending_line_break:
            if (
                previous_target_run is not None
                and previous_run_properties == properties_key
                and _run_can_carry_soft_break(previous_target_run)
            ):
                previous_target_run.add_break()
            else:
                paragraph.add_run().add_break()
                previous_target_run = None
                previous_run_properties = None
            pending_line_break = False
        if (
            previous_target_run is not None
            and previous_run_properties == properties_key
        ):
            append_text(previous_target_run, text)
            return previous_target_run
        target_run = paragraph.add_run(text)
        if copied_properties is not None and target_run._r.rPr is None:
            target_run._r.insert(0, copied_properties)
        previous_target_run = target_run
        previous_run_properties = properties_key
        return target_run

    for line in block.lines:
        if previous_line is not None:
            # Preserve the source visual baselines inside one editable
            # paragraph.  This is intentionally a soft line break, not another
            # positioned paragraph: copying the whole region remains usable,
            # while equations and inline visual islands keep their coordinates.
            emitted_text += "\n"
            pending_line_break = True
        previous_frame: _FrameLine | None = None
        first_prefix = ""
        for frame in sorted(line.frames, key=lambda item: item.x):
            _validate_source_paragraph_for_region(frame.paragraph)
            joiner = first_prefix
            if previous_frame is not None:
                gap = frame.x - previous_frame.x1
                tab_threshold = _frame_tab_threshold(
                    previous_frame,
                    frame,
                    precision=bool(block.precision),
                )
                if gap >= tab_threshold:
                    tab_position = max(20, frame.x - region.x)
                    paragraph.paragraph_format.tab_stops.add_tab_stop(
                        Twips(tab_position)
                    )
                    joiner += "\t"
                else:
                    joiner += _text_joiner(
                        previous_frame.text,
                        frame.text,
                        gap,
                        max(previous_frame.font_size, frame.font_size),
                    )
            source_runs = list(frame.paragraph.runs)
            if not source_runs:
                frame_text = (
                    _normalized_source_text(frame.text)
                    if normalize_text
                    else frame.text
                )
                text = joiner + frame_text
                if normalize_text:
                    text = _normalize_text_boundary(emitted_text, text)
                emit_run(text, b"", None)
                emitted_text += text
            else:
                for source_run in source_runs:
                    unsupported = [
                        str(child.tag).rsplit("}", 1)[-1]
                        for child in source_run._r
                        if str(child.tag).rsplit("}", 1)[-1]
                        not in {"rPr", "t"}
                    ]
                    if unsupported:
                        raise ValidationError(
                            "固定坐标源 Word 包含无法无损迁移的复杂文字节点（"
                            f"{', '.join(sorted(set(unsupported)))}），"
                            "已保留原固定坐标结果"
                        )
                first_run = True
                run_segments = (
                    _normalized_frame_runs(frame)
                    if normalize_text
                    else [(run, str(run.text or "")) for run in source_runs]
                )
                for source_run, normalized_text in run_segments:
                    text = normalized_text
                    if first_run:
                        text = joiner + text
                        first_run = False
                    if normalize_text:
                        text = _normalize_text_boundary(emitted_text, text)
                    properties = source_run._r.rPr
                    copied = None
                    if properties is not None:
                        from copy import deepcopy
                        from docx.oxml.ns import qn

                        copied = deepcopy(properties)
                        for name in ("w:fitText", "w:w"):
                            for element in list(copied.findall(qn(name))):
                                copied.remove(element)
                    properties_key = (
                        etree.tostring(copied, with_tail=False)
                        if copied is not None
                        else b""
                    )
                    emit_run(text, properties_key, copied)
                    emitted_text += text
            first_prefix = ""
            previous_frame = frame
        previous_line = line
    element = paragraph._p
    element.getparent().remove(element)
    return element


def _append_text_box(
    document: Any,
    carrier: Any,
    region: _TextRegion,
    *,
    page_index: int,
    region_index: int,
    word_lexicon: set[str],
    font_scale: int,
    normalize_text: bool,
) -> None:
    from lxml import etree
    from docx.oxml import OxmlElement

    emu = 635
    x = max(0, region.x * emu)
    y = max(0, region.y * emu)
    width = max(12700, region.width * emu)
    height = max(12700, region.height * emu)
    drawing = OxmlElement("w:drawing")
    anchor = OxmlElement("wp:anchor")
    for name, value in (
        ("distT", "0"),
        ("distB", "0"),
        ("distL", "0"),
        ("distR", "0"),
        ("simplePos", "0"),
        ("relativeHeight", str(251659264 + region_index)),
        ("behindDoc", "0"),
        ("locked", "0"),
        ("layoutInCell", "1"),
        ("allowOverlap", "1"),
    ):
        anchor.set(name, value)
    simple = OxmlElement("wp:simplePos")
    simple.set("x", "0")
    simple.set("y", "0")
    anchor.append(simple)
    for axis, offset in (("H", x), ("V", y)):
        position = OxmlElement(f"wp:position{axis}")
        position.set("relativeFrom", "page")
        value = OxmlElement("wp:posOffset")
        value.text = str(offset)
        position.append(value)
        anchor.append(position)
    extent = OxmlElement("wp:extent")
    extent.set("cx", str(width))
    extent.set("cy", str(height))
    anchor.append(extent)
    effect = OxmlElement("wp:effectExtent")
    for side in ("l", "t", "r", "b"):
        effect.set(side, "0")
    anchor.append(effect)
    anchor.append(OxmlElement("wp:wrapNone"))
    properties = OxmlElement("wp:docPr")
    properties.set("id", str(1_000_000 + page_index * 10_000 + region_index))
    properties.set("name", f"LayoutLoom region page {page_index + 1} text {region_index + 1}")
    properties.set("descr", "LayoutLoom editable region")
    anchor.append(properties)
    anchor.append(OxmlElement("wp:cNvGraphicFramePr"))

    graphic = OxmlElement("a:graphic")
    graphic_data = OxmlElement("a:graphicData")
    graphic_data.set("uri", _WPS_GRAPHIC_URI)
    shape = etree.Element(etree.QName(_WPS_NS, "wsp"), nsmap={"wps": _WPS_NS})
    non_visual = etree.SubElement(shape, etree.QName(_WPS_NS, "cNvSpPr"))
    non_visual.set("txBox", "1")
    shape_properties = etree.SubElement(shape, etree.QName(_WPS_NS, "spPr"))
    transform = OxmlElement("a:xfrm")
    offset = OxmlElement("a:off")
    offset.set("x", "0")
    offset.set("y", "0")
    transform.append(offset)
    inner_extent = OxmlElement("a:ext")
    inner_extent.set("cx", str(width))
    inner_extent.set("cy", str(height))
    transform.append(inner_extent)
    shape_properties.append(transform)
    geometry = OxmlElement("a:prstGeom")
    geometry.set("prst", "rect")
    geometry.append(OxmlElement("a:avLst"))
    shape_properties.append(geometry)
    shape_properties.append(OxmlElement("a:noFill"))
    line = OxmlElement("a:ln")
    line.append(OxmlElement("a:noFill"))
    shape_properties.append(line)

    text_box = etree.SubElement(shape, etree.QName(_WPS_NS, "txbx"))
    content = OxmlElement("w:txbxContent")
    previous: _TextBlock | None = None
    for block in region.blocks:
        for emitted_block in _emission_blocks(block):
            content.append(
                _append_region_paragraph(
                    document,
                    emitted_block,
                    region,
                    previous_block=previous,
                    word_lexicon=word_lexicon,
                    normalize_text=normalize_text,
                )
            )
            previous = emitted_block
    text_box.append(content)
    body = etree.SubElement(shape, etree.QName(_WPS_NS, "bodyPr"))
    for name, value in (
        ("rot", "0"),
        ("vertOverflow", "clip"),
        ("horzOverflow", "clip"),
        ("vert", "horz"),
        ("wrap", "square"),
        ("lIns", "0"),
        ("tIns", "0"),
        ("rIns", "0"),
        ("bIns", "0"),
        ("numCol", "1"),
        ("anchor", "t"),
        ("compatLnSpc", "1"),
    ):
        body.set(name, value)
    if font_scale >= 100000:
        # A 100% normAutofit still makes WPS recalculate every text box while
        # opening.  Explicit no-autofit renders identically at this scale and
        # leaves text editing responsive; the correction pass can still use a
        # bounded normAutofit when the first rendered audit fails.
        autofit = OxmlElement("a:noAutofit")
    else:
        autofit = OxmlElement("a:normAutofit")
        autofit.set("fontScale", str(max(90000, min(100000, font_scale))))
        autofit.set("lnSpcReduction", "0")
    body.append(autofit)
    graphic_data.append(shape)
    graphic.append(graphic_data)
    anchor.append(graphic)
    drawing.append(anchor)
    carrier.add_run()._r.append(drawing)


def _inline_picture_to_anchor(
    inline: Any,
    *,
    x_offset: int,
    y_offset: int,
    page_index: int,
    visual_index: int,
    visual_kind: str,
) -> None:
    from docx.oxml import OxmlElement

    children = list(inline)
    by_name = {str(child.tag).rsplit("}", 1)[-1]: child for child in children}
    required = ("extent", "docPr", "cNvGraphicFramePr", "graphic")
    if any(name not in by_name for name in required):
        raise ValidationError("无法建立区域级 Word 的定位图片")
    anchor = OxmlElement("wp:anchor")
    for name, value in (
        ("distT", "0"),
        ("distB", "0"),
        ("distL", "0"),
        ("distR", "0"),
        ("simplePos", "0"),
        ("relativeHeight", str(251658240 + visual_index)),
        ("behindDoc", "0"),
        ("locked", "1"),
        ("layoutInCell", "1"),
        ("allowOverlap", "1"),
    ):
        anchor.set(name, value)
    simple = OxmlElement("wp:simplePos")
    simple.set("x", "0")
    simple.set("y", "0")
    anchor.append(simple)
    for axis, offset in (("H", x_offset), ("V", y_offset)):
        position = OxmlElement(f"wp:position{axis}")
        position.set("relativeFrom", "page")
        value = OxmlElement("wp:posOffset")
        value.text = str(max(0, int(offset)))
        position.append(value)
        anchor.append(position)
    anchor.append(by_name["extent"])
    effect = by_name.get("effectExtent") or OxmlElement("wp:effectExtent")
    if not effect.attrib:
        for side in ("l", "t", "r", "b"):
            effect.set(side, "0")
    anchor.append(effect)
    anchor.append(OxmlElement("wp:wrapNone"))
    doc_properties = by_name["docPr"]
    doc_properties.set(
        "id",
        str(100_000_000 + page_index * 10_000 + visual_index),
    )
    safe_kind = re.sub(r"[^a-z_]", "_", str(visual_kind or "figure").casefold())
    doc_properties.set(
        "name",
        f"LayoutLoom region page {page_index + 1} visual {visual_index + 1} {safe_kind}",
    )
    doc_properties.set("descr", "LayoutLoom region visual")
    anchor.append(doc_properties)
    anchor.append(by_name["cNvGraphicFramePr"])
    anchor.append(by_name["graphic"])
    drawing = inline.getparent()
    if drawing is None:
        raise ValidationError("无法定位区域级 Word 图片容器")
    drawing.replace(inline, anchor)


def _append_visual(
    carrier: Any,
    region: _VisualRegion,
    *,
    page_index: int,
    visual_index: int,
) -> None:
    from docx.shared import Twips

    run = carrier.add_run()
    inline = run.add_picture(
        io.BytesIO(region.blob),
        width=Twips(region.width),
        height=Twips(region.height),
    )._inline
    _inline_picture_to_anchor(
        inline,
        x_offset=region.x * 635,
        y_offset=region.y * 635,
        page_index=page_index,
        visual_index=visual_index,
        visual_kind=str(getattr(region, "kind", "figure")),
    )


def _reading_order(items: list[Any], page_width: int) -> list[Any]:
    page_backgrounds = [
        item
        for item in items
        if isinstance(item, _VisualRegion)
        and str(getattr(item, "kind", "figure")) == "background_safe_page"
    ]
    items = [item for item in items if id(item) not in {id(x) for x in page_backgrounds}]
    two_columns = any(getattr(item, "column", "main") in {"left", "right"} for item in items)
    for item in items:
        if isinstance(item, _VisualRegion):
            item.column = _assign_column(item, page_width, two_columns)
    if not two_columns:
        return [*page_backgrounds, *sorted(items, key=lambda item: (item.y, item.x))]
    full = sorted(
        (item for item in items if item.column in {"full", "main"}),
        key=lambda item: (item.y, item.x),
    )
    remaining = [item for item in items if item.column not in {"full", "main"}]
    ordered: list[Any] = []
    for separator in full:
        before = [item for item in remaining if item.center_y < separator.center_y]
        for column in ("left", "right"):
            ordered.extend(sorted((item for item in before if item.column == column), key=lambda item: (item.y, item.x)))
        before_ids = {id(item) for item in before}
        remaining = [item for item in remaining if id(item) not in before_ids]
        ordered.append(separator)
    for column in ("left", "right"):
        ordered.extend(sorted((item for item in remaining if item.column == column), key=lambda item: (item.y, item.x)))
    return [*page_backgrounds, *ordered]


def _build_document(
    pages: list[_RegionPage],
    source_text: str,
    *,
    font_scale: int,
    normalize_text: bool,
) -> Any:
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
    except ImportError as exc:
        raise MissingEngineError("区域级 Word 重建需要 python-docx") from exc

    document = Document()
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    settings = document.settings.element
    compatibility = settings.find(qn("w:compat"))
    if compatibility is None:
        compatibility = OxmlElement("w:compat")
        settings.append(compatibility)
    if compatibility.find(qn("w:doNotExpandShiftReturn")) is None:
        compatibility.append(OxmlElement("w:doNotExpandShiftReturn"))
    if settings.find(qn("w:doNotAutoCompressPictures")) is None:
        settings.append(OxmlElement("w:doNotAutoCompressPictures"))
    proof_state = settings.find(qn("w:proofState"))
    if proof_state is None:
        proof_state = OxmlElement("w:proofState")
        settings.append(proof_state)
    proof_state.set(qn("w:spelling"), "clean")
    proof_state.set(qn("w:grammar"), "clean")
    word_lexicon = {
        token.casefold() for token in re.findall(r"[A-Za-z]{4,}", source_text)
    }
    previous_size: tuple[int, int] | None = None
    for page_index, page in enumerate(pages):
        page_size = (page.width, page.height)
        if page_index == 0:
            section = document.sections[0]
            carrier = _new_page_carrier(document, page_break=False)
        elif page_size != previous_size:
            section = document.add_section(WD_SECTION.NEW_PAGE)
            carrier = _new_page_carrier(document, page_break=False)
        else:
            section = document.sections[-1]
            carrier = _new_page_carrier(document, page_break=True)
        _configure_section(section, page.width, page.height)

        items = _reading_order([*page.text_regions, *page.visual_regions], page.width)
        text_index = 0
        visual_index = 0
        for item in items:
            if isinstance(item, _TextRegion):
                _append_text_box(
                    document,
                    carrier,
                    item,
                    page_index=page_index,
                    region_index=text_index,
                    word_lexicon=word_lexicon,
                    font_scale=font_scale,
                    normalize_text=normalize_text,
                )
                text_index += 1
            else:
                _append_visual(
                    carrier,
                    item,
                    page_index=page_index,
                    visual_index=visual_index,
                )
                visual_index += 1
        previous_size = page_size
    return document


def _planned_editable_text(
    pages: list[_RegionPage],
    *,
    normalize_text: bool,
) -> str:
    def frame_text(frame: _FrameLine) -> str:
        runs = list(frame.paragraph.runs)
        if not runs:
            return (
                _normalized_source_text(frame.text)
                if normalize_text
                else frame.text
            )
        combined = "".join(run.text for run in runs)
        return _normalized_source_text(combined) if normalize_text else combined

    return "\n".join(
        frame_text(frame)
        for page in pages
        for region in (
            item
            for item in _reading_order(
                [*page.text_regions, *page.visual_regions], page.width
            )
            if isinstance(item, _TextRegion)
        )
        for block in region.blocks
        for line in block.lines
        for frame in line.frames
    )


def _text_from_document(document: Any) -> str:
    paragraphs: list[str] = []
    for paragraph in document.element.body.xpath(".//w:txbxContent/w:p"):
        pieces: list[str] = []
        for node in paragraph.iter():
            local_name = str(node.tag).rsplit("}", 1)[-1]
            if local_name == "t":
                pieces.append(str(node.text or ""))
            elif local_name == "tab":
                pieces.append("\t")
            elif local_name in {"br", "cr"}:
                pieces.append("\n")
        paragraphs.append("".join(pieces))
    return "\n".join(paragraphs)


def _text_counter_from_document(document: Any) -> Counter[str]:
    return _word_counter(_text_from_document(document))


def _counter_recall(source: Counter[str], output: Counter[str]) -> tuple[int, int, float]:
    total = sum(source.values())
    retained = sum(min(count, output.get(token, 0)) for token, count in source.items())
    return total, retained, retained / total if total else 1.0


def _normalized_integrity_text(text: str) -> str:
    normalized = unicodedata.normalize("NFKC", str(text or ""))
    return "".join(
        character
        for character in normalized
        if not character.isspace() and unicodedata.category(character) != "Cf"
    )


def _sequence_coverage(source: str, output: str) -> float:
    if not source:
        return 1.0
    if not output:
        return 0.0
    width = 3 if len(source) >= 12 else 1
    source_sequences = {
        source[index : index + width]
        for index in range(max(1, len(source) - width + 1))
    }
    output_sequences = {
        output[index : index + width]
        for index in range(max(1, len(output) - width + 1))
    }
    return len(source_sequences & output_sequences) / max(1, len(source_sequences))


def _integrity_english_words(text: str) -> list[str]:
    return [
        token.casefold()
        for token in _INTEGRITY_ENGLISH_WORD.findall(
            unicodedata.normalize("NFKC", str(text or ""))
        )
    ]


def _multiset_recall(source: Iterable[Any], output: Iterable[Any]) -> float:
    source_counter = Counter(source)
    if not source_counter:
        return 1.0
    output_counter = Counter(output)
    return sum((source_counter & output_counter).values()) / sum(
        source_counter.values()
    )


def _text_integrity_metrics(source_text: str, output_text: str) -> _TextIntegrityMetrics:
    source_tokens = _word_counter(source_text)
    output_tokens = _word_counter(output_text)
    token_total, token_retained, token_recall = _counter_recall(
        source_tokens, output_tokens
    )

    source_compact = _normalized_integrity_text(source_text)
    output_compact = _normalized_integrity_text(output_text)
    source_characters = Counter(source_compact)
    output_characters = Counter(output_compact)
    matched_characters = sum((source_characters & output_characters).values())
    source_character_count = sum(source_characters.values())
    output_character_count = sum(output_characters.values())
    character_recall = (
        matched_characters / source_character_count
        if source_character_count
        else 1.0
    )
    character_precision = (
        matched_characters / output_character_count
        if output_character_count
        else (1.0 if not source_character_count else 0.0)
    )

    source_words = _integrity_english_words(source_text)
    output_words = _integrity_english_words(output_text)
    return _TextIntegrityMetrics(
        token_total,
        token_retained,
        token_recall,
        source_character_count,
        matched_characters,
        character_recall,
        character_precision,
        _sequence_coverage(source_compact, output_compact),
        _multiset_recall(source_words, output_words),
        _multiset_recall(
            zip(source_words, source_words[1:]),
            zip(output_words, output_words[1:]),
        ),
    )


def _validate_text_integrity(
    source_text: str,
    output_text: str,
) -> _TextIntegrityMetrics:
    """Reject real content loss without mistaking coordinate fragments for loss.

    Fixed-layout sources frequently split one visible word into several
    positioned frames (for example ``adap`` + ``tive``).  The region rebuild
    correctly joins those fragments, so a whitespace-delimited token counter
    alone reports false losses.  Non-whitespace character, character-order,
    English-word, and adjacent-word checks retain the strict safety guarantee
    while allowing only boundary reconstruction.
    """

    metrics = _text_integrity_metrics(source_text, output_text)
    exact_boundary_reconstruction = bool(
        metrics.character_recall >= 0.9995
        and metrics.character_precision >= 0.9995
        and metrics.sequence_coverage >= 0.995
    )
    if (
        metrics.character_recall < _MIN_CHARACTER_RECALL
        or metrics.character_precision < _MIN_CHARACTER_PRECISION
        or metrics.sequence_coverage < _MIN_SEQUENCE_COVERAGE
        or (
            not exact_boundary_reconstruction
            and (
                metrics.english_word_recall < _MIN_ENGLISH_WORD_RECALL
                or metrics.adjacent_word_recall < _MIN_ADJACENT_WORD_RECALL
            )
        )
    ):
        raise ValidationError(
            "区域级 Word 文字完整度检查未通过："
            f"字符保留 {metrics.character_recall:.1%}、"
            f"字符纯度 {metrics.character_precision:.1%}、"
            f"字符序列 {metrics.sequence_coverage:.1%}、"
            f"英文词 {metrics.english_word_recall:.1%}、"
            f"相邻词序 {metrics.adjacent_word_recall:.1%}"
        )
    return metrics


def _validate_planned_frame_visual_overlap(pages: Iterable[_RegionPage]) -> None:
    """Reject substantive glyph/precise-visual collisions, not broad bboxes.

    Figures, coloured panels, and table backgrounds often surround editable
    glyph holes, so their bounding rectangles are not proof of a collision.
    Exact formula/visual islands are checked in both axes and later again by
    the WPS rendered-pixel audit.
    """

    for page in pages:
        frames = {
            id(frame): frame
            for region in page.text_regions
            for block in region.blocks
            for line in block.lines
            for frame in line.frames
        }.values()
        for frame in frames:
            for visual in page.visual_regions:
                kind = str(getattr(visual, "kind", "figure"))
                if kind.startswith(_BACKGROUND_SAFE_VISUAL_PREFIX):
                    # These crops come from the redacted page background, not
                    # the WPS composite render.  Editable spans are absent by
                    # construction; the unchanged local pixel audit remains
                    # the final proof against accidental visual duplication.
                    continue
                if kind not in {
                    "formula",
                    "formula_unconfirmed",
                    "formula_line",
                    "formula_exact",
                    "inline_math",
                    "text_visual",
                }:
                    continue
                intersection_width = min(frame.x1, visual.x1) - max(
                    frame.x, visual.x
                )
                intersection_height = min(frame.y1, visual.y1) - max(
                    frame.y, visual.y
                )
                if intersection_width <= 0 or intersection_height <= 0:
                    continue
                frame_center_inside = bool(
                    visual.x <= frame.x + frame.width / 2.0 <= visual.x1
                    and visual.y <= frame.y + frame.height / 2.0 <= visual.y1
                )
                visual_center_inside = bool(
                    frame.x <= visual.x + visual.width / 2.0 <= frame.x1
                    and frame.y <= visual.y + visual.height / 2.0 <= frame.y1
                )
                horizontal_fraction = intersection_width / max(
                    1, min(frame.width, visual.width)
                )
                vertical_fraction = intersection_height / max(
                    1, min(frame.height, visual.height)
                )
                area_fraction = (
                    intersection_width * intersection_height
                ) / max(
                    1,
                    min(
                        frame.width * frame.height,
                        visual.width * visual.height,
                    ),
                )
                substantive = bool(
                    horizontal_fraction >= 0.18
                    and vertical_fraction >= 0.30
                    and area_fraction >= 0.06
                )
                if frame_center_inside or visual_center_inside or substantive:
                    raise ValidationError(
                        f"区域级 Word 第 {page.index + 1} 页可编辑字形与"
                        f"{kind} 视觉层重叠（水平 {horizontal_fraction:.1%}、"
                        f"垂直 {vertical_fraction:.1%}、面积 {area_fraction:.1%}）"
                    )


def _text_region_anchor_rectangle(
    region: _TextRegion,
) -> tuple[int, int, int, int]:
    """Return the exact EMU rectangle written by ``_append_text_box``."""

    emu = 635
    x = max(0, region.x * emu)
    y = max(0, region.y * emu)
    width = max(12700, region.width * emu)
    height = max(12700, region.height * emu)
    return x, y, x + width, y + height


def _text_regions_have_substantive_frame_overlap(
    first: _TextRegion,
    second: _TextRegion,
) -> bool:
    """Return whether transparent regions contain colliding text frames."""

    first_frames = {
        id(frame): frame
        for block in first.blocks
        for line in block.lines
        for frame in line.frames
    }.values()
    second_frames = {
        id(frame): frame
        for block in second.blocks
        for line in block.lines
        for frame in line.frames
    }.values()
    for first_frame in first_frames:
        for second_frame in second_frames:
            if id(first_frame) == id(second_frame):
                return True
            intersection_width = min(first_frame.x1, second_frame.x1) - max(
                first_frame.x, second_frame.x
            )
            intersection_height = min(first_frame.y1, second_frame.y1) - max(
                first_frame.y, second_frame.y
            )
            if intersection_width <= 0 or intersection_height <= 0:
                continue
            horizontal_fraction = intersection_width / max(
                1, min(first_frame.width, second_frame.width)
            )
            vertical_fraction = intersection_height / max(
                1, min(first_frame.height, second_frame.height)
            )
            area_fraction = (
                intersection_width * intersection_height
            ) / max(
                1,
                min(
                    first_frame.width * first_frame.height,
                    second_frame.width * second_frame.height,
                ),
            )
            if (
                horizontal_fraction >= 0.18
                and vertical_fraction >= 0.30
                and area_fraction >= 0.06
            ):
                return True
    return False


def _validate_page_text_region_reading_order(
    pages: Iterable[_RegionPage],
    actual_rectangles: Mapping[int, Iterable[tuple[int, int, int, int]]],
) -> None:
    """Verify each page's serialized text-box order against its layout plan.

    Word exposes every DrawingML text box as a separate text-frame story.  The
    order in ``document.xml`` therefore determines the order seen by selection,
    accessibility, and callers enumerating those stories.  Visual anchors may
    be interleaved freely and are deliberately ignored here; they only help the
    existing layout planner decide where a genuine full-width separator sits.
    """

    for page in pages:
        planned_regions = [
            item
            for item in _reading_order(
                [*page.text_regions, *page.visual_regions], page.width
            )
            if isinstance(item, _TextRegion)
        ]
        expected = [
            _text_region_anchor_rectangle(region) for region in planned_regions
        ]
        actual = list(actual_rectangles.get(page.index + 1, ()))
        if Counter(actual) != Counter(expected):
            raise ValidationError(
                f"区域级 Word 第 {page.index + 1} 页文本区域坐标与规划不一致，"
                "无法确认阅读顺序"
            )
        if actual == expected:
            continue

        planned_by_rectangle = {
            rectangle: region
            for rectangle, region in zip(expected, planned_regions)
        }
        mismatch = next(
            index
            for index, (expected_box, actual_box) in enumerate(
                zip(expected, actual)
            )
            if expected_box != actual_box
        )
        expected_region = planned_by_rectangle[expected[mismatch]]
        actual_region = planned_by_rectangle[actual[mismatch]]
        expected_column = str(expected_region.column or "main")
        actual_column = str(actual_region.column or "main")
        expected_heading = bool(
            expected_column in {"full", "main"}
            and any(block.heading for block in expected_region.blocks)
        )
        actual_heading = bool(
            actual_column in {"full", "main"}
            and any(block.heading for block in actual_region.blocks)
        )
        if {
            expected_column,
            actual_column,
        } == {"left", "right"}:
            issue = "双栏/局部分栏区域整体逆序"
        elif expected_heading != actual_heading:
            issue = "跨栏标题顺序异常"
        else:
            issue = "文本区域阅读顺序异常"
        raise ValidationError(
            f"区域级 Word 第 {page.index + 1} 页{issue}："
            f"第 {mismatch + 1} 个文本区域应为 {expected_column} 栏，"
            f"实际为 {actual_column} 栏"
        )


def _inspect_and_validate_structure(
    output: Path,
    pages: list[_RegionPage],
    source_text: str,
    source_frames: int,
) -> RegionWordInspection:
    try:
        from docx import Document
        from lxml import etree
    except ImportError as exc:
        raise MissingEngineError("区域级 Word 结构检查需要 python-docx 与 lxml") from exc

    document = Document(output)
    body = document.element.body
    if body.xpath(".//w:pPr/w:framePr"):
        raise ValidationError("区域级 Word 仍包含旧式逐行定位框")
    if body.xpath(".//w:fitText"):
        raise ValidationError("区域级 Word 仍包含会导致移动端拉伸的 fitText")
    drawing_ids = [
        str(item.get("id", ""))
        for item in body.xpath(".//wp:docPr")
        if str(item.get("id", ""))
    ]
    if len(drawing_ids) != len(set(drawing_ids)):
        raise ValidationError("区域级 Word 存在重复的绘图对象 ID")
    text_boxes = body.xpath(
        ".//*[namespace-uri()='http://schemas.microsoft.com/office/word/2010/wordprocessingShape' and local-name()='txbx']"
    )
    visual_properties = body.xpath(
        ".//wp:docPr[@descr='LayoutLoom region visual' or "
        "@descr='DocuForge region visual']"
    )
    page_numbers: set[int] = set()
    rectangles: dict[int, list[tuple[int, int, int, int]]] = {}
    visual_rectangles: dict[
        int, list[tuple[tuple[int, int, int, int], str]]
    ] = {}
    for anchor in body.xpath(".//wp:anchor"):
        properties = anchor.find("{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr")
        name = str(properties.get("name", "")) if properties is not None else ""
        match = _PAGE_NAME.match(name)
        if match is None:
            continue
        page_number = int(match.group(1))
        page_numbers.add(page_number)
        description = str(properties.get("descr", "")) if properties is not None else ""
        if description not in {
            "LayoutLoom editable region",
            "LayoutLoom region visual",
            "DocuForge editable region",
            "DocuForge region visual",
        }:
            continue
        x_node = anchor.find("{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionH/{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}posOffset")
        y_node = anchor.find("{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionV/{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}posOffset")
        extent = anchor.find("{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent")
        if x_node is None or y_node is None or extent is None:
            raise ValidationError("区域级 Word 存在缺少坐标的文本区域")
        x = int(x_node.text or 0)
        y = int(y_node.text or 0)
        width = int(extent.get("cx", "0"))
        height = int(extent.get("cy", "0"))
        if width <= 0 or height <= 0:
            raise ValidationError("区域级 Word 存在无效尺寸的文本区域")
        rectangle = (x, y, x + width, y + height)
        if description in {"LayoutLoom editable region", "DocuForge editable region"}:
            rectangles.setdefault(page_number, []).append(rectangle)
        else:
            visual_match = _VISUAL_NAME.match(name)
            kind = (
                str(visual_match.group(3) or "figure")
                if visual_match is not None
                else "figure"
            )
            visual_rectangles.setdefault(page_number, []).append((rectangle, kind))
    expected_pages = set(range(1, len(pages) + 1))
    if page_numbers != expected_pages:
        raise ValidationError(
            f"区域级 Word 页面锚点不完整：应有 {len(pages)} 页，实际识别 {len(page_numbers)} 页"
        )
    _validate_page_text_region_reading_order(pages, rectangles)
    planned_regions_by_page = {
        page.index + 1: [
            item
            for item in _reading_order(
                [*page.text_regions, *page.visual_regions], page.width
            )
            if isinstance(item, _TextRegion)
        ]
        for page in pages
    }
    for page_number, boxes in rectangles.items():
        planned_regions = planned_regions_by_page.get(page_number, ())
        if len(planned_regions) != len(boxes):
            raise ValidationError(
                f"Region Word page {page_number} text-region plan count mismatch"
            )
        for index, first in enumerate(boxes):
            for second_index, second in enumerate(
                boxes[index + 1 :],
                start=index + 1,
            ):
                x0 = max(first[0], second[0])
                y0 = max(first[1], second[1])
                x1 = min(first[2], second[2])
                y1 = min(first[3], second[3])
                if x1 <= x0 or y1 <= y0:
                    continue
                overlap = (x1 - x0) * (y1 - y0)
                minimum = min(
                    (first[2] - first[0]) * (first[3] - first[1]),
                    (second[2] - second[0]) * (second[3] - second[1]),
                )
                if overlap / max(1, minimum) > 0.12:
                    first_region = planned_regions[index]
                    second_region = planned_regions[second_index]
                    transparent_precision_overlap = bool(
                        first_region.precision
                        and second_region.precision
                        and not _text_regions_have_substantive_frame_overlap(
                            first_region,
                            second_region,
                        )
                    )
                    if transparent_precision_overlap:
                        # Stacked fractions and table fragments may have
                        # intersecting transparent shape rectangles while the
                        # actual positioned glyph frames stay disjoint.  The
                        # rendered overlap detector and WPS pixel audit remain
                        # the second proof against double printing.
                        continue
                    raise ValidationError(
                        f"区域级 Word 第 {page_number} 页存在明显重叠的文本区域"
                    )
        page_visuals = visual_rectangles.get(page_number, ())
        protected_visuals = [
            (rectangle, kind)
            for rectangle, kind in page_visuals
            if not kind.startswith("table") and kind != "figure"
        ]
        for text_box in boxes:
            for visual_box, visual_kind in protected_visuals:
                if visual_kind.startswith(_BACKGROUND_SAFE_VISUAL_PREFIX) or visual_kind in {
                    "formula_exact",
                    "inline_math",
                    "text_visual",
                }:
                    # These visual islands intentionally sit inside a
                    # tab-aligned precision row.  Actual glyph conflicts were
                    # rejected during planning and are verified again by the
                    # local rendered-pixel audit.
                    continue
                x0 = max(text_box[0], visual_box[0])
                y0 = max(text_box[1], visual_box[1])
                x1 = min(text_box[2], visual_box[2])
                y1 = min(text_box[3], visual_box[3])
                if x1 <= x0 or y1 <= y0:
                    continue
                overlap = (x1 - x0) * (y1 - y0)
                minimum = min(
                    (text_box[2] - text_box[0]) * (text_box[3] - text_box[1]),
                    (visual_box[2] - visual_box[0])
                    * (visual_box[3] - visual_box[1]),
                )
                overlap_limit = (
                    0.08
                    if visual_kind.startswith("formula")
                    else 0.60
                )
                if overlap / max(1, minimum) > overlap_limit:
                    raise ValidationError(
                        f"区域级 Word 第 {page_number} 页文字区域与"
                        f"{visual_kind} 视觉层重叠"
                    )
        for index, (first, first_kind) in enumerate(page_visuals):
            for second, second_kind in page_visuals[index + 1 :]:
                precise_kinds = {"formula_exact", "inline_math", "text_visual"}
                if not (
                    {first_kind, second_kind} & precise_kinds
                    or first_kind.startswith(_BACKGROUND_SAFE_VISUAL_PREFIX)
                    or second_kind.startswith(_BACKGROUND_SAFE_VISUAL_PREFIX)
                ):
                    continue
                x0 = max(first[0], second[0])
                y0 = max(first[1], second[1])
                x1 = min(first[2], second[2])
                y1 = min(first[3], second[3])
                if x1 <= x0 or y1 <= y0:
                    continue
                overlap = (x1 - x0) * (y1 - y0)
                minimum = min(
                    (first[2] - first[0]) * (first[3] - first[1]),
                    (second[2] - second[0]) * (second[3] - second[1]),
                )
                if overlap / max(1, minimum) >= 0.85:
                    raise ValidationError(
                        f"区域级 Word 第 {page_number} 页存在重复视觉层"
                    )
    output_text = _text_from_document(document)
    integrity = _validate_text_integrity(source_text, output_text)
    region_count = len(text_boxes)
    if source_frames >= 8 and region_count >= source_frames:
        raise ValidationError("区域级 Word 未能减少逐行定位节点")
    expected_visuals = sum(len(page.visual_regions) for page in pages)
    if len(visual_properties) != expected_visuals:
        raise ValidationError(
            f"区域级 Word 图像/公式回填不完整：应有 {expected_visuals} 个，实际 {len(visual_properties)} 个"
        )
    paragraph_count = len(body.xpath(".//w:txbxContent/w:p"))
    if paragraph_count < region_count:
        raise ValidationError("区域级 Word 存在没有普通段落的文本区域")
    return RegionWordInspection(
        source_pages=len(pages),
        source_frames=source_frames,
        output_text_regions=region_count,
        output_visual_regions=expected_visuals,
        output_paragraphs=paragraph_count,
        source_words=integrity.source_tokens,
        retained_words=integrity.retained_tokens,
        word_recall=integrity.token_recall,
        structural_passed=True,
        rendered_passed=False,
        used_correction_pass=False,
        character_recall=integrity.character_recall,
        character_precision=integrity.character_precision,
        sequence_coverage=integrity.sequence_coverage,
        english_word_recall=integrity.english_word_recall,
        adjacent_word_recall=integrity.adjacent_word_recall,
    )


def _page_difference(first: Any, second: Any) -> tuple[float, float]:
    from PIL import Image, ImageChops

    first_pixmap = first.get_pixmap(dpi=96, alpha=False)
    second_pixmap = second.get_pixmap(dpi=96, alpha=False)
    if (first_pixmap.width, first_pixmap.height) != (
        second_pixmap.width,
        second_pixmap.height,
    ):
        return 255.0, 1.0
    first_image = Image.frombytes(
        "RGB", (first_pixmap.width, first_pixmap.height), first_pixmap.samples
    )
    second_image = Image.frombytes(
        "RGB", (second_pixmap.width, second_pixmap.height), second_pixmap.samples
    )
    difference = ImageChops.difference(first_image, second_image).convert("L")
    histogram = difference.histogram()
    pixels = max(1, first_pixmap.width * first_pixmap.height)
    mean = sum(value * count for value, count in enumerate(histogram)) / pixels
    changed = sum(histogram[25:]) / pixels
    return float(mean), float(changed)


def _local_visual_difference(
    first_page: Any,
    second_page: Any,
    regions: Iterable[tuple[int, int, int, int]],
) -> tuple[float, float, float]:
    """Return worst changed, extra-ink, and missing-ink fractions."""

    try:
        import pymupdf
        from PIL import Image, ImageChops
    except ImportError as exc:
        raise MissingEngineError("区域级 Word 局部视觉复检需要 PyMuPDF 与 Pillow") from exc

    worst_changed = 0.0
    worst_extra = 0.0
    worst_missing = 0.0
    page_rect = first_page.rect
    for raw_x0, raw_y0, raw_x1, raw_y1 in regions:
        margin = 3.0
        clip = pymupdf.Rect(
            float(page_rect.x0) + raw_x0 / 20.0 - margin,
            float(page_rect.y0) + raw_y0 / 20.0 - margin,
            float(page_rect.x0) + raw_x1 / 20.0 + margin,
            float(page_rect.y0) + raw_y1 / 20.0 + margin,
        ) & page_rect
        if clip.is_empty or clip.is_infinite:
            continue
        first_pixmap = first_page.get_pixmap(dpi=144, clip=clip, alpha=False)
        second_pixmap = second_page.get_pixmap(dpi=144, clip=clip, alpha=False)
        if (first_pixmap.width, first_pixmap.height) != (
            second_pixmap.width,
            second_pixmap.height,
        ):
            return 1.0, 1.0, 1.0
        first_image = Image.frombytes(
            "RGB",
            (first_pixmap.width, first_pixmap.height),
            first_pixmap.samples,
        ).convert("L")
        second_image = Image.frombytes(
            "RGB",
            (second_pixmap.width, second_pixmap.height),
            second_pixmap.samples,
        ).convert("L")
        pixels = max(1, first_image.width * first_image.height)
        changed = ImageChops.difference(first_image, second_image).histogram()
        extra = ImageChops.subtract(first_image, second_image).histogram()
        missing = ImageChops.subtract(second_image, first_image).histogram()
        worst_changed = max(worst_changed, sum(changed[25:]) / pixels)
        worst_extra = max(worst_extra, sum(extra[25:]) / pixels)
        worst_missing = max(worst_missing, sum(missing[25:]) / pixels)
    return float(worst_changed), float(worst_extra), float(worst_missing)


def _has_rendered_text_overlap(page: Any) -> bool:
    lines: list[tuple[float, float, float, float, str]] = []
    payload = page.get_text("dict")
    for block in payload.get("blocks", []):
        if int(block.get("type", 0)) != 0:
            continue
        for line in block.get("lines", []):
            spans = list(line.get("spans", []))
            text = "".join(str(span.get("text", "")) for span in spans).strip()
            if len(text) < 3 or not spans:
                continue
            x0 = min(float(span.get("bbox", (0, 0, 0, 0))[0]) for span in spans)
            x1 = max(float(span.get("bbox", (0, 0, 0, 0))[2]) for span in spans)
            baselines = [float(span.get("origin", (0, 0))[1]) for span in spans]
            sizes = [max(1.0, float(span.get("size", 1.0))) for span in spans]
            baseline = statistics.median(baselines)
            size = statistics.median(sizes)
            if x1 > x0:
                lines.append((x0, x1, baseline, size, text))
    for index, first in enumerate(lines):
        for second in lines[index + 1 :]:
            if abs(first[2] - second[2]) > min(first[3], second[3]) * 0.30:
                continue
            overlap = min(first[1], second[1]) - max(first[0], second[0])
            minimum_width = min(first[1] - first[0], second[1] - second[0])
            if overlap / max(1.0, minimum_width) > 0.35:
                if first[4].casefold() != second[4].casefold():
                    return True
    return False


def _render_source_for_regions(
    source: Path,
    folder: Path,
    *,
    engine: str,
    timeout: float,
    progress: Callable[[float, str], None] | None = None,
) -> Path:
    from docuforge.processors.office import convert_with_office

    output = folder / "source"
    output.mkdir(parents=True, exist_ok=True)
    _progress_reporter(progress)(
        0.07,
        "使用 WPS/Office 预渲染原 Word，提取精确公式行",
    )
    return convert_with_office(
        source,
        output,
        "pdf",
        engine=engine,
        overwrite=True,
        timeout=timeout,
    )[0]


def _rasterized_source_pages(source_pdf: Path) -> list[bytes]:
    try:
        import pymupdf
    except ImportError as exc:
        raise MissingEngineError("区域级 Word 公式行提取需要 PyMuPDF") from exc

    document = pymupdf.open(source_pdf)
    try:
        return [
            page.get_pixmap(dpi=192, alpha=False).tobytes("png")
            for page in document
        ]
    finally:
        document.close()


def _render_audit(
    source: Path,
    output: Path,
    *,
    engine: str,
    timeout: float,
    expected_pages: int,
    expected_text: str,
    source_pdf: Path | None = None,
    progress: Callable[[float, str], None] | None = None,
    max_mean_difference: float = 32.0,
    max_changed_fraction: float = 0.58,
    precise_visual_regions_by_page: Mapping[
        int, Iterable[tuple[int, int, int, int]]
    ]
    | None = None,
) -> _RenderAudit:
    try:
        import pymupdf
    except ImportError as exc:
        raise MissingEngineError("区域级 Word 二重渲染检查需要 PyMuPDF") from exc

    from docuforge.processors.office import convert_with_office
    emit_progress = _progress_reporter(progress)

    with tempfile.TemporaryDirectory(prefix="docuforge-word-region-check-") as name:
        root = Path(name)
        output_dir = root / "output"
        output_dir.mkdir()
        if source_pdf is None:
            source_dir = root / "source"
            source_dir.mkdir()
            emit_progress(0.82, "二重检查：使用 WPS/Office 渲染原 Word")
            source_pdf = convert_with_office(
                source,
                source_dir,
                "pdf",
                engine=engine,
                overwrite=True,
                timeout=timeout,
            )[0]
        emit_progress(0.88, "二重检查：渲染区域级 Word 并检测重叠、缺字和分页")
        output_pdf = convert_with_office(
            output,
            output_dir,
            "pdf",
            engine=engine,
            overwrite=True,
            timeout=timeout,
        )[0]
        source_document = pymupdf.open(source_pdf)
        output_document = pymupdf.open(output_pdf)
        try:
            if source_document.page_count != expected_pages:
                return _RenderAudit(False, f"原 Word 渲染页数异常：{source_document.page_count}")
            if output_document.page_count != expected_pages:
                return _RenderAudit(
                    False,
                    f"区域级 Word 分页异常：应有 {expected_pages} 页，实际 {output_document.page_count} 页",
                    int(output_document.page_count),
                )
            # The fixed-layout source may deliberately use fitText, and WPS PDF
            # extraction can then split one visual word into individual letters.
            # Compare the rendered result with the source DOCX text nodes instead.
            source_counter = _word_counter(expected_text)
            output_counter = _word_counter(
                "\n".join(page.get_text("text") for page in output_document)
            )
            _total, _retained, recall = _counter_recall(source_counter, output_counter)
            if source_counter and recall < 0.965:
                return _RenderAudit(
                    False,
                    f"渲染后文字保留率不足：{recall:.1%}",
                    int(output_document.page_count),
                    recall,
                )
            worst_mean = 0.0
            worst_changed = 0.0
            for index in range(expected_pages):
                emit_progress(
                    0.90 + 0.07 * ((index + 1) / max(1, expected_pages)),
                    f"二重检查：逐页核对 {index + 1}/{expected_pages}",
                )
                if _has_rendered_text_overlap(output_document[index]):
                    return _RenderAudit(
                        False,
                        f"第 {index + 1} 页检测到疑似文字重叠",
                        expected_pages,
                        recall,
                    )
                local_regions = tuple(
                    (precise_visual_regions_by_page or {}).get(index, ())
                )
                if local_regions:
                    local_changed, extra_ink, missing_ink = _local_visual_difference(
                        source_document[index],
                        output_document[index],
                        local_regions,
                    )
                    if (
                        local_changed > 0.08
                        or extra_ink > 0.04
                        or missing_ink > 0.04
                    ):
                        return _RenderAudit(
                            False,
                            f"第 {index + 1} 页局部视觉复检未通过："
                            f"变化 {local_changed:.1%}、"
                            f"额外墨迹 {extra_ink:.1%}、"
                            f"缺失墨迹 {missing_ink:.1%}",
                            expected_pages,
                            recall,
                        )
                mean, changed = _page_difference(
                    source_document[index], output_document[index]
                )
                worst_mean = max(worst_mean, mean)
                worst_changed = max(worst_changed, changed)
            if (
                worst_mean > float(max_mean_difference)
                or worst_changed > float(max_changed_fraction)
            ):
                return _RenderAudit(
                    False,
                    "区域级版面与原 Word 差异过大："
                    f"平均像素差 {worst_mean:.2f}，显著变化区域 {worst_changed:.1%}",
                    expected_pages,
                    recall,
                    worst_mean,
                    worst_changed,
                )
            return _RenderAudit(
                True,
                "",
                expected_pages,
                recall,
                worst_mean,
                worst_changed,
            )
        finally:
            source_document.close()
            output_document.close()


def build_region_compatible_word(
    input_docx: str | Path,
    output_path: str | Path,
    *,
    verification_engine: str = "auto",
    timeout: float = 300,
    overwrite: bool = False,
    progress: Callable[[float, str], None] | None = None,
    max_mean_difference: float = 32.0,
    max_changed_fraction: float = 0.58,
    preserve_editable_text: bool = False,
    normalize_text: bool = True,
    formula_hints_by_page: Mapping[
        int, Iterable[tuple[int, int, int, int]]
    ]
    | None = None,
    table_hints_by_page: Mapping[
        int, Iterable[tuple[int, int, int, int]]
    ]
    | None = None,
    visual_hints_by_page: Mapping[int, Iterable[_PreciseVisualHint]] | None = None,
) -> list[Path]:
    """Create a page-faithful, region-copyable Word copy with two-stage checks."""

    try:
        from docx import Document
    except ImportError as exc:
        raise MissingEngineError("区域级 Word 重建需要 python-docx") from exc

    emit_progress = _progress_reporter(progress)

    source = Path(input_docx).expanduser().resolve()
    if not source.is_file():
        raise ValidationError(f"Word 文件不存在：{source}")
    if source.suffix.casefold() != ".docx":
        raise ValidationError("区域级 Word 重建仅支持 DOCX 文件")
    target = unique_path(Path(output_path).expanduser().resolve(), overwrite)
    if source == target:
        raise ValidationError("区域级 Word 必须输出为新文件，不能覆盖原 Word")
    engine = str(verification_engine or "auto").casefold().strip()
    if engine not in {"auto", "wps", "microsoft_office", "libreoffice", "none"}:
        raise ValidationError("区域级 Word 复检引擎参数无效")
    try:
        timeout_value = float(timeout)
    except (TypeError, ValueError, OverflowError) as exc:
        raise ValidationError("区域级 Word 复检超时必须是大于 0 的秒数") from exc
    if timeout_value <= 0:
        raise ValidationError("区域级 Word 复检超时必须是大于 0 的秒数")
    try:
        mean_limit = float(max_mean_difference)
        changed_limit = float(max_changed_fraction)
    except (TypeError, ValueError, OverflowError) as exc:
        raise ValidationError("区域级 Word 版面差异阈值无效") from exc
    if not 0 < mean_limit <= 255 or not 0 < changed_limit <= 1:
        raise ValidationError("区域级 Word 版面差异阈值超出有效范围")

    formula_hints_by_page = _materialize_page_hints(formula_hints_by_page)
    table_hints_by_page = _materialize_page_hints(table_hints_by_page)
    visual_hints_by_page = _materialize_page_hints(visual_hints_by_page)

    emit_progress(0.03, "读取高保真 Word 的页面、坐标、文字和视觉层")
    try:
        source_document = Document(source)
    except Exception as exc:
        raise ValidationError(f"无法读取 Word 文件：{exc}") from exc
    raw_pages = _source_pages(source_document)
    source_frames = sum(len(page.frames) for page in raw_pages)
    if source_frames < 1:
        raise ValidationError(
            "未检测到旧版固定版面定位文字；该文件无需区域级重建"
        )
    render_workspace: Path | None = None
    render_workspace_token: _WorkspaceCleanupToken | None = None
    render_workspace_cleanup: weakref.finalize | None = None
    source_pdf: Path | None = None
    rendered_pages: list[bytes] | None = None
    has_formula_hints = bool(
        formula_hints_by_page is not None
        and any(tuple(values) for values in formula_hints_by_page.values())
    )
    has_visual_hints = bool(
        visual_hints_by_page is not None
        and any(tuple(values) for values in visual_hints_by_page.values())
    )
    if engine != "none":
        render_workspace = Path(
            tempfile.mkdtemp(prefix="docuforge-word-region-source-")
        )
        render_workspace_token = _WorkspaceCleanupToken()
        render_workspace_cleanup = weakref.finalize(
            render_workspace_token,
            shutil.rmtree,
            render_workspace,
            True,
        )
        source_pdf = _render_source_for_regions(
            source,
            render_workspace,
            engine=engine,
            timeout=timeout_value,
            progress=emit_progress,
        )
        if not preserve_editable_text or has_formula_hints or has_visual_hints:
            rendered_pages = _rasterized_source_pages(source_pdf)
    emit_progress(0.10, f"识别 {len(raw_pages)} 页的段落、分栏、图片和公式区域")
    pages = _plan_pages(
        source_document,
        rendered_pages,
        preserve_editable_text=bool(preserve_editable_text),
        formula_hints_by_page=formula_hints_by_page,
        table_hints_by_page=table_hints_by_page,
        visual_hints_by_page=visual_hints_by_page,
    )
    _validate_planned_frame_visual_overlap(pages)
    if not any(page.text_regions for page in pages):
        raise ValidationError("没有识别到可重建的文字区域")
    source_text = _planned_editable_text(
        pages,
        normalize_text=bool(normalize_text),
    )
    precise_visual_regions_by_page = {
        page.index: tuple(
            sorted(
                {
                    *page.precise_audit_regions,
                    *(
                        (visual.x, visual.y, visual.x1, visual.y1)
                        for visual in page.visual_regions
                        if str(getattr(visual, "kind", "figure"))
                        in {
                            "formula_exact",
                            "formula_line",
                            "inline_math",
                            "text_visual",
                        }
                        or str(getattr(visual, "kind", "figure")).startswith(
                            _BACKGROUND_SAFE_VISUAL_PREFIX
                        )
                    ),
                },
                key=lambda box: (box[1], box[0]),
            )
        )
        for page in pages
    }

    used_correction = False
    last_audit = _RenderAudit(engine == "none")
    inspection: RegionWordInspection | None = None
    with atomic_output(target) as temporary:
        for pass_index, font_scale in enumerate((100000, 96500), start=1):
            emit_progress(
                0.24 if pass_index == 1 else 0.72,
                "按坐标生成区域文本框、普通段落及定位图片"
                if pass_index == 1
                else "二重检查未通过，使用保守字体适配进行第二次重建",
            )
            document = _build_document(
                pages,
                source_text,
                font_scale=font_scale,
                normalize_text=bool(normalize_text),
            )
            document.save(temporary)
            if not temporary.is_file() or temporary.stat().st_size == 0:
                raise ValidationError("没有生成有效的区域级 Word")
            emit_progress(
                0.68,
                "第一重检查：校验区域、文字、图片、分页和旧定位节点",
            )
            inspection = _inspect_and_validate_structure(
                temporary, pages, source_text, source_frames
            )
            if engine == "none":
                break
            last_audit = _render_audit(
                source,
                temporary,
                engine=engine,
                timeout=timeout_value,
                expected_pages=len(pages),
                expected_text=source_text,
                source_pdf=source_pdf,
                progress=emit_progress,
                max_mean_difference=mean_limit,
                max_changed_fraction=changed_limit,
                precise_visual_regions_by_page=precise_visual_regions_by_page,
            )
            if last_audit.passed:
                break
            if pass_index == 1:
                used_correction = True
                continue
            raise ValidationError(f"区域级 Word 二重检查未通过：{last_audit.reason}")

    assert inspection is not None
    emit_progress(
        1.0,
        "区域级 Word 完成："
        f"{source_frames} 个逐行定位节点重建为 {inspection.output_text_regions} 个可复制区域，"
        f"回填 {inspection.output_visual_regions} 个图片/公式区域，"
        f"非空白字符保留率 {inspection.character_recall:.1%}，"
        f"字符序列保留率 {inspection.sequence_coverage:.1%}"
        + ("，已执行自动纠正重建" if used_correction else ""),
    )
    if render_workspace_cleanup is not None:
        render_workspace_cleanup()
    return [target]


__all__ = [
    "RegionWordInspection",
    "build_region_compatible_word",
]
