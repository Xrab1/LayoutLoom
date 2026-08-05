from __future__ import annotations

import contextlib
import html
import io
import json
import math
import re
import shutil
import tempfile
import threading
import time
import unicodedata
import warnings
import xml.etree.ElementTree as ET
from collections import Counter
from dataclasses import dataclass
from functools import lru_cache
from pathlib import Path
from statistics import median
from typing import Any, Callable, Iterable, Iterator, Mapping
from zipfile import BadZipFile, ZipFile

from PIL import Image

from ..models import CancelledError, DocuForgeError, MissingEngineError, ValidationError
from ..utils import atomic_output, parse_page_spec, unique_path

_SVG_URL_PATTERN = re.compile(r"url\(\s*(['\"]?)(.*?)\1\s*\)", re.IGNORECASE)
_ENGLISH_WORD_PATTERN = re.compile(r"[A-Za-z]+")
_LINE_END_HYPHENATED_ENGLISH_PATTERN = re.compile(
    r"(?P<left>[A-Za-z]+)-[^\S\r\n]*\r?\n[^\S\r\n]*(?P<right>[a-z]+)"
)
_MISENCODED_PDF_WORD_HYPHEN = "\uff27"
_MISENCODED_PDF_APOSTROPHE = "\U001001b3"
_OVERLAPPING_PDF_FRAGMENT_LETTERS = frozenset("gjpqy")

_MIN_EDITABLE_SEQUENCE_COVERAGE = 0.80
_MIN_EDITABLE_ENGLISH_WORD_RECALL = 0.80
_MIN_EDITABLE_ADJACENT_WORD_COVERAGE = 0.70
_WARN_EDITABLE_SEQUENCE_COVERAGE = 0.95
_WARN_EDITABLE_ENGLISH_WORD_RECALL = 0.95
_WARN_EDITABLE_ADJACENT_WORD_COVERAGE = 0.90

_MIN_HYBRID_PAGE_SEQUENCE_COVERAGE = 0.90
_MIN_HYBRID_PAGE_ENGLISH_WORD_RECALL = 0.90
_MIN_HYBRID_PAGE_ADJACENT_WORD_COVERAGE = 0.90
_MIN_HYBRID_HIGH_RECALL_SEQUENCE_COVERAGE = 0.90
_MIN_HYBRID_HIGH_RECALL_ENGLISH_WORD_RECALL = 0.98
_MIN_HYBRID_HIGH_RECALL_ADJACENT_WORD_COVERAGE = 0.85
_HYBRID_MIN_EDITABLE_CHARACTERS = 8
_HYBRID_MAX_SINGLE_IMAGE_RATIO = 0.10
_HYBRID_MAX_TOTAL_IMAGE_RATIO = 0.15

_PDF2DOCX_COLUMN_LAYOUTS = frozenset({"auto", "single", "double", "mixed"})
_PDF2DOCX_DEFAULT_MIN_SECTION_HEIGHT = 50.0
_PDF2DOCX_COLUMN_AWARE_MIN_SECTION_HEIGHT = 35.0
_PDF2DOCX_WORD_FLOW_BEFORE_SLACK_TWIPS = 30
_PDF2DOCX_WORD_FLOW_SECTION_SLACK_TWIPS = 60
_PDF2DOCX_WORD_FLOW_GEOMETRY_SLACK_TWIPS = 60
_PDF2DOCX_DENSE_GRAPHIC_FOOTER_BEFORE_TWIPS = 480
_PDF2DOCX_DENSE_GRAPHIC_FOOTER_SLACK_TWIPS = 960
_PDF2DOCX_DENSE_GRAPHIC_FOOTER_MIN_TEXT_CHARACTERS = 3500
_PDF2DOCX_DENSE_GRAPHIC_FOOTER_MIN_PARAGRAPH_CHARACTERS = 1200
_PDF2DOCX_DENSE_GRAPHIC_FOOTER_MIN_AVERAGE_CHARACTERS = 400
_PDF2DOCX_WPS_RENDER_MIN_PAGES = 8
_FIXED_LAYOUT_REGION_MIN_FRAMES = 48
_FIXED_LAYOUT_REGION_MIN_FRAMES_LONG_DOCUMENT = 24
_FIXED_LAYOUT_REGION_MIN_PAGES = 8
_FIXED_LAYOUT_REGION_MIN_SEQUENCE_RETENTION = 0.985
_FIXED_LAYOUT_REGION_MIN_WORD_RETENTION = 0.985
_FIXED_LAYOUT_REGION_MIN_ADJACENT_WORD_RETENTION = 0.97
_FIXED_LAYOUT_REGION_MIN_CHARACTER_RETENTION = 0.995
_FIXED_LAYOUT_REGION_DESIGNED_MIN_SEQUENCE_RETENTION = 0.90
_FIXED_LAYOUT_REGION_DESIGNED_MIN_ADJACENT_RETENTION = 0.85
_PDF_COLUMN_GUTTER_MARGIN_RATIO = 0.004

_HYBRID_FULL_PAGE_REGION_RATIO = 0.85
_HYBRID_REGION_PADDING = 2.5
_HYBRID_FIGURE_LABEL_GAP = 16.0
_HYBRID_FORMULA_LABEL_GAP = 2.0
_HYBRID_FULL_IMAGE_REDACTION_RATIO = 0.98
_MATH_FONT_TOKENS = (
    "cmmi",
    "cmsy",
    "cmex",
    "texcmmath",
    "rtxmi",
    "txmi",
    "txsy",
    "yhcmex",
    "msam",
    "msbm",
    "wasy",
    "rsfs",
)
_CAPTION_PATTERN = re.compile(
    r"^(?:fig(?:ure)?\.?|table)\s*\d+|^(?:图|表)\s*[一二三四五六七八九十百\d]+",
    re.IGNORECASE,
)
_TABLE_TITLE_PATTERN = re.compile(
    r"^(?:table)\s*[ivxlcdm\d]+\s*[:.]?|^(?:表)\s*[一二三四五六七八九十百\d]+",
    re.IGNORECASE,
)
_EQUATION_NUMBER_PATTERN = re.compile(r"^\(?\s*\d+[a-z]?\s*\)?$", re.IGNORECASE)
_DISPLAY_FORMULA_OPERATOR_PATTERN = re.compile(r"[=∑∫√⇒→≈≠≤≥±]")
_NUMBERED_FORMULA_REASON = "编号显示公式"
_UNRELIABLE_PDF_ENCODING_REASON = "无法可靠映射的异常字符"
_OVERLAPPING_PDF_GLYPH_REASON = "重叠字形碎片"
_FORMULA_FUNCTION_WORDS = frozenset(
    {
        "attention",
        "arctan",
        "avg",
        "ciou",
        "concat",
        "conv",
        "cos",
        "det",
        "diag",
        "downsample",
        "exp",
        "head",
        "iou",
        "log",
        "max",
        "mean",
        "min",
        "multihead",
        "norm",
        "outer",
        "pool",
        "relu",
        "sigmoid",
        "sin",
        "softmax",
        "sqrt",
        "tan",
        "upsample",
    }
)


@dataclass(frozen=True)
class _HybridRegion:
    page_index: int
    rect: tuple[float, float, float, float]
    kind: str
    reasons: tuple[str, ...] = ()
    dpi: int = 300


@dataclass
class _HybridPageAssessment:
    page_index: int
    source_text: str
    editable_source_text: str
    editable_text_blocks: tuple[str, ...]
    draw_items: int
    draw_bbox_max_ratio: float
    reasons: list[str]
    visual_regions: list[_HybridRegion]
    detected_two_columns: bool = False
    column_split_x: float | None = None
    column_pair_bands: tuple[tuple[float, float], ...] = ()
    anchored_visual_regions: tuple[_HybridRegion, ...] = ()


@dataclass(frozen=True)
class _PdfColumnProfile:
    split_x: float
    left_lines: tuple[tuple[tuple[float, float, float, float], int], ...]
    right_lines: tuple[tuple[tuple[float, float, float, float], int], ...]
    crossing_lines: tuple[tuple[tuple[float, float, float, float], int], ...]


@dataclass(frozen=True)
class _RenderedPageContentAssessment:
    text_characters: int
    median_font_size: float
    visual_ratio: float
    required_characters: int
    substantive: bool


@dataclass(frozen=True)
class _RenderedPageCounterConflict:
    page_number: int
    counters: tuple[tuple[int, int], ...]


@dataclass(frozen=True)
class _PdfWordLayoutProfile:
    page_count: int
    text_pages: int
    two_column_pages: int
    designed_pages: tuple[int, ...]
    fixed_layout_recommended: bool
    reasons: tuple[str, ...]


_TEMPORARY_DIRECTORY_CLEANUP_DELAYS = (0.0, 0.05, 0.15, 0.35)


def _cleanup_temporary_working_directory(folder: Path) -> None:
    """Best-effort cleanup that never replaces the real conversion result/error."""

    last_error: OSError | None = None
    for delay in _TEMPORARY_DIRECTORY_CLEANUP_DELAYS:
        if delay:
            time.sleep(delay)
        try:
            shutil.rmtree(folder)
        except FileNotFoundError:
            return
        except OSError as exc:
            last_error = exc
            retryable = isinstance(exc, PermissionError) or getattr(
                exc, "winerror", None
            ) in {5, 32}
            if not retryable:
                break
        else:
            return

    if last_error is not None:
        # Warning filters may be configured as errors by callers. A cleanup
        # diagnostic must never hide a successful conversion or its real error.
        with contextlib.suppress(Exception):
            warnings.warn(
                "PDF 转 Word 临时目录仍被系统短暂占用，已保留供系统稍后清理："
                f"{folder}（{last_error}）",
                stacklevel=2,
            )


@contextlib.contextmanager
def _temporary_working_directory(
    *,
    prefix: str,
    before_cleanup: Callable[[], None] | None = None,
) -> Iterator[Path]:
    """Create a working directory and release engines before deleting it."""

    folder = Path(tempfile.mkdtemp(prefix=prefix))
    try:
        yield folder
    finally:
        if before_cleanup is not None:
            with contextlib.suppress(Exception):
                before_cleanup()
        _cleanup_temporary_working_directory(folder)


def _validate_svg_resources(svg_bytes: bytes) -> None:
    upper = svg_bytes[:4096].upper()
    if b"<!DOCTYPE" in upper or b"<!ENTITY" in upper:
        raise ValidationError("SVG 不允许包含 DOCTYPE 或外部实体")
    try:
        root = ET.fromstring(svg_bytes)
    except ET.ParseError as exc:
        raise ValidationError(f"SVG XML 格式无效：{exc}") from exc

    def validate_reference(value: str) -> None:
        reference = value.strip().strip("'\"")
        if (
            not reference
            or reference.startswith("#")
            or reference.lower().startswith("data:")
        ):
            return
        raise ValidationError("SVG 不允许引用外部文件或网络资源；请先把资源嵌入 SVG")

    if re.search(rb"@import\s+", svg_bytes, re.IGNORECASE):
        raise ValidationError("SVG 不允许通过 CSS @import 引用外部样式")

    for element in root.iter():
        for name, value in element.attrib.items():
            local_name = name.rsplit("}", 1)[-1].lower()
            if local_name in {"href", "src"}:
                validate_reference(value)
            if local_name == "style":
                for match in _SVG_URL_PATTERN.finditer(value):
                    validate_reference(match.group(2))
        if element.text and element.tag.rsplit("}", 1)[-1].lower() == "style":
            for match in _SVG_URL_PATTERN.finditer(element.text):
                validate_reference(match.group(2))


def _unlock_reader(path: Path, password: str | None = None):
    from pypdf import PdfReader

    reader = PdfReader(path)
    if reader.is_encrypted:
        if not password or not reader.decrypt(password):
            raise ValidationError(f"PDF 已加密或密码不正确：{path.name}")
    return reader


def _require_pymupdf():
    try:
        import pymupdf
    except ImportError as exc:
        raise MissingEngineError(
            "PDF 转 Word 需要 PyMuPDF；请重新运行安装程序补齐依赖"
        ) from exc
    return pymupdf


def _open_pymupdf_document(source: Path, password: str | None):
    pymupdf = _require_pymupdf()
    try:
        document = pymupdf.open(str(source))
    except Exception as exc:
        raise ValidationError(f"无法打开 PDF：{source.name}（{exc}）") from exc
    if document.needs_pass:
        authenticated = bool(password and document.authenticate(password))
        if not authenticated:
            document.close()
            raise ValidationError(f"PDF 已加密或密码不正确：{source.name}")
    if document.page_count < 1:
        document.close()
        raise ValidationError("PDF 没有可转换的页面")
    return pymupdf, document


def _is_ascii_letter(character: str) -> bool:
    return len(character) == 1 and character.isascii() and character.isalpha()


def _repair_known_pdf_text_encoding(text: str) -> str:
    """Repair only PDF character mappings whose surrounding text proves meaning.

    Some embedded teaching-material fonts expose a visual hyphen as full-width
    ``G`` and a visual apostrophe as a supplementary private-use character.
    Both substitutions are safe only inside an ASCII word; standalone symbols
    remain untouched and are handled by the visual-fallback detector instead.
    """

    if not text or not any(
        character in text
        for character in (_MISENCODED_PDF_WORD_HYPHEN, _MISENCODED_PDF_APOSTROPHE)
    ):
        return text
    characters = list(text)
    for index, character in enumerate(characters):
        if character not in {
            _MISENCODED_PDF_WORD_HYPHEN,
            _MISENCODED_PDF_APOSTROPHE,
        }:
            continue
        previous = characters[index - 1] if index > 0 else ""
        following = characters[index + 1] if index + 1 < len(characters) else ""
        previous_nonspace = next(
            (
                characters[candidate]
                for candidate in range(index - 1, -1, -1)
                if not characters[candidate].isspace()
            ),
            "",
        )
        following_nonspace = next(
            (
                characters[candidate]
                for candidate in range(index + 1, len(characters))
                if not characters[candidate].isspace()
            ),
            "",
        )
        attached_word_context = (
            _is_ascii_letter(previous) and _is_ascii_letter(following_nonspace)
        ) or (_is_ascii_letter(following) and _is_ascii_letter(previous_nonspace))
        if not attached_word_context:
            continue
        characters[index] = "-" if character == _MISENCODED_PDF_WORD_HYPHEN else "'"
    return "".join(characters)


def _normalize_validation_text(text: str) -> str:
    normalized = unicodedata.normalize("NFKC", _repair_known_pdf_text_encoding(text))
    return "".join(
        character
        for character in normalized
        if not character.isspace() and unicodedata.category(character) != "Cf"
    )


def _text_sequence_coverage(source_text: str, output_text: str) -> float:
    source = _normalize_validation_text(source_text)
    output = _normalize_validation_text(output_text)
    if not source:
        return 1.0
    if not output:
        return 0.0
    width = 3 if len(source) >= 12 else 1
    source_sequences = {
        source[index : index + width]
        for index in range(max(1, len(source) - width + 1))
    }
    output_sequences = {
        output[index : index + width]
        for index in range(max(1, len(output) - width + 1))
    }
    return len(source_sequences & output_sequences) / len(source_sequences)


def _normalize_english_line_break_hyphens(text: str) -> str:
    normalized = unicodedata.normalize(
        "NFKC", _repair_known_pdf_text_encoding(text)
    ).casefold()
    source_words = set(_ENGLISH_WORD_PATTERN.findall(normalized))

    def replace(match: re.Match[str]) -> str:
        candidate = f"{match.group('left')}{match.group('right')}"
        return candidate if candidate in source_words else match.group(0)

    return _LINE_END_HYPHENATED_ENGLISH_PATTERN.sub(replace, normalized)


def _english_words(text: str) -> list[str]:
    normalized = _normalize_english_line_break_hyphens(text)
    return _ENGLISH_WORD_PATTERN.findall(normalized)


def _english_word_multiset_recall(source_text: str, output_text: str) -> float:
    source_words = Counter(_english_words(source_text))
    if not source_words:
        return 1.0
    output_words = Counter(_english_words(output_text))
    matched = sum((source_words & output_words).values())
    return matched / sum(source_words.values())


def _adjacent_english_word_coverage(source_text: str, output_text: str) -> float:
    source_words = _english_words(source_text)
    if len(source_words) < 2:
        return 1.0
    output_words = _english_words(output_text)
    source_pairs = Counter(zip(source_words, source_words[1:]))
    output_pairs = Counter(zip(output_words, output_words[1:]))
    matched = sum((source_pairs & output_pairs).values())
    return matched / sum(source_pairs.values())


def _block_local_adjacent_english_word_coverage(
    source_blocks: Iterable[str], output_text: str
) -> float:
    source_pairs: Counter[tuple[str, str]] = Counter()
    for block in source_blocks:
        words = _english_words(block)
        source_pairs.update(zip(words, words[1:]))
    if not source_pairs:
        return 1.0
    output_words = _english_words(output_text)
    output_pairs = Counter(zip(output_words, output_words[1:]))
    matched = sum((source_pairs & output_pairs).values())
    return matched / sum(source_pairs.values())


def _compact_english_sequence_coverage(source_text: str, output_text: str) -> float:
    source = "".join(_english_words(source_text))
    output = "".join(_english_words(output_text))
    return _text_sequence_coverage(source, output)


def _pdf_english_boundary_recovery_pass(
    source_text: str,
    output_text: str,
    *,
    character_coverage: float | None = None,
) -> bool:
    """Accept proven space recovery when the PDF text layer lost word boundaries.

    The compact letter sequence must remain virtually unchanged, the source must
    contain abnormally long joined tokens, and the output must expose materially
    more words.  This never treats arbitrary rewording or missing letters as a
    successful conversion.
    """

    coverage = (
        _text_sequence_coverage(source_text, output_text)
        if character_coverage is None
        else float(character_coverage)
    )
    if (
        coverage < 0.92
        or _compact_english_sequence_coverage(source_text, output_text) < 0.98
    ):
        return False
    source_words = _english_words(source_text)
    output_words = _english_words(output_text)
    if not source_words or len(output_words) < len(source_words) + 2:
        return False
    joined_source_words = sum(len(word) >= 18 for word in source_words)
    return joined_source_words >= 1 and len(output_words) >= math.ceil(
        len(source_words) * 1.20
    )


def _pdf2docx_span_text(span: Any) -> str:
    return str(getattr(span, "text", "") or "")


def _is_pdf2docx_text_span(span: Any) -> bool:
    return all(hasattr(span, name) for name in ("chars", "size", "bbox", "text"))


def _pdf2docx_bbox_value(element: Any, name: str, index: int) -> float:
    bbox = element.bbox
    value = getattr(bbox, name, None)
    return float(value if value is not None else bbox[index])


def _pdf2docx_space_gap_threshold(*spans: Any) -> float:
    sizes = [float(getattr(span, "size", 0.0) or 0.0) for span in spans]
    positive_sizes = [size for size in sizes if size > 0]
    font_size = min(positive_sizes) if positive_sizes else 10.0
    return max(0.8, font_size * 0.18)


def _pdf2docx_visible_word_gap_threshold(*spans: Any) -> float:
    """Minimum visible outline gap that reliably represents a word boundary.

    Some teaching-material PDFs replace every embedded glyph advance with a full
    em and then position each character manually.  In those files the character
    bboxes overlap heavily, so ``next.x0 - previous.x1`` is always negative even
    when a clearly visible word space exists.  The actual glyph outlines remain
    reliable; a 22% em visible gap cleanly separates ordinary side bearings from
    a word space while remaining conservative for tracked text.
    """

    sizes = [float(getattr(span, "size", 0.0) or 0.0) for span in spans]
    positive_sizes = [size for size in sizes if size > 0]
    font_size = min(positive_sizes) if positive_sizes else 10.0
    return max(1.0, font_size * 0.22)


def _pdf2docx_font_trace_key(font_name: str) -> str:
    without_subset = re.sub(r"^[A-Za-z]{6}\+", "", str(font_name or "").strip())
    return re.sub(r"[^0-9A-Za-z]+", "", without_subset).casefold()


def _pdf2docx_geometry_character(text: str) -> str:
    return next(
        (character for character in str(text or "") if not character.isspace()), ""
    )


def _pdf2docx_glyph_geometry_key(
    font_name: str, size: float, character: Any
) -> tuple[Any, ...] | None:
    text = _pdf2docx_geometry_character(str(getattr(character, "c", "") or ""))
    origin = getattr(character, "origin", None)
    if not text or not origin:
        return None
    try:
        return (
            _pdf2docx_font_trace_key(font_name),
            round(float(size), 3),
            round(float(origin[0]), 2),
            round(float(origin[1]), 2),
            ord(text),
        )
    except (IndexError, TypeError, ValueError):
        return None


class _Pdf2DocxFontOutline:
    """Lazy embedded-font outline reader used only while restoring spaces."""

    def __init__(self, font: Any) -> None:
        self.font = font
        self.units_per_em = float(font["head"].unitsPerEm)
        self.glyph_order = list(font.getGlyphOrder())
        self._bounds: dict[int, tuple[float, float] | None] = {}

    def bounds(self, glyph_id: int) -> tuple[float, float] | None:
        if glyph_id in self._bounds:
            return self._bounds[glyph_id]
        if glyph_id < 0 or glyph_id >= len(self.glyph_order) or self.units_per_em <= 0:
            self._bounds[glyph_id] = None
            return None

        glyph_name = self.glyph_order[glyph_id]
        result: tuple[float, float] | None = None
        try:
            if "glyf" in self.font:
                glyph = self.font["glyf"][glyph_name]
                result = (float(glyph.xMin), float(glyph.xMax))
            else:
                from fontTools.pens.boundsPen import BoundsPen

                glyph_set = self.font.getGlyphSet()
                pen = BoundsPen(glyph_set)
                glyph_set[glyph_name].draw(pen)
                if pen.bounds is not None:
                    result = (float(pen.bounds[0]), float(pen.bounds[2]))
        except (AttributeError, KeyError, TypeError, ValueError):
            result = None
        self._bounds[glyph_id] = result
        return result

    def close(self) -> None:
        with contextlib.suppress(Exception):
            self.font.close()


class _Pdf2DocxGlyphGeometryResolver:
    """Map pdf2docx characters back to their visible embedded-glyph outlines."""

    def __init__(self, document: Any) -> None:
        self.document = document
        self._font_outlines: dict[int, _Pdf2DocxFontOutline | None] = {}
        self._page_geometry: dict[int, dict[tuple[Any, ...], tuple[float, float]]] = {}

    def _font_outline(self, xref: int) -> _Pdf2DocxFontOutline | None:
        if xref in self._font_outlines:
            return self._font_outlines[xref]
        outline: _Pdf2DocxFontOutline | None = None
        try:
            from fontTools.ttLib import TTFont

            _name, _extension, _font_type, data = self.document.extract_font(xref)
            if data:
                outline = _Pdf2DocxFontOutline(TTFont(io.BytesIO(data), lazy=True))
        except Exception:
            outline = None
        self._font_outlines[xref] = outline
        return outline

    def page_geometry(
        self, page_index: int
    ) -> Mapping[tuple[Any, ...], tuple[float, float]]:
        if page_index in self._page_geometry:
            return self._page_geometry[page_index]

        geometry: dict[tuple[Any, ...], tuple[float, float]] = {}
        self._page_geometry[page_index] = geometry
        try:
            source_page = self.document[page_index]
            font_xrefs: dict[str, int] = {}
            for font_record in source_page.get_fonts(full=True):
                if not font_record:
                    continue
                xref = int(font_record[0])
                for name_index in (3, 4):
                    if len(font_record) <= name_index:
                        continue
                    key = _pdf2docx_font_trace_key(str(font_record[name_index] or ""))
                    if key:
                        font_xrefs.setdefault(key, xref)

            for trace_span in source_page.get_texttrace():
                direction = trace_span.get("dir", (1.0, 0.0))
                try:
                    if abs(float(direction[0])) < abs(float(direction[1])):
                        continue
                except (IndexError, TypeError, ValueError):
                    continue
                font_name = str(trace_span.get("font", "") or "")
                font_key = _pdf2docx_font_trace_key(font_name)
                xref = font_xrefs.get(font_key)
                if xref is None:
                    continue
                outline = self._font_outline(xref)
                if outline is None:
                    continue
                size = float(trace_span.get("size", 0.0) or 0.0)
                if size <= 0:
                    continue
                for raw_character in trace_span.get("chars", ()):
                    try:
                        codepoint, glyph_id, origin, _bbox = raw_character
                        glyph_bounds = outline.bounds(int(glyph_id))
                        if glyph_bounds is None:
                            continue
                        origin_x = float(origin[0])
                        key = (
                            font_key,
                            round(size, 3),
                            round(origin_x, 2),
                            round(float(origin[1]), 2),
                            int(codepoint),
                        )
                        scale = size / outline.units_per_em
                        geometry[key] = (
                            origin_x + glyph_bounds[0] * scale,
                            origin_x + glyph_bounds[1] * scale,
                        )
                        geometry.setdefault(
                            (
                                "",
                                round(size, 3),
                                round(origin_x, 2),
                                round(float(origin[1]), 2),
                                None,
                            ),
                            geometry[key],
                        )
                    except (IndexError, TypeError, ValueError):
                        continue
        except Exception:
            geometry.clear()
        return geometry

    def close(self) -> None:
        for outline in self._font_outlines.values():
            if outline is not None:
                outline.close()


def _pdf2docx_character_visible_box(
    span: Any,
    character: Any,
    glyph_geometry: Mapping[tuple[Any, ...], tuple[float, float]] | None,
) -> tuple[float, float] | None:
    if not glyph_geometry:
        return None
    key = _pdf2docx_glyph_geometry_key(
        str(getattr(span, "font", "") or ""),
        float(getattr(span, "size", 0.0) or 0.0),
        character,
    )
    if key is None:
        return None
    exact = glyph_geometry.get(key)
    if exact is not None:
        return exact
    return glyph_geometry.get(("", key[1], key[2], key[3], None))


_PDF2DOCX_NO_SPACE_BEFORE = frozenset(
    ",.;:!?)]}%\u3001\u3002\uff0c\uff0e\uff1b\uff1a\uff01\uff1f\u2019\u201d"
)
_PDF2DOCX_NO_SPACE_AFTER = frozenset("([{\u2018\u201c")


def _pdf2docx_coordinate_space_candidate(
    previous_text: str,
    current_text: str,
    *,
    previous_context: str = "",
) -> bool:
    previous = str(previous_text or "").rstrip()
    current = str(current_text or "").lstrip()
    if not previous or not current:
        return False
    left = unicodedata.normalize("NFKC", previous[-1])
    right = unicodedata.normalize("NFKC", current[0])
    left = left[-1] if left else previous[-1]
    right = right[0] if right else current[0]
    if left in ".,:/-" and right.isdigit():
        numeric_source = previous[-2] if len(previous) >= 2 else previous_context[-1:]
        numeric_left = unicodedata.normalize("NFKC", numeric_source)
        if numeric_left and numeric_left[-1].isdigit():
            return False
    if right in _PDF2DOCX_NO_SPACE_BEFORE or left in _PDF2DOCX_NO_SPACE_AFTER:
        return False
    left_word = left.isascii() and left.isalnum()
    right_word = right.isascii() and right.isalnum()
    if left_word and right_word:
        return True
    if right_word and (unicodedata.category(left).startswith("P") or left in "）】》"):
        return True
    if left_word and right in "\"'\u2018\u201c":
        return True
    return False


def _pdf2docx_visible_character_gap(
    previous_span: Any,
    previous_character: Any,
    current_span: Any,
    current_character: Any,
    glyph_geometry: Mapping[tuple[Any, ...], tuple[float, float]] | None,
) -> float | None:
    previous_box = _pdf2docx_character_visible_box(
        previous_span, previous_character, glyph_geometry
    )
    current_box = _pdf2docx_character_visible_box(
        current_span, current_character, glyph_geometry
    )
    if previous_box is None or current_box is None:
        return None
    return current_box[0] - previous_box[1]


def _set_pdf2docx_span_text(span: Any, text: str) -> None:
    # After parsing, character boxes are no longer needed by make_docx(). Clearing
    # them makes TextSpan.text use the corrected text while retaining its bbox/style.
    span.chars = []
    span.text = text


def _append_pdf2docx_span_space(span: Any) -> int:
    text = _pdf2docx_span_text(span)
    if not text or text[-1].isspace():
        return 0
    _set_pdf2docx_span_text(span, f"{text} ")
    return 1


def _join_pdf2docx_line_break_hyphen(
    previous: Any,
    current: Any,
    source_words: set[str] | frozenset[str],
) -> bool:
    """Join a line-broken English word only when the source corpus proves it."""

    previous_text = _pdf2docx_span_text(previous)
    current_text = _pdf2docx_span_text(current)
    left_match = re.search(r"([A-Za-z]+)-$", previous_text)
    right_match = re.match(r"([a-z]+)", current_text)
    if left_match is None or right_match is None:
        return False
    candidate = f"{left_match.group(1)}{right_match.group(1)}".casefold()
    if candidate not in source_words:
        return False
    _set_pdf2docx_span_text(span=previous, text=previous_text[:-1])
    return True


def _restore_pdf2docx_span_spaces(
    span: Any,
    glyph_geometry: Mapping[tuple[Any, ...], tuple[float, float]] | None = None,
) -> int:
    """Restore coordinate-implied spaces missing inside a parsed text span."""

    chars = list(getattr(span, "chars", ()) or ())
    if len(chars) < 2:
        return 0

    restored = [str(chars[0].c or "")]
    inserted = 0
    threshold = _pdf2docx_space_gap_threshold(span)
    for boundary_index, (previous, current) in enumerate(zip(chars, chars[1:])):
        previous_text = str(previous.c or "")
        current_text = str(current.c or "")
        visible_gap = _pdf2docx_visible_character_gap(
            span,
            previous,
            span,
            current,
            glyph_geometry,
        )
        boundary_candidate = _pdf2docx_coordinate_space_candidate(
            previous_text,
            current_text,
            previous_context=(
                str(chars[boundary_index - 1].c or "") if boundary_index > 0 else ""
            ),
        )
        if visible_gap is None:
            gap = _pdf2docx_bbox_value(current, "x0", 0) - _pdf2docx_bbox_value(
                previous, "x1", 2
            )
            should_insert = boundary_candidate and gap > threshold
        else:
            should_insert = (
                boundary_candidate
                and visible_gap > _pdf2docx_visible_word_gap_threshold(span)
            )
        if should_insert:
            restored.append(" ")
            inserted += 1
        restored.append(current_text)

    if inserted:
        _set_pdf2docx_span_text(span, "".join(restored))
    return inserted


def _pdf2docx_line_is_horizontal(line: Any) -> bool:
    direction = getattr(line, "dir", (1.0, 0.0))
    try:
        return abs(float(direction[0])) >= abs(float(direction[1]))
    except (IndexError, TypeError, ValueError):
        return True


def _first_pdf2docx_text_span(line: Any) -> Any | None:
    return next(
        (
            span
            for span in line.spans
            if _is_pdf2docx_text_span(span) and _pdf2docx_span_text(span)
        ),
        None,
    )


def _last_pdf2docx_text_span(line: Any) -> Any | None:
    return next(
        (
            span
            for span in reversed(list(line.spans))
            if _is_pdf2docx_text_span(span) and _pdf2docx_span_text(span)
        ),
        None,
    )


def _restore_pdf2docx_text_block_spaces(
    block: Any,
    source_words: set[str] | frozenset[str] = frozenset(),
    glyph_geometry: Mapping[tuple[Any, ...], tuple[float, float]] | None = None,
) -> int:
    """Restore spaces between parsed spans and English words joined across lines."""

    lines = list(block.lines)
    inserted = 0
    for line in lines:
        if not _pdf2docx_line_is_horizontal(line):
            continue
        spans = list(line.spans)
        boundary_characters: dict[int, tuple[Any | None, Any | None]] = {}
        for span in spans:
            if not _is_pdf2docx_text_span(span):
                continue
            characters = [
                character
                for character in list(getattr(span, "chars", ()) or ())
                if str(getattr(character, "c", "") or "")
            ]
            boundary_characters[id(span)] = (
                characters[0] if characters else None,
                characters[-1] if characters else None,
            )
        for span in spans:
            if _is_pdf2docx_text_span(span):
                inserted += _restore_pdf2docx_span_spaces(
                    span, glyph_geometry=glyph_geometry
                )

        for previous, current in zip(spans, spans[1:]):
            if not (
                _is_pdf2docx_text_span(previous) and _is_pdf2docx_text_span(current)
            ):
                continue
            previous_text = _pdf2docx_span_text(previous)
            current_text = _pdf2docx_span_text(current)
            if (
                not previous_text
                or not current_text
                or previous_text[-1].isspace()
                or current_text[0].isspace()
            ):
                continue
            _previous_first, previous_character = boundary_characters.get(
                id(previous), (None, None)
            )
            current_character, _current_last = boundary_characters.get(
                id(current), (None, None)
            )
            visible_gap = None
            if previous_character is not None and current_character is not None:
                visible_gap = _pdf2docx_visible_character_gap(
                    previous,
                    previous_character,
                    current,
                    current_character,
                    glyph_geometry,
                )
            if visible_gap is None:
                gap = _pdf2docx_bbox_value(current, "x0", 0) - _pdf2docx_bbox_value(
                    previous, "x1", 2
                )
                should_insert = _pdf2docx_coordinate_space_candidate(
                    previous_text, current_text
                ) and gap > _pdf2docx_space_gap_threshold(previous, current)
            else:
                should_insert = _pdf2docx_coordinate_space_candidate(
                    previous_text, current_text
                ) and visible_gap > _pdf2docx_visible_word_gap_threshold(
                    previous, current
                )
            if should_insert:
                inserted += _append_pdf2docx_span_space(previous)

    no_join_characters = set("-‐‑‒–—/")
    for previous_line, current_line in zip(lines, lines[1:]):
        if not (
            _pdf2docx_line_is_horizontal(previous_line)
            and _pdf2docx_line_is_horizontal(current_line)
        ):
            continue
        if getattr(previous_line, "line_break", 0):
            continue
        previous = _last_pdf2docx_text_span(previous_line)
        current = _first_pdf2docx_text_span(current_line)
        if previous is None or current is None:
            continue
        previous_text = _pdf2docx_span_text(previous)
        current_text = _pdf2docx_span_text(current)
        if not previous_text or not current_text:
            continue
        previous_character = previous_text[-1]
        current_character = current_text[0]
        if _join_pdf2docx_line_break_hyphen(previous, current, source_words):
            inserted += 1
            continue
        portable_previous = unicodedata.normalize("NFKC", previous_character)
        portable_current = unicodedata.normalize("NFKC", current_character)
        portable_previous = (
            portable_previous[-1] if portable_previous else previous_character
        )
        portable_current = (
            portable_current[0] if portable_current else current_character
        )
        previous_context = previous_text[:-1].rstrip()
        if not previous_context:
            text_spans = [
                span
                for span in previous_line.spans
                if _is_pdf2docx_text_span(span) and _pdf2docx_span_text(span)
            ]
            with contextlib.suppress(ValueError):
                previous_index = text_spans.index(previous)
                for preceding_span in reversed(text_spans[:previous_index]):
                    preceding_text = _pdf2docx_span_text(preceding_span).rstrip()
                    if preceding_text:
                        previous_context = preceding_text
                        break
        if (
            previous_character.isspace()
            or current_character.isspace()
            or portable_previous in no_join_characters
        ):
            continue
        if _pdf2docx_coordinate_space_candidate(
            previous_text,
            current_text,
            previous_context=previous_context,
        ):
            inserted += _append_pdf2docx_span_space(previous)

    return inserted


def _iter_pdf2docx_text_blocks(blocks: Iterable[Any]) -> Iterable[Any]:
    for block in blocks:
        if hasattr(block, "lines"):
            yield block
            continue
        if not hasattr(block, "__iter__"):
            continue
        for row in block:
            if not hasattr(row, "__iter__"):
                continue
            for cell in row:
                if cell and hasattr(cell, "blocks"):
                    yield from _iter_pdf2docx_text_blocks(cell.blocks)


def _restore_pdf2docx_spaces(
    pages: Iterable[Any],
    source_words: set[str] | frozenset[str] = frozenset(),
    source_document: Any | None = None,
) -> int:
    inserted = 0
    geometry_resolver = (
        _Pdf2DocxGlyphGeometryResolver(source_document)
        if source_document is not None
        else None
    )
    try:
        for page in pages:
            if not getattr(page, "finalized", False):
                continue
            page_index = int(getattr(page, "id", 0) or 0)
            glyph_geometry = (
                geometry_resolver.page_geometry(page_index)
                if geometry_resolver is not None
                else None
            )
            for section in page.sections:
                for column in section:
                    for block in _iter_pdf2docx_text_blocks(column.blocks):
                        inserted += _restore_pdf2docx_text_block_spaces(
                            block,
                            source_words=source_words,
                            glyph_geometry=glyph_geometry,
                        )
    finally:
        if geometry_resolver is not None:
            geometry_resolver.close()
    return inserted


def _repair_pdf2docx_known_encoding(pages: Iterable[Any]) -> int:
    """Apply proven one-for-one encoding repairs to parsed editable spans."""

    repaired = 0
    for page in pages:
        if not getattr(page, "finalized", False):
            continue
        for section in page.sections:
            for column in section:
                spans = [
                    span
                    for block in _iter_pdf2docx_text_blocks(column.blocks)
                    for line in block.lines
                    for span in line.spans
                    if _is_pdf2docx_text_span(span)
                ]
                original_parts = [_pdf2docx_span_text(span) for span in spans]
                original_text = "".join(original_parts)
                corrected_text = _repair_known_pdf_text_encoding(original_text)
                if corrected_text == original_text:
                    continue
                offset = 0
                for span, original_part in zip(spans, original_parts, strict=True):
                    corrected_part = corrected_text[
                        offset : offset + len(original_part)
                    ]
                    offset += len(original_part)
                    if corrected_part == original_part:
                        continue
                    repaired += sum(
                        left != right
                        for left, right in zip(
                            original_part, corrected_part, strict=True
                        )
                    )
                    _set_pdf2docx_span_text(span, corrected_part)
    return repaired


def _inspect_pdf_text_layers(
    source: Path, password: str | None
) -> tuple[list[str], list[int]]:
    _pymupdf, pdf_document = _open_pymupdf_document(source, password)
    try:
        page_texts: list[str] = []
        pages_without_text: list[int] = []
        for page_number, page in enumerate(pdf_document, start=1):
            text = page.get_text("text", sort=True) or ""
            page_texts.append(text)
            if not _normalize_validation_text(text):
                pages_without_text.append(page_number)
        return page_texts, pages_without_text
    finally:
        pdf_document.close()


def _clipped_bbox_area_ratio(bbox: Any, page_rect: Any) -> float:
    try:
        x0, y0, x1, y1 = (float(value) for value in bbox)
        page_x0 = float(page_rect.x0)
        page_y0 = float(page_rect.y0)
        page_x1 = float(page_rect.x1)
        page_y1 = float(page_rect.y1)
    except (AttributeError, TypeError, ValueError):
        return 0.0
    width = max(0.0, min(x1, page_x1) - max(x0, page_x0))
    height = max(0.0, min(y1, page_y1) - max(y0, page_y0))
    page_area = max(1.0, (page_x1 - page_x0) * (page_y1 - page_y0))
    return width * height / page_area


def _pdf_page_image_area_ratios(
    page_dict: Mapping[str, Any], page_rect: Any
) -> tuple[float, float]:
    ratios: list[float] = []
    seen: set[tuple[float, float, float, float]] = set()
    for block in page_dict.get("blocks", ()):
        if int(block.get("type", 0) or 0) != 1:
            continue
        bbox = block.get("bbox", ())
        try:
            key = tuple(round(float(value), 2) for value in bbox)
        except (TypeError, ValueError):
            continue
        if len(key) != 4 or key in seen:
            continue
        seen.add(key)
        ratio = _clipped_bbox_area_ratio(key, page_rect)
        if ratio > 0:
            ratios.append(ratio)
    return (max(ratios, default=0.0), min(1.0, sum(ratios)))


def _has_suspicious_pdf_characters(text: str) -> bool:
    for character in text:
        category = unicodedata.category(character)
        if (
            character == "\ufffd"
            or category in {"Co", "Cs"}
            or (category == "Cc" and character not in "\n\r\t")
        ):
            return True
    return False


def _coerce_pdf_rect(
    bbox: Any,
    page_rect: Any,
    *,
    padding: float = 0.0,
) -> tuple[float, float, float, float] | None:
    try:
        x0, y0, x1, y1 = (float(value) for value in bbox)
        page_x0 = float(page_rect.x0)
        page_y0 = float(page_rect.y0)
        page_x1 = float(page_rect.x1)
        page_y1 = float(page_rect.y1)
    except (AttributeError, TypeError, ValueError):
        return None
    if x1 < x0:
        x0, x1 = x1, x0
    if y1 < y0:
        y0, y1 = y1, y0
    if x1 - x0 < 0.5:
        center = (x0 + x1) / 2.0
        x0, x1 = center - 0.25, center + 0.25
    if y1 - y0 < 0.5:
        center = (y0 + y1) / 2.0
        y0, y1 = center - 0.25, center + 0.25
    x0 = max(page_x0, x0 - padding)
    y0 = max(page_y0, y0 - padding)
    x1 = min(page_x1, x1 + padding)
    y1 = min(page_y1, y1 + padding)
    if x1 <= x0 or y1 <= y0:
        return None
    return (x0, y0, x1, y1)


def _rect_area(rect: tuple[float, float, float, float]) -> float:
    return max(0.0, rect[2] - rect[0]) * max(0.0, rect[3] - rect[1])


def _rect_union(
    left: tuple[float, float, float, float],
    right: tuple[float, float, float, float],
) -> tuple[float, float, float, float]:
    return (
        min(left[0], right[0]),
        min(left[1], right[1]),
        max(left[2], right[2]),
        max(left[3], right[3]),
    )


def _rect_distance(
    left: tuple[float, float, float, float],
    right: tuple[float, float, float, float],
) -> float:
    horizontal = max(left[0] - right[2], right[0] - left[2], 0.0)
    vertical = max(left[1] - right[3], right[1] - left[3], 0.0)
    return math.hypot(horizontal, vertical)


def _rect_intersection_ratio(
    subject: tuple[float, float, float, float],
    region: tuple[float, float, float, float],
) -> float:
    width = max(0.0, min(subject[2], region[2]) - max(subject[0], region[0]))
    height = max(0.0, min(subject[3], region[3]) - max(subject[1], region[1]))
    return width * height / max(1.0, _rect_area(subject))


def _rect_center_inside(
    subject: tuple[float, float, float, float],
    region: tuple[float, float, float, float],
) -> bool:
    center_x = (subject[0] + subject[2]) / 2.0
    center_y = (subject[1] + subject[3]) / 2.0
    return region[0] <= center_x <= region[2] and region[1] <= center_y <= region[3]


def _cluster_pdf_rects(
    rects: Iterable[tuple[float, float, float, float]],
    *,
    gap: float,
) -> list[list[tuple[float, float, float, float]]]:
    pending = list(rects)
    clusters: list[list[tuple[float, float, float, float]]] = []
    while pending:
        cluster = [pending.pop(0)]
        union = cluster[0]
        changed = True
        while changed:
            changed = False
            remaining: list[tuple[float, float, float, float]] = []
            for rect in pending:
                if _rect_distance(union, rect) <= gap:
                    cluster.append(rect)
                    union = _rect_union(union, rect)
                    changed = True
                else:
                    remaining.append(rect)
            pending = remaining
        clusters.append(cluster)
    return clusters


def _merge_hybrid_regions(
    regions: Iterable[_HybridRegion],
    page_rect: Any,
    *,
    gap: float = 1.0,
) -> list[_HybridRegion]:
    pending = list(regions)
    merged: list[_HybridRegion] = []
    while pending:
        seed = pending.pop(0)
        rect = seed.rect
        kinds = [seed.kind]
        reasons = list(seed.reasons)
        dpi = seed.dpi
        changed = True
        while changed:
            changed = False
            remaining: list[_HybridRegion] = []
            for candidate in pending:
                if _rect_distance(rect, candidate.rect) <= gap:
                    rect = _rect_union(rect, candidate.rect)
                    kinds.append(candidate.kind)
                    reasons.extend(candidate.reasons)
                    dpi = max(dpi, candidate.dpi)
                    changed = True
                else:
                    remaining.append(candidate)
            pending = remaining
        normalized = _coerce_pdf_rect(rect, page_rect)
        if normalized is None:
            continue
        kind = kinds[0] if len(set(kinds)) == 1 else "complex"
        merged.append(
            _HybridRegion(
                page_index=seed.page_index,
                rect=normalized,
                kind=kind,
                reasons=tuple(dict.fromkeys(reasons)),
                dpi=dpi,
            )
        )
    return sorted(merged, key=lambda item: (item.rect[1], item.rect[0]))


def _math_font(font_name: str) -> bool:
    key = _word_font_key(font_name)
    return any(token in key for token in _MATH_FONT_TOKENS)


def _text_line_rect(line: Mapping[str, Any], page_rect: Any):
    bbox = line.get("bbox")
    if bbox:
        return _coerce_pdf_rect(bbox, page_rect)
    span_rects = [
        rect
        for span in line.get("spans", ())
        if (rect := _coerce_pdf_rect(span.get("bbox", ()), page_rect)) is not None
    ]
    if not span_rects:
        return None
    rect = span_rects[0]
    for span_rect in span_rects[1:]:
        rect = _rect_union(rect, span_rect)
    return rect


def _text_line_value(line: Mapping[str, Any]) -> str:
    return "".join(str(span.get("text", "") or "") for span in line.get("spans", ()))


def _pdf_line_has_overlapping_lowercase_fragments(
    line: Mapping[str, Any],
    page_rect: Any,
) -> bool:
    """Detect split descenders that become isolated ``p/g/y/j`` text in Word.

    Affected PDF fonts expose selected glyphs as separate one-letter spans whose
    boxes overlap the surrounding word.  Requiring at least two overlapping
    lowercase fragments avoids classifying ordinary standalone articles,
    option labels, or drop caps as broken text.
    """

    entries: list[
        tuple[
            str,
            tuple[float, float, float, float],
            bool,
            str,
            float,
        ]
    ] = []
    for span in line.get("spans", ()):
        text = str(span.get("text", "") or "")
        rect = _coerce_pdf_rect(span.get("bbox", ()), page_rect)
        if rect is None or not text.strip():
            continue
        stripped = text.strip()
        entries.append(
            (
                text,
                rect,
                len(stripped) == 1 and stripped in _OVERLAPPING_PDF_FRAGMENT_LETTERS,
                _word_font_key(str(span.get("font", "") or "")),
                float(span.get("size", 0.0) or 0.0),
            )
        )
    if sum(1 for _text, _rect, single, _font, _size in entries if single) < 2:
        return False

    overlapping_fragments = 0
    fragment_style_matches = 0
    for entry_index, (_text, rect, single, font, size) in enumerate(entries):
        if not single:
            continue
        fragment_width = max(1.0, rect[2] - rect[0])
        fragment_height = max(1.0, rect[3] - rect[1])
        fragment_overlaps = False
        fragment_style_match = False
        for other_index, (
            other_text,
            other_rect,
            _other_single,
            other_font,
            other_size,
        ) in enumerate(entries):
            if other_index == entry_index or not any(
                _is_ascii_letter(character) for character in other_text
            ):
                continue
            horizontal_overlap = max(
                0.0, min(rect[2], other_rect[2]) - max(rect[0], other_rect[0])
            )
            vertical_overlap = max(
                0.0, min(rect[3], other_rect[3]) - max(rect[1], other_rect[1])
            )
            if (
                horizontal_overlap >= max(0.75, fragment_width * 0.12)
                and vertical_overlap >= fragment_height * 0.35
            ):
                fragment_overlaps = True
                if (
                    abs(other_index - entry_index) <= 1
                    and font == other_font
                    and abs(size - other_size) <= max(0.25, size * 0.03)
                ):
                    fragment_style_match = True
        if fragment_overlaps:
            overlapping_fragments += 1
        if fragment_style_match:
            fragment_style_matches += 1
    if overlapping_fragments < 2:
        return False

    # Some embedded fonts store the descender of p/g/y/j as a separate span
    # that slightly overlaps the rest of the same word. When these spans remain
    # in left-to-right order and share one style, pdf2docx writes ordinary
    # sequential runs; rasterizing them only destroys editability.
    repaired_text = _repair_known_pdf_text_encoding(
        "".join(text for text, _rect, _single, _font, _size in entries)
    )
    substantive = [character for character in repaired_text if not character.isspace()]
    ascii_letters = sum(_is_ascii_letter(character) for character in substantive)
    raw_direction = tuple(line.get("dir", (1.0, 0.0)))
    direction = tuple(float(value) for value in raw_direction[:2])
    horizontally_ordered = all(
        current[1][0] + 0.75 >= previous[1][0]
        for previous, current in zip(entries, entries[1:])
    )
    recoverable_english_descenders = (
        len(substantive) >= 3
        and ascii_letters / max(1, len(substantive)) >= 0.60
        and fragment_style_matches == overlapping_fragments
        and horizontally_ordered
        and len(direction) == 2
        and direction[0] >= 0.95
        and abs(direction[1]) <= 0.05
        and not _DISPLAY_FORMULA_OPERATOR_PATTERN.search(repaired_text)
        and not _has_suspicious_pdf_characters(repaired_text)
    )
    return not recoverable_english_descenders


def _detect_pdf_encoding_regions(
    page_index: int,
    page_dict: Mapping[str, Any],
    page_rect: Any,
    *,
    dpi: int,
) -> list[_HybridRegion]:
    """Rasterize only lines whose font encoding cannot be repaired reliably."""

    regions: list[_HybridRegion] = []
    visual_dpi = min(600, max(int(dpi), 450))
    page_y0 = float(getattr(page_rect, "y0", 0.0) or 0.0)
    page_height = max(1.0, float(getattr(page_rect, "height", 0.0) or 0.0))
    for block in page_dict.get("blocks", ()):
        if int(block.get("type", 0) or 0) != 0:
            continue
        for line in block.get("lines", ()):
            line_rect = _text_line_rect(line, page_rect)
            if line_rect is None:
                continue
            # Running furniture is handled by the dedicated header/footer pass;
            # turning it into a local image would interfere with page floating.
            if (
                line_rect[3] <= page_y0 + page_height * 0.06
                or line_rect[1] >= page_y0 + page_height * 0.97
            ):
                continue
            raw_text = _text_line_value(line)
            repaired_text = _repair_known_pdf_text_encoding(raw_text)
            reasons: list[str] = []
            if _has_suspicious_pdf_characters(repaired_text):
                reasons.append(_UNRELIABLE_PDF_ENCODING_REASON)
            if _pdf_line_has_overlapping_lowercase_fragments(line, page_rect):
                reasons.append(_OVERLAPPING_PDF_GLYPH_REASON)
            if not reasons:
                continue
            normalized = _coerce_pdf_rect(line_rect, page_rect, padding=1.5)
            if normalized is None:
                continue
            regions.append(
                _HybridRegion(
                    page_index=page_index,
                    rect=normalized,
                    kind="complex",
                    reasons=tuple(reasons),
                    dpi=visual_dpi,
                )
            )

    # The destructive mask for the remaining genuinely unsafe overlap cases is
    # contracted away from neighboring tall glyphs later. Do not recursively
    # expand through adjacent baselines here: one bad line could otherwise turn
    # a whole paragraph into a non-editable image.
    return _merge_hybrid_regions(regions, page_rect, gap=1.0)


def _hybrid_text_line_is_exempt(
    line_rect: tuple[float, float, float, float],
    text: str,
    page_rect: Any,
) -> bool:
    """Keep running heads, footers, and figure/table captions editable."""

    page_y0 = float(getattr(page_rect, "y0", 0.0) or 0.0)
    page_height = max(1.0, float(getattr(page_rect, "height", 0.0) or 0.0))
    if line_rect[3] <= page_y0 + page_height * 0.06:
        return True
    if line_rect[1] >= page_y0 + page_height * 0.97:
        return True
    normalized_text = " ".join(text.split())
    return bool(
        normalized_text
        and (
            _CAPTION_PATTERN.match(normalized_text)
            or _TABLE_TITLE_PATTERN.match(normalized_text)
            or re.match(
                r"^fig(?:ure)?\.?\s*[ivxlcdm]+(?:\s|[:.])",
                normalized_text,
                re.IGNORECASE,
            )
        )
    )


def _rect_fully_contains(
    container: tuple[float, float, float, float],
    subject: tuple[float, float, float, float],
    *,
    tolerance: float = 0.01,
) -> bool:
    return (
        container[0] <= subject[0] + tolerance
        and container[1] <= subject[1] + tolerance
        and container[2] >= subject[2] - tolerance
        and container[3] >= subject[3] - tolerance
    )


def _hybrid_text_rect_is_excluded(
    text_rect: tuple[float, float, float, float],
    region_rect: tuple[float, float, float, float],
) -> bool:
    return _rect_center_inside(text_rect, region_rect) or (
        _rect_intersection_ratio(text_rect, region_rect) >= 0.35
    )


def _close_hybrid_regions_over_text(
    regions: Iterable[_HybridRegion],
    page_dict: Mapping[str, Any],
    page_rect: Any,
) -> list[_HybridRegion]:
    """Close final regions over every non-caption text rect they suppress.

    Expanding one rectangle can make another line or span newly cross the same
    exclusion threshold.  Iterate through merge-and-expand cycles until the
    geometry reaches a fixed point, so rasterization never clips a glyph that
    has already been removed from the editable layer.
    """

    candidates: list[tuple[float, float, float, float]] = []
    for block in page_dict.get("blocks", ()):
        if int(block.get("type", 0) or 0) != 0:
            continue
        for line in block.get("lines", ()):
            line_rect = _text_line_rect(line, page_rect)
            if line_rect is None:
                continue
            if _hybrid_text_line_is_exempt(
                line_rect,
                _text_line_value(line),
                page_rect,
            ):
                continue
            candidates.append(line_rect)
            candidates.extend(
                span_rect
                for span in line.get("spans", ())
                if (span_rect := _coerce_pdf_rect(span.get("bbox", ()), page_rect))
                is not None
            )

    closed = _merge_hybrid_regions(regions, page_rect, gap=1.0)
    while closed:
        expanded: list[_HybridRegion] = []
        for region in closed:
            rect = region.rect
            changed = True
            while changed:
                changed = False
                for text_rect in candidates:
                    if not _hybrid_text_rect_is_excluded(text_rect, rect):
                        continue
                    if _rect_fully_contains(rect, text_rect):
                        continue
                    rect = _rect_union(rect, text_rect)
                    changed = True
            normalized = _coerce_pdf_rect(rect, page_rect)
            if normalized is None:
                continue
            expanded.append(
                _HybridRegion(
                    page_index=region.page_index,
                    rect=normalized,
                    kind=region.kind,
                    reasons=region.reasons,
                    dpi=region.dpi,
                )
            )

        next_closed = _merge_hybrid_regions(expanded, page_rect, gap=1.0)
        current_signature = [
            (region.rect, region.kind, region.reasons, region.dpi) for region in closed
        ]
        next_signature = [
            (region.rect, region.kind, region.reasons, region.dpi)
            for region in next_closed
        ]
        if next_signature == current_signature:
            return next_closed
        closed = next_closed
    return []


def _constrain_numbered_formula_regions_to_body_prose(
    regions: Iterable[_HybridRegion],
    page_dict: Mapping[str, Any],
    page_rect: Any,
) -> list[_HybridRegion]:
    """Trim malformed formula bboxes where they intrude into adjacent prose."""

    body_prose_rects: list[tuple[float, float, float, float]] = []
    for block in page_dict.get("blocks", ()):
        if int(block.get("type", 0) or 0) != 0:
            continue
        for line in block.get("lines", ()):
            line_rect = _text_line_rect(line, page_rect)
            if line_rect is None:
                continue
            text = _text_line_value(line)
            if _hybrid_text_line_is_exempt(line_rect, text, page_rect):
                continue
            repaired = _repair_known_pdf_text_encoding(text)
            normalized = _normalize_validation_text(repaired)
            prose_words = [
                word
                for word in re.findall(r"[A-Za-z]{3,}", repaired)
                if word.casefold() not in _FORMULA_FUNCTION_WORDS
            ]
            prose_letters = sum(len(word) for word in prose_words)
            cjk_characters = sum(
                "\u3400" <= character <= "\u9fff" for character in repaired
            )
            letter_ratio = (
                sum(character.isalpha() for character in normalized)
                / max(1, len(normalized))
            )
            # A body sentence can legitimately contain one or more italic/math
            # runs (for example ``where D_loc denotes ...``).  Rejecting the
            # complete line merely because one span uses a math font prevented
            # the old guard from clipping formula rectangles before the prose
            # below them.  Judge the ordinary-language share of the full line
            # instead of its most exotic span.
            prose_like = bool(
                (len(prose_words) >= 4 and prose_letters >= 18)
                or (
                    len(prose_words) >= 3
                    and len(normalized) >= 28
                    and letter_ratio >= 0.48
                )
                or (cjk_characters >= 8)
            )
            if not prose_like:
                continue
            body_prose_rects.append(line_rect)

    constrained: list[_HybridRegion] = []
    for region in regions:
        if region.kind != "formula" or _NUMBERED_FORMULA_REASON not in region.reasons:
            constrained.append(region)
            continue
        x0, y0, x1, y1 = region.rect
        center_y = (y0 + y1) / 2.0
        previous_bottoms: list[float] = []
        following_tops: list[float] = []
        for prose_rect in body_prose_rects:
            horizontal_overlap = max(
                0.0,
                min(x1, prose_rect[2]) - max(x0, prose_rect[0]),
            )
            if horizontal_overlap <= 0.0:
                continue
            prose_center_y = (prose_rect[1] + prose_rect[3]) / 2.0
            vertical_overlap = max(
                0.0,
                min(y1, prose_rect[3]) - max(y0, prose_rect[1]),
            )
            if vertical_overlap <= 0.0:
                continue
            if prose_center_y < center_y - 0.5:
                previous_bottoms.append(prose_rect[3])
            elif prose_center_y > center_y + 0.5:
                following_tops.append(prose_rect[1])

        constrained_y0 = max([y0, *previous_bottoms])
        constrained_y1 = min([y1, *following_tops])
        if constrained_y1 - constrained_y0 < 1.0:
            constrained.append(region)
            continue
        constrained.append(
            _HybridRegion(
                page_index=region.page_index,
                rect=(x0, constrained_y0, x1, constrained_y1),
                kind=region.kind,
                reasons=region.reasons,
                dpi=region.dpi,
            )
        )
    return constrained


def _expand_region_to_nearby_labels(
    region: _HybridRegion,
    page_dict: Mapping[str, Any],
    page_rect: Any,
) -> _HybridRegion:
    base_rect = region.rect
    if region.kind == "formula" or any(
        reason
        in {
            _UNRELIABLE_PDF_ENCODING_REASON,
            _OVERLAPPING_PDF_GLYPH_REASON,
        }
        for reason in region.reasons
    ):
        # Formula detection already includes a sub-point glyph safety margin.
        # Encoding fallbacks already contain the complete source line.  Extra
        # generic padding can pull adjacent prose or a page number into either
        # raster island, so keep both kinds tightly bounded.
        normalized = _coerce_pdf_rect(base_rect, page_rect)
        if normalized is None:
            return region
        return _HybridRegion(
            page_index=region.page_index,
            rect=normalized,
            kind=region.kind,
            reasons=region.reasons,
            dpi=region.dpi,
        )

    rect = base_rect
    nearby_gap = _HYBRID_FIGURE_LABEL_GAP
    maximum_label_height = 24.0
    for block in page_dict.get("blocks", ()):
        if int(block.get("type", 0) or 0) != 0:
            continue
        for line in block.get("lines", ()):
            line_rect = _text_line_rect(line, page_rect)
            if line_rect is None:
                continue
            text = " ".join(_text_line_value(line).split())
            if not text:
                continue
            intersection = _rect_intersection_ratio(line_rect, base_rect)
            inside = _rect_center_inside(line_rect, base_rect)
            line_height = line_rect[3] - line_rect[1]
            horizontal_gap = max(
                line_rect[0] - base_rect[2],
                base_rect[0] - line_rect[2],
                0.0,
            )
            vertical_gap = max(
                line_rect[1] - base_rect[3],
                base_rect[1] - line_rect[3],
                0.0,
            )
            horizontal_overlap = max(
                0.0,
                min(line_rect[2], base_rect[2]) - max(line_rect[0], base_rect[0]),
            )
            vertical_overlap = max(
                0.0,
                min(line_rect[3], base_rect[3]) - max(line_rect[1], base_rect[1]),
            )
            near_above_or_below = (
                vertical_gap <= nearby_gap and horizontal_overlap > 0.0
            )
            near_short_side_label = (
                horizontal_gap <= nearby_gap
                and vertical_overlap > 0.0
                and len(text) <= 32
            )
            is_near_label = (
                (near_above_or_below or near_short_side_label)
                and line_height <= maximum_label_height
                and len(text) <= 180
            )
            if (
                is_near_label
                and not inside
                and intersection <= 0.0
                and _CAPTION_PATTERN.match(text)
            ):
                continue
            if inside or intersection >= 0.10 or is_near_label:
                rect = _rect_union(rect, line_rect)
    normalized = _coerce_pdf_rect(
        rect,
        page_rect,
        padding=_HYBRID_REGION_PADDING,
    )
    if normalized is None:
        return region
    return _HybridRegion(
        page_index=region.page_index,
        rect=normalized,
        kind=region.kind,
        reasons=region.reasons,
        dpi=region.dpi,
    )


def _image_rects_share_figure(
    left: tuple[float, float, float, float],
    right: tuple[float, float, float, float],
    page_dict: Mapping[str, Any],
    page_rect: Any,
) -> bool:
    horizontal_gap = max(left[0] - right[2], right[0] - left[2], 0.0)
    vertical_gap = max(left[1] - right[3], right[1] - left[3], 0.0)
    vertical_overlap = max(0.0, min(left[3], right[3]) - max(left[1], right[1]))
    horizontal_overlap = max(0.0, min(left[2], right[2]) - max(left[0], right[0]))
    vertical_overlap_ratio = vertical_overlap / max(
        1.0,
        min(left[3] - left[1], right[3] - right[1]),
    )
    horizontal_overlap_ratio = horizontal_overlap / max(
        1.0,
        min(left[2] - left[0], right[2] - right[0]),
    )

    if horizontal_gap <= 36.0 and vertical_overlap_ratio >= 0.45:
        return True
    if vertical_gap > 20.0 or horizontal_overlap_ratio < 0.45:
        return False

    upper, lower = (left, right) if left[1] <= right[1] else (right, left)
    if lower[1] < upper[3]:
        return True
    for block in page_dict.get("blocks", ()):
        if int(block.get("type", 0) or 0) != 0:
            continue
        for line in block.get("lines", ()):
            line_rect = _text_line_rect(line, page_rect)
            if line_rect is None:
                continue
            center_y = (line_rect[1] + line_rect[3]) / 2.0
            if not (upper[3] <= center_y <= lower[1]):
                continue
            text = " ".join(_text_line_value(line).split())
            if text and _CAPTION_PATTERN.match(text):
                return False
    return True


def _cluster_pdf_image_rects(
    rects: Iterable[tuple[float, float, float, float]],
    page_dict: Mapping[str, Any],
    page_rect: Any,
) -> list[list[tuple[float, float, float, float]]]:
    pending = list(rects)
    clusters: list[list[tuple[float, float, float, float]]] = []
    while pending:
        cluster = [pending.pop(0)]
        changed = True
        while changed:
            changed = False
            remaining: list[tuple[float, float, float, float]] = []
            for rect in pending:
                if any(
                    _image_rects_share_figure(member, rect, page_dict, page_rect)
                    for member in cluster
                ):
                    cluster.append(rect)
                    changed = True
                else:
                    remaining.append(rect)
            pending = remaining
        clusters.append(cluster)
    return clusters


def _detect_pdf_image_regions(
    page_index: int,
    page_dict: Mapping[str, Any],
    page_rect: Any,
    *,
    dpi: int,
) -> list[_HybridRegion]:
    page_area = max(1.0, float(page_rect.width) * float(page_rect.height))
    image_rects: list[tuple[float, float, float, float]] = []
    for block in page_dict.get("blocks", ()):
        if int(block.get("type", 0) or 0) != 1:
            continue
        rect = _coerce_pdf_rect(block.get("bbox", ()), page_rect, padding=0.75)
        if rect is None or _rect_area(rect) / page_area < 0.0015:
            continue
        image_rects.append(rect)

    regions: list[_HybridRegion] = []
    for cluster in _cluster_pdf_image_rects(image_rects, page_dict, page_rect):
        union = cluster[0]
        for rect in cluster[1:]:
            union = _rect_union(union, rect)
        ratio = _rect_area(union) / page_area
        if len(cluster) < 2 and ratio < 0.01:
            continue
        reason = "复合图片区域" if len(cluster) > 1 else "图片与标签区域"
        regions.append(
            _HybridRegion(
                page_index=page_index,
                rect=union,
                kind="figure",
                reasons=(reason,),
                dpi=dpi,
            )
        )
    return regions


def _detect_pdf_vector_regions(
    page_index: int,
    drawings: Iterable[Mapping[str, Any]],
    page_rect: Any,
    *,
    dpi: int,
) -> list[_HybridRegion]:
    page_area = max(1.0, float(page_rect.width) * float(page_rect.height))
    entries: list[tuple[tuple[float, float, float, float], int]] = []
    for drawing in drawings:
        rect = _coerce_pdf_rect(drawing.get("rect", ()), page_rect, padding=0.75)
        if rect is None:
            continue
        items = max(1, len(drawing.get("items", ()) or ()))
        fill = drawing.get("fill")
        is_white_background = (
            items <= 1
            and isinstance(fill, (tuple, list))
            and len(fill) >= 3
            and all(float(component) >= 0.98 for component in fill[:3])
            and _rect_area(rect) / page_area >= 0.05
        )
        if is_white_background:
            continue
        if (
            rect[3] <= float(page_rect.y0) + 80.0
            and rect[3] - rect[1] <= 3.0
            and items < 40
        ):
            continue
        entries.append((rect, items))

    pending = list(entries)
    regions: list[_HybridRegion] = []
    while pending:
        rect, score = pending.pop(0)
        component_count = 1
        changed = True
        while changed:
            changed = False
            remaining: list[tuple[tuple[float, float, float, float], int]] = []
            for candidate_rect, candidate_score in pending:
                if _rect_distance(rect, candidate_rect) <= 4.0:
                    rect = _rect_union(rect, candidate_rect)
                    score += candidate_score
                    component_count += 1
                    changed = True
                else:
                    remaining.append((candidate_rect, candidate_score))
            pending = remaining
        ratio = _rect_area(rect) / page_area
        if not (
            score >= 100 or component_count >= 40 or (score >= 40 and ratio >= 0.02)
        ):
            continue
        regions.append(
            _HybridRegion(
                page_index=page_index,
                rect=rect,
                kind="vector",
                reasons=("高密度矢量图形",),
                dpi=dpi,
            )
        )
    return _merge_hybrid_regions(regions, page_rect, gap=32.0)


def _pymupdf_table_bboxes(page: Any, draw_items: int) -> list[Any]:
    if draw_items < 6 or not hasattr(page, "find_tables"):
        return []
    try:
        with warnings.catch_warnings(), contextlib.redirect_stdout(
            io.StringIO()
        ), contextlib.redirect_stderr(io.StringIO()):
            warnings.simplefilter("ignore")
            finder = page.find_tables()
        return [table.bbox for table in (getattr(finder, "tables", ()) or ())]
    except Exception:
        return []


def _pymupdf_table_count(page: Any, draw_items: int) -> int:
    return len(_pymupdf_table_bboxes(page, draw_items))


def _merge_pdf_horizontal_rule_segments(
    drawings: Iterable[Mapping[str, Any]],
    page_rect: Any,
) -> list[tuple[float, float, float, float]]:
    rows: list[tuple[float, list[tuple[float, float, float, float]]]] = []
    for drawing in drawings:
        rect = _coerce_pdf_rect(drawing.get("rect", ()), page_rect, padding=0.25)
        if rect is None:
            continue
        if rect[3] - rect[1] > 3.0 or rect[2] - rect[0] < 3.0:
            continue
        center_y = (rect[1] + rect[3]) / 2.0
        for row_index, (row_y, segments) in enumerate(rows):
            if abs(row_y - center_y) <= 1.5:
                segments.append(rect)
                rows[row_index] = (
                    (row_y * (len(segments) - 1) + center_y) / len(segments),
                    segments,
                )
                break
        else:
            rows.append((center_y, [rect]))

    merged: list[tuple[float, float, float, float]] = []
    for _row_y, segments in rows:
        ordered = sorted(segments, key=lambda rect: rect[0])
        current = ordered[0]
        for segment in ordered[1:]:
            if segment[0] - current[2] <= 3.0:
                current = _rect_union(current, segment)
            else:
                merged.append(current)
                current = segment
        merged.append(current)
    return merged


def _pdf_running_footer_separator_y(
    page_dict: Mapping[str, Any],
    page_rect: Any,
    rules: Iterable[tuple[float, float, float, float]],
) -> float | None:
    """Locate a journal footer rule so it cannot be mistaken for a table row.

    Papers commonly draw a full-width separator immediately above a DOI, date,
    and page-number footer.  Its width and alignment look exactly like a table
    border, but the nearby footer text provides strong page-furniture evidence.
    """

    page_y0 = float(getattr(page_rect, "y0", 0.0) or 0.0)
    page_width = max(1.0, float(getattr(page_rect, "width", 0.0) or 0.0))
    page_height = max(1.0, float(getattr(page_rect, "height", 0.0) or 0.0))
    footer_lines: list[tuple[tuple[float, float, float, float], str]] = []
    footer_start = page_y0 + page_height * 0.92
    for block in page_dict.get("blocks", ()):
        if int(block.get("type", 0) or 0) != 0:
            continue
        for line in block.get("lines", ()):
            line_rect = _text_line_rect(line, page_rect)
            if line_rect is None or line_rect[1] < footer_start:
                continue
            text = " ".join(_text_line_value(line).split())
            if text:
                footer_lines.append((line_rect, text))
    if not footer_lines:
        return None

    footer_text = " ".join(text for _rect, text in footer_lines)
    has_footer_marker = bool(
        re.search(
            r"https?://|\bdoi\b|\bpage\s+\d+|\b\d+\s*/\s*\d+\b|" r"\b(?:19|20)\d{2}\b",
            footer_text,
            re.IGNORECASE,
        )
    )
    left_present = any(
        rect[0] <= float(page_rect.x0) + page_width * 0.35 for rect, _ in footer_lines
    )
    right_present = any(
        rect[2] >= float(page_rect.x0) + page_width * 0.65 for rect, _ in footer_lines
    )
    if not has_footer_marker and not (
        left_present and right_present and len(footer_lines) >= 2
    ):
        return None

    footer_text_top = min(rect[1] for rect, _text in footer_lines)
    candidates = []
    for rule in rules:
        center_y = (rule[1] + rule[3]) / 2.0
        width_ratio = max(0.0, rule[2] - rule[0]) / page_width
        if (
            width_ratio >= 0.70
            and center_y <= footer_text_top
            and footer_text_top - center_y <= page_height * 0.04
        ):
            candidates.append(center_y)
    return max(candidates) if candidates else None


def _horizontal_rule_overlap_ratio(
    left: tuple[float, float, float, float],
    right: tuple[float, float, float, float],
) -> float:
    overlap = max(0.0, min(left[2], right[2]) - max(left[0], right[0]))
    minimum_width = max(
        1.0,
        min(left[2] - left[0], right[2] - right[0]),
    )
    return overlap / minimum_width


def _cluster_pdf_table_rule_groups(
    rules: Iterable[tuple[float, float, float, float]],
    page_rect: Any,
    *,
    min_rules: int = 2,
) -> list[list[tuple[float, float, float, float]]]:
    """Group aligned horizontal rules without merging adjacent tables.

    Scientific papers commonly use borderless tables with only three or four
    horizontal rules.  A column-local table is usually much narrower than 45%
    of the full page, so grouping is based on relative rule geometry instead of
    a full-page width threshold.
    """

    page_width = max(1.0, float(page_rect.width))
    page_height = max(1.0, float(page_rect.height))
    minimum_width = max(96.0, page_width * 0.25)
    pending = sorted(
        (rect for rect in rules if rect[2] - rect[0] >= minimum_width),
        key=lambda rect: ((rect[1] + rect[3]) / 2.0, rect[0]),
    )
    groups: list[list[tuple[float, float, float, float]]] = []
    while pending:
        group = [pending.pop(0)]
        changed = True
        while changed:
            changed = False
            remaining: list[tuple[float, float, float, float]] = []
            group_top = min(rect[1] for rect in group)
            group_bottom = max(rect[3] for rect in group)
            group_left = min(rect[0] for rect in group)
            group_right = max(rect[2] for rect in group)
            group_width = max(1.0, group_right - group_left)
            group_center = (group_left + group_right) / 2.0
            for candidate in pending:
                candidate_width = max(1.0, candidate[2] - candidate[0])
                width_ratio = min(group_width, candidate_width) / max(
                    group_width, candidate_width
                )
                candidate_center = (candidate[0] + candidate[2]) / 2.0
                vertical_gap = max(
                    candidate[1] - group_bottom,
                    group_top - candidate[3],
                    0.0,
                )
                group_span = (group_left, group_top, group_right, group_bottom)
                if (
                    width_ratio >= 0.70
                    and abs(candidate_center - group_center) <= page_width * 0.08
                    and _horizontal_rule_overlap_ratio(group_span, candidate) >= 0.75
                    and vertical_gap <= page_height * 0.30
                ):
                    group.append(candidate)
                    changed = True
                else:
                    remaining.append(candidate)
            pending = remaining
        distinct_rows = {round((rect[1] + rect[3]) / 2.0, 1) for rect in group}
        if len(distinct_rows) >= min_rules:
            groups.append(sorted(group, key=lambda rect: rect[1]))
    return groups


def _table_rect_from_rule_group(
    rule_group: Iterable[tuple[float, float, float, float]],
    page_dict: Mapping[str, Any],
    page_rect: Any,
) -> tuple[float, float, float, float] | None:
    rules = list(rule_group)
    if not rules:
        return None
    rect = rules[0]
    for rule in rules[1:]:
        rect = _rect_union(rect, rule)
    top = min((rule[1] + rule[3]) / 2.0 for rule in rules)
    bottom = max((rule[1] + rule[3]) / 2.0 for rule in rules)
    row_bands: set[float] = set()
    for block in page_dict.get("blocks", ()):
        if int(block.get("type", 0) or 0) != 0:
            continue
        for line in block.get("lines", ()):
            line_rect = _text_line_rect(line, page_rect)
            if line_rect is None:
                continue
            center_y = (line_rect[1] + line_rect[3]) / 2.0
            horizontal_overlap = max(
                0.0,
                min(line_rect[2], rect[2]) - max(line_rect[0], rect[0]),
            )
            if (
                top - 3.0 <= center_y <= bottom + 3.0
                and horizontal_overlap
                >= min(line_rect[2] - line_rect[0], rect[2] - rect[0]) * 0.20
            ):
                rect = _rect_union(rect, line_rect)
                row_bands.add(round(center_y / 2.0) * 2.0)
    if not row_bands:
        return None
    return _coerce_pdf_rect(rect, page_rect, padding=1.5)


def _detect_pdf_table_regions(
    page: Any,
    page_index: int,
    page_dict: Mapping[str, Any],
    drawings: list[Mapping[str, Any]],
    *,
    dpi: int,
) -> list[_HybridRegion]:
    page_rect = page.rect
    draw_items = sum(len(drawing.get("items", ()) or ()) for drawing in drawings)
    regions: list[_HybridRegion] = []
    for bbox in _pymupdf_table_bboxes(page, draw_items):
        rect = _coerce_pdf_rect(bbox, page_rect, padding=1.5)
        if rect is not None:
            regions.append(
                _HybridRegion(
                    page_index=page_index,
                    rect=rect,
                    kind="table",
                    reasons=("复杂表格或网格图",),
                    dpi=dpi,
                )
            )

    horizontal_rules = _merge_pdf_horizontal_rule_segments(drawings, page_rect)
    footer_separator_y = _pdf_running_footer_separator_y(
        page_dict,
        page_rect,
        horizontal_rules,
    )
    table_rules = [
        rule
        for rule in horizontal_rules
        if footer_separator_y is None
        or (rule[1] + rule[3]) / 2.0 < footer_separator_y - 0.5
    ]
    rule_groups = _cluster_pdf_table_rule_groups(
        table_rules,
        page_rect,
        min_rules=2,
    )

    # Three or more aligned rules are a strong table signal even when the table
    # occupies only one column and PyMuPDF's lattice detector returns nothing.
    for rule_group in rule_groups:
        if len(rule_group) < 3:
            continue
        table_rect = _table_rect_from_rule_group(rule_group, page_dict, page_rect)
        if table_rect is None:
            continue
        regions.append(
            _HybridRegion(
                page_index=page_index,
                rect=table_rect,
                kind="table",
                reasons=("横线表格",),
                dpi=dpi,
            )
        )

    for block in page_dict.get("blocks", ()):
        if int(block.get("type", 0) or 0) != 0:
            continue
        block_text = " ".join(_text_line_value(line) for line in block.get("lines", ()))
        normalized_title = " ".join(block_text.split())
        if not _TABLE_TITLE_PATTERN.match(normalized_title):
            continue
        title_rect = _coerce_pdf_rect(block.get("bbox", ()), page_rect)
        if title_rect is None:
            continue
        nearby_groups: list[tuple[float, list[tuple[float, float, float, float]]]] = []
        for rule_group in rule_groups:
            group_rect = rule_group[0]
            for rule in rule_group[1:]:
                group_rect = _rect_union(group_rect, rule)
            horizontal_overlap = max(
                0.0,
                min(title_rect[2], group_rect[2]) - max(title_rect[0], group_rect[0]),
            )
            if (
                horizontal_overlap
                < min(
                    title_rect[2] - title_rect[0],
                    group_rect[2] - group_rect[0],
                )
                * 0.35
            ):
                continue
            vertical_gap = max(
                title_rect[1] - group_rect[3],
                group_rect[1] - title_rect[3],
                0.0,
            )
            if vertical_gap <= 240.0:
                nearby_groups.append((vertical_gap, rule_group))
        if not nearby_groups:
            continue
        _gap, nearest_group = min(nearby_groups, key=lambda item: item[0])
        table_rect = _table_rect_from_rule_group(
            nearest_group,
            page_dict,
            page_rect,
        )
        if table_rect is None:
            continue
        regions.append(
            _HybridRegion(
                page_index=page_index,
                rect=table_rect,
                kind="table",
                reasons=("横线表格",),
                dpi=dpi,
            )
        )
    return _merge_hybrid_regions(regions, page_rect, gap=2.0)


def _pdf_layout_text_lines(
    page_dict: Mapping[str, Any],
    page_rect: Any,
) -> tuple[tuple[tuple[float, float, float, float], int], ...]:
    """Return body text lines suitable for geometric column analysis."""

    page_height = max(1.0, float(page_rect.height))
    top_limit = float(page_rect.y0) + page_height * 0.05
    bottom_limit = float(page_rect.y1) - page_height * 0.07
    entries: list[tuple[tuple[float, float, float, float], int]] = []
    for block in page_dict.get("blocks", ()):
        if int(block.get("type", 0) or 0) != 0:
            continue
        for line in block.get("lines", ()):
            line_rect = _text_line_rect(line, page_rect)
            if line_rect is None:
                continue
            center_y = (line_rect[1] + line_rect[3]) / 2.0
            if not (top_limit <= center_y <= bottom_limit):
                continue
            text = " ".join(_text_line_value(line).split())
            weight = len(_normalize_validation_text(text))
            if weight >= 4:
                entries.append((line_rect, weight))
    return tuple(entries)


def _pdf_column_profile_for_split(
    lines: Iterable[tuple[tuple[float, float, float, float], int]],
    *,
    split_x: float,
    gutter_margin: float,
) -> _PdfColumnProfile:
    left: list[tuple[tuple[float, float, float, float], int]] = []
    right: list[tuple[tuple[float, float, float, float], int]] = []
    crossing: list[tuple[tuple[float, float, float, float], int]] = []
    for entry in lines:
        rect, _weight = entry
        if rect[2] <= split_x - gutter_margin:
            left.append(entry)
        elif rect[0] >= split_x + gutter_margin:
            right.append(entry)
        else:
            crossing.append(entry)
    return _PdfColumnProfile(
        split_x=float(split_x),
        left_lines=tuple(left),
        right_lines=tuple(right),
        crossing_lines=tuple(crossing),
    )


def _pdf_standard_two_column_profile(
    page_dict: Mapping[str, Any],
    page_rect: Any,
) -> _PdfColumnProfile:
    page_width = max(1.0, float(page_rect.width))
    midpoint = float(page_rect.x0) + page_width / 2.0
    return _pdf_column_profile_for_split(
        _pdf_layout_text_lines(page_dict, page_rect),
        split_x=midpoint,
        gutter_margin=max(1.0, page_width * _PDF_COLUMN_GUTTER_MARGIN_RATIO),
    )


def _pdf_column_profile_looks_two_column(profile: _PdfColumnProfile) -> bool:
    left_weight = sum(weight for _rect, weight in profile.left_lines)
    right_weight = sum(weight for _rect, weight in profile.right_lines)
    crossing_weight = sum(weight for _rect, weight in profile.crossing_lines)
    confined_weight = left_weight + right_weight
    total_weight = confined_weight + crossing_weight
    if (
        len(profile.left_lines) < 6
        or len(profile.right_lines) < 6
        or total_weight < 160
    ):
        return False
    return (
        left_weight / max(1, total_weight) >= 0.18
        and right_weight / max(1, total_weight) >= 0.18
        and confined_weight / max(1, total_weight) >= 0.62
    )


def _pdf_page_looks_two_column(
    page_dict: Mapping[str, Any],
    page_rect: Any,
) -> bool:
    """Infer a two-column body while tolerating full-width titles and figures."""

    return _pdf_column_profile_looks_two_column(
        _pdf_standard_two_column_profile(page_dict, page_rect)
    )


def _pdf_asymmetric_column_profile(
    page_dict: Mapping[str, Any],
    page_rect: Any,
) -> tuple[_PdfColumnProfile, bool] | None:
    """Return the strongest narrow/wide parallel-flow profile on a PDF page.

    Word columns are especially unstable for this topology: the short/narrow side is
    rebalanced into the main body and lines at the same height are often interleaved.
    The boolean result is true when the narrow flow is on the left.
    """

    lines = _pdf_layout_text_lines(page_dict, page_rect)
    if len(lines) < 10:
        return None
    page_width = max(1.0, float(page_rect.width))
    page_height = max(1.0, float(page_rect.height))
    page_x0 = float(page_rect.x0)
    gutter_margin = max(1.0, page_width * 0.004)
    best: tuple[float, _PdfColumnProfile] | None = None
    for step in range(22, 79):
        split_x = page_x0 + page_width * step / 100.0
        profile = _pdf_column_profile_for_split(
            lines,
            split_x=split_x,
            gutter_margin=gutter_margin,
        )
        if len(profile.left_lines) < 5 or len(profile.right_lines) < 5:
            continue
        left_weight = sum(weight for _rect, weight in profile.left_lines)
        right_weight = sum(weight for _rect, weight in profile.right_lines)
        crossing_weight = sum(weight for _rect, weight in profile.crossing_lines)
        total_weight = left_weight + right_weight + crossing_weight
        if total_weight < 500:
            continue
        confined_ratio = (left_weight + right_weight) / total_weight
        minority_ratio = min(left_weight, right_weight) / total_weight
        score = confined_ratio + minority_ratio * 0.35
        if best is None or score > best[0]:
            best = (score, profile)
    if best is None:
        return None

    profile = best[1]
    left_weight = sum(weight for _rect, weight in profile.left_lines)
    right_weight = sum(weight for _rect, weight in profile.right_lines)
    crossing_weight = sum(weight for _rect, weight in profile.crossing_lines)
    total_weight = left_weight + right_weight + crossing_weight
    if (left_weight + right_weight) / total_weight < 0.88:
        return None
    if min(left_weight, right_weight) / total_weight < 0.04:
        return None

    left_width = median(rect[2] - rect[0] for rect, _weight in profile.left_lines)
    right_width = median(rect[2] - rect[0] for rect, _weight in profile.right_lines)
    narrow_width = min(left_width, right_width)
    wide_width = max(left_width, right_width)
    if (
        narrow_width < page_width * 0.14
        or wide_width < page_width * 0.38
        or narrow_width / max(1.0, wide_width) > 0.58
    ):
        return None

    left_y0 = min(rect[1] for rect, _weight in profile.left_lines)
    left_y1 = max(rect[3] for rect, _weight in profile.left_lines)
    right_y0 = min(rect[1] for rect, _weight in profile.right_lines)
    right_y1 = max(rect[3] for rect, _weight in profile.right_lines)
    overlap_height = max(0.0, min(left_y1, right_y1) - max(left_y0, right_y0))
    if overlap_height < page_height * 0.08:
        return None
    return profile, left_width <= right_width


def _pdf_asymmetric_column_risk_reason(
    page_dict: Mapping[str, Any],
    page_rect: Any,
) -> str | None:
    if _pdf_asymmetric_column_profile(page_dict, page_rect) is None:
        return None
    return "不等宽侧栏与主正文混排"


def _pdf_asymmetric_sidebar_region(
    page_index: int,
    page_dict: Mapping[str, Any],
    drawings: Iterable[Mapping[str, Any]],
    page_rect: Any,
    *,
    dpi: int,
) -> _HybridRegion | None:
    """Locate the short narrow flow that should be restored as a page overlay.

    The detector deliberately keeps the title, running header/footer and any later
    regular two-column body out of the crop.  A nearby narrow background rectangle
    may extend the crop (for example an author-information card), but it may never
    grow into the wide editable body.
    """

    result = _pdf_asymmetric_column_profile(page_dict, page_rect)
    if result is None:
        return None
    profile, narrow_on_left = result
    narrow_lines = profile.left_lines if narrow_on_left else profile.right_lines
    wide_lines = profile.right_lines if narrow_on_left else profile.left_lines
    if not narrow_lines or not wide_lines:
        return None

    page_width = max(1.0, float(page_rect.width))
    page_height = max(1.0, float(page_rect.height))
    page_x0 = float(page_rect.x0)
    page_y0 = float(page_rect.y0)
    page_x1 = float(page_rect.x1)
    page_y1 = float(page_rect.y1)
    line_heights = [
        rect[3] - rect[1] for rect, _weight in narrow_lines if rect[3] > rect[1]
    ]
    typical_line_height = median(line_heights) if line_heights else 10.0
    vertical_tolerance = max(5.0, typical_line_height * 1.5)
    wide_y0 = min(rect[1] for rect, _weight in wide_lines)
    wide_y1 = max(rect[3] for rect, _weight in wide_lines)
    candidates = [
        entry
        for entry in narrow_lines
        if wide_y0 - vertical_tolerance
        <= (entry[0][1] + entry[0][3]) / 2.0
        <= wide_y1 + vertical_tolerance
    ]
    if len(candidates) < 3:
        return None

    # A running logo can occupy the same side as the sidebar.  Split the narrow
    # flow into vertical clusters so that the heavier metadata/author card wins,
    # while a later ordinary two-column paragraph remains editable.
    merge_gap = max(10.0, typical_line_height * 2.25)
    clusters: list[list[tuple[tuple[float, float, float, float], int]]] = []
    for entry in sorted(candidates, key=lambda item: (item[0][1], item[0][0])):
        if not clusters:
            clusters.append([entry])
            continue
        previous_y1 = max(item[0][3] for item in clusters[-1])
        if entry[0][1] <= previous_y1 + merge_gap:
            clusters[-1].append(entry)
        else:
            clusters.append([entry])
    clusters = [cluster for cluster in clusters if len(cluster) >= 3]
    if not clusters:
        return None
    cluster = max(
        clusters,
        key=lambda items: (
            sum(weight for _rect, weight in items),
            max(rect[3] for rect, _weight in items)
            - min(rect[1] for rect, _weight in items),
            len(items),
        ),
    )
    neighbor_gap = max(merge_gap, page_height * 0.075)
    selected_y0 = min(rect[1] for rect, _weight in cluster)
    selected_y1 = max(rect[3] for rect, _weight in cluster)
    for neighboring_cluster in clusters:
        if neighboring_cluster is cluster:
            continue
        neighbor_y0 = min(rect[1] for rect, _weight in neighboring_cluster)
        neighbor_y1 = max(rect[3] for rect, _weight in neighboring_cluster)
        vertical_gap = max(
            0.0,
            max(selected_y0, neighbor_y0) - min(selected_y1, neighbor_y1),
        )
        if vertical_gap <= neighbor_gap:
            cluster.extend(neighboring_cluster)
            selected_y0 = min(selected_y0, neighbor_y0)
            selected_y1 = max(selected_y1, neighbor_y1)
    cluster_weight = sum(weight for _rect, weight in cluster)
    total_weight = sum(weight for _rect, weight in (*narrow_lines, *wide_lines))
    if cluster_weight < max(35, int(total_weight * 0.03)):
        return None

    padding = max(3.0, min(9.0, typical_line_height * 0.5))
    x0 = min(rect[0] for rect, _weight in cluster) - padding
    y0 = min(rect[1] for rect, _weight in cluster) - padding
    x1 = max(rect[2] for rect, _weight in cluster) + padding
    y1 = max(rect[3] for rect, _weight in cluster) + padding
    overlapping_wide_lines = [
        rect
        for rect, _weight in wide_lines
        if rect[3] >= y0 - vertical_tolerance and rect[1] <= y1 + vertical_tolerance
    ]
    if not overlapping_wide_lines:
        return None
    if narrow_on_left:
        wide_edge = min(rect[0] for rect in overlapping_wide_lines)
        side_limit = wide_edge - max(2.0, page_width * 0.004)
        x1 = min(x1, side_limit)
    else:
        wide_edge = max(rect[2] for rect in overlapping_wide_lines)
        side_limit = wide_edge + max(2.0, page_width * 0.004)
        x0 = max(x0, side_limit)

    # Preserve a card background or sidebar decoration with the text.  Reject
    # page-spanning rules and shapes that enter the main editable column.
    seed = (x0, y0, x1, y1)
    for drawing in drawings:
        drawing_rect = _coerce_pdf_rect(
            drawing.get("rect", ()),
            page_rect,
            padding=0.0,
        )
        if drawing_rect is None:
            continue
        dx0, dy0, dx1, dy1 = drawing_rect
        if (dx1 - dx0) * (dy1 - dy0) > page_width * page_height * 0.24:
            continue
        if narrow_on_left and dx1 > side_limit + max(2.0, page_width * 0.01):
            continue
        if not narrow_on_left and dx0 < side_limit - max(2.0, page_width * 0.01):
            continue
        overlap_width = max(0.0, min(seed[2], dx1) - max(seed[0], dx0))
        overlap_height = max(0.0, min(seed[3], dy1) - max(seed[1], dy0))
        required_overlap_width = max(
            1.0,
            min(seed[2] - seed[0], dx1 - dx0) * 0.08,
        )
        required_overlap_height = max(
            1.0,
            min(seed[3] - seed[1], dy1 - dy0) * 0.02,
        )
        if (
            overlap_width < required_overlap_width
            or overlap_height < required_overlap_height
        ):
            continue
        x0 = min(x0, dx0)
        y0 = min(y0, dy0)
        x1 = max(x1, dx1)
        y1 = max(y1, dy1)

    x0 = max(page_x0, x0)
    y0 = max(page_y0, y0)
    x1 = min(page_x1, x1)
    y1 = min(page_y1, y1)
    width = max(0.0, x1 - x0)
    height = max(0.0, y1 - y0)
    if (
        width <= 0.0
        or height <= 0.0
        or width / page_width > 0.42
        or width * height / (page_width * page_height) > 0.35
    ):
        return None
    return _HybridRegion(
        page_index=page_index,
        rect=(x0, y0, x1, y1),
        kind="sidebar",
        reasons=("不等宽侧栏局部高清保真",),
        dpi=min(600, max(180, int(dpi))),
    )


def _filter_hybrid_regions_for_anchored_overlays(
    regions: Iterable[_HybridRegion],
    anchored_regions: Iterable[_HybridRegion],
    page_rect: Any,
) -> list[_HybridRegion]:
    """Remove visual detections already represented by a floating page crop."""

    overlays = tuple(anchored_regions)
    if not overlays:
        return list(regions)
    page_width = max(1.0, float(page_rect.width))
    filtered: list[_HybridRegion] = []
    for region in regions:
        region_area = max(1.0, _rect_area(region.rect))
        covered_ratio = max(
            (
                max(
                    0.0,
                    min(region.rect[2], overlay.rect[2])
                    - max(region.rect[0], overlay.rect[0]),
                )
                * max(
                    0.0,
                    min(region.rect[3], overlay.rect[3])
                    - max(region.rect[1], overlay.rect[1]),
                )
                / region_area
            )
            for overlay in overlays
        )
        if covered_ratio >= 0.72:
            continue
        is_sidebar_vector_grouping_artifact = (
            region.kind == "vector"
            and "高密度矢量图形" in region.reasons
            and (region.rect[2] - region.rect[0]) / page_width >= 0.72
            and covered_ratio >= 0.08
        )
        if is_sidebar_vector_grouping_artifact:
            continue
        filtered.append(region)
    return filtered


def _pdf_two_column_pair_bands(
    profile: _PdfColumnProfile,
) -> tuple[tuple[float, float], ...]:
    """Locate rows where independent left/right lines coexist around a gutter."""

    if not _pdf_column_profile_looks_two_column(profile):
        return ()
    heights = [
        rect[3] - rect[1]
        for rect, _weight in (*profile.left_lines, *profile.right_lines)
        if rect[3] > rect[1]
    ]
    typical_line_height = median(heights) if heights else 10.0
    tolerance = max(3.0, typical_line_height * 0.55)
    merge_gap = max(6.0, typical_line_height * 2.5)
    bands: list[tuple[float, float]] = []
    for left_rect, _left_weight in profile.left_lines:
        left_center = (left_rect[1] + left_rect[3]) / 2.0
        matches = [
            right_rect
            for right_rect, _right_weight in profile.right_lines
            if abs((right_rect[1] + right_rect[3]) / 2.0 - left_center) <= tolerance
        ]
        if not matches:
            continue
        nearest = min(
            matches,
            key=lambda rect: abs((rect[1] + rect[3]) / 2.0 - left_center),
        )
        bands.append((min(left_rect[1], nearest[1]), max(left_rect[3], nearest[3])))
    if not bands:
        return ()
    bands.sort()
    merged: list[tuple[float, float]] = [bands[0]]
    for y0, y1 in bands[1:]:
        previous_y0, previous_y1 = merged[-1]
        if y0 <= previous_y1 + merge_gap:
            merged[-1] = (previous_y0, max(previous_y1, y1))
        else:
            merged.append((y0, y1))
    return tuple(merged)


def _resolve_hybrid_auto_column_layout(
    column_layout: str,
    assessments: Iterable[_HybridPageAssessment],
) -> str:
    normalized = _normalize_pdf2docx_column_layout(column_layout)
    if normalized != "auto":
        return normalized
    text_pages = [
        assessment
        for assessment in assessments
        if len(_normalize_validation_text(assessment.source_text)) >= 160
    ]
    if not text_pages:
        return normalized
    detected = sum(1 for assessment in text_pages if assessment.detected_two_columns)
    if detected >= max(1, math.ceil(len(text_pages) * 0.35)):
        # "mixed" retains full-width paper titles, figures and tables while
        # enabling column-local formula detection and the more sensitive
        # pdf2docx section analysis used for two-column papers.
        return "mixed"
    return normalized


def _pdf_word_layout_profile(
    source: Path,
    password: str | None,
    *,
    column_layout: str,
) -> _PdfWordLayoutProfile:
    """Choose fixed coordinates for papers and design-heavy single pages.

    Ordinary prose remains on the flow-based pdf2docx path.  Dense two-column
    papers and pages whose appearance is carried by many vector decorations do
    not survive Word reflow reliably, so they use the coordinate-editable path.
    """

    normalized_layout = _normalize_pdf2docx_column_layout(column_layout)
    _pymupdf, document = _open_pymupdf_document(source, password)
    text_pages = 0
    two_column_pages = 0
    designed_pages: list[int] = []
    try:
        for page_index, page in enumerate(document):
            try:
                page_dict = page.get_text("dict") or {}
                drawings = page.get_drawings() or ()
            except Exception:
                continue
            text_lines = sum(
                len(block.get("lines", ()))
                for block in page_dict.get("blocks", ())
                if int(block.get("type", 0) or 0) == 0
            )
            text_characters = len(
                _normalize_validation_text(str(page.get_text() or ""))
            )
            if text_characters >= 80:
                text_pages += 1
            profile = _pdf_standard_two_column_profile(page_dict, page.rect)
            if _pdf_column_profile_looks_two_column(profile):
                two_column_pages += 1

            image_blocks = sum(
                int(block.get("type", 0) or 0) == 1
                for block in page_dict.get("blocks", ())
            )
            drawing_items = sum(len(drawing.get("items", ())) for drawing in drawings)
            design_heavy = bool(
                text_lines >= 12
                and (
                    drawing_items >= 80
                    or (
                        image_blocks >= 1 and len(drawings) >= 8 and drawing_items >= 24
                    )
                )
            )
            if design_heavy:
                designed_pages.append(page_index)
        page_count = int(document.page_count)
    finally:
        document.close()

    reasons: list[str] = []
    explicit_multi_column = normalized_layout in {"double", "mixed"}
    detected_paper = bool(
        normalized_layout != "single"
        and text_pages
        and two_column_pages >= max(1, math.ceil(text_pages * 0.35))
    )
    if explicit_multi_column or detected_paper:
        reasons.append("双栏/混合分栏论文")
    if designed_pages:
        reasons.append("设计型复杂页面")
    return _PdfWordLayoutProfile(
        page_count=page_count,
        text_pages=text_pages,
        two_column_pages=two_column_pages,
        designed_pages=tuple(designed_pages),
        fixed_layout_recommended=bool(reasons),
        reasons=tuple(reasons),
    )


def _detect_pdf_formula_regions(
    page_index: int,
    page_dict: Mapping[str, Any],
    page_rect: Any,
    *,
    dpi: int,
    column_layout: str = "auto",
) -> list[_HybridRegion]:
    normalized_column_layout = _normalize_pdf2docx_column_layout(column_layout)
    column_aware = normalized_column_layout in {"double", "mixed"}
    page_width = max(1.0, float(page_rect.width))
    page_x0 = float(page_rect.x0)
    page_center_x = float(page_rect.x0) + page_width / 2.0

    def line_column_id(
        rect: tuple[float, float, float, float],
        *,
        left_number_anchor: bool = False,
        right_number_anchor: bool = False,
    ) -> str:
        if not column_aware:
            return "page"
        if left_number_anchor:
            return "left"
        if right_number_anchor:
            return "right"
        x0_ratio = (rect[0] - page_x0) / page_width
        x1_ratio = (rect[2] - page_x0) / page_width
        center_ratio = ((rect[0] + rect[2]) / 2.0 - page_x0) / page_width
        if x0_ratio < 0.42 and x1_ratio > 0.58:
            return "full"
        return "left" if center_ratio <= 0.52 else "right"

    def columns_are_compatible(left: str, right: str) -> bool:
        if not column_aware:
            return True
        return left == right or "full" in {left, right}

    def horizontal_gap(
        left: tuple[float, float, float, float],
        right: tuple[float, float, float, float],
    ) -> float:
        return max(left[0] - right[2], right[0] - left[2], 0.0)

    line_entries: list[dict[str, Any]] = []
    for block in page_dict.get("blocks", ()):
        if int(block.get("type", 0) or 0) != 0:
            continue
        for line in block.get("lines", ()):
            line_rect = _text_line_rect(line, page_rect)
            if line_rect is None or line_rect[3] - line_rect[1] > 36.0:
                continue
            spans = list(line.get("spans", ()))
            text = _repair_known_pdf_text_encoding(_text_line_value(line))
            normalized_text = _normalize_validation_text(text)
            if not normalized_text or len(normalized_text) > 180:
                continue
            if _hybrid_text_line_is_exempt(line_rect, text, page_rect):
                continue
            math_characters = sum(
                len(_normalize_validation_text(str(span.get("text", "") or "")))
                for span in spans
                if _math_font(str(span.get("font", "")))
            )
            plain_text = _repair_known_pdf_text_encoding(
                "".join(
                    str(span.get("text", "") or "")
                    for span in spans
                    if not _math_font(str(span.get("font", "")))
                )
            )
            math_ratio = math_characters / max(1, len(normalized_text))
            has_unportable_symbol = any(
                _font_requires_visual_fallback(str(span.get("font", "")))
                for span in spans
            )
            has_suspicious_character = _has_suspicious_pdf_characters(text)
            prose_words = [
                word
                for word in re.findall(r"[A-Za-z]{3,}", plain_text)
                if word.casefold() not in _FORMULA_FUNCTION_WORDS
            ]
            prose_letters = sum(len(word) for word in prose_words)
            prose_like = (len(prose_words) >= 4 and len(normalized_text) >= 36) or (
                prose_letters >= 4 and math_ratio < 0.75
            )
            normalized_number = " ".join(text.split())
            parenthesized_number = bool(
                normalized_number.startswith("(")
                and normalized_number.endswith(")")
                and _EQUATION_NUMBER_PATTERN.fullmatch(normalized_number)
            )
            nonempty_spans = [
                span
                for span in spans
                if _normalize_validation_text(str(span.get("text", "") or ""))
            ]
            encoded_number = False
            if len(nonempty_spans) >= 3:
                middle_text = "".join(
                    str(span.get("text", "") or "") for span in nonempty_spans[1:-1]
                )
                middle_number = _normalize_validation_text(middle_text)
                first_wrapper_text = str(nonempty_spans[0].get("text", "") or "")
                last_wrapper_text = str(nonempty_spans[-1].get("text", "") or "")
                first_wrapper_font = str(nonempty_spans[0].get("font", ""))
                last_wrapper_font = str(nonempty_spans[-1].get("font", ""))
                encoded_number = bool(
                    re.fullmatch(r"\d{1,3}[a-z]?", middle_number, re.IGNORECASE)
                    and (
                        _math_font(first_wrapper_font)
                        or _font_requires_visual_fallback(first_wrapper_font)
                        or _has_suspicious_pdf_characters(first_wrapper_text)
                    )
                    and (
                        _math_font(last_wrapper_font)
                        or _font_requires_visual_fallback(last_wrapper_font)
                        or _has_suspicious_pdf_characters(last_wrapper_text)
                    )
                )
            equation_number_text = parenthesized_number or encoded_number
            line_center_ratio = (
                (line_rect[0] + line_rect[2]) / 2.0 - page_x0
            ) / page_width
            right_number_anchor = bool(
                equation_number_text and line_rect[0] >= page_x0 + page_width * 0.70
            )
            left_number_anchor = bool(
                column_aware
                and equation_number_text
                and 0.42 <= line_center_ratio <= 0.62
            )
            equation_number = right_number_anchor or left_number_anchor
            display_operator = bool(_DISPLAY_FORMULA_OPERATOR_PATTERN.search(text))
            line_width_ratio = (line_rect[2] - line_rect[0]) / page_width
            line_center_x = (line_rect[0] + line_rect[2]) / 2.0
            independently_centered = (
                line_width_ratio <= 0.65
                and abs(line_center_x - page_center_x) <= page_width * 0.30
            )
            has_italic_math_font = any(
                "ital" in _word_font_key(str(span.get("font", ""))) for span in spans
            )
            short_symbolic = (
                len(normalized_text) <= 48
                and len(prose_words) <= 2
                and (
                    math_characters >= 1
                    or has_italic_math_font
                    or bool(re.search(r"[^\w\s.,;:'\"-]", text, re.UNICODE))
                )
            )
            formula_support = not equation_number and (
                display_operator
                or math_characters >= 1
                or has_unportable_symbol
                or has_suspicious_character
                or short_symbolic
            )
            is_seed = (
                has_unportable_symbol
                or has_suspicious_character
                or (
                    display_operator
                    and not prose_like
                    and (math_characters >= 1 or independently_centered)
                )
                or (independently_centered and math_ratio >= 0.30 and not prose_like)
            )
            normalized_rect = _coerce_pdf_rect(
                line_rect,
                page_rect,
                padding=0.75,
            )
            if normalized_rect is not None:
                line_entries.append(
                    {
                        "rect": normalized_rect,
                        "center_y": (line_rect[1] + line_rect[3]) / 2.0,
                        "equation_number": equation_number,
                        "column_id": line_column_id(
                            normalized_rect,
                            left_number_anchor=left_number_anchor,
                            right_number_anchor=right_number_anchor,
                        ),
                        "formula_support": formula_support,
                        "seed": is_seed,
                        "prose_like": prose_like,
                        "prose_word_count": len(prose_words),
                        "normalized_length": len(normalized_text),
                        "body_prose_signal": (
                            not display_operator
                            and math_characters == 0
                            and not has_unportable_symbol
                            and not has_suspicious_character
                            and not has_italic_math_font
                            and bool(prose_words)
                        ),
                        "same_baseline_prose": False,
                    }
                )

    for entry in line_entries:
        baseline_prose = [
            other
            for other in line_entries
            if other is not entry
            and other["body_prose_signal"]
            and columns_are_compatible(entry["column_id"], other["column_id"])
            and horizontal_gap(entry["rect"], other["rect"])
            <= max(12.0, page_width * 0.04)
            and (
                abs(float(entry["center_y"]) - float(other["center_y"])) <= 6.0
                or _rect_intersection_ratio(entry["rect"], other["rect"]) >= 0.35
            )
        ]
        entry["same_baseline_prose"] = (
            sum(int(other["prose_word_count"]) for other in baseline_prose) >= 4
            or sum(int(other["normalized_length"]) for other in baseline_prose) >= 36
        )

    formula_dpi = min(600, max(int(dpi), 450))
    numbered_regions: dict[str, list[_HybridRegion]] = {}
    seed_regions: dict[str, list[_HybridRegion]] = {}

    # A right-side equation number is the most reliable anchor for formulas
    # split by the PDF text layer into many numerator, denominator, radical,
    # subscript, and superscript lines.  Keep the vertical band fixed around
    # the number so the closure cannot drift into the explanatory prose above
    # or below the display equation.
    claimed_by_number: set[int] = set()
    for number_index, number_entry in enumerate(line_entries):
        if not number_entry["equation_number"]:
            continue
        rect = number_entry["rect"]
        attached = False
        number_center_y = float(number_entry["center_y"])
        number_column_id = str(number_entry["column_id"])
        attached_indexes: list[int] = []
        for candidate_index, candidate in enumerate(line_entries):
            if (
                not candidate["formula_support"]
                or candidate["same_baseline_prose"]
                or not columns_are_compatible(
                    number_column_id,
                    str(candidate["column_id"]),
                )
            ):
                continue
            if abs(float(candidate["center_y"]) - number_center_y) > 22.0:
                continue
            rect = _rect_union(rect, candidate["rect"])
            attached = True
            attached_indexes.append(candidate_index)
        if attached:
            claimed_by_number.add(number_index)
            claimed_by_number.update(attached_indexes)
            numbered_regions.setdefault(number_column_id, []).append(
                _HybridRegion(
                    page_index=page_index,
                    rect=rect,
                    kind="formula",
                    reasons=(_NUMBERED_FORMULA_REASON,),
                    dpi=formula_dpi,
                )
            )

    support_entries = [
        (entry_index, entry)
        for entry_index, entry in enumerate(line_entries)
        if entry_index not in claimed_by_number
        and entry["formula_support"]
        and not entry["prose_like"]
        and not entry["same_baseline_prose"]
    ]
    for seed_index, seed_entry in enumerate(line_entries):
        if (
            seed_index in claimed_by_number
            or not seed_entry["seed"]
            or seed_entry["same_baseline_prose"]
        ):
            continue
        rect = seed_entry["rect"]
        seed_center_y = float(seed_entry["center_y"])
        seed_column_id = str(seed_entry["column_id"])
        for support_index, support_entry in support_entries:
            if support_index == seed_index:
                continue
            if not columns_are_compatible(
                seed_column_id,
                str(support_entry["column_id"]),
            ):
                continue
            if abs(float(support_entry["center_y"]) - seed_center_y) > 22.0:
                continue
            support_rect = support_entry["rect"]
            if _rect_distance(rect, support_rect) > 12.0:
                continue
            rect = _rect_union(rect, support_rect)
        seed_regions.setdefault(seed_column_id, []).append(
            _HybridRegion(
                page_index=page_index,
                rect=rect,
                kind="formula",
                reasons=("复杂公式或符号字形",),
                dpi=formula_dpi,
            )
        )
    merged_regions: list[_HybridRegion] = []
    for column_id in sorted(set(numbered_regions) | set(seed_regions)):
        merged_numbered = _merge_hybrid_regions(
            numbered_regions.get(column_id, ()),
            page_rect,
            gap=1.0,
        )
        merged_seeds = _merge_hybrid_regions(
            seed_regions.get(column_id, ()),
            page_rect,
            gap=8.0,
        )
        merged_regions.extend(
            _merge_hybrid_regions(
                [*merged_numbered, *merged_seeds],
                page_rect,
                gap=1.0,
            )
        )
    merged_regions.sort(key=lambda item: (item.rect[1], item.rect[0]))
    return _constrain_numbered_formula_regions_to_body_prose(
        merged_regions,
        page_dict,
        page_rect,
    )


def _extract_editable_pdf_text_blocks(
    page_dict: Mapping[str, Any],
    page_rect: Any,
    regions: Iterable[_HybridRegion],
) -> tuple[str, ...]:
    region_rects = [region.rect for region in regions]
    blocks: list[str] = []
    for block in page_dict.get("blocks", ()):
        if int(block.get("type", 0) or 0) != 0:
            continue
        lines: list[str] = []
        for line in block.get("lines", ()):
            line_rect = _text_line_rect(line, page_rect)
            line_is_exempt = line_rect is not None and _hybrid_text_line_is_exempt(
                line_rect,
                _text_line_value(line),
                page_rect,
            )
            # Captions and running furniture are normally kept editable even
            # when they merely touch a nearby visual island.  If the island
            # actually covers almost the whole line, however, redaction will
            # remove that line from the editable PDF.  Do not include such
            # already-rasterized text in the page quality baseline.
            if line_is_exempt and any(
                _rect_intersection_ratio(line_rect, region_rect) >= 0.85
                for region_rect in region_rects
            ):
                line_is_exempt = False
            parts: list[str] = []
            for span in line.get("spans", ()):
                text = str(span.get("text", "") or "")
                span_rect = _coerce_pdf_rect(span.get("bbox", ()), page_rect)
                if (
                    not line_is_exempt
                    and span_rect is not None
                    and any(
                        _hybrid_text_rect_is_excluded(span_rect, region_rect)
                        for region_rect in region_rects
                    )
                ):
                    continue
                parts.append(text)
            line_text = _repair_known_pdf_text_encoding("".join(parts))
            if line_text.strip():
                lines.append(line_text)
        block_text = "\n".join(lines)
        if block_text.strip():
            blocks.append(block_text)
    return tuple(blocks)


def _add_hybrid_visual_regions(
    assessment: _HybridPageAssessment,
    page_dict: Mapping[str, Any],
    page_rect: Any,
    regions: Iterable[_HybridRegion],
) -> None:
    """Add late-discovered local fallbacks and refresh the quality baseline."""

    combined = _merge_hybrid_regions(
        [*assessment.visual_regions, *regions],
        page_rect,
        gap=1.0,
    )
    combined = _close_hybrid_regions_over_text(combined, page_dict, page_rect)
    combined = _constrain_numbered_formula_regions_to_body_prose(
        combined,
        page_dict,
        page_rect,
    )
    assessment.visual_regions = combined
    editable_blocks = _extract_editable_pdf_text_blocks(
        page_dict,
        page_rect,
        combined,
    )
    assessment.editable_text_blocks = editable_blocks
    assessment.editable_source_text = "\n".join(editable_blocks)


def _hybrid_region_area_ratio(
    regions: Iterable[_HybridRegion], page_rect: Any
) -> float:
    page_area = max(1.0, float(page_rect.width) * float(page_rect.height))
    return min(1.0, sum(_rect_area(region.rect) for region in regions) / page_area)


def _hybrid_dense_two_column_figure_risk(
    regions: Iterable[_HybridRegion],
    page_rect: Any,
    *,
    column_layout: str,
) -> bool:
    """Detect multi-figure column pages that Word/WPS tends to rebalance badly."""

    normalized_layout = _normalize_pdf2docx_column_layout(column_layout)
    if normalized_layout == "single":
        return False
    page_width = max(1.0, float(page_rect.width))
    page_area = max(1.0, page_width * float(page_rect.height))
    page_midpoint = float(page_rect.x0) + page_width / 2.0
    figure_regions = [region for region in regions if region.kind == "figure"]
    if len(figure_regions) < 4:
        return False
    left_figures = 0
    right_figures = 0
    figure_area = 0.0
    for region in figure_regions:
        width = max(0.0, region.rect[2] - region.rect[0])
        if width > page_width * 0.58:
            return False
        figure_area += _rect_area(region.rect)
        center_x = (region.rect[0] + region.rect[2]) / 2.0
        if center_x <= page_midpoint:
            left_figures += 1
        else:
            right_figures += 1
    return left_figures >= 2 and right_figures >= 2 and figure_area / page_area >= 0.28


def _hybrid_bottom_full_width_visual_overflow_risk(
    regions: Iterable[_HybridRegion],
    page_rect: Any,
) -> bool:
    """Detect only an exceptionally large indivisible island at page bottom.

    Ordinary bottom tables are intentionally kept as local visual islands.  A
    whole-page fallback is justified only when the island itself occupies about
    a third of the page and reaches the physical bottom; otherwise rasterizing
    the complete page sacrifices far more editability than it protects.
    """

    page_width = max(1.0, float(page_rect.width))
    page_height = max(1.0, float(page_rect.height))
    page_x0 = float(page_rect.x0)
    page_y0 = float(page_rect.y0)
    for region in regions:
        if region.kind not in {"table", "complex"}:
            continue
        x0, y0, x1, y1 = region.rect
        width_ratio = max(0.0, x1 - x0) / page_width
        height_ratio = max(0.0, y1 - y0) / page_height
        bottom_ratio = (y1 - page_y0) / page_height
        centered_page_coverage = (
            x0 <= page_x0 + page_width * 0.12 and x1 >= page_x0 + page_width * 0.88
        )
        if (
            width_ratio >= 0.88
            and centered_page_coverage
            and height_ratio >= 0.34
            and bottom_ratio >= 0.94
        ):
            return True
    return False


def _assess_pdf_pages_for_hybrid(
    source: Path,
    password: str | None,
    *,
    dpi: int = 300,
    column_layout: str = "auto",
) -> list[_HybridPageAssessment]:
    """Plan page-local visual islands while keeping reliable body text editable."""

    from docuforge.runner import check_cancelled, report_progress

    normalized_column_layout = _normalize_pdf2docx_column_layout(column_layout)
    _pymupdf, pdf_document = _open_pymupdf_document(source, password)
    assessments: list[_HybridPageAssessment] = []
    page_count = max(1, int(pdf_document.page_count))
    try:
        for page_index, page in enumerate(pdf_document):
            check_cancelled("任务已取消；已完成的文件会保留")
            report_progress(
                0.03 + 0.22 * (page_index / page_count),
                f"分析 PDF 页面 {page_index + 1}/{page_count}",
            )
            reasons: list[str] = []
            try:
                source_text = _repair_known_pdf_text_encoding(
                    page.get_text("text", sort=True) or ""
                )
                page_dict = page.get_text("dict") or {}
                drawings = page.get_drawings() or []
            except Exception:
                assessments.append(
                    _HybridPageAssessment(
                        page_index=page_index,
                        source_text="",
                        editable_source_text="",
                        editable_text_blocks=(),
                        draw_items=0,
                        draw_bbox_max_ratio=0.0,
                        reasons=["页面结构检测失败"],
                        visual_regions=[],
                        detected_two_columns=False,
                    )
                )
                continue
            column_profile = _pdf_standard_two_column_profile(page_dict, page.rect)
            detected_two_columns = _pdf_column_profile_looks_two_column(column_profile)
            column_pair_bands = (
                _pdf_two_column_pair_bands(column_profile)
                if detected_two_columns
                else ()
            )
            asymmetric_column_reason = _pdf_asymmetric_column_risk_reason(
                page_dict,
                page.rect,
            )
            asymmetric_sidebar = (
                _pdf_asymmetric_sidebar_region(
                    page_index,
                    page_dict,
                    drawings,
                    page.rect,
                    dpi=int(dpi),
                )
                if asymmetric_column_reason
                else None
            )
            anchored_visual_regions = (
                (asymmetric_sidebar,) if asymmetric_sidebar is not None else ()
            )
            if asymmetric_column_reason and asymmetric_sidebar is None:
                reasons.append(asymmetric_column_reason)
            page_column_layout = normalized_column_layout
            if normalized_column_layout == "auto" and detected_two_columns:
                page_column_layout = "mixed"

            draw_items = sum(len(drawing.get("items", ())) for drawing in drawings)
            draw_bbox_max_ratio = max(
                (
                    _clipped_bbox_area_ratio(drawing.get("rect", ()), page.rect)
                    for drawing in drawings
                ),
                default=0.0,
            )
            visual_regions = [
                *_detect_pdf_image_regions(
                    page_index,
                    page_dict,
                    page.rect,
                    dpi=int(dpi),
                ),
                *_detect_pdf_vector_regions(
                    page_index,
                    drawings,
                    page.rect,
                    dpi=int(dpi),
                ),
                *_detect_pdf_table_regions(
                    page,
                    page_index,
                    page_dict,
                    drawings,
                    dpi=int(dpi),
                ),
                *_detect_pdf_formula_regions(
                    page_index,
                    page_dict,
                    page.rect,
                    dpi=int(dpi),
                    column_layout=page_column_layout,
                ),
                *_detect_pdf_encoding_regions(
                    page_index,
                    page_dict,
                    page.rect,
                    dpi=int(dpi),
                ),
            ]
            expanded_regions = [
                _expand_region_to_nearby_labels(region, page_dict, page.rect)
                for region in visual_regions
            ]
            expanded_vector_regions = _merge_hybrid_regions(
                (region for region in expanded_regions if region.kind == "vector"),
                page.rect,
                gap=32.0,
            )
            visual_regions = _merge_hybrid_regions(
                [
                    *(region for region in expanded_regions if region.kind != "vector"),
                    *expanded_vector_regions,
                ],
                page.rect,
                gap=1.0,
            )
            visual_regions = _close_hybrid_regions_over_text(
                visual_regions,
                page_dict,
                page.rect,
            )
            visual_regions = _constrain_numbered_formula_regions_to_body_prose(
                visual_regions,
                page_dict,
                page.rect,
            )
            visual_regions = _filter_hybrid_regions_for_anchored_overlays(
                visual_regions,
                anchored_visual_regions,
                page.rect,
            )
            editable_text_blocks = _extract_editable_pdf_text_blocks(
                page_dict,
                page.rect,
                [*visual_regions, *anchored_visual_regions],
            )
            editable_source_text = "\n".join(editable_text_blocks)
            normalized_characters = len(
                _normalize_validation_text(editable_source_text)
            )
            source_characters = len(_normalize_validation_text(source_text))
            region_ratio = _hybrid_region_area_ratio(visual_regions, page.rect)
            if source_characters < _HYBRID_MIN_EDITABLE_CHARACTERS:
                reasons.append("无或极少可编辑文字")
            elif region_ratio >= _HYBRID_FULL_PAGE_REGION_RATIO:
                reasons.append("视觉内容接近整页")
            elif normalized_characters < _HYBRID_MIN_EDITABLE_CHARACTERS:
                reasons.append("视觉区域外无可靠文字")
            elif _has_suspicious_pdf_characters(editable_source_text):
                reasons.append("区域外仍有异常字符")
            elif _hybrid_dense_two_column_figure_risk(
                visual_regions,
                page.rect,
                column_layout=normalized_column_layout,
            ):
                reasons.append("双栏多图密集，已整页保真以避免栏平衡产生大面积空白")
            elif _hybrid_bottom_full_width_visual_overflow_risk(
                visual_regions,
                page.rect,
            ):
                reasons.append("页底全宽复杂内容可能溢页，已整页保真")

            assessments.append(
                _HybridPageAssessment(
                    page_index=page_index,
                    source_text=source_text,
                    editable_source_text=editable_source_text,
                    editable_text_blocks=editable_text_blocks,
                    draw_items=draw_items,
                    draw_bbox_max_ratio=draw_bbox_max_ratio,
                    reasons=reasons,
                    visual_regions=visual_regions,
                    detected_two_columns=detected_two_columns,
                    column_split_x=(
                        column_profile.split_x if detected_two_columns else None
                    ),
                    column_pair_bands=column_pair_bands,
                    anchored_visual_regions=anchored_visual_regions,
                )
            )
    finally:
        pdf_document.close()
    report_progress(0.25, f"完成 {len(assessments)} 页版面分析")
    return assessments


def _extract_docx_text(path: Path) -> str:
    word_namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    tag = lambda name: f"{{{word_namespace}}}{name}"

    def paragraph_text(paragraph: ET.Element) -> str:
        parts: list[str] = []
        for node in paragraph.iter():
            if node.tag == tag("t") and node.text:
                parts.append(node.text)
            elif node.tag == tag("tab"):
                parts.append("\t")
            elif node.tag in {tag("br"), tag("cr")}:
                parts.append("\n")
        return "".join(parts)

    def table_text(table: ET.Element) -> str:
        rows: list[str] = []
        for row in table.findall(tag("tr")):
            cells: list[str] = []
            for cell in row.findall(tag("tc")):
                parts: list[str] = []
                for child in list(cell):
                    if child.tag == tag("p"):
                        parts.append(paragraph_text(child))
                    elif child.tag == tag("tbl"):
                        parts.append(table_text(child))
                cells.append("\n".join(parts))
            rows.append("\t".join(cells))
        return "\n".join(rows)

    try:
        with ZipFile(path) as archive, archive.open("word/document.xml") as stream:
            document_root = ET.parse(stream).getroot()
    except (BadZipFile, KeyError, ET.ParseError, OSError) as exc:
        raise ValidationError(f"无法读取转换后的 Word 文本：{exc}") from exc

    body = document_root.find(tag("body"))
    if body is None:
        return ""
    blocks: list[str] = []
    for child in list(body):
        if child.tag == tag("p"):
            blocks.append(paragraph_text(child))
        elif child.tag == tag("tbl"):
            blocks.append(table_text(child))
    return "\n".join(blocks)


_DOCX_PDF2DOCX_DATE_PREFIX_PATTERN = re.compile(
    r"^[（(]?\s*(?:19|20)\d{2}\s*年"
    r"(?:\s*(?:0?[1-9]|1[0-2])\s*月)?\s*[）)]?\s*[:：]?$"
)
_DOCX_PDF2DOCX_FIELD_PATTERN = re.compile(
    r"\b(?:PAGE|NUMPAGES|PAGEREF|TOC)\b",
    re.IGNORECASE,
)
_DOCX_PDF2DOCX_JUSTIFIED_VALUES = frozenset({"both", "distribute", "thaidistribute"})


def _pdf2docx_docx_layout_quality_reasons(path: Path) -> tuple[str, ...]:
    """Inspect the serialized DOCX for WPS-sensitive layout failures.

    Text-recall checks cannot detect a short Chinese label stretched across a
    justified line or a hidden leading tab.  This second-stage verifier reads
    the final OOXML, resolves paragraph-style inheritance, and rejects only the
    strong structures known to render incorrectly in Word/WPS.
    """

    word_namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    tag = lambda name: f"{{{word_namespace}}}{name}"
    attribute = lambda name: f"{{{word_namespace}}}{name}"

    try:
        with ZipFile(path) as archive:
            document_root = ET.fromstring(archive.read("word/document.xml"))
            settings_root = ET.fromstring(archive.read("word/settings.xml"))
            try:
                styles_root = ET.fromstring(archive.read("word/styles.xml"))
            except KeyError:
                styles_root = None
    except (BadZipFile, KeyError, ET.ParseError, OSError) as exc:
        return (f"无法读取最终 Word 排版结构：{exc}",)

    reasons: list[str] = []
    compatibility = settings_root.find(tag("compat"))
    compatibility_flags = (
        compatibility.findall(tag("doNotExpandShiftReturn"))
        if compatibility is not None
        else []
    )
    compatible_values = {"1", "true", "on"}
    if len(compatibility_flags) != 1 or (
        compatibility_flags
        and compatibility_flags[0].get(attribute("val")) is not None
        and str(compatibility_flags[0].get(attribute("val"))).casefold()
        not in compatible_values
    ):
        reasons.append("WPS 软换行防拉伸兼容设置缺失、无效或重复")

    style_alignment: dict[str, str | None] = {}
    style_parent: dict[str, str | None] = {}
    style_has_tabs: dict[str, bool] = {}
    default_alignment: str | None = None
    default_has_tabs = False
    default_style_id: str | None = None
    if styles_root is not None:
        doc_defaults = styles_root.find(tag("docDefaults"))
        if doc_defaults is not None:
            paragraph_defaults = doc_defaults.find(tag("pPrDefault"))
            default_properties = (
                paragraph_defaults.find(tag("pPr"))
                if paragraph_defaults is not None
                else None
            )
            if default_properties is not None:
                default_justification = default_properties.find(tag("jc"))
                if default_justification is not None:
                    default_alignment = str(
                        default_justification.get(attribute("val")) or ""
                    ).casefold()
                default_tabs = default_properties.find(tag("tabs"))
                default_has_tabs = bool(
                    default_tabs is not None
                    and any(
                        str(item.get(attribute("val")) or "").casefold() != "clear"
                        and item.get(attribute("pos")) is not None
                        for item in default_tabs.findall(tag("tab"))
                    )
                )
        for style in styles_root.findall(tag("style")):
            if str(style.get(attribute("type")) or "paragraph") != "paragraph":
                continue
            style_id = str(style.get(attribute("styleId")) or "")
            if not style_id:
                continue
            if (
                str(style.get(attribute("default")) or "").casefold()
                in compatible_values
            ):
                default_style_id = style_id
            based_on = style.find(tag("basedOn"))
            style_parent[style_id] = (
                str(based_on.get(attribute("val")) or "") or None
                if based_on is not None
                else None
            )
            properties = style.find(tag("pPr"))
            justification = (
                properties.find(tag("jc")) if properties is not None else None
            )
            style_alignment[style_id] = (
                str(justification.get(attribute("val")) or "").casefold()
                if justification is not None
                else None
            )
            tabs = properties.find(tag("tabs")) if properties is not None else None
            style_has_tabs[style_id] = bool(
                tabs is not None
                and any(
                    str(item.get(attribute("val")) or "").casefold() != "clear"
                    and item.get(attribute("pos")) is not None
                    for item in tabs.findall(tag("tab"))
                )
            )

    def resolved_style_value(
        style_id: str | None,
        values: Mapping[str, Any],
        default: Any,
    ) -> Any:
        current = style_id or default_style_id
        seen: set[str] = set()
        while current and current not in seen:
            seen.add(current)
            value = values.get(current)
            if value not in {None, False}:
                return value
            current = style_parent.get(current)
        return default

    excluded_paragraphs: set[int] = set()
    for container_name in ("tbl", "txbxContent"):
        for container in document_root.iter(tag(container_name)):
            excluded_paragraphs.update(
                id(paragraph) for paragraph in container.iter(tag("p"))
            )

    ignored_subtrees = {
        tag("del"),
        tag("moveFrom"),
        tag("txbxContent"),
        tag("drawing"),
        tag("pict"),
        tag("object"),
        tag("oMath"),
        tag("oMathPara"),
    }

    def paragraph_tokens(paragraph: ET.Element) -> list[tuple[str, str]]:
        tokens: list[tuple[str, str]] = []

        def walk(node: ET.Element) -> None:
            if node.tag in ignored_subtrees:
                return
            if node.tag == tag("t") and node.text:
                tokens.append(("text", node.text))
                return
            if node.tag == tag("tab"):
                tokens.append(("tab", "\t"))
                return
            if node.tag in {tag("br"), tag("cr")}:
                break_type = str(node.get(attribute("type")) or "textWrapping")
                if break_type in {"", "textWrapping"}:
                    tokens.append(("break", "\n"))
                return
            for child in list(node):
                walk(child)

        walk(paragraph)
        return tokens

    def prefix_is_short_cjk_or_date(text: str) -> bool:
        normalized = unicodedata.normalize("NFKC", text).strip(" \t\r\n")
        if not normalized:
            return False
        if _DOCX_PDF2DOCX_DATE_PREFIX_PATTERN.fullmatch(normalized):
            return True
        visible = [character for character in normalized if not character.isspace()]
        cjk = sum(
            unicodedata.name(character, "").startswith(
                ("CJK UNIFIED IDEOGRAPH", "CJK COMPATIBILITY IDEOGRAPH")
            )
            for character in visible
        )
        ascii_letters = sum(_is_ascii_letter(character) for character in visible)
        return bool(
            cjk >= 2
            and len(visible) <= 18
            and cjk / max(1, cjk + ascii_letters) >= 0.60
        )

    def suffix_is_english_flow(text: str) -> bool:
        normalized = unicodedata.normalize("NFKC", text)
        visible = sum(not character.isspace() for character in normalized)
        letters = sum(_is_ascii_letter(character) for character in normalized)
        return bool(
            letters >= 12
            and len(_english_words(normalized)) >= 3
            and letters / max(1, visible) >= 0.45
        )

    mixed_justify_count = 0
    mixed_default_tab_count = 0
    embedded_list_count = 0
    for paragraph in document_root.iter(tag("p")):
        if id(paragraph) in excluded_paragraphs:
            continue
        properties = paragraph.find(tag("pPr"))
        if properties is not None and properties.find(tag("framePr")) is not None:
            continue
        field_text = " ".join(
            [str(node.text or "") for node in paragraph.iter(tag("instrText"))]
            + [
                str(node.get(attribute("instr")) or "")
                for node in paragraph.iter(tag("fldSimple"))
            ]
        )
        if _DOCX_PDF2DOCX_FIELD_PATTERN.search(field_text):
            continue
        tokens = paragraph_tokens(paragraph)
        if not tokens or not any(kind == "break" for kind, _value in tokens):
            continue

        segments: list[dict[str, Any]] = [
            {"text": "", "leading_tabs": 0, "visible": False}
        ]
        for kind, value in tokens:
            if kind == "break":
                segments.append({"text": "", "leading_tabs": 0, "visible": False})
                continue
            segment = segments[-1]
            if kind == "tab":
                if not segment["visible"]:
                    segment["leading_tabs"] += 1
                segment["text"] += value
                continue
            segment["text"] += value
            if value.strip():
                segment["visible"] = True

        full_text = "".join(str(segment["text"]) for segment in segments).strip()
        if not full_text or _PDF2DOCX_CLEAR_PAGE_NUMBER_PATTERN.fullmatch(
            unicodedata.normalize("NFKC", full_text)
        ):
            continue
        style_id: str | None = None
        direct_alignment: str | None = None
        direct_has_tabs = False
        if properties is not None:
            paragraph_style = properties.find(tag("pStyle"))
            if paragraph_style is not None:
                style_id = str(paragraph_style.get(attribute("val")) or "") or None
            justification = properties.find(tag("jc"))
            if justification is not None:
                direct_alignment = str(
                    justification.get(attribute("val")) or ""
                ).casefold()
            tabs = properties.find(tag("tabs"))
            direct_has_tabs = bool(
                tabs is not None
                and any(
                    str(item.get(attribute("val")) or "").casefold() != "clear"
                    and item.get(attribute("pos")) is not None
                    for item in tabs.findall(tag("tab"))
                )
            )
        effective_alignment = direct_alignment or resolved_style_value(
            style_id,
            style_alignment,
            default_alignment,
        )
        has_explicit_tabs = direct_has_tabs or bool(
            resolved_style_value(style_id, style_has_tabs, default_has_tabs)
        )

        for segment_index in range(len(segments) - 1):
            prefix = str(segments[segment_index]["text"])
            suffix = "".join(
                str(segment["text"]) for segment in segments[segment_index + 1 :]
            )
            mixed_structure = prefix_is_short_cjk_or_date(
                prefix
            ) and suffix_is_english_flow(suffix)
            if (
                mixed_structure
                and effective_alignment in _DOCX_PDF2DOCX_JUSTIFIED_VALUES
            ):
                mixed_justify_count += 1
            if (
                mixed_structure
                and not has_explicit_tabs
                and (
                    int(segments[segment_index]["leading_tabs"]) >= 2
                    or int(segments[segment_index + 1]["leading_tabs"]) >= 1
                )
            ):
                mixed_default_tab_count += 1

            previous_text = "".join(
                str(segment["text"]) for segment in segments[: segment_index + 1]
            ).strip()
            next_text = str(segments[segment_index + 1]["text"]).strip()
            if len(
                _normalize_validation_text(previous_text)
            ) >= 20 and _pdf2docx_text_starts_list_item(next_text):
                embedded_list_count += 1

    if mixed_justify_count:
        reasons.append(
            f"仍有 {mixed_justify_count} 个短中文/日期与英文共用两端对齐软换行段落"
        )
    if mixed_default_tab_count:
        reasons.append(
            f"仍有 {mixed_default_tab_count} 个中英混合段落使用危险的默认前导 Tab"
        )
    if embedded_list_count:
        reasons.append(f"仍有 {embedded_list_count} 个软换行后的编号项未独立成段")
    return tuple(reasons)


def _rectangles_union_area(
    rectangles: Iterable[tuple[float, float, float, float]],
) -> float:
    """Return exact union area for a small collection of axis-aligned rectangles."""

    normalized = [
        rect for rect in rectangles if rect[2] > rect[0] and rect[3] > rect[1]
    ]
    if not normalized:
        return 0.0
    x_values = sorted({value for rect in normalized for value in (rect[0], rect[2])})
    area = 0.0
    for x0, x1 in zip(x_values, x_values[1:]):
        if x1 <= x0:
            continue
        midpoint = (x0 + x1) / 2.0
        intervals = sorted(
            (rect[1], rect[3]) for rect in normalized if rect[0] <= midpoint < rect[2]
        )
        if not intervals:
            continue
        covered_y = 0.0
        current_y0, current_y1 = intervals[0]
        for y0, y1 in intervals[1:]:
            if y0 <= current_y1:
                current_y1 = max(current_y1, y1)
            else:
                covered_y += current_y1 - current_y0
                current_y0, current_y1 = y0, y1
        covered_y += current_y1 - current_y0
        area += (x1 - x0) * covered_y
    return area


def _rendered_pdf_page_content_assessment(page: Any) -> _RenderedPageContentAssessment:
    """Estimate whether a WPS-rendered page contains useful body content."""

    page_rect = page.rect
    page_width = max(1.0, float(page_rect.width))
    page_height = max(1.0, float(page_rect.height))
    body_rect = (
        float(page_rect.x0) + page_width * 0.04,
        float(page_rect.y0) + page_height * 0.04,
        float(page_rect.x1) - page_width * 0.04,
        float(page_rect.y1) - page_height * 0.055,
    )
    body_area = max(1.0, _rect_area(body_rect))

    def clipped_rect(value: Any) -> tuple[float, float, float, float] | None:
        try:
            raw = tuple(float(item) for item in value)
        except (TypeError, ValueError):
            return None
        if len(raw) != 4:
            return None
        clipped = (
            max(body_rect[0], raw[0]),
            max(body_rect[1], raw[1]),
            min(body_rect[2], raw[2]),
            min(body_rect[3], raw[3]),
        )
        return clipped if clipped[2] > clipped[0] and clipped[3] > clipped[1] else None

    page_dict = page.get_text("dict") or {}
    text_characters: list[str] = []
    font_weights: list[tuple[float, int]] = []
    visual_rects: list[tuple[float, float, float, float]] = []
    for block in page_dict.get("blocks", ()):
        block_type = int(block.get("type", 0) or 0)
        if block_type == 1:
            if image_rect := clipped_rect(block.get("bbox", ())):
                visual_rects.append(image_rect)
            continue
        if block_type != 0:
            continue
        for line in block.get("lines", ()):
            for span in line.get("spans", ()):
                span_rect = clipped_rect(span.get("bbox", ()))
                if span_rect is None:
                    continue
                characters = [
                    character
                    for character in str(span.get("text", "") or "")
                    if character.isalnum()
                ]
                if not characters:
                    continue
                text_characters.extend(characters)
                try:
                    font_size = float(span.get("size", 0.0) or 0.0)
                except (TypeError, ValueError):
                    font_size = 0.0
                if font_size > 0.0:
                    font_weights.append((font_size, len(characters)))

    drawings = page.get_drawings() or ()
    for drawing in drawings:
        drawing_rect = clipped_rect(drawing.get("rect"))
        if drawing_rect is None:
            continue
        # A single unfilled rectangle is commonly just a page border or text
        # frame; treating its bounding box as occupied area would let an almost
        # empty page pass.  Filled shapes and genuinely complex vector paths are
        # visual content.  Ruled tables are added from ``find_tables()`` below.
        drawing_items = drawing.get("items") or ()
        visually_dense = drawing.get("fill") is not None or len(drawing_items) >= 4
        if visually_dense and _rect_area(drawing_rect) / body_area >= 0.003:
            visual_rects.append(drawing_rect)

    if len(drawings) >= 4 and hasattr(page, "find_tables"):
        with warnings.catch_warnings(), contextlib.suppress(Exception):
            warnings.simplefilter("ignore")
            finder = page.find_tables()
            for table in getattr(finder, "tables", ()):
                if table_rect := clipped_rect(getattr(table, "bbox", ())):
                    visual_rects.append(table_rect)

    total_font_weight = sum(weight for _size, weight in font_weights)
    median_font_size = 10.5
    if total_font_weight:
        midpoint = total_font_weight / 2.0
        running_weight = 0
        for font_size, weight in sorted(font_weights):
            running_weight += weight
            if running_weight >= midpoint:
                median_font_size = font_size
                break
    median_font_size = min(36.0, max(6.0, median_font_size))

    visual_ratio = min(1.0, _rectangles_union_area(visual_rects) / body_area)
    cjk_characters = sum(
        any(
            token in unicodedata.name(character, "")
            for token in ("CJK UNIFIED IDEOGRAPH", "CJK COMPATIBILITY IDEOGRAPH")
        )
        for character in text_characters
    )
    cjk_ratio = cjk_characters / max(1, len(text_characters))
    average_character_width = median_font_size * (0.52 + 0.48 * cjk_ratio)
    line_height = median_font_size * 1.35
    estimated_capacity = body_area / max(1.0, average_character_width * line_height)
    remaining_ratio = max(0.0, 1.0 - visual_ratio)
    if visual_ratio >= 0.28:
        required_characters = 0
    elif visual_ratio >= 0.10:
        required_characters = max(
            12,
            round(estimated_capacity * remaining_ratio * 0.022),
        )
    else:
        required_characters = max(48, round(estimated_capacity * 0.035))

    large_title_page = len(text_characters) >= 8 and median_font_size >= 16.0
    substantive = bool(
        visual_ratio >= 0.28
        or len(text_characters) >= required_characters
        or large_title_page
    )
    return _RenderedPageContentAssessment(
        text_characters=len(text_characters),
        median_font_size=median_font_size,
        visual_ratio=visual_ratio,
        required_characters=required_characters,
        substantive=substantive,
    )


def _pdf_page_is_visually_blank(page: Any) -> bool:
    """Recognize an intentionally blank PDF canvas, including a white scan."""

    text = str(page.get_text() or "")
    if any(character.isalnum() for character in text):
        return False
    if not page.get_images(full=True) and not page.get_drawings():
        return True
    try:
        pixmap = page.get_pixmap(dpi=36, alpha=False)
        samples = memoryview(pixmap.samples)
        components = max(1, int(pixmap.n))
        pixels = max(1, int(pixmap.width) * int(pixmap.height))
        ink_pixels = 0
        for offset in range(0, len(samples), components):
            channels = samples[offset : offset + min(3, components)]
            if channels and min(channels) < 245:
                ink_pixels += 1
        return ink_pixels / pixels <= 0.0005
    except Exception:
        return False


def _match_source_and_rendered_blank_pages(
    source_blank_indexes: Iterable[int],
    rendered_blank_indexes: Iterable[int],
    *,
    source_page_count: int,
    rendered_page_count: int,
) -> dict[int, int]:
    """Pair blank pages monotonically while allowing ordinary Word reflow."""

    source_indexes = sorted(set(int(index) for index in source_blank_indexes))
    rendered_indexes = sorted(set(int(index) for index in rendered_blank_indexes))
    if not source_indexes or not rendered_indexes:
        return {}
    tolerance = max(
        2.0,
        float(abs(rendered_page_count - source_page_count) + 2),
        rendered_page_count * 0.04,
    )
    matched: dict[int, int] = {}
    last_rendered_index = -1
    for source_index in source_indexes:
        expected = (
            (source_index + 0.5) / max(1, source_page_count)
        ) * rendered_page_count - 0.5
        candidates = [
            rendered_index
            for rendered_index in rendered_indexes
            if rendered_index > last_rendered_index
            and rendered_index not in matched
            and abs(rendered_index - expected) <= tolerance
        ]
        if not candidates:
            continue
        rendered_index = min(candidates, key=lambda index: abs(index - expected))
        matched[rendered_index] = source_index
        last_rendered_index = rendered_index
    return matched


def _rendered_low_page_is_source_backed_fragment(
    rendered_text: str,
    source_page_texts: Iterable[str],
) -> bool:
    """Allow a short but real reflow fragment that is demonstrably from the PDF."""

    normalized = _normalize_validation_text(rendered_text)
    if len(normalized) < 12:
        return False
    normalized_text = unicodedata.normalize("NFKC", rendered_text)
    if _PDF2DOCX_PAGE_FURNITURE_PATTERN.search(normalized_text):
        return False
    return any(
        _text_sequence_coverage(rendered_text, source_text) >= 0.70
        for source_text in source_page_texts
        if source_text.strip()
    )


def _pdf2docx_rendered_page_counter_conflicts(
    rendered_pdf: Path,
) -> tuple[_RenderedPageCounterConflict, ...]:
    """Locate physical pages containing multiple source-page counters."""

    _pymupdf, document = _open_pymupdf_document(rendered_pdf, None)
    try:
        conflicts: list[_RenderedPageCounterConflict] = []
        for page_index in range(int(document.page_count)):
            text = str(document[page_index].get_text() or "")
            counters = _pdf2docx_conflicting_page_counter_signatures(text)
            if counters:
                conflicts.append(
                    _RenderedPageCounterConflict(
                        page_number=page_index + 1,
                        counters=counters,
                    )
                )
        return tuple(conflicts)
    finally:
        document.close()


def _pdf2docx_rendered_pdf_content_quality_result(
    source_pdf: Path,
    rendered_pdf: Path,
    *,
    expected_pages: int,
    password: str | None,
) -> tuple[str | None, str | None]:
    """Judge real WPS output by page substance instead of exact page count."""

    _pymupdf, source_document = _open_pymupdf_document(source_pdf, password)
    _pymupdf, rendered_document = _open_pymupdf_document(rendered_pdf, None)
    try:
        actual_pages = int(rendered_document.page_count)
        expected_page_count = max(1, int(expected_pages))
        source_page_count = int(source_document.page_count)
        source_page_texts = [
            str(source_document[index].get_text() or "")
            for index in range(source_page_count)
        ]
        source_blank_indexes = [
            index
            for index in range(source_page_count)
            if _pdf_page_is_visually_blank(source_document[index])
        ]
        rendered_assessments: list[_RenderedPageContentAssessment] = []
        rendered_texts: list[str] = []
        rendered_blank_indexes: list[int] = []
        isolated_page_numbers: list[tuple[int, str]] = []
        page_counter_conflicts: list[_RenderedPageCounterConflict] = []
        for page_index in range(actual_pages):
            rendered_page = rendered_document[page_index]
            assessment = _rendered_pdf_page_content_assessment(rendered_page)
            rendered_text = str(rendered_page.get_text() or "")
            rendered_assessments.append(assessment)
            rendered_texts.append(rendered_text)
            conflicting_counters = _pdf2docx_conflicting_page_counter_signatures(
                rendered_text
            )
            if conflicting_counters:
                page_counter_conflicts.append(
                    _RenderedPageCounterConflict(
                        page_number=page_index + 1,
                        counters=conflicting_counters,
                    )
                )
            if _pdf2docx_isolated_page_number_text(rendered_text):
                isolated_page_numbers.append(
                    (page_index + 1, " ".join(rendered_text.split()))
                )
            if not assessment.substantive and _pdf_page_is_visually_blank(
                rendered_page
            ):
                rendered_blank_indexes.append(page_index)

        if page_counter_conflicts:
            details = "、".join(
                f"{conflict.page_number}（"
                + "、".join(
                    f"第{current}页共{total}页" for current, total in conflict.counters
                )
                + "）"
                for conflict in page_counter_conflicts[:8]
            )
            if len(page_counter_conflicts) > 8:
                details += f" 等 {len(page_counter_conflicts)} 页"
            return f"WPS 实际渲染检测到同一页面存在多个源页码：{details}", None

        if isolated_page_numbers:
            details = "、".join(
                f"{page_number}（{text}）"
                for page_number, text in isolated_page_numbers[:8]
            )
            if len(isolated_page_numbers) > 8:
                details += f" 等 {len(isolated_page_numbers)} 页"
            return f"WPS 实际渲染检测到页码或页脚独占页面：{details}", None

        matched_blank_pages = _match_source_and_rendered_blank_pages(
            source_blank_indexes,
            rendered_blank_indexes,
            source_page_count=source_page_count,
            rendered_page_count=actual_pages,
        )
        low_content_pages: list[tuple[int, _RenderedPageContentAssessment]] = []
        source_backed_fragment_pages: list[int] = []
        for page_index, assessment in enumerate(rendered_assessments):
            if assessment.substantive:
                continue
            if page_index in matched_blank_pages:
                continue
            if _rendered_low_page_is_source_backed_fragment(
                rendered_texts[page_index],
                source_page_texts,
            ):
                source_backed_fragment_pages.append(page_index + 1)
                continue
            low_content_pages.append((page_index + 1, assessment))

        if low_content_pages:
            details = []
            for page_number, assessment in low_content_pages[:8]:
                details.append(
                    f"{page_number}（有效字符 {assessment.text_characters}/"
                    f"要求 {assessment.required_characters}，"
                    f"图像/表格占比 {assessment.visual_ratio:.0%}）"
                )
            preview = "、".join(details)
            if len(low_content_pages) > 8:
                preview += f" 等 {len(low_content_pages)} 页"
            return f"WPS 实际渲染出现疑似空白或低内容溢出页：{preview}", None

        # Source PDF pages are fixed canvases, while editable Word pages reflow.
        # A substantial consolidation (for example 56 source pages rendered as
        # 42 content-rich Word pages) is therefore valid.  Keep only a broad
        # catastrophe guard for a structure that has clearly collapsed or
        # exploded; ordinary page-count differences are warnings, not failures.
        maximum_drift = max(4, math.ceil(expected_page_count * 0.35))
        if abs(actual_pages - expected_page_count) > maximum_drift:
            return (
                "WPS 实际渲染页数重排幅度异常"
                f"（源文件 {expected_page_count} 页，Word 渲染 {actual_pages} 页，"
                f"安全差值 {maximum_drift} 页）；虽然各页均有内容，"
                "但整体分页结构可能已经大幅收缩或膨胀",
                None,
            )

        warning_parts: list[str] = []
        if actual_pages != expected_page_count:
            warning_parts.append(
                "WPS 实际渲染页数发生合理重排"
                f"（源文件 {expected_page_count} 页，Word 渲染 {actual_pages} 页）；"
                "所有页面均已按字号、有效字符数及图像/表格占用空间通过内容检测，"
                "已允许保存。"
            )
        if matched_blank_pages:
            blank_mappings = "、".join(
                f"源第 {source_index + 1} 页→Word 第 {rendered_index + 1} 页"
                for rendered_index, source_index in sorted(matched_blank_pages.items())
            )
            warning_parts.append(f"已识别并保留源 PDF 的真实空白页：{blank_mappings}。")
        if source_backed_fragment_pages:
            page_preview = "、".join(
                str(page_number) for page_number in source_backed_fragment_pages[:12]
            )
            if len(source_backed_fragment_pages) > 12:
                page_preview += f" 等 {len(source_backed_fragment_pages)} 页"
            warning_parts.append(
                f"Word 第 {page_preview} 页虽内容较少，但文字已与源 PDF 内容匹配，"
                "判定为有效重排片段并允许保存。"
            )
        return None, "".join(warning_parts) or None
    finally:
        source_document.close()
        rendered_document.close()


def _pdf2docx_rendered_pdf_quality_reason(
    source_pdf: Path,
    rendered_pdf: Path,
    *,
    expected_pages: int,
    password: str | None,
) -> str | None:
    """Backward-compatible reason-only wrapper for rendered PDF validation."""

    reason, _warning = _pdf2docx_rendered_pdf_content_quality_result(
        source_pdf,
        rendered_pdf,
        expected_pages=expected_pages,
        password=password,
    )
    return reason


def _pdf2docx_wps_render_worker_entry(
    connection: Any,
    docx_path: str,
    output_dir: str,
) -> None:
    """Isolate WPS COM so a broken automation server cannot crash the job."""

    try:
        from .wps import convert_with_wps

        rendered = convert_with_wps(
            Path(docx_path),
            Path(output_dir),
            "pdf",
            overwrite=True,
        )[0]
        connection.send({"ok": True, "path": str(rendered)})
    except Exception as exc:  # pragma: no cover - defensive process boundary
        with contextlib.suppress(Exception):
            connection.send({"ok": False, "error": str(exc)})
    finally:
        with contextlib.suppress(Exception):
            connection.close()


def _pdf2docx_wps_render_quality_result(
    source_pdf: Path,
    docx_path: Path,
    *,
    expected_pages: int,
    password: str | None,
    progress: Callable[[float, str], None] | None = None,
) -> tuple[str | None, str | None]:
    """Render a longer final DOCX through WPS and verify physical pagination.

    The OOXML pass catches structural faults, but only the real layout engine
    can prove that a footer or section carrier did not spill onto an extra page.
    WPS is intentionally the first-choice renderer.  Its COM call runs in a
    spawned process because a broken automation server must not destabilize the
    main conversion worker.  Short documents keep the fast structural path.
    """

    if int(expected_pages) < _PDF2DOCX_WPS_RENDER_MIN_PAGES:
        return None, None
    try:
        from .wps import detect_wps_engines

        status = detect_wps_engines()["writer"]
    except Exception as exc:
        return None, f"WPS 实际分页复检未能初始化：{exc}"
    if not status.available:
        return None, None

    import multiprocessing
    import time

    from docuforge.runner import check_cancelled

    try:
        with tempfile.TemporaryDirectory(prefix="docuforge-wps-layout-check-") as name:
            context = multiprocessing.get_context("spawn")

            def emit_progress(ratio: float, message: str) -> None:
                if progress is None:
                    return
                with contextlib.suppress(Exception):
                    progress(max(0.0, min(0.995, float(ratio))), message)

            def render_once(attempt: int) -> tuple[Path | None, str | None]:
                parent_connection, child_connection = context.Pipe(duplex=False)
                process = context.Process(
                    target=_pdf2docx_wps_render_worker_entry,
                    args=(child_connection, str(docx_path), name),
                    name="docuforge-wps-layout-check",
                    daemon=False,
                )
                try:
                    attempt_ranges = ((0.0, 0.84), (0.84, 0.94), (0.94, 0.99))
                    attempt_start, attempt_end = attempt_ranges[
                        max(0, min(len(attempt_ranges) - 1, int(attempt)))
                    ]
                    emit_progress(
                        attempt_start,
                        "WPS 正在启动最终版面复检"
                        if attempt == 0
                        else f"WPS 正在进行第 {attempt + 1} 次分页复检",
                    )
                    process.start()
                    child_connection.close()
                    started_at = time.monotonic()
                    estimated_wait = max(
                        18.0,
                        min(120.0, 14.0 + float(expected_pages) * 0.9),
                    )
                    last_progress_at = started_at - 1.0
                    deadline = time.monotonic() + max(
                        90.0,
                        min(240.0, float(expected_pages) * 4.0),
                    )
                    result: Mapping[str, Any] | None = None
                    while result is None:
                        check_cancelled("任务已取消；正在终止 WPS 分页复检")
                        now = time.monotonic()
                        if now - last_progress_at >= 0.75:
                            elapsed = max(0.0, now - started_at)
                            local_ratio = min(0.96, elapsed / estimated_wait)
                            phase_ratio = attempt_start + (
                                attempt_end - attempt_start
                            ) * local_ratio
                            if elapsed < estimated_wait * 0.45:
                                message = "WPS 正在渲染最终 Word"
                            elif elapsed < estimated_wait * 0.82:
                                message = "WPS 正在核对实际分页"
                            else:
                                message = "正在检查空白页、文字与字体结果"
                            emit_progress(phase_ratio, message)
                            last_progress_at = now
                        remaining = deadline - time.monotonic()
                        if remaining <= 0:
                            return None, "WPS 实际分页复检超时，已跳过该项检查"
                        if parent_connection.poll(min(0.1, remaining)):
                            try:
                                message = parent_connection.recv()
                            except (EOFError, OSError):
                                message = None
                            if isinstance(message, Mapping):
                                result = message
                            else:
                                break
                        elif not process.is_alive():
                            break

                    process.join(2.0)
                    if not result or not result.get("ok"):
                        detail = (
                            str(result.get("error") or "子进程未返回结果")
                            if result
                            else f"子进程退出码 {process.exitcode}"
                        )
                        return None, f"WPS 实际分页复检未完成：{detail}"
                    rendered_path = Path(str(result.get("path") or ""))
                    if not rendered_path.is_file():
                        return None, "WPS 实际分页复检未生成有效 PDF"
                    emit_progress(attempt_end, "WPS 已完成本轮版面渲染")
                    return rendered_path, None
                finally:
                    with contextlib.suppress(Exception):
                        parent_connection.close()
                    with contextlib.suppress(Exception):
                        child_connection.close()
                    if process.is_alive():
                        process.terminate()
                        process.join(2.0)

            # WPS is also the authoritative repair oracle: if Word reflow has
            # consolidated two source-page counters onto one physical page,
            # remove only the uniquely identifiable older framed counter and
            # render again.  Two bounded repair rounds prevent oscillation.
            for attempt in range(3):
                rendered_path, render_error = render_once(attempt)
                if render_error:
                    return None, render_error
                assert rendered_path is not None
                conflicts = _pdf2docx_rendered_page_counter_conflicts(rendered_path)
                if conflicts and attempt < 2:
                    repaired = _repair_pdf2docx_rendered_page_counter_conflicts(
                        docx_path,
                        conflicts,
                    )
                    if repaired:
                        continue
                return _pdf2docx_rendered_pdf_content_quality_result(
                    source_pdf,
                    rendered_path,
                    expected_pages=expected_pages,
                    password=password,
                )
            return None, "WPS 实际分页复检未完成：页码冲突自动修复未收敛"
    except Exception as exc:
        return None, f"WPS 实际分页复检未完成：{exc}"


def _hybrid_clip_image_res_ratio(
    assessments: Iterable[_HybridPageAssessment],
    requested_dpi: int,
) -> float:
    visual_dpis = [
        max(72, int(region.dpi or requested_dpi))
        for assessment in assessments
        for region in assessment.visual_regions
    ]
    maximum_dpi = max([int(requested_dpi), *visual_dpis])
    return max(2.5, maximum_dpi / 72.0)


def _normalize_pdf2docx_column_layout(column_layout: str) -> str:
    normalized = str(column_layout or "auto").lower().strip()
    if normalized not in _PDF2DOCX_COLUMN_LAYOUTS:
        raise ValidationError("PDF 分栏结构必须是 auto、single、double 或 mixed")
    return normalized


def _pdf2docx_settings(
    converter: Any,
    *,
    resilient: bool = False,
    parse_lattice_table: bool = True,
    clip_image_res_ratio: float | None = None,
    column_layout: str = "auto",
) -> dict[str, Any]:
    normalized_column_layout = _normalize_pdf2docx_column_layout(column_layout)
    min_section_height = (
        _PDF2DOCX_COLUMN_AWARE_MIN_SECTION_HEIGHT
        if normalized_column_layout in {"double", "mixed"}
        else _PDF2DOCX_DEFAULT_MIN_SECTION_HEIGHT
    )
    settings = dict(converter.default_settings)
    settings.update(
        {
            "min_section_height": min_section_height,
            "multi_processing": False,
            "ignore_page_error": resilient,
            "parse_lattice_table": bool(parse_lattice_table),
            "parse_stream_table": False,
            "clip_image_res_ratio": max(
                2.5,
                float(clip_image_res_ratio or 0.0),
            ),
            "delete_end_line_hyphen": False,
            "raw_exceptions": not resilient,
        }
    )
    return settings


_PDF2DOCX_LIST_PREFIX_PATTERN = re.compile(
    r"^\s*(?:"
    r"[A-HＡ-Ｈ][)）.．、]|"
    r"\(?[0-9０-９]{1,4}\)?[)）.．、]|"
    r"[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]|"
    r"[•●○▪■◆◇]\s*"
    r")"
)
_PDF2DOCX_OPTION_PREFIX_PATTERN = re.compile(r"^\s*[A-HＡ-Ｈ][)）.．、]")
_PDF2DOCX_SENTENCE_END_CHARACTERS = frozenset(".!?。！？．")


def _pdf2docx_block_rows(block: Any) -> list[Any]:
    try:
        return list(block.lines.group_by_physical_rows(sorted=True))
    except (AttributeError, TypeError):
        return [line for line in getattr(block, "lines", ())]


def _pdf2docx_row_text(row: Any) -> str:
    return "".join(
        str(getattr(span, "text", "") or "")
        for line in row
        for span in getattr(line, "spans", ())
    )


def _pdf2docx_row_bbox(row: Any) -> tuple[float, float, float, float]:
    bbox = row.bbox
    return tuple(float(bbox[index]) for index in range(4))  # type: ignore[return-value]


def _pdf2docx_line_reference_y(line: Any) -> float:
    spans = [
        span
        for span in getattr(line, "spans", ())
        if _is_pdf2docx_text_span(span) and _pdf2docx_span_text(span).strip()
    ]
    if spans:
        primary = max(
            spans,
            key=lambda span: (
                sum(
                    character.isascii() and character.isalpha()
                    for character in _pdf2docx_span_text(span)
                ),
                len(_pdf2docx_span_text(span)),
            ),
        )
        origins = [
            float(character.origin[1])
            for character in list(getattr(primary, "chars", ()) or ())
            if getattr(character, "origin", None)
        ]
        if origins:
            return float(median(origins))
        return (
            _pdf2docx_bbox_value(primary, "y0", 1)
            + _pdf2docx_bbox_value(primary, "y1", 3)
        ) / 2.0
    bbox = line.bbox
    return (float(bbox[1]) + float(bbox[3])) / 2.0


def _pdf2docx_row_reference_y(row: Any) -> float:
    lines = list(row)
    return float(median(_pdf2docx_line_reference_y(line) for line in lines))


def _pdf2docx_block_font_size(block: Any) -> float:
    sizes: Counter[float] = Counter()
    for line in getattr(block, "lines", ()):
        for span in getattr(line, "spans", ()):
            if not _is_pdf2docx_text_span(span):
                continue
            text = _pdf2docx_span_text(span).strip()
            size = float(getattr(span, "size", 0.0) or 0.0)
            if text and size > 0:
                sizes[round(size, 2)] += max(1, len(text))
    if not sizes:
        return 0.0
    return max(sizes.items(), key=lambda item: (item[1], item[0]))[0]


def _pdf2docx_block_is_english_prose(block: Any) -> bool:
    if not bool(getattr(block, "is_text_block", False)):
        return False
    for line in getattr(block, "lines", ()):
        if getattr(line, "image_spans", ()):
            return False
    text = unicodedata.normalize("NFKC", str(getattr(block, "text", "") or ""))
    if "<image>" in text:
        return False
    letters = sum(character.isascii() and character.isalpha() for character in text)
    visible = sum(not character.isspace() for character in text)
    size = _pdf2docx_block_font_size(block)
    return letters >= 12 and letters / max(1, visible) >= 0.45 and 7.0 <= size <= 15.5


def _pdf2docx_block_is_short_english_flow(block: Any) -> bool:
    if not bool(getattr(block, "is_text_block", False)):
        return False
    for line in getattr(block, "lines", ()):
        if getattr(line, "image_spans", ()):
            return False
    text = unicodedata.normalize("NFKC", str(getattr(block, "text", "") or ""))
    letters = sum(character.isascii() and character.isalpha() for character in text)
    visible = sum(not character.isspace() for character in text)
    size = _pdf2docx_block_font_size(block)
    return letters >= 3 and letters / max(1, visible) >= 0.45 and 7.0 <= size <= 15.5


def _pdf2docx_text_starts_list_item(text: str) -> bool:
    return _pdf2docx_list_prefix_match(text) is not None


def _pdf2docx_list_prefix_match(text: str) -> re.Match[str] | None:
    """Return a real list marker while rejecting decimal-number lookalikes."""

    original = str(text or "")
    normalized = unicodedata.normalize("NFKC", original)
    match = _PDF2DOCX_LIST_PREFIX_PATTERN.match(original)
    candidate = original
    if match is None:
        match = _PDF2DOCX_LIST_PREFIX_PATTERN.match(normalized)
        candidate = normalized
    if match is None:
        return None

    remainder = candidate[match.end() :]
    if not remainder.strip():
        return None
    matched = unicodedata.normalize("NFKC", candidate[: match.end()])
    # ``5.4 million`` and ``19.39%`` are decimal values, not numbered rows.
    # A genuine compact marker such as ``1.For`` remains valid because the
    # first character after the punctuation is not another digit.
    if matched.rstrip().endswith(".") and remainder[:1].isdigit():
        return None
    return match


def _pdf2docx_text_starts_option(text: str) -> bool:
    return (
        _PDF2DOCX_OPTION_PREFIX_PATTERN.match(unicodedata.normalize("NFKC", text))
        is not None
    )


def _pdf2docx_text_has_cjk_ideograph(text: str) -> bool:
    return any(
        unicodedata.name(character, "").startswith(
            ("CJK UNIFIED IDEOGRAPH", "CJK COMPATIBILITY IDEOGRAPH")
        )
        for character in str(text or "")
    )


def _clear_pdf2docx_block_tabs(block: Any) -> int:
    """Clear both block-level stops and line-level tab requests.

    ``TextBlock.tab_stops`` alone is not sufficient: ``Line.make_docx()`` can
    still emit ``<w:tab/>`` from ``line.tab_stop``.  Those hidden line flags are
    the source of several large title/body indent shifts in WPS.
    """

    changed = 0
    if list(getattr(block, "tab_stops", ()) or ()):
        block.tab_stops = []
        changed += 1
    for line in list(getattr(block, "lines", ()) or ()):
        if int(getattr(line, "tab_stop", 0) or 0):
            line.tab_stop = 0
            changed += 1
    return changed


def _pdf2docx_text_ends_sentence(text: str) -> bool:
    normalized = unicodedata.normalize("NFKC", str(text or "")).rstrip()
    normalized = normalized.rstrip("'\"’”)]}）】》")
    return bool(normalized and normalized[-1] in _PDF2DOCX_SENTENCE_END_CHARACTERS)


def _pdf2docx_column_line_pitch(column: Any) -> float | None:
    references: list[tuple[float, float]] = []
    for block in column.blocks:
        if not _pdf2docx_block_is_english_prose(block):
            continue
        size = _pdf2docx_block_font_size(block)
        for row in _pdf2docx_block_rows(block):
            references.append((_pdf2docx_row_reference_y(row), size))
    references.sort()
    steps: list[float] = []
    for (previous_y, previous_size), (current_y, current_size) in zip(
        references, references[1:]
    ):
        size = min(previous_size, current_size)
        step = current_y - previous_y
        if size > 0 and size * 1.15 <= step <= size * 2.8:
            steps.append(step)
    return float(median(steps)) if steps else None


def _split_pdf2docx_structural_break_blocks(column: Any) -> int:
    """Detach short CJK labels/dates from following English flow.

    Teaching materials frequently store a short label such as ``第二段`` or
    ``2025 年 12 月`` and the following English paragraph in one PDF text block,
    separated only by a native line break.  When that combined block is mapped
    to a justified Word paragraph, Word/WPS expands the short line across the
    full measure and also treats its source X offset as the prose first-line
    indent.  Split only strong, geometry-backed structural boundaries so each
    side can receive its own alignment and indentation.
    """

    try:
        from pdf2docx.common.share import TextAlignment
        from pdf2docx.text.TextBlock import TextBlock
    except ImportError as exc:
        raise MissingEngineError("PDF 转 Word 的结构段落重建需要 pdf2docx") from exc

    column_bbox = column.working_bbox
    column_width = max(1.0, float(column_bbox[2]) - float(column_bbox[0]))
    changed = 0
    rebuilt: list[Any] = []
    for block in column.blocks:
        if not bool(getattr(block, "is_text_block", False)):
            rebuilt.append(block)
            continue
        rows = _pdf2docx_block_rows(block)
        if len(rows) < 2:
            rebuilt.append(block)
            continue

        structural_indexes: list[int] = []
        for row_index, (row, following_row) in enumerate(zip(rows, rows[1:])):
            row_lines = list(row)
            if not row_lines or not getattr(row_lines[-1], "line_break", 0):
                continue
            raw_text = _pdf2docx_row_text(row)
            text = unicodedata.normalize("NFKC", raw_text).strip()
            following_text = unicodedata.normalize(
                "NFKC", _pdf2docx_row_text(following_row)
            ).strip()
            visible = sum(not character.isspace() for character in text)
            ascii_letters = sum(
                character.isascii() and character.isalpha() for character in text
            )
            following_visible = sum(
                not character.isspace() for character in following_text
            )
            following_letters = sum(
                character.isascii() and character.isalpha()
                for character in following_text
            )
            row_bbox = _pdf2docx_row_bbox(row)
            width_ratio = (row_bbox[2] - row_bbox[0]) / column_width
            if not (
                text
                and _pdf2docx_text_has_cjk_ideograph(text)
                and visible <= 24
                and ascii_letters <= 4
                and width_ratio <= 0.55
                and following_letters >= 6
                and following_letters / max(1, following_visible) >= 0.35
            ):
                continue
            structural_indexes.append(row_index)

        if not structural_indexes:
            rebuilt.append(block)
            continue

        boundaries = {0, len(rows)}
        for row_index in structural_indexes:
            boundaries.update((row_index, row_index + 1))
        ordered_boundaries = sorted(boundaries)
        groups = [
            (start, rows[start:end])
            for start, end in zip(ordered_boundaries, ordered_boundaries[1:])
            if end > start
        ]
        structural_index_set = set(structural_indexes)
        for group_index, (start, grouped_rows) in enumerate(groups):
            replacement = TextBlock()
            for grouped_row in grouped_rows:
                replacement.add(list(grouped_row))
            replacement.alignment = block.alignment
            replacement.left_space = block.left_space
            replacement.right_space = block.right_space
            replacement.first_line_space = block.first_line_space
            replacement.line_space = block.line_space
            replacement.line_space_type = block.line_space_type
            replacement.before_space = block.before_space if group_index == 0 else 0.0
            replacement.after_space = (
                block.after_space if group_index == len(groups) - 1 else 0.0
            )
            _clear_pdf2docx_block_tabs(replacement)

            if start in structural_index_set and len(grouped_rows) == 1:
                row_bbox = _pdf2docx_row_bbox(grouped_rows[0])
                left_gap = max(0.0, row_bbox[0] - float(column_bbox[0]))
                right_gap = max(0.0, float(column_bbox[2]) - row_bbox[2])
                row_width = max(0.0, row_bbox[2] - row_bbox[0])
                font_size = max(1.0, _pdf2docx_block_font_size(replacement))
                centered = bool(
                    row_width / column_width <= 0.72
                    and abs(left_gap - right_gap)
                    <= max(font_size * 1.5, column_width * 0.06)
                )
                right_aligned = bool(
                    not centered
                    and left_gap >= column_width * 0.20
                    and right_gap <= max(font_size, left_gap * 0.25)
                )
                replacement.alignment = (
                    TextAlignment.CENTER
                    if centered
                    else TextAlignment.RIGHT if right_aligned else TextAlignment.LEFT
                )
                replacement.left_space = 0.0 if centered else left_gap
                replacement.right_space = right_gap if right_aligned else 0.0
                replacement.first_line_space = 0.0
                replacement._docuforge_structural_heading = True
            if getattr(block, "_docuforge_list_item", False) and (
                _pdf2docx_text_starts_list_item(
                    str(getattr(replacement, "text", "") or "")
                )
            ):
                replacement._docuforge_list_item = True
            rebuilt.append(replacement)
        changed += len(groups) - 1

    if changed:
        column.blocks.reset(rebuilt)
    return changed


def _split_pdf2docx_list_blocks(column: Any) -> int:
    """Split A-D / numbered physical rows that pdf2docx merged into one paragraph."""

    try:
        from pdf2docx.text.TextBlock import TextBlock
    except ImportError as exc:
        raise MissingEngineError("PDF 转 Word 的列表重建需要 pdf2docx") from exc

    changed = 0
    rebuilt: list[Any] = []
    for block in column.blocks:
        if not bool(getattr(block, "is_text_block", False)):
            rebuilt.append(block)
            continue
        rows = _pdf2docx_block_rows(block)
        marker_indexes = [
            index
            for index, row in enumerate(rows)
            if _pdf2docx_text_starts_list_item(_pdf2docx_row_text(row))
        ]
        if not marker_indexes:
            rebuilt.append(block)
            continue

        if len(marker_indexes) == 1:
            marker_index = marker_indexes[0]
            if marker_index <= 0:
                rebuilt.append(block)
                continue
            previous_row = rows[marker_index - 1]
            marker_row = rows[marker_index]
            previous_bbox = _pdf2docx_row_bbox(previous_row)
            marker_bbox = _pdf2docx_row_bbox(marker_row)
            font_size = max(1.0, _pdf2docx_block_font_size(block))
            row_step = _pdf2docx_row_reference_y(
                marker_row
            ) - _pdf2docx_row_reference_y(previous_row)
            column_bbox = column.working_bbox
            column_width = max(
                1.0,
                float(column_bbox[2]) - float(column_bbox[0]),
            )
            marker_indent = marker_bbox[0] - float(column_bbox[0])
            # A lone marker inside a larger block is split only when it begins
            # a new nearby physical row at a plausible list anchor.  This
            # prevents decimal measurements or distant floating labels from
            # becoming false paragraph boundaries.
            if not (
                max(2.0, font_size * 0.70) <= row_step <= max(42.0, font_size * 3.8)
                and -font_size <= marker_indent <= max(72.0, column_width * 0.30)
                and marker_bbox[2] > marker_bbox[0]
                and previous_bbox[2] > previous_bbox[0]
            ):
                rebuilt.append(block)
                continue

        boundaries = sorted({0, *marker_indexes, len(rows)})
        groups = [
            rows[start:end]
            for start, end in zip(boundaries, boundaries[1:])
            if end > start
        ]
        if len(groups) < 2:
            rebuilt.append(block)
            continue

        for group_index, grouped_rows in enumerate(groups):
            replacement = TextBlock()
            for row in grouped_rows:
                replacement.add(list(row))
            replacement.alignment = block.alignment
            replacement.left_space = block.left_space
            replacement.right_space = block.right_space
            replacement.first_line_space = block.first_line_space
            replacement.line_space = block.line_space
            replacement.line_space_type = block.line_space_type
            replacement.before_space = block.before_space if group_index == 0 else 0.0
            replacement.after_space = (
                block.after_space if group_index == len(groups) - 1 else 0.0
            )
            _clear_pdf2docx_block_tabs(replacement)
            if _pdf2docx_text_starts_list_item(_pdf2docx_row_text(grouped_rows[0])):
                replacement._docuforge_list_item = True
            rebuilt.append(replacement)
        changed += len(groups) - 1

    if changed:
        column.blocks.reset(rebuilt)
    return changed


def _should_merge_pdf2docx_prose_blocks(
    previous: Any,
    current: Any,
    column: Any,
    line_pitch: float | None,
) -> bool:
    if not (
        _pdf2docx_block_is_english_prose(previous)
        and (
            _pdf2docx_block_is_english_prose(current)
            or _pdf2docx_block_is_short_english_flow(current)
        )
    ):
        return False
    if getattr(current, "_docuforge_list_item", False):
        return False
    current_text = str(getattr(current, "text", "") or "")
    if _pdf2docx_text_starts_list_item(current_text):
        return False
    previous_text = str(getattr(previous, "text", "") or "")
    if getattr(
        previous, "_docuforge_list_item", False
    ) and _pdf2docx_text_ends_sentence(previous_text):
        return False

    previous_size = _pdf2docx_block_font_size(previous)
    current_size = _pdf2docx_block_font_size(current)
    size = min(previous_size, current_size)
    if size <= 0 or abs(previous_size - current_size) > max(0.8, size * 0.1):
        return False

    previous_rows = _pdf2docx_block_rows(previous)
    current_rows = _pdf2docx_block_rows(current)
    if not previous_rows or not current_rows:
        return False
    previous_last = previous_rows[-1]
    current_first = current_rows[0]
    step = _pdf2docx_row_reference_y(current_first) - _pdf2docx_row_reference_y(
        previous_last
    )
    expected_pitch = line_pitch or size * 1.7
    if not (
        expected_pitch * 0.72 <= step <= expected_pitch * 1.32
        or abs(step - expected_pitch) <= max(2.5, size * 0.35)
    ):
        return False

    column_bbox = column.working_bbox
    column_width = max(1.0, float(column_bbox[2]) - float(column_bbox[0]))
    previous_last_bbox = _pdf2docx_row_bbox(previous_last)
    current_first_bbox = _pdf2docx_row_bbox(current_first)
    previous_width_ratio = (
        previous_last_bbox[2] - previous_last_bbox[0]
    ) / column_width
    if len(previous_rows) == 1 and previous_width_ratio < 0.72:
        return False

    if _pdf2docx_text_ends_sentence(previous_text):
        right_gap = float(column_bbox[2]) - previous_last_bbox[2]
        current_indent = current_first_bbox[0] - float(column_bbox[0])
        if right_gap > max(size * 1.1, column_width * 0.06):
            return False
        if current_indent > size * 0.85:
            return False
    return True


def _reconstruct_pdf2docx_prose_blocks(pages: Iterable[Any]) -> int:
    """Rejoin prose split by tall punctuation bboxes before Word generation."""

    changed = 0
    for page in pages:
        if not getattr(page, "finalized", False):
            continue
        for section in page.sections:
            for column in section:
                changed += _split_pdf2docx_structural_break_blocks(column)
                changed += _split_pdf2docx_list_blocks(column)
                line_pitch = _pdf2docx_column_line_pitch(column)
                rebuilt: list[Any] = []
                for block in column.blocks:
                    previous = rebuilt[-1] if rebuilt else None
                    if previous is None or not _should_merge_pdf2docx_prose_blocks(
                        previous, block, column, line_pitch
                    ):
                        rebuilt.append(block)
                        continue
                    previous.add(list(block.lines))
                    previous.after_space = block.after_space
                    _clear_pdf2docx_block_tabs(previous)
                    previous._docuforge_reconstructed = True
                    changed += 1
                if len(rebuilt) != len(column.blocks):
                    column.blocks.reset(rebuilt)
                for block in rebuilt:
                    if not getattr(block, "_docuforge_reconstructed", False):
                        continue
                    rows = _pdf2docx_block_rows(block)
                    reference_steps = [
                        _pdf2docx_row_reference_y(current)
                        - _pdf2docx_row_reference_y(previous)
                        for previous, current in zip(rows, rows[1:])
                    ]
                    positive_steps = [step for step in reference_steps if step > 0]
                    if positive_steps:
                        block.line_space_type = 0
                        block.line_space = float(median(positive_steps))
    return changed


def _normalize_pdf2docx_list_prefixes(pages: Iterable[Any]) -> int:
    """Use portable half-width list markers and guarantee marker/body spacing."""

    changed = 0
    for page in pages:
        if not getattr(page, "finalized", False):
            continue
        for section in page.sections:
            for column in section:
                for block in _iter_pdf2docx_text_blocks(column.blocks):
                    lines = list(getattr(block, "lines", ()) or ())
                    if not lines:
                        continue
                    spans = [
                        span for span in lines[0].spans if _is_pdf2docx_text_span(span)
                    ]
                    original = "".join(_pdf2docx_span_text(span) for span in spans)
                    match = _pdf2docx_list_prefix_match(original)
                    if match is None:
                        continue
                    prefix_end = match.end()
                    normalized_prefix = "".join(
                        (
                            unicodedata.normalize("NFKC", character)
                            if "\uff01" <= character <= "\uff5e"
                            else character
                        )
                        for character in original[:prefix_end]
                    )
                    corrected = normalized_prefix + original[prefix_end:]
                    if (
                        prefix_end < len(corrected)
                        and not corrected[prefix_end].isspace()
                    ):
                        corrected = (
                            corrected[:prefix_end] + " " + corrected[prefix_end:]
                        )
                    if corrected == original:
                        continue
                    offset = 0
                    remaining_insert = len(corrected) - len(original)
                    for span in spans:
                        span_text = _pdf2docx_span_text(span)
                        end = offset + len(span_text)
                        take = len(span_text)
                        if remaining_insert and offset <= prefix_end <= end:
                            take += remaining_insert
                            remaining_insert = 0
                        replacement = corrected[offset : offset + take]
                        if replacement != span_text:
                            _set_pdf2docx_span_text(span, replacement)
                            changed += 1
                        offset += take
    return changed


def _normalize_pdf2docx_english_widths(pages: Iterable[Any]) -> int:
    """Use portable half-width ASCII glyphs inside English-dominant paragraphs.

    A number of teaching-material PDFs encode ordinary Latin punctuation and
    digits as U+FF01-U+FF5E.  Their embedded font draws those glyphs narrowly,
    but Word/WPS substitutes a normal CJK font and keeps the full-em advance.
    The resulting punctuation appears detached from the preceding word and
    dates such as ``２０２５`` become visibly over-spaced.  Limit the conversion
    to English prose/list blocks so genuine full-width typography in Chinese
    content and running page numbers remains untouched.
    """

    changed = 0
    for page in pages:
        if not getattr(page, "finalized", False):
            continue
        for section in page.sections:
            for column in section:
                for block in _iter_pdf2docx_text_blocks(column.blocks):
                    text = str(getattr(block, "text", "") or "")
                    if not (
                        _pdf2docx_block_is_english_prose(block)
                        or _pdf2docx_block_is_short_english_flow(block)
                    ):
                        continue
                    for line in getattr(block, "lines", ()):
                        for span in getattr(line, "spans", ()):
                            if not _is_pdf2docx_text_span(span):
                                continue
                            original = _pdf2docx_span_text(span)
                            corrected = "".join(
                                (
                                    unicodedata.normalize("NFKC", character)
                                    if "\uff01" <= character <= "\uff5e"
                                    else character
                                )
                                for character in original
                            )
                            if corrected == original:
                                continue
                            _set_pdf2docx_span_text(span, corrected)
                            changed += 1
    return changed


def _normalize_pdf2docx_prose_alignment(pages: Iterable[Any]) -> int:
    """Correct false alignment while allowing Word/WPS to reflow prose safely."""

    try:
        from pdf2docx.common.share import TextAlignment
    except ImportError as exc:
        raise MissingEngineError("PDF 转 Word 的段落重建需要 pdf2docx") from exc

    changed = 0
    for page in pages:
        if not getattr(page, "finalized", False):
            continue
        for section in page.sections:
            for column in section:
                # A soft break after the final physical line is redundant with
                # the paragraph mark itself.  pdf2docx can retain that trailing
                # break on short headings and answer choices, which adds an
                # invisible blank line and may push only a footer onto a new
                # page.  Internal breaks in non-English structured content are
                # left intact; only the terminal one is universally safe to drop.
                for text_block in _iter_pdf2docx_text_blocks(column.blocks):
                    terminal_lines = list(getattr(text_block, "lines", ()) or ())
                    if terminal_lines and getattr(terminal_lines[-1], "line_break", 0):
                        terminal_lines[-1].line_break = 0
                        changed += 1
                column_bbox = column.working_bbox
                column_width = max(1.0, float(column_bbox[2]) - float(column_bbox[0]))
                list_groups: list[list[Any]] = []
                active_list_group: list[Any] = []
                for block in column.blocks:
                    text = str(getattr(block, "text", "") or "")
                    is_list = bool(
                        getattr(block, "_docuforge_list_item", False)
                        or _pdf2docx_text_starts_list_item(text)
                    )
                    rows = _pdf2docx_block_rows(block)
                    if not rows:
                        continue
                    size = _pdf2docx_block_font_size(block)
                    row_boxes = [_pdf2docx_row_bbox(row) for row in rows]

                    # pdf2docx occasionally labels a geometrically centered
                    # one-line title/date as LEFT or RIGHT and carries source
                    # tab requests into Word.  Recover alignment from the PDF
                    # coordinates before the English-prose-only branch below.
                    if len(rows) == 1 and not is_list:
                        row_width = row_boxes[0][2] - row_boxes[0][0]
                        row_center = (row_boxes[0][0] + row_boxes[0][2]) / 2.0
                        column_center = (
                            float(column_bbox[0]) + float(column_bbox[2])
                        ) / 2.0
                        geometrically_centered = bool(
                            row_width / column_width <= 0.75
                            and abs(row_center - column_center)
                            <= max(size, column_width * 0.025)
                        )
                        if geometrically_centered:
                            for attribute, desired in (
                                ("alignment", TextAlignment.CENTER),
                                ("left_space", 0.0),
                                ("right_space", 0.0),
                                ("first_line_space", 0.0),
                            ):
                                current = getattr(block, attribute)
                                different = (
                                    current != desired
                                    if attribute == "alignment"
                                    else abs(float(current) - float(desired)) > 0.05
                                )
                                if different:
                                    setattr(block, attribute, desired)
                                    changed += 1
                            changed += _clear_pdf2docx_block_tabs(block)
                            if active_list_group:
                                list_groups.append(active_list_group)
                                active_list_group = []
                            continue

                    if not (_pdf2docx_block_is_english_prose(block) or is_list):
                        if active_list_group:
                            list_groups.append(active_list_group)
                            active_list_group = []
                        if getattr(block, "_docuforge_structural_heading", False):
                            changed += _clear_pdf2docx_block_tabs(block)
                        continue
                    if is_list:
                        active_list_group.append(block)
                    elif active_list_group:
                        list_groups.append(active_list_group)
                        active_list_group = []

                    if len(rows) == 1:
                        width_ratio = (row_boxes[0][2] - row_boxes[0][0]) / column_width
                        if width_ratio >= 0.72 or is_list:
                            if block.alignment != TextAlignment.LEFT:
                                block.alignment = TextAlignment.LEFT
                                changed += 1
                            desired_left = max(
                                0.0, row_boxes[0][0] - float(column_bbox[0])
                            )
                            if abs(float(block.left_space) - desired_left) > 0.5:
                                block.left_space = desired_left
                                changed += 1
                            if block.right_space != 0:
                                block.right_space = 0.0
                                changed += 1
                            if block.first_line_space != 0:
                                block.first_line_space = 0.0
                                changed += 1
                            changed += _clear_pdf2docx_block_tabs(block)
                        continue

                    continuation_x0 = float(median(box[0] for box in row_boxes[1:]))
                    first_line_space = row_boxes[0][0] - continuation_x0
                    if abs(first_line_space) < max(1.0, size * 0.18):
                        first_line_space = 0.0
                    desired_left = max(0.0, continuation_x0 - float(column_bbox[0]))
                    full_row_boxes = row_boxes[:-1] or row_boxes
                    right_anchor = float(median(box[2] for box in full_row_boxes))
                    desired_right = max(0.0, float(column_bbox[2]) - right_anchor)
                    normalized = unicodedata.normalize("NFKC", text)
                    ascii_letters = sum(
                        character.isascii() and character.isalpha()
                        for character in normalized
                    )
                    spaces = normalized.count(" ")
                    reliable_spaces = spaces / max(1, ascii_letters) >= 0.055
                    full_width_rows = sum(
                        (box[2] - continuation_x0)
                        / max(1.0, float(column_bbox[2]) - continuation_x0)
                        >= 0.72
                        for box in full_row_boxes
                    )
                    desired_alignment = (
                        TextAlignment.JUSTIFY
                        if reliable_spaces
                        and full_width_rows
                        >= max(1, math.ceil(len(full_row_boxes) * 0.6))
                        and not is_list
                        else TextAlignment.LEFT
                    )
                    for attribute, desired in (
                        ("alignment", desired_alignment),
                        ("left_space", desired_left),
                        (
                            "right_space",
                            (
                                desired_right
                                if desired_alignment == TextAlignment.JUSTIFY
                                else 0.0
                            ),
                        ),
                        ("first_line_space", first_line_space),
                    ):
                        current = getattr(block, attribute)
                        different = (
                            current != desired
                            if attribute == "alignment"
                            else abs(float(current) - float(desired)) > 0.5
                        )
                        if different:
                            setattr(block, attribute, desired)
                            changed += 1
                    changed += _clear_pdf2docx_block_tabs(block)
                    # Do not synthesize a break for every physical PDF row.
                    # pdf2docx's native internal breaks are intentionally kept:
                    # they mark short headings, examples, addresses, and similar
                    # structure.  Ordinary prose rows remain break-free and can
                    # reflow after Word/WPS font substitution.  The universally
                    # redundant terminal break was removed above.
                if active_list_group:
                    list_groups.append(active_list_group)

                for group in list_groups:
                    option_blocks = [
                        block
                        for block in group
                        if _pdf2docx_text_starts_option(
                            str(getattr(block, "text", "") or "")
                        )
                    ]
                    if len(option_blocks) < 2:
                        continue
                    anchors = [
                        _pdf2docx_row_bbox(_pdf2docx_block_rows(block)[0])[0]
                        for block in option_blocks
                    ]
                    shared_left = max(
                        0.0, float(median(anchors)) - float(column_bbox[0])
                    )
                    for block in option_blocks:
                        if block.alignment != TextAlignment.LEFT:
                            block.alignment = TextAlignment.LEFT
                            changed += 1
                        if abs(float(block.left_space) - shared_left) > 0.05:
                            block.left_space = shared_left
                            changed += 1
                        if block.right_space != 0:
                            block.right_space = 0.0
                            changed += 1
                        changed += _clear_pdf2docx_block_tabs(block)
    return changed


def _apply_pdf2docx_column_layout_hint(
    pages: Iterable[Any],
    column_layout: str,
) -> int:
    """Apply a user column hint while page contents are still line-level objects."""

    normalized = _normalize_pdf2docx_column_layout(column_layout)
    if normalized != "single":
        # pdf2docx already supports mixed one/two-column sections.  For explicit
        # double or mixed hints we retain those sections and use a more sensitive
        # min_section_height while parsing instead of destructively forcing
        # full-width titles, figures, or tables into a column.
        return 0

    try:
        from pdf2docx.layout.Column import Column
        from pdf2docx.layout.Section import Section
    except ImportError as exc:
        raise MissingEngineError("PDF 转 Word 的分栏重建需要 pdf2docx") from exc

    changed = 0
    for page in pages:
        if bool(getattr(page, "skip_parsing", False)):
            continue
        sections = list(getattr(page, "sections", ()) or ())
        if not sections:
            continue

        blocks: list[Any] = []
        shapes: list[Any] = []
        seen_blocks: set[int] = set()
        seen_shapes: set[int] = set()
        for section in sections:
            for source_column in section:
                for block in source_column.blocks:
                    identity = id(block)
                    if identity not in seen_blocks:
                        seen_blocks.add(identity)
                        blocks.append(block)
                for shape in source_column.shapes:
                    identity = id(shape)
                    if identity not in seen_shapes:
                        seen_shapes.add(identity)
                        shapes.append(shape)

        if not blocks and not shapes:
            continue

        try:
            left, right, top, bottom = (
                float(value) for value in getattr(page, "margin", (0, 0, 0, 0))
            )
        except (TypeError, ValueError):
            left = right = top = bottom = 0.0
        page_width = max(1.0, float(getattr(page, "width", 0.0) or 0.0))
        page_height = max(1.0, float(getattr(page, "height", 0.0) or 0.0))
        x0 = max(0.0, left)
        y0 = max(0.0, top)
        x1 = max(x0 + 1.0, page_width - max(0.0, right))
        y1 = max(y0 + 1.0, page_height - max(0.0, bottom))

        merged_column = Column().update_bbox((x0, y0, x1, y1))
        merged_column.assign_blocks(blocks)
        merged_column.assign_shapes(shapes)
        merged_section = Section(
            space=0,
            columns=[merged_column],
            parent=page.sections,
        )
        merged_section.before_space = max(
            0.0,
            float(getattr(sections[0], "before_space", 0.0) or 0.0),
        )
        page.sections.reset([merged_section])
        changed += 1
    return changed


def _subtract_pdf_vertical_intervals(
    bands: Iterable[tuple[float, float]],
    cuts: Iterable[tuple[float, float]],
    *,
    minimum_height: float = 18.0,
) -> list[tuple[float, float]]:
    remaining = [tuple(map(float, band)) for band in bands]
    for cut_y0, cut_y1 in cuts:
        if cut_y1 <= cut_y0:
            continue
        next_remaining: list[tuple[float, float]] = []
        for band_y0, band_y1 in remaining:
            if cut_y1 <= band_y0 or cut_y0 >= band_y1:
                next_remaining.append((band_y0, band_y1))
                continue
            if cut_y0 - band_y0 >= minimum_height:
                next_remaining.append((band_y0, min(band_y1, cut_y0)))
            if band_y1 - cut_y1 >= minimum_height:
                next_remaining.append((max(band_y0, cut_y1), band_y1))
        remaining = next_remaining
    return sorted(remaining)


def _hybrid_editable_column_bands(
    assessment: _HybridPageAssessment,
) -> list[tuple[float, float]]:
    split_x = assessment.column_split_x
    if split_x is None:
        return []
    crossing_visual_cuts = [
        (region.rect[1] - 2.0, region.rect[3] + 2.0)
        for region in assessment.visual_regions
        if region.rect[0] < split_x - 12.0 and region.rect[2] > split_x + 12.0
    ]
    return _subtract_pdf_vertical_intervals(
        assessment.column_pair_bands,
        crossing_visual_cuts,
    )


def _pdf2docx_element_rect(element: Any) -> tuple[float, float, float, float]:
    bbox = element.bbox
    return tuple(float(bbox[index]) for index in range(4))  # type: ignore[return-value]


def _build_pdf2docx_column_sections(
    raw_page: Any,
    assessment: _HybridPageAssessment,
) -> list[Any]:
    """Rebuild reliable mixed sections directly from source line geometry.

    pdf2docx can mistake a centered page number for the second column and then
    silently discard the real right column.  The source detector already knows
    the gutter and paired-row bands, so use those facts before paragraph merging.
    """

    try:
        from pdf2docx.common.Collection import Collection
        from pdf2docx.layout.Column import Column
        from pdf2docx.layout.Section import Section
    except ImportError as exc:
        raise MissingEngineError("PDF 转 Word 的分栏重建需要 pdf2docx") from exc

    split_x = assessment.column_split_x
    bands = _hybrid_editable_column_bands(assessment)
    if split_x is None or not bands:
        return []

    page_x0, page_y0, page_x1, _page_y1 = (
        float(value) for value in raw_page.working_bbox
    )
    page_width = max(1.0, page_x1 - page_x0)
    gutter_margin = max(1.0, page_width * 0.004)

    # A full-width heading or image inside a provisional band is a genuine
    # single-column interruption.  Cut it out before assigning elements.
    crossing_element_cuts: list[tuple[float, float]] = []
    for block in raw_page.blocks:
        x0, y0, x1, y1 = _pdf2docx_element_rect(block)
        center_y = (y0 + y1) / 2.0
        if not any(band_y0 <= center_y <= band_y1 for band_y0, band_y1 in bands):
            continue
        if (
            x0 < split_x - gutter_margin
            and x1 > split_x + gutter_margin
            and x1 - x0 >= page_width * 0.45
        ):
            crossing_element_cuts.append((y0 - 1.0, y1 + 1.0))
    bands = _subtract_pdf_vertical_intervals(bands, crossing_element_cuts)
    if not bands:
        return []

    elements = Collection()
    elements.extend(raw_page.blocks)
    elements.extend(raw_page.shapes)
    grouped: dict[tuple[str, int], list[Any]] = {}
    for element in elements:
        x0, y0, x1, y1 = _pdf2docx_element_rect(element)
        center_y = (y0 + y1) / 2.0
        band_index = next(
            (
                index
                for index, (band_y0, band_y1) in enumerate(bands)
                if band_y0 <= center_y <= band_y1
            ),
            None,
        )
        confined = x1 <= split_x - gutter_margin or x0 >= split_x + gutter_margin
        if band_index is not None and confined:
            key = ("double", band_index)
        else:
            single_position = sum(
                1 for _band_y0, band_y1 in bands if center_y > band_y1
            )
            key = ("single", single_position)
        grouped.setdefault(key, []).append(element)

    ordered_keys = sorted(
        grouped,
        key=lambda key: min(_pdf2docx_element_rect(item)[1] for item in grouped[key]),
    )
    sections: list[Any] = []
    y_reference = page_y0
    for key in ordered_keys:
        items = grouped[key]
        render_as_double = key[0] == "double"
        left_items: list[Any] = []
        right_items: list[Any] = []
        if render_as_double:
            for item in items:
                x0, _y0, x1, _y1 = _pdf2docx_element_rect(item)
                if x1 <= split_x - gutter_margin:
                    left_items.append(item)
                elif x0 >= split_x + gutter_margin:
                    right_items.append(item)
            render_as_double = bool(left_items and right_items)

        if not render_as_double:
            y0 = min(_pdf2docx_element_rect(item)[1] for item in items)
            y1 = max(_pdf2docx_element_rect(item)[3] for item in items)
            column = Column().update_bbox((page_x0, y0, page_x1, y1))
            column.add_elements(items)
            section = Section(space=0, columns=[column])
            section.before_space = max(0.0, round(y0 - y_reference, 1))
        else:
            left_y0 = min(_pdf2docx_element_rect(item)[1] for item in left_items)
            left_y1 = max(_pdf2docx_element_rect(item)[3] for item in left_items)
            right_y0 = min(_pdf2docx_element_rect(item)[1] for item in right_items)
            right_y1 = max(_pdf2docx_element_rect(item)[3] for item in right_items)
            left_column = Column().update_bbox((page_x0, left_y0, split_x, left_y1))
            left_column.add_elements(left_items)
            right_column = Column().update_bbox((split_x, right_y0, page_x1, right_y1))
            right_column.add_elements(right_items)
            section = Section(
                space=0,
                columns=[left_column, right_column],
            )
            section.before_space = max(
                0.0,
                round(min(left_y0, right_y0) - y_reference, 1),
            )
        sections.append(section)
        y_reference = max(y_reference, float(section.bbox[3]))
    return sections


def _parse_pdf2docx_document_with_layout_hints(
    converter: Any,
    settings: Mapping[str, Any],
    page_assessments: Mapping[int, _HybridPageAssessment],
) -> None:
    """Equivalent to pdf2docx Pages.parse(), with page-local gutter guidance."""

    try:
        from pdf2docx.font.Fonts import Fonts
        from pdf2docx.page.RawPageFactory import RawPageFactory
    except ImportError as exc:
        raise MissingEngineError("PDF 转 Word 的分栏重建需要 pdf2docx") from exc

    settings_dict = dict(settings)
    fonts = Fonts.extract(converter.fitz_doc)
    parsed_pairs: list[tuple[Any, Any]] = []
    for page in converter.pages:
        if bool(getattr(page, "skip_parsing", False)):
            continue
        raw_page = RawPageFactory.create(
            page_engine=converter.fitz_doc[page.id],
            backend="PyMuPDF",
        )
        raw_page.restore(**settings_dict)
        raw_page.clean_up(**settings_dict)
        raw_page.process_font(fonts)
        page.width = raw_page.width
        page.height = raw_page.height
        page.float_images.reset().extend(raw_page.blocks.floating_image_blocks)
        raw_page.margin = page.margin = raw_page.calculate_margin(**settings_dict)
        parsed_pairs.append((page, raw_page))

    for page, raw_page in parsed_pairs:
        assessment = page_assessments.get(int(page.id))
        sections: list[Any] = []
        if assessment is not None and assessment.detected_two_columns:
            with contextlib.suppress(Exception):
                sections = _build_pdf2docx_column_sections(raw_page, assessment)
        if not sections:
            sections = list(raw_page.parse_section(**settings_dict) or ())
        page.sections.reset(sections)


def _parse_pdf2docx_pages(
    converter: Any,
    *,
    settings: Mapping[str, Any],
    column_layout: str,
    page_assessments: Mapping[int, _HybridPageAssessment] | None = None,
    start: int = 0,
    end: int | None = None,
    pages: list[int] | None = None,
) -> None:
    """Parse in two phases so a column hint can affect the raw page sections."""

    converter.load_pages(start=start, end=end, pages=pages)
    if page_assessments:
        _parse_pdf2docx_document_with_layout_hints(
            converter,
            settings,
            page_assessments,
        )
    else:
        converter.parse_document(**dict(settings))
    _apply_pdf2docx_column_layout_hint(converter.pages, column_layout)
    converter.parse_pages(**dict(settings))
    _reconstruct_pdf2docx_prose_blocks(converter.pages)


def _extract_pdf2docx_page_text(page: Any) -> str:
    parts: list[str] = []
    for section in page.sections:
        for column in section:
            for block in _iter_pdf2docx_text_blocks(column.blocks):
                text = str(getattr(block, "text", "") or "")
                if text:
                    parts.append(text)
    return "\n".join(parts)


def _pdf2docx_page_layout_risk_reason(page: Any) -> str | None:
    """Detect section transitions that are unstable in Word/WPS flow layout."""

    sections = list(page.sections)
    if not sections:
        return "未识别到可编辑版面"
    column_counts = [len(section) for section in sections]
    if any(count > 2 for count in column_counts):
        return "三栏及以上复杂分栏"

    def bbox_value(bbox: Any, name: str, index: int) -> float:
        value = getattr(bbox, name, None)
        if value is not None:
            return float(value)
        return float(bbox[index])

    for section, column_count in zip(sections, column_counts):
        if column_count != 2:
            continue
        column_sizes: list[tuple[float, float]] = []
        for column in section:
            bbox = getattr(column, "bbox", None)
            if bbox is None:
                column_sizes = []
                break
            width = getattr(bbox, "width", None)
            height = getattr(bbox, "height", None)
            if width is None:
                width = bbox_value(bbox, "x1", 2) - bbox_value(bbox, "x0", 0)
            if height is None:
                height = bbox_value(bbox, "y1", 3) - bbox_value(bbox, "y0", 1)
            column_sizes.append((max(0.0, float(width)), max(0.0, float(height))))
        if len(column_sizes) != 2:
            continue
        width_ratio = min(size[0] for size in column_sizes) / max(
            1.0, max(size[0] for size in column_sizes)
        )
        height_ratio = min(size[1] for size in column_sizes) / max(
            1.0, max(size[1] for size in column_sizes)
        )
        if width_ratio < 0.62 or (width_ratio < 0.80 and height_ratio < 0.55):
            return "不等宽短侧栏与主正文混排"

    multi_column_indexes = [
        index for index, count in enumerate(column_counts) if count >= 2
    ]
    if not multi_column_indexes:
        return None
    if len(multi_column_indexes) > 1:
        return "页面内多次分栏切换"

    page_height = max(1.0, float(getattr(page, "height", 0.0) or 0.0))
    for section, column_count in zip(sections, column_counts):
        if column_count != 1:
            continue
        bbox = section.bbox
        try:
            y0 = float(bbox.y0)
            y1 = float(bbox.y1)
        except AttributeError:
            y0 = float(bbox[1])
            y1 = float(bbox[3])
        height_ratio = max(0.0, y1 - y0) / page_height
        is_header = y1 <= page_height * 0.06
        is_footer = y0 >= page_height * 0.90
        if height_ratio >= 0.08 and not is_header and not is_footer:
            return "全宽内容与多栏正文混排"
    return None


def _pdf2docx_intercolumn_text_merge_rects(
    page: Any,
    *,
    split_x: float | None,
    paired_bands: Iterable[tuple[float, float]],
) -> list[tuple[float, float, float, float]]:
    """Return parsed paragraphs that wrongly combine simultaneous columns."""

    bands = tuple(paired_bands)
    if split_x is None or not bands:
        return []
    page_width = max(1.0, float(getattr(page, "width", 0.0) or 0.0))
    side_margin = page_width * 0.10
    matches: list[tuple[float, float, float, float]] = []
    for section in page.sections:
        if len(section) != 1:
            continue
        for column in section:
            for block in _iter_pdf2docx_text_blocks(column.blocks):
                text = str(getattr(block, "text", "") or "")
                if (
                    len(_normalize_validation_text(text)) < 24
                    or text.strip() == "<image>"
                ):
                    continue
                bbox = getattr(block, "bbox", None)
                if bbox is None:
                    continue
                try:
                    x0 = float(getattr(bbox, "x0") if hasattr(bbox, "x0") else bbox[0])
                    y0 = float(getattr(bbox, "y0") if hasattr(bbox, "y0") else bbox[1])
                    x1 = float(getattr(bbox, "x1") if hasattr(bbox, "x1") else bbox[2])
                    y1 = float(getattr(bbox, "y1") if hasattr(bbox, "y1") else bbox[3])
                except (IndexError, TypeError, ValueError):
                    continue
                if x0 > split_x - side_margin or x1 < split_x + side_margin:
                    continue
                if any(
                    min(y1, band_y1) - max(y0, band_y0) > 0.5
                    for band_y0, band_y1 in bands
                ):
                    matches.append((x0, y0, x1, y1))
    return matches


def _pdf2docx_page_has_intercolumn_text_merge(
    page: Any,
    *,
    split_x: float | None,
    paired_bands: Iterable[tuple[float, float]],
) -> bool:
    """Find parsed paragraphs that span both source columns on the same rows."""

    return bool(
        _pdf2docx_intercolumn_text_merge_rects(
            page,
            split_x=split_x,
            paired_bands=paired_bands,
        )
    )


def _pdf2docx_page_quality_reason(
    source_text: str,
    parsed_text: str,
    *,
    source_blocks: Iterable[str] | None = None,
) -> str | None:
    coverage = _text_sequence_coverage(source_text, parsed_text)
    english_word_recall = _english_word_multiset_recall(source_text, parsed_text)
    adjacent_word_coverage = (
        _block_local_adjacent_english_word_coverage(source_blocks, parsed_text)
        if source_blocks is not None
        else _adjacent_english_word_coverage(source_text, parsed_text)
    )
    strict_quality_pass = (
        coverage >= _MIN_HYBRID_PAGE_SEQUENCE_COVERAGE
        and english_word_recall >= _MIN_HYBRID_PAGE_ENGLISH_WORD_RECALL
        and adjacent_word_coverage >= _MIN_HYBRID_PAGE_ADJACENT_WORD_COVERAGE
    )
    high_recall_tolerance_pass = (
        coverage >= _MIN_HYBRID_HIGH_RECALL_SEQUENCE_COVERAGE
        and english_word_recall >= _MIN_HYBRID_HIGH_RECALL_ENGLISH_WORD_RECALL
        and adjacent_word_coverage >= _MIN_HYBRID_HIGH_RECALL_ADJACENT_WORD_COVERAGE
    )
    boundary_recovery_pass = _pdf_english_boundary_recovery_pass(
        source_text,
        parsed_text,
        character_coverage=coverage,
    )
    if strict_quality_pass or high_recall_tolerance_pass or boundary_recovery_pass:
        return None
    return (
        "页级文字恢复不足"
        f"（字符 {coverage:.0%}、英文词 {english_word_recall:.0%}、词序 {adjacent_word_coverage:.0%}）"
    )


def _word_font_key(font_name: str) -> str:
    without_subset = re.sub(r"^[A-Za-z]{6}\+", "", str(font_name or "").strip())
    return re.sub(r"[^a-z0-9]", "", without_subset.casefold())


def _font_key_variants(font_name: str) -> set[str]:
    key = _word_font_key(font_name)
    variants = {key} if key else set()
    style_suffixes = (
        "bolditalic",
        "semibolditalic",
        "boldoblique",
        "semibold",
        "regular",
        "italic",
        "oblique",
        "medium",
        "light",
        "narrow",
        "black",
        "bold",
    )
    for suffix in style_suffixes:
        if key.endswith(suffix) and len(key) > len(suffix) + 2:
            variants.add(key[: -len(suffix)])
    return variants


@lru_cache(maxsize=1)
def _installed_word_font_keys() -> frozenset[str]:
    """Return font-family keys reported by the Windows font registry."""

    names: set[str] = set()
    try:
        import winreg

        registry_path = r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts"
        views = [0]
        for flag_name in ("KEY_WOW64_64KEY", "KEY_WOW64_32KEY"):
            flag = getattr(winreg, flag_name, 0)
            if flag and flag not in views:
                views.append(flag)
        for root in (winreg.HKEY_LOCAL_MACHINE, winreg.HKEY_CURRENT_USER):
            for view in views:
                try:
                    key = winreg.OpenKey(
                        root,
                        registry_path,
                        0,
                        winreg.KEY_READ | view,
                    )
                except OSError:
                    continue
                with key:
                    index = 0
                    while True:
                        try:
                            display_name = str(winreg.EnumValue(key, index)[0])
                        except OSError:
                            break
                        index += 1
                        display_name = re.sub(r"\s*\([^)]*\)\s*$", "", display_name)
                        names.update(
                            part.strip()
                            for part in display_name.split("&")
                            if part.strip()
                        )
    except (ImportError, OSError):
        pass

    keys: set[str] = set()
    for name in names:
        keys.update(_font_key_variants(name))
    return frozenset(keys)


def _standard_word_font_keys() -> frozenset[str]:
    names = {
        "Arial",
        "Calibri",
        "Cambria",
        "Cambria Math",
        "Consolas",
        "Courier New",
        "Georgia",
        "Microsoft JhengHei",
        "Microsoft YaHei",
        "MS Gothic",
        "MS Mincho",
        "Palatino Linotype",
        "Segoe UI",
        "SimHei",
        "SimSun",
        "Symbol",
        "Tahoma",
        "Times New Roman",
        "Verdana",
        "Wingdings",
    }
    keys: set[str] = set()
    for name in names:
        keys.update(_font_key_variants(name))
    return frozenset(keys)


def _fallback_word_font_name(font_name: str, *, east_asia: bool = False) -> str:
    compact = _word_font_key(font_name)
    if any(
        token in compact
        for token in (
            "math",
            "symbol",
            "cmmi",
            "cmsy",
            "cmex",
            "texcmmath",
            "rtxmi",
            "txmi",
            "txsy",
            "yhcmex",
            "msam",
            "msbm",
            "wasy",
            "rsfs",
        )
    ):
        return "Cambria Math"
    if any(token in compact for token in ("song", "sung", "ming", "mincho")):
        return "SimSun"
    if any(token in compact for token in ("hei", "cjk", "yahei")):
        return "Microsoft YaHei"
    if east_asia:
        return "Microsoft YaHei"
    if any(
        token in compact
        for token in (
            "sans",
            "arial",
            "helvetica",
            "formata",
            "grotesk",
            "univers",
            "frutiger",
        )
    ):
        return "Arial"
    if any(token in compact for token in ("mono", "courier", "typewriter", "console")):
        return "Courier New"
    return "Times New Roman"


def _resolve_word_font_name(
    font_name: str,
    *,
    east_asia: bool,
    installed_fonts: frozenset[str],
) -> str:
    available_fonts = installed_fonts or _standard_word_font_keys()
    normalized_name = _normalized_word_font_name(font_name)
    if _font_key_variants(normalized_name) & available_fonts:
        return normalized_name

    preferred = _fallback_word_font_name(normalized_name, east_asia=east_asia)
    fallback_candidates: dict[str, tuple[str, ...]] = {
        "Cambria Math": ("Cambria Math", "Times New Roman", "Cambria"),
        "Microsoft YaHei": (
            "Microsoft YaHei",
            "SimSun",
            "Microsoft JhengHei",
            "Arial",
        ),
        "SimSun": ("SimSun", "Microsoft YaHei", "Microsoft JhengHei", "Arial"),
        "Arial": ("Arial", "Calibri", "Segoe UI", "Times New Roman"),
        "Courier New": ("Courier New", "Consolas", "Times New Roman"),
        "Times New Roman": ("Times New Roman", "Cambria", "Arial"),
    }
    for candidate in fallback_candidates.get(preferred, (preferred,)):
        if _font_key_variants(candidate) & available_fonts:
            return candidate
    return preferred


def _word_font_width_scale_percent(source_name: str, target_name: str) -> int:
    """Compensate common PDF/PostScript font aliases after Word substitution.

    The PostScript Times families used by many papers are measurably narrower than
    the Windows Times New Roman renderer.  Without preserving that metric
    difference, 9--11 pt body paragraphs gain extra wrapped lines and can spill
    past a source-page section break in Word/WPS.
    """

    source = re.sub(r"^[A-Za-z]{6}\+", "", str(source_name or "").strip())
    source_key = _word_font_key(source)
    target_key = _word_font_key(target_name)
    if target_key == _word_font_key("Times New Roman") and source_key.startswith(
        (
            "timesnewromanps",
            "nimbusromno9l",
            "texgyretermes",
            "texgyretermesx",
        )
    ):
        return 96
    return 100


def _run_font_size_half_points(run_properties: Any, qn: Any) -> int | None:
    for tag_name in ("w:sz", "w:szCs"):
        element = run_properties.find(qn(tag_name))
        if element is None:
            continue
        try:
            return int(str(element.get(qn("w:val")) or ""))
        except ValueError:
            continue
    return None


def _set_word_run_width_scale(
    run_properties: Any,
    percent: int,
    *,
    qn: Any,
    element_factory: Any,
) -> bool:
    """Set OOXML character scaling without compounding an existing adjustment."""

    percent = max(50, min(100, int(percent)))
    width_element = run_properties.find(qn("w:w"))
    if width_element is None:
        width_element = element_factory("w:w")
        run_properties.append(width_element)
        current = 100
    else:
        try:
            current = int(str(width_element.get(qn("w:val")) or "100"))
        except ValueError:
            current = 100
    desired = min(current, percent)
    if width_element.get(qn("w:val")) == str(desired):
        return False
    width_element.set(qn("w:val"), str(desired))
    return True


def _font_requires_visual_fallback(font_name: str) -> bool:
    compact = _word_font_key(font_name)
    if not any(
        token in compact
        for token in (
            "awesome",
            "dingbat",
            "icon",
            "symbolmt",
            "webdings",
            "wingdings",
            "zapf",
        )
    ):
        return False
    installed_fonts = _installed_word_font_keys() or _standard_word_font_keys()
    return not bool(_font_key_variants(font_name) & installed_fonts)


def _normalized_word_font_name(font_name: str) -> str:
    """Map non-portable PDF font names to reliable Word/WPS font families."""

    original = str(font_name or "").strip()
    if not original:
        return original
    without_subset = re.sub(r"^[A-Za-z]{6}\+", "", original)
    compact = _word_font_key(without_subset)
    canonical_aliases = {
        "microsoftyahei": "Microsoft YaHei",
        "microsoftyaheibold": "Microsoft YaHei",
        "microsoftyaheiui": "Microsoft YaHei",
        "microsoftyaheiuibold": "Microsoft YaHei",
        "microsoftjhenghei": "Microsoft JhengHei",
        "microsoftjhengheibold": "Microsoft JhengHei",
        "segoeui": "Segoe UI",
        "segoeuibold": "Segoe UI",
        "simsun": "SimSun",
        "simhei": "SimHei",
        "timesnewroman": "Times New Roman",
        "timesnewromanbold": "Times New Roman",
        "arial": "Arial",
        "arialbold": "Arial",
    }
    if compact in canonical_aliases:
        return canonical_aliases[compact]
    if compact.startswith(("fzssk", "fzdbsk")):
        return "SimSun"
    if compact.startswith(("fzhtk", "fzdhtk")):
        return "SimHei"
    if compact in {"eb1"}:
        return "Times New Roman"
    if compact in {"ebz", "efz"}:
        return "SimSun"
    if compact in {"ehz", "eh7"}:
        return "SimHei"
    if compact in {"eyt1"}:
        return "Microsoft YaHei"
    if compact.startswith("lmroman"):
        return "Times New Roman"
    if compact.startswith("lmsans"):
        return "Arial"
    if compact.startswith("lmmono"):
        return "Courier New"
    if compact.startswith(("cmr", "cmbx", "texcmroman")):
        return "Times New Roman"
    if compact.startswith(("cmss", "texcmsans")):
        return "Arial"
    if compact.startswith(("cmtt", "texcmtypewriter")):
        return "Courier New"
    if compact.startswith(
        (
            "cmmi",
            "cmsy",
            "cmex",
            "texcmmath",
            "rtxmi",
            "txmi",
            "txsy",
            "yhcmex",
            "msam",
            "msbm",
            "wasy",
            "rsfs",
        )
    ):
        return "Cambria Math"
    if compact.startswith(
        (
            "arialmt",
            "helvetica",
            "formata",
            "nimbussan",
            "texgyreheros",
        )
    ):
        return "Arial"
    if (
        compact.startswith(
            (
                "timesnewromanps",
                "timesroman",
                "timesltstd",
                "nimbusromno9l",
                "urwpalladiol",
                "texgyretermes",
                "texgyrepagella",
                "rtxr",
            )
        )
        or compact == "times"
    ):
        return "Times New Roman"
    if compact.startswith(("couriernewps", "courierstd")) or compact == "courier":
        return "Courier New"
    return without_subset


def _normalize_docx_fonts(document: Any) -> int:
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from lxml import etree
    except ImportError as exc:
        raise MissingEngineError("PDF 转 Word 需要 python-docx") from exc

    mapped_attributes = 0
    installed_fonts = _installed_word_font_keys()
    east_asia_attribute = qn("w:eastAsia")
    font_attributes = tuple(
        qn(name) for name in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs")
    )
    for part in document.part.package.parts:
        element = getattr(part, "element", None)
        if element is None:
            element = getattr(part, "_element", None)
        if element is None or not hasattr(element, "xpath"):
            continue
        try:
            font_elements = element.xpath(".//w:rFonts")
        except (AttributeError, TypeError, ValueError):
            continue
        for font_element in font_elements:
            source_names: list[str] = []
            resolved_names: list[str] = []
            for attribute in font_attributes:
                source_name = font_element.get(attribute)
                if not source_name:
                    continue
                source_names.append(source_name)
                normalized_name = _resolve_word_font_name(
                    source_name,
                    east_asia=attribute == east_asia_attribute,
                    installed_fonts=installed_fonts,
                )
                resolved_names.append(normalized_name)
                if normalized_name != source_name:
                    font_element.set(attribute, normalized_name)
                    mapped_attributes += 1

            compact_names = [_word_font_key(name) for name in source_names]
            run_properties = font_element.getparent()
            if run_properties is None:
                continue
            target_keys = {_word_font_key(name) for name in resolved_names if name}
            font_size = _run_font_size_half_points(run_properties, qn)
            if target_keys == {_word_font_key("Times New Roman")} and (
                font_size is not None and 18 <= font_size <= 22
            ):
                width_scale = min(
                    _word_font_width_scale_percent(source_name, target_name)
                    for source_name, target_name in zip(source_names, resolved_names)
                )
                if width_scale < 100 and _set_word_run_width_scale(
                    run_properties,
                    width_scale,
                    qn=qn,
                    element_factory=OxmlElement,
                ):
                    mapped_attributes += 1
            if any(
                name.startswith(("cmmi", "rtxmi", "txmi")) for name in compact_names
            ):
                if run_properties.find(qn("w:i")) is None:
                    run_properties.append(OxmlElement("w:i"))
            if any(name.startswith("cmbx") for name in compact_names):
                if run_properties.find(qn("w:b")) is None:
                    run_properties.append(OxmlElement("w:b"))

    word_namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for part in document.part.package.parts:
        if not str(part.partname).lower().endswith(".xml"):
            continue
        if (
            getattr(part, "element", None) is not None
            or getattr(part, "_element", None) is not None
        ):
            continue
        blob = getattr(part, "blob", b"")
        if not blob:
            continue
        try:
            root = etree.fromstring(blob)
        except (etree.XMLSyntaxError, TypeError, ValueError):
            continue
        changed = False
        for font_element in root.xpath(".//w:rFonts", namespaces={"w": word_namespace}):
            source_names: list[str] = []
            resolved_names: list[str] = []
            for attribute in font_attributes:
                source_name = font_element.get(attribute)
                if not source_name:
                    continue
                source_names.append(source_name)
                normalized_name = _resolve_word_font_name(
                    source_name,
                    east_asia=attribute == east_asia_attribute,
                    installed_fonts=installed_fonts,
                )
                resolved_names.append(normalized_name)
                if normalized_name != source_name:
                    font_element.set(attribute, normalized_name)
                    mapped_attributes += 1
                    changed = True
            run_properties = font_element.getparent()
            if run_properties is None:
                continue
            target_keys = {_word_font_key(name) for name in resolved_names if name}
            font_size = _run_font_size_half_points(run_properties, qn)
            if target_keys == {_word_font_key("Times New Roman")} and (
                font_size is not None and 18 <= font_size <= 22
            ):
                width_scale = min(
                    _word_font_width_scale_percent(source_name, target_name)
                    for source_name, target_name in zip(source_names, resolved_names)
                )
                if width_scale < 100 and _set_word_run_width_scale(
                    run_properties,
                    width_scale,
                    qn=qn,
                    element_factory=OxmlElement,
                ):
                    mapped_attributes += 1
                    changed = True
        if changed:
            part._blob = etree.tostring(
                root,
                encoding="UTF-8",
                xml_declaration=True,
                standalone=True,
            )
    return mapped_attributes


def _stabilize_pdf2docx_paragraph_layout(document: Any) -> int:
    """Prevent Word/WPS compatibility rules from expanding PDF-derived flow."""

    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError as exc:
        raise MissingEngineError("PDF 转 Word 需要 python-docx") from exc

    changed = 0
    settings = document.settings.element
    compatibility = settings.find(qn("w:compat"))
    if compatibility is None:
        compatibility = OxmlElement("w:compat")
        settings.insert_element_before(
            compatibility,
            "w:rsids",
            "m:mathPr",
            "w:themeFontLang",
            "w:clrSchemeMapping",
            "w:doNotAutoCompressPictures",
        )
        changed += 1
    do_not_expand = compatibility.find(qn("w:doNotExpandShiftReturn"))
    if do_not_expand is None:
        do_not_expand = OxmlElement("w:doNotExpandShiftReturn")
        compatibility.append(do_not_expand)
        changed += 1
    if do_not_expand.get(qn("w:val")) != "1":
        do_not_expand.set(qn("w:val"), "1")
        changed += 1

    for paragraph in document.element.body.xpath(".//w:p"):
        properties = paragraph.get_or_add_pPr()
        for tag_name in ("w:snapToGrid", "w:widowControl"):
            element = properties.find(qn(tag_name))
            if element is None:
                element = OxmlElement(tag_name)
                section_properties = properties.find(qn("w:sectPr"))
                if section_properties is None:
                    properties.append(element)
                else:
                    properties.insert(properties.index(section_properties), element)
            if element.get(qn("w:val")) != "0":
                element.set(qn("w:val"), "0")
                changed += 1

        # pdf2docx positions source-page content with exact line boxes.  Some
        # PDFs encode an ordinary 9--12 pt body line with a box more than twice
        # the font height.  Word and WPS honor that oversized box literally,
        # causing dense source pages to spill into page-number-only sheets.  A
        # tightly scoped compatibility reserve fixes the reflow without
        # changing titles, formulas, visual carriers, or section boundaries.
        if (
            properties.find(qn("w:framePr")) is not None
            or properties.find(qn("w:sectPr")) is not None
            or paragraph.xpath(
                ".//w:drawing|.//w:pict|.//w:object|.//m:oMath|.//m:oMathPara"
            )
        ):
            continue

        visible_text = "".join(paragraph.xpath(".//w:t/text()"))
        normalized_text = unicodedata.normalize("NFKC", visible_text).strip()
        if not normalized_text or _PDF2DOCX_CLEAR_PAGE_NUMBER_PATTERN.fullmatch(
            normalized_text
        ):
            continue

        size_weights: Counter[int] = Counter()
        text_runs: list[tuple[Any, Any, int]] = []
        for run in paragraph.xpath(".//w:r"):
            run_text = "".join(run.xpath(".//w:t/text()"))
            if not run_text.strip():
                continue
            run_properties = run.find(qn("w:rPr"))
            if run_properties is None:
                continue
            size_element = run_properties.find(qn("w:sz"))
            if size_element is None:
                continue
            try:
                half_points = int(str(size_element.get(qn("w:val")) or ""))
            except ValueError:
                continue
            size_weights[half_points] += max(1, len(run_text.strip()))
            text_runs.append((run, run_properties, half_points))
        if not size_weights:
            continue

        dominant_half_points = max(
            size_weights.items(),
            key=lambda item: (item[1], item[0]),
        )[0]
        if not 18 <= dominant_half_points <= 24:
            continue

        added_width_reserve = False
        for _run, run_properties, half_points in text_runs:
            if not 18 <= half_points <= 24:
                continue
            # Preserve pdf2docx's own character scaling and the narrower
            # PostScript-font compensation installed by _normalize_docx_fonts.
            if run_properties.find(qn("w:w")) is not None:
                continue
            width = OxmlElement("w:w")
            width.set(qn("w:val"), "98")
            run_properties.append(width)
            added_width_reserve = True
            changed += 1

        spacing = properties.find(qn("w:spacing"))
        if spacing is None:
            continue

        abnormal_exact_line = False
        if (spacing.get(qn("w:lineRule")) or "").lower() == "exact":
            try:
                original_line = int(str(spacing.get(qn("w:line")) or "0"))
            except ValueError:
                original_line = 0
            font_height = dominant_half_points * 10
            maximum_line = dominant_half_points * 20
            abnormal_exact_line = original_line > maximum_line
            if original_line > 0 and (abnormal_exact_line or added_width_reserve):
                compatible_line = min(original_line, maximum_line)
                compatible_line = max(font_height, round(compatible_line * 0.98))
                if compatible_line < original_line:
                    spacing.set(qn("w:line"), str(compatible_line))
                    changed += 1

        # Tie the one-time 2% paragraph reserve either to a newly added width
        # marker or to the naturally idempotent oversized-line correction.  A
        # second stabilizer pass therefore cannot compound the adjustment.
        if added_width_reserve or abnormal_exact_line:
            try:
                original_before = int(str(spacing.get(qn("w:before")) or "0"))
            except ValueError:
                original_before = 0
            compatible_before = max(0, round(original_before * 0.98))
            if compatible_before < original_before:
                spacing.set(qn("w:before"), str(compatible_before))
                changed += 1

    return changed


_PDF2DOCX_PAGE_FURNITURE_PATTERN = re.compile(
    r"(?:https?://|www\.|doi\.org|(?:^|\s)\d+\s*/\s*\d+(?:\s|$)|"
    r"^\s*[-‐‑‒–—―−﹘﹣－]\s*(?:\d\s*){1,4}[-‐‑‒–—―−﹘﹣－]\s*$|"
    r"^\s*(?:\d\s*){1,4}$|"
    r"^\s*第\s*(?:\d\s*){1,6}页(?:\s*共\s*(?:\d\s*){1,6}页)?\s*$|"
    r"^\s*共\s*(?:\d\s*){1,6}页\s*第\s*(?:\d\s*){1,6}页\s*$|"
    r"\bpage\s+\d+(?:\s+of\s+\d+)?)",
    re.IGNORECASE,
)
_PDF2DOCX_CLEAR_PAGE_NUMBER_PATTERN = re.compile(
    r"(?:^|\s)\d+\s*/\s*\d+(?:\s|$)|"
    r"^\s*[-‐‑‒–—―−﹘﹣－]\s*(?:\d\s*){1,4}[-‐‑‒–—―−﹘﹣－]\s*$|"
    r"^\s*(?:\d\s*){1,4}$|"
    r"^\s*第\s*(?:\d\s*){1,6}页(?:\s*共\s*(?:\d\s*){1,6}页)?\s*$|"
    r"^\s*共\s*(?:\d\s*){1,6}页\s*第\s*(?:\d\s*){1,6}页\s*$|"
    r"\bpage\s+\d+(?:\s+of\s+\d+)?",
    re.IGNORECASE,
)
_PDF2DOCX_FORWARD_FULL_PAGE_COUNTER_PATTERN = re.compile(
    r"第\s*(?P<current>(?:\d\s*){1,6})页\s*" r"共\s*(?P<total>(?:\d\s*){1,6})页",
    re.IGNORECASE,
)
_PDF2DOCX_REVERSED_FULL_PAGE_COUNTER_PATTERN = re.compile(
    r"共\s*(?P<total>(?:\d\s*){1,6})页\s*" r"第\s*(?P<current>(?:\d\s*){1,6})页",
    re.IGNORECASE,
)


def _pdf2docx_page_counter_signatures(text: str) -> tuple[tuple[int, int], ...]:
    """Return ordered ``(current, total)`` counters embedded in rendered text."""

    normalized = unicodedata.normalize("NFKC", str(text or ""))
    matches: list[tuple[int, int, int]] = []
    for pattern in (
        _PDF2DOCX_FORWARD_FULL_PAGE_COUNTER_PATTERN,
        _PDF2DOCX_REVERSED_FULL_PAGE_COUNTER_PATTERN,
    ):
        for match in pattern.finditer(normalized):
            current_digits = "".join(
                character for character in match["current"] if character.isdigit()
            )
            total_digits = "".join(
                character for character in match["total"] if character.isdigit()
            )
            if not current_digits or not total_digits:
                continue
            current = int(current_digits)
            total = int(total_digits)
            if current <= 0 or total <= 0 or current > total:
                continue
            matches.append((match.start(), current, total))

    signatures: list[tuple[int, int]] = []
    for _offset, current, total in sorted(matches):
        signature = (current, total)
        if signature not in signatures:
            signatures.append(signature)
    return tuple(signatures)


def _pdf2docx_conflicting_page_counter_signatures(
    text: str,
) -> tuple[tuple[int, int], ...]:
    """Find two or more different full counters for one logical page total."""

    counters = _pdf2docx_page_counter_signatures(text)
    totals: dict[int, list[tuple[int, int]]] = {}
    for counter in counters:
        totals.setdefault(counter[1], []).append(counter)
    conflicts = [
        counter
        for total in sorted(totals)
        if len({current for current, _total in totals[total]}) > 1
        for counter in totals[total]
    ]
    return tuple(conflicts)


def _pdf2docx_isolated_page_number_text(text: str) -> bool:
    normalized = unicodedata.normalize("NFKC", str(text or "")).strip()
    return bool(
        normalized and _PDF2DOCX_CLEAR_PAGE_NUMBER_PATTERN.fullmatch(normalized)
    )


def _deduplicate_adjacent_pdf2docx_page_counters(document: Any) -> int:
    """Remove an older counter when two framed source counters are adjacent.

    This is a safe structural fast path for source pages that Word has already
    consolidated.  It only touches standalone, page-anchored full counters with
    the same declared total and never crosses ordinary text, tables, drawings,
    or a section carrier.  The later WPS render pass handles non-adjacent cases.
    """

    from docx.oxml.ns import qn

    body = document.element.body
    children = [child for child in body if child.tag != qn("w:sectPr")]

    def counter_signature(paragraph: Any) -> tuple[int, int] | None:
        if paragraph.tag != qn("w:p"):
            return None
        properties = paragraph.pPr
        if (
            properties is None
            or properties.sectPr is not None
            or properties.find(qn("w:framePr")) is None
            or paragraph.xpath(".//w:drawing|.//w:pict|.//w:object")
        ):
            return None
        text = " ".join(
            "".join(paragraph.xpath(".//w:t/text()")).replace("\u200b", "").split()
        )
        signatures = _pdf2docx_page_counter_signatures(text)
        if len(signatures) != 1 or not _pdf2docx_isolated_page_number_text(text):
            return None
        return signatures[0]

    changed = 0
    index = 0
    while index < len(children):
        signature = counter_signature(children[index])
        if signature is None:
            index += 1
            continue
        group = [(children[index], signature)]
        group_index = index + 1
        while group_index < len(children):
            next_signature = counter_signature(children[group_index])
            if next_signature is None or next_signature[1] != signature[1]:
                break
            group.append((children[group_index], next_signature))
            group_index += 1
        if len(group) > 1:
            keep_element, _keep_signature = max(
                group,
                key=lambda item: item[1][0],
            )
            for paragraph, _counter in group:
                if paragraph is keep_element or paragraph.getparent() is None:
                    continue
                paragraph.getparent().remove(paragraph)
                changed += 1
        index = group_index
    return changed


def _repair_pdf2docx_rendered_page_counter_conflicts(
    docx_path: Path,
    conflicts: Iterable[_RenderedPageCounterConflict],
) -> int:
    """Remove uniquely identifiable older counters reported by WPS rendering."""

    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError as exc:
        raise MissingEngineError("PDF 转 Word 需要 python-docx") from exc

    conflict_list = list(conflicts)
    removable_signatures: set[tuple[int, int]] = set()
    for conflict in conflict_list:
        by_total: dict[int, list[tuple[int, int]]] = {}
        for counter in conflict.counters:
            by_total.setdefault(counter[1], []).append(counter)
        for counters in by_total.values():
            if len(counters) < 2:
                continue
            keep = max(counters, key=lambda counter: counter[0])
            removable_signatures.update(
                counter for counter in counters if counter != keep
            )
    if not removable_signatures:
        return 0

    document = Document(docx_path)
    candidates: dict[tuple[int, int], list[Any]] = {}
    for paragraph in _body_content_children(document):
        if paragraph.tag != qn("w:p"):
            continue
        properties = paragraph.pPr
        if (
            properties is None
            or properties.sectPr is not None
            or properties.find(qn("w:framePr")) is None
            or paragraph.xpath(".//w:drawing|.//w:pict|.//w:object")
        ):
            continue
        text = " ".join(
            "".join(paragraph.xpath(".//w:t/text()")).replace("\u200b", "").split()
        )
        signatures = _pdf2docx_page_counter_signatures(text)
        if len(signatures) != 1 or not _pdf2docx_isolated_page_number_text(text):
            continue
        candidates.setdefault(signatures[0], []).append(paragraph)

    changed = 0
    for signature in removable_signatures:
        matching_paragraphs = candidates.get(signature, ())
        # A source page counter is expected to be unique.  Ambiguous repeated
        # text is left untouched so an ordinary body reference cannot be lost.
        if len(matching_paragraphs) != 1:
            continue
        paragraph = matching_paragraphs[0]
        parent = paragraph.getparent()
        if parent is None:
            continue
        parent.remove(paragraph)
        changed += 1
    if changed:
        document.save(docx_path)
    return changed


def _float_pdf2docx_isolated_page_counters(document: Any) -> int:
    """Anchor standalone page counters so they cannot occupy a Word page alone."""

    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError as exc:
        raise MissingEngineError("PDF 转 Word 需要 python-docx") from exc

    body = document.element.body
    content_children = [child for child in body if child.tag != qn("w:sectPr")]

    def paragraph_text(paragraph: Any) -> str:
        return " ".join(
            "".join(paragraph.xpath(".//w:t/text()")).replace("\u200b", "").split()
        )

    def next_page_section_properties(start_index: int) -> Any:
        for child in content_children[start_index:]:
            if child.tag != qn("w:p"):
                continue
            properties = child.pPr
            section_properties = properties.sectPr if properties is not None else None
            if section_properties is not None and _section_break_is_next_page(
                section_properties
            ):
                return section_properties
        return body.sectPr

    changed = 0
    for index, child in enumerate(content_children):
        if child.getparent() is None or child.tag != qn("w:p"):
            continue
        properties = child.pPr
        if properties is None:
            continue
        if child.xpath(".//w:drawing|.//w:pict|.//w:object"):
            continue
        text = paragraph_text(child)
        if not _pdf2docx_isolated_page_number_text(text):
            continue

        attached_section_properties = properties.sectPr
        existing_frame = properties.find(qn("w:framePr"))
        split_from_section_carrier = False
        section_carrier = child
        if existing_frame is not None and attached_section_properties is None:
            continue
        if attached_section_properties is not None and _section_break_is_next_page(
            attached_section_properties
        ):
            floating_paragraph = _split_footer_from_next_page_section_carrier(
                document,
                child,
            )
            if floating_paragraph is not child:
                child = floating_paragraph
                properties = child.pPr
                existing_frame = properties.find(qn("w:framePr"))
                split_from_section_carrier = True
            section_properties = attached_section_properties
        else:
            section_properties = next_page_section_properties(index)
        if section_properties is None:
            continue
        page_size = section_properties.find(qn("w:pgSz"))
        page_margins = section_properties.find(qn("w:pgMar"))
        if page_size is None or page_margins is None:
            continue
        try:
            page_width = int(page_size.get(qn("w:w")))
            page_height = int(page_size.get(qn("w:h")))
            left_margin = int(page_margins.get(qn("w:left"), "0"))
            bottom_margin = int(page_margins.get(qn("w:bottom"), "0"))
        except (TypeError, ValueError):
            continue
        if page_width <= 0 or page_height <= 0:
            continue

        font_sizes: list[int] = []
        for size_element in child.xpath(".//w:rPr/w:sz"):
            try:
                font_sizes.append(int(size_element.get(qn("w:val"))))
            except (TypeError, ValueError):
                continue
        if max(font_sizes, default=18) > 24:
            continue
        line_height = max(120, max(font_sizes, default=18) * 12)
        spacing = properties.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            properties.append(spacing)
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")

        frame = existing_frame
        if frame is None:
            frame = OxmlElement("w:framePr")
            properties.insert(0, frame)
        frame.set(qn("w:wrap"), "none")
        frame.set(qn("w:hAnchor"), "page")
        frame.set(qn("w:vAnchor"), "page")
        frame.set(qn("w:x"), str(max(0, left_margin)))
        frame.set(
            qn("w:y"),
            str(max(0, page_height - bottom_margin - line_height - 40)),
        )
        frame.set(qn("w:w"), str(max(1, page_width - left_margin)))
        frame.set(qn("w:hRule"), "auto")
        frame.set(qn("w:anchorLock"), "1")

        for tag_name in (
            "w:keepNext",
            "w:keepLines",
            "w:pageBreakBefore",
            "w:snapToGrid",
        ):
            element = properties.find(qn(tag_name))
            if element is None:
                element = OxmlElement(tag_name)
                properties.append(element)
            element.set(qn("w:val"), "0")
        if split_from_section_carrier:
            section_anchor = section_carrier.getprevious()
            scanned = 0
            drawing_anchor = None
            moved_to_stable_anchor = False
            while section_anchor is not None and scanned < 24:
                scanned += 1
                if section_anchor is child:
                    section_anchor = section_anchor.getprevious()
                    continue
                if section_anchor.tag == qn("w:tbl"):
                    section_anchor.addprevious(child)
                    moved_to_stable_anchor = True
                    break
                if section_anchor.tag == qn("w:p"):
                    anchor_text = paragraph_text(section_anchor)
                    if _pdf2docx_isolated_page_number_text(anchor_text):
                        break
                    if anchor_text:
                        section_anchor.addprevious(child)
                        moved_to_stable_anchor = True
                        break
                    if drawing_anchor is None and section_anchor.xpath(
                        ".//w:drawing|.//w:pict|.//w:object"
                    ):
                        drawing_anchor = section_anchor
                section_anchor = section_anchor.getprevious()
            if not moved_to_stable_anchor and drawing_anchor is not None:
                # If no nearby editable text exists, a drawing paragraph is the
                # best available anchor.  Text is preferred because page-anchored
                # drawings can leave their own empty anchor on the following page.
                if child.getnext() is section_carrier:
                    drawing_anchor.addprevious(child)
        changed += 1
    return changed


def _split_footer_from_next_page_section_carrier(
    document: Any,
    paragraph: Any,
) -> Any:
    """Detach footer content from a paragraph-level next-page section break.

    During page-by-page assembly, the incoming page's section properties are
    normally moved onto the preceding page's final paragraph.  When that final
    paragraph is a real running footer, framing the combined paragraph also
    frames the section break and WPS can paginate the two source pages as three.
    Move the footer payload (including fields, hyperlinks, and drawings) into a
    sibling paragraph, while the original paragraph remains the compact section
    carrier.  Moving rather than recreating the payload preserves relationships.
    """

    from copy import deepcopy

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph

    properties = paragraph.pPr
    section_properties = properties.sectPr if properties is not None else None
    if section_properties is None or not _section_break_is_next_page(
        section_properties
    ):
        return paragraph

    payload = [child for child in paragraph if child is not properties]
    if not payload:
        return paragraph

    floating_paragraph = OxmlElement("w:p")
    floating_properties = deepcopy(properties)
    copied_section_properties = floating_properties.find(qn("w:sectPr"))
    if copied_section_properties is not None:
        floating_properties.remove(copied_section_properties)
    floating_paragraph.append(floating_properties)
    for child in payload:
        floating_paragraph.append(child)
    paragraph.addprevious(floating_paragraph)

    original_frame = properties.find(qn("w:framePr"))
    if original_frame is not None:
        properties.remove(original_frame)
    _compact_section_break_paragraph(Paragraph(paragraph, document._body))
    return floating_paragraph


def _float_pdf2docx_running_footers(document: Any) -> int:
    """Keep page furniture out of the editable document's vertical text flow.

    pdf2docx emits running footers as ordinary final body paragraphs.  Dense
    journal pages can then push only that footer onto a new Word/WPS page.  A
    page-anchored frame preserves the editable footer, links, rules, and page
    number without consuming body height.  Detection is deliberately limited
    to strong page-furniture evidence at the end of a real next-page section.
    """

    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError as exc:
        raise MissingEngineError("PDF 转 Word 需要 python-docx") from exc

    body = document.element.body
    content_children = [child for child in body if child.tag != qn("w:sectPr")]
    sections: list[tuple[list[Any], Any, bool]] = []
    section_start = 0
    for index, child in enumerate(content_children):
        if child.tag != qn("w:p"):
            continue
        properties = child.pPr
        section_properties = properties.sectPr if properties is not None else None
        if section_properties is None or not _section_break_is_next_page(
            section_properties
        ):
            continue
        # A page-by-page hybrid build can place the section properties directly
        # on the preceding page's footer paragraph.  Keep the boundary paragraph
        # in that source page so the footer detector can safely split it later.
        sections.append(
            (content_children[section_start : index + 1], section_properties, False)
        )
        section_start = index + 1
    sections.append((content_children[section_start:], body.sectPr, True))

    def paragraph_text(paragraph: Any) -> str:
        return " ".join(
            "".join(paragraph.xpath(".//w:t/text()")).replace("\u200b", "").split()
        )

    def maximum_font_size(paragraph: Any) -> int:
        sizes: list[int] = []
        for size_element in paragraph.xpath(".//w:rPr/w:sz"):
            try:
                sizes.append(int(size_element.get(qn("w:val"))))
            except (TypeError, ValueError):
                continue
        return max(sizes, default=0)

    def spacing_value(properties: Any, name: str, default: int = 0) -> int:
        spacing = properties.find(qn("w:spacing"))
        if spacing is None:
            return default
        try:
            return int(spacing.get(qn(name), str(default)))
        except (TypeError, ValueError):
            return default

    def graphic_extent_twips(paragraph: Any) -> tuple[int, int]:
        widths: list[int] = []
        heights: list[int] = []
        for extent in paragraph.xpath(".//wp:extent"):
            try:
                # 914400 EMU == 1440 twips, therefore 1 twip == 635 EMU.
                widths.append(max(0, round(int(extent.get("cx")) / 635)))
                heights.append(max(0, round(int(extent.get("cy")) / 635)))
            except (TypeError, ValueError):
                continue
        return max(widths, default=0), max(heights, default=0)

    def is_shallow_wide_footer_graphic(
        paragraph: Any,
        section_properties: Any,
    ) -> bool:
        if paragraph_text(paragraph):
            return False
        if len(paragraph.xpath(".//wp:inline")) != 1:
            return False
        if paragraph.xpath(
            ".//wp:anchor|.//w:pict|.//w:object|.//w:br|.//w:cr|"
            ".//w:fldChar|.//w:fldSimple|.//w:hyperlink"
        ):
            return False
        extents = paragraph.xpath(".//wp:extent")
        if len(extents) != 1:
            return False
        page_size = section_properties.find(qn("w:pgSz"))
        if page_size is None:
            return False
        try:
            page_width = int(page_size.get(qn("w:w")))
            page_height = int(page_size.get(qn("w:h")))
        except (TypeError, ValueError):
            return False
        graphic_width, graphic_height = graphic_extent_twips(paragraph)
        return bool(
            page_width > 0
            and page_height > 0
            and graphic_width >= page_width * 0.72
            and 20 <= graphic_height <= page_height * 0.12
        )

    def graphic_footer_body_is_dense(
        section_children: Iterable[Any],
        footer_ids: set[int],
    ) -> bool:
        text_lengths: list[int] = []
        for paragraph in section_children:
            if id(paragraph) in footer_ids or paragraph.tag != qn("w:p"):
                continue
            properties = paragraph.pPr
            if properties is not None and properties.sectPr is not None:
                continue
            text = paragraph_text(paragraph)
            if text:
                text_lengths.append(len(text))
        if not text_lengths:
            return False
        total_characters = sum(text_lengths)
        average_characters = total_characters / len(text_lengths)
        return bool(
            total_characters >= _PDF2DOCX_DENSE_GRAPHIC_FOOTER_MIN_TEXT_CHARACTERS
            and (
                max(text_lengths)
                >= _PDF2DOCX_DENSE_GRAPHIC_FOOTER_MIN_PARAGRAPH_CHARACTERS
                or average_characters
                >= _PDF2DOCX_DENSE_GRAPHIC_FOOTER_MIN_AVERAGE_CHARACTERS
            )
        )

    candidates: list[tuple[list[Any], Any, bool, list[Any], str, bool, bool]] = []
    signature_counts: Counter[str] = Counter()
    graphic_footer_candidates = 0
    for section_children, section_properties, is_final_section in sections:
        if not section_children:
            continue
        if any(
            child.xpath(
                ".//wp:anchor/wp:docPr[@descr='LayoutLoom running footer' or "
                "@descr='DocuForge running footer']"
            )
            for child in section_children
            if child.tag == qn("w:p")
        ):
            continue

        footer_cluster: list[Any] = []
        trailing_artifacts: list[Any] = []
        has_strong_evidence = False
        graphic_only_footer = False
        for child in reversed(section_children):
            if len(trailing_artifacts) >= 2 or child.tag != qn("w:p"):
                break
            properties = child.pPr
            if properties is None:
                break
            is_boundary_carrier = bool(
                child is section_children[-1]
                and properties.sectPr is section_properties
            )
            if properties.sectPr is not None and not is_boundary_carrier:
                break
            if properties.find(qn("w:framePr")) is not None:
                # Already normalized: do not apply page-fit slack a second time.
                footer_cluster = []
                break
            text = paragraph_text(child)
            has_graphic = bool(child.xpath(".//w:drawing|.//w:pict|.//w:object"))
            if not text:
                # A normal pdf2docx page boundary is an empty sectPr carrier
                # immediately after the footer.  It must not hide the preceding
                # marker from the reverse scan.
                if is_boundary_carrier and not has_graphic:
                    continue
                if has_graphic and is_shallow_wide_footer_graphic(
                    child, section_properties
                ):
                    footer_cluster = [child]
                    has_strong_evidence = True
                    graphic_only_footer = True
                    break
                break
            if len(text) > 240:
                break
            if maximum_font_size(child) > 20:
                break
            before = spacing_value(properties, "w:before")
            # PDF text layers and pdf2docx frequently preserve page counters as
            # full-width digits (for example ``—４—``).  Normalize compatibility
            # characters before matching so these counters receive the same
            # page-anchored treatment as plain ``4`` or ``4 / 20`` footers.
            normalized_text = unicodedata.normalize("NFKC", text)
            has_marker = bool(_PDF2DOCX_PAGE_FURNITURE_PATTERN.search(normalized_text))
            if has_marker:
                footer_cluster = [*trailing_artifacts, child]
                has_strong_evidence = True
                # Only the first strong marker from the physical page bottom is
                # eligible.  Continuing upward would absorb a figure-specific
                # DOI, caption, or bottom drawing into the running footer.
                break
            # Some publishers place a short invisible/logo-font artefact after
            # the real footer (for example ``Vol.:(0123456789)``).  Permit only
            # a very short trailing suffix while searching for the real marker;
            # a drawing or ordinary caption stops detection immediately.
            if has_graphic or before < 80 or len(text) > 48:
                break
            trailing_artifacts.append(child)

        if not footer_cluster or not has_strong_evidence:
            continue

        logical_text = " ".join(
            paragraph_text(paragraph) for paragraph in reversed(footer_cluster)
        )
        signature = ""
        if not graphic_only_footer:
            signature = unicodedata.normalize("NFKC", logical_text).casefold()
            signature = re.sub(r"\d+", "#", signature)
            signature = re.sub(r"\s+", " ", signature).strip()
        clear_page_number = bool(
            _PDF2DOCX_CLEAR_PAGE_NUMBER_PATTERN.search(
                unicodedata.normalize("NFKC", logical_text)
            )
        )
        candidates.append(
            (
                section_children,
                section_properties,
                is_final_section,
                footer_cluster,
                signature,
                clear_page_number,
                graphic_only_footer,
            )
        )
        if graphic_only_footer:
            graphic_footer_candidates += 1
        elif signature and has_strong_evidence:
            signature_counts[signature] += 1

    def anchor_graphic_footer_to_page(
        section_children: list[Any],
        footer_paragraph: Any,
        *,
        page_width: int,
        page_height: int,
        left_margin: int,
        right_margin: int,
        bottom_margin: int,
        relative_height: int,
    ) -> bool:
        inline_elements = footer_paragraph.xpath(".//wp:inline")
        if len(inline_elements) != 1:
            return False
        inline = inline_elements[0]
        graphic_width, graphic_height = graphic_extent_twips(footer_paragraph)
        if graphic_width <= 0 or graphic_height <= 0:
            return False

        anchor_paragraph = next(
            (
                paragraph
                for paragraph in section_children
                if paragraph is not footer_paragraph
                and paragraph.getparent() is not None
                and paragraph.tag == qn("w:p")
                and (paragraph.pPr is None or paragraph.pPr.sectPr is None)
            ),
            None,
        )
        if anchor_paragraph is None:
            first_attached = next(
                (child for child in section_children if child.getparent() is not None),
                None,
            )
            if first_attached is None:
                return False
            anchor_paragraph = OxmlElement("w:p")
            anchor_properties = OxmlElement("w:pPr")
            anchor_spacing = OxmlElement("w:spacing")
            anchor_spacing.set(qn("w:before"), "0")
            anchor_spacing.set(qn("w:after"), "0")
            anchor_spacing.set(qn("w:line"), "20")
            anchor_spacing.set(qn("w:lineRule"), "exact")
            anchor_properties.append(anchor_spacing)
            anchor_paragraph.append(anchor_properties)
            first_attached.addprevious(anchor_paragraph)
            section_children.insert(0, anchor_paragraph)

        properties = footer_paragraph.get_or_add_pPr()
        indentation = properties.find(qn("w:ind"))

        def indentation_value(*names: str) -> int:
            if indentation is None:
                return 0
            for name in names:
                raw_value = indentation.get(qn(name))
                if raw_value is None:
                    continue
                try:
                    return int(raw_value)
                except (TypeError, ValueError):
                    continue
            return 0

        left_indent = indentation_value("w:left", "w:start")
        right_indent = indentation_value("w:right", "w:end")
        content_left = max(0, left_margin + left_indent)
        content_right = max(content_left, page_width - right_margin - right_indent)
        alignment = properties.find(qn("w:jc"))
        alignment_value = (
            str(alignment.get(qn("w:val")) or "").casefold()
            if alignment is not None
            else ""
        )
        if alignment_value in {"center", "both"}:
            x_offset = content_left + max(
                0,
                (content_right - content_left - graphic_width) // 2,
            )
        elif alignment_value in {"right", "end"}:
            x_offset = max(content_left, content_right - graphic_width)
        else:
            x_offset = content_left
        y_offset = max(0, page_height - bottom_margin - graphic_height)

        drawing = inline.getparent()
        run = drawing
        while run is not None and run.tag != qn("w:r"):
            run = run.getparent()
        if drawing is None or run is None:
            return False
        anchor_paragraph.append(run)
        _inline_picture_to_page_anchor(
            inline,
            x_offset=max(0, x_offset) * 635,
            y_offset=y_offset * 635,
            relative_height=relative_height,
        )
        anchored_doc_properties = anchor_paragraph.xpath(".//wp:anchor/wp:docPr")
        if anchored_doc_properties:
            anchored_doc_properties[-1].set("descr", "LayoutLoom running footer")

        parent = footer_paragraph.getparent()
        if parent is not None and all(
            child is properties for child in footer_paragraph
        ):
            parent.remove(footer_paragraph)
        return True

    graphic_repeat_threshold = max(2, min(4, math.ceil(len(sections) * 0.35)))
    changed = 0
    for (
        section_children,
        section_properties,
        is_final_section,
        footer_cluster,
        signature,
        clear_page_number,
        graphic_only_footer,
    ) in candidates:
        repeated_furniture = bool(
            (
                graphic_only_footer
                and graphic_footer_candidates >= graphic_repeat_threshold
            )
            or (signature and signature_counts.get(signature, 0) >= 2)
        )
        if not repeated_furniture and not clear_page_number:
            continue
        # The final section has no following source page to create an
        # inter-page blank.  Only normalize it when it clearly belongs to the
        # repeated running furniture or contains an explicit page counter.
        if is_final_section and not repeated_furniture and not clear_page_number:
            continue

        resolved_footer_cluster: list[Any] = []
        for paragraph in footer_cluster:
            properties = paragraph.pPr
            if properties is not None and properties.sectPr is not None:
                floating_paragraph = _split_footer_from_next_page_section_carrier(
                    document,
                    paragraph,
                )
                if floating_paragraph is not paragraph:
                    carrier_index = section_children.index(paragraph)
                    section_children.insert(carrier_index, floating_paragraph)
                    paragraph = floating_paragraph
            resolved_footer_cluster.append(paragraph)
        footer_cluster = resolved_footer_cluster

        page_size = section_properties.find(qn("w:pgSz"))
        page_margins = section_properties.find(qn("w:pgMar"))
        if page_size is None or page_margins is None:
            continue
        try:
            page_width = int(page_size.get(qn("w:w")))
            page_height = int(page_size.get(qn("w:h")))
            left_margin = int(page_margins.get(qn("w:left"), "0"))
            right_margin = int(page_margins.get(qn("w:right"), "0"))
            bottom_margin = int(page_margins.get(qn("w:bottom"), "0"))
        except (TypeError, ValueError):
            continue

        original_footer_before = min(
            (
                spacing_value(paragraph.get_or_add_pPr(), "w:before")
                for paragraph in footer_cluster
            ),
            default=0,
        )
        maximum_footer_graphic_height = max(
            (graphic_extent_twips(paragraph)[1] for paragraph in footer_cluster),
            default=0,
        )
        if graphic_only_footer:
            anchored_graphics = 0
            for paragraph in footer_cluster:
                if anchor_graphic_footer_to_page(
                    section_children,
                    paragraph,
                    page_width=page_width,
                    page_height=page_height,
                    left_margin=left_margin,
                    right_margin=right_margin,
                    bottom_margin=bottom_margin,
                    relative_height=251659264 + changed + anchored_graphics,
                ):
                    anchored_graphics += 1
            if anchored_graphics != len(footer_cluster):
                continue
            changed += anchored_graphics
        else:
            y_cursor = max(0, page_height - bottom_margin - 40)
            for paragraph in footer_cluster:
                properties = paragraph.get_or_add_pPr()
                spacing = properties.find(qn("w:spacing"))
                if spacing is None:
                    spacing = OxmlElement("w:spacing")
                    properties.append(spacing)
                line_rule = str(spacing.get(qn("w:lineRule")) or "auto").casefold()
                raw_line_height = spacing_value(properties, "w:line", 0)
                if line_rule in {"exact", "atleast"} and raw_line_height > 0:
                    line_height = raw_line_height
                else:
                    line_height = max(
                        120,
                        maximum_font_size(paragraph) * 12,
                        raw_line_height,
                    )
                spacing.set(qn("w:before"), "0")
                spacing.set(qn("w:after"), "0")

                frame = OxmlElement("w:framePr")
                frame.set(qn("w:wrap"), "none")
                frame.set(qn("w:hAnchor"), "page")
                frame.set(qn("w:vAnchor"), "page")
                frame.set(qn("w:x"), str(max(0, left_margin)))
                frame.set(qn("w:y"), str(max(0, y_cursor - line_height)))
                frame.set(
                    qn("w:w"),
                    # pdf2docx footer tab stops are frequently measured nearly to
                    # the physical page edge (for example a right-aligned page
                    # number).  Let the floating frame use the right margin too so
                    # that a final ``1 / 19`` does not wrap onto two lines in WPS.
                    str(max(1, page_width - left_margin)),
                )
                frame.set(qn("w:hRule"), "auto")
                frame.set(qn("w:anchorLock"), "1")
                properties.insert(0, frame)

                for tag_name in (
                    "w:keepNext",
                    "w:keepLines",
                    "w:pageBreakBefore",
                    "w:snapToGrid",
                ):
                    element = properties.find(qn(tag_name))
                    if element is None:
                        element = OxmlElement(tag_name)
                        properties.append(element)
                    element.set(qn("w:val"), "0")
                y_cursor = max(0, y_cursor - line_height - 20)
                changed += 1

        # Word and WPS need a small compatibility reserve on dense pages.
        # pdf2docx often represents the page's initial Y offset as an otherwise
        # empty paragraph whose gap is stored in ``space-after``.  The earlier
        # text-only reserve skipped that geometry paragraph completely, so an
        # image-heavy page could still push its footer and section marker onto a
        # nearly blank page.  Tighten at most three points of geometry-only gap
        # first; this moves the whole reconstructed page uniformly and preserves
        # all internal text/image spacing.  The ordinary bottom-up text-gap
        # reserve remains independently capped below.
        geometry_slack = _PDF2DOCX_WORD_FLOW_GEOMETRY_SLACK_TWIPS
        for paragraph in reversed(section_children):
            if geometry_slack <= 0:
                break
            if paragraph.tag != qn("w:p"):
                continue
            properties = paragraph.pPr
            if properties is None or properties.sectPr is not None:
                continue
            if properties.find(qn("w:framePr")) is not None:
                continue
            if paragraph_text(paragraph):
                continue
            if paragraph.xpath(
                ".//w:drawing|.//w:pict|.//w:object|.//w:br|.//w:cr|"
                ".//w:fldChar|.//w:fldSimple|.//w:hyperlink|.//w:sdt|"
                ".//w:footnoteReference|.//w:endnoteReference|"
                ".//w:commentReference|.//w:annotationRef"
            ):
                continue
            spacing = properties.find(qn("w:spacing"))
            if spacing is None:
                continue
            for spacing_name in ("w:after", "w:before"):
                try:
                    value = int(spacing.get(qn(spacing_name), "0"))
                except (TypeError, ValueError):
                    value = 0
                reduction = min(max(0, value), geometry_slack)
                if reduction <= 0:
                    continue
                spacing.set(qn(spacing_name), str(value - reduction))
                geometry_slack -= reduction
                changed += 1
                if geometry_slack <= 0:
                    break

        # Reduce existing text gaps from the bottom upward, never touch line
        # height, and cap the total correction per source page. This is only
        # done once, in sections where a repeated/numbered footer was newly
        # anchored.
        footer_ids = {id(paragraph) for paragraph in footer_cluster}
        remaining_slack = _PDF2DOCX_WORD_FLOW_SECTION_SLACK_TWIPS
        dense_graphic_footer = bool(
            graphic_only_footer
            and original_footer_before <= _PDF2DOCX_DENSE_GRAPHIC_FOOTER_BEFORE_TWIPS
            and graphic_footer_body_is_dense(section_children, footer_ids)
        )
        if dense_graphic_footer:
            remaining_slack = min(
                _PDF2DOCX_DENSE_GRAPHIC_FOOTER_SLACK_TWIPS,
                max(remaining_slack, maximum_footer_graphic_height),
            )
        for paragraph in reversed(section_children):
            if remaining_slack <= 0:
                break
            if id(paragraph) in footer_ids or paragraph.tag != qn("w:p"):
                continue
            properties = paragraph.pPr
            if properties is None or properties.sectPr is not None:
                continue
            if properties.find(qn("w:framePr")) is not None:
                continue
            text = paragraph_text(paragraph)
            if not text:
                continue
            if paragraph.xpath(".//w:drawing|.//w:pict|.//w:object") and not (
                paragraph.xpath(
                    ".//wp:anchor/wp:docPr[@descr='LayoutLoom running footer' or "
                    "@descr='DocuForge running footer']"
                )
            ):
                continue
            spacing = properties.find(qn("w:spacing"))
            if spacing is None:
                continue
            try:
                before = int(spacing.get(qn("w:before"), "0"))
            except (TypeError, ValueError):
                continue
            reduction = min(
                before,
                (
                    remaining_slack
                    if dense_graphic_footer
                    else _PDF2DOCX_WORD_FLOW_BEFORE_SLACK_TWIPS
                ),
                remaining_slack,
            )
            tightened = before - reduction
            if tightened != before:
                spacing.set(qn("w:before"), str(tightened))
                remaining_slack -= reduction
                changed += 1

    return changed


def _normalize_pdf2docx_table_grids(document: Any) -> int:
    """Repair reliable pdf2docx table grids from cell width constraints.

    ``python-docx`` initially creates every table with equal columns spanning the
    available page width.  pdf2docx then writes the source-derived width only to
    each ``tcW`` and leaves that original ``tblGrid`` untouched.  Word and WPS
    prefer the grid for fixed-layout tables, so asymmetric forms can be rendered
    with completely different column proportions.

    Treat every dxa ``tcW`` as an equation over the columns covered by its
    ``gridSpan``.  A table is changed only when those equations uniquely
    determine every column, the rounded solution fits them closely, and it is a
    material improvement over the existing grid.  Ambiguous or inconsistent
    tables are deliberately left alone.
    """

    try:
        import numpy as np
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        return 0

    def positive_integer(element: Any, attribute: str) -> int | None:
        try:
            value = int(element.get(qn(attribute)))
        except (TypeError, ValueError):
            return None
        return value if value > 0 else None

    def row_cells(row: Any, column_count: int) -> list[tuple[Any, int, int]]:
        row_properties = row.find(qn("w:trPr"))
        grid_before = (
            row_properties.find(qn("w:gridBefore"))
            if row_properties is not None
            else None
        )
        column_index = 0
        if grid_before is not None:
            try:
                column_index = max(0, int(grid_before.get(qn("w:val"), "0")))
            except (TypeError, ValueError):
                return []

        positioned: list[tuple[Any, int, int]] = []
        for cell in row.findall(qn("w:tc")):
            cell_properties = cell.find(qn("w:tcPr"))
            grid_span = (
                cell_properties.find(qn("w:gridSpan"))
                if cell_properties is not None
                else None
            )
            try:
                span = max(
                    1,
                    (
                        int(grid_span.get(qn("w:val"), "1"))
                        if grid_span is not None
                        else 1
                    ),
                )
            except (TypeError, ValueError):
                return []
            if column_index + span > column_count:
                return []
            positioned.append((cell, column_index, span))
            column_index += span
        return positioned

    def constraint_errors(
        widths: Iterable[int],
        constraints: Iterable[tuple[int, int, float]],
    ) -> tuple[float, float]:
        resolved_widths = tuple(widths)
        relative_errors: list[float] = []
        for start, span, target_width in constraints:
            predicted = sum(resolved_widths[start : start + span])
            relative_errors.append(
                abs(float(predicted) - target_width) / max(target_width, 1.0)
            )
        if not relative_errors:
            return float("inf"), float("inf")
        root_mean_square = math.sqrt(
            sum(error * error for error in relative_errors) / len(relative_errors)
        )
        return root_mean_square, max(relative_errors)

    changed = 0
    for table in document.element.body.xpath(".//w:tbl"):
        table_properties = table.find(qn("w:tblPr"))
        table_layout = (
            table_properties.find(qn("w:tblLayout"))
            if table_properties is not None
            else None
        )
        if table_layout is None or table_layout.get(qn("w:type")) != "fixed":
            continue
        table_grid = table.find(qn("w:tblGrid"))
        if table_grid is None:
            continue
        grid_columns = table_grid.findall(qn("w:gridCol"))
        column_count = len(grid_columns)
        if column_count < 2 or column_count > 128:
            continue

        current_widths: list[int] = []
        for grid_column in grid_columns:
            width = positive_integer(grid_column, "w:w")
            if width is None:
                current_widths = []
                break
            current_widths.append(width)
        if len(current_widths) != column_count:
            continue

        constraint_widths: dict[tuple[int, int], list[int]] = {}
        raw_constraints: list[tuple[int, int, float]] = []
        positioned_rows: list[list[tuple[Any, int, int]]] = []
        invalid_structure = False
        for row in table.findall(qn("w:tr")):
            positioned = row_cells(row, column_count)
            if not positioned:
                invalid_structure = True
                break
            positioned_rows.append(positioned)
            for cell, start, span in positioned:
                cell_properties = cell.find(qn("w:tcPr"))
                cell_width = (
                    cell_properties.find(qn("w:tcW"))
                    if cell_properties is not None
                    else None
                )
                if cell_width is None or cell_width.get(qn("w:type")) != "dxa":
                    continue
                width = positive_integer(cell_width, "w:w")
                if width is None:
                    continue
                constraint_widths.setdefault((start, span), []).append(width)
                raw_constraints.append((start, span, float(width)))
        if invalid_structure or not constraint_widths:
            continue

        solving_constraints = [
            (start, span, float(median(widths)))
            for (start, span), widths in sorted(constraint_widths.items())
        ]
        if len(solving_constraints) < column_count:
            continue

        matrix = np.zeros((len(solving_constraints), column_count), dtype=float)
        targets = np.zeros(len(solving_constraints), dtype=float)
        for index, (start, span, target_width) in enumerate(solving_constraints):
            matrix[index, start : start + span] = 1.0
            targets[index] = target_width
        if int(np.linalg.matrix_rank(matrix)) < column_count:
            continue

        try:
            solution, *_ = np.linalg.lstsq(matrix, targets, rcond=None)
        except np.linalg.LinAlgError:
            continue
        candidate_widths = [int(round(float(value))) for value in solution]
        if any(width < 20 for width in candidate_widths):
            continue

        candidate_rms, candidate_max = constraint_errors(
            candidate_widths, raw_constraints
        )
        current_rms, current_max = constraint_errors(current_widths, raw_constraints)
        if candidate_rms > 0.008 or candidate_max > 0.02:
            continue
        if current_max <= 0.02:
            continue
        if candidate_rms >= current_rms * 0.35:
            continue

        for grid_column, width in zip(grid_columns, candidate_widths):
            grid_column.set(qn("w:w"), str(width))

        for positioned in positioned_rows:
            for cell, start, span in positioned:
                cell_properties = cell.find(qn("w:tcPr"))
                cell_width = (
                    cell_properties.find(qn("w:tcW"))
                    if cell_properties is not None
                    else None
                )
                if cell_width is None or cell_width.get(qn("w:type")) != "dxa":
                    continue
                cell_width.set(
                    qn("w:w"), str(sum(candidate_widths[start : start + span]))
                )

        table_width = table_properties.find(qn("w:tblW"))
        if table_width is None:
            table_width = OxmlElement("w:tblW")
            table_properties.insert_element_before(
                table_width,
                "w:jc",
                "w:tblCellSpacing",
                "w:tblInd",
                "w:tblBorders",
                "w:shd",
                "w:tblLayout",
                "w:tblCellMar",
                "w:tblLook",
                "w:tblCaption",
                "w:tblDescription",
                "w:tblPrChange",
            )
        table_width.set(qn("w:type"), "dxa")
        table_width.set(qn("w:w"), str(sum(candidate_widths)))
        changed += 1

    return changed


def _postprocess_pdf2docx_document(document: Any) -> int:
    """Run the shared, idempotent Word compatibility pipeline."""

    mapped = _normalize_docx_fonts(document)
    isolated_page_counters = _float_pdf2docx_isolated_page_counters(document)
    floated_footers = _float_pdf2docx_running_footers(document)
    merged_orphans = _merge_cross_page_english_orphan_fragments(document)
    _normalize_empty_next_page_section_paragraphs(document)
    stabilized = _stabilize_pdf2docx_paragraph_layout(document)
    table_grids = _normalize_pdf2docx_table_grids(document)
    isolated_page_counters += _float_pdf2docx_isolated_page_counters(document)
    floated_footers += _float_pdf2docx_running_footers(document)
    duplicate_page_counters = _deduplicate_adjacent_pdf2docx_page_counters(document)
    _normalize_empty_next_page_section_paragraphs(document)
    stabilized += _stabilize_pdf2docx_paragraph_layout(document)
    section_carriers = any(
        child.pPr is not None
        and child.pPr.sectPr is not None
        and _section_break_is_next_page(child.pPr.sectPr)
        for child in _body_content_children(document)
        if getattr(child, "tag", "").endswith("}p")
    )
    return (
        mapped
        + isolated_page_counters
        + floated_footers
        + duplicate_page_counters
        + merged_orphans
        + stabilized
        + table_grids
        + int(section_carriers)
    )


def _normalize_docx_file_fonts(path: Path) -> int:
    try:
        from docx import Document
    except ImportError as exc:
        raise MissingEngineError("PDF 转 Word 需要 python-docx") from exc
    document = Document(path)
    changed = _postprocess_pdf2docx_document(document)
    if changed:
        document.save(path)
    return changed


def _fixed_layout_region_optimization_worthwhile(stats: Any) -> bool:
    frames = max(0, int(getattr(stats, "editable_frames", 0) or 0))
    pages = max(0, int(getattr(stats, "pages", 0) or 0))
    return bool(
        frames >= _FIXED_LAYOUT_REGION_MIN_FRAMES
        or (
            pages >= _FIXED_LAYOUT_REGION_MIN_PAGES
            and frames >= _FIXED_LAYOUT_REGION_MIN_FRAMES_LONG_DOCUMENT
        )
    )


def _fixed_layout_region_hints(
    pdf_document: Any,
    *,
    dpi: int,
) -> tuple[
    dict[int, tuple[tuple[int, int, int, int], ...]],
    dict[int, tuple[tuple[int, int, int, int], ...]],
]:
    """Return high-confidence formula and table geometry in Word twips.

    Region finalization must not infer display equations from sparse pixels
    alone: table ticks, underlines, and inline math symbols have the same
    visual shape.  The original PDF text/font geometry is still available at
    this point, so reuse the mature PDF detectors and pass only confirmed
    numbered display equations to the Word postprocessor.  Lower-confidence
    cases remain on the legacy fixed-coordinate baseline unless the rendered
    candidate proves safe.
    """

    formula_hints: dict[int, tuple[tuple[int, int, int, int], ...]] = {}
    table_hints: dict[int, tuple[tuple[int, int, int, int], ...]] = {}

    def as_twips(rect: tuple[float, float, float, float], page_rect: Any) -> tuple[int, int, int, int]:
        return (
            round((rect[0] - float(page_rect.x0)) * 20.0),
            round((rect[1] - float(page_rect.y0)) * 20.0),
            round((rect[2] - float(page_rect.x0)) * 20.0),
            round((rect[3] - float(page_rect.y0)) * 20.0),
        )

    for page_index in range(int(pdf_document.page_count)):
        page = pdf_document[page_index]
        page_dict = page.get_text("dict") or {}
        try:
            tables = _detect_pdf_table_regions(
                page,
                page_index,
                page_dict,
                page.get_drawings() or [],
                dpi=int(dpi),
            )
        except Exception:
            tables = []
        table_hints[page_index] = tuple(
            as_twips(region.rect, page.rect) for region in tables
        )

        try:
            formulas = _detect_pdf_formula_regions(
                page_index,
                page_dict,
                page.rect,
                dpi=int(dpi),
                column_layout=(
                    "mixed"
                    if _pdf_page_looks_two_column(page_dict, page.rect)
                    else "auto"
                ),
            )
        except Exception:
            formulas = []

        confirmed: list[tuple[int, int, int, int]] = []
        for formula in formulas:
            if _NUMBERED_FORMULA_REASON not in formula.reasons:
                continue
            formula_area = max(1.0, _rect_area(formula.rect))
            center_x = (formula.rect[0] + formula.rect[2]) / 2.0
            center_y = (formula.rect[1] + formula.rect[3]) / 2.0
            inside_table = False
            for table in tables:
                intersection = (
                    max(0.0, min(formula.rect[2], table.rect[2]) - max(formula.rect[0], table.rect[0]))
                    * max(0.0, min(formula.rect[3], table.rect[3]) - max(formula.rect[1], table.rect[1]))
                )
                if intersection / formula_area >= 0.55 or (
                    table.rect[0] <= center_x <= table.rect[2]
                    and table.rect[1] <= center_y <= table.rect[3]
                ):
                    inside_table = True
                    break
            if not inside_table:
                confirmed.append(as_twips(formula.rect, page.rect))
        formula_hints[page_index] = tuple(confirmed)
    return formula_hints, table_hints


def _filter_fixed_layout_visual_hints(
    visual_hints_by_page: Mapping[
        int, Iterable[tuple[str, int, int, int, int]]
    ]
    | None,
    formula_hints_by_page: Mapping[
        int, Iterable[tuple[int, int, int, int]]
    ]
    | None,
    table_hints_by_page: Mapping[
        int, Iterable[tuple[int, int, int, int]]
    ]
    | None,
) -> dict[int, tuple[tuple[str, int, int, int, int], ...]]:
    """Drop visual spans already covered by a confirmed formula or table."""

    filtered: dict[int, tuple[tuple[str, int, int, int, int], ...]] = {}
    for raw_page_index, raw_hints in (visual_hints_by_page or {}).items():
        try:
            page_index = int(raw_page_index)
        except (TypeError, ValueError):
            continue
        confirmed_regions = (
            tuple((formula_hints_by_page or {}).get(page_index, ()))
            + tuple((table_hints_by_page or {}).get(page_index, ()))
        )
        retained: list[tuple[str, int, int, int, int]] = []
        for raw_hint in raw_hints:
            try:
                kind, raw_x0, raw_y0, raw_x1, raw_y1 = raw_hint
                x0, y0, x1, y1 = (
                    int(raw_x0),
                    int(raw_y0),
                    int(raw_x1),
                    int(raw_y1),
                )
            except (TypeError, ValueError):
                continue
            if x1 <= x0 or y1 <= y0:
                continue
            hint_area = max(1, (x1 - x0) * (y1 - y0))
            center_x = (x0 + x1) / 2.0
            center_y = (y0 + y1) / 2.0
            covered = False
            for region in confirmed_regions:
                try:
                    region_x0, region_y0, region_x1, region_y1 = (
                        int(value) for value in region
                    )
                except (TypeError, ValueError):
                    continue
                intersection = (
                    max(0, min(x1, region_x1) - max(x0, region_x0))
                    * max(0, min(y1, region_y1) - max(y0, region_y0))
                )
                if intersection / hint_area >= 0.45 or (
                    region_x0 <= center_x <= region_x1
                    and region_y0 <= center_y <= region_y1
                ):
                    covered = True
                    break
            if not covered:
                retained.append((str(kind), x0, y0, x1, y1))
        filtered[page_index] = tuple(retained)
    return filtered


def _fixed_layout_expected_editable_text(
    path: Path,
    formula_hints_by_page: Mapping[
        int, Iterable[tuple[int, int, int, int]]
    ],
) -> str:
    """Extract baseline text excluding glyph frames replaced by exact formulas."""

    try:
        from docx import Document
    except ImportError as exc:
        raise MissingEngineError("固定坐标区域文字复检需要 python-docx") from exc
    from .word_flow import _source_pages

    document = Document(path)
    retained: list[str] = []
    for page in _source_pages(document):
        # When the fixed-layout page has a redacted background, the region
        # builder crops confirmed formulas from that background and keeps all
        # editable frames.  The outer adoption baseline must mirror that
        # contract instead of pre-emptively discounting covered text.
        hints = (
            ()
            if page.background is not None
            else tuple(formula_hints_by_page.get(page.index, ()))
        )
        for frame in page.frames:
            excluded = False
            for x0, y0, x1, y1 in hints:
                horizontal = max(0, min(frame.x1, x1) - max(frame.x, x0))
                vertical = max(0, min(frame.y1, y1) - max(frame.y, y0))
                overlap = (horizontal * vertical) / max(
                    1, frame.width * frame.height
                )
                # A small formula island can sit at the centre of a much
                # wider editable frame.  Centre containment alone must not
                # remove that whole frame from the adoption baseline: the
                # region builder deliberately rejects this ambiguous case.
                # Only a hint covering most of the frame is authoritative.
                if overlap >= 0.80:
                    excluded = True
                    break
            if not excluded:
                retained.append(frame.text)
    return "\n".join(retained)


def _docx_raster_pixel_budget(path: Path) -> int:
    total = 0
    try:
        with ZipFile(path) as archive:
            for item in archive.infolist():
                if item.is_dir() or not item.filename.startswith("word/media/"):
                    continue
                try:
                    with Image.open(io.BytesIO(archive.read(item.filename))) as image:
                        total += max(0, int(image.width)) * max(0, int(image.height))
                except Exception:
                    continue
    except (BadZipFile, OSError):
        return 0
    return total


def _docx_document_xml_size(path: Path) -> int:
    try:
        with ZipFile(path) as archive:
            return len(archive.read("word/document.xml"))
    except (BadZipFile, KeyError, OSError):
        return 0


def _fixed_layout_region_object_budget(stats: Any) -> tuple[bool, int, int]:
    """Return safe text-region and total-anchor limits for one candidate."""

    compact_designed_layout = bool(
        stats.pages <= 2 and stats.editable_frames <= 120
    )
    region_ratio = 0.65 if compact_designed_layout else 0.55
    anchor_ratio = 0.85 if compact_designed_layout else 0.75
    maximum_regions = max(
        1,
        math.ceil(stats.editable_frames * region_ratio),
    )
    maximum_anchors = max(
        stats.pages,
        math.ceil((stats.editable_frames + stats.pages) * anchor_ratio),
    )
    return compact_designed_layout, maximum_regions, maximum_anchors


def _optimize_pdf_fixed_layout_docx_for_wps(
    target: Path,
    stats: Any,
    *,
    formula_hints_by_page: Mapping[
        int, Iterable[tuple[int, int, int, int]]
    ]
    | None = None,
    table_hints_by_page: Mapping[
        int, Iterable[tuple[int, int, int, int]]
    ]
    | None = None,
    visual_hints_by_page: Mapping[
        int, Iterable[tuple[str, int, int, int, int]]
    ]
    | None = None,
) -> Any:
    """Adopt a region-level candidate only after structural and WPS checks.

    The legacy per-line coordinate document remains the precision baseline and
    automatic fallback.  Large fixed-layout documents are rebuilt with fewer
    modern DrawingML text regions and cropped visual islands; the candidate is
    committed only when the region builder's two-stage validation succeeds and
    its structure is materially lighter.
    """

    from docuforge.runner import report_progress

    if not _fixed_layout_region_optimization_worthwhile(stats):
        return stats
    try:
        from docx import Document

        from .pdf_word_layout import FixedLayoutBuildStats
        from .word_region import build_region_compatible_word
        from .wps import detect_wps_engines

        writer = detect_wps_engines()["writer"]
        if not writer.available:
            return stats
        retained_visual_hints = _filter_fixed_layout_visual_hints(
            (
                visual_hints_by_page
                if visual_hints_by_page is not None
                else getattr(stats, "visual_hints_by_page", {})
            ),
            formula_hints_by_page,
            table_hints_by_page,
        )
        source_pixels = _docx_raster_pixel_budget(target)
        source_xml_size = _docx_document_xml_size(target)
        with _temporary_working_directory(
            prefix="docuforge-fixed-region-finalize-"
        ) as folder:
            candidate = folder / "region-optimized.docx"
            timeout = max(240.0, min(600.0, float(stats.pages) * 18.0))
            build_region_compatible_word(
                target,
                candidate,
                verification_engine="wps",
                timeout=timeout,
                overwrite=True,
                max_mean_difference=28.0,
                max_changed_fraction=0.35,
                preserve_editable_text=False,
                normalize_text=False,
                formula_hints_by_page=formula_hints_by_page,
                table_hints_by_page=table_hints_by_page,
                visual_hints_by_page=retained_visual_hints,
                progress=lambda ratio, message: report_progress(
                    0.89 + min(1.0, max(0.0, float(ratio))) * 0.05,
                    f"固定坐标区域优化：{message}",
                ),
            )
            document = Document(candidate)
            body = document.element.body
            if body.xpath(".//w:pPr/w:framePr") or body.xpath(".//w:fitText"):
                raise ValidationError("区域优化候选仍包含旧式逐行定位节点")
            region_count = len(
                body.xpath(
                    ".//*[namespace-uri()='http://schemas.microsoft.com/office/word/2010/wordprocessingShape' "
                    "and local-name()='txbx']"
                )
            )
            (
                compact_designed_layout,
                maximum_regions,
                maximum_anchors,
            ) = _fixed_layout_region_object_budget(stats)
            if region_count < 1 or region_count > maximum_regions:
                raise ValidationError(
                    "区域优化候选未能充分减少定位对象："
                    f"原 {stats.editable_frames} 个，候选 {region_count} 个"
                )
            anchor_count = len(body.xpath(".//wp:anchor"))
            if anchor_count > maximum_anchors:
                raise ValidationError(
                    "区域优化候选的总定位对象仍然过多："
                    f"候选 {anchor_count} 个，安全上限 {maximum_anchors} 个"
                )
            source_text = (
                _fixed_layout_expected_editable_text(
                    target,
                    formula_hints_by_page,
                )
                if formula_hints_by_page is not None
                else _extract_docx_text(target)
            )
            candidate_text = _extract_docx_text(candidate)
            sequence_retention = _text_sequence_coverage(
                source_text,
                candidate_text,
            )
            word_retention = _english_word_multiset_recall(
                source_text,
                candidate_text,
            )
            adjacent_retention = _adjacent_english_word_coverage(
                source_text,
                candidate_text,
            )
            source_characters = Counter(_normalize_validation_text(source_text))
            candidate_characters = Counter(
                _normalize_validation_text(candidate_text)
            )
            character_retention = (
                sum((source_characters & candidate_characters).values())
                / sum(source_characters.values())
                if source_characters
                else 1.0
            )
            sequence_limit = (
                _FIXED_LAYOUT_REGION_DESIGNED_MIN_SEQUENCE_RETENTION
                if compact_designed_layout
                else _FIXED_LAYOUT_REGION_MIN_SEQUENCE_RETENTION
            )
            adjacent_limit = (
                _FIXED_LAYOUT_REGION_DESIGNED_MIN_ADJACENT_RETENTION
                if compact_designed_layout
                else _FIXED_LAYOUT_REGION_MIN_ADJACENT_WORD_RETENTION
            )
            if (
                character_retention < _FIXED_LAYOUT_REGION_MIN_CHARACTER_RETENTION
                or sequence_retention < sequence_limit
                or word_retention < _FIXED_LAYOUT_REGION_MIN_WORD_RETENTION
                or adjacent_retention < adjacent_limit
            ):
                raise ValidationError(
                    "区域优化候选的可编辑文字相对原固定版发生回退："
                    f"字符保留 {character_retention:.1%}、"
                    f"字符序列 {sequence_retention:.1%}、"
                    f"英文词 {word_retention:.1%}、"
                    f"相邻词序 {adjacent_retention:.1%}"
                )
            candidate_pixels = _docx_raster_pixel_budget(candidate)
            if (
                source_pixels > 0
                and candidate_pixels > math.ceil(source_pixels * 1.05)
            ):
                raise ValidationError("区域优化候选增加了视觉层解码像素")
            candidate_xml_size = _docx_document_xml_size(candidate)
            # Modern DrawingML text boxes are verbose in XML even when they
            # replace many legacy line frames.  On a one- or two-page designed
            # document the absolute XML is still tiny, so allow the checked
            # candidate to grow to 2x while retaining the stricter file-size,
            # anchor, pixel, text, and WPS-render gates.
            xml_growth_limit = 2.00 if compact_designed_layout else 1.05
            if (
                source_xml_size > 0
                and candidate_xml_size
                > math.ceil(source_xml_size * xml_growth_limit)
            ):
                raise ValidationError("区域优化候选增加了 Word 主体结构体积")
            if candidate.stat().st_size > math.ceil(target.stat().st_size * 1.05):
                raise ValidationError("区域优化候选增加了最终文件体积")
            with atomic_output(target) as temporary:
                shutil.copy2(candidate, temporary)
        report_progress(
            0.94,
            "固定坐标区域优化已通过 WPS 二重检查，已减少逐行定位节点",
        )
        return FixedLayoutBuildStats(
            pages=stats.pages,
            editable_frames=stats.editable_frames,
            editable_spans=stats.editable_spans,
            visual_spans=stats.visual_spans,
            region_optimized=True,
            region_text_boxes=region_count,
            visual_hints_by_page=getattr(stats, "visual_hints_by_page", {}),
        )
    except CancelledError:
        raise
    except Exception:
        report_progress(
            0.94,
            "区域候选未通过全部检查，已自动保留原固定坐标高精度结果",
        )
        return stats


def _build_pdf_fixed_layout_docx(
    source: Path,
    target: Path,
    *,
    password: str | None,
    dpi: int,
    page_plans: Mapping[int, Any] | None = None,
) -> Any:
    """Build and post-process the layout-critical coordinate-editable path."""

    from docuforge.runner import report_progress

    try:
        from docx import Document
    except ImportError as exc:
        raise MissingEngineError("PDF 转 Word 需要 python-docx") from exc
    from .pdf_word_layout import build_fixed_layout_docx

    installed_fonts = _installed_word_font_keys()

    def resolve_font(font_name: str, east_asia: bool) -> str:
        return _resolve_word_font_name(
            font_name,
            east_asia=east_asia,
            installed_fonts=installed_fonts,
        )

    _pymupdf, pdf_document = _open_pymupdf_document(source, password)
    stats = None
    formula_hints_by_page: dict[
        int, tuple[tuple[int, int, int, int], ...]
    ] = {}
    table_hints_by_page: dict[
        int, tuple[tuple[int, int, int, int], ...]
    ] = {}
    try:
        with atomic_output(target) as temporary:
            stats = build_fixed_layout_docx(
                pdf_document,
                temporary,
                dpi=int(dpi),
                page_plans=page_plans,
                repair_text=_repair_known_pdf_text_encoding,
                font_requires_visual=_font_requires_visual_fallback,
                math_font=_math_font,
                suspicious_text=_has_suspicious_pdf_characters,
                resolve_font=resolve_font,
                progress=lambda ratio, message: report_progress(
                    0.10 + ratio * 0.76,
                    message,
                ),
            )
            report_progress(0.88, "稳定固定版面 Word 兼容性")
            document = Document(temporary)
            _postprocess_pdf2docx_document(document)
            document.save(temporary)
            if not temporary.is_file() or temporary.stat().st_size == 0:
                raise ValidationError("固定版面 PDF 转 Word 没有生成有效文件")
        if _fixed_layout_region_optimization_worthwhile(stats):
            formula_hints_by_page, table_hints_by_page = (
                _fixed_layout_region_hints(
                    pdf_document,
                    dpi=int(dpi),
                )
            )
    finally:
        pdf_document.close()
    assert stats is not None
    return _optimize_pdf_fixed_layout_docx_for_wps(
        target,
        stats,
        formula_hints_by_page=formula_hints_by_page,
        table_hints_by_page=table_hints_by_page,
        visual_hints_by_page=getattr(stats, "visual_hints_by_page", {}),
    )


def _convert_pdf_to_fixed_editable_docx(
    source: Path,
    target: Path,
    *,
    password: str | None,
    low_quality_policy: str,
    profile: _PdfWordLayoutProfile,
    dpi: int = 300,
) -> None:
    """Convert a paper or designed page with editable coordinate text frames."""

    from docuforge.runner import report_progress

    page_texts, pages_without_text = _inspect_pdf_text_layers(source, password)
    build_stats = _build_pdf_fixed_layout_docx(
        source,
        target,
        password=password,
        dpi=int(dpi),
    )

    pending_warnings: list[str] = []
    source_text = "\n".join(page_texts)
    output_text = _extract_docx_text(target)
    coverage = _text_sequence_coverage(source_text, output_text)
    english_word_recall = _english_word_multiset_recall(source_text, output_text)
    adjacent_word_coverage = _adjacent_english_word_coverage(source_text, output_text)
    quality_details = (
        f"字符序列 {coverage:.0%}、英文词召回 {english_word_recall:.0%}、"
        f"相邻词序 {adjacent_word_coverage:.0%}"
    )
    boundary_recovery_pass = _pdf_english_boundary_recovery_pass(
        source_text,
        output_text,
        character_coverage=coverage,
    )
    low_quality = not boundary_recovery_pass and bool(
        coverage < _MIN_EDITABLE_SEQUENCE_COVERAGE
        or english_word_recall < _MIN_EDITABLE_ENGLISH_WORD_RECALL
        or adjacent_word_coverage < _MIN_EDITABLE_ADJACENT_WORD_COVERAGE
    )
    layout_quality_reasons = _pdf2docx_docx_layout_quality_reasons(target)
    if (low_quality or layout_quality_reasons) and low_quality_policy == "discard":
        details: list[str] = []
        if low_quality:
            details.append(f"文字完整度自动校验未通过（{quality_details}）")
        if layout_quality_reasons:
            details.append(
                "排版结构自动校验未通过（" f"{'；'.join(layout_quality_reasons)}）"
            )
        target.unlink(missing_ok=True)
        raise ValidationError(
            "固定版面可编辑 Word 的最终成品二次校验未通过："
            f"{'；'.join(details)}。本次结果未保存；也可选择“仍保留并警告”。"
        )
    if low_quality:
        pending_warnings.append(
            f"固定版面可编辑文字校验未通过（{quality_details}）；"
            "已按用户选择保留，请重点复核特殊字体和公式。"
        )
    if layout_quality_reasons:
        pending_warnings.append(
            "固定版面 Word 排版结构复检未通过（"
            f"{'；'.join(layout_quality_reasons)}）；已按用户选择保留。"
        )

    report_progress(0.95, "使用 WPS 复核固定版面实际分页")
    pagination_reason, pagination_warning = _pdf2docx_wps_render_quality_result(
        source,
        target,
        expected_pages=len(page_texts),
        password=password,
        progress=lambda ratio, message: report_progress(
            0.95 + ratio * 0.04,
            message,
        ),
    )
    if pagination_warning:
        pending_warnings.append(pagination_warning)
    if pagination_reason:
        if low_quality_policy == "discard":
            target.unlink(missing_ok=True)
            raise ValidationError(
                "固定版面可编辑 Word 的最终成品分页复检未通过："
                f"{pagination_reason}。本次结果未保存；"
                "也可选择“仍保留并警告”。"
            )
        pending_warnings.append(
            f"固定版面可编辑 Word 的分页复检未通过：{pagination_reason}；"
            "已按用户选择保留。"
        )
    if pages_without_text:
        page_preview = "、".join(str(item) for item in pages_without_text[:12])
        pending_warnings.append(
            f"第 {page_preview} 页没有可靠文字层，已保留为高清视觉背景。"
        )
    if build_stats.region_optimized:
        pending_warnings.append(
            "检测到"
            f"{'、'.join(profile.reasons)}，已自动使用固定坐标区域优化布局："
            f"原 {build_stats.editable_frames} 个逐行定位节点已压缩为 "
            f"{build_stats.region_text_boxes} 个可编辑文字区域；"
            "公式、特殊符号、图片与页面装饰仍按原位置高清保留，"
            "并已通过真实 WPS 版面复检。"
        )
    else:
        pending_warnings.append(
            "检测到"
            f"{'、'.join(profile.reasons)}，已自动使用固定坐标可编辑布局："
            "普通正文可编辑，公式、特殊符号、图片与页面装饰按原位置高清保留。"
        )
    for message in pending_warnings:
        warnings.warn(message, stacklevel=2)


def _convert_pdf_to_fixed_hybrid_docx(
    source: Path,
    target: Path,
    *,
    password: str | None,
    dpi: int,
    low_quality_policy: str,
    assessments: list[_HybridPageAssessment],
    forced_indexes: set[int],
    profile: _PdfWordLayoutProfile,
) -> None:
    """Hybrid fixed layout: editable prose over exact visual islands."""

    from docuforge.runner import report_progress

    from .pdf_word_layout import FixedLayoutPagePlan

    risk_reasons: dict[int, list[str]] = {
        assessment.page_index: list(assessment.reasons)
        for assessment in assessments
        if assessment.reasons
    }
    for page_index in forced_indexes:
        risk_reasons.setdefault(page_index, []).append("用户指定原样保留")
    plans = {
        assessment.page_index: FixedLayoutPagePlan(
            excluded_regions=tuple(
                region.rect
                for region in (
                    *assessment.visual_regions,
                    *assessment.anchored_visual_regions,
                )
            ),
            full_visual=assessment.page_index in risk_reasons,
        )
        for assessment in assessments
    }
    build_stats = _build_pdf_fixed_layout_docx(
        source,
        target,
        password=password,
        dpi=int(dpi),
        page_plans=plans,
    )

    pending_warnings: list[str] = []
    layout_quality_reasons = _pdf2docx_docx_layout_quality_reasons(target)
    if layout_quality_reasons:
        message = (
            "固定版面混合 Word 排版结构复检未通过（"
            f"{'；'.join(layout_quality_reasons)}）。"
        )
        if low_quality_policy == "discard":
            target.unlink(missing_ok=True)
            raise ValidationError(f"{message}本次结果未保存。")
        pending_warnings.append(f"{message}已按用户选择保留。")

    report_progress(0.95, "使用 WPS 复核固定版面实际分页")
    pagination_reason, pagination_warning = _pdf2docx_wps_render_quality_result(
        source,
        target,
        expected_pages=len(assessments),
        password=password,
        progress=lambda ratio, message: report_progress(
            0.95 + ratio * 0.04,
            message,
        ),
    )
    if pagination_warning:
        pending_warnings.append(pagination_warning)
    if pagination_reason:
        if low_quality_policy == "discard":
            target.unlink(missing_ok=True)
            raise ValidationError(
                "固定版面混合 Word 的最终成品二次检测未通过："
                f"{pagination_reason}。本次结果未保存。"
            )
        pending_warnings.append(
            f"固定版面混合 Word 的二次检测未通过：{pagination_reason}；"
            "已按用户选择保留。"
        )

    editable_indexes = [
        assessment.page_index
        for assessment in assessments
        if assessment.page_index not in risk_reasons
    ]
    if editable_indexes:
        source_text = "\n".join(
            assessments[page_index].editable_source_text
            for page_index in editable_indexes
        )
        source_blocks = [
            block
            for page_index in editable_indexes
            for block in assessments[page_index].editable_text_blocks
        ]
        output_text = _extract_docx_text(target)
        coverage = _text_sequence_coverage(source_text, output_text)
        english_word_recall = _english_word_multiset_recall(source_text, output_text)
        adjacent_word_coverage = _block_local_adjacent_english_word_coverage(
            source_blocks,
            output_text,
        )
        boundary_recovery_pass = _pdf_english_boundary_recovery_pass(
            source_text,
            output_text,
            character_coverage=coverage,
        )
        low_quality = not boundary_recovery_pass and bool(
            coverage < _MIN_HYBRID_PAGE_SEQUENCE_COVERAGE
            or english_word_recall < _MIN_HYBRID_PAGE_ENGLISH_WORD_RECALL
            or adjacent_word_coverage < _MIN_HYBRID_PAGE_ADJACENT_WORD_COVERAGE
        )
        if low_quality:
            quality_details = (
                f"字符序列 {coverage:.0%}、英文词召回 {english_word_recall:.0%}、"
                f"相邻词序 {adjacent_word_coverage:.0%}"
            )
            if low_quality_policy == "discard":
                target.unlink(missing_ok=True)
                raise ValidationError(
                    "固定版面混合 Word 的可编辑正文校验未通过（"
                    f"{quality_details}），本次结果未保存。"
                )
            pending_warnings.append(
                f"固定版面混合正文校验未通过（{quality_details}）；"
                "已按用户选择保留。"
            )

    region_summary = _hybrid_region_summary(assessments, risk_reasons)
    if region_summary:
        pending_warnings.append(
            "固定版面混合模式已将以下公式、图表或复杂表格高清保留，"
            f"其余正文保持可编辑：{region_summary}。"
        )
    if risk_reasons:
        pending_warnings.append(
            "以下高风险页面按原样保留："
            f"{_hybrid_visual_page_summary(risk_reasons)}。"
        )
    if build_stats.region_optimized:
        pending_warnings.append(
            "检测到"
            f"{'、'.join(profile.reasons)}，已自动采用一源页一 Word 页的"
            "固定坐标区域优化布局："
            f"原 {build_stats.editable_frames} 个逐行定位节点已压缩为 "
            f"{build_stats.region_text_boxes} 个可编辑文字区域，"
            "避免双栏重排、整页拆分和大面积空白，并已通过真实 WPS 版面复检。"
        )
    else:
        pending_warnings.append(
            "检测到"
            f"{'、'.join(profile.reasons)}，已自动采用一源页一 Word 页的固定坐标混合布局，"
            "避免双栏重排、整页拆分和大面积空白。"
        )
    for message in pending_warnings:
        warnings.warn(message, stacklevel=2)


def _convert_pdf_to_editable_docx(
    source: Path,
    target: Path,
    *,
    password: str | None,
    low_quality_policy: str,
    column_layout: str,
) -> None:
    from docuforge.runner import report_progress

    profile = _pdf_word_layout_profile(
        source,
        password,
        column_layout=column_layout,
    )
    if profile.fixed_layout_recommended:
        report_progress(0.03, "检测到复杂版面，启用固定坐标可编辑重建")
        _convert_pdf_to_fixed_editable_docx(
            source,
            target,
            password=password,
            low_quality_policy=low_quality_policy,
            profile=profile,
            dpi=300,
        )
        return

    try:
        from pdf2docx import Converter
    except ImportError as exc:
        raise MissingEngineError(
            "高精度可编辑模式需要 pdf2docx；请重新运行安装程序补齐依赖"
        ) from exc

    report_progress(0.03, "检查 PDF 文字层")
    page_texts, pages_without_text = _inspect_pdf_text_layers(source, password)
    if len(pages_without_text) == len(page_texts):
        raise ValidationError(
            "该 PDF 没有可编辑文字层，可能是扫描件。请先执行 OCR 识别，"
            "或选择“整篇高清原样（不可编辑）”。"
        )

    pending_warnings: list[str] = []
    with atomic_output(target) as temporary:
        converter = None
        try:
            converter = Converter(str(source), password=password or None)
            settings = _pdf2docx_settings(
                converter,
                column_layout=column_layout,
            )
            report_progress(0.18, f"重建 {len(page_texts)} 页可编辑版面")
            _parse_pdf2docx_pages(
                converter,
                settings=settings,
                column_layout=column_layout,
                start=0,
                end=None,
            )
            source_words = set(_english_words("\n".join(page_texts)))
            _restore_pdf2docx_spaces(
                converter.pages,
                source_words=source_words,
                source_document=converter.fitz_doc,
            )
            _repair_pdf2docx_known_encoding(converter.pages)
            _normalize_pdf2docx_english_widths(converter.pages)
            _normalize_pdf2docx_list_prefixes(converter.pages)
            _normalize_pdf2docx_prose_alignment(converter.pages)
            report_progress(0.76, "写入可编辑 Word 文档")
            converter.make_docx(str(temporary), **settings)
        except Exception as exc:
            raise ValidationError(f"PDF 可编辑版面重建失败：{exc}") from exc
        finally:
            if converter is not None:
                with contextlib.suppress(Exception):
                    converter.close()

        if not temporary.is_file() or temporary.stat().st_size == 0:
            raise ValidationError("PDF 可编辑版面重建没有生成有效的 Word 文件")
        report_progress(0.86, "统一 Word 字体与分页兼容性")
        _normalize_docx_file_fonts(temporary)
        layout_quality_reasons = _pdf2docx_docx_layout_quality_reasons(temporary)
        source_text = "\n".join(page_texts)
        output_text = _extract_docx_text(temporary)
        coverage = _text_sequence_coverage(source_text, output_text)
        english_word_recall = _english_word_multiset_recall(source_text, output_text)
        adjacent_word_coverage = _adjacent_english_word_coverage(
            source_text, output_text
        )
        quality_details = (
            f"字符序列 {coverage:.0%}、英文词召回 {english_word_recall:.0%}、"
            f"相邻词序 {adjacent_word_coverage:.0%}"
        )
        boundary_recovery_pass = _pdf_english_boundary_recovery_pass(
            source_text,
            output_text,
            character_coverage=coverage,
        )
        low_quality = not boundary_recovery_pass and (
            coverage < _MIN_EDITABLE_SEQUENCE_COVERAGE
            or english_word_recall < _MIN_EDITABLE_ENGLISH_WORD_RECALL
            or adjacent_word_coverage < _MIN_EDITABLE_ADJACENT_WORD_COVERAGE
        )
        if (low_quality or layout_quality_reasons) and low_quality_policy == "discard":
            failed_checks: list[str] = []
            if low_quality:
                failed_checks.append(f"文本完整度自动校验未通过（{quality_details}）")
            if layout_quality_reasons:
                failed_checks.append(
                    "排版结构自动校验未通过（" f"{'；'.join(layout_quality_reasons)}）"
                )
            raise ValidationError(
                "转换后的最终成品二次校验未通过："
                f"{'；'.join(failed_checks)}。"
                "为避免输出内容缺失，未保存该结果；可选择保留低质量结果，"
                "或改用高清原样模式。"
            )
        if layout_quality_reasons:
            pending_warnings.append(
                "转换后的最终成品排版结构复检未通过（"
                f"{'；'.join(layout_quality_reasons)}）。"
                "已按用户选择保留该结果，请重点复核中英混排、软换行与缩进。"
            )
        if low_quality:
            pending_warnings.append(
                f"转换后的文本完整度自动校验未通过（{quality_details}）。"
                "已按用户选择保留该结果，请重点复核单词间距、公式和多栏内容。"
            )
        elif (
            coverage < _WARN_EDITABLE_SEQUENCE_COVERAGE
            or english_word_recall < _WARN_EDITABLE_ENGLISH_WORD_RECALL
            or adjacent_word_coverage < _WARN_EDITABLE_ADJACENT_WORD_COVERAGE
        ):
            pending_warnings.append(
                f"可编辑 Word 的质量自动校验为：{quality_details}。"
                "建议重点复核特殊字体、公式和多栏内容。"
            )
        report_progress(0.98, "完成文字完整度校验")
        if pages_without_text:
            page_preview = "、".join(str(item) for item in pages_without_text[:12])
            if len(pages_without_text) > 12:
                page_preview += f" 等 {len(pages_without_text)} 页"
            pending_warnings.append(
                f"第 {page_preview} 页没有文字层；页面图像会尽量保留，"
                "但其中的文字不能单独编辑。"
            )
        report_progress(0.985, "使用 WPS 复核最终实际分页")
        pagination_reason, pagination_warning = _pdf2docx_wps_render_quality_result(
            source,
            temporary,
            expected_pages=len(page_texts),
            password=password,
            progress=lambda ratio, message: report_progress(
                0.985 + ratio * 0.01,
                message,
            ),
        )
        if pagination_warning:
            pending_warnings.append(pagination_warning)
        if pagination_reason:
            if low_quality_policy == "discard":
                raise ValidationError(
                    f"全文可编辑 Word 的最终成品分页复检未通过：{pagination_reason}。"
                    "为避免交付异常分页，本次结果未保存；也可选择“仍保留并警告”。"
                )
            pending_warnings.append(
                f"全文可编辑 Word 的最终成品分页复检未通过：{pagination_reason}。"
                "已按用户选择保留该结果，请人工复核对应页码。"
            )

    for message in pending_warnings:
        warnings.warn(message, stacklevel=2)


def _partition_pdf_redaction_regions(
    page: Any,
    regions: Iterable[_HybridRegion],
    page_rect: Any,
) -> tuple[list[_HybridRegion], list[_HybridRegion]]:
    """Split regions into safe image-removal and pixel-masking phases.

    PyMuPDF applies one image policy to every redaction annotation currently on
    a page.  A page-wide PIXELS policy preserves partially covered placements,
    but also leaves almost-fully-covered source images behind replacement
    images.  Conversely, REMOVE would delete the outside portion of every
    partially covered placement.  Classifying each region against all image
    placements lets the caller apply REMOVE only where it is safe, followed by
    PIXELS for regions that must preserve image content outside their bounds.
    """

    normalized_regions = list(regions)
    region_rects = [region.rect for region in normalized_regions]
    # ``get_image_info()`` reports the full placement matrix.  A PDF clip path
    # can make only part of that placement visible, so comparing a replacement
    # region with the un-clipped bbox incorrectly classifies a visually complete
    # replacement as partial.  The old image then survives PIXELS redaction and
    # pdf2docx may place it again behind or beside the replacement.  Text-dict
    # image blocks expose the actually visible / clipped bbox and are therefore
    # the authoritative geometry for this decision.
    try:
        page_dict = page.get_text("dict") or {}
    except Exception:
        page_dict = {}
    image_rects = [
        image_rect
        for block in page_dict.get("blocks", ())
        if int(block.get("type", 0) or 0) == 1
        if (
            image_rect := _coerce_pdf_rect(
                block.get("bbox", ()),
                page_rect,
            )
        )
        is not None
    ]
    if not image_rects:
        try:
            image_info = page.get_image_info(xrefs=True) or ()
        except Exception:
            return [], normalized_regions
        image_rects = [
            image_rect
            for info in image_info
            if (
                image_rect := _coerce_pdf_rect(
                    info.get("bbox", ()),
                    page_rect,
                )
            )
            is not None
        ]

    fully_covered_images: list[tuple[float, float, float, float]] = []
    partially_covered_images: list[tuple[float, float, float, float]] = []
    for image_rect in image_rects:
        coverage = min(
            1.0,
            sum(
                _rect_intersection_ratio(image_rect, region_rect)
                for region_rect in region_rects
            ),
        )
        if coverage <= 0.0:
            continue
        if coverage < _HYBRID_FULL_IMAGE_REDACTION_RATIO:
            partially_covered_images.append(image_rect)
        else:
            fully_covered_images.append(image_rect)

    remove_regions: list[_HybridRegion] = []
    pixel_regions: list[_HybridRegion] = []
    for region in normalized_regions:
        intersects_partial_image = any(
            _rect_intersection_ratio(image_rect, region.rect) > 0.0
            for image_rect in partially_covered_images
        )
        intersects_fully_covered_image = any(
            _rect_intersection_ratio(image_rect, region.rect) > 0.0
            for image_rect in fully_covered_images
        )
        if intersects_fully_covered_image and not intersects_partial_image:
            remove_regions.append(region)
        else:
            pixel_regions.append(region)
    return remove_regions, pixel_regions


def _contract_pdf_redaction_rect_away_from_external_text(
    page: Any,
    region: _HybridRegion,
) -> tuple[float, float, float, float]:
    """Keep PyMuPDF redaction from deleting a neighboring tall glyph.

    Redaction removes a complete glyph when its bounding box merely touches the
    annotation.  Some embedded fonts expose boxes several points beyond the
    visible baseline, so a region beginning on the next physical line can erase
    the preceding list number.  The replacement image still uses the original
    rectangle; only the destructive redaction edge is contracted inward, and
    only for the known overlapping-glyph fallback.
    """

    if _OVERLAPPING_PDF_GLYPH_REASON not in region.reasons:
        return region.rect
    try:
        page_dict = page.get_text("dict") or {}
        page_rect = page.rect
    except Exception:
        return region.rect

    intended: list[tuple[float, float, float, float]] = []
    external_above: list[float] = []
    external_below: list[float] = []
    region_rect = region.rect
    for block in page_dict.get("blocks", ()):
        if int(block.get("type", 0) or 0) != 0:
            continue
        for line in block.get("lines", ()):
            for span in line.get("spans", ()):
                if not str(span.get("text", "") or "").strip():
                    continue
                span_rect = _coerce_pdf_rect(span.get("bbox", ()), page_rect)
                if span_rect is None:
                    continue
                intersection = _rect_intersection_ratio(span_rect, region_rect)
                if intersection <= 0.0:
                    continue
                if _rect_center_inside(span_rect, region_rect) or intersection >= 0.35:
                    intended.append(span_rect)
                    continue
                center_y = (span_rect[1] + span_rect[3]) / 2.0
                if center_y < region_rect[1]:
                    external_above.append(span_rect[3])
                elif center_y > region_rect[3]:
                    external_below.append(span_rect[1])

    if not intended or (not external_above and not external_below):
        return region_rect
    x0, y0, x1, y1 = region_rect
    if external_above:
        y0 = max(y0, max(external_above) + 0.02)
    if external_below:
        y1 = min(y1, min(external_below) - 0.02)
    contracted = (x0, y0, x1, y1)
    if y1 - y0 < 0.5:
        return region_rect
    # The contracted mask must still touch every selected text island.  If an
    # unusual page geometry cannot satisfy that invariant, retain the original
    # behavior rather than silently leaving duplicate editable text behind.
    if any(
        _rect_intersection_ratio(span_rect, contracted) <= 0.0 for span_rect in intended
    ):
        return region_rect
    return contracted


def _flatten_pdf_regions(
    source: Path | str,
    target: Path | str,
    regions: Iterable[_HybridRegion],
    *,
    password: str | None = None,
    dpi: int = 300,
) -> None:
    """Replace selected PDF rectangles with pixel-faithful images in place."""

    if int(dpi) < 72 or int(dpi) > 600:
        raise ValidationError("局部高清区域的 DPI 必须在 72–600 之间")
    source_path = Path(source)
    target_path = Path(target)
    pymupdf, document = _open_pymupdf_document(source_path, password)
    try:
        grouped: dict[int, list[_HybridRegion]] = {}
        for region in regions:
            if region.page_index < 0 or region.page_index >= document.page_count:
                raise ValidationError(
                    f"局部高清区域页码超出范围：{region.page_index + 1}"
                )
            grouped.setdefault(region.page_index, []).append(region)

        for page_index, page_regions in grouped.items():
            page = document[page_index]
            normalized_regions = _merge_hybrid_regions(
                (
                    _HybridRegion(
                        page_index=page_index,
                        rect=normalized,
                        kind=region.kind,
                        reasons=region.reasons,
                        dpi=region.dpi,
                    )
                    for region in page_regions
                    if (
                        normalized := _coerce_pdf_rect(
                            region.rect,
                            page.rect,
                        )
                    )
                    is not None
                ),
                page.rect,
                gap=0.5,
            )
            remove_regions, pixel_regions = _partition_pdf_redaction_regions(
                page,
                normalized_regions,
                page.rect,
            )
            rendered: list[tuple[tuple[float, float, float, float], bytes]] = []
            for region in normalized_regions:
                effective_dpi = min(
                    600,
                    max(72, int(region.dpi or dpi)),
                )
                rect = pymupdf.Rect(region.rect)
                pixmap = page.get_pixmap(
                    dpi=effective_dpi,
                    clip=rect,
                    alpha=False,
                    annots=True,
                )
                rendered.append((region.rect, pixmap.tobytes("png")))
            if rendered:
                for phase_regions, image_mode in (
                    (remove_regions, pymupdf.PDF_REDACT_IMAGE_REMOVE),
                    (pixel_regions, pymupdf.PDF_REDACT_IMAGE_PIXELS),
                ):
                    if not phase_regions:
                        continue
                    for region in phase_regions:
                        page.add_redact_annot(
                            pymupdf.Rect(
                                _contract_pdf_redaction_rect_away_from_external_text(
                                    page,
                                    region,
                                )
                            ),
                            fill=(1.0, 1.0, 1.0),
                            cross_out=False,
                        )
                    page.apply_redactions(
                        images=image_mode,
                        graphics=pymupdf.PDF_REDACT_LINE_ART_REMOVE_IF_COVERED,
                        text=pymupdf.PDF_REDACT_TEXT_REMOVE,
                    )
                for rect, image_bytes in rendered:
                    page.insert_image(
                        pymupdf.Rect(rect),
                        stream=image_bytes,
                        keep_proportion=False,
                        overlay=True,
                    )

        target_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(
            str(target_path),
            garbage=4,
            deflate=True,
            deflate_images=True,
            use_objstms=True,
        )
    except ValidationError:
        raise
    except Exception as exc:
        raise ValidationError(f"PDF 局部高清保真处理失败：{exc}") from exc
    finally:
        document.close()
    if not target_path.is_file() or target_path.stat().st_size == 0:
        raise ValidationError("PDF 局部高清保真处理没有生成有效文件")


def _mask_pdf_regions_for_editable_conversion(
    source: Path,
    target: Path,
    regions: Iterable[_HybridRegion],
    *,
    password: str | None,
) -> None:
    """Blank page-local overlays before pdf2docx sees the editable body.

    Unlike ``_flatten_pdf_regions()``, this helper never inserts the rendered
    pixels back into the intermediate PDF.  The original crop is restored only
    after Word construction as a page-relative floating drawing, so it cannot
    consume line height or force an extra page.
    """

    source_path = Path(source)
    target_path = Path(target)
    pymupdf, document = _open_pymupdf_document(source_path, password)
    try:
        grouped: dict[int, list[_HybridRegion]] = {}
        for region in regions:
            if region.page_index < 0 or region.page_index >= document.page_count:
                raise ValidationError(
                    f"PDF 浮动保真区域页码超出范围：{region.page_index + 1}"
                )
            grouped.setdefault(region.page_index, []).append(region)

        for page_index, page_regions in grouped.items():
            page = document[page_index]
            normalized_regions = _merge_hybrid_regions(
                (
                    _HybridRegion(
                        page_index=page_index,
                        rect=normalized,
                        kind=region.kind,
                        reasons=region.reasons,
                        dpi=region.dpi,
                    )
                    for region in page_regions
                    if (
                        normalized := _coerce_pdf_rect(
                            region.rect,
                            page.rect,
                            padding=0.0,
                        )
                    )
                    is not None
                ),
                page.rect,
                gap=0.5,
            )
            if not normalized_regions:
                continue
            for region in normalized_regions:
                page.add_redact_annot(
                    pymupdf.Rect(region.rect),
                    fill=(1.0, 1.0, 1.0),
                    cross_out=False,
                )
            page.apply_redactions(
                images=pymupdf.PDF_REDACT_IMAGE_PIXELS,
                graphics=pymupdf.PDF_REDACT_LINE_ART_REMOVE_IF_COVERED,
                text=pymupdf.PDF_REDACT_TEXT_REMOVE,
            )

        target_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(
            str(target_path),
            garbage=4,
            deflate=True,
            deflate_images=True,
            use_objstms=True,
        )
    except ValidationError:
        raise
    except Exception as exc:
        raise ValidationError(f"PDF 侧栏遮罩处理失败：{exc}") from exc
    finally:
        document.close()
    if not target_path.is_file() or target_path.stat().st_size == 0:
        raise ValidationError("PDF 侧栏遮罩处理没有生成有效文件")


def _body_content_children(document: Any) -> list[Any]:
    from docx.oxml.ns import qn

    return [child for child in document.element.body if child.tag != qn("w:sectPr")]


def _section_break_is_continuous(section_properties: Any) -> bool:
    from docx.oxml.ns import qn

    section_type = section_properties.find(qn("w:type"))
    return (
        section_type is not None
        and str(section_type.get(qn("w:val")) or "").casefold() == "continuous"
    )


def _section_break_is_next_page(section_properties: Any) -> bool:
    from docx.oxml.ns import qn

    section_type = section_properties.find(qn("w:type"))
    if section_type is None:
        return True
    return str(section_type.get(qn("w:val")) or "").casefold() == "nextpage"


def _compact_section_break_paragraph(paragraph: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(1)
    paragraph.paragraph_format.keep_together = False
    paragraph.paragraph_format.keep_with_next = False
    paragraph.paragraph_format.page_break_before = False
    paragraph.paragraph_format.widow_control = False
    properties = paragraph._p.get_or_add_pPr()
    snap_to_grid = properties.find(qn("w:snapToGrid"))
    if snap_to_grid is None:
        snap_to_grid = OxmlElement("w:snapToGrid")
        section_properties = properties.find(qn("w:sectPr"))
        if section_properties is None:
            properties.append(snap_to_grid)
        else:
            properties.insert(properties.index(section_properties), snap_to_grid)
    snap_to_grid.set(qn("w:val"), "0")
    if not paragraph._p.findall(qn("w:r")):
        run = paragraph.add_run()
        run.font.size = Pt(1)
        run.font.hidden = True


def _relocate_or_compact_page_section_break(
    document: Any,
    previous_body_child: Any,
    new_body_children: Iterable[Any],
) -> None:
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph

    break_element = None
    section_properties = None
    for child in new_body_children:
        if child.tag != qn("w:p"):
            continue
        properties = child.pPr
        candidate = properties.sectPr if properties is not None else None
        if candidate is None:
            continue
        # Page.make_docx() first inserts a section carrier for the section that
        # ended on the previous source page.  On a mixed-layout next page that
        # carrier can be ``continuous``; the later, internal single-to-double
        # transition then carries the real ``nextPage`` start type.  Selecting
        # by type therefore moves the wrong section onto the previous page and
        # can make its footer start on a new page.  The first newly appended
        # paragraph-level sectPr is the unambiguous cross-page carrier.
        break_element = child
        section_properties = candidate
        break
    if break_element is None or section_properties is None:
        return

    previous_is_paragraph = (
        previous_body_child is not None and previous_body_child.tag == qn("w:p")
    )
    previous_has_embedded_object = bool(
        previous_is_paragraph
        and previous_body_child.xpath(".//w:drawing|.//w:pict|.//w:object")
    )
    break_has_structural_content = any(
        child.tag != qn("w:pPr") for child in break_element
    )
    break_is_empty = (
        not break_has_structural_content
        and not any(
            str(value or "").strip() for value in break_element.xpath(".//w:t/text()")
        )
        and not break_element.xpath(".//w:drawing|.//w:pict|.//w:object|.//w:br")
    )
    if previous_is_paragraph and not previous_has_embedded_object and break_is_empty:
        previous_properties = previous_body_child.get_or_add_pPr()
        if previous_properties.sectPr is None:
            previous_properties.append(section_properties)
            parent = break_element.getparent()
            if parent is not None:
                parent.remove(break_element)
            return

    _compact_section_break_paragraph(Paragraph(break_element, document._body))


def _section_carrier_has_substantive_content(paragraph: Any) -> bool:
    if any(
        str(value or "").strip()
        for value in paragraph.xpath(
            ".//w:t/text()|.//w:delText/text()|.//w:instrText/text()|"
            ".//w:delInstrText/text()"
        )
    ):
        return True
    if paragraph.xpath(
        ".//w:drawing|.//w:pict|.//w:object|.//w:br|.//w:cr|"
        ".//w:lastRenderedPageBreak|.//w:tab|.//w:ptab|.//w:sym|"
        ".//w:noBreakHyphen|.//w:softHyphen|.//w:fldChar|.//w:fldSimple|"
        ".//w:hyperlink|.//w:sdt|.//w:footnoteReference|"
        ".//w:endnoteReference|.//w:commentReference|.//w:annotationRef|"
        ".//w:altChunk|.//w:subDoc|.//w:contentPart"
    ):
        return True
    return False


def _normalize_empty_next_page_section_paragraphs(document: Any) -> None:
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn

    for child in list(_body_content_children(document)):
        if child.getparent() is None or child.tag != qn("w:p"):
            continue
        properties = child.pPr
        section_properties = properties.sectPr if properties is not None else None
        if section_properties is None or not _section_break_is_next_page(
            section_properties
        ):
            continue
        if _section_carrier_has_substantive_content(child):
            continue
        # Cross-source-page NEW_PAGE boundaries are handled while appending each
        # parsed page, where their identity is unambiguous.  A global relocation
        # would also move pdf2docx's legitimate in-page bootstrap sectPr (for
        # example, a full-width paper title followed by two-column content) onto
        # the title paragraph and change pagination.  Keep such internal section
        # markers independent, but make their paragraph footprint negligible.
        _compact_section_break_paragraph(Paragraph(child, document._body))


_PDF2DOCX_SHORT_ENGLISH_FRAGMENT_PATTERN = re.compile(
    r"^[a-z][A-Za-z]*(?:['’\-][A-Za-z]+)*"
    r"(?:\s+[A-Za-z]+(?:['’\-][A-Za-z]+)*){0,3}"
    r"[.!?](?:[\"'’”\)\]])?$"
)
_PDF2DOCX_ENGLISH_WORD_WITH_APOSTROPHE_PATTERN = re.compile(
    r"[A-Za-z]+(?:['’\-][A-Za-z]+)*"
)


def _merge_cross_page_english_orphan_fragments(document: Any) -> int:
    """Join a very short lowercase sentence tail stranded after a page boundary.

    PDF page-by-page reconstruction can force the final one-to-four words of an
    English sentence onto the next source-page section even when Word has room
    for them on the preceding physical page.  Only a terminal, lowercase,
    style-compatible fragment is moved.  Section carriers, tables, figures and
    ordinary next-page paragraphs remain untouched.
    """

    from docx.oxml.ns import qn

    body_children = list(_body_content_children(document))
    sections: list[list[Any]] = []
    section_start = 0
    for index, child in enumerate(body_children):
        if child.tag != qn("w:p"):
            continue
        properties = child.pPr
        section_properties = properties.sectPr if properties is not None else None
        if section_properties is None or not _section_break_is_next_page(
            section_properties
        ):
            continue
        sections.append(body_children[section_start : index + 1])
        section_start = index + 1
    sections.append(body_children[section_start:])
    if len(sections) < 2:
        return 0

    def paragraph_text(paragraph: Any) -> str:
        return " ".join(
            "".join(paragraph.xpath(".//w:t/text()")).replace("\u200b", "").split()
        )

    def is_floating_furniture(paragraph: Any) -> bool:
        return bool(
            paragraph.xpath("./w:pPr/w:framePr")
            or paragraph.xpath(
                ".//wp:anchor/wp:docPr[@descr='LayoutLoom running footer' or "
                "@descr='DocuForge running footer']"
            )
        )

    def has_embedded_visual(paragraph: Any) -> bool:
        return bool(paragraph.xpath(".//w:drawing|.//w:pict|.//w:object"))

    def is_page_furniture(text: str) -> bool:
        return bool(
            text and len(text) <= 180 and _PDF2DOCX_PAGE_FURNITURE_PATTERN.search(text)
        )

    def boundary_paragraph(
        children: Iterable[Any],
        *,
        reverse: bool,
    ) -> Any | None:
        decorative_paragraphs = 0
        ordered = reversed(list(children)) if reverse else iter(children)
        for child in ordered:
            if child.getparent() is None:
                continue
            if child.tag == qn("w:tbl"):
                return None
            if child.tag != qn("w:p"):
                continue
            properties = child.pPr
            if (
                properties is not None
                and properties.sectPr is not None
                and not _section_carrier_has_substantive_content(child)
            ):
                continue
            if is_floating_furniture(child):
                continue
            text = paragraph_text(child)
            visual = has_embedded_visual(child)
            if not text:
                if visual:
                    decorative_paragraphs += 1
                    if decorative_paragraphs > 2:
                        return None
                continue
            if is_page_furniture(text):
                continue
            if visual:
                return None
            return child
        return None

    def paragraph_style_id(paragraph: Any) -> str:
        properties = paragraph.pPr
        style = properties.find(qn("w:pStyle")) if properties is not None else None
        return str(style.get(qn("w:val")) or "") if style is not None else ""

    def terminal_run_signature(
        paragraph: Any,
        *,
        reverse: bool,
    ) -> tuple[frozenset[str], int | None]:
        runs = list(paragraph.xpath(".//w:r"))
        if reverse:
            runs.reverse()
        for run in runs:
            if not "".join(run.xpath(".//w:t/text()")).strip():
                continue
            properties = run.find(qn("w:rPr"))
            if properties is None:
                return frozenset(), None
            fonts = properties.find(qn("w:rFonts"))
            font_names = frozenset(
                str(fonts.get(qn(attribute)) or "").strip().casefold()
                for attribute in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs")
                if fonts is not None and str(fonts.get(qn(attribute)) or "").strip()
            )
            size = properties.find(qn("w:sz"))
            try:
                half_points = int(size.get(qn("w:val"))) if size is not None else None
            except (TypeError, ValueError):
                half_points = None
            return font_names, half_points
        return frozenset(), None

    def formatting_is_compatible(previous: Any, following: Any) -> bool:
        previous_style = paragraph_style_id(previous)
        following_style = paragraph_style_id(following)
        if previous_style and following_style and previous_style != following_style:
            return False
        previous_fonts, previous_size = terminal_run_signature(previous, reverse=True)
        following_fonts, following_size = terminal_run_signature(
            following, reverse=False
        )
        if (
            previous_fonts
            and following_fonts
            and previous_fonts.isdisjoint(following_fonts)
        ):
            return False
        if (
            previous_size is not None
            and following_size is not None
            and abs(previous_size - following_size) > 2
        ):
            return False
        return True

    changed = 0
    for previous_section, following_section in zip(sections, sections[1:]):
        previous = boundary_paragraph(previous_section, reverse=True)
        following = boundary_paragraph(following_section, reverse=False)
        if previous is None or following is None:
            continue
        previous_text = paragraph_text(previous)
        following_text = paragraph_text(following)
        previous_words = _PDF2DOCX_ENGLISH_WORD_WITH_APOSTROPHE_PATTERN.findall(
            previous_text
        )
        following_words = _PDF2DOCX_ENGLISH_WORD_WITH_APOSTROPHE_PATTERN.findall(
            following_text
        )
        if (
            len(previous_text) < 24
            or len(previous_words) < 4
            or not re.search(r"[A-Za-z]$", previous_text)
            or len(following_text) > 48
            or not 1 <= len(following_words) <= 4
            or not _PDF2DOCX_SHORT_ENGLISH_FRAGMENT_PATTERN.fullmatch(following_text)
            or following.xpath(".//w:br|.//w:tab")
            or not formatting_is_compatible(previous, following)
        ):
            continue

        text_nodes = following.xpath(".//w:t")
        if not text_nodes:
            continue
        first_text = text_nodes[0]
        first_text.text = " " + str(first_text.text or "").lstrip()
        first_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        following_properties = following.pPr
        for payload in list(following):
            if payload is following_properties:
                continue
            previous.append(payload)
        parent = following.getparent()
        if parent is not None:
            parent.remove(following)
            changed += 1
    return changed


def _make_pdf2docx_page_with_hard_column_breaks(
    document: Any,
    parsed_page: Any,
) -> None:
    """Render pdf2docx sections using real column breaks instead of nextColumn sections.

    WPS can rebalance ``nextColumn`` section transitions and move the tail of the
    left source column above the start of the right one. A plain column break keeps
    both columns in one section and preserves their explicit reading order.
    """

    try:
        from docx.enum.section import WD_SECTION
        from docx.enum.text import WD_BREAK
        from docx.shared import Pt
        from pdf2docx.common import constants
        from pdf2docx.common.docx import reset_paragraph_format, set_columns
    except ImportError as exc:
        raise MissingEngineError(
            "PDF 转 Word 的稳定分栏输出需要 pdf2docx 与 python-docx"
        ) from exc

    if document.paragraphs:
        page_section = document.add_section(WD_SECTION.NEW_PAGE)
    else:
        page_section = document.sections[0]
    page_section.page_width = Pt(float(parsed_page.width))
    page_section.page_height = Pt(float(parsed_page.height))
    left, right, top, bottom = (float(value) for value in parsed_page.margin)
    page_section.left_margin = Pt(left)
    page_section.right_margin = Pt(right)
    page_section.top_margin = Pt(top)
    page_section.bottom_margin = Pt(bottom)

    page_paragraph_start = len(document.paragraphs)
    sections = list(parsed_page.sections)
    if not sections:
        return

    def add_spacing_paragraph(before_space: float) -> None:
        paragraph = document.add_paragraph()
        line_height = min(max(0.0, float(before_space)), 11.0)
        paragraph_format = reset_paragraph_format(
            paragraph,
            line_spacing=Pt(max(1.0, line_height)),
        )
        paragraph_format.space_after = Pt(max(0.0, float(before_space) - line_height))

    def add_hard_column_break() -> None:
        paragraph = document.add_paragraph()
        paragraph_format = reset_paragraph_format(paragraph, line_spacing=Pt(1))
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph.add_run().add_break(WD_BREAK.COLUMN)

    def render_layout_section(layout_section: Any) -> None:
        word_section = document.sections[-1]
        width_list = [
            float(column.bbox[2] - column.bbox[0]) for column in layout_section
        ]
        set_columns(word_section, width_list, float(layout_section.space))
        for column_index, column in enumerate(layout_section):
            if column_index:
                add_hard_column_break()
            column.make_docx(document)

    first = sections[0]
    if float(first.before_space) > float(constants.MINOR_DIST):
        add_spacing_paragraph(float(first.before_space))
    if len(first) == 2:
        document.add_section(WD_SECTION.CONTINUOUS)
    render_layout_section(first)

    for layout_section in sections[1:]:
        document.add_section(WD_SECTION.CONTINUOUS)
        if len(document.paragraphs) >= 2:
            spacing_paragraph = document.paragraphs[-2]
        else:
            spacing_paragraph = document.paragraphs[-1]
        if (
            not spacing_paragraph.text.strip()
            and "graphicData" in spacing_paragraph._p.xml
        ):
            spacing_paragraph = document.paragraphs[-1]
        spacing_paragraph.paragraph_format.space_after = Pt(
            max(0.0, float(layout_section.before_space))
        )
        render_layout_section(layout_section)

    paragraphs = document.paragraphs
    if page_paragraph_start < len(paragraphs):
        for float_image in parsed_page.float_images:
            float_image.make_docx(paragraphs[page_paragraph_start])


def _inline_picture_to_page_anchor(
    inline: Any,
    *,
    x_offset: int,
    y_offset: int,
    relative_height: int,
) -> None:
    """Turn a python-docx inline picture into a physical-page anchored drawing."""

    try:
        from docx.oxml import OxmlElement
    except ImportError as exc:
        raise MissingEngineError("PDF 转 Word 需要 python-docx") from exc

    children = list(inline)
    by_local_name = {str(child.tag).rsplit("}", 1)[-1]: child for child in children}
    required = ("extent", "docPr", "cNvGraphicFramePr", "graphic")
    if any(name not in by_local_name for name in required):
        raise ValidationError("无法建立 Word 页面对齐的浮动图片")

    anchor = OxmlElement("wp:anchor")
    for name, value in (
        ("distT", "0"),
        ("distB", "0"),
        ("distL", "0"),
        ("distR", "0"),
        ("simplePos", "0"),
        ("relativeHeight", str(max(0, int(relative_height)))),
        ("behindDoc", "0"),
        ("locked", "1"),
        ("layoutInCell", "1"),
        ("allowOverlap", "1"),
    ):
        anchor.set(name, value)

    simple_position = OxmlElement("wp:simplePos")
    simple_position.set("x", "0")
    simple_position.set("y", "0")
    anchor.append(simple_position)

    for axis, offset in (("H", x_offset), ("V", y_offset)):
        position = OxmlElement(f"wp:position{axis}")
        position.set("relativeFrom", "page")
        position_offset = OxmlElement("wp:posOffset")
        position_offset.text = str(max(0, int(offset)))
        position.append(position_offset)
        anchor.append(position)

    anchor.append(by_local_name["extent"])
    effect_extent = by_local_name.get("effectExtent")
    if effect_extent is None:
        effect_extent = OxmlElement("wp:effectExtent")
        for side in ("l", "t", "r", "b"):
            effect_extent.set(side, "0")
    anchor.append(effect_extent)
    anchor.append(OxmlElement("wp:wrapNone"))
    anchor.append(by_local_name["docPr"])
    anchor.append(by_local_name["cNvGraphicFramePr"])
    anchor.append(by_local_name["graphic"])

    drawing = inline.getparent()
    if drawing is None:
        raise ValidationError("无法定位 Word 浮动图片容器")
    drawing.replace(inline, anchor)


def _append_page_relative_pdf_regions(
    document: Any,
    page: Any,
    regions: Iterable[_HybridRegion],
    folder: Path,
    *,
    existing_body_ids: set[int],
) -> int:
    """Restore masked PDF crops without adding height to the editable text flow."""

    try:
        from docx.oxml.ns import qn
        from docx.shared import Pt
        from docx.text.paragraph import Paragraph
    except ImportError as exc:
        raise MissingEngineError("PDF 转 Word 需要 python-docx") from exc

    normalized_regions = [
        (region, normalized)
        for region in regions
        if (
            normalized := _coerce_pdf_rect(
                region.rect,
                page.rect,
                padding=0.0,
            )
        )
        is not None
    ]
    if not normalized_regions:
        return 0

    new_children = [
        child
        for child in _body_content_children(document)
        if id(child) not in existing_body_ids
    ]
    paragraph_elements = [child for child in new_children if child.tag == qn("w:p")]
    substantive_paragraphs = [
        child
        for child in paragraph_elements
        if child.xpath(".//w:t[string-length(normalize-space(.)) > 0]|.//w:drawing")
    ]
    anchor_element = next(
        (
            child
            for child in substantive_paragraphs
            if child.pPr is None or child.pPr.sectPr is None
        ),
        None,
    )
    if anchor_element is None:
        anchor_element = next(iter(substantive_paragraphs), None)
    if anchor_element is None:
        anchor_element = next(iter(paragraph_elements), None)
    if anchor_element is None:
        raise ValidationError("可编辑页面没有可承载侧栏保真图的段落")
    anchor_paragraph = Paragraph(anchor_element, document._body)

    page_x0 = float(page.rect.x0)
    page_y0 = float(page.rect.y0)
    restored = 0
    for region_index, (region, rect) in enumerate(normalized_regions):
        image_path = folder / (
            f"page-{region.page_index + 1:06d}-{region.kind}-{region_index + 1:02d}.png"
        )
        if not image_path.is_file():
            pixmap = page.get_pixmap(
                dpi=min(600, max(180, int(region.dpi))),
                clip=rect,
                alpha=False,
                annots=True,
            )
            pixmap.save(str(image_path))
        width = max(0.1, float(rect[2] - rect[0]))
        height = max(0.1, float(rect[3] - rect[1]))
        inline_shape = anchor_paragraph.add_run().add_picture(
            str(image_path),
            width=Pt(width),
            height=Pt(height),
        )
        _inline_picture_to_page_anchor(
            inline_shape._inline,
            x_offset=int(Pt(float(rect[0]) - page_x0)),
            y_offset=int(Pt(float(rect[1]) - page_y0)),
            relative_height=251659264 + region_index,
        )
        restored += 1
    return restored


def _append_pdf2docx_page(
    document: Any,
    parsed_page: Any,
    *,
    is_first_page: bool,
) -> None:
    previous_children = _body_content_children(document)
    previous_body_child = previous_children[-1] if previous_children else None
    parsed_sections = getattr(parsed_page, "sections", ()) or ()
    if any(len(section) == 2 for section in parsed_sections):
        _make_pdf2docx_page_with_hard_column_breaks(document, parsed_page)
    else:
        parsed_page.make_docx(document)
    if is_first_page or previous_body_child is None:
        return
    current_children = _body_content_children(document)
    try:
        previous_index = current_children.index(previous_body_child)
    except ValueError:
        return
    _relocate_or_compact_page_section_break(
        document,
        previous_body_child,
        current_children[previous_index + 1 :],
    )


def _append_visual_pdf_page(
    document: Any,
    page: Any,
    image_path: Path,
    *,
    dpi: int,
    is_first_page: bool,
) -> None:
    try:
        from docx.enum.section import WD_SECTION
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise MissingEngineError("PDF 转 Word 需要 python-docx") from exc

    if not is_first_page:
        previous_children = _body_content_children(document)
        previous_body_child = previous_children[-1] if previous_children else None
        document.add_section(WD_SECTION.NEW_PAGE)
        current_children = _body_content_children(document)
        if previous_body_child is not None:
            previous_index = current_children.index(previous_body_child)
            _relocate_or_compact_page_section_break(
                document,
                previous_body_child,
                current_children[previous_index + 1 :],
            )

    section = document.sections[-1]
    section.page_width = Pt(float(page.rect.width))
    section.page_height = Pt(float(page.rect.height))
    section.top_margin = Inches(0.05)
    section.bottom_margin = Inches(0.05)
    section.left_margin = Inches(0.05)
    section.right_margin = Inches(0.05)
    section.header_distance = Pt(0)
    section.footer_distance = Pt(0)

    if image_path.is_file():
        with Image.open(image_path) as rendered_image:
            image_width, image_height = rendered_image.size
    else:
        pixmap = page.get_pixmap(dpi=int(dpi), alpha=False)
        pixmap.save(str(image_path))
        image_width, image_height = pixmap.width, pixmap.height

    usable_width = int(section.page_width - section.left_margin - section.right_margin)
    usable_height = int(
        section.page_height - section.top_margin - section.bottom_margin
    )
    scale = min(usable_width / image_width, usable_height / image_height)
    draw_width = max(1, round(image_width * scale))
    draw_height = max(1, round(image_height * scale))

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    # An exact 1 pt line box clips an inline full-page image completely in WPS
    # (and in some Word compatibility modes).  Auto/single spacing lets the
    # line box expand to the drawing height while the zero paragraph spacing
    # still prevents an extra blank line around the page image.
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(
        str(image_path), width=draw_width, height=draw_height
    )


def _build_visual_docx_document(
    pdf_document: Any,
    *,
    dpi: int,
    folder: Path,
    progress_start: float = 0.10,
    progress_span: float = 0.80,
) -> Any:
    from docuforge.runner import check_cancelled, report_progress

    try:
        from docx import Document
    except ImportError as exc:
        raise MissingEngineError("PDF 转 Word 需要 python-docx") from exc
    document = Document()
    page_count = max(1, int(pdf_document.page_count))
    for page_index in range(pdf_document.page_count):
        check_cancelled("任务已取消；已完成的文件会保留")
        report_progress(
            progress_start + progress_span * ((page_index + 1) / page_count),
            f"渲染高清页面 {page_index + 1}/{page_count}",
        )
        _append_visual_pdf_page(
            document,
            pdf_document[page_index],
            folder / f"page-{page_index + 1:06d}.png",
            dpi=dpi,
            is_first_page=page_index == 0,
        )
    return document


def _hybrid_visual_page_summary(reasons: Mapping[int, list[str]]) -> str:
    entries: list[str] = []
    ordered = sorted(reasons)
    for page_index in ordered[:12]:
        unique_reasons = list(dict.fromkeys(reasons[page_index]))
        detail = "、".join(unique_reasons[:2])
        entries.append(f"{page_index + 1}（{detail}）")
    if len(ordered) > 12:
        entries.append(f"另 {len(ordered) - 12} 页")
    return "；".join(entries)


def _hybrid_region_summary(
    assessments: Iterable[_HybridPageAssessment],
    full_page_reasons: Mapping[int, list[str]],
) -> str:
    kind_names = {
        "figure": "图像/图表",
        "vector": "矢量图",
        "table": "表格",
        "formula": "公式",
        "complex": "复合区域",
    }
    entries: list[str] = []
    for assessment in assessments:
        if assessment.page_index in full_page_reasons or not assessment.visual_regions:
            continue
        counts = Counter(
            kind_names.get(region.kind, "复杂区域")
            for region in assessment.visual_regions
        )
        detail = "、".join(f"{name}{count}处" for name, count in counts.items())
        entries.append(f"{assessment.page_index + 1}（{detail}）")
        if len(entries) >= 12:
            break
    remaining = sum(
        1
        for assessment in assessments
        if assessment.page_index not in full_page_reasons and assessment.visual_regions
    ) - len(entries)
    if remaining > 0:
        entries.append(f"另 {remaining} 页")
    return "；".join(entries)


def _convert_pdf_to_hybrid_docx(
    source: Path,
    target: Path,
    *,
    password: str | None,
    dpi: int,
    low_quality_policy: str,
    force_visual_pages: str,
    column_layout: str,
) -> None:
    from docuforge.runner import check_cancelled, report_progress

    try:
        from docx import Document
        from pdf2docx import Converter
    except ImportError as exc:
        raise MissingEngineError(
            "版式优先混合模式需要 pdf2docx、PyMuPDF 与 python-docx；"
            "请重新运行安装程序补齐依赖"
        ) from exc
    if int(dpi) < 72 or int(dpi) > 600:
        raise ValidationError("混合模式图像清晰度 DPI 必须在 72–600 之间")

    report_progress(0.01, "读取 PDF 并准备混合保真分析")
    assessments = _assess_pdf_pages_for_hybrid(
        source,
        password,
        dpi=int(dpi),
        column_layout=column_layout,
    )
    effective_column_layout = _resolve_hybrid_auto_column_layout(
        column_layout,
        assessments,
    )
    clip_image_res_ratio = _hybrid_clip_image_res_ratio(assessments, int(dpi))
    page_count = len(assessments)
    report_progress(0.27, "规划可编辑正文与高清保真区域")
    forced_indexes = (
        set(parse_page_spec(force_visual_pages, page_count))
        if str(force_visual_pages).strip()
        else set()
    )
    profile = _pdf_word_layout_profile(
        source,
        password,
        column_layout=column_layout,
    )
    if profile.fixed_layout_recommended:
        report_progress(0.29, "检测到复杂版面，启用固定坐标混合重建")
        _convert_pdf_to_fixed_hybrid_docx(
            source,
            target,
            password=password,
            dpi=int(dpi),
            low_quality_policy=low_quality_policy,
            assessments=assessments,
            forced_indexes=forced_indexes,
            profile=profile,
        )
        return

    risk_reasons: dict[int, list[str]] = {
        assessment.page_index: list(assessment.reasons)
        for assessment in assessments
        if assessment.reasons
    }
    for page_index in forced_indexes:
        risk_reasons.setdefault(page_index, []).append("用户指定原样保留")

    pending_warnings: list[str] = []
    converter = None
    pdf_document = None

    def close_converter() -> None:
        nonlocal converter
        active_converter = converter
        converter = None
        if active_converter is not None:
            with contextlib.suppress(Exception):
                active_converter.close()

    try:
        _pymupdf, pdf_document = _open_pymupdf_document(source, password)
        with atomic_output(target) as temporary, _temporary_working_directory(
            prefix="docuforge-pdf-docx-",
            before_cleanup=close_converter,
        ) as folder:
            region_source = folder / "region-hybrid-source.pdf"
            flatten_regions = [
                region
                for assessment in assessments
                if assessment.page_index not in risk_reasons
                for region in assessment.visual_regions
            ]
            anchored_regions = [
                region
                for assessment in assessments
                if assessment.page_index not in risk_reasons
                for region in assessment.anchored_visual_regions
            ]
            converter_source = source
            converter_password = password
            if flatten_regions:
                _flatten_pdf_regions(
                    source,
                    region_source,
                    flatten_regions,
                    password=password,
                    dpi=int(dpi),
                )
                converter_source = region_source
                converter_password = None
            if anchored_regions:
                masked_source = folder / "region-hybrid-masked-source.pdf"
                _mask_pdf_regions_for_editable_conversion(
                    converter_source,
                    masked_source,
                    anchored_regions,
                    password=converter_password,
                )
                converter_source = masked_source
                converter_password = None
            report_progress(0.33, "准备可编辑转换源文件")

            candidate_indexes = [
                assessment.page_index
                for assessment in assessments
                if assessment.page_index not in risk_reasons
            ]
            settings: dict[str, Any] = {}
            if candidate_indexes:
                try:

                    def parse_candidate_source(
                        candidate_source: Path,
                        candidate_password: str | None,
                    ) -> None:
                        nonlocal converter, settings
                        close_converter()
                        converter = Converter(
                            str(candidate_source),
                            password=candidate_password or None,
                        )
                        settings = _pdf2docx_settings(
                            converter,
                            resilient=True,
                            parse_lattice_table=False,
                            clip_image_res_ratio=clip_image_res_ratio,
                            column_layout=effective_column_layout,
                        )
                        report_progress(
                            0.38,
                            f"重建 {len(candidate_indexes)} 页可编辑版面",
                        )
                        _parse_pdf2docx_pages(
                            converter,
                            settings=settings,
                            column_layout=effective_column_layout,
                            page_assessments={
                                assessment.page_index: assessment
                                for assessment in assessments
                            },
                            pages=candidate_indexes,
                        )
                        source_words = set(
                            _english_words(
                                "\n".join(
                                    assessment.editable_source_text
                                    for assessment in assessments
                                    if assessment.page_index in candidate_indexes
                                )
                            )
                        )
                        _restore_pdf2docx_spaces(
                            converter.pages,
                            source_words=source_words,
                            source_document=converter.fitz_doc,
                        )
                        _repair_pdf2docx_known_encoding(converter.pages)
                        _normalize_pdf2docx_english_widths(converter.pages)
                        _normalize_pdf2docx_list_prefixes(converter.pages)
                        _normalize_pdf2docx_prose_alignment(converter.pages)
                        report_progress(0.54, "完成英文、列表与分栏重建")

                    parse_candidate_source(converter_source, converter_password)

                    # A short paragraph can occasionally be merged from the
                    # left and right source columns before pdf2docx creates its
                    # sections.  Preserve just that narrow row as pixels and
                    # parse once more, instead of sacrificing the complete page.
                    localized_merge_pages: set[int] = set()
                    for page_index in candidate_indexes:
                        parsed_page = converter.pages[page_index]
                        if not getattr(parsed_page, "finalized", False):
                            continue
                        assessment = assessments[page_index]
                        merge_rects = _pdf2docx_intercolumn_text_merge_rects(
                            parsed_page,
                            split_x=assessment.column_split_x,
                            paired_bands=_hybrid_editable_column_bands(assessment),
                        )
                        if not merge_rects:
                            continue
                        source_page = pdf_document[page_index]
                        page_area = max(
                            1.0,
                            float(source_page.rect.width)
                            * float(source_page.rect.height),
                        )
                        local_height_limit = max(
                            36.0,
                            float(source_page.rect.height) * 0.08,
                        )
                        if (
                            len(merge_rects) > 4
                            or any(
                                rect[3] - rect[1] > local_height_limit
                                for rect in merge_rects
                            )
                            or sum(_rect_area(rect) for rect in merge_rects)
                            > page_area * 0.12
                        ):
                            continue
                        local_regions = [
                            _HybridRegion(
                                page_index=page_index,
                                rect=normalized,
                                kind="complex",
                                reasons=("局部跨栏文字合并保护",),
                                dpi=int(dpi),
                            )
                            for rect in merge_rects
                            if (
                                normalized := _coerce_pdf_rect(
                                    rect,
                                    source_page.rect,
                                    padding=1.5,
                                )
                            )
                            is not None
                        ]
                        if not local_regions:
                            continue
                        page_dict = source_page.get_text("dict") or {}
                        _add_hybrid_visual_regions(
                            assessment,
                            page_dict,
                            source_page.rect,
                            local_regions,
                        )
                        localized_merge_pages.add(page_index)

                    if localized_merge_pages:
                        retry_flattened_source = (
                            folder / "region-hybrid-source-retry.pdf"
                        )
                        retry_regions = [
                            region
                            for assessment in assessments
                            if assessment.page_index not in risk_reasons
                            for region in assessment.visual_regions
                        ]
                        _flatten_pdf_regions(
                            source,
                            retry_flattened_source,
                            retry_regions,
                            password=password,
                            dpi=int(dpi),
                        )
                        retry_source = retry_flattened_source
                        if anchored_regions:
                            retry_source = folder / "region-hybrid-masked-retry.pdf"
                            _mask_pdf_regions_for_editable_conversion(
                                retry_flattened_source,
                                retry_source,
                                anchored_regions,
                                password=None,
                            )
                        parse_candidate_source(retry_source, None)
                except Exception:
                    for page_index in candidate_indexes:
                        risk_reasons.setdefault(page_index, []).append(
                            "可编辑引擎初始化或版面分析失败，已自动原样回退"
                        )
                else:
                    for candidate_position, page_index in enumerate(
                        candidate_indexes, start=1
                    ):
                        check_cancelled("任务已取消；已完成的文件会保留")
                        report_progress(
                            0.55
                            + 0.10
                            * (candidate_position / max(1, len(candidate_indexes))),
                            f"校验页面 {page_index + 1}/{page_count}",
                        )
                        try:
                            parsed_page = converter.pages[page_index]
                            if not getattr(parsed_page, "finalized", False):
                                risk_reasons.setdefault(page_index, []).append(
                                    "可编辑版面分析失败"
                                )
                                continue
                            assessment = assessments[page_index]
                            quality_reason = _pdf2docx_page_quality_reason(
                                assessment.editable_source_text,
                                _extract_pdf2docx_page_text(parsed_page),
                                source_blocks=assessment.editable_text_blocks,
                            )
                            if quality_reason:
                                risk_reasons.setdefault(page_index, []).append(
                                    quality_reason
                                )
                            layout_reason = _pdf2docx_page_layout_risk_reason(
                                parsed_page
                            )
                            if layout_reason in {
                                "三栏及以上复杂分栏",
                                "页面内多次分栏切换",
                                "不等宽短侧栏与主正文混排",
                            }:
                                risk_reasons.setdefault(page_index, []).append(
                                    layout_reason
                                )
                            if _pdf2docx_page_has_intercolumn_text_merge(
                                parsed_page,
                                split_x=assessment.column_split_x,
                                paired_bands=_hybrid_editable_column_bands(assessment),
                            ):
                                risk_reasons.setdefault(page_index, []).append(
                                    "检测到跨栏文字合并，已整页保真避免阅读顺序混乱"
                                )
                        except Exception:
                            risk_reasons.setdefault(page_index, []).append(
                                "页级版面校验失败，已自动原样回退"
                            )

            if converter is not None:
                for candidate_position, page_index in enumerate(
                    candidate_indexes, start=1
                ):
                    check_cancelled("任务已取消；已完成的文件会保留")
                    report_progress(
                        0.65
                        + 0.09 * (candidate_position / max(1, len(candidate_indexes))),
                        f"预检 Word 页面 {page_index + 1}/{page_count}",
                    )
                    if page_index in risk_reasons:
                        continue
                    probe_path = folder / f"probe-page-{page_index + 1:06d}.docx"
                    try:
                        probe_document = Document()
                        _append_pdf2docx_page(
                            probe_document,
                            converter.pages[page_index],
                            is_first_page=True,
                        )
                        _stabilize_pdf2docx_paragraph_layout(probe_document)
                        probe_document.save(probe_path)
                        assessment = assessments[page_index]
                        postbuild_reason = _pdf2docx_page_quality_reason(
                            assessment.editable_source_text,
                            _extract_docx_text(probe_path),
                            source_blocks=assessment.editable_text_blocks,
                        )
                        if postbuild_reason:
                            risk_reasons.setdefault(page_index, []).append(
                                f"Word 构建后{postbuild_reason}，已自动原样回退"
                            )
                        layout_reasons = _pdf2docx_docx_layout_quality_reasons(
                            probe_path
                        )
                        if layout_reasons:
                            risk_reasons.setdefault(page_index, []).append(
                                "Word 页级排版结构预检未通过（"
                                f"{'；'.join(layout_reasons)}），已自动原样回退"
                            )
                    except Exception:
                        risk_reasons.setdefault(page_index, []).append(
                            "Word 页面构建预检失败，已自动原样回退"
                        )
                    finally:
                        probe_path.unlink(missing_ok=True)

            def build_once() -> tuple[Any, tuple[int, Exception] | None]:
                candidate_document = Document()
                build_failure: tuple[int, Exception] | None = None
                for page_index in range(page_count):
                    check_cancelled("任务已取消；已完成的文件会保留")
                    report_progress(
                        0.75 + 0.15 * ((page_index + 1) / max(1, page_count)),
                        f"生成 Word 页面 {page_index + 1}/{page_count}",
                    )
                    if page_index in risk_reasons:
                        _append_visual_pdf_page(
                            candidate_document,
                            pdf_document[page_index],
                            folder / f"page-{page_index + 1:06d}.png",
                            dpi=int(dpi),
                            is_first_page=page_index == 0,
                        )
                        continue
                    try:
                        existing_body_ids = {
                            id(child)
                            for child in _body_content_children(candidate_document)
                        }
                        _append_pdf2docx_page(
                            candidate_document,
                            converter.pages[page_index],
                            is_first_page=page_index == 0,
                        )
                        _append_page_relative_pdf_regions(
                            candidate_document,
                            pdf_document[page_index],
                            assessments[page_index].anchored_visual_regions,
                            folder,
                            existing_body_ids=existing_body_ids,
                        )
                    except Exception as exc:
                        build_failure = (page_index, exc)
                        break
                return candidate_document, build_failure

            document, build_failure = build_once()
            attempted_failures: set[int] = set()
            while build_failure is not None:
                failed_page, _failure = build_failure
                if failed_page in attempted_failures:
                    break
                attempted_failures.add(failed_page)
                risk_reasons.setdefault(failed_page, []).append(
                    "Word 页面构建失败，已自动原样回退"
                )
                document, build_failure = build_once()
            if build_failure is not None:
                for page_index in range(page_count):
                    risk_reasons.setdefault(page_index, []).append(
                        "混合页面构建仍不稳定，已自动整篇原样回退"
                    )
                document = _build_visual_docx_document(
                    pdf_document,
                    dpi=int(dpi),
                    folder=folder,
                    progress_start=0.75,
                    progress_span=0.15,
                )

            # Page construction has copied everything needed into python-docx.
            # Release the temporary PDF before later validation can fail and
            # before the working directory is removed.
            close_converter()

            # Font substitution can change glyph widths and therefore page
            # wrapping.  Resolve fonts before applying any pagination-sensitive
            # footer/section cleanup, then run a second structural pass after
            # table-grid repair.  Every helper is idempotent: already floated
            # furniture is skipped and spacing reserves are never compounded.
            report_progress(0.91, "稳定字体、表格与分页结构")
            _postprocess_pdf2docx_document(document)
            report_progress(0.96, "保存并复核混合保真 Word")
            document.save(temporary)
            if not temporary.is_file() or temporary.stat().st_size == 0:
                raise ValidationError("PDF 混合转换没有生成有效的 Word 文件")
            final_layout_reasons = _pdf2docx_docx_layout_quality_reasons(temporary)
            if final_layout_reasons:
                layout_message = (
                    "PDF 混合转换的最终成品排版结构二次校验未通过（"
                    f"{'；'.join(final_layout_reasons)}）。"
                )
                if low_quality_policy == "discard":
                    raise ValidationError(
                        f"{layout_message}为避免静默交付异常文档，本次结果未保存；"
                        "也可选择“仍保留并警告”。"
                    )
                pending_warnings.append(
                    f"{layout_message}已按用户选择保留该结果，请人工复核。"
                )
            report_progress(0.975, "使用 WPS 复核最终实际分页")
            pagination_reason, pagination_warning = _pdf2docx_wps_render_quality_result(
                source,
                temporary,
                expected_pages=page_count,
                password=password,
                progress=lambda ratio, message: report_progress(
                    0.975 + ratio * 0.02,
                    message,
                ),
            )
            if pagination_warning:
                pending_warnings.append(pagination_warning)
            if pagination_reason:
                if low_quality_policy == "discard":
                    raise ValidationError(
                        f"PDF 混合转换的最终成品二次检测未通过：{pagination_reason}。"
                        "为避免交付异常分页，本次结果未保存；"
                        "也可选择“仍保留并警告”。"
                    )
                pending_warnings.append(
                    f"PDF 混合转换的最终成品二次检测未通过：{pagination_reason}。"
                    "已按用户选择保留该结果，请人工复核对应页码。"
                )

            editable_indexes = [
                page_index
                for page_index in range(page_count)
                if page_index not in risk_reasons
            ]
            if editable_indexes:
                source_text = "\n".join(
                    assessments[page_index].editable_source_text
                    for page_index in editable_indexes
                )
                source_blocks = [
                    block
                    for page_index in editable_indexes
                    for block in assessments[page_index].editable_text_blocks
                ]
                output_text = _extract_docx_text(temporary)
                coverage = _text_sequence_coverage(source_text, output_text)
                english_word_recall = _english_word_multiset_recall(
                    source_text, output_text
                )
                adjacent_word_coverage = _block_local_adjacent_english_word_coverage(
                    source_blocks,
                    output_text,
                )
                quality_details = (
                    f"字符序列 {coverage:.0%}、英文词召回 {english_word_recall:.0%}、"
                    f"相邻词序 {adjacent_word_coverage:.0%}"
                )
                boundary_recovery_pass = _pdf_english_boundary_recovery_pass(
                    source_text,
                    output_text,
                    character_coverage=coverage,
                )
                low_quality = not boundary_recovery_pass and (
                    coverage < _MIN_EDITABLE_SEQUENCE_COVERAGE
                    or english_word_recall < _MIN_EDITABLE_ENGLISH_WORD_RECALL
                    or adjacent_word_coverage < _MIN_EDITABLE_ADJACENT_WORD_COVERAGE
                )
                if low_quality and low_quality_policy == "discard":
                    pending_warnings.append(
                        f"混合转换汇总文字校验未通过（{quality_details}）；"
                        "已保留通过页级校验的局部混合结果，避免因双栏汇总顺序差异将整篇误改为图片。",
                    )
                elif low_quality:
                    pending_warnings.append(
                        f"混合转换最终文字校验未通过（{quality_details}）；"
                        "已按用户选择保留可编辑结果，请人工复核。",
                    )
            report_progress(0.99, "完成混合保真质量校验")

            region_summary = _hybrid_region_summary(assessments, risk_reasons)
            if region_summary:
                pending_warnings.append(
                    "版式优先混合已仅将以下页内的公式、图表或复杂表格高清图像化，"
                    f"其余正文仍可编辑：{region_summary}。"
                )

            if risk_reasons:
                pending_warnings.append(
                    "版式优先混合仅将以下高风险页面整页高清原样保留："
                    f"{_hybrid_visual_page_summary(risk_reasons)}。"
                )
    finally:
        if pdf_document is not None:
            with contextlib.suppress(Exception):
                pdf_document.close()
        close_converter()

    for message in pending_warnings:
        warnings.warn(message, stacklevel=2)


def _convert_pdf_to_visual_docx(
    source: Path,
    target: Path,
    *,
    password: str | None,
    dpi: int,
) -> None:
    from docuforge.runner import report_progress

    if int(dpi) < 72 or int(dpi) > 600:
        raise ValidationError("高清原样模式的 DPI 必须在 72–600 之间")

    report_progress(0.04, "读取 PDF 页面")
    _pymupdf, pdf_document = _open_pymupdf_document(source, password)
    try:
        with atomic_output(target) as temporary, tempfile.TemporaryDirectory(
            prefix="docuforge-pdf-docx-"
        ) as folder_name:
            document = _build_visual_docx_document(
                pdf_document, dpi=int(dpi), folder=Path(folder_name)
            )
            report_progress(0.94, "保存高清原样 Word")
            document.save(temporary)
    finally:
        pdf_document.close()


def _execute_pdf_to_docx(
    source: Path,
    target: Path,
    *,
    password: str | None,
    mode: str,
    dpi: int,
    low_quality_policy: str,
    hybrid_force_visual_pages: str,
    column_layout: str,
) -> None:
    """Execute one already-normalized conversion in the current process."""

    normalized_mode = mode
    normalized_quality_policy = low_quality_policy
    normalized_column_layout = column_layout
    if normalized_mode == "editable":
        _convert_pdf_to_editable_docx(
            source,
            target,
            password=password,
            low_quality_policy=normalized_quality_policy,
            column_layout=normalized_column_layout,
        )
    elif normalized_mode == "hybrid":
        _convert_pdf_to_hybrid_docx(
            source,
            target,
            password=password,
            dpi=int(dpi),
            low_quality_policy=normalized_quality_policy,
            force_visual_pages=hybrid_force_visual_pages,
            column_layout=normalized_column_layout,
        )
    else:
        _convert_pdf_to_visual_docx(
            source,
            target,
            password=password,
            dpi=int(dpi),
        )


def _pdf_to_docx_worker(
    connection: object,
    cancel_event: object,
    payload: Mapping[str, Any],
) -> None:
    """Run conversion in a disposable process so Stop can be immediate."""

    from docuforge.runner import (
        _CURRENT_PROGRESS_REPORTER,
        TaskRunner,
        task_runner_context,
    )

    def send(message: Mapping[str, Any]) -> None:
        try:
            connection.send(dict(message))  # type: ignore[attr-defined]
        except (BrokenPipeError, EOFError, OSError):
            pass

    def forward_progress(
        fraction: float,
        stage: str,
        current_file: int | None,
        total_files: int | None,
    ) -> None:
        send(
            {
                "type": "progress",
                "fraction": float(fraction),
                "stage": str(stage),
                "current_file": current_file,
                "total_files": total_files,
            }
        )

    token = _CURRENT_PROGRESS_REPORTER.set(forward_progress)
    child_runner = TaskRunner(cancel_event=cancel_event)  # type: ignore[arg-type]

    def monitor_cancel() -> None:
        cancel_event.wait()  # type: ignore[attr-defined]
        child_runner.cancel()

    cancel_monitor = threading.Thread(
        target=monitor_cancel,
        name="docuforge-pdf-word-cancel-monitor",
        daemon=True,
    )
    cancel_monitor.start()
    try:
        with task_runner_context(child_runner):
            with warnings.catch_warnings(record=True) as records:
                warnings.simplefilter("always")
                _execute_pdf_to_docx(
                    Path(str(payload["source"])),
                    Path(str(payload["target"])),
                    password=(
                        str(payload["password"])
                        if payload.get("password") is not None
                        else None
                    ),
                    mode=str(payload["mode"]),
                    dpi=int(payload["dpi"]),
                    low_quality_policy=str(payload["low_quality_policy"]),
                    hybrid_force_visual_pages=str(
                        payload["hybrid_force_visual_pages"]
                    ),
                    column_layout=str(payload["column_layout"]),
                )
        send(
            {
                "type": "result",
                "ok": True,
                "warnings": [str(record.message) for record in records],
            }
        )
    except BaseException as exc:
        send(
            {
                "type": "result",
                "ok": False,
                "kind": type(exc).__name__,
                "message": str(exc) or type(exc).__name__,
            }
        )
    finally:
        _CURRENT_PROGRESS_REPORTER.reset(token)
        try:
            connection.close()  # type: ignore[attr-defined]
        except OSError:
            pass


def _run_pdf_to_docx_supervised(
    source: Path,
    target: Path,
    *,
    password: str | None,
    mode: str,
    dpi: int,
    low_quality_policy: str,
    hybrid_force_visual_pages: str,
    column_layout: str,
) -> None:
    import multiprocessing

    from docuforge.runner import (
        cancellation_callback,
        check_cancelled,
        report_progress,
    )

    context = multiprocessing.get_context("spawn")
    parent_connection, child_connection = context.Pipe(duplex=False)
    child_cancel_event = context.Event()
    process = context.Process(
        target=_pdf_to_docx_worker,
        args=(
            child_connection,
            child_cancel_event,
            {
                "source": str(source),
                "target": str(target),
                "password": password,
                "mode": mode,
                "dpi": int(dpi),
                "low_quality_policy": low_quality_policy,
                "hybrid_force_visual_pages": hybrid_force_visual_pages,
                "column_layout": column_layout,
            },
        ),
        name="docuforge-pdf-to-word",
        daemon=False,
    )
    check_cancelled("任务已取消；PDF 转 Word 尚未启动")
    try:
        process.start()
    except Exception as exc:
        parent_connection.close()
        child_connection.close()
        raise MissingEngineError(f"无法启动 PDF 转 Word 隔离进程：{exc}") from exc
    child_connection.close()

    result: Mapping[str, Any] | None = None
    deadline = time.monotonic() + 6 * 60 * 60

    def request_worker_stop() -> None:
        child_cancel_event.set()

    try:
        with cancellation_callback(request_worker_stop):
            while result is None:
                check_cancelled("任务已取消；正在立即终止 PDF 转 Word")
                if time.monotonic() >= deadline:
                    request_worker_stop()
                    raise MissingEngineError("PDF 转 Word 超时（6 小时）")
                try:
                    has_message = parent_connection.poll(0.05)
                except (EOFError, OSError):
                    has_message = False
                if has_message:
                    while True:
                        try:
                            message = parent_connection.recv()
                        except (EOFError, OSError):
                            message = None
                        if isinstance(message, Mapping):
                            if message.get("type") == "progress":
                                report_progress(
                                    float(message.get("fraction", 0.0)),
                                    str(message.get("stage") or "PDF 转 Word 处理中"),
                                    current_file=message.get("current_file"),
                                    total_files=message.get("total_files"),
                                )
                            elif message.get("type") == "result":
                                result = message
                        try:
                            if result is not None or not parent_connection.poll(0):
                                break
                        except (EOFError, OSError):
                            break
                if result is None and not process.is_alive():
                    try:
                        if parent_connection.poll(0):
                            continue
                    except (EOFError, OSError):
                        pass
                    break
    finally:
        parent_connection.close()
        process.join(0.6)
        if process.is_alive():
            process.terminate()
            process.join(0.8)
        if process.is_alive() and hasattr(process, "kill"):
            process.kill()
            process.join(0.8)

    check_cancelled("任务已取消；PDF 转 Word 已终止")
    if result is None:
        raise MissingEngineError("PDF 转 Word 隔离进程意外退出")
    if bool(result.get("ok")):
        for message in result.get("warnings", ()):
            warnings.warn(str(message), stacklevel=2)
        return
    message = str(result.get("message") or "PDF 转 Word 失败")
    kind = str(result.get("kind") or "DocuForgeError")
    if kind == "ValidationError":
        raise ValidationError(message)
    if kind == "MissingEngineError":
        raise MissingEngineError(message)
    if kind == "CancelledError":
        raise CancelledError(message)
    raise DocuForgeError(message)


def pdf_to_docx(
    input_pdf: str | Path,
    output_path: str | Path,
    *,
    password: str | None = None,
    mode: str = "hybrid",
    dpi: int = 300,
    low_quality_policy: str = "discard",
    hybrid_force_visual_pages: str = "",
    column_layout: str = "auto",
    overwrite: bool = False,
) -> list[Path]:
    """Convert a PDF to an editable, hybrid, or pixel-faithful DOCX."""

    from docuforge.runner import report_progress, task_runner_active

    source = Path(input_pdf)
    target = unique_path(output_path, overwrite)
    normalized_mode = mode.lower().strip()
    if normalized_mode not in {"editable", "hybrid", "visual"}:
        raise ValidationError("Word 内容模式必须是 editable、hybrid 或 visual")
    normalized_quality_policy = str(low_quality_policy).lower().strip()
    if normalized_quality_policy not in {"discard", "keep"}:
        raise ValidationError("低质量结果处理策略必须是 discard 或 keep")
    normalized_column_layout = _normalize_pdf2docx_column_layout(column_layout)
    arguments = {
        "password": password,
        "mode": normalized_mode,
        "dpi": int(dpi),
        "low_quality_policy": normalized_quality_policy,
        "hybrid_force_visual_pages": hybrid_force_visual_pages,
        "column_layout": normalized_column_layout,
    }
    if task_runner_active():
        _run_pdf_to_docx_supervised(source, target, **arguments)
    else:
        _execute_pdf_to_docx(source, target, **arguments)
    report_progress(1.0, "PDF 转 Word 完成")
    return [target]


def pdf_to_html(
    input_pdf: str | Path,
    output_path: str | Path,
    *,
    password: str | None = None,
    include_tables: bool = True,
    overwrite: bool = False,
) -> list[Path]:
    """Create a semantic, readable HTML document; it intentionally does not mimic fixed PDF layout."""

    try:
        import pdfplumber
    except ImportError as exc:
        raise MissingEngineError("PDF 转 HTML 需要 pdfplumber") from exc
    source = Path(input_pdf)
    target = unique_path(output_path, overwrite)
    sections: list[str] = []
    with pdfplumber.open(source, password=password) as pdf:
        for page_number, page in enumerate(pdf.pages, 1):
            found_tables = page.find_tables() if include_tables else []
            table_boxes = [table.bbox for table in found_tables]

            def outside_tables(obj: Mapping[str, Any]) -> bool:
                if obj.get("object_type") != "char":
                    return True
                center_x = (float(obj.get("x0", 0)) + float(obj.get("x1", 0))) / 2
                center_y = (float(obj.get("top", 0)) + float(obj.get("bottom", 0))) / 2
                return not any(
                    left <= center_x <= right and top <= center_y <= bottom
                    for left, top, right, bottom in table_boxes
                )

            text_page = page.filter(outside_tables) if table_boxes else page
            text = text_page.extract_text(layout=False) or ""
            paragraphs = "\n".join(
                f"<p>{html.escape(line)}</p>"
                for line in text.splitlines()
                if line.strip()
            )
            tables_html = ""
            if include_tables:
                rendered_tables: list[str] = []
                for found_table in found_tables:
                    table = found_table.extract()
                    rows = []
                    for row in table:
                        cells = "".join(
                            f"<td>{html.escape(cell or '')}</td>" for cell in row
                        )
                        rows.append(f"<tr>{cells}</tr>")
                    rendered_tables.append(f"<table>{''.join(rows)}</table>")
                tables_html = "".join(rendered_tables)
            if not paragraphs and not tables_html:
                paragraphs = '<p class="warning">本页未检测到文字，可能需要 OCR。</p>'
            sections.append(
                f'<section class="page"><h2>第 {page_number} 页</h2>{paragraphs}{tables_html}</section>'
            )
    document = f"""<!doctype html>
<html lang="zh-CN"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width">
<title>{html.escape(source.stem)}</title><style>
body{{font-family:"Microsoft YaHei",system-ui,sans-serif;max-width:960px;margin:32px auto;padding:0 20px;color:#172033;line-height:1.7}}
.page{{border-bottom:1px solid #dde3ec;padding:12px 0 28px}} h2{{font-size:18px;color:#2563eb}}
table{{border-collapse:collapse;margin:14px 0;max-width:100%;overflow:auto}}td{{border:1px solid #aeb8c8;padding:5px 8px}}
.warning{{color:#b54708;background:#fff3e0;padding:10px}}</style></head><body>
<h1>{html.escape(source.stem)}</h1>{''.join(sections)}</body></html>"""
    with atomic_output(target) as temporary:
        temporary.write_text(document, encoding="utf-8")
    return [target]


def fill_pdf_form(
    input_pdf: str | Path,
    output_path: str | Path,
    fields: Mapping[str, Any] | str,
    *,
    password: str | None = None,
    flatten: bool = False,
    overwrite: bool = False,
) -> list[Path]:
    from pypdf import PdfWriter

    if isinstance(fields, str):
        try:
            values = json.loads(fields)
        except json.JSONDecodeError as exc:
            raise ValidationError(f"表单字段必须是 JSON 对象：{exc}") from exc
    else:
        values = dict(fields)
    if not isinstance(values, dict):
        raise ValidationError("表单字段必须是 JSON 对象")
    reader = _unlock_reader(Path(input_pdf), password)
    writer = PdfWriter()
    writer.clone_document_from_reader(reader)
    fields_found = reader.get_fields() or {}
    if not fields_found:
        raise ValidationError("该 PDF 没有可填写的 AcroForm 表单字段")
    unknown = sorted(set(values) - set(fields_found))
    if unknown:
        raise ValidationError(f"未找到表单字段：{unknown[0]}")
    writer.update_page_form_field_values(
        None, values, auto_regenerate=True, flatten=flatten
    )
    target = unique_path(output_path, overwrite)
    with atomic_output(target) as temporary:
        with temporary.open("wb") as stream:
            writer.write(stream)
    return [target]


def add_pdf_note(
    input_pdf: str | Path,
    output_path: str | Path,
    *,
    page: int,
    text: str,
    x: float = 36,
    y: float = 36,
    width: float = 220,
    height: float = 90,
    password: str | None = None,
    overwrite: bool = False,
) -> list[Path]:
    from pypdf import PdfWriter
    from pypdf.annotations import Text

    reader = _unlock_reader(Path(input_pdf), password)
    page_index = int(page) - 1
    if page_index < 0 or page_index >= len(reader.pages):
        raise ValidationError(f"页码超出范围（共 {len(reader.pages)} 页）")
    writer = PdfWriter()
    writer.clone_document_from_reader(reader)
    annotation = Text(
        text=str(text),
        rect=(float(x), float(y), float(x) + float(width), float(y) + float(height)),
        open=False,
    )
    writer.add_annotation(page_number=page_index, annotation=annotation)
    target = unique_path(output_path, overwrite)
    with atomic_output(target) as temporary:
        with temporary.open("wb") as stream:
            writer.write(stream)
    return [target]


def add_pdf_markup(
    input_pdf: str | Path,
    output_path: str | Path,
    *,
    page: int,
    kind: str = "highlight",
    x: float = 36,
    y: float = 36,
    width: float = 220,
    height: float = 24,
    color: str = "#ffff00",
    opacity: float = 0.45,
    comment: str = "",
    password: str | None = None,
    overwrite: bool = False,
) -> list[Path]:
    """Add a standard PDF highlight, underline, or strikeout annotation."""

    from PIL import ImageColor
    from pypdf import PdfWriter
    from pypdf.generic import (
        ArrayObject,
        DictionaryObject,
        FloatObject,
        NameObject,
        NumberObject,
        TextStringObject,
    )

    normalized_kind = kind.strip().lower()
    subtype = {
        "highlight": "/Highlight",
        "underline": "/Underline",
        "strikeout": "/StrikeOut",
    }.get(normalized_kind)
    if subtype is None:
        raise ValidationError("标记类型仅支持 highlight、underline 或 strikeout")
    if float(width) <= 0 or float(height) <= 0:
        raise ValidationError("标记区域宽度和高度必须大于 0")
    if not 0 <= float(opacity) <= 1:
        raise ValidationError("标记透明度必须在 0 到 1 之间")
    try:
        red, green, blue = ImageColor.getrgb(color)
    except ValueError as exc:
        raise ValidationError(f"无效标记颜色：{color}") from exc

    reader = _unlock_reader(Path(input_pdf), password)
    try:
        page_index = int(page) - 1
        if page_index < 0 or page_index >= len(reader.pages):
            raise ValidationError(f"页码超出范围（共 {len(reader.pages)} 页）")
        left = float(x)
        bottom = float(y)
        right = left + float(width)
        top = bottom + float(height)
        target_page = reader.pages[page_index]
        page_width = float(target_page.mediabox.width)
        page_height = float(target_page.mediabox.height)
        if left < 0 or bottom < 0 or right > page_width or top > page_height:
            raise ValidationError(
                f"标记区域超出页面边界；该页尺寸为 {page_width:g} × {page_height:g} 点"
            )
        annotation = DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Annot"),
                NameObject("/Subtype"): NameObject(subtype),
                NameObject("/Rect"): ArrayObject(
                    [
                        FloatObject(left),
                        FloatObject(bottom),
                        FloatObject(right),
                        FloatObject(top),
                    ]
                ),
                NameObject("/QuadPoints"): ArrayObject(
                    [
                        FloatObject(left),
                        FloatObject(top),
                        FloatObject(right),
                        FloatObject(top),
                        FloatObject(left),
                        FloatObject(bottom),
                        FloatObject(right),
                        FloatObject(bottom),
                    ]
                ),
                NameObject("/C"): ArrayObject(
                    [
                        FloatObject(red / 255),
                        FloatObject(green / 255),
                        FloatObject(blue / 255),
                    ]
                ),
                NameObject("/CA"): FloatObject(float(opacity)),
                NameObject("/F"): NumberObject(4),
                NameObject("/T"): TextStringObject("LayoutLoom"),
                NameObject("/Contents"): TextStringObject(str(comment)),
            }
        )
        writer = PdfWriter()
        writer.clone_document_from_reader(reader)
        writer.add_annotation(page_number=page_index, annotation=annotation)
        target = unique_path(output_path, overwrite)
        with atomic_output(target) as temporary:
            with temporary.open("wb") as stream:
                writer.write(stream)
        writer.close()
        return [target]
    finally:
        reader.close()


def ocr_pdf(
    input_pdf: str | Path,
    output_path: str | Path,
    *,
    output_format: str = "docx",
    language: str = "chi_sim+eng",
    dpi: int = 300,
    password: str | None = None,
    overwrite: bool = False,
) -> list[Path]:
    """OCR a scanned PDF to editable TXT or DOCX using a locally installed engine."""

    try:
        from pdf2image import convert_from_path
    except ImportError as exc:
        raise MissingEngineError("OCR 需要 pdf2image 和 Poppler") from exc
    from ..engines import find_executable, has_module, poppler_bin_path

    engine = ""
    pytesseract = None
    paddle = None
    if find_executable("tesseract") and has_module("pytesseract"):
        import pytesseract as pytesseract_module

        pytesseract = pytesseract_module
        engine = "tesseract"
    elif has_module("paddleocr"):
        try:
            from paddleocr import PaddleOCR

            paddle_lang = (
                "ch" if language.lower().startswith(("chi", "zh", "ch")) else "en"
            )
            paddle = PaddleOCR(use_angle_cls=True, lang=paddle_lang, show_log=False)
            engine = "paddleocr"
        except Exception as exc:
            raise MissingEngineError(f"PaddleOCR 初始化失败：{exc}") from exc
    else:
        raise MissingEngineError(
            "未检测到 OCR 引擎；请安装 Tesseract + pytesseract 或 PaddleOCR"
        )

    source = Path(input_pdf)
    if not source.is_file():
        raise ValidationError(f"PDF 不存在：{source}")
    if int(dpi) < 150 or int(dpi) > 600:
        raise ValidationError("OCR DPI 建议并限制在 150–600 之间")
    reader = _unlock_reader(source, password)
    try:
        page_count = len(reader.pages)
    finally:
        reader.close()
    if page_count < 1:
        raise ValidationError("PDF 没有可识别页面")
    texts: list[str] = []
    for page_number in range(1, page_count + 1):
        rendered = convert_from_path(
            source,
            dpi=int(dpi),
            fmt="png",
            poppler_path=poppler_bin_path(),
            thread_count=1,
            userpw=password,
            first_page=page_number,
            last_page=page_number,
        )
        if len(rendered) != 1:
            raise ValidationError(f"PDF 第 {page_number} 页渲染失败")
        page = rendered[0]
        try:
            if engine == "tesseract":
                assert pytesseract is not None
                try:
                    text = pytesseract.image_to_string(page, lang=language)
                except Exception as exc:
                    raise MissingEngineError(
                        f"Tesseract 识别失败：{exc}。请检查语言包 {language} 是否已安装。"
                    ) from exc
            else:
                assert paddle is not None
                try:
                    import numpy as np

                    with page.convert("RGB") as rgb_page:
                        raw = paddle.ocr(np.asarray(rgb_page), cls=True)
                    lines = raw[0] if raw and isinstance(raw, list) else raw
                    extracted: list[str] = []
                    for item in lines or []:
                        if isinstance(item, (list, tuple)) and len(item) >= 2:
                            payload = item[1]
                            if isinstance(payload, (list, tuple)) and payload:
                                extracted.append(str(payload[0]))
                    text = "\n".join(extracted)
                except Exception as exc:
                    raise MissingEngineError(
                        f"PaddleOCR 识别失败或 API 不兼容：{exc}"
                    ) from exc
            texts.append(text.strip())
        finally:
            page.close()

    fmt = output_format.lower().lstrip(".")
    target = unique_path(output_path, overwrite)
    if fmt == "txt":
        with atomic_output(target) as temporary:
            temporary.write_text("\n\n\f\n\n".join(texts), encoding="utf-8")
    elif fmt == "docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise MissingEngineError("OCR 输出 Word 需要 python-docx") from exc
        document = Document()
        for index, text in enumerate(texts, 1):
            document.add_heading(f"第 {index} 页", level=2)
            for line in text.splitlines():
                document.add_paragraph(line)
            if index < len(texts):
                document.add_page_break()
        with atomic_output(target) as temporary:
            document.save(temporary)
    else:
        raise ValidationError("OCR 输出格式仅支持 txt 或 docx")
    return [target]


def add_visual_signature(
    input_pdf: str | Path,
    signature_image: str | Path,
    output_path: str | Path,
    *,
    pages: str = "1",
    x: float = 36,
    y: float = 36,
    width: float = 120,
    opacity: float = 1.0,
    password: str | None = None,
    overwrite: bool = False,
) -> list[Path]:
    from pypdf import PdfWriter
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen.canvas import Canvas

    image_path = Path(signature_image)
    if not image_path.is_file():
        raise ValidationError(f"签名图片不存在：{image_path}")
    reader = _unlock_reader(Path(input_pdf), password)
    selected = set(parse_page_spec(pages, len(reader.pages)))
    writer = PdfWriter()
    writer.clone_document_from_reader(reader)
    with Image.open(image_path) as signature:
        aspect = signature.height / max(signature.width, 1)
    height = float(width) * aspect
    for index, target_page in enumerate(writer.pages):
        if index in selected:
            buffer = io.BytesIO()
            page_width = float(target_page.mediabox.width)
            page_height = float(target_page.mediabox.height)
            canvas = Canvas(buffer, pagesize=(page_width, page_height))
            if hasattr(canvas, "setFillAlpha"):
                canvas.setFillAlpha(max(0.0, min(1.0, float(opacity))))
            canvas.drawImage(
                ImageReader(str(image_path)),
                float(x),
                float(y),
                width=float(width),
                height=height,
                preserveAspectRatio=True,
                mask="auto",
            )
            canvas.save()
            buffer.seek(0)
            overlay = _unlock_reader_from_stream(buffer).pages[0]
            target_page.merge_page(overlay)
    target = unique_path(output_path, overwrite)
    with atomic_output(target) as temporary:
        with temporary.open("wb") as stream:
            writer.write(stream)
    return [target]


def _unlock_reader_from_stream(stream: io.BytesIO):
    from pypdf import PdfReader

    return PdfReader(stream)


def remove_background(
    inputs: Iterable[str | Path], output_dir: str | Path, *, overwrite: bool = False
) -> list[Path]:
    try:
        from rembg import new_session, remove
    except ImportError as exc:
        raise MissingEngineError("AI 抠图需要安装 rembg 和本地模型") from exc
    try:
        session = new_session()
    except Exception as exc:
        raise MissingEngineError(
            "AI 抠图模型初始化失败；首次使用可能需要联网下载模型，之后可离线复用。"
        ) from exc
    output = Path(output_dir)
    output.mkdir(parents=True, exist_ok=True)
    results: list[Path] = []
    for raw in inputs:
        source = Path(raw)
        target = unique_path(output / f"{source.stem}_透明背景.png", overwrite)
        with Image.open(source) as image:
            rgba = image.convert("RGBA")
            try:
                result = remove(rgba, session=session)
            finally:
                rgba.close()
        if isinstance(result, bytes):
            with Image.open(io.BytesIO(result)) as rendered:
                with atomic_output(target) as temporary:
                    rendered.save(temporary, "PNG")
        else:
            try:
                with atomic_output(target) as temporary:
                    result.save(temporary, "PNG")
            finally:
                result.close()
        results.append(target)
    return results


def heic_to_images(
    inputs: Iterable[str | Path], output_dir: str | Path, *, target_format: str = "jpg"
) -> list[Path]:
    try:
        from pillow_heif import register_heif_opener
    except ImportError as exc:
        raise MissingEngineError("HEIC 转换需要 pillow-heif") from exc
    register_heif_opener()
    from .image import convert_format

    return convert_format(list(inputs), target_format, output_dir)


def raw_to_images(
    inputs: Iterable[str | Path], output_dir: str | Path, *, target_format: str = "jpg"
) -> list[Path]:
    try:
        import rawpy
    except ImportError as exc:
        raise MissingEngineError("RAW 转换需要 rawpy/LibRaw") from exc
    output = Path(output_dir)
    output.mkdir(parents=True, exist_ok=True)
    fmt = target_format.lower()
    results: list[Path] = []
    for raw in inputs:
        source = Path(raw)
        suffix = ".png" if fmt == "png" else ".jpg"
        target = unique_path(output / f"{source.stem}{suffix}")
        with rawpy.imread(str(source)) as raw_image:
            rgb = raw_image.postprocess(
                use_camera_wb=True, no_auto_bright=False, output_bps=8
            )
        image = Image.fromarray(rgb)
        with atomic_output(target) as temporary:
            image.save(temporary, "PNG" if fmt == "png" else "JPEG", quality=95)
        image.close()
        results.append(target)
    return results


def svg_to_images(
    inputs: Iterable[str | Path],
    output_dir: str | Path,
    *,
    target_format: str = "png",
    scale: float = 1.0,
) -> list[Path]:
    try:
        import cairosvg
    except ImportError as exc:
        raise MissingEngineError("SVG 转位图需要 CairoSVG") from exc
    output = Path(output_dir)
    output.mkdir(parents=True, exist_ok=True)
    normalized_format = target_format.lower().lstrip(".")
    if normalized_format not in {"png", "jpg", "jpeg"}:
        raise ValidationError("SVG 目标格式仅支持 PNG 或 JPG")
    scale_value = float(scale)
    if not math.isfinite(scale_value) or scale_value <= 0 or scale_value > 20:
        raise ValidationError("SVG 渲染倍率必须是 0 到 20 之间的有限数字")
    results: list[Path] = []
    for raw in inputs:
        source = Path(raw)
        if not source.is_file():
            raise ValidationError(f"SVG 不存在：{source}")
        svg_bytes = source.read_bytes()
        if len(svg_bytes) > 50 * 1024 * 1024:
            raise ValidationError("SVG 文件超过 50 MB 安全上限")
        _validate_svg_resources(svg_bytes)
        if normalized_format in {"jpg", "jpeg"}:
            jpg_target = unique_path(output / f"{source.stem}.jpg")
            with tempfile.TemporaryDirectory(
                prefix="docuforge-svg-", dir=output
            ) as folder_name:
                temporary_png = Path(folder_name) / "rendered.png"
                cairosvg.svg2png(
                    bytestring=svg_bytes,
                    write_to=str(temporary_png),
                    scale=scale_value,
                    unsafe=False,
                )
                with Image.open(temporary_png) as image:
                    flattened = Image.new("RGB", image.size, "white")
                    if image.mode == "RGBA":
                        flattened.paste(image, mask=image.getchannel("A"))
                    else:
                        flattened.paste(image.convert("RGB"))
                    with atomic_output(jpg_target) as temporary:
                        flattened.save(temporary, "JPEG", quality=95)
                    flattened.close()
            results.append(jpg_target)
        else:
            png_target = unique_path(output / f"{source.stem}.png")
            with atomic_output(png_target) as temporary:
                cairosvg.svg2png(
                    bytestring=svg_bytes,
                    write_to=str(temporary),
                    scale=scale_value,
                    unsafe=False,
                )
            results.append(png_target)
    return results
