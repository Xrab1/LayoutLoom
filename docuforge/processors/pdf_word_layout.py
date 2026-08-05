"""Fixed-coordinate editable Word reconstruction for layout-critical PDFs.

Flow-based PDF converters are the right default for ordinary prose, but they
cannot reliably rebalance scientific two-column pages or design-heavy resumes.
This module provides the complementary path: preserve non-text page artwork as
a high-resolution background and place reliable source text back at its PDF
coordinates as editable Word frames.  One PDF page therefore remains one Word
page without sacrificing the editability of ordinary text.
"""

from __future__ import annotations

import io
import re
from collections import defaultdict
from dataclasses import dataclass, field
from pathlib import Path
from statistics import median
from typing import Any, Callable, Iterable, Mapping

from docuforge.models import MissingEngineError, ValidationError

PdfRect = tuple[float, float, float, float]
VisualHint = tuple[str, int, int, int, int]


@dataclass(frozen=True)
class FixedLayoutPagePlan:
    """Describe text that must remain visual on one source page."""

    excluded_regions: tuple[PdfRect, ...] = ()
    full_visual: bool = False


@dataclass(frozen=True)
class FixedLayoutBuildStats:
    pages: int
    editable_frames: int
    editable_spans: int
    visual_spans: int
    region_optimized: bool = False
    region_text_boxes: int = 0
    visual_hints_by_page: Mapping[int, tuple[VisualHint, ...]] = field(
        default_factory=dict
    )


@dataclass(frozen=True)
class _EditableSpan:
    text: str
    font: str
    size: float
    flags: int
    color: int
    bbox: PdfRect


@dataclass(frozen=True)
class _EditableCluster:
    spans: tuple[_EditableSpan, ...]
    bbox: PdfRect


@dataclass(frozen=True)
class _VisualSpanHint:
    """A source glyph island intentionally retained in the visual layer."""

    kind: str
    bbox: PdfRect


@dataclass(frozen=True)
class _CharacterAdvanceModel:
    by_character: Mapping[tuple[str, float, str], float]
    by_font: Mapping[tuple[str, float], float]

    def baseline(self, font: str, size: float, character: str) -> float:
        bucket = _font_size_bucket(size)
        return self.by_character.get(
            (font, bucket, character),
            self.by_font.get((font, bucket), max(1.0, float(size) * 0.5)),
        )


def _rect_area(rect: PdfRect) -> float:
    return max(0.0, rect[2] - rect[0]) * max(0.0, rect[3] - rect[1])


def _intersection_ratio(rect: PdfRect, region: PdfRect) -> float:
    intersection = (
        max(rect[0], region[0]),
        max(rect[1], region[1]),
        min(rect[2], region[2]),
        min(rect[3], region[3]),
    )
    return _rect_area(intersection) / max(1.0, _rect_area(rect))


def _span_is_excluded(rect: PdfRect, regions: Iterable[PdfRect]) -> bool:
    center_x = (rect[0] + rect[2]) / 2.0
    center_y = (rect[1] + rect[3]) / 2.0
    return any(
        _intersection_ratio(rect, region) >= 0.18
        or (region[0] <= center_x <= region[2] and region[1] <= center_y <= region[3])
        for region in regions
    )


def _coerce_bbox(value: Any) -> PdfRect | None:
    try:
        rect = tuple(float(item) for item in value)
    except (TypeError, ValueError):
        return None
    if len(rect) != 4 or rect[2] <= rect[0] or rect[3] <= rect[1]:
        return None
    return rect  # type: ignore[return-value]


def _merge_visual_span_hints(
    hints: Iterable[_VisualSpanHint],
) -> list[_VisualSpanHint]:
    """Conservatively join adjacent visual glyphs from the same source line."""

    ordered = sorted(
        hints,
        key=lambda hint: (
            hint.bbox[0],
            hint.bbox[1],
            hint.bbox[2],
            hint.bbox[3],
        ),
    )
    merged: list[_VisualSpanHint] = []
    for hint in ordered:
        if not merged:
            merged.append(hint)
            continue
        previous = merged[-1]
        previous_height = max(0.1, previous.bbox[3] - previous.bbox[1])
        hint_height = max(0.1, hint.bbox[3] - hint.bbox[1])
        maximum_height = max(previous_height, hint_height)
        center_delta = abs(
            (previous.bbox[1] + previous.bbox[3]) / 2.0
            - (hint.bbox[1] + hint.bbox[3]) / 2.0
        )
        horizontal_gap = hint.bbox[0] - previous.bbox[2]
        same_line = center_delta <= max(2.0, maximum_height * 0.68)
        nearby = horizontal_gap <= max(3.0, min(10.0, maximum_height * 0.8))
        if previous.kind == hint.kind and same_line and nearby:
            merged[-1] = _VisualSpanHint(
                kind=previous.kind,
                bbox=(
                    min(previous.bbox[0], hint.bbox[0]),
                    min(previous.bbox[1], hint.bbox[1]),
                    max(previous.bbox[2], hint.bbox[2]),
                    max(previous.bbox[3], hint.bbox[3]),
                ),
            )
        else:
            merged.append(hint)
    return merged


def _visual_hints_to_page_twips(
    hints: Iterable[_VisualSpanHint],
    page_rect: Any,
) -> tuple[VisualHint, ...]:
    """Convert PDF coordinates into the page-local twips used by Word."""

    page_x0 = float(page_rect.x0)
    page_y0 = float(page_rect.y0)
    page_width = max(1, round(float(page_rect.width) * 20.0))
    page_height = max(1, round(float(page_rect.height) * 20.0))
    converted: list[VisualHint] = []
    for hint in hints:
        x0 = max(0, min(page_width, round((hint.bbox[0] - page_x0) * 20.0)))
        y0 = max(0, min(page_height, round((hint.bbox[1] - page_y0) * 20.0)))
        x1 = max(0, min(page_width, round((hint.bbox[2] - page_x0) * 20.0)))
        y1 = max(0, min(page_height, round((hint.bbox[3] - page_y0) * 20.0)))
        if x1 <= x0 or y1 <= y0:
            continue
        converted.append((hint.kind, x0, y0, x1, y1))
    return tuple(converted)


def _font_size_bucket(size: float) -> float:
    return round(max(1.0, float(size)) * 2.0) / 2.0


def _lower_quantile(values: Iterable[float], ratio: float) -> float:
    ordered = sorted(float(value) for value in values)
    if not ordered:
        return 0.0
    index = max(0, min(len(ordered) - 1, int((len(ordered) - 1) * ratio)))
    return ordered[index]


def _line_character_records(line: Mapping[str, Any]) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for span_index, raw_span in enumerate(line.get("spans", ())):
        font = str(raw_span.get("font", "") or "")
        size = max(1.0, float(raw_span.get("size", 10.0) or 10.0))
        for character_index, raw_character in enumerate(raw_span.get("chars", ())):
            text = str(raw_character.get("c", "") or "")
            origin = raw_character.get("origin", ())
            try:
                origin_x, origin_y = (float(value) for value in origin)
            except (TypeError, ValueError):
                continue
            records.append(
                {
                    "span_index": span_index,
                    "character_index": character_index,
                    "text": text,
                    "origin": (origin_x, origin_y),
                    "font": font,
                    "size": size,
                }
            )
    return records


def _projected_character_gap(
    previous: Mapping[str, Any],
    current: Mapping[str, Any],
    direction: tuple[float, float],
) -> float:
    direction_x, direction_y = direction
    length = max(0.001, (direction_x**2 + direction_y**2) ** 0.5)
    delta_x = float(current["origin"][0]) - float(previous["origin"][0])
    delta_y = float(current["origin"][1]) - float(previous["origin"][1])
    return (delta_x * direction_x + delta_y * direction_y) / length


def _build_character_advance_model(
    page_dicts: Iterable[Mapping[str, Any]],
) -> _CharacterAdvanceModel:
    by_character_values: dict[tuple[str, float, str], list[float]] = defaultdict(list)
    by_font_values: dict[tuple[str, float], list[float]] = defaultdict(list)
    for page_dict in page_dicts:
        for block in page_dict.get("blocks", ()):
            if int(block.get("type", 0) or 0) != 0:
                continue
            for line in block.get("lines", ()):
                records = _line_character_records(line)
                raw_direction = line.get("dir", (1.0, 0.0))
                try:
                    direction = (float(raw_direction[0]), float(raw_direction[1]))
                except (IndexError, TypeError, ValueError):
                    direction = (1.0, 0.0)
                for previous, current in zip(records, records[1:]):
                    previous_text = str(previous["text"])
                    current_text = str(current["text"])
                    if (
                        not previous_text
                        or not current_text
                        or previous_text.isspace()
                        or current_text.isspace()
                    ):
                        continue
                    advance = _projected_character_gap(previous, current, direction)
                    size = float(previous["size"])
                    if not 0.05 <= advance <= size * 1.8:
                        continue
                    font_key = (str(previous["font"]), _font_size_bucket(size))
                    by_font_values[font_key].append(advance)
                    by_character_values[(*font_key, previous_text)].append(advance)

    by_font = {
        key: _lower_quantile(values, 0.35)
        for key, values in by_font_values.items()
    }
    by_character = {
        key: _lower_quantile(values, 0.25)
        for key, values in by_character_values.items()
        if len(values) >= 3
    }
    return _CharacterAdvanceModel(by_character=by_character, by_font=by_font)


def _english_segmentation_boundaries(text: str) -> set[int]:
    """Return likely word boundaries without trusting language statistics alone."""

    try:
        import wordninja
    except ImportError:
        return set()

    boundaries: set[int] = set()
    for match in re.finditer(r"[A-Za-z]{6,}", text):
        token = match.group(0)
        if token.isupper() and len(token) <= 12:
            continue
        try:
            parts = [part for part in wordninja.split(token) if part]
        except Exception:
            continue
        if len(parts) < 2 or sum(len(part) for part in parts) != len(token):
            continue
        offset = match.start()
        for part in parts[:-1]:
            offset += len(part)
            boundaries.add(offset)
    return boundaries


def _space_boundary_context(previous: str, current: str) -> bool:
    if not previous or not current or previous.isspace() or current.isspace():
        return False
    if not any(character.isascii() and character.isalpha() for character in (previous, current)):
        return False
    opening_characters = "([{'\"（【《“‘"
    previous_is_word_end = previous.isalnum() or (
        not previous.isspace() and previous not in opening_characters
    )
    current_is_word_start = current.isalnum() or current in opening_characters
    return previous_is_word_end and current_is_word_start


def _prepare_raw_line_spans(
    line: Mapping[str, Any],
    model: _CharacterAdvanceModel,
) -> list[dict[str, Any]]:
    """Restore visually present word gaps omitted by a PDF's text mapping."""

    prepared = [dict(raw_span) for raw_span in line.get("spans", ())]
    records = _line_character_records(line)
    if not records:
        return prepared

    raw_direction = line.get("dir", (1.0, 0.0))
    try:
        direction = (float(raw_direction[0]), float(raw_direction[1]))
    except (IndexError, TypeError, ValueError):
        direction = (1.0, 0.0)
    flattened_text = "".join(str(record["text"]) for record in records)
    language_boundaries = _english_segmentation_boundaries(flattened_text)
    insert_before: set[tuple[int, int]] = set()
    for index, (previous, current) in enumerate(zip(records, records[1:]), start=1):
        previous_text = str(previous["text"])
        current_text = str(current["text"])
        if not _space_boundary_context(previous_text, current_text):
            continue
        advance = _projected_character_gap(previous, current, direction)
        size = float(previous["size"])
        excess = advance - model.baseline(
            str(previous["font"]),
            size,
            previous_text,
        )
        statistically_segmented = index in language_boundaries and excess >= max(
            2.2,
            size * 0.20,
        )
        adjacent_ascii_letters = (
            previous_text.isascii()
            and previous_text.isalpha()
            and current_text.isascii()
            and current_text.isalpha()
        )
        # Letter kerning can make pairs such as "mn" look like a large gap.
        # For an all-letter token, require the language segmenter to agree;
        # geometry alone remains sufficient at punctuation/script boundaries.
        visually_separated = (
            not adjacent_ascii_letters and excess >= max(2.4, size * 0.28)
        )
        punctuation_separated = (
            previous_text in ",;:!?，；：！？"
            and current_text.isascii()
            and current_text.isalpha()
            and advance >= max(3.0, size * 0.45)
        )
        if statistically_segmented or visually_separated or punctuation_separated:
            insert_before.add(
                (int(current["span_index"]), int(current["character_index"]))
            )

    for span_index, raw_span in enumerate(prepared):
        raw_characters = raw_span.get("chars", ())
        if raw_characters:
            text_parts: list[str] = []
            for character_index, raw_character in enumerate(raw_characters):
                if (span_index, character_index) in insert_before:
                    text_parts.append(" ")
                text_parts.append(str(raw_character.get("c", "") or ""))
            raw_span["text"] = "".join(text_parts)
        else:
            raw_span["text"] = str(raw_span.get("text", "") or "")
    return prepared


def _editable_line_clusters(
    spans: Iterable[Mapping[str, Any]],
    *,
    plan: FixedLayoutPagePlan,
    repair_text: Callable[[str], str],
    font_requires_visual: Callable[[str], bool],
    math_font: Callable[[str], bool],
    suspicious_text: Callable[[str], bool],
    visual_hint_collector: list[_VisualSpanHint] | None = None,
) -> tuple[list[_EditableCluster], int]:
    """Return contiguous reliable spans while preserving visual-only gaps."""

    clusters: list[list[_EditableSpan]] = []
    active: list[_EditableSpan] = []
    active_visual: list[_VisualSpanHint] = []
    visual_spans = 0
    pending_space = False

    def flush() -> None:
        nonlocal active
        if active:
            clusters.append(active)
            active = []

    def flush_visual() -> None:
        nonlocal active_visual
        if active_visual and visual_hint_collector is not None:
            visual_hint_collector.extend(
                _merge_visual_span_hints(active_visual)
            )
        active_visual = []

    for raw_span in spans:
        raw_text = str(raw_span.get("text", "") or "")
        text = repair_text(raw_text)
        bbox = _coerce_bbox(raw_span.get("bbox", ()))
        font = str(raw_span.get("font", "") or "")
        if not text.strip():
            # PDF producers often emit an isolated whitespace span between
            # every word.  It belongs to the active editable line and must not
            # split that line into dozens of separate Word frames.
            pending_space = pending_space or bool(text)
            continue
        ordinary_punctuation = any(
            character in ".,;:!?()[]{}<>+-=/%&'\"，。．；：！？（）【】《》“”‘’—–…"
            for character in text
        )
        portable_font = not font_requires_visual(font)
        uses_math_font = math_font(font)
        suspicious = suspicious_text(text)
        editable_content = bool(
            any(character.isalnum() for character in text) or ordinary_punctuation
        )
        excluded = bool(
            bbox is not None
            and (
                plan.full_visual
                or _span_is_excluded(bbox, plan.excluded_regions)
            )
        )
        reliable = bool(
            bbox is not None
            and editable_content
            and portable_font
            and not uses_math_font
            and not suspicious
            and not excluded
        )
        if not reliable or bbox is None:
            if text.strip():
                visual_spans += 1
                if bbox is not None:
                    active_visual.append(
                        _VisualSpanHint(
                            kind=("inline_math" if uses_math_font else "text_visual"),
                            bbox=bbox,
                        )
                    )
            flush()
            pending_space = False
            continue

        flush_visual()
        span_text = text
        if active:
            previous = active[-1]
            gap = bbox[0] - previous.bbox[2]
            typical_size = median(item.size for item in active)
            # A large geometric gap is a separate text island (for example the
            # other half of a resume information row), not an ordinary space.
            if gap > max(7.0, typical_size * 1.45):
                flush()
        if (
            active
            and pending_space
            and not active[-1].text.endswith((" ", "\t"))
            and not span_text.startswith((" ", "\t"))
        ):
            span_text = " " + span_text
        pending_space = False

        span = _EditableSpan(
            text=span_text,
            font=font,
            size=max(1.0, float(raw_span.get("size", 10.0) or 10.0)),
            flags=int(raw_span.get("flags", 0) or 0),
            color=int(raw_span.get("color", 0) or 0),
            bbox=bbox,
        )
        active.append(span)
    flush()
    flush_visual()

    result: list[_EditableCluster] = []
    for cluster in clusters:
        result.append(
            _EditableCluster(
                spans=tuple(cluster),
                bbox=(
                    min(span.bbox[0] for span in cluster),
                    min(span.bbox[1] for span in cluster),
                    max(span.bbox[2] for span in cluster),
                    max(span.bbox[3] for span in cluster),
                ),
            )
        )
    return result, visual_spans


def _inline_picture_to_background_anchor(
    inline: Any,
    *,
    width_emu: int,
    height_emu: int,
) -> None:
    from docx.oxml import OxmlElement

    children = list(inline)
    by_local_name = {str(child.tag).rsplit("}", 1)[-1]: child for child in children}
    required = ("extent", "docPr", "cNvGraphicFramePr", "graphic")
    if any(name not in by_local_name for name in required):
        raise ValidationError("无法建立固定版面 Word 背景")

    anchor = OxmlElement("wp:anchor")
    for name, value in (
        ("distT", "0"),
        ("distB", "0"),
        ("distL", "0"),
        ("distR", "0"),
        ("simplePos", "0"),
        ("relativeHeight", "0"),
        ("behindDoc", "1"),
        ("locked", "1"),
        ("layoutInCell", "1"),
        ("allowOverlap", "1"),
    ):
        anchor.set(name, value)

    simple_position = OxmlElement("wp:simplePos")
    simple_position.set("x", "0")
    simple_position.set("y", "0")
    anchor.append(simple_position)
    for axis in ("H", "V"):
        position = OxmlElement(f"wp:position{axis}")
        position.set("relativeFrom", "page")
        position_offset = OxmlElement("wp:posOffset")
        position_offset.text = "0"
        position.append(position_offset)
        anchor.append(position)

    extent = by_local_name["extent"]
    extent.set("cx", str(max(1, int(width_emu))))
    extent.set("cy", str(max(1, int(height_emu))))
    anchor.append(extent)
    effect_extent = by_local_name.get("effectExtent")
    if effect_extent is None:
        effect_extent = OxmlElement("wp:effectExtent")
        for side in ("l", "t", "r", "b"):
            effect_extent.set(side, "0")
    anchor.append(effect_extent)
    anchor.append(OxmlElement("wp:wrapNone"))
    document_properties = by_local_name["docPr"]
    document_properties.set("descr", "LayoutLoom fixed-layout background")
    anchor.append(document_properties)
    anchor.append(by_local_name["cNvGraphicFramePr"])
    anchor.append(by_local_name["graphic"])

    drawing = inline.getparent()
    if drawing is None:
        raise ValidationError("无法定位固定版面 Word 背景容器")
    drawing.replace(inline, anchor)


def _set_zero_page_geometry(section: Any, page_rect: Any) -> None:
    from docx.shared import Pt

    section.page_width = Pt(float(page_rect.width))
    section.page_height = Pt(float(page_rect.height))
    section.top_margin = Pt(0)
    section.bottom_margin = Pt(0)
    section.left_margin = Pt(0)
    section.right_margin = Pt(0)
    section.header_distance = Pt(0)
    section.footer_distance = Pt(0)


def _append_background(document: Any, png: bytes, page_rect: Any) -> None:
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(1)
    inline = paragraph.add_run().add_picture(
        io.BytesIO(png),
        width=Pt(float(page_rect.width)),
        height=Pt(float(page_rect.height)),
    )
    _inline_picture_to_background_anchor(
        inline._inline,
        width_emu=int(Pt(float(page_rect.width))),
        height_emu=int(Pt(float(page_rect.height))),
    )


def _append_editable_cluster(
    document: Any,
    cluster: _EditableCluster,
    *,
    frame_id: int,
    resolve_font: Callable[[str, bool], str],
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    x0, y0, x1, y1 = cluster.bbox
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = Pt(max(1.0, y1 - y0))
    paragraph_format.keep_together = False
    paragraph_format.keep_with_next = False
    paragraph_format.page_break_before = False

    properties = paragraph._p.get_or_add_pPr()
    frame = OxmlElement("w:framePr")
    for name, value in (
        ("w:wrap", "none"),
        ("w:hAnchor", "page"),
        ("w:vAnchor", "page"),
        ("w:x", str(round(x0 * 20))),
        ("w:y", str(round(y0 * 20))),
        ("w:w", str(max(20, round((x1 - x0 + 1.2) * 20)))),
        ("w:h", str(max(20, round((y1 - y0 + 0.8) * 20)))),
        ("w:hRule", "atLeast"),
        ("w:anchorLock", "1"),
    ):
        frame.set(qn(name), value)
    properties.insert(0, frame)
    snap_to_grid = OxmlElement("w:snapToGrid")
    snap_to_grid.set(qn("w:val"), "0")
    properties.append(snap_to_grid)

    fit_width = str(max(20, round((x1 - x0) * 20)))
    text_parts: list[str] = []
    previous: _EditableSpan | None = None
    for span in cluster.spans:
        text = span.text
        if previous is not None:
            gap = max(0.0, span.bbox[0] - previous.bbox[2])
            if (
                gap > max(0.8, span.size * 0.18)
                and not str(previous.text).endswith((" ", "\t"))
                and not text.startswith((" ", "\t"))
                and _space_boundary_context(
                    str(previous.text)[-1:],
                    text[:1],
                )
            ):
                text = " " + text
        text_parts.append(text)
        previous = span

    combined_text = "".join(text_parts)
    dominant = max(
        cluster.spans,
        key=lambda item: (
            sum(character.isalnum() for character in item.text),
            len(item.text.strip()),
            item.size,
        ),
    )
    east_asia_characters = sum(
        "\u2e80" <= character <= "\u9fff"
        or "\uf900" <= character <= "\ufaff"
        or "\uff00" <= character <= "\uffef"
        for character in combined_text
    )
    latin_characters = sum(
        character.isascii() and character.isalnum() for character in combined_text
    )
    font_name = resolve_font(
        dominant.font,
        east_asia_characters > 0 and east_asia_characters >= latin_characters,
    )
    run = paragraph.add_run(combined_text)
    run.font.size = Pt(dominant.size)
    run.bold = bool(dominant.flags & 16)
    run.italic = bool(dominant.flags & 2)
    run.font.color.rgb = RGBColor(
        (dominant.color >> 16) & 255,
        (dominant.color >> 8) & 255,
        dominant.color & 255,
    )
    run_properties = run._r.get_or_add_rPr()
    fonts = run_properties.get_or_add_rFonts()
    for attribute in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(attribute), font_name)
    # WPS does not consistently merge adjacent runs that share one fitText id.
    # One fit operation per positioned line preserves the source width without
    # stretching every punctuation/font fragment to the entire frame.
    fit_text = OxmlElement("w:fitText")
    fit_text.set(qn("w:val"), fit_width)
    fit_text.set(qn("w:id"), str(frame_id % 65535))
    run_properties.append(fit_text)
    # Thousands of positioned runs can otherwise trigger a lengthy background
    # spell-check pass when WPS first opens an English-heavy document.
    run_properties.append(OxmlElement("w:noProof"))


def build_fixed_layout_docx(
    pdf_document: Any,
    target: Path,
    *,
    dpi: int,
    page_plans: Mapping[int, FixedLayoutPagePlan] | None,
    repair_text: Callable[[str], str],
    font_requires_visual: Callable[[str], bool],
    math_font: Callable[[str], bool],
    suspicious_text: Callable[[str], bool],
    resolve_font: Callable[[str, bool], str],
    progress: Callable[[float, str], None] | None = None,
) -> FixedLayoutBuildStats:
    """Build a one-source-page/one-Word-page fixed-coordinate document."""

    try:
        import pymupdf
        from docx import Document
        from docx.enum.section import WD_SECTION
    except ImportError as exc:
        raise MissingEngineError(
            "固定版面 PDF 转 Word 需要 PyMuPDF 与 python-docx"
        ) from exc

    requested_dpi = int(dpi)
    if requested_dpi < 72 or requested_dpi > 600:
        raise ValidationError("固定版面可编辑 Word 清晰度必须在 72–600 DPI 之间")
    # The public API historically accepts 72 DPI, while the fixed-layout path
    # relies on its background for figures and decoration.  Honor the request
    # without allowing that visual layer to fall below a readable 150 DPI.
    requested_dpi = max(150, requested_dpi)

    plans = dict(page_plans or {})
    page_count = int(pdf_document.page_count)
    if page_count <= 0:
        raise ValidationError("PDF 没有可转换页面")

    page_dicts = [
        pdf_document[page_index].get_text("rawdict") or {}
        for page_index in range(page_count)
    ]
    advance_model = _build_character_advance_model(page_dicts)
    working_document = pymupdf.open()
    working_document.insert_pdf(pdf_document)
    document = Document()
    editable_frames = 0
    editable_spans = 0
    visual_spans = 0
    visual_hints_by_page: dict[int, tuple[VisualHint, ...]] = {}
    frame_id = 1
    try:
        for page_index in range(page_count):
            if progress is not None:
                progress(
                    (page_index + 1) / page_count,
                    f"固定版面重建 {page_index + 1}/{page_count}",
                )
            source_page = pdf_document[page_index]
            background_page = working_document[page_index]
            plan = plans.get(page_index, FixedLayoutPagePlan())
            page_dict = page_dicts[page_index]
            page_clusters: list[_EditableCluster] = []
            page_visual_hints: list[_VisualSpanHint] = []
            page_visual_spans = 0
            for block in page_dict.get("blocks", ()):
                if int(block.get("type", 0) or 0) != 0:
                    continue
                for line in block.get("lines", ()):
                    clusters, excluded_count = _editable_line_clusters(
                        _prepare_raw_line_spans(line, advance_model),
                        plan=plan,
                        repair_text=repair_text,
                        font_requires_visual=font_requires_visual,
                        math_font=math_font,
                        suspicious_text=suspicious_text,
                        visual_hint_collector=page_visual_hints,
                    )
                    page_clusters.extend(clusters)
                    page_visual_spans += excluded_count

            for cluster in page_clusters:
                for span in cluster.spans:
                    background_page.add_redact_annot(
                        pymupdf.Rect(span.bbox),
                        fill=False,
                        cross_out=False,
                    )
            if page_clusters:
                background_page.apply_redactions(
                    images=pymupdf.PDF_REDACT_IMAGE_NONE,
                    graphics=pymupdf.PDF_REDACT_LINE_ART_NONE,
                    text=pymupdf.PDF_REDACT_TEXT_REMOVE,
                )

            pixmap = background_page.get_pixmap(
                dpi=requested_dpi,
                alpha=False,
                annots=True,
            )
            background_png = pixmap.tobytes("png")

            if page_index:
                section = document.add_section(WD_SECTION.NEW_PAGE)
            else:
                section = document.sections[0]
            _set_zero_page_geometry(section, source_page.rect)
            _append_background(document, background_png, source_page.rect)
            for cluster in page_clusters:
                _append_editable_cluster(
                    document,
                    cluster,
                    frame_id=frame_id,
                    resolve_font=resolve_font,
                )
                frame_id += 1
                editable_frames += 1
                editable_spans += len(cluster.spans)
            visual_spans += page_visual_spans
            visual_hints_by_page[page_index] = _visual_hints_to_page_twips(
                page_visual_hints,
                source_page.rect,
            )
    finally:
        working_document.close()

    document.save(target)
    if not target.is_file() or target.stat().st_size == 0:
        raise ValidationError("固定版面 PDF 转 Word 没有生成有效文件")
    return FixedLayoutBuildStats(
        pages=page_count,
        editable_frames=editable_frames,
        editable_spans=editable_spans,
        visual_spans=visual_spans,
        visual_hints_by_page=visual_hints_by_page,
    )


__all__ = [
    "FixedLayoutBuildStats",
    "FixedLayoutPagePlan",
    "build_fixed_layout_docx",
]
