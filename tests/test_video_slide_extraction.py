from __future__ import annotations

import tempfile
from pathlib import Path

import cv2
import numpy as np
import pytest
from docx import Document
from pptx import Presentation

from docuforge.models import ValidationError
from docuforge.processors import video_slides


def _lecture_frame(index: int, fps: int, width: int, height: int) -> np.ndarray:
    second = index / fps
    first_page = second < 3.0
    background = (75, 38, 12) if first_page else (24, 70, 24)
    frame = np.full((height, width, 3), background, dtype=np.uint8)
    title = "PAGE ONE" if first_page else "PAGE TWO"
    cv2.putText(
        frame,
        title,
        (45, 115),
        cv2.FONT_HERSHEY_SIMPLEX,
        1.25,
        (245, 245, 245),
        3,
        cv2.LINE_AA,
    )
    page_start = 0.0 if first_page else 3.0
    cv2.putText(
        frame,
        "PRINTED BASE",
        (55, 175),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.85,
        (235, 235, 235),
        2,
        cv2.LINE_AA,
    )
    if second - page_start >= 1.0:
        cv2.putText(
            frame,
            "FINAL PRINT",
            (55, 235),
            cv2.FONT_HERSHEY_SIMPLEX,
            0.9,
            (255, 255, 255),
            2,
            cv2.LINE_AA,
        )
    if second - page_start >= 1.5:
        progress = min(1.0, (second - page_start - 1.5) / 1.2)
        end_x = int(70 + progress * 320)
        cv2.line(frame, (70, 270), (end_x, 270), (255, 210, 0), 5)
        if progress > 0.20:
            arc = int(round(360 * min(1.0, (progress - 0.20) / 0.80)))
            cv2.ellipse(
                frame,
                (360, 245),
                (25, 25),
                0,
                0,
                arc,
                (255, 210, 0),
                4,
            )
    watermark_x = 20 + (index * 37) % (width - 250)
    cv2.putText(
        frame,
        "MOVING WM",
        (watermark_x, 38),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.7,
        (248, 248, 248),
        2,
        cv2.LINE_AA,
    )
    return frame


def _write_synthetic_lecture(path: Path) -> None:
    fps = 10
    width, height = 640, 360
    writer = cv2.VideoWriter(
        str(path), cv2.VideoWriter_fourcc(*"MJPG"), fps, (width, height)
    )
    if not writer.isOpened():
        pytest.skip("本机 OpenCV 没有可用的 MJPG 测试编码器")
    try:
        for index in range(60):
            writer.write(_lecture_frame(index, fps, width, height))
    finally:
        writer.release()


def _write_no_presenter_lecture(path: Path) -> np.ndarray:
    fps = 10
    width, height = 480, 270
    frame = np.full((height, width, 3), 245, dtype=np.uint8)
    cv2.putText(
        frame,
        "STATIC SLIDE",
        (35, 72),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.9,
        (25, 25, 25),
        2,
        cv2.LINE_AA,
    )
    cv2.putText(
        frame,
        "KEEP",
        (390, 235),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.55,
        (15, 15, 15),
        2,
        cv2.LINE_AA,
    )
    writer = cv2.VideoWriter(
        str(path), cv2.VideoWriter_fourcc(*"MJPG"), fps, (width, height)
    )
    if not writer.isOpened():
        pytest.skip("本机 OpenCV 没有可用的 MJPG 测试编码器")
    try:
        for _index in range(12):
            writer.write(frame)
    finally:
        writer.release()
    return frame


def _candidate_record(index: int, *, edge_pixels: int = 500) -> video_slides._ScanFrame:
    gray = np.zeros((90, 160), dtype=np.uint8)
    return video_slides._ScanFrame(
        frame_index=index,
        timestamp=float(index),
        gray=gray,
        packed_edges=np.packbits(np.zeros(gray.size, dtype=np.uint8)),
        edge_pixels=edge_pixels,
        phash=np.zeros(64, dtype=bool),
        sharpness=0.0,
    )


def _scan_record(index: int, image: np.ndarray) -> video_slides._ScanFrame:
    return video_slides._scan_frame_from_image(index, float(index), image, [])


def test_percent_region_validation_is_explicit_and_bounded() -> None:
    assert video_slides._parse_percent_boxes(
        "0,0,100,18; 70,62,30,38", "区域", multiple=True
    ) == [(0.0, 0.0, 1.0, 0.18), (0.7, 0.62, 0.3, 0.38)]
    with pytest.raises(ValidationError, match="不能超出"):
        video_slides._parse_percent_boxes("90,0,20,20", "区域", multiple=False)
    with pytest.raises(ValidationError, match="格式"):
        video_slides._parse_percent_boxes("1,2,3", "区域", multiple=False)


def test_colour_options_accept_multicolour_hex_and_rgb() -> None:
    mode, colours, tolerance = video_slides._validate_annotation_colour_options(
        "manual", "#00AEEF;255,60,60", 28
    )
    assert mode == "manual"
    assert colours == ((239, 174, 0), (60, 60, 255))
    assert tolerance == 28
    fill_mode, fill_colour = video_slides._validate_fixed_fill_options(
        "color", "#F8F7F2"
    )
    assert fill_mode == "color"
    assert fill_colour == (242, 247, 248)
    with pytest.raises(ValidationError, match="0–255"):
        video_slides._validate_annotation_colour_options(
            "manual", "300,20,20", 24
        )


def test_temporal_restore_uses_real_uncovered_pixels() -> None:
    height, width = 100, 180
    base = np.zeros((height, width, 3), dtype=np.uint8)
    base[:, :] = (30, 55, 90)
    cv2.putText(
        base,
        "SOURCE",
        (24, 62),
        cv2.FONT_HERSHEY_SIMPLEX,
        1.0,
        (220, 220, 220),
        2,
        cv2.LINE_AA,
    )
    frames = []
    for left in (0, 45, 90, 125):
        frame = base.copy()
        cv2.rectangle(
            frame, (left, 5), (min(width - 1, left + 45), 30), (250, 250, 250), -1
        )
        frames.append(frame)
    output = frames[-1].copy()
    restored, uncertain = video_slides._temporal_restore_region(
        output,
        frames,
        (0, 0, width, 38),
        variation_threshold=12,
        force=False,
    )
    assert restored > 500
    assert uncertain == 0
    mean_error = float(np.mean(np.abs(output[:38].astype(int) - base[:38].astype(int))))
    assert mean_error < 8.0


def test_temporal_restore_preserves_late_but_persistent_printed_text() -> None:
    height, width = 100, 240
    background = np.full((height, width, 3), (35, 52, 80), dtype=np.uint8)
    frames: list[np.ndarray] = []
    for index in range(6):
        frame = background.copy()
        if index >= 3:
            cv2.putText(
                frame,
                "FINAL",
                (25, 70),
                cv2.FONT_HERSHEY_SIMPLEX,
                1.0,
                (245, 245, 245),
                2,
                cv2.LINE_AA,
            )
        cv2.putText(
            frame,
            "WM",
            (5 + index * 32, 24),
            cv2.FONT_HERSHEY_SIMPLEX,
            0.55,
            (250, 250, 250),
            2,
            cv2.LINE_AA,
        )
        frames.append(frame)
    output = frames[-1].copy()
    video_slides._temporal_restore_region(
        output,
        frames,
        (0, 0, width, height),
        variation_threshold=15,
        force=False,
        selected_index=5,
    )
    printed = output[45:78, 20:125]
    assert int(np.count_nonzero(np.max(printed, axis=2) > 220)) > 180
    watermark_band = output[:32]
    assert int(np.count_nonzero(np.max(watermark_band, axis=2) > 220)) < 220


def test_candidate_selection_avoids_single_bright_watermark_outlier() -> None:
    height, width = 180, 360
    base = np.full((height, width, 3), (38, 46, 62), dtype=np.uint8)
    cv2.putText(
        base,
        "PRINTED BODY",
        (25, 120),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.8,
        (235, 235, 235),
        2,
        cv2.LINE_AA,
    )
    frames = [base.copy() for _index in range(5)]
    cv2.rectangle(frames[-1], (25, 8), (335, 34), (245, 245, 245), -1)
    candidates = [
        (_candidate_record(index), frame) for index, frame in enumerate(frames)
    ]

    selected = video_slides._select_final_candidate(
        candidates,
        excluded_boxes=[],
        watermark_boxes=[(0, 0, width, 55)],
    )

    assert selected == 3


def test_candidate_selection_preserves_late_persistent_white_title() -> None:
    height, width = 180, 360
    base = np.full((height, width, 3), (38, 46, 62), dtype=np.uint8)
    frames: list[np.ndarray] = []
    candidates: list[tuple[video_slides._ScanFrame, np.ndarray]] = []
    for index in range(5):
        frame = base.copy()
        if index >= 3:
            cv2.rectangle(frame, (25, 8), (335, 34), (245, 245, 245), -1)
        frames.append(frame)
        candidates.append(
            (
                _candidate_record(index, edge_pixels=1200 if index >= 3 else 100),
                frame,
            )
        )

    selected = video_slides._select_final_candidate(
        candidates,
        excluded_boxes=[],
        watermark_boxes=[(0, 0, width, 55)],
    )

    assert selected == 4


def test_candidate_selection_prefers_clean_frame_over_slow_white_ticker() -> None:
    height, width = 220, 520
    base = np.full((height, width, 3), (34, 48, 66), dtype=np.uint8)
    cv2.putText(
        base,
        "COMPLETE PRINTED QUESTION",
        (26, 112),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.72,
        (238, 238, 238),
        2,
        cv2.LINE_AA,
    )
    candidates: list[tuple[video_slides._ScanFrame, np.ndarray]] = []
    for index in range(9):
        frame = base.copy()
        if index >= 5:
            cv2.putText(
                frame,
                "WM 3849178",
                (24 + (index - 5) * 2, 30),
                cv2.FONT_HERSHEY_SIMPLEX,
                0.46,
                (248, 248, 248),
                1,
                cv2.LINE_AA,
            )
        candidates.append((_candidate_record(index), frame))

    selected = video_slides._select_final_candidate(
        candidates,
        excluded_boxes=[],
        watermark_boxes=[(0, 0, width, height)],
        stability_support_override=np.full(len(candidates), 5, dtype=np.int16),
    )

    assert selected < 5


def test_manual_colour_prior_reincludes_sparse_clean_history_and_keeps_print() -> None:
    height, width = 220, 480
    ink = (239, 174, 0)  # #00AEEF in OpenCV BGR order.
    clean_page = np.full((height, width, 3), (34, 48, 66), dtype=np.uint8)
    cv2.putText(
        clean_page,
        "PRINTED CYAN",
        (24, 58),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.82,
        ink,
        2,
        cv2.LINE_AA,
    )
    cv2.arrowedLine(clean_page, (30, 92), (115, 92), ink, 4, cv2.LINE_AA)
    cv2.putText(
        clean_page,
        "AUTHENTIC BODY",
        (28, 142),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.8,
        (240, 240, 240),
        2,
        cv2.LINE_AA,
    )
    timestamps = (10.0, 15.0, 20.0, 95.0, 95.22, 95.44, 95.66)
    candidates: list[tuple[video_slides._ScanFrame, np.ndarray]] = []
    for index, timestamp in enumerate(timestamps):
        frame = clean_page.copy()
        if index >= 3:
            points = np.asarray(
                [[40, 178], [125, 165], [220, 188], [330, 166], [430, 192]],
                dtype=np.int32,
            )
            cv2.polylines(frame, [points], False, ink, 6, cv2.LINE_AA)
            cv2.circle(frame, (390, 116), 30, ink, 5, cv2.LINE_AA)
        scan = video_slides._scan_frame_from_image(index, timestamp, frame, [])
        candidates.append((scan, frame))

    clean = video_slides._clean_slide(
        candidates,
        watermark_boxes=[],
        fixed_boxes=[],
        presenter_boxes=[],
        watermark_hint="",
        annotation_mode="manual",
        annotation_colours=(ink,),
        annotation_tolerance=24,
    )

    assert clean is not None
    assert clean.selected_timestamp <= 20.0
    printed_mask = (
        np.max(
            np.abs(
                clean_page.astype(np.int16)
                - np.full_like(clean_page, (34, 48, 66)).astype(np.int16)
            ),
            axis=2,
        )
        >= 24
    )
    retained_print = (
        np.max(
            np.abs(
                clean.image.astype(np.int16)
                - np.full_like(clean.image, (34, 48, 66)).astype(np.int16)
            ),
            axis=2,
        )
        >= 24
    )
    assert int(np.count_nonzero(retained_print & printed_mask)) >= int(
        np.count_nonzero(printed_mask) * 0.96
    )
    assert float(np.mean(np.abs(clean.image.astype(int) - clean_page.astype(int)))) < 1.0


def test_high_sharpness_handwriting_cannot_overpower_manual_cleanliness() -> None:
    height, width = 210, 480
    ink = (239, 174, 0)
    base = np.full((height, width, 3), (34, 48, 66), dtype=np.uint8)
    cv2.putText(
        base,
        "COMPLETE PRINTED PAGE",
        (28, 92),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.82,
        (240, 240, 240),
        2,
        cv2.LINE_AA,
    )
    candidates: list[tuple[video_slides._ScanFrame, np.ndarray]] = []
    for index in range(6):
        frame = base.copy()
        if index >= 3:
            for offset in range(0, 120, 12):
                cv2.line(
                    frame,
                    (35 + offset, 145),
                    (125 + offset, 178),
                    ink,
                    5,
                    cv2.LINE_AA,
                )
            cv2.putText(
                frame,
                "WM 3849173 IMPORTANT",
                (8, 32),
                cv2.FONT_HERSHEY_SIMPLEX,
                0.55,
                (248, 248, 248),
                2,
                cv2.LINE_AA,
            )
        record = video_slides._scan_frame_from_image(
            index, index * 0.25, frame, []
        )
        record = video_slides._ScanFrame(
            frame_index=record.frame_index,
            timestamp=record.timestamp,
            gray=record.gray,
            packed_edges=record.packed_edges,
            edge_pixels=record.edge_pixels,
            phash=record.phash,
            sharpness=1_000.0 if index < 3 else 50_000.0,
        )
        candidates.append((record, frame))

    selected = video_slides._select_final_candidate(
        candidates,
        excluded_boxes=[],
        watermark_boxes=[(0, 0, width, 55)],
        annotation_mode="manual",
        annotation_colours=(ink,),
        annotation_tolerance=24,
        stability_support_override=np.full(len(candidates), 5, dtype=np.int16),
    )

    assert selected == 2


def test_full_history_modal_watermark_restore_keeps_late_printed_answer() -> None:
    height, width = 210, 480
    base = np.full((height, width, 3), (34, 48, 66), dtype=np.uint8)
    cv2.putText(
        base,
        "QUESTION BODY",
        (28, 92),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.85,
        (235, 235, 235),
        2,
        cv2.LINE_AA,
    )
    printed = base.copy()
    cv2.putText(
        printed,
        "ANSWER: CD",
        (42, 158),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.9,
        (245, 245, 245),
        2,
        cv2.LINE_AA,
    )
    timestamps = (10.0, 20.0, 30.0, 95.0, 95.24, 95.48, 95.72)
    candidates: list[tuple[video_slides._ScanFrame, np.ndarray]] = []
    for index, timestamp in enumerate(timestamps):
        frame = base.copy() if index < 3 else printed.copy()
        if index >= 3:
            cv2.putText(
                frame,
                "WM",
                (8 + (index - 3) * 92, 32),
                cv2.FONT_HERSHEY_SIMPLEX,
                0.62,
                (248, 248, 248),
                2,
                cv2.LINE_AA,
            )
        scan = video_slides._scan_frame_from_image(index, timestamp, frame, [])
        candidates.append((scan, frame))

    clean = video_slides._clean_slide(
        candidates,
        watermark_boxes=[(0.0, 0.0, 1.0, 0.24)],
        fixed_boxes=[],
        presenter_boxes=[],
        watermark_hint="WM",
        annotation_mode="off",
    )

    assert clean is not None
    assert clean.selected_timestamp >= 95.0
    assert int(np.count_nonzero(np.max(clean.image[:48], axis=2) >= 220)) < 90
    answer_mask = (
        np.max(np.abs(printed.astype(np.int16) - base.astype(np.int16)), axis=2) >= 24
    )
    answer_error = float(
        np.mean(
            np.abs(
                clean.image[answer_mask].astype(np.int16)
                - printed[answer_mask].astype(np.int16)
            )
        )
    )
    assert answer_error < 6.0


def test_slow_scrolling_white_watermark_is_not_protected_as_print() -> None:
    height, width = 240, 640
    base = np.full((height, width, 3), (30, 46, 72), dtype=np.uint8)
    cv2.putText(
        base,
        "AUTHENTIC SLIDE BODY",
        (36, 112),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.9,
        (235, 235, 235),
        2,
        cv2.LINE_AA,
    )
    printed = base.copy()
    cv2.putText(
        printed,
        "ANSWER CD",
        (48, 198),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.9,
        (245, 245, 245),
        2,
        cv2.LINE_AA,
    )
    frames: list[np.ndarray] = []
    for index in range(24):
        frame = printed.copy() if index >= 20 else base.copy()
        if index >= 12:
            cv2.putText(
                frame,
                "WM 3849173 IMPORTANT",
                (-180 + (index - 12) * 34, 34),
                cv2.FONT_HERSHEY_SIMPLEX,
                0.58,
                (248, 248, 248),
                2,
                cv2.LINE_AA,
            )
        frames.append(frame)
    output = frames[-1].copy()
    watermark_before = int(
        np.count_nonzero(np.max(output[:52], axis=2) >= 220)
    )

    restored, uncertain = video_slides._temporal_restore_region(
        output,
        frames,
        (0, 0, width, height),
        variation_threshold=15,
        force=False,
        selected_index=len(frames) - 1,
    )

    watermark_after = int(np.count_nonzero(np.max(output[:52], axis=2) >= 220))
    assert watermark_before > 250
    assert restored > 500
    assert uncertain == 0
    assert watermark_after <= int(watermark_before * 0.12)
    answer_mask = (
        np.max(np.abs(printed.astype(np.int16) - base.astype(np.int16)), axis=2) >= 24
    )
    assert float(
        np.mean(
            np.abs(
                output[answer_mask].astype(np.int16)
                - printed[answer_mask].astype(np.int16)
            )
        )
    ) < 7.0


def test_very_slow_scrolling_white_watermark_is_not_protected_as_print() -> None:
    """A nearly stationary ticker must not look like late printed text.

    Real lecture ads can move only one pixel between dense candidates.  The
    same bright glyph then survives three or more consecutive frames, which is
    exactly the persistence signal used to protect a late printed answer.  Its
    surrounding local motion must veto that protection.
    """

    height, width = 240, 640
    base = np.full((height, width, 3), (30, 46, 72), dtype=np.uint8)
    cv2.putText(
        base,
        "AUTHENTIC SLIDE BODY",
        (36, 112),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.9,
        (235, 235, 235),
        2,
        cv2.LINE_AA,
    )
    printed = base.copy()
    cv2.putText(
        printed,
        "ANSWER CD",
        (48, 198),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.9,
        (245, 245, 245),
        2,
        cv2.LINE_AA,
    )
    frames: list[np.ndarray] = []
    for index in range(24):
        frame = printed.copy() if index >= 20 else base.copy()
        if index >= 6:
            cv2.putText(
                frame,
                "WM 3849173 IMPORTANT",
                (16 + (index - 6), 34),
                cv2.FONT_HERSHEY_SIMPLEX,
                0.58,
                (248, 248, 248),
                2,
                cv2.LINE_AA,
            )
        frames.append(frame)
    output = frames[-1].copy()
    watermark_before = int(np.count_nonzero(np.max(output[:52], axis=2) >= 220))

    restored, uncertain = video_slides._temporal_restore_region(
        output,
        frames,
        (0, 0, width, height),
        variation_threshold=15,
        force=False,
        selected_index=len(frames) - 1,
    )

    watermark_after = int(np.count_nonzero(np.max(output[:52], axis=2) >= 220))
    assert watermark_before > 250
    assert restored > 500
    assert uncertain == 0
    assert watermark_after <= int(watermark_before * 0.12)
    answer_mask = (
        np.max(np.abs(printed.astype(np.int16) - base.astype(np.int16)), axis=2) >= 24
    )
    assert float(
        np.mean(
            np.abs(
                output[answer_mask].astype(np.int16)
                - printed[answer_mask].astype(np.int16)
            )
        )
    ) < 7.0


def test_slow_ticker_cleanup_preserves_progressively_revealed_body_text() -> None:
    height, width = 300, 720
    background = np.full((height, width, 3), (28, 44, 70), dtype=np.uint8)
    final_print = background.copy()
    lines = (
        ("QUESTION BODY LINE", 110),
        ("OPTION A PRINTED", 170),
        ("OPTION B PRINTED", 230),
    )
    for text_value, baseline in lines:
        cv2.putText(
            final_print,
            text_value,
            (55, baseline),
            cv2.FONT_HERSHEY_SIMPLEX,
            0.85,
            (245, 245, 245),
            2,
            cv2.LINE_AA,
        )
    frames: list[np.ndarray] = []
    for index in range(18):
        frame = background.copy()
        visible_lines = min(len(lines), max(0, (index - 3) // 3 + 1))
        for text_value, baseline in lines[:visible_lines]:
            cv2.putText(
                frame,
                text_value,
                (55, baseline),
                cv2.FONT_HERSHEY_SIMPLEX,
                0.85,
                (245, 245, 245),
                2,
                cv2.LINE_AA,
            )
        cv2.putText(
            frame,
            "LONG MOVING WATERMARK 3849178",
            (-110 + index * 5, 28),
            cv2.FONT_HERSHEY_SIMPLEX,
            0.52,
            (248, 248, 248),
            2,
            cv2.LINE_AA,
        )
        frames.append(frame)
    output = frames[-1].copy()

    video_slides._temporal_restore_region(
        output,
        frames,
        (0, 0, width, height),
        variation_threshold=15,
        force=False,
        selected_index=len(frames) - 1,
    )

    print_mask = (
        np.max(
            np.abs(final_print.astype(np.int16) - background.astype(np.int16)),
            axis=2,
        )
        >= 20
    )
    assert float(
        np.mean(
            np.abs(
                output[print_mask].astype(np.int16)
                - final_print[print_mask].astype(np.int16)
            )
        )
    ) < 7.0


def test_fixed_watermark_background_and_colour_fill_are_feathered() -> None:
    height, width = 150, 260
    background = np.full((height, width, 3), (40, 72, 104), dtype=np.uint8)
    logo_box = (0.04, 0.05, 0.26, 0.24)
    left, top, right, bottom = video_slides._pixel_box(logo_box, width, height)
    logo = background.copy()
    cv2.rectangle(logo, (left, top), (right - 1, bottom - 1), (245, 245, 245), -1)
    candidates = []
    for index, timestamp in enumerate((0.0, 0.3, 0.6, 0.9)):
        scan = video_slides._scan_frame_from_image(index, timestamp, logo, [logo_box])
        candidates.append((scan, logo.copy()))

    modelled = video_slides._clean_slide(
        candidates,
        watermark_boxes=[],
        fixed_boxes=[logo_box],
        presenter_boxes=[],
        watermark_hint="",
        fixed_fill="background",
    )
    coloured = video_slides._clean_slide(
        candidates,
        watermark_boxes=[],
        fixed_boxes=[logo_box],
        presenter_boxes=[],
        watermark_hint="",
        fixed_fill="color",
        fixed_fill_colour=(18, 36, 54),
    )

    assert modelled is not None and coloured is not None
    centre = (slice(top + 6, bottom - 6), slice(left + 6, right - 6))
    assert float(np.mean(np.abs(modelled.image[centre].astype(int) - background[centre].astype(int)))) < 5.0
    expected_colour = np.asarray((18, 36, 54), dtype=np.int16)
    assert float(np.mean(np.abs(coloured.image[centre].astype(np.int16) - expected_colour))) < 2.0
    assert modelled.low_confidence_pixels == (right - left) * (bottom - top)
    assert coloured.low_confidence_pixels == (right - left) * (bottom - top)


def test_candidate_selection_rejects_high_edge_crossfade_before_stable_page() -> None:
    height, width = 180, 360
    old_page = np.full((height, width, 3), (28, 42, 66), dtype=np.uint8)
    for top in range(25, 170, 22):
        cv2.putText(
            old_page,
            f"OLD PAGE {top}",
            (10, top),
            cv2.FONT_HERSHEY_SIMPLEX,
            0.45,
            (245, 245, 245),
            1,
            cv2.LINE_AA,
        )
    new_page = np.full((height, width, 3), 235, dtype=np.uint8)
    cv2.putText(
        new_page,
        "NEW PAGE",
        (35, 105),
        cv2.FONT_HERSHEY_SIMPLEX,
        1.2,
        (20, 20, 20),
        3,
        cv2.LINE_AA,
    )
    crossfade = cv2.addWeighted(old_page, 0.55, new_page, 0.45, 0)
    candidates = [
        (_candidate_record(0, edge_pixels=7000), crossfade),
        (_candidate_record(1, edge_pixels=700), new_page.copy()),
        (_candidate_record(2, edge_pixels=700), new_page.copy()),
        (_candidate_record(3, edge_pixels=700), new_page.copy()),
    ]

    support, quiet_pairs = video_slides._candidate_stability_support(
        candidates, excluded_boxes=[]
    )
    selected = video_slides._select_final_candidate(
        candidates,
        excluded_boxes=[],
    )

    assert quiet_pairs.tolist() == [False, True, True]
    assert int(support[0]) == 0
    assert bool(np.all(support[1:] > 0))
    assert selected == 3


def test_pure_transition_candidates_have_no_stable_support() -> None:
    height, width = 180, 360
    old_page = np.full((height, width, 3), (20, 25, 35), dtype=np.uint8)
    new_page = np.full((height, width, 3), (235, 240, 245), dtype=np.uint8)
    cv2.putText(
        old_page,
        "OLD",
        (25, 110),
        cv2.FONT_HERSHEY_SIMPLEX,
        2.0,
        (250, 250, 250),
        4,
        cv2.LINE_AA,
    )
    cv2.putText(
        new_page,
        "NEW",
        (180, 110),
        cv2.FONT_HERSHEY_SIMPLEX,
        2.0,
        (10, 10, 10),
        4,
        cv2.LINE_AA,
    )
    candidates = [
        (
            _candidate_record(index, edge_pixels=1800),
            cv2.addWeighted(old_page, 1.0 - alpha, new_page, alpha, 0),
        )
        for index, alpha in enumerate((0.20, 0.45, 0.70, 0.95))
    ]

    support, quiet_pairs = video_slides._candidate_stability_support(
        candidates, excluded_boxes=[]
    )

    assert not bool(np.any(quiet_pairs))
    assert not bool(np.any(support))


def test_scan_segmentation_detects_slow_cumulative_transition(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    scans = [_candidate_record(index) for index in range(9)]

    def synthetic_difference(
        first: video_slides._ScanFrame,
        second: video_slides._ScanFrame,
        _valid: np.ndarray,
    ) -> tuple[float, float, float, int]:
        if second.frame_index - first.frame_index == 1:
            return (0.020, 0.008, 0.010, 2)
        crosses_dissolve = first.frame_index <= 2 and second.frame_index >= 4
        if crosses_dissolve:
            return (0.120, 0.070, 0.100, 16)
        return (0.025, 0.009, 0.012, 3)

    monkeypatch.setattr(video_slides, "_scan_difference", synthetic_difference)
    monkeypatch.setattr(video_slides, "_edge_containment", lambda _a, _b: 0.25)

    segments = video_slides._segment_scans(
        scans, ignored_boxes=[], sensitivity="balanced"
    )

    assert [(segment.first, segment.last) for segment in segments] == [(0, 3), (4, 8)]


def test_incremental_printed_reveal_is_merged_into_final_slide() -> None:
    base = np.full((180, 320, 3), 245, dtype=np.uint8)
    cv2.putText(
        base, "QUESTION", (24, 52), cv2.FONT_HERSHEY_SIMPLEX, 0.9, (20, 20, 20), 2
    )
    final = base.copy()
    for row, text_value in enumerate(("FIRST LINE", "SECOND LINE", "FINAL ANSWER")):
        cv2.putText(
            final,
            text_value,
            (28, 90 + row * 30),
            cv2.FONT_HERSHEY_SIMPLEX,
            0.65,
            (25, 25, 25),
            2,
            cv2.LINE_AA,
        )
    scans = [
        _scan_record(0, base),
        _scan_record(1, base),
        _scan_record(2, base),
        _scan_record(3, final),
        _scan_record(4, final),
        _scan_record(5, final),
    ]
    valid = np.ones(scans[0].gray.shape, dtype=bool)

    merged = video_slides._merge_short_and_incremental_segments(
        [video_slides._SlideSegment(0, 2), video_slides._SlideSegment(3, 5)],
        scans,
        valid,
    )

    assert [(segment.first, segment.last) for segment in merged] == [(0, 5)]


def test_replaced_printed_page_is_not_misclassified_as_incremental_reveal() -> None:
    first = np.full((180, 320, 3), 245, dtype=np.uint8)
    second = first.copy()
    cv2.putText(
        first, "OLD SUBJECT", (25, 105), cv2.FONT_HERSHEY_SIMPLEX, 1.0, (20, 20, 20), 3
    )
    cv2.putText(
        second, "NEW RESULT", (85, 105), cv2.FONT_HERSHEY_SIMPLEX, 1.0, (20, 20, 20), 3
    )
    scans = [
        _scan_record(0, first),
        _scan_record(1, first),
        _scan_record(2, second),
        _scan_record(3, second),
    ]
    valid = np.ones(scans[0].gray.shape, dtype=bool)

    merged = video_slides._merge_short_and_incremental_segments(
        [video_slides._SlideSegment(0, 1), video_slides._SlideSegment(2, 3)],
        scans,
        valid,
    )

    assert [(segment.first, segment.last) for segment in merged] == [(0, 1), (2, 3)]


def test_template_similar_pages_use_pre_boundary_stable_frame(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    scans = [_candidate_record(index) for index in range(5)]
    scans = [
        video_slides._ScanFrame(
            frame_index=record.frame_index,
            timestamp=float(index),
            gray=record.gray,
            packed_edges=record.packed_edges,
            edge_pixels=record.edge_pixels,
            phash=record.phash,
            sharpness=record.sharpness,
        )
        for index, record in enumerate(scans)
    ]

    def synthetic_difference(
        first: video_slides._ScanFrame,
        _second: video_slides._ScanFrame,
        _valid: np.ndarray,
    ) -> tuple[float, float, float, int]:
        # Frame 2 is already contaminated by the next page and would satisfy
        # the old progressive-reveal rule.  Frame 1 is the real stable tail.
        if first.frame_index == 2:
            return (0.10, 0.035, 0.08, 8)
        return (0.16, 0.058, 0.10, 13)

    monkeypatch.setattr(video_slides, "_scan_difference", synthetic_difference)
    monkeypatch.setattr(
        video_slides,
        "_edge_containment",
        lambda first, _second: 0.98 if first.frame_index == 2 else 0.93,
    )
    valid = np.ones(scans[0].gray.shape, dtype=bool)

    merged = video_slides._merge_short_and_incremental_segments(
        [video_slides._SlideSegment(0, 2), video_slides._SlideSegment(3, 4)],
        scans,
        valid,
    )

    assert [(segment.first, segment.last) for segment in merged] == [(0, 2), (3, 4)]


def test_slow_sparse_crossfade_selects_only_the_settled_endpoint() -> None:
    height, width = 180, 360
    old_page = np.full((height, width, 3), 242, dtype=np.uint8)
    new_page = old_page.copy()
    cv2.putText(
        old_page,
        "OLD CONTENT",
        (12, 72),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.8,
        (20, 20, 20),
        2,
        cv2.LINE_AA,
    )
    cv2.putText(
        new_page,
        "NEW CONTENT",
        (118, 128),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.8,
        (20, 20, 20),
        2,
        cv2.LINE_AA,
    )
    alphas = [index / 10.0 for index in range(11)] + [1.0] * 6
    candidates = []
    for index, alpha in enumerate(alphas):
        frame = cv2.addWeighted(old_page, 1.0 - alpha, new_page, alpha, 0)
        candidates.append(
            (
                video_slides._scan_frame_from_image(
                    index,
                    index / 8.0,
                    frame,
                    ignored_boxes=[],
                ),
                frame,
            )
        )

    clean = video_slides._clean_slide(
        candidates,
        watermark_boxes=[],
        fixed_boxes=[],
        presenter_boxes=[],
        watermark_hint="",
    )

    assert clean is not None
    assert clean.selected_timestamp == pytest.approx((len(alphas) - 1) / 8.0)
    assert float(np.mean(np.abs(clean.image.astype(int) - new_page.astype(int)))) < 1.0


def test_sparse_candidates_several_seconds_apart_do_not_form_a_plateau() -> None:
    frame = np.full((180, 360, 3), 235, dtype=np.uint8)
    cv2.putText(
        frame,
        "SAME LOOK",
        (45, 105),
        cv2.FONT_HERSHEY_SIMPLEX,
        1.0,
        (20, 20, 20),
        2,
        cv2.LINE_AA,
    )
    timestamps = (0.0, 0.12, 0.24, 6.0, 6.12, 6.24)
    candidates = [
        (
            video_slides._scan_frame_from_image(
                index,
                timestamp,
                frame,
                ignored_boxes=[],
            ),
            frame.copy(),
        )
        for index, timestamp in enumerate(timestamps)
    ]

    support, quiet_pairs = video_slides._candidate_stability_support(
        candidates, excluded_boxes=[]
    )

    assert not bool(quiet_pairs[2])
    assert not bool(np.any(support))


def test_decoded_frame_index_never_accepts_one_frame_backend_drift() -> None:
    class Capture:
        def __init__(self, position: float) -> None:
            self.position = position

        def get(self, _property: int) -> float:
            return self.position

    assert video_slides._decoded_frame_index(Capture(101.0), 100) == 100
    assert video_slides._decoded_frame_index(Capture(102.0), 100) == 100
    assert video_slides._decoded_frame_index(Capture(100.0), 100) == 100


def test_stable_cyan_printed_title_is_not_treated_as_handwriting() -> None:
    height, width = 180, 360
    printed = np.full((height, width, 3), (34, 48, 66), dtype=np.uint8)
    cv2.putText(
        printed,
        "CYAN TITLE",
        (20, 100),
        cv2.FONT_HERSHEY_SIMPLEX,
        1.0,
        (235, 180, 20),
        3,
        cv2.LINE_AA,
    )
    frames = [printed.copy() for _index in range(4)]
    output = printed.copy()

    static_mask = video_slides._colored_handwriting_mask(printed, excluded_boxes=[])
    restored = video_slides._restore_transient_colored_marks(
        output,
        frames,
        selected_index=len(frames) - 1,
        excluded_boxes=[],
    )

    assert int(np.count_nonzero(static_mask)) == 0
    assert restored == 0
    assert np.array_equal(output, printed)


def test_clean_slide_keeps_late_printed_answer_and_removes_later_cyan_ink() -> None:
    height, width = 180, 420
    base = np.full((height, width, 3), (34, 48, 66), dtype=np.uint8)
    cv2.putText(
        base,
        "QUESTION",
        (18, 48),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.85,
        (245, 245, 245),
        2,
        cv2.LINE_AA,
    )
    printed = base.copy()
    cv2.putText(
        printed,
        "ANSWER:",
        (45, 118),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.95,
        (245, 245, 245),
        2,
        cv2.LINE_AA,
    )
    cv2.putText(
        printed,
        "CD",
        (245, 118),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.95,
        (190, 225, 245),
        2,
        cv2.LINE_AA,
    )
    ink_points = np.asarray(
        [[55, 146], [115, 153], [180, 143], [245, 154], [315, 142]],
        dtype=np.int32,
    )
    frames: list[np.ndarray] = []
    for index in range(10):
        frame = printed.copy() if index >= 2 else base.copy()
        if index >= 5:
            point_count = min(len(ink_points), index - 3)
            if point_count >= 2:
                cv2.polylines(
                    frame,
                    [ink_points[:point_count]],
                    False,
                    (235, 180, 20),
                    4,
                    cv2.LINE_AA,
                )
            if index >= 7:
                cv2.circle(frame, (350, 92), 22, (235, 180, 20), 4, cv2.LINE_AA)
        frames.append(frame)

    printed_mask = (
        np.max(np.abs(printed.astype(np.int16) - base.astype(np.int16)), axis=2) >= 25
    )
    final_hsv = cv2.cvtColor(frames[-1], cv2.COLOR_BGR2HSV)
    cyan_ink = (
        (final_hsv[:, :, 0] >= 74)
        & (final_hsv[:, :, 0] <= 108)
        & (final_hsv[:, :, 1] >= 90)
        & (final_hsv[:, :, 2] >= 70)
    )
    temporal_mask = video_slides._annotation_mask(frames, excluded_boxes=[])
    assert int(np.count_nonzero((temporal_mask > 0) & printed_mask)) <= int(
        np.count_nonzero(printed_mask) * 0.02
    )
    assert int(np.count_nonzero((temporal_mask > 0) & cyan_ink)) > 500

    candidates = []
    for index, frame in enumerate(frames):
        if index < 2:
            edge_pixels = 500
        elif index < 5:
            edge_pixels = 1500
        else:
            edge_pixels = 4000 + min(4, index - 5) * 1000
        candidates.append(
            (_candidate_record(index, edge_pixels=edge_pixels), frame.copy())
        )
    clean = video_slides._clean_slide(
        candidates,
        watermark_boxes=[],
        fixed_boxes=[],
        presenter_boxes=[],
        watermark_hint="",
    )

    output_print = (
        np.max(np.abs(clean.image.astype(np.int16) - base.astype(np.int16)), axis=2)
        >= 25
    )
    retained_print = int(np.count_nonzero(output_print & printed_mask))
    assert retained_print >= int(np.count_nonzero(printed_mask) * 0.95)
    output_hsv = cv2.cvtColor(clean.image, cv2.COLOR_BGR2HSV)
    residual_cyan = (
        (output_hsv[:, :, 0] >= 74)
        & (output_hsv[:, :, 0] <= 108)
        & (output_hsv[:, :, 1] >= 90)
        & (output_hsv[:, :, 2] >= 70)
    )
    assert int(np.count_nonzero(residual_cyan)) <= int(
        np.count_nonzero(cyan_ink) * 0.35
    )


def test_full_page_watermark_cleanup_keeps_answer_from_last_dense_frames() -> None:
    height, width = 180, 420
    base = np.full((height, width, 3), (34, 48, 66), dtype=np.uint8)
    cv2.putText(
        base,
        "QUESTION",
        (18, 72),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.85,
        (225, 225, 225),
        2,
        cv2.LINE_AA,
    )
    printed = base.copy()
    cv2.putText(
        printed,
        "ANSWER:",
        (45, 140),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.9,
        (245, 245, 245),
        2,
        cv2.LINE_AA,
    )
    cv2.putText(
        printed,
        "CD",
        (240, 140),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.9,
        (190, 225, 245),
        2,
        cv2.LINE_AA,
    )
    candidates = []
    for index in range(32):
        frame = printed.copy() if index >= 28 else base.copy()
        watermark_x = 5 + (index * 47) % (width - 80)
        watermark_y = 24 + (index % 2) * 18
        cv2.putText(
            frame,
            "WM",
            (watermark_x, watermark_y),
            cv2.FONT_HERSHEY_SIMPLEX,
            0.55,
            (250, 250, 250),
            2,
            cv2.LINE_AA,
        )
        record = video_slides._scan_frame_from_image(
            index,
            index / 8.0,
            frame,
            ignored_boxes=[],
        )
        if index >= 28:
            record = video_slides._ScanFrame(
                frame_index=record.frame_index,
                timestamp=record.timestamp,
                gray=record.gray,
                packed_edges=record.packed_edges,
                edge_pixels=record.edge_pixels + 6000,
                phash=record.phash,
                sharpness=record.sharpness,
            )
        candidates.append((record, frame))

    clean = video_slides._clean_slide(
        candidates,
        watermark_boxes=[(0.0, 0.0, 1.0, 1.0)],
        fixed_boxes=[],
        presenter_boxes=[],
        watermark_hint="WM",
    )

    assert clean is not None
    printed_mask = (
        np.max(np.abs(printed.astype(np.int16) - base.astype(np.int16)), axis=2) >= 25
    )
    retained = (
        np.max(np.abs(clean.image.astype(np.int16) - base.astype(np.int16)), axis=2)
        >= 25
    )
    assert int(np.count_nonzero(retained & printed_mask)) >= int(
        np.count_nonzero(printed_mask) * 0.95
    )
    answer_error = float(
        np.mean(
            np.abs(
                clean.image[printed_mask].astype(np.int16)
                - printed[printed_mask].astype(np.int16)
            )
        )
    )
    assert answer_error < 8.0


def test_refined_candidate_indices_keep_frame_time_mapping() -> None:
    fps = 10.0
    image = np.full((90, 160, 3), 220, dtype=np.uint8)
    coarse_indices = (100, 115, 130)
    scans = [
        video_slides._scan_frame_from_image(
            frame_index,
            frame_index / fps,
            image,
            ignored_boxes=[],
        )
        for frame_index in coarse_indices
    ]
    segment = video_slides._SlideSegment(first=0, last=len(scans) - 1)

    refined_indices = sorted(
        video_slides._refined_candidate_frame_indices(segment, scans, fps)
    )
    refined_records = [
        video_slides._scan_frame_from_image(
            frame_index,
            frame_index / fps,
            image,
            ignored_boxes=[],
        )
        for frame_index in refined_indices
    ]

    assert set(coarse_indices).issubset(refined_indices)
    assert refined_indices[0] == coarse_indices[0]
    assert refined_indices[-1] == coarse_indices[-1]
    assert all(
        first.frame_index < second.frame_index and first.timestamp < second.timestamp
        for first, second in zip(refined_records, refined_records[1:])
    )
    assert all(
        record.timestamp == pytest.approx(record.frame_index / fps)
        for record in refined_records
    )


@pytest.mark.parametrize("candidate_count", [2, 3])
def test_temporal_restore_preserves_late_printed_title_with_short_history(
    candidate_count: int,
) -> None:
    height, width = 110, 280
    background = np.full((height, width, 3), (32, 46, 64), dtype=np.uint8)
    frames: list[np.ndarray] = []
    for index in range(candidate_count):
        frame = background.copy()
        if index >= 1:
            cv2.putText(
                frame,
                "PRINTED TITLE",
                (18, 78),
                cv2.FONT_HERSHEY_SIMPLEX,
                0.8,
                (245, 245, 245),
                2,
                cv2.LINE_AA,
            )
        cv2.putText(
            frame,
            "WM",
            (5 + index * 70, 22),
            cv2.FONT_HERSHEY_SIMPLEX,
            0.45,
            (250, 250, 250),
            1,
            cv2.LINE_AA,
        )
        frames.append(frame)

    expected = frames[-1]
    output = expected.copy()
    video_slides._temporal_restore_region(
        output,
        frames,
        (0, 0, width, height),
        variation_threshold=21,
        force=False,
        selected_index=candidate_count - 1,
    )

    title_region = (slice(50, 86), slice(12, 245))
    expected_title = expected[title_region]
    output_title = output[title_region]
    expected_bright = np.max(expected_title, axis=2) >= 220
    retained_bright = expected_bright & (np.max(output_title, axis=2) >= 220)
    assert int(np.count_nonzero(expected_bright)) > 300
    assert int(np.count_nonzero(retained_bright)) >= int(
        np.count_nonzero(expected_bright) * 0.95
    )


def test_transient_color_cleanup_preserves_late_persistent_cyan_print() -> None:
    height, width = 110, 300
    background = np.full((height, width, 3), (34, 48, 66), dtype=np.uint8)
    frames: list[np.ndarray] = []
    for index in range(5):
        frame = background.copy()
        if index >= 3:
            cv2.putText(
                frame,
                "CYAN PRINT",
                (16, 76),
                cv2.FONT_HERSHEY_SIMPLEX,
                0.85,
                (240, 190, 20),
                3,
                cv2.LINE_AA,
            )
        frames.append(frame)

    expected = frames[-1]
    output = expected.copy()
    video_slides._restore_transient_colored_marks(
        output,
        frames,
        selected_index=len(frames) - 1,
        excluded_boxes=[],
    )

    expected_print = (
        np.max(np.abs(expected.astype(np.int16) - background.astype(np.int16)), axis=2)
        >= 35
    )
    retained_print = expected_print & (
        np.max(np.abs(output.astype(np.int16) - background.astype(np.int16)), axis=2)
        >= 35
    )
    assert int(np.count_nonzero(expected_print)) > 500
    assert int(np.count_nonzero(retained_print)) >= int(
        np.count_nonzero(expected_print) * 0.95
    )


def test_multicolour_cursor_is_restored_from_nearest_clean_frame() -> None:
    height, width = 360, 640
    base = np.full((height, width, 3), (34, 48, 66), dtype=np.uint8)
    cv2.putText(
        base,
        "UNDER CURSOR",
        (210, 202),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.75,
        (245, 245, 245),
        2,
        cv2.LINE_AA,
    )
    frames = [base.copy() for _index in range(5)]
    center = (320, 195)
    cv2.circle(frames[-1], center, 11, (5, 5, 5), -1)
    cv2.circle(frames[-1], center, 9, (230, 175, 15), -1)
    cv2.circle(frames[-1], center, 7, (34, 180, 160), -1)
    output = frames[-1].copy()

    restored = video_slides._restore_transient_colored_marks(
        output,
        frames,
        selected_index=len(frames) - 1,
        excluded_boxes=[],
    )

    region = (slice(180, 211), slice(305, 336))
    assert restored > 300
    assert (
        float(np.mean(np.abs(output[region].astype(int) - base[region].astype(int))))
        < 1.0
    )


def test_static_yellow_icon_and_cyan_print_are_preserved() -> None:
    height, width = 360, 640
    frame = np.full((height, width, 3), (34, 48, 66), dtype=np.uint8)
    cv2.circle(frame, (180, 190), 7, (0, 255, 255), -1)
    cv2.putText(
        frame,
        "CYAN PRINT",
        (205, 198),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.65,
        (235, 180, 20),
        2,
        cv2.LINE_AA,
    )
    frames = [frame.copy() for _index in range(6)]
    output = frame.copy()

    restored = video_slides._restore_transient_colored_marks(
        output,
        frames,
        selected_index=len(frames) - 1,
        excluded_boxes=[],
    )

    assert restored == 0
    assert np.array_equal(output, frame)


def test_parked_olive_cursor_uses_quiet_local_background_fallback() -> None:
    height, width = 360, 640
    base = np.full((height, width, 3), (34, 48, 66), dtype=np.uint8)
    frames: list[np.ndarray] = []
    for _index in range(6):
        frame = base.copy()
        cv2.circle(frame, (320, 190), 9, (5, 5, 5), -1)
        cv2.circle(frame, (320, 190), 7, (34, 180, 160), -1)
        frames.append(frame)
    output = frames[-1].copy()

    restored = video_slides._restore_transient_colored_marks(
        output,
        frames,
        selected_index=len(frames) - 1,
        excluded_boxes=[],
    )

    region = (slice(172, 209), slice(302, 339))
    assert restored > 300
    assert (
        float(np.mean(np.abs(output[region].astype(int) - base[region].astype(int))))
        < 4.0
    )


def test_parked_cursor_fallback_reconnects_clear_table_line() -> None:
    height, width = 360, 640
    base = np.full((height, width, 3), (34, 48, 66), dtype=np.uint8)
    cv2.line(base, (120, 190), (520, 190), (235, 235, 235), 2)
    frames: list[np.ndarray] = []
    for _index in range(6):
        frame = base.copy()
        cv2.circle(frame, (320, 190), 9, (5, 5, 5), -1)
        cv2.circle(frame, (320, 190), 7, (34, 180, 160), -1)
        frames.append(frame)
    output = frames[-1].copy()

    video_slides._restore_transient_colored_marks(
        output,
        frames,
        selected_index=len(frames) - 1,
        excluded_boxes=[],
    )

    restored_line = np.min(output[189:192, 305:336], axis=2) >= 190
    assert int(np.count_nonzero(restored_line)) >= int(restored_line.size * 0.80)


def test_auto_presenter_policy_preserves_static_bottom_right_slide_content(
    tmp_path: Path,
) -> None:
    source = tmp_path / "no-presenter.avi"
    expected = _write_no_presenter_lecture(source)
    outputs = video_slides.extract_slides_to_pptx(
        source,
        tmp_path / "output",
        scan_mode="accurate",
        change_sensitivity="balanced",
        crop_mode="full",
        watermark_search="off",
        presenter_policy="auto_crop",
        keep_images=True,
        keep_report=False,
    )

    png = next(path for path in outputs if path.suffix == ".png")
    actual = video_slides._read_cv_image(png)
    assert actual is not None
    if actual.shape != expected.shape:
        actual = cv2.resize(
            actual,
            (expected.shape[1], expected.shape[0]),
            interpolation=cv2.INTER_AREA,
        )
    region = (slice(205, 250), slice(380, 478))
    expected_dark = np.min(expected[region], axis=2) < 90
    actual_dark = np.min(actual[region], axis=2) < 90
    assert int(np.count_nonzero(expected_dark)) > 100
    assert int(np.count_nonzero(actual_dark & expected_dark)) >= int(
        np.count_nonzero(expected_dark) * 0.85
    )


def _duplicate_test_page() -> np.ndarray:
    page = np.full((360, 640, 3), (34, 48, 66), dtype=np.uint8)
    cv2.putText(
        page,
        "PRINTED SLIDE",
        (38, 65),
        cv2.FONT_HERSHEY_SIMPLEX,
        1.05,
        (245, 245, 245),
        3,
        cv2.LINE_AA,
    )
    for index, text in enumerate(
        ("First printed line", "Second printed line", "Third printed line")
    ):
        cv2.putText(
            page,
            text,
            (55, 125 + index * 48),
            cv2.FONT_HERSHEY_SIMPLEX,
            0.72,
            (235, 235, 235),
            2,
            cv2.LINE_AA,
        )
    cv2.rectangle(page, (420, 115), (585, 285), (210, 210, 210), 2)
    return page


def test_near_duplicate_merges_clean_and_heavily_annotated_occurrences() -> None:
    clean = _duplicate_test_page()
    annotated = clean.copy()
    cyan = (255, 210, 0)
    strokes = (
        np.asarray([[50, 145], [130, 118], [220, 160], [310, 125]], dtype=np.int32),
        np.asarray([[80, 238], [180, 205], [290, 250], [410, 215]], dtype=np.int32),
        np.asarray([[340, 305], [430, 270], [540, 315]], dtype=np.int32),
    )
    for points in strokes:
        cv2.polylines(annotated, [points], False, cyan, 6, cv2.LINE_AA)
    cv2.circle(annotated, (500, 190), 58, cyan, 6, cv2.LINE_AA)

    annotation = video_slides._colored_handwriting_mask(annotated, excluded_boxes=[])

    assert int(np.count_nonzero(annotation)) >= int(annotation.size * 0.01)
    assert video_slides._near_duplicate(clean, annotated)
    assert video_slides._near_duplicate(annotated, clean)


def test_near_duplicate_tolerates_small_watermark_and_cleanup_noise() -> None:
    clean = _duplicate_test_page()
    noisy = clean.copy()
    overlay = noisy.copy()
    cv2.putText(
        overlay,
        "WM",
        (530, 32),
        cv2.FONT_HERSHEY_SIMPLEX,
        0.5,
        (245, 245, 245),
        1,
        cv2.LINE_AA,
    )
    cv2.addWeighted(overlay, 0.35, noisy, 0.65, 0, noisy)
    random = np.random.default_rng(7)
    rows = random.integers(300, 330, 120)
    columns = random.integers(40, 150, 120)
    cleanup_noise = random.integers(-10, 11, (120, 3))
    noisy[rows, columns] = np.clip(
        noisy[rows, columns].astype(np.int16) + cleanup_noise,
        0,
        255,
    ).astype(np.uint8)

    assert not np.array_equal(clean, noisy)
    assert video_slides._near_duplicate(clean, noisy)
    assert video_slides._near_duplicate(noisy, clean)


def test_near_duplicate_keeps_pages_with_different_legal_blue_text() -> None:
    first = np.full((360, 640, 3), 245, dtype=np.uint8)
    cv2.putText(
        first,
        "TOPIC",
        (40, 70),
        cv2.FONT_HERSHEY_SIMPLEX,
        1.2,
        (10, 10, 10),
        3,
        cv2.LINE_AA,
    )
    second = first.copy()
    cv2.putText(
        first,
        "ALPHA CONTENT",
        (60, 200),
        cv2.FONT_HERSHEY_SIMPLEX,
        1.1,
        (255, 180, 0),
        3,
        cv2.LINE_AA,
    )
    cv2.putText(
        second,
        "BETA CONTENT",
        (60, 200),
        cv2.FONT_HERSHEY_SIMPLEX,
        1.1,
        (255, 180, 0),
        3,
        cv2.LINE_AA,
    )

    assert not video_slides._near_duplicate(first, second)


def test_static_handwriting_mask_keeps_repeated_printed_cyan_icons() -> None:
    image = np.full((360, 640, 3), (34, 48, 66), dtype=np.uint8)
    for top in (70, 130, 190, 250):
        points = np.asarray([[55, top], [78, top + 12], [55, top + 24]], dtype=np.int32)
        cv2.fillPoly(image, [points], (235, 180, 20), cv2.LINE_AA)

    mask = video_slides._colored_handwriting_mask(image, excluded_boxes=[])

    assert int(np.count_nonzero(mask)) == 0


def test_end_to_end_outputs_final_printed_state_without_incremental_ink() -> None:
    with tempfile.TemporaryDirectory() as folder:
        root = Path(folder)
        source = root / "lecture.avi"
        output = root / "output"
        _write_synthetic_lecture(source)

        outputs = video_slides.extract_slides_to_pptx(
            source,
            output,
            scan_mode="accurate",
            change_sensitivity="balanced",
            crop_mode="full",
            watermark_search="top",
            presenter_policy="keep",
            keep_images=True,
            keep_report=True,
        )

        pptx_path = next(path for path in outputs if path.suffix == ".pptx")
        pngs = sorted(path for path in outputs if path.suffix == ".png")
        report_path = next(path for path in outputs if path.suffix == ".docx")
        assert len(Presentation(pptx_path).slides) == 2
        assert len(pngs) == 2
        report = Document(report_path)
        assert report.core_properties.keywords == video_slides._REPORT_SCHEMA
        detail_table = next(
            table
            for table in report.tables
            if table.rows[0].cells[0].text == "PPT 页码"
        )
        assert len(detail_table.rows) == 3
        assert [row.cells[0].text for row in detail_table.rows[1:]] == ["1", "2"]

        first = video_slides._read_cv_image(pngs[0])
        assert first is not None
        if first.shape[:2] != (360, 640):
            first = cv2.resize(first, (640, 360), interpolation=cv2.INTER_AREA)
        final_print_region = first[205:245, 45:300]
        assert int(np.count_nonzero(np.max(final_print_region, axis=2) > 210)) > 200
        ink_region = first[250:285, 55:410]
        cyan_like = (
            (ink_region[:, :, 0] > 180)
            & (ink_region[:, :, 1] > 140)
            & (ink_region[:, :, 2] < 80)
        )
        assert int(np.count_nonzero(cyan_like)) < 80


def test_pptx_second_check_proves_ordered_images_are_embedded(tmp_path: Path) -> None:
    first = np.full((180, 320, 3), (30, 60, 180), dtype=np.uint8)
    second = np.full((180, 320, 3), (180, 80, 30), dtype=np.uint8)
    cv2.putText(
        first, "FIRST", (35, 105), cv2.FONT_HERSHEY_SIMPLEX, 1.1, (255, 255, 255), 3
    )
    cv2.putText(
        second, "SECOND", (25, 105), cv2.FONT_HERSHEY_SIMPLEX, 1.1, (255, 255, 255), 3
    )
    first_path = tmp_path / "0001.png"
    second_path = tmp_path / "0002.png"
    pptx_path = tmp_path / "ordered.pptx"
    video_slides._write_png(first_path, first)
    video_slides._write_png(second_path, second)
    video_slides._write_pptx([first_path, second_path], pptx_path)

    video_slides._validate_pptx(pptx_path, 2, expected_images=[first_path, second_path])
    with pytest.raises(ValidationError, match="时间排序清单不一致"):
        video_slides._validate_pptx(
            pptx_path, 2, expected_images=[second_path, first_path]
        )


def test_short_preview_is_deferred_and_all_outputs_reorder_together(
    tmp_path: Path,
) -> None:
    previewed_page = np.full((120, 220, 3), (30, 70, 190), dtype=np.uint8)
    preceding_page = np.full((120, 220, 3), (180, 80, 35), dtype=np.uint8)
    first_path = tmp_path / "幻灯片_0001.png"
    second_path = tmp_path / "幻灯片_0002.png"
    video_slides._write_png(first_path, previewed_page)
    video_slides._write_png(second_path, preceding_page)
    preview = {"start_seconds": 10.0, "end_seconds": 11.5}
    formal = {"start_seconds": 40.0, "end_seconds": 65.0}
    preceding = {"start_seconds": 12.0, "end_seconds": 35.0}
    records = [
        {
            "slide": 1,
            "start_seconds": 10.0,
            "end_seconds": 11.5,
            "candidate_frames": 6,
            "occurrences": [preview, formal],
            "merged_duplicate_segments": 1,
            "selected_occurrence": formal,
            "selected_timestamp_seconds": 55.0,
        },
        {
            "slide": 2,
            "start_seconds": 12.0,
            "end_seconds": 35.0,
            "candidate_frames": 6,
            "occurrences": [preceding],
            "merged_duplicate_segments": 0,
            "selected_occurrence": preceding,
            "selected_timestamp_seconds": 20.0,
        },
    ]
    signatures = [
        video_slides._slide_signature(previewed_page),
        video_slides._slide_signature(preceding_page),
    ]

    paths, ordered_records, ordered_signatures, scores = (
        video_slides._reorder_slide_outputs_by_effective_timeline(
            [first_path, second_path],
            records,
            signatures,
            [1.0, 2.0],
        )
    )

    assert [record["ordering_start_seconds"] for record in ordered_records] == [
        12.0,
        40.0,
    ]
    assert [record["slide"] for record in ordered_records] == [1, 2]
    assert ordered_records[1]["ordering_reason"] == "short_preview_deferred"
    assert scores == [2.0, 1.0]
    assert ordered_signatures[0] is signatures[1]
    assert ordered_signatures[1] is signatures[0]
    assert [path.name for path in paths] == ["幻灯片_0001.png", "幻灯片_0002.png"]
    assert np.array_equal(video_slides._read_cv_image(paths[0]), preceding_page)
    assert np.array_equal(video_slides._read_cv_image(paths[1]), previewed_page)
    video_slides._validate_slide_timeline(ordered_records, 70.0)


def test_unique_short_page_keeps_its_real_first_occurrence() -> None:
    only_occurrence = {"start_seconds": 5.0, "end_seconds": 6.5}
    anchor, reason = video_slides._effective_ordering_occurrence(
        {"occurrences": [only_occurrence]}
    )

    assert anchor is only_occurrence
    assert reason == "first_occurrence"


def test_word_report_keeps_first_seen_order_and_all_duplicate_times(
    tmp_path: Path,
) -> None:
    first_occurrence = {"start_seconds": 0.0, "end_seconds": 2.0}
    repeated_occurrence = {"start_seconds": 8.0, "end_seconds": 10.0}
    middle_occurrence = {"start_seconds": 3.0, "end_seconds": 6.0}
    records = [
        {
            "slide": 1,
            "start_seconds": 0.0,
            "end_seconds": 2.0,
            "candidate_frames": 6,
            "annotation_pixels_restored": 120,
            "residual_annotation_pixels": 0,
            "dynamic_watermark_pixels_restored": 240,
            "low_confidence_pixels": 0,
            "occurrences": [first_occurrence, repeated_occurrence],
            "merged_duplicate_segments": 1,
            "selected_occurrence": repeated_occurrence,
            "selected_timestamp_seconds": 9.0,
        },
        {
            "slide": 2,
            "start_seconds": 3.0,
            "end_seconds": 6.0,
            "candidate_frames": 5,
            "annotation_pixels_restored": 0,
            "residual_annotation_pixels": 18,
            "dynamic_watermark_pixels_restored": 20,
            "low_confidence_pixels": 0,
            "occurrences": [middle_occurrence],
            "merged_duplicate_segments": 0,
            "selected_occurrence": middle_occurrence,
            "selected_timestamp_seconds": 5.0,
        },
    ]
    payload = {
        "schema": video_slides._REPORT_SCHEMA,
        "source": str(tmp_path / "lecture.mp4"),
        "mode": "final_printed_state_only",
        "metadata": {
            "fps": 25.0,
            "frame_count": 300,
            "width": 1920,
            "height": 1080,
            "duration": 12.0,
            "sample_interval": 0.5,
        },
        "scan_frames": 24,
        "detected_segments": 3,
        "output_slides": 2,
        "presentation_crop_pixels": [0, 0, 1920, 1080],
        "watermark_search": "auto",
        "watermark_text_hint": "",
        "presenter_policy": "keep",
        "auto_presenter_detected": False,
        "auto_presenter_cropped": False,
        "slides": records,
        "warnings": [],
    }
    report_path = tmp_path / "提取报告.docx"

    video_slides._write_report_docx(report_path, payload)
    video_slides._validate_report_docx(report_path, 2, payload)

    report = Document(report_path)
    detail_table = next(
        table for table in report.tables if table.rows[0].cells[0].text == "PPT 页码"
    )
    assert detail_table.rows[1].cells[1].text == "00:00:00.000–00:00:02.000"
    assert detail_table.rows[1].cells[2].text == "00:00:09.000"
    assert detail_table.rows[1].cells[3].text == "2"
    assert detail_table.rows[1].cells[4].text.splitlines() == [
        "00:00:00.000–00:00:02.000",
        "00:00:08.000–00:00:10.000",
    ]


def test_duplicate_signature_keeps_only_compact_page_features() -> None:
    first = np.full((360, 640, 3), 238, dtype=np.uint8)
    cv2.putText(
        first,
        "COMPACT FEATURES",
        (35, 180),
        cv2.FONT_HERSHEY_SIMPLEX,
        1.1,
        (20, 20, 20),
        3,
        cv2.LINE_AA,
    )
    second = first.copy()
    first_signature = video_slides._slide_signature(first)
    second_signature = video_slides._slide_signature(second)

    retained_bytes = (
        first_signature.phash.nbytes
        + first_signature.gray.nbytes
        + first_signature.edges.nbytes
        + first_signature.dilated_edges.nbytes
        + first_signature.colored_ink.nbytes
        + first_signature.annotation_ink.nbytes
    )
    assert retained_bytes < first.nbytes // 10
    assert video_slides._signatures_near_duplicate(
        first_signature, second_signature
    ) == video_slides._near_duplicate(first, second)


def test_task_staging_does_not_publish_half_results(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    source = tmp_path / "lecture.avi"
    target = tmp_path / "result"
    _write_synthetic_lecture(source)

    def fail_pptx(_images: object, _target: object) -> None:
        raise ValidationError("forced PPT failure")

    monkeypatch.setattr(video_slides, "_write_pptx", fail_pptx)
    with pytest.raises(ValidationError, match="forced PPT failure"):
        video_slides.extract_slides_to_pptx(
            source,
            target,
            scan_mode="accurate",
            change_sensitivity="balanced",
            crop_mode="full",
            watermark_search="top",
            presenter_policy="keep",
            keep_images=True,
            keep_report=True,
        )

    assert not target.exists()
    assert not list(tmp_path.glob(".dfvs-*"))


def test_publish_staged_directory_retries_transient_windows_lock(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    staging = tmp_path / ".dfvs-stage"
    target = tmp_path / "published"
    staging.mkdir()
    (staging / "slide.png").write_bytes(b"validated")
    real_replace = video_slides.os.replace
    attempts = 0

    def transient_replace(source: object, destination: object) -> None:
        nonlocal attempts
        attempts += 1
        if attempts < 3:
            error = PermissionError("temporarily locked")
            error.winerror = 5  # type: ignore[attr-defined]
            raise error
        real_replace(source, destination)

    monkeypatch.setattr(video_slides.os, "name", "nt")
    monkeypatch.setattr(video_slides.os, "replace", transient_replace)
    monkeypatch.setattr(video_slides.time, "sleep", lambda _seconds: None)

    video_slides._publish_staged_directory(staging, target)

    assert attempts == 3
    assert not staging.exists()
    assert (target / "slide.png").read_bytes() == b"validated"
