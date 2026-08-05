from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from PIL import Image
from docx import Document
from openpyxl import Workbook, load_workbook
from pypdf import PdfReader, PdfWriter
from pypdf.generic import (
    ArrayObject,
    DictionaryObject,
    NameObject,
    NumberObject,
    TextStringObject,
)
from reportlab.pdfgen.canvas import Canvas

from docuforge import engines
from docuforge.registry import CORE_OPERATION_IDS, get_operations
from docuforge.runner import TaskRunner


class RegistryIntegrationTests(unittest.TestCase):
    def setUp(self) -> None:
        self.operations = {item.id: item for item in get_operations()}

    def test_catalog_is_unique_and_curated(self) -> None:
        operations = get_operations()
        self.assertEqual(len(operations), len(CORE_OPERATION_IDS))
        self.assertEqual(len(operations), 73)
        self.assertEqual(len(operations), len({item.id for item in operations}))
        self.assertTrue(
            all(item.name and item.group and item.description for item in operations)
        )
        removed = {
            "pdf.to_excel",
            "pdf.to_html",
            "pdf.ocr",
            "pdf.digital_signature",
            "word.to_epub",
            "excel.formulas_to_values",
            "ppt.to_video",
            "ppt.master",
            "image.remove_background",
            "image.raw",
        }
        self.assertTrue(removed.isdisjoint(self.operations))

    def test_lecture_video_slide_extraction_is_local_final_state_reconstruction(
        self,
    ) -> None:
        operation = self.operations["video.extract_slides_ppt"]
        self.assertEqual(operation.group, "视频生成")
        self.assertEqual(operation.fidelity, "visual")
        self.assertIn("不可编辑 PPT", operation.name)
        self.assertIn("最完整的最终状态", operation.description)
        self.assertIn("手写批注", operation.notes)
        self.assertIn("获授权", operation.notes)
        self.assertEqual(
            [parameter.key for parameter in operation.parameters],
            [
                "scan_mode",
                "change_sensitivity",
                "crop_mode",
                "crop_rect",
                "watermark_search",
                "watermark_rect",
                "watermark_text_hint",
                "annotation_color_mode",
                "annotation_colors",
                "annotation_color_tolerance",
                "fixed_watermark_regions",
                "fixed_watermark_fill",
                "fixed_watermark_fill_color",
                "presenter_policy",
                "presenter_rect",
                "enhancement_mode",
                "keep_images",
                "keep_report",
            ],
        )
        defaults = operation.normalize_parameters({})
        self.assertEqual(defaults["scan_mode"], "accurate")
        self.assertEqual(defaults["watermark_search"], "auto")
        self.assertEqual(defaults["annotation_color_mode"], "auto")
        self.assertEqual(defaults["annotation_colors"], "#00AEEF")
        self.assertEqual(defaults["annotation_color_tolerance"], 24)
        self.assertEqual(defaults["fixed_watermark_fill"], "temporal")
        self.assertEqual(defaults["fixed_watermark_fill_color"], "#FFFFFF")
        self.assertEqual(defaults["presenter_policy"], "auto_crop")
        self.assertTrue(defaults["keep_images"])
        self.assertTrue(defaults["keep_report"])

        annotation_colors = next(
            item for item in operation.parameters if item.key == "annotation_colors"
        )
        self.assertEqual(annotation_colors.kind, "colors")
        self.assertIn("悬停放大观察", annotation_colors.help_text)
        self.assertEqual(
            annotation_colors.visible_when,
            ("annotation_color_mode", ("manual",)),
        )
        self.assertEqual(annotation_colors.section, "② 手写批注清理")
        region_parameters = {
            item.key: item for item in operation.parameters if item.kind == "region"
        }
        self.assertEqual(
            set(region_parameters),
            {
                "crop_rect",
                "watermark_rect",
                "fixed_watermark_regions",
                "presenter_rect",
            },
        )
        self.assertTrue(region_parameters["fixed_watermark_regions"].advanced)

    def test_video_ppt_manual_repair_is_a_separate_guarded_operation(self) -> None:
        operation = self.operations["video.repair_slides_ppt"]
        self.assertEqual(operation.group, "视频生成")
        self.assertEqual(operation.min_inputs, 2)
        self.assertEqual(operation.max_inputs, 2)
        self.assertIn(".pptx", operation.extensions)
        self.assertIn(".mp4", operation.extensions)
        self.assertIn("未框选区域", operation.description)
        self.assertIn("逐像素", operation.notes)
        self.assertEqual([item.key for item in operation.parameters], ["repair_plan"])
        self.assertTrue(operation.parameters[0].required)

    def test_lecture_video_slide_extraction_forwards_color_guidance(self) -> None:
        operation = self.operations["video.extract_slides_ppt"]
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "讲义.mp4"
            source.write_bytes(b"video")
            output = root / "输出"
            output.mkdir()
            parameters = operation.normalize_parameters(
                {
                    "annotation_color_mode": "manual",
                    "annotation_colors": "#00AEEF;255,60,60",
                    "annotation_color_tolerance": 31,
                    "fixed_watermark_regions": "0,0,16,10",
                    "fixed_watermark_fill": "color",
                    "fixed_watermark_fill_color": "#F8F7F2",
                }
            )
            mocked_output = root / "输出.pptx"
            with patch(
                "docuforge.processors.video_slides.extract_slides_to_pptx",
                return_value=[mocked_output],
            ) as extract:
                results = operation.handler([source], output, parameters)

        self.assertEqual(results, [mocked_output])
        self.assertEqual(extract.call_args.kwargs["annotation_color_mode"], "manual")
        self.assertEqual(
            extract.call_args.kwargs["annotation_colors"], "#00AEEF;255,60,60"
        )
        self.assertEqual(extract.call_args.kwargs["annotation_color_tolerance"], 31)
        self.assertEqual(extract.call_args.kwargs["fixed_watermark_fill"], "color")
        self.assertEqual(
            extract.call_args.kwargs["fixed_watermark_fill_color"], "#F8F7F2"
        )

    def test_office_auto_engine_ui_documents_wps_first_and_strict_explicit_mode(
        self,
    ) -> None:
        word_engine = self.operations["word.to_pdf"].parameters[0]
        choices = dict(word_engine.choices)
        self.assertIn("WPS → Microsoft Office → LibreOffice", choices["auto"])
        self.assertIn("仅显式选择", choices["microsoft_office"])
        self.assertIn("不会偷换", self.operations["word.to_pdf"].notes)

        excel_engine = self.operations["excel.to_pdf"].parameters[0]
        self.assertIn("WPS、Microsoft Office 与 LibreOffice", excel_engine.help_text)
        self.assertIn("非空有效文件", excel_engine.help_text)

        ppt_images = self.operations["ppt.to_images"]
        self.assertIn("优先由 WPS", ppt_images.description)
        self.assertIn("显式选择 Microsoft Office", ppt_images.notes)

    def test_word_full_compatibility_is_a_legacy_advanced_upgrade(self) -> None:
        operation = self.operations["word.full_compatibility"]
        self.assertEqual(operation.group, "兼容修复 / 高级工具")
        self.assertEqual(operation.name, "旧版固定坐标 Word 兼容升级")
        self.assertEqual(operation.extensions, (".docx",))
        self.assertEqual(operation.parameters[0].key, "verification_engine")
        self.assertEqual(operation.parameters[0].default, "auto")
        self.assertIn("WPS", operation.parameters[0].help_text)
        self.assertIn("不会把页面转成图片", operation.notes)
        self.assertIn("新文件通常无需再次处理", operation.notes)

        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "转换结果.docx"
            Document().save(source)
            output = root / "输出"
            output.mkdir()
            params = operation.normalize_parameters({"verification_engine": "none"})
            expected = output / "转换结果_旧版固定坐标Word_兼容升级.docx"
            with patch(
                "docuforge.processors.word_compat.optimize_word_full_compatibility",
                return_value=[expected],
            ) as optimize:
                results = operation.handler([source], output, params)

            self.assertEqual(results, [expected])
            self.assertEqual(optimize.call_args.args[:2], (source, expected))
            self.assertEqual(
                optimize.call_args.kwargs["verification_engine"],
                "none",
            )

    def test_image_pdf_and_render_pipeline(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            image_path = root / "中文图片.png"
            Image.new("RGBA", (120, 80), (20, 100, 220, 180)).save(image_path)
            output = root / "输出"

            resized = TaskRunner().run(
                self.operations["image.resize"],
                [image_path],
                output,
                {"width": "60", "height": "40", "keep_aspect": "false"},
            )
            with Image.open(resized.outputs[0]) as resized_image:
                self.assertEqual(resized_image.size, (60, 40))

            combined = TaskRunner().run(
                self.operations["image.to_pdf"],
                [image_path],
                output,
                {"filename": "图片文档", "dpi": "96", "background": "white"},
            )
            self.assertTrue(combined.outputs[0].is_file())

            rendered = TaskRunner().run(
                self.operations["pdf.to_images"],
                combined.outputs,
                output,
                {"format": "png", "dpi": "150"},
            )
            self.assertEqual(len(rendered.outputs), 1)
            self.assertTrue(rendered.outputs[0].is_file())

            video = TaskRunner().run(
                self.operations["image.to_video"],
                [image_path, resized.outputs[0]],
                output,
                {
                    "filename": "图片演示",
                    "slide_duration": "0.5",
                    "fps": "5",
                    "resolution": "720p",
                    "transition": "none",
                    "transition_duration": "0.1",
                    "background": "black",
                    "quality": "28",
                },
            )
            self.assertGreater(video.outputs[0].stat().st_size, 1000)

            transcoded = TaskRunner().run(
                self.operations["video.transcode"],
                video.outputs,
                output,
                {
                    "format": "mkv",
                    "video_codec": "auto",
                    "audio_codec": "none",
                    "quality": "28",
                    "resolution": "original",
                    "target_fps": "",
                    "audio_bitrate": "128",
                },
            )
            self.assertEqual(transcoded.outputs[0].suffix, ".mkv")
            self.assertGreater(transcoded.outputs[0].stat().st_size, 1000)

            trimmed = TaskRunner().run(
                self.operations["video.trim"],
                video.outputs,
                output,
                {
                    "format": "mp4",
                    "start": "0",
                    "end": "0.4",
                    "duration": "",
                    "quality": "28",
                },
            )
            self.assertGreater(trimmed.outputs[0].stat().st_size, 1000)

    def test_pdf_registry_mapping(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "样例.pdf"
            canvas = Canvas(str(source), pagesize=(300, 200))
            canvas.drawString(40, 150, "page one")
            canvas.showPage()
            canvas.drawString(40, 150, "page two")
            canvas.save()
            output = root / "输出"

            rotated = TaskRunner().run(
                self.operations["pdf.rotate"],
                [source],
                output,
                {"pages": "全部", "angle": "90"},
            )
            self.assertEqual(len(rotated.outputs), 1)
            extracted = TaskRunner().run(
                self.operations["pdf.extract_pages"],
                [source],
                output,
                {"pages": "2"},
            )
            self.assertEqual(len(extracted.outputs), 1)

    def test_pdf_extract_images_registry_mapping(self) -> None:
        operation = self.operations["pdf.extract_images"]
        self.assertEqual(operation.group, "PDF 格式转换")
        self.assertEqual(operation.fidelity, "extract")
        self.assertEqual(operation.extensions, (".pdf",))
        self.assertFalse(operation.reject_encrypted_pdf_inputs)
        self.assertFalse(operation.reject_signed_pdf_inputs)
        self.assertEqual(
            [parameter.key for parameter in operation.parameters],
            [
                "mode",
                "pages",
                "format",
                "dpi",
                "jpeg_quality",
                "min_width",
                "min_height",
                "merge_gap",
                "region_padding",
                "deduplicate",
                "include_annotations",
                "write_manifest",
                "password",
            ],
        )
        self.assertEqual(
            dict(operation.parameters[0].choices),
            {
                "original": "原始资源无重采样（最快，推荐）",
                "visible": "逐个可见位置高清还原",
                "smart": "相邻碎片智能合并",
                "both": "原始资源 + 智能合并",
                "all": "原始 + 可见 + 智能（输出最多）",
            },
        )
        self.assertEqual(
            dict(operation.parameters[2].choices),
            {"auto": "自动（推荐）", "png": "PNG", "jpg": "JPG"},
        )
        defaults = operation.normalize_parameters({})
        self.assertEqual(
            defaults,
            {
                "mode": "original",
                "pages": "全部",
                "format": "auto",
                "dpi": 300,
                "jpeg_quality": 95,
                "min_width": 1,
                "min_height": 1,
                "merge_gap": 4.0,
                "region_padding": 2.0,
                "deduplicate": True,
                "include_annotations": False,
                "write_manifest": True,
                "password": "",
            },
        )
        parameter_by_key = {item.key: item for item in operation.parameters}
        self.assertEqual(
            (parameter_by_key["dpi"].minimum, parameter_by_key["dpi"].maximum),
            (72, 1200),
        )
        self.assertEqual(
            (
                parameter_by_key["jpeg_quality"].minimum,
                parameter_by_key["jpeg_quality"].maximum,
            ),
            (30, 100),
        )
        self.assertEqual(
            (
                parameter_by_key["merge_gap"].minimum,
                parameter_by_key["merge_gap"].maximum,
            ),
            (0, 72),
        )
        self.assertIn(
            "仅在输出或转换为 JPG 时生效", parameter_by_key["jpeg_quality"].help_text
        )
        self.assertIn("渲染宽度", parameter_by_key["min_width"].help_text)
        self.assertIn("1 PDF 点 = 1/72 英寸", parameter_by_key["merge_gap"].help_text)
        self.assertIn(
            "整个 PDF 和所选模式间去重", parameter_by_key["deduplicate"].help_text
        )
        self.assertIn("密码不会写入", parameter_by_key["password"].help_text)
        self.assertIn("<文件名>_PDF图片", operation.notes)
        self.assertIn("不覆盖历史结果", operation.notes)
        self.assertEqual(operation.capability().engine, "PyMuPDF + Pillow")

        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "论文.pdf"
            source.touch()
            output = root / "输出"
            output.mkdir()
            expected = output / "论文_PDF图片" / "图片_001.jpg"
            params = operation.normalize_parameters(
                {
                    "mode": "all",
                    "pages": "2-4",
                    "format": "jpg",
                    "dpi": "420",
                    "jpeg_quality": "88",
                    "min_width": "32",
                    "min_height": "24",
                    "merge_gap": "6.5",
                    "region_padding": "3.5",
                    "deduplicate": "false",
                    "include_annotations": "true",
                    "write_manifest": "false",
                    "password": "secret",
                }
            )
            with patch(
                "docuforge.processors.pdf_images.extract_pdf_images",
                return_value=[expected],
            ) as extract_images:
                results = operation.handler([source], output, params)

        self.assertEqual(results, [expected])
        extract_images.assert_called_once_with(
            source,
            output / "论文_PDF图片",
            mode="all",
            pages="2-4",
            image_format="jpg",
            dpi=420,
            jpeg_quality=88,
            min_width=32,
            min_height=24,
            merge_gap=6.5,
            region_padding=3.5,
            deduplicate=False,
            include_annotations=True,
            write_manifest=False,
            password="secret",
            overwrite=False,
        )

    def test_pdf_watermark_count_registry_defaults_bounds_and_forwarding(self) -> None:
        operation = self.operations["pdf.watermark"]
        count_spec = next(item for item in operation.parameters if item.key == "count")
        self.assertEqual(count_spec.default, 1)
        self.assertEqual((count_spec.minimum, count_spec.maximum), (1, 100))
        self.assertEqual(operation.normalize_parameters({})["count"], 1)
        for invalid in ("0", "101"):
            with self.subTest(count=invalid), self.assertRaisesRegex(
                Exception, "每页水印数量"
            ):
                operation.normalize_parameters({"count": invalid})

        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "样例.pdf"
            source.touch()
            output = root / "输出"
            output.mkdir()
            params = operation.normalize_parameters({"count": "9"})
            expected = output / "样例_水印.pdf"
            with patch(
                "docuforge.processors.pdf.add_watermark", return_value=[expected]
            ) as add_watermark:
                results = operation.handler([source], output, params)
            self.assertEqual(results, [expected])
            self.assertEqual(add_watermark.call_args.kwargs["count"], 9)

    def test_pdf_lossy_compression_defaults_to_structure_preserving_strategy(
        self,
    ) -> None:
        operation = self.operations["pdf.compress_lossy"]
        self.assertEqual(
            [parameter.key for parameter in operation.parameters],
            ["strategy", "dpi", "jpeg_quality", "color_mode", "password"],
        )
        strategy = operation.parameters[0]
        self.assertEqual(strategy.default, "smart")
        self.assertEqual(
            dict(strategy.choices),
            {
                "smart": "结构保留智能压缩（推荐）",
                "raster": "整页栅格兼容压缩（扫描件）",
            },
        )
        self.assertIn("文字仍可搜索复制", strategy.help_text)
        self.assertIn("不整页栅格化", operation.notes)

        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "样例.pdf"
            source.touch()
            output = root / "输出"
            output.mkdir()
            params = operation.normalize_parameters({})
            expected = output / "样例_高精度有损压缩.pdf"
            with patch(
                "docuforge.processors.pdf.compress_pdf_lossy",
                return_value=[expected],
            ) as compress:
                results = operation.handler([source], output, params)

            self.assertEqual(results, [expected])
            self.assertEqual(compress.call_args.kwargs["strategy"], "smart")
            self.assertEqual(compress.call_args.kwargs["dpi"], 220)
            self.assertEqual(compress.call_args.kwargs["jpeg_quality"], 88)

    def test_pdf_to_word_exposes_and_routes_modes_and_quality_policies(self) -> None:
        operation = self.operations["pdf.to_word"]
        self.assertEqual(
            [parameter.key for parameter in operation.parameters],
            [
                "mode",
                "column_layout",
                "hybrid_force_visual_pages",
                "low_quality_policy",
                "dpi",
                "password",
            ],
        )
        mode = operation.parameters[0]
        self.assertEqual(mode.default, "hybrid")
        self.assertEqual(
            dict(mode.choices),
            {
                "hybrid": "版式优先混合（推荐）",
                "editable": "全文可编辑重建",
                "visual": "整篇高清原样（不可编辑）",
            },
        )
        column_layout = operation.parameters[1]
        self.assertEqual(column_layout.default, "auto")
        self.assertEqual(
            dict(column_layout.choices),
            {
                "auto": "自动识别（推荐）",
                "single": "全文单栏",
                "double": "全文双栏",
                "mixed": "混合分栏（单双栏并存）",
            },
        )
        self.assertIn("阅读顺序", column_layout.help_text)
        self.assertIn("标题、摘要为单栏而正文为双栏", column_layout.help_text)
        force_visual_pages = operation.parameters[2]
        self.assertEqual(force_visual_pages.default, "")
        self.assertIn("1,3-5", force_visual_pages.help_text)
        low_quality_policy = operation.parameters[3]
        self.assertEqual(low_quality_policy.default, "discard")
        self.assertEqual(
            dict(low_quality_policy.choices),
            {
                "discard": "不保留（推荐）",
                "keep": "仍保留并警告",
            },
        )
        self.assertIn("版式优先混合模式生效", low_quality_policy.help_text)
        self.assertIn("WPS 实际分页", low_quality_policy.help_text)
        self.assertIn("仍保留时", low_quality_policy.help_text)
        self.assertIn("可靠页面保持可编辑", operation.notes)

        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "合同.pdf"
            source.touch()
            output = root / "输出"
            output.mkdir()
            for (
                selected_mode,
                selected_layout,
                selected_policy,
                forced_pages,
                filename_part,
            ) in (
                ("hybrid", "double", "discard", "1,3-5", "混合保真Word"),
                ("editable", "single", "keep", "", "可编辑Word"),
                ("visual", "mixed", "discard", "", "高清原样Word"),
            ):
                with self.subTest(
                    mode=selected_mode,
                    column_layout=selected_layout,
                    policy=selected_policy,
                ):
                    params = operation.normalize_parameters(
                        {
                            "mode": selected_mode,
                            "column_layout": selected_layout,
                            "hybrid_force_visual_pages": forced_pages,
                            "low_quality_policy": selected_policy,
                            "dpi": "360",
                            "password": "secret",
                        }
                    )
                    with patch(
                        "docuforge.processors.conversion.pdf_to_docx",
                        side_effect=lambda _source, target, **_kwargs: [target],
                    ) as convert:
                        outputs = operation.handler([source], output, params)
                    self.assertEqual(len(outputs), 1)
                    self.assertIn(filename_part, outputs[0].name)
                    self.assertEqual(convert.call_args.kwargs["mode"], selected_mode)
                    self.assertEqual(
                        convert.call_args.kwargs["column_layout"],
                        selected_layout,
                    )
                    self.assertEqual(
                        convert.call_args.kwargs["low_quality_policy"],
                        selected_policy,
                    )
                    self.assertEqual(
                        convert.call_args.kwargs["hybrid_force_visual_pages"],
                        forced_pages,
                    )
                    self.assertEqual(convert.call_args.kwargs["dpi"], 360)
                    self.assertEqual(convert.call_args.kwargs["password"], "secret")

            legacy_params = operation.normalize_parameters(
                {
                    "mode": "hybrid",
                    "hybrid_force_visual_pages": "",
                    "low_quality_policy": "discard",
                    "dpi": "300",
                    "password": "",
                }
            )
            self.assertEqual(legacy_params["column_layout"], "auto")

    def test_pdf_insert_pages_uses_a_new_name_on_repeat(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            main_pdf = root / "正文.pdf"
            inserted_pdf = root / "附录.pdf"
            for path in (main_pdf, inserted_pdf):
                writer = PdfWriter()
                writer.add_blank_page(width=300, height=200)
                with path.open("wb") as stream:
                    writer.write(stream)
                writer.close()

            operation = self.operations["pdf.insert_pages"]
            output = root / "输出"
            first = TaskRunner().run(
                operation,
                [main_pdf, inserted_pdf],
                output,
                {"position": "1", "pages": ""},
            )
            second = TaskRunner().run(
                operation,
                [main_pdf, inserted_pdf],
                output,
                {"position": "1", "pages": ""},
            )

            self.assertNotEqual(first.outputs[0], second.outputs[0])
            self.assertTrue(first.outputs[0].is_file())
            self.assertTrue(second.outputs[0].is_file())
            self.assertEqual(second.outputs[0].stem, "正文_已插页_1")

    def test_video_registry_uses_precise_capabilities_and_safe_names(self) -> None:
        self.assertIs(
            self.operations["image.to_video"].capability_probe,
            engines.slideshow_video_capability,
        )
        for operation_id in ("video.transcode", "video.compress", "video.trim"):
            self.assertIs(
                self.operations[operation_id].capability_probe,
                engines.video_transform_capability,
            )
        audio_operation = self.operations["video.extract_audio"]
        self.assertIs(
            audio_operation.capability_probe,
            engines.audio_extraction_capability,
        )
        self.assertEqual(
            [item.key for item in audio_operation.parameters],
            ["sample_rate", "channels"],
        )
        self.assertEqual(audio_operation.name, "提取无损 WAV 音频")
        self.assertIn(".ppt", self.operations["ppt.to_images"].extensions)

        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "frame.png"
            Image.new("RGB", (16, 16), "white").save(source)
            output = root / "输出"
            output.mkdir()
            operation = self.operations["image.to_video"]
            params = operation.normalize_parameters({"filename": "demo.mp4"})
            expected = output / "demo.mp4"
            with patch(
                "docuforge.processors.video.images_to_video", return_value=[expected]
            ) as convert:
                operation.handler([source], output, params)
            self.assertEqual(convert.call_args.args[1], expected)

    def test_pdf_mutations_reject_encrypted_inputs(self) -> None:
        protected_operations = {
            "pdf.merge",
            "pdf.split",
            "pdf.extract_pages",
            "pdf.delete_pages",
            "pdf.insert_pages",
            "pdf.rotate",
            "pdf.compress",
            "pdf.watermark",
            "pdf.header_footer",
        }
        for operation_id in protected_operations:
            with self.subTest(operation=operation_id):
                self.assertTrue(
                    self.operations[operation_id].reject_encrypted_pdf_inputs
                )

        self.assertFalse(self.operations["pdf.decrypt"].reject_encrypted_pdf_inputs)

    def test_blank_signature_field_is_not_treated_as_applied_signature(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "unsigned-signature-field.pdf"
            writer = PdfWriter()
            page = writer.add_blank_page(width=300, height=200)
            field = DictionaryObject(
                {
                    NameObject("/FT"): NameObject("/Sig"),
                    NameObject("/T"): TextStringObject("Signature1"),
                    NameObject("/Type"): NameObject("/Annot"),
                    NameObject("/Subtype"): NameObject("/Widget"),
                    NameObject("/Rect"): ArrayObject(
                        [
                            NumberObject(36),
                            NumberObject(36),
                            NumberObject(220),
                            NumberObject(100),
                        ]
                    ),
                    NameObject("/F"): NumberObject(4),
                }
            )
            field_reference = writer._add_object(field)
            page[NameObject("/Annots")] = ArrayObject([field_reference])
            writer._root_object[NameObject("/AcroForm")] = writer._add_object(
                DictionaryObject(
                    {
                        NameObject("/Fields"): ArrayObject([field_reference]),
                        NameObject("/SigFlags"): NumberObject(1),
                    }
                )
            )
            with source.open("wb") as stream:
                writer.write(stream)
            writer.close()

            result = TaskRunner().run(
                self.operations["pdf.rotate"],
                [source],
                root / "output",
                {"pages": "1", "angle": "90"},
            )
            self.assertEqual(len(result.outputs), 1)
            reader = PdfReader(result.outputs[0])
            try:
                fields = reader.get_fields() or {}
                self.assertIn("Signature1", fields)
                self.assertNotIn("/V", fields["Signature1"])
            finally:
                reader.close()

    def test_word_and_excel_registry_mapping(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            output = root / "输出"
            word_path = root / "合同.docx"
            document = Document()
            document.add_paragraph("甲方：旧公司")
            document.save(word_path)
            word_result = TaskRunner().run(
                self.operations["word.replace"],
                [word_path],
                output,
                {"replacements": '{"旧公司":"新公司"}', "case_sensitive": "true"},
            )
            changed = Document(word_result.outputs[0])
            self.assertIn("新公司", "\n".join(item.text for item in changed.paragraphs))

            excel_path = root / "数据.xlsx"
            workbook = Workbook()
            sheet = workbook.active
            sheet.append(["名称", "分数"])
            sheet.append(["乙", 2])
            sheet.append(["甲", 1])
            workbook.save(excel_path)
            excel_result = TaskRunner().run(
                self.operations["excel.sort"],
                [excel_path],
                output,
                {"column": "B", "sheet": "", "header": "true", "reverse": "false"},
            )
            sorted_book = load_workbook(excel_result.outputs[0], data_only=False)
            self.assertEqual(sorted_book.active["A2"].value, "甲")
            sorted_book.close()


if __name__ == "__main__":
    unittest.main()
