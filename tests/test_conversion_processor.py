from __future__ import annotations

import math
import sys
import tempfile
import unittest
import warnings
from xml.etree import ElementTree
from pathlib import Path
from types import ModuleType, SimpleNamespace
from unittest.mock import Mock, patch
from zipfile import ZipFile

from PIL import Image, ImageDraw
from docx import Document
from pypdf import PdfReader, PdfWriter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen.canvas import Canvas

from docuforge.models import MissingEngineError, ValidationError
from docuforge.processors import conversion as conversion_processor
from docuforge.processors.conversion import (
    _adjacent_english_word_coverage,
    _english_word_multiset_recall,
    _extract_docx_text,
    _join_pdf2docx_line_break_hyphen,
    _normalize_english_line_break_hyphens,
    _restore_pdf2docx_span_spaces,
    _restore_pdf2docx_text_block_spaces,
    add_pdf_markup,
    add_pdf_note,
    add_visual_signature,
    fill_pdf_form,
    pdf_to_docx,
    pdf_to_html,
    remove_background,
    svg_to_images,
)


class ConversionProcessorTests(unittest.TestCase):
    def _make_pdf(self, path: Path) -> Path:
        canvas = Canvas(str(path), pagesize=(320, 220))
        canvas.drawString(30, 170, "Editable contract text")
        canvas.save()
        return path

    def _make_three_page_text_pdf(self, path: Path) -> Path:
        canvas = Canvas(str(path), pagesize=(320, 220))
        for text in (
            "Editable page one",
            "Forced visual page two",
            "Editable page three",
        ):
            canvas.drawString(30, 170, text)
            canvas.showPage()
        canvas.save()
        return path

    @staticmethod
    def _make_dense_fixed_layout_docx(path: Path, frames: int = 60) -> Path:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt

        document = Document()
        for index in range(frames):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(
                f"Dense editable line {index + 1} keeps normal copy order."
            )
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            properties = paragraph._p.get_or_add_pPr()
            frame = OxmlElement("w:framePr")
            for name, value in (
                ("w:wrap", "none"),
                ("w:hAnchor", "page"),
                ("w:vAnchor", "page"),
                ("w:x", "900"),
                ("w:y", str(500 + index * 205)),
                ("w:w", "7200"),
                ("w:h", "210"),
                ("w:hRule", "atLeast"),
            ):
                frame.set(qn(name), value)
            properties.insert(0, frame)
            fit = OxmlElement("w:fitText")
            fit.set(qn("w:val"), "7100")
            fit.set(qn("w:id"), str(index + 1))
            run._r.get_or_add_rPr().append(fit)
        document.save(path)
        return path

    @staticmethod
    def _set_table_grid_widths(table, widths: list[int]) -> None:
        from docx.oxml.ns import qn

        grid_columns = table._tbl.tblGrid.findall(qn("w:gridCol"))
        if len(grid_columns) != len(widths):
            raise AssertionError("test table grid width count mismatch")
        for grid_column, width in zip(grid_columns, widths):
            grid_column.set(qn("w:w"), str(width))

    @staticmethod
    def _set_cell_dxa_width(cell, width: int) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        cell_properties = cell._tc.get_or_add_tcPr()
        cell_width = cell_properties.find(qn("w:tcW"))
        if cell_width is None:
            cell_width = OxmlElement("w:tcW")
            cell_properties.insert(0, cell_width)
        cell_width.set(qn("w:type"), "dxa")
        cell_width.set(qn("w:w"), str(width))

    def _set_merged_row_widths(
        self,
        table,
        row_index: int,
        cells: list[tuple[int, int, int]],
    ) -> None:
        for start, span, width in sorted(cells, reverse=True):
            cell = table.cell(row_index, start)
            if span > 1:
                cell = cell.merge(table.cell(row_index, start + span - 1))
            self._set_cell_dxa_width(cell, width)

    def test_pdf_to_editable_docx_and_semantic_html(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "合同.pdf")
            docx_path = root / "合同.docx"
            html_path = root / "合同.html"
            self.assertEqual(pdf_to_docx(source, docx_path), [docx_path])
            document = Document(docx_path)
            self.assertIn(
                "Editable contract text", "\n".join(p.text for p in document.paragraphs)
            )
            self.assertEqual(pdf_to_html(source, html_path), [html_path])
            self.assertIn(
                "Editable contract text", html_path.read_text(encoding="utf-8")
            )

            visual_path = root / "合同_原样.docx"
            self.assertEqual(
                pdf_to_docx(source, visual_path, mode="visual", dpi=110),
                [visual_path],
            )
            self.assertEqual(len(Document(visual_path).inline_shapes), 1)

    def test_hybrid_pdf_to_docx_keeps_editable_pages_and_rasterizes_forced_page(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_three_page_text_pdf(root / "three-pages.pdf")
            output = root / "hybrid.docx"

            self.assertEqual(
                pdf_to_docx(
                    source,
                    output,
                    mode="hybrid",
                    hybrid_force_visual_pages="2",
                    dpi=110,
                ),
                [output],
            )

            document = Document(output)
            self.assertEqual(len(document.sections), 3)
            self.assertEqual(len(document.inline_shapes), 1)
            editable_text = _extract_docx_text(output)
            self.assertIn("Editable page one", editable_text)
            self.assertNotIn("Forced visual page two", editable_text)
            self.assertIn("Editable page three", editable_text)
            self.assertLess(
                editable_text.index("Editable page one"),
                editable_text.index("Editable page three"),
            )

    def test_multi_page_hybrid_has_no_default_height_empty_next_page_sections(
        self,
    ) -> None:
        from docx.oxml.ns import qn

        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_three_page_text_pdf(root / "three-editable-pages.pdf")
            output = root / "three-editable-pages.docx"

            self.assertEqual(
                pdf_to_docx(source, output, mode="hybrid", dpi=110),
                [output],
            )

            document = Document(output)
            self.assertEqual(len(document.sections), 3)
            self.assertEqual(len(document.inline_shapes), 0)

            next_page_breaks = []
            default_height_empty_breaks = []
            for paragraph in document.element.body.xpath("./w:p[w:pPr/w:sectPr]"):
                section_properties = paragraph.xpath("./w:pPr/w:sectPr")[0]
                section_type = section_properties.xpath("./w:type")
                section_type_value = (
                    section_type[0].get(qn("w:val")) if section_type else None
                )
                if section_type_value not in {None, "nextPage"}:
                    continue
                next_page_breaks.append(paragraph)

                has_content = bool(
                    paragraph.xpath(
                        ".//w:t[normalize-space(.)] | .//w:drawing | .//w:pict | "
                        ".//w:object | .//w:br"
                    )
                )
                if has_content:
                    continue
                spacing = paragraph.xpath("./w:pPr/w:spacing")
                is_exact_one_point = bool(spacing) and (
                    spacing[0].get(qn("w:before")) == "0"
                    and spacing[0].get(qn("w:after")) == "0"
                    and spacing[0].get(qn("w:line")) == "20"
                    and spacing[0].get(qn("w:lineRule")) == "exact"
                )
                if not is_exact_one_point:
                    default_height_empty_breaks.append(paragraph.xml)

            self.assertEqual(len(next_page_breaks), 2)
            self.assertEqual(default_height_empty_breaks, [])

    def test_office_native_routes_directly_to_word_reflow(self) -> None:
        quality = conversion_processor._PdfWordCandidateQuality(
            0.99, 0.99, 0.99, (), None, None, 0, 12, 0, True
        )
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "paper.pdf")
            output = root / "office-native.docx"

            def build_office(
                _source: Path, candidate: Path, **_kwargs: object
            ) -> object:
                document = Document()
                document.add_paragraph("Microsoft Word native PDF Reflow.")
                document.save(candidate)
                return quality

            with patch.object(
                conversion_processor,
                "_convert_pdf_with_microsoft_reflow_candidate",
                side_effect=build_office,
            ) as office_converter:
                self.assertEqual(
                    pdf_to_docx(source, output, mode="office_native"),
                    [output],
                )

            office_converter.assert_called_once()
            self.assertIn("native PDF Reflow", _extract_docx_text(output))

    def test_editable_mode_keeps_the_builtin_reconstruction_pipeline(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "paper.pdf")
            output = root / "editable.docx"

            def build_builtin(
                _source: Path, target: Path, **_kwargs: object
            ) -> None:
                document = Document()
                document.add_paragraph("LayoutLoom editable reconstruction.")
                document.save(target)

            with patch.object(
                conversion_processor,
                "_execute_layoutloom_pdf_to_docx",
                side_effect=build_builtin,
            ) as builtin_converter, patch.object(
                conversion_processor,
                "_convert_pdf_with_microsoft_reflow_candidate",
            ) as office_converter:
                self.assertEqual(
                    pdf_to_docx(source, output, mode="editable"),
                    [output],
                )

            builtin_converter.assert_called_once()
            office_converter.assert_not_called()
            self.assertIn("LayoutLoom editable", _extract_docx_text(output))
    def test_hybrid_pdf_to_docx_rasterizes_a_scan_while_editable_rejects_it(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            image_path = root / "scan.png"
            Image.new("RGB", (500, 300), "white").save(image_path)
            source = root / "scan.pdf"
            canvas = Canvas(str(source), pagesize=(500, 300))
            canvas.drawImage(str(image_path), 0, 0, width=500, height=300)
            canvas.save()

            hybrid = root / "scan-hybrid.docx"
            self.assertEqual(
                pdf_to_docx(source, hybrid, mode="hybrid", dpi=110), [hybrid]
            )
            hybrid_document = Document(hybrid)
            self.assertEqual(len(hybrid_document.sections), 1)
            self.assertEqual(len(hybrid_document.inline_shapes), 1)

            editable = root / "scan-editable.docx"
            with self.assertRaises(ValidationError):
                pdf_to_docx(source, editable, mode="editable")
            self.assertFalse(editable.exists())

    def test_hybrid_pdf_to_docx_falls_back_when_editable_engine_cannot_start(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "engine-failure.pdf")
            output = root / "fallback.docx"

            with patch(
                "pdf2docx.Converter",
                side_effect=RuntimeError("simulated engine failure"),
            ), self.assertWarnsRegex(UserWarning, "可编辑引擎初始化"):
                self.assertEqual(
                    pdf_to_docx(source, output, mode="hybrid", dpi=110), [output]
                )

            document = Document(output)
            self.assertEqual(len(document.inline_shapes), 1)
            self.assertNotIn("Editable contract text", _extract_docx_text(output))

    def test_hybrid_warning_as_error_does_not_delete_committed_output(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "warning.pdf")
            output = root / "warning.docx"

            with warnings.catch_warnings():
                warnings.simplefilter("error", UserWarning)
                with self.assertRaisesRegex(UserWarning, "用户指定原样保留"):
                    pdf_to_docx(
                        source,
                        output,
                        mode="hybrid",
                        hybrid_force_visual_pages="1",
                        dpi=110,
                    )

            self.assertTrue(output.is_file())
            self.assertEqual(len(Document(output).inline_shapes), 1)

    def test_temporary_working_directory_releases_then_retries_cleanup(self) -> None:
        events: list[str] = []
        attempts = 0
        real_rmtree = conversion_processor.shutil.rmtree

        def flaky_rmtree(path, *args, **kwargs) -> None:
            nonlocal attempts
            attempts += 1
            events.append("cleanup")
            if attempts == 1:
                raise PermissionError(13, "simulated delayed Windows handle", str(path))
            real_rmtree(path, *args, **kwargs)

        with patch.object(
            conversion_processor.shutil,
            "rmtree",
            side_effect=flaky_rmtree,
        ), patch.object(conversion_processor.time, "sleep"):
            with conversion_processor._temporary_working_directory(
                prefix="docuforge-test-release-",
                before_cleanup=lambda: events.append("release"),
            ) as working:
                (working / "payload.bin").write_bytes(b"payload")

        self.assertEqual(events, ["release", "cleanup", "cleanup"])
        self.assertFalse(working.exists())

    def test_hybrid_closes_temp_pdf_before_cleanup_when_validation_stops(
        self,
    ) -> None:
        from pdf2docx import Converter as RealConverter

        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "locked-region-source.pdf")
            output = root / "locked-region-output.docx"
            assessment = conversion_processor._HybridPageAssessment(
                page_index=0,
                source_text="Editable contract text",
                editable_source_text="Editable contract text",
                editable_text_blocks=("Editable contract text",),
                draw_items=0,
                draw_bbox_max_ratio=0.0,
                reasons=[],
                visual_regions=[
                    conversion_processor._HybridRegion(
                        page_index=0,
                        rect=(30.0, 90.0, 95.0, 140.0),
                        kind="figure",
                        reasons=("test local visual",),
                        dpi=144,
                    )
                ],
            )
            events: list[str] = []
            opened_paths: list[Path] = []
            real_rmtree = conversion_processor.shutil.rmtree

            class TrackingConverter:
                def __init__(self, *args, **kwargs) -> None:
                    self._inner = RealConverter(*args, **kwargs)
                    self._closed = False
                    opened_paths.append(Path(args[0]))
                    events.append("open")

                def __getattr__(self, name: str):
                    return getattr(self._inner, name)

                def close(self) -> None:
                    if self._closed:
                        return
                    self._closed = True
                    self._inner.close()
                    events.append("close")

            def tracked_rmtree(path, *args, **kwargs) -> None:
                if Path(path).name.startswith("docuforge-pdf-docx-"):
                    events.append("cleanup")
                real_rmtree(path, *args, **kwargs)

            def stop_during_page_validation(_value: float, message: str) -> None:
                if message.startswith("校验页面"):
                    raise ValidationError("forced page validation stop")

            with patch(
                "docuforge.processors.conversion._assess_pdf_pages_for_hybrid",
                return_value=[assessment],
            ), patch("pdf2docx.Converter", TrackingConverter), patch(
                "docuforge.runner.report_progress",
                side_effect=stop_during_page_validation,
            ), patch.object(
                conversion_processor.shutil,
                "rmtree",
                side_effect=tracked_rmtree,
            ):
                with self.assertRaisesRegex(
                    ValidationError, "forced page validation stop"
                ):
                    pdf_to_docx(source, output, mode="hybrid", dpi=144)

            self.assertEqual(opened_paths[0].name, "region-hybrid-source.pdf")
            self.assertLess(events.index("close"), events.index("cleanup"))
            self.assertFalse(output.exists())

    def test_hybrid_rasterizes_a_page_when_word_postbuild_quality_drops(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "postbuild-quality.pdf")
            output = root / "postbuild-quality.docx"

            with patch(
                "docuforge.processors.conversion._extract_docx_text",
                return_value="",
            ), self.assertWarnsRegex(UserWarning, "Word 构建后"):
                self.assertEqual(
                    pdf_to_docx(source, output, mode="hybrid", dpi=110), [output]
                )

            document = Document(output)
            self.assertEqual(len(document.inline_shapes), 1)
            self.assertEqual(len(document.sections), 1)

    def test_hybrid_retries_a_narrow_cross_column_merge_as_a_local_image(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "local-merge.pdf"
            output = root / "local-merge.docx"
            canvas = Canvas(str(source), pagesize=(320, 220))
            canvas.drawString(30, 170, "Merged heading row")
            canvas.drawString(30, 130, "Editable body remains")
            canvas.save()
            responses = iter([[(20.0, 35.0, 250.0, 60.0)], []])

            with patch(
                "docuforge.processors.conversion._pdf2docx_intercolumn_text_merge_rects",
                side_effect=lambda *_args, **_kwargs: next(responses, []),
            ), warnings.catch_warnings(record=True) as caught:
                warnings.simplefilter("always")
                self.assertEqual(
                    pdf_to_docx(source, output, mode="hybrid", dpi=144),
                    [output],
                )

            self.assertEqual(len(Document(output).inline_shapes), 1)
            editable_text = _extract_docx_text(output)
            self.assertNotIn("Merged heading row", editable_text)
            self.assertIn("Editable body remains", editable_text)
            self.assertFalse(
                any("整页高清原样保留" in str(item.message) for item in caught)
            )

    def test_pdf2docx_settings_keep_lattice_for_editable_only(self) -> None:
        converter = SimpleNamespace(default_settings={})

        editable_settings = conversion_processor._pdf2docx_settings(converter)
        hybrid_settings = conversion_processor._pdf2docx_settings(
            converter,
            resilient=True,
            parse_lattice_table=False,
            clip_image_res_ratio=450 / 72,
        )
        double_column_settings = conversion_processor._pdf2docx_settings(
            converter,
            column_layout="double",
        )
        mixed_column_settings = conversion_processor._pdf2docx_settings(
            converter,
            column_layout="mixed",
        )

        self.assertTrue(editable_settings["parse_lattice_table"])
        self.assertFalse(hybrid_settings["parse_lattice_table"])
        self.assertAlmostEqual(hybrid_settings["clip_image_res_ratio"], 450 / 72)
        self.assertLess(
            double_column_settings["min_section_height"],
            editable_settings["min_section_height"],
        )
        self.assertEqual(
            mixed_column_settings["min_section_height"],
            double_column_settings["min_section_height"],
        )

    def test_hybrid_clip_image_ratio_uses_highest_visual_region_dpi(self) -> None:
        assessment = SimpleNamespace(
            visual_regions=[
                conversion_processor._HybridRegion(
                    page_index=0,
                    rect=(20.0, 20.0, 80.0, 80.0),
                    kind="figure",
                    dpi=300,
                ),
                conversion_processor._HybridRegion(
                    page_index=0,
                    rect=(100.0, 100.0, 180.0, 140.0),
                    kind="formula",
                    dpi=450,
                ),
            ]
        )

        ratio = conversion_processor._hybrid_clip_image_res_ratio(
            [assessment],
            144,
        )

        self.assertAlmostEqual(ratio, 450 / 72)
        self.assertGreaterEqual(
            conversion_processor._hybrid_clip_image_res_ratio([assessment], 600),
            600 / 72,
        )

    def test_normalized_word_font_name_maps_pdf_fonts_to_portable_word_fonts(
        self,
    ) -> None:
        normalize = conversion_processor._normalized_word_font_name
        self.assertEqual(normalize("LMRoman10"), "Times New Roman")
        self.assertEqual(normalize("ABCDEF+LMRoman8-Regular"), "Times New Roman")
        self.assertEqual(normalize("CMR10"), "Times New Roman")
        self.assertEqual(normalize("CMBX12"), "Times New Roman")
        self.assertEqual(normalize("TeX_CM_Roman"), "Times New Roman")
        self.assertEqual(normalize("CMMI10"), "Cambria Math")
        self.assertEqual(normalize("CMSY7"), "Cambria Math")
        self.assertEqual(normalize("CMEX10"), "Cambria Math")
        self.assertEqual(normalize("TeX_CM_Maths"), "Cambria Math")
        self.assertEqual(normalize("TeX_CM_Maths_Extension"), "Cambria Math")
        self.assertEqual(normalize("rtxmi"), "Cambria Math")
        self.assertEqual(normalize("txsy"), "Cambria Math")
        self.assertEqual(normalize("Yhcmex"), "Cambria Math")
        self.assertEqual(normalize("rtxr"), "Times New Roman")
        self.assertEqual(normalize("TimesLTStd"), "Times New Roman")
        self.assertEqual(normalize("ArialMT"), "Arial")
        self.assertEqual(normalize("FormataOTFMd"), "Arial")
        self.assertEqual(normalize("Calibri"), "Calibri")
        self.assertEqual(normalize("MicrosoftYaHei"), "Microsoft YaHei")
        self.assertEqual(normalize("MicrosoftYaHei-Bold"), "Microsoft YaHei")
        self.assertEqual(normalize("FZSSK--GBK1-0"), "SimSun")
        self.assertEqual(normalize("FZHTK--GBK1-0"), "SimHei")
        self.assertEqual(normalize("E-B1"), "Times New Roman")
        self.assertEqual(normalize("E-BZ"), "SimSun")

    def test_fixed_layout_restores_geometry_confirmed_english_word_spaces(
        self,
    ) -> None:
        from docuforge.processors import pdf_word_layout

        text = "Forthispart"
        origins: list[float] = []
        current_x = 10.0
        for index, character in enumerate(text):
            origins.append(current_x)
            advance = 5.0
            if index in {2, 6}:
                advance += 4.0
            current_x += advance
        line = {
            "dir": (1.0, 0.0),
            "spans": [
                {
                    "font": "E-B1",
                    "size": 10.0,
                    "bbox": (10.0, 20.0, current_x, 32.0),
                    "chars": [
                        {
                            "c": character,
                            "origin": (origin, 30.0),
                            "bbox": (origin, 20.0, origin + 6.0, 32.0),
                        }
                        for character, origin in zip(text, origins)
                    ],
                }
            ],
        }
        model = pdf_word_layout._CharacterAdvanceModel(
            by_character={
                ("E-B1", 10.0, character): 5.0 for character in set(text)
            },
            by_font={("E-B1", 10.0): 5.0},
        )

        prepared = pdf_word_layout._prepare_raw_line_spans(line, model)

        self.assertEqual(prepared[0]["text"], "For this part")

    def test_fixed_layout_keeps_isolated_whitespace_between_editable_spans(
        self,
    ) -> None:
        from docuforge.processors import pdf_word_layout

        clusters, visual_spans = pdf_word_layout._editable_line_clusters(
            [
                {
                    "text": "Hello",
                    "font": "Arial",
                    "size": 10.0,
                    "flags": 0,
                    "color": 0,
                    "bbox": (10.0, 10.0, 35.0, 22.0),
                },
                {
                    "text": " ",
                    "font": "Arial",
                    "size": 10.0,
                    "flags": 0,
                    "color": 0,
                    "bbox": (35.0, 10.0, 39.0, 22.0),
                },
                {
                    "text": "World",
                    "font": "Arial",
                    "size": 10.0,
                    "flags": 0,
                    "color": 0,
                    "bbox": (39.0, 10.0, 64.0, 22.0),
                },
            ],
            plan=pdf_word_layout.FixedLayoutPagePlan(),
            repair_text=lambda text: text,
            font_requires_visual=lambda _font: False,
            math_font=lambda _font: False,
            suspicious_text=lambda _text: False,
        )

        self.assertEqual(visual_spans, 0)
        self.assertEqual(len(clusters), 1)
        self.assertEqual("".join(span.text for span in clusters[0].spans), "Hello World")

    def test_fixed_layout_collects_merged_visual_only_span_hints(self) -> None:
        from docuforge.processors import pdf_word_layout

        visual_hints: list[pdf_word_layout._VisualSpanHint] = []
        clusters, visual_spans = pdf_word_layout._editable_line_clusters(
            [
                {
                    "text": "x",
                    "font": "Math",
                    "size": 10.0,
                    "bbox": (10.0, 10.0, 16.0, 22.0),
                },
                {
                    "text": "2",
                    "font": "Math",
                    "size": 8.0,
                    "bbox": (16.5, 8.0, 20.0, 18.0),
                },
                {
                    "text": "=",
                    "font": "Arial",
                    "size": 10.0,
                    "bbox": (25.0, 10.0, 30.0, 22.0),
                },
                {
                    "text": "bad",
                    "font": "Legacy",
                    "size": 10.0,
                    "bbox": (34.0, 10.0, 48.0, 22.0),
                },
                {
                    "text": "font",
                    "font": "Legacy",
                    "size": 10.0,
                    "bbox": (49.0, 10.0, 65.0, 22.0),
                },
                {
                    "text": "\ufffd",
                    "font": "Arial",
                    "size": 10.0,
                    "bbox": (70.0, 10.0, 75.0, 22.0),
                },
            ],
            plan=pdf_word_layout.FixedLayoutPagePlan(),
            repair_text=lambda text: text,
            font_requires_visual=lambda font: font == "Legacy",
            math_font=lambda font: font == "Math",
            suspicious_text=lambda text: "\ufffd" in text,
            visual_hint_collector=visual_hints,
        )

        self.assertEqual(visual_spans, 5)
        self.assertEqual(len(clusters), 1)
        self.assertEqual(clusters[0].spans[0].text, "=")
        self.assertEqual(
            visual_hints,
            [
                pdf_word_layout._VisualSpanHint(
                    kind="inline_math",
                    bbox=(10.0, 8.0, 20.0, 22.0),
                ),
                pdf_word_layout._VisualSpanHint(
                    kind="text_visual",
                    bbox=(34.0, 10.0, 75.0, 22.0),
                ),
            ],
        )

        converted = pdf_word_layout._visual_hints_to_page_twips(
            visual_hints,
            SimpleNamespace(x0=10.0, y0=8.0, width=100.0, height=200.0),
        )
        self.assertEqual(
            converted,
            (
                ("inline_math", 0, 0, 200, 280),
                ("text_visual", 480, 40, 1300, 280),
            ),
        )

    def test_fixed_layout_uses_one_wps_width_fit_per_positioned_line(self) -> None:
        from docuforge.processors import pdf_word_layout

        document = Document()
        cluster = pdf_word_layout._EditableCluster(
            spans=(
                pdf_word_layout._EditableSpan(
                    text="For this part, you are allowed",
                    font="E-B1",
                    size=11.0,
                    flags=4,
                    color=0,
                    bbox=(40.0, 40.0, 195.0, 55.0),
                ),
                pdf_word_layout._EditableSpan(
                    text=".",
                    font="E-BZ",
                    size=11.0,
                    flags=4,
                    color=0,
                    bbox=(195.0, 40.0, 205.0, 55.0),
                ),
            ),
            bbox=(40.0, 40.0, 205.0, 55.0),
        )
        installed = conversion_processor._standard_word_font_keys()

        pdf_word_layout._append_editable_cluster(
            document,
            cluster,
            frame_id=1,
            resolve_font=lambda name, east_asia: conversion_processor._resolve_word_font_name(
                name,
                east_asia=east_asia,
                installed_fonts=installed,
            ),
        )

        paragraph = document.paragraphs[-1]
        self.assertEqual(len(paragraph.runs), 1)
        self.assertEqual(paragraph.text, "For this part, you are allowed.")
        self.assertEqual(len(paragraph._p.xpath(".//w:fitText")), 1)
        self.assertEqual(len(paragraph._p.xpath(".//w:noProof")), 1)
        self.assertEqual(paragraph.runs[0].font.name, "Times New Roman")

    def test_dense_fixed_layout_adopts_checked_region_candidate(self) -> None:
        from docuforge.processors import pdf_word_layout, word_region

        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            target = self._make_dense_fixed_layout_docx(root / "fixed.docx")
            stats = pdf_word_layout.FixedLayoutBuildStats(
                pages=1,
                editable_frames=60,
                editable_spans=60,
                visual_spans=0,
                visual_hints_by_page={
                    0: (
                        ("inline_math", 100, 100, 200, 200),
                        ("text_visual", 300, 300, 350, 350),
                        ("text_visual", 500, 500, 560, 560),
                    )
                },
            )
            real_builder = word_region.build_region_compatible_word
            received_visual_hints: list[object] = []

            def build_without_real_wps(
                source: Path,
                output: Path,
                **kwargs,
            ) -> list[Path]:
                received_visual_hints.append(kwargs.get("visual_hints_by_page"))
                return real_builder(
                    source,
                    output,
                    verification_engine="none",
                    overwrite=True,
                    progress=kwargs.get("progress"),
                    preserve_editable_text=kwargs.get(
                        "preserve_editable_text",
                        False,
                    ),
                    normalize_text=kwargs.get("normalize_text", True),
                )

            with (
                patch(
                    "docuforge.processors.wps.detect_wps_engines",
                    return_value={"writer": SimpleNamespace(available=True)},
                ),
                patch.object(
                    word_region,
                    "build_region_compatible_word",
                    side_effect=build_without_real_wps,
                ),
                patch.object(
                    conversion_processor,
                    "_docx_raster_pixel_budget",
                    side_effect=(1_000_000, 100_000),
                ),
                patch.object(
                    conversion_processor,
                    "_docx_document_xml_size",
                    side_effect=(500_000, 100_000),
                ),
            ):
                result = conversion_processor._optimize_pdf_fixed_layout_docx_for_wps(
                    target,
                    stats,
                    formula_hints_by_page={0: ((90, 90, 210, 210),)},
                    table_hints_by_page={0: ((280, 280, 380, 380),)},
                )

            document = Document(target)
            body = document.element.body
            self.assertTrue(result.region_optimized)
            self.assertGreater(result.region_text_boxes, 0)
            self.assertLess(result.region_text_boxes, stats.editable_frames)
            self.assertEqual(result.visual_hints_by_page, stats.visual_hints_by_page)
            self.assertEqual(
                received_visual_hints,
                [{0: (("text_visual", 500, 500, 560, 560),)}],
            )
            self.assertFalse(body.xpath(".//w:pPr/w:framePr"))
            self.assertFalse(body.xpath(".//w:fitText"))
            self.assertTrue(body.xpath(".//w:txbxContent/w:p"))

    def test_compact_designed_layout_allows_verified_moderate_region_reduction(
        self,
    ) -> None:
        compact, maximum_regions, maximum_anchors = (
            conversion_processor._fixed_layout_region_object_budget(
                SimpleNamespace(pages=1, editable_frames=49)
            )
        )
        paper, paper_regions, paper_anchors = (
            conversion_processor._fixed_layout_region_object_budget(
                SimpleNamespace(pages=9, editable_frames=1081)
            )
        )

        self.assertTrue(compact)
        self.assertGreaterEqual(maximum_regions, 29)
        self.assertGreaterEqual(maximum_anchors, 39)
        self.assertFalse(paper)
        self.assertEqual(paper_regions, math.ceil(1081 * 0.55))
        self.assertEqual(paper_anchors, math.ceil((1081 + 9) * 0.75))

    def test_fixed_layout_expected_text_keeps_wide_frame_around_small_formula_hint(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as folder:
            source = self._make_dense_fixed_layout_docx(
                Path(folder) / "fixed.docx",
                frames=1,
            )

            retained = conversion_processor._fixed_layout_expected_editable_text(
                source,
                {0: ((4400, 500, 4600, 710),)},
            )
            replaced = conversion_processor._fixed_layout_expected_editable_text(
                source,
                {0: ((900, 500, 8100, 710),)},
            )

            self.assertIn("Dense editable line 1", retained)
            self.assertNotIn("Dense editable line 1", replaced)

    def test_fixed_layout_expected_text_keeps_formula_frame_with_redacted_background(
        self,
    ) -> None:
        from docuforge.processors import word_flow

        with tempfile.TemporaryDirectory() as folder:
            source = self._make_dense_fixed_layout_docx(
                Path(folder) / "fixed.docx",
                frames=1,
            )
            document = Document(source)
            frame = word_flow._paragraph_frame(document.paragraphs[0])
            self.assertIsNotNone(frame)
            page = SimpleNamespace(index=0, frames=[frame], background=b"redacted")

            with patch.object(word_flow, "_source_pages", return_value=[page]):
                retained = conversion_processor._fixed_layout_expected_editable_text(
                    source,
                    {0: ((900, 500, 8100, 710),)},
                )

            self.assertIn("Dense editable line 1", retained)

    def test_dense_fixed_layout_keeps_baseline_when_region_check_fails(self) -> None:
        from docuforge.processors import pdf_word_layout, word_region

        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            target = self._make_dense_fixed_layout_docx(root / "fixed.docx")
            original = target.read_bytes()
            stats = pdf_word_layout.FixedLayoutBuildStats(
                pages=1,
                editable_frames=60,
                editable_spans=60,
                visual_spans=0,
            )

            with (
                patch(
                    "docuforge.processors.wps.detect_wps_engines",
                    return_value={"writer": SimpleNamespace(available=True)},
                ),
                patch.object(
                    word_region,
                    "build_region_compatible_word",
                    side_effect=ValidationError("render mismatch"),
                ),
            ):
                result = conversion_processor._optimize_pdf_fixed_layout_docx_for_wps(
                    target,
                    stats,
                )

            self.assertFalse(result.region_optimized)
            self.assertEqual(target.read_bytes(), original)

    def test_dense_fixed_layout_rejects_candidate_with_editable_text_loss(self) -> None:
        from docuforge.processors import pdf_word_layout, word_region

        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            target = self._make_dense_fixed_layout_docx(root / "fixed.docx")
            original = target.read_bytes()
            stats = pdf_word_layout.FixedLayoutBuildStats(
                pages=1,
                editable_frames=60,
                editable_spans=60,
                visual_spans=0,
            )
            real_builder = word_region.build_region_compatible_word

            def build_without_real_wps(
                source: Path,
                output: Path,
                **kwargs,
            ) -> list[Path]:
                return real_builder(
                    source,
                    output,
                    verification_engine="none",
                    overwrite=True,
                    preserve_editable_text=True,
                    normalize_text=False,
                )

            with (
                patch(
                    "docuforge.processors.wps.detect_wps_engines",
                    return_value={"writer": SimpleNamespace(available=True)},
                ),
                patch.object(
                    word_region,
                    "build_region_compatible_word",
                    side_effect=build_without_real_wps,
                ),
                patch.object(
                    conversion_processor,
                    "_extract_docx_text",
                    side_effect=(
                        "one two three four five six seven eight",
                        "one two three",
                    ),
                ),
            ):
                result = conversion_processor._optimize_pdf_fixed_layout_docx_for_wps(
                    target,
                    stats,
                )

            self.assertFalse(result.region_optimized)
            self.assertEqual(target.read_bytes(), original)

    def test_docx_font_normalizer_falls_back_for_an_uninstalled_family(self) -> None:
        document = Document()
        run = document.add_paragraph().add_run("Portable text")
        run.font.name = "UninstalledCustomSerif"

        installed = frozenset(
            {
                conversion_processor._word_font_key("Arial"),
                conversion_processor._word_font_key("Cambria Math"),
                conversion_processor._word_font_key("Courier New"),
                conversion_processor._word_font_key("Times New Roman"),
            }
        )
        with patch.object(
            conversion_processor,
            "_installed_word_font_keys",
            return_value=installed,
        ):
            self.assertGreater(conversion_processor._normalize_docx_fonts(document), 0)
        self.assertEqual(run.font.name, "Times New Roman")

    def test_docx_font_normalizer_preserves_postscript_times_body_width(self) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt

        document = Document()
        body_run = document.add_paragraph().add_run("Body text must keep its PDF wrap")
        body_run.font.name = "TimesNewRomanPSMT"
        body_run.font.size = Pt(10)
        body_fonts = body_run._r.get_or_add_rPr().get_or_add_rFonts()
        body_fonts.set(qn("w:eastAsia"), "TimesNewRomanPSMT")

        reference_run = document.add_paragraph().add_run("Compact reference text")
        reference_run.font.name = "TimesNewRomanPSMT"
        reference_run.font.size = Pt(8)
        reference_fonts = reference_run._r.get_or_add_rPr().get_or_add_rFonts()
        reference_fonts.set(qn("w:eastAsia"), "TimesNewRomanPSMT")

        pre_scaled_run = document.add_paragraph().add_run("Existing narrower text")
        pre_scaled_run.font.name = "NimbusRomNo9L-Regu"
        pre_scaled_run.font.size = Pt(10)
        pre_scaled_fonts = pre_scaled_run._r.get_or_add_rPr().get_or_add_rFonts()
        pre_scaled_fonts.set(qn("w:eastAsia"), "NimbusRomNo9L-Regu")
        width = OxmlElement("w:w")
        width.set(qn("w:val"), "92")
        pre_scaled_run._r.get_or_add_rPr().append(width)

        tex_gyre_run = document.add_paragraph().add_run(
            "Dense references must keep their source pagination"
        )
        tex_gyre_run.font.name = "TeXGyreTermesX-Regular"
        tex_gyre_run.font.size = Pt(10)
        tex_gyre_fonts = tex_gyre_run._r.get_or_add_rPr().get_or_add_rFonts()
        tex_gyre_fonts.set(qn("w:eastAsia"), "TeXGyreTermesX-Regular")

        with patch.object(
            conversion_processor,
            "_installed_word_font_keys",
            return_value=conversion_processor._standard_word_font_keys(),
        ):
            self.assertGreater(conversion_processor._normalize_docx_fonts(document), 0)

        self.assertEqual(body_run.font.name, "Times New Roman")
        body_width = body_run._r.xpath("./w:rPr/w:w")
        self.assertEqual(len(body_width), 1)
        self.assertEqual(body_width[0].get(qn("w:val")), "96")
        self.assertFalse(reference_run._r.xpath("./w:rPr/w:w"))
        existing_width = pre_scaled_run._r.xpath("./w:rPr/w:w")
        self.assertEqual(existing_width[0].get(qn("w:val")), "92")
        tex_gyre_width = tex_gyre_run._r.xpath("./w:rPr/w:w")
        self.assertEqual(tex_gyre_run.font.name, "Times New Roman")
        self.assertEqual(tex_gyre_width[0].get(qn("w:val")), "96")

        with patch.object(
            conversion_processor,
            "_installed_word_font_keys",
            return_value=conversion_processor._standard_word_font_keys(),
        ):
            conversion_processor._normalize_docx_fonts(document)
        self.assertEqual(body_width[0].get(qn("w:val")), "96")

    def test_pdf2docx_layout_stabilizer_disables_grid_and_widow_reflow(self) -> None:
        from docx.oxml.ns import qn
        from docx.shared import Pt

        document = Document()
        paragraph = document.add_paragraph("A PDF-derived paragraph")
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.paragraph_format.line_spacing = 1.1

        self.assertGreater(
            conversion_processor._stabilize_pdf2docx_paragraph_layout(document), 0
        )

        properties = paragraph._p.get_or_add_pPr()
        self.assertEqual(properties.find(qn("w:snapToGrid")).get(qn("w:val")), "0")
        self.assertEqual(properties.find(qn("w:widowControl")).get(qn("w:val")), "0")
        spacing = properties.find(qn("w:spacing"))
        self.assertEqual(spacing.get(qn("w:before")), "200")
        self.assertEqual(spacing.get(qn("w:after")), "100")
        self.assertEqual(spacing.get(qn("w:line")), "264")
        compatibility = document.settings.element.find(qn("w:compat"))
        do_not_expand = compatibility.find(qn("w:doNotExpandShiftReturn"))
        self.assertEqual(do_not_expand.get(qn("w:val")), "1")

    def test_pdf2docx_layout_stabilizer_caps_oversized_body_lines_once(self) -> None:
        from docx.oxml.ns import qn
        from docx.shared import Pt

        document = Document()
        paragraph = document.add_paragraph()
        run = paragraph.add_run("Dense 10 point body text must stay on its source page")
        run.font.size = Pt(10)
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.line_spacing = Pt(30)

        first = conversion_processor._stabilize_pdf2docx_paragraph_layout(document)
        properties = paragraph._p.get_or_add_pPr()
        spacing = properties.find(qn("w:spacing"))
        width = run._r.get_or_add_rPr().find(qn("w:w"))

        self.assertGreater(first, 0)
        self.assertEqual(width.get(qn("w:val")), "98")
        self.assertEqual(spacing.get(qn("w:lineRule")), "exact")
        self.assertEqual(spacing.get(qn("w:line")), "392")
        self.assertEqual(spacing.get(qn("w:before")), "196")

        second = conversion_processor._stabilize_pdf2docx_paragraph_layout(document)
        self.assertEqual(second, 0)
        self.assertEqual(width.get(qn("w:val")), "98")
        self.assertEqual(spacing.get(qn("w:line")), "392")
        self.assertEqual(spacing.get(qn("w:before")), "196")

    def test_pdf2docx_layout_stabilizer_keeps_visual_and_section_carriers(self) -> None:
        from docx.enum.section import WD_SECTION
        from docx.oxml.ns import qn
        from docx.shared import Pt

        document = Document()
        visual_carrier = document.add_paragraph()
        visual_carrier.paragraph_format.space_before = Pt(10)
        visual_carrier.paragraph_format.line_spacing = 1.1
        document.add_section(WD_SECTION.CONTINUOUS)
        section_carrier = document.paragraphs[-1]
        section_carrier.add_run("Hidden section carrier")
        section_carrier.paragraph_format.space_before = Pt(10)
        section_carrier.paragraph_format.line_spacing = 1.1

        conversion_processor._stabilize_pdf2docx_paragraph_layout(document)

        visual_spacing = visual_carrier._p.xpath("./w:pPr/w:spacing")[0]
        section_spacing = section_carrier._p.xpath("./w:pPr/w:spacing")[0]
        self.assertEqual(visual_spacing.get(qn("w:before")), "200")
        self.assertEqual(visual_spacing.get(qn("w:line")), "264")
        self.assertEqual(section_spacing.get(qn("w:before")), "200")
        self.assertEqual(section_spacing.get(qn("w:line")), "264")

    def test_docx_font_normalizer_updates_styles_with_effects_after_save(
        self,
    ) -> None:
        styles_with_effects = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:style w:type="paragraph" w:styleId="EffectStyle">
    <w:name w:val="Effect Style"/>
    <w:rPr>
      <w:rFonts w:ascii="CMR10" w:hAnsi="CMBX12" w:eastAsia="TeX_CM_Roman" w:cs="CMMI10"/>
    </w:rPr>
  </w:style>
</w:styles>
"""
        installed = conversion_processor._standard_word_font_keys()
        effects_relationship = (
            "http://schemas.microsoft.com/office/2007/relationships/"
            "stylesWithEffects"
        )

        with tempfile.TemporaryDirectory() as folder:
            path = Path(folder) / "styles-with-effects.docx"
            document = Document()
            effects_part = next(
                relationship.target_part
                for relationship in document.part.rels.values()
                if relationship.reltype == effects_relationship
            )
            effects_part._blob = styles_with_effects
            document.save(path)

            with patch.object(
                conversion_processor,
                "_installed_word_font_keys",
                return_value=installed,
            ):
                self.assertGreaterEqual(
                    conversion_processor._normalize_docx_file_fonts(path), 4
                )

            with ZipFile(path) as archive:
                saved_xml = archive.read("word/stylesWithEffects.xml")
                self.assertIsNone(archive.testzip())

            root = ElementTree.fromstring(saved_xml)
            namespace = {
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            }
            font_element = root.find(".//w:rFonts", namespace)
            self.assertIsNotNone(font_element)
            word = namespace["w"]
            self.assertEqual(font_element.get(f"{{{word}}}ascii"), "Times New Roman")
            self.assertEqual(font_element.get(f"{{{word}}}hAnsi"), "Times New Roman")
            self.assertEqual(font_element.get(f"{{{word}}}eastAsia"), "Times New Roman")
            self.assertEqual(font_element.get(f"{{{word}}}cs"), "Cambria Math")

    def test_editable_docx_postprocess_compacts_empty_next_page_carriers(self) -> None:
        from docx.enum.section import WD_SECTION
        from docx.oxml.ns import qn

        with tempfile.TemporaryDirectory() as folder:
            path = Path(folder) / "editable-pages.docx"
            document = Document()
            document.add_paragraph("Dense first source page")
            document.add_section(WD_SECTION.NEW_PAGE)
            carrier = document.paragraphs[-1]
            document.add_paragraph("Second source page")
            document.save(path)

            self.assertFalse(carrier._p.xpath("./w:pPr/w:spacing"))
            conversion_processor._normalize_docx_file_fonts(path)

            normalized = Document(path)
            carriers = [
                paragraph
                for paragraph in normalized.paragraphs
                if paragraph._p.xpath("./w:pPr/w:sectPr")
            ]
            self.assertEqual(len(carriers), 1)
            spacing = carriers[0]._p.xpath("./w:pPr/w:spacing")
            self.assertEqual(len(spacing), 1)
            self.assertEqual(spacing[0].get(qn("w:before")), "0")
            self.assertEqual(spacing[0].get(qn("w:after")), "0")
            self.assertEqual(spacing[0].get(qn("w:line")), "20")
            self.assertEqual(spacing[0].get(qn("w:lineRule")), "exact")

    def test_editable_docx_postprocess_floats_running_footer_without_scaling_lines(
        self,
    ) -> None:
        from docx.enum.section import WD_SECTION
        from docx.oxml.ns import qn
        from docx.shared import Pt

        with tempfile.TemporaryDirectory() as folder:
            path = Path(folder) / "running-footer.docx"
            document = Document()
            body = document.add_paragraph("Dense editable body paragraph")
            body.paragraph_format.space_before = Pt(3)
            body.paragraph_format.line_spacing = Pt(11)
            footer = document.add_paragraph(
                "Journal | https://doi.org/10.1000/example 1 / 2"
            )
            footer.runs[0].font.size = Pt(8)
            footer.paragraph_format.space_before = Pt(24)
            footer.paragraph_format.line_spacing = Pt(8)
            document.add_section(WD_SECTION.NEW_PAGE)
            document.add_paragraph("Second source page")
            document.save(path)

            conversion_processor._normalize_docx_file_fonts(path)
            normalized = Document(path)
            body = next(
                paragraph
                for paragraph in normalized.paragraphs
                if paragraph.text == "Dense editable body paragraph"
            )
            footer = next(
                paragraph
                for paragraph in normalized.paragraphs
                if paragraph.text.startswith("Journal |")
            )

            frame = footer._p.xpath("./w:pPr/w:framePr")
            self.assertEqual(len(frame), 1)
            self.assertEqual(frame[0].get(qn("w:hAnchor")), "page")
            self.assertEqual(frame[0].get(qn("w:vAnchor")), "page")
            self.assertEqual(frame[0].get(qn("w:wrap")), "none")
            footer_spacing = footer._p.xpath("./w:pPr/w:spacing")[0]
            self.assertEqual(footer_spacing.get(qn("w:before")), "0")
            self.assertEqual(footer_spacing.get(qn("w:after")), "0")
            self.assertEqual(footer_spacing.get(qn("w:line")), "160")

            body_spacing = body._p.xpath("./w:pPr/w:spacing")[0]
            self.assertEqual(body_spacing.get(qn("w:before")), "30")
            self.assertEqual(body_spacing.get(qn("w:line")), "220")

            # Re-running the postprocessor must not keep shrinking body gaps.
            conversion_processor._normalize_docx_file_fonts(path)
            rerun = Document(path)
            rerun_body = next(
                paragraph
                for paragraph in rerun.paragraphs
                if paragraph.text == "Dense editable body paragraph"
            )
            rerun_spacing = rerun_body._p.xpath("./w:pPr/w:spacing")[0]
            self.assertEqual(rerun_spacing.get(qn("w:before")), "30")
            self.assertEqual(rerun_spacing.get(qn("w:line")), "220")

    def test_footer_split_keeps_next_page_section_on_a_compact_carrier(self) -> None:
        from docx.enum.section import WD_SECTION
        from docx.oxml.ns import qn
        from docx.shared import Pt
        from docx.text.paragraph import Paragraph

        document = Document()
        document.add_paragraph("Dense first source page")
        footer = document.add_paragraph("PLOS ONE | December 12, 2024 1 / 2")
        footer.runs[0].font.size = Pt(8)
        footer.paragraph_format.space_before = Pt(12)
        footer.paragraph_format.line_spacing = Pt(8)
        document.add_section(WD_SECTION.NEW_PAGE)
        empty_carrier = document.paragraphs[-1]
        section_properties = empty_carrier._p.get_or_add_pPr().sectPr
        footer._p.get_or_add_pPr().append(section_properties)
        document.element.body.remove(empty_carrier._p)
        document.add_paragraph("Second source page")

        self.assertGreater(
            conversion_processor._float_pdf2docx_running_footers(document),
            0,
        )

        carrier = footer._p
        self.assertEqual(footer.text, "")
        self.assertTrue(carrier.xpath("./w:pPr/w:sectPr"))
        self.assertFalse(carrier.xpath("./w:pPr/w:framePr"))
        carrier_spacing = carrier.xpath("./w:pPr/w:spacing")
        self.assertEqual(carrier_spacing[0].get(qn("w:line")), "20")
        self.assertEqual(carrier_spacing[0].get(qn("w:lineRule")), "exact")

        floating_element = carrier.getprevious()
        floating_footer = Paragraph(floating_element, document._body)
        self.assertEqual(
            floating_footer.text,
            "PLOS ONE | December 12, 2024 1 / 2",
        )
        self.assertTrue(floating_element.xpath("./w:pPr/w:framePr"))
        self.assertFalse(floating_element.xpath("./w:pPr/w:sectPr"))
        self.assertEqual(len(document.sections), 2)

    def test_standalone_chinese_page_counters_are_anchored_to_the_page(self) -> None:
        from docx.enum.section import WD_SECTION
        from docx.oxml.ns import qn
        from docx.shared import Pt

        document = Document()
        document.add_paragraph("Dense first source page body")
        first_counter = document.add_paragraph("第1 5页共53页")
        first_counter.runs[0].font.size = Pt(9)
        first_counter.paragraph_format.space_before = Pt(18)
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_paragraph("Dense second source page body")
        second_counter = document.add_paragraph("第26页共53页")
        second_counter.runs[0].font.size = Pt(9)
        second_counter.paragraph_format.space_before = Pt(24)
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_paragraph("Third source page body")

        self.assertEqual(
            conversion_processor._float_pdf2docx_isolated_page_counters(document),
            2,
        )
        for counter in (first_counter, second_counter):
            frames = counter._p.xpath("./w:pPr/w:framePr")
            self.assertEqual(len(frames), 1)
            self.assertEqual(frames[0].get(qn("w:hAnchor")), "page")
            self.assertEqual(frames[0].get(qn("w:vAnchor")), "page")
            spacing = counter._p.xpath("./w:pPr/w:spacing")[0]
            self.assertEqual(spacing.get(qn("w:before")), "0")
            self.assertEqual(spacing.get(qn("w:after")), "0")

        self.assertEqual(
            conversion_processor._float_pdf2docx_isolated_page_counters(document),
            0,
        )

    def test_page_counter_anchor_does_not_float_body_text_with_trailing_counter(
        self,
    ) -> None:
        from docx.shared import Pt

        document = Document()
        body = document.add_paragraph(
            "D. This is substantive answer text 第1 8页共53页"
        )
        large_title = document.add_paragraph("第1页")
        large_title.runs[0].font.size = Pt(18)

        self.assertEqual(
            conversion_processor._float_pdf2docx_isolated_page_counters(document),
            0,
        )
        self.assertFalse(body._p.xpath("./w:pPr/w:framePr"))
        self.assertFalse(large_title._p.xpath("./w:pPr/w:framePr"))

    def test_adjacent_framed_source_page_counters_keep_the_newer_counter(
        self,
    ) -> None:
        from docx.shared import Pt

        document = Document()
        document.add_paragraph("Dense source page body")
        older = document.add_paragraph("第24页共53页")
        newer = document.add_paragraph("第26页共53页")
        older.paragraph_format.space_before = Pt(1)
        newer.paragraph_format.space_before = Pt(1)
        document.add_paragraph("Following editable body")

        self.assertEqual(
            conversion_processor._float_pdf2docx_isolated_page_counters(document),
            2,
        )
        self.assertEqual(
            conversion_processor._deduplicate_adjacent_pdf2docx_page_counters(document),
            1,
        )
        self.assertNotIn(older._p, document.element.body)
        self.assertIn(newer._p, document.element.body)
        self.assertEqual(
            conversion_processor._deduplicate_adjacent_pdf2docx_page_counters(document),
            0,
        )

    def test_adjacent_counter_deduplication_does_not_cross_body_text(self) -> None:
        from docx.shared import Pt

        document = Document()
        first = document.add_paragraph("第24页共53页")
        document.add_paragraph("Real editable content between source pages")
        second = document.add_paragraph("第26页共53页")
        first.paragraph_format.space_before = Pt(1)
        second.paragraph_format.space_before = Pt(1)

        self.assertEqual(
            conversion_processor._float_pdf2docx_isolated_page_counters(document),
            2,
        )
        self.assertEqual(
            conversion_processor._deduplicate_adjacent_pdf2docx_page_counters(document),
            0,
        )
        self.assertIn(first._p, document.element.body)
        self.assertIn(second._p, document.element.body)

    def test_rendered_counter_conflict_repair_removes_only_unique_older_counter(
        self,
    ) -> None:
        from docx.shared import Pt

        with tempfile.TemporaryDirectory() as folder:
            path = Path(folder) / "conflict.docx"
            document = Document()
            document.add_paragraph("Dense source page body")
            older = document.add_paragraph("第24页共53页")
            document.add_paragraph("Intervening editable content")
            newer = document.add_paragraph("第26页共53页")
            older.paragraph_format.space_before = Pt(1)
            newer.paragraph_format.space_before = Pt(1)
            conversion_processor._float_pdf2docx_isolated_page_counters(document)
            document.save(path)

            conflict = conversion_processor._RenderedPageCounterConflict(
                page_number=27,
                counters=((24, 53), (26, 53)),
            )
            self.assertEqual(
                conversion_processor._repair_pdf2docx_rendered_page_counter_conflicts(
                    path,
                    (conflict,),
                ),
                1,
            )
            repaired = Document(path)
            texts = [paragraph.text for paragraph in repaired.paragraphs]
            self.assertNotIn("第24页共53页", texts)
            self.assertIn("第26页共53页", texts)
            self.assertIn("Intervening editable content", texts)
            self.assertEqual(
                conversion_processor._repair_pdf2docx_rendered_page_counter_conflicts(
                    path,
                    (conflict,),
                ),
                0,
            )

    def test_page_counter_anchor_splits_an_attached_next_page_section(self) -> None:
        from docx.enum.section import WD_SECTION

        document = Document()
        document.add_paragraph("Dense source page body")
        counter = document.add_paragraph("第26页共53页")
        document.add_section(WD_SECTION.NEW_PAGE)
        empty_carrier = document.paragraphs[-1]
        section_properties = empty_carrier._p.get_or_add_pPr().sectPr
        counter._p.get_or_add_pPr().append(section_properties)
        document.element.body.remove(empty_carrier._p)
        document.add_paragraph("Next source page body")

        self.assertEqual(
            conversion_processor._float_pdf2docx_isolated_page_counters(document),
            1,
        )

        carrier = counter._p
        self.assertEqual(counter.text, "")
        self.assertTrue(carrier.xpath("./w:pPr/w:sectPr"))
        self.assertFalse(carrier.xpath("./w:pPr/w:framePr"))
        floating_counter = next(
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text == "第26页共53页"
        )
        self.assertEqual(floating_counter.text, "第26页共53页")
        self.assertTrue(floating_counter._p.xpath("./w:pPr/w:framePr"))
        self.assertFalse(floating_counter._p.xpath("./w:pPr/w:sectPr"))
        self.assertIs(
            floating_counter._p.getnext(),
            next(
                paragraph._p
                for paragraph in document.paragraphs
                if paragraph.text == "Dense source page body"
            ),
        )

    def test_fullwidth_wrapped_page_number_is_floated_and_is_idempotent(
        self,
    ) -> None:
        from docx.enum.section import WD_SECTION
        from docx.oxml.ns import qn
        from docx.shared import Pt

        document = Document()
        document.add_paragraph("Dense reconstructed body")
        footer = document.add_paragraph("—４—")
        footer.runs[0].font.size = Pt(8)
        footer.paragraph_format.space_before = Pt(12)
        document.add_section(WD_SECTION.NEW_PAGE)
        empty_carrier = document.paragraphs[-1]
        section_properties = empty_carrier._p.get_or_add_pPr().sectPr
        footer._p.get_or_add_pPr().append(section_properties)
        document.element.body.remove(empty_carrier._p)
        document.add_paragraph("Next source page")

        self.assertGreater(
            conversion_processor._float_pdf2docx_running_footers(document),
            0,
        )
        framed = [
            paragraph for paragraph in document.paragraphs if paragraph.text == "—４—"
        ]
        self.assertEqual(len(framed), 1)
        self.assertTrue(framed[0]._p.xpath("./w:pPr/w:framePr"))
        carriers = document.element.body.xpath("./w:p[w:pPr/w:sectPr]")
        self.assertEqual(len(carriers), 1)
        self.assertFalse(carriers[0].xpath(".//w:t[normalize-space(.)]"))
        spacing = carriers[0].xpath("./w:pPr/w:spacing")[0]
        self.assertEqual(spacing.get(qn("w:line")), "20")

        snapshot = document.element.body.xml
        self.assertEqual(
            conversion_processor._float_pdf2docx_running_footers(document),
            0,
        )
        self.assertEqual(document.element.body.xml, snapshot)

    def test_spaced_multidigit_fullwidth_page_number_is_floated(self) -> None:
        from docx.enum.section import WD_SECTION
        from docx.shared import Pt

        document = Document()
        document.add_paragraph("Dense reconstructed body")
        footer = document.add_paragraph("—１ ０—")
        footer.runs[0].font.size = Pt(8)
        footer.paragraph_format.space_before = Pt(12)
        document.add_section(WD_SECTION.NEW_PAGE)
        empty_carrier = document.paragraphs[-1]
        section_properties = empty_carrier._p.get_or_add_pPr().sectPr
        footer._p.get_or_add_pPr().append(section_properties)
        document.element.body.remove(empty_carrier._p)
        document.add_paragraph("Next source page")

        self.assertGreater(
            conversion_processor._float_pdf2docx_running_footers(document),
            0,
        )
        framed = [
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text == "—１ ０—"
        ]
        self.assertEqual(len(framed), 1)
        self.assertTrue(framed[0]._p.xpath("./w:pPr/w:framePr"))

    def test_footer_fit_reserve_uses_empty_geometry_spacing_once(self) -> None:
        from docx.enum.section import WD_SECTION
        from docx.oxml.ns import qn
        from docx.shared import Pt

        document = Document()
        geometry_spacer = document.add_paragraph()
        geometry_spacer.paragraph_format.space_after = Pt(30)
        body = document.add_paragraph("Dense body line")
        body.paragraph_format.space_before = Pt(6)
        footer = document.add_paragraph("—４—")
        footer.runs[0].font.size = Pt(8)
        footer.paragraph_format.space_before = Pt(12)
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_paragraph("Next source page")

        self.assertGreater(
            conversion_processor._float_pdf2docx_running_footers(document),
            0,
        )
        geometry_spacing = geometry_spacer._p.xpath("./w:pPr/w:spacing")[0]
        self.assertEqual(geometry_spacing.get(qn("w:after")), "540")
        first_snapshot = document.element.body.xml

        self.assertEqual(
            conversion_processor._float_pdf2docx_running_footers(document),
            0,
        )
        self.assertEqual(document.element.body.xml, first_snapshot)
        self.assertEqual(geometry_spacing.get(qn("w:after")), "540")

    def test_footer_page_fit_slack_is_capped_at_sixty_twips_per_section(
        self,
    ) -> None:
        from docx.enum.section import WD_SECTION
        from docx.oxml.ns import qn
        from docx.shared import Pt

        document = Document()
        body_paragraphs = []
        for index in range(4):
            paragraph = document.add_paragraph(f"Dense body line {index + 1}")
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.line_spacing = Pt(11)
            body_paragraphs.append(paragraph)
        footer = document.add_paragraph("Journal footer 1 / 2")
        footer.runs[0].font.size = Pt(8)
        footer.paragraph_format.space_before = Pt(12)
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_paragraph("Second source page")

        before_total = sum(
            int(paragraph._p.xpath("./w:pPr/w:spacing")[0].get(qn("w:before")))
            for paragraph in body_paragraphs
        )
        conversion_processor._float_pdf2docx_running_footers(document)
        after_total = sum(
            int(paragraph._p.xpath("./w:pPr/w:spacing")[0].get(qn("w:before")))
            for paragraph in body_paragraphs
        )

        self.assertEqual(before_total - after_total, 60)
        for paragraph in body_paragraphs:
            spacing = paragraph._p.xpath("./w:pPr/w:spacing")[0]
            self.assertEqual(spacing.get(qn("w:line")), "220")

    def test_repeated_shallow_wide_graphic_footers_become_page_anchors(
        self,
    ) -> None:
        from docx.enum.section import WD_SECTION
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt

        with tempfile.TemporaryDirectory() as folder:
            image_path = Path(folder) / "running-footer.png"
            Image.new("RGB", (1400, 80), "white").save(image_path)
            document = Document()
            bodies = []
            footers = []
            for page_index in range(2):
                body = document.add_paragraph(f"Editable source page {page_index + 1}")
                body.paragraph_format.space_before = Pt(6)
                body.paragraph_format.line_spacing = Pt(11)
                bodies.append(body)
                footer = document.add_paragraph()
                footer.paragraph_format.left_indent = Pt(36 if page_index == 0 else 0)
                footer.paragraph_format.space_before = Pt(18)
                footer.add_run().add_picture(
                    str(image_path),
                    width=Inches(7),
                    height=Pt(32),
                )
                footers.append(footer)
                document.add_section(WD_SECTION.NEW_PAGE)
            document.add_paragraph("Third source page")
            for section in document.sections:
                section.left_margin = Pt(0)
                section.right_margin = Pt(0)
                section.top_margin = Pt(0)
                section.bottom_margin = Pt(0)

            self.assertGreaterEqual(
                conversion_processor._float_pdf2docx_running_footers(document),
                2,
            )

            anchors = document.element.body.xpath(".//wp:anchor")
            self.assertEqual(len(anchors), 2)
            self.assertTrue(all(footer._p.getparent() is None for footer in footers))
            self.assertFalse(document.element.body.xpath(".//w:pPr/w:framePr"))
            first_x = anchors[0].xpath("./wp:positionH/wp:posOffset")[0]
            second_x = anchors[1].xpath("./wp:positionH/wp:posOffset")[0]
            self.assertEqual(first_x.text, str(int(Pt(36))))
            self.assertEqual(second_x.text, "0")
            for anchor in anchors:
                self.assertEqual(anchor.get("behindDoc"), "0")
                self.assertEqual(anchor.get("allowOverlap"), "1")
                self.assertEqual(
                    anchor.xpath("./wp:positionV")[0].get("relativeFrom"),
                    "page",
                )

            before_rerun = [
                body._p.xpath("./w:pPr/w:spacing")[0].get(qn("w:before"))
                for body in bodies
            ]
            self.assertEqual(
                conversion_processor._float_pdf2docx_running_footers(document),
                0,
            )
            after_rerun = [
                body._p.xpath("./w:pPr/w:spacing")[0].get(qn("w:before"))
                for body in bodies
            ]
            self.assertEqual(before_rerun, after_rerun)
            self.assertEqual(len(document.element.body.xpath(".//wp:anchor")), 2)

    def test_unique_shallow_wide_graphic_is_not_treated_as_running_footer(
        self,
    ) -> None:
        from docx.enum.section import WD_SECTION
        from docx.shared import Inches, Pt

        with tempfile.TemporaryDirectory() as folder:
            image_path = Path(folder) / "unique-wide-figure.png"
            Image.new("RGB", (1400, 80), "white").save(image_path)
            document = Document()
            document.add_paragraph("Editable body")
            figure = document.add_paragraph()
            figure.add_run().add_picture(
                str(image_path),
                width=Inches(7),
                height=Pt(32),
            )
            document.add_section(WD_SECTION.NEW_PAGE)
            document.add_paragraph("Second source page")
            for section in document.sections:
                section.left_margin = Pt(0)
                section.right_margin = Pt(0)

            self.assertEqual(
                conversion_processor._float_pdf2docx_running_footers(document),
                0,
            )
            self.assertIs(figure._p.getparent(), document.element.body)
            self.assertTrue(figure._p.xpath(".//wp:inline"))
            self.assertFalse(figure._p.xpath(".//wp:anchor"))

    def test_repeated_graphic_footer_uses_dynamic_slack_only_for_dense_flow(
        self,
    ) -> None:
        from docx.enum.section import WD_SECTION
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt

        with tempfile.TemporaryDirectory() as folder:
            image_path = Path(folder) / "graphic-footer.png"
            Image.new("RGB", (1400, 80), "white").save(image_path)

            def build_document(*, dense: bool):
                document = Document()
                body_groups = []
                for page_index in range(2):
                    body_group = []
                    lengths = (1300, 1200, 1200) if dense else (80, 80, 80, 80)
                    for paragraph_index, text_length in enumerate(lengths):
                        paragraph = document.add_paragraph(
                            f"{page_index}-{paragraph_index} " + "x" * text_length
                        )
                        paragraph.paragraph_format.space_before = Pt(15 if dense else 6)
                        paragraph.paragraph_format.line_spacing = Pt(11)
                        body_group.append(paragraph)
                    body_groups.append(body_group)
                    footer = document.add_paragraph()
                    footer.paragraph_format.space_before = Pt(18)
                    footer.add_run().add_picture(
                        str(image_path),
                        width=Inches(7),
                        height=Pt(36),
                    )
                    document.add_section(WD_SECTION.NEW_PAGE)
                document.add_paragraph("Final source page")
                for section in document.sections:
                    section.left_margin = Pt(0)
                    section.right_margin = Pt(0)
                    section.top_margin = Pt(0)
                    section.bottom_margin = Pt(0)
                return document, body_groups

            for dense, expected_reduction in ((False, 60), (True, 720)):
                with self.subTest(dense=dense):
                    document, body_groups = build_document(dense=dense)
                    before_totals = [
                        sum(
                            int(
                                paragraph._p.xpath("./w:pPr/w:spacing")[0].get(
                                    qn("w:before")
                                )
                            )
                            for paragraph in body_group
                        )
                        for body_group in body_groups
                    ]

                    conversion_processor._float_pdf2docx_running_footers(document)

                    after_totals = [
                        sum(
                            int(
                                paragraph._p.xpath("./w:pPr/w:spacing")[0].get(
                                    qn("w:before")
                                )
                            )
                            for paragraph in body_group
                        )
                        for body_group in body_groups
                    ]
                    self.assertEqual(
                        [
                            before - after
                            for before, after in zip(before_totals, after_totals)
                        ],
                        [expected_reduction, expected_reduction],
                    )
                    for body_group in body_groups:
                        for paragraph in body_group:
                            spacing = paragraph._p.xpath("./w:pPr/w:spacing")[0]
                            self.assertEqual(spacing.get(qn("w:line")), "220")

    def test_visual_page_after_trailing_table_keeps_section_break_after_table(
        self,
    ) -> None:
        from docx.oxml.ns import qn

        class FakePixmap:
            width = 320
            height = 220

            @staticmethod
            def save(path: str) -> None:
                Image.new("RGB", (320, 220), "white").save(path)

        class FakePage:
            rect = SimpleNamespace(width=320, height=220)

            @staticmethod
            def get_pixmap(*, dpi: int, alpha: bool) -> FakePixmap:
                self.assertEqual(dpi, 110)
                self.assertFalse(alpha)
                return FakePixmap()

        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            output = root / "table-then-visual.docx"
            document = Document()
            document.add_paragraph("Editable page heading")
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "Last editable-page block"

            conversion_processor._append_visual_pdf_page(
                document,
                FakePage(),
                root / "visual-page.png",
                dpi=110,
                is_first_page=False,
            )
            document.save(output)

            reloaded = Document(output)
            body_blocks = [
                child
                for child in reloaded.element.body
                if child.tag.rsplit("}", 1)[-1] != "sectPr"
            ]
            self.assertEqual(
                [child.tag.rsplit("}", 1)[-1] for child in body_blocks],
                ["p", "tbl", "p", "p"],
            )
            self.assertEqual(len(reloaded.sections), 2)
            self.assertEqual(len(reloaded.tables), 1)
            self.assertEqual(
                reloaded.tables[0].cell(0, 0).text, "Last editable-page block"
            )
            self.assertTrue(body_blocks[2].xpath("./w:pPr/w:sectPr"))
            self.assertTrue(body_blocks[3].xpath(".//w:drawing"))
            visual_spacing = body_blocks[3].xpath("./w:pPr/w:spacing")
            self.assertEqual(len(visual_spacing), 1)
            self.assertEqual(visual_spacing[0].get(qn("w:line")), "240")
            self.assertEqual(visual_spacing[0].get(qn("w:lineRule")), "auto")
            self.assertEqual(len(reloaded.inline_shapes), 1)

    def test_pdf2docx_page_append_moves_next_page_section_break_to_previous_paragraph(
        self,
    ) -> None:
        from docx.enum.section import WD_SECTION
        from docx.oxml.ns import qn

        class FakeParsedPage:
            @staticmethod
            def make_docx(document: Document) -> None:
                document.add_section(WD_SECTION.NEW_PAGE)
                document.add_paragraph("Second page body")

        document = Document()
        previous_paragraph = document.add_paragraph("First page body")

        conversion_processor._append_pdf2docx_page(
            document,
            FakeParsedPage(),
            is_first_page=False,
        )

        body_blocks = [
            child for child in document.element.body if child.tag != qn("w:sectPr")
        ]
        self.assertEqual(
            [child.tag.rsplit("}", 1)[-1] for child in body_blocks],
            ["p", "p"],
        )
        self.assertEqual(
            [paragraph.text for paragraph in document.paragraphs],
            [
                "First page body",
                "Second page body",
            ],
        )
        self.assertEqual(len(document.sections), 2)

        relocated = previous_paragraph._p.xpath("./w:pPr/w:sectPr")
        self.assertEqual(len(relocated), 1)
        section_type = relocated[0].xpath("./w:type")
        if section_type:
            self.assertEqual(section_type[0].get(qn("w:val")), "nextPage")

    def test_pdf2docx_page_append_compacts_next_page_section_break_after_table(
        self,
    ) -> None:
        from docx.enum.section import WD_SECTION
        from docx.oxml.ns import qn

        class FakeParsedPage:
            @staticmethod
            def make_docx(document: Document) -> None:
                document.add_section(WD_SECTION.NEW_PAGE)
                document.add_paragraph("Second page body")

        document = Document()
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Last block on first page"

        conversion_processor._append_pdf2docx_page(
            document,
            FakeParsedPage(),
            is_first_page=False,
        )

        body_blocks = [
            child for child in document.element.body if child.tag != qn("w:sectPr")
        ]
        self.assertEqual(
            [child.tag.rsplit("}", 1)[-1] for child in body_blocks],
            ["tbl", "p", "p"],
        )
        compact_break = body_blocks[1]
        self.assertTrue(compact_break.xpath("./w:pPr/w:sectPr"))
        self.assertFalse(compact_break.xpath(".//w:t[normalize-space(.)]"))

        spacing = compact_break.xpath("./w:pPr/w:spacing")
        self.assertEqual(len(spacing), 1)
        self.assertEqual(spacing[0].get(qn("w:before")), "0")
        self.assertEqual(spacing[0].get(qn("w:after")), "0")
        self.assertEqual(spacing[0].get(qn("w:line")), "20")
        self.assertEqual(spacing[0].get(qn("w:lineRule")), "exact")

    def test_pdf2docx_page_append_keeps_boundary_separate_from_an_image(self) -> None:
        from docx.enum.section import WD_SECTION
        from docx.oxml.ns import qn

        class FakeParsedPage:
            @staticmethod
            def make_docx(document: Document) -> None:
                document.add_section(WD_SECTION.NEW_PAGE)
                document.add_paragraph("Next editable page")

        with tempfile.TemporaryDirectory() as folder:
            image_path = Path(folder) / "visual-page.png"
            Image.new("RGB", (320, 220), "white").save(image_path)

            document = Document()
            image_paragraph = document.add_paragraph()
            image_paragraph.add_run().add_picture(str(image_path))

            conversion_processor._append_pdf2docx_page(
                document,
                FakeParsedPage(),
                is_first_page=False,
            )

        body_blocks = [
            child for child in document.element.body if child.tag != qn("w:sectPr")
        ]
        self.assertEqual(
            [child.tag.rsplit("}", 1)[-1] for child in body_blocks],
            ["p", "p", "p"],
        )
        self.assertTrue(body_blocks[0].xpath(".//w:drawing"))
        self.assertFalse(body_blocks[0].xpath("./w:pPr/w:sectPr"))
        self.assertTrue(body_blocks[1].xpath("./w:pPr/w:sectPr"))
        self.assertEqual(document.paragraphs[-1].text, "Next editable page")

    def test_pdf2docx_page_append_preserves_continuous_section_break(self) -> None:
        from docx.enum.section import WD_SECTION
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        class FakeParsedPage:
            continuous_break = None

            def make_docx(self, document: Document) -> None:
                document.add_section(WD_SECTION.NEW_PAGE)
                document.add_paragraph("Second page heading")
                continuous_paragraph = document.add_paragraph()
                continuous_section = OxmlElement("w:sectPr")
                section_type = OxmlElement("w:type")
                section_type.set(qn("w:val"), "continuous")
                continuous_section.append(section_type)
                continuous_paragraph._p.get_or_add_pPr().append(continuous_section)
                self.continuous_break = continuous_paragraph._p
                document.add_paragraph("Second page column text")

        document = Document()
        previous_paragraph = document.add_paragraph("First page body")
        parsed_page = FakeParsedPage()

        conversion_processor._append_pdf2docx_page(
            document,
            parsed_page,
            is_first_page=False,
        )

        self.assertIsNotNone(parsed_page.continuous_break)
        self.assertIs(parsed_page.continuous_break.getparent(), document.element.body)
        continuous_section = parsed_page.continuous_break.xpath(
            "./w:pPr/w:sectPr[w:type/@w:val='continuous']"
        )
        self.assertEqual(len(continuous_section), 1)
        self.assertFalse(parsed_page.continuous_break.xpath("./w:pPr/w:spacing"))
        self.assertTrue(previous_paragraph._p.xpath("./w:pPr/w:sectPr"))

        all_continuous = document.element.body.xpath(
            ".//w:sectPr[w:type/@w:val='continuous']"
        )
        self.assertEqual(len(all_continuous), 1)
        self.assertEqual(
            continuous_section[0].xpath("./w:type")[0].get(qn("w:val")),
            "continuous",
        )

    def test_pdf2docx_page_append_moves_the_first_boundary_carrier_for_mixed_pages(
        self,
    ) -> None:
        from docx.enum.section import WD_SECTION
        from docx.oxml.ns import qn

        class FakeParsedPage:
            heading_element = None
            next_page_section_element = None

            def make_docx(self, document: Document) -> None:
                document.add_section(WD_SECTION.NEW_PAGE)
                heading = document.add_paragraph("Second page single-column heading")
                self.heading_element = heading._p
                document.add_section(WD_SECTION.CONTINUOUS)
                self.next_page_section_element = document.paragraphs[-1]._p
                document.add_paragraph("Second page two-column body")

        document = Document()
        document.add_paragraph("First page body")
        document.add_section(WD_SECTION.CONTINUOUS)
        previous_footer = document.add_paragraph("First page footer")
        parsed_page = FakeParsedPage()

        conversion_processor._append_pdf2docx_page(
            document,
            parsed_page,
            is_first_page=False,
        )

        footer_section = previous_footer._p.xpath("./w:pPr/w:sectPr")
        self.assertEqual(len(footer_section), 1)
        footer_type = footer_section[0].xpath("./w:type")
        self.assertEqual(len(footer_type), 1)
        self.assertEqual(footer_type[0].get(qn("w:val")), "continuous")

        self.assertIs(previous_footer._p.getnext(), parsed_page.heading_element)
        self.assertIsNotNone(parsed_page.next_page_section_element)
        self.assertIs(
            parsed_page.next_page_section_element.getparent(),
            document.element.body,
        )
        next_page_section = parsed_page.next_page_section_element.xpath(
            "./w:pPr/w:sectPr"
        )
        self.assertEqual(len(next_page_section), 1)
        self.assertTrue(
            conversion_processor._section_break_is_next_page(next_page_section[0])
        )

    def test_global_section_cleanup_preserves_bookmarks(self) -> None:
        from docx.enum.section import WD_SECTION
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        document = Document()
        document.add_paragraph("First page")
        document.add_section(WD_SECTION.NEW_PAGE)
        break_paragraph = document.paragraphs[-1]

        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_start.set(qn("w:id"), "7")
        bookmark_start.set(qn("w:name"), "page_boundary")
        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), "7")
        break_paragraph._p.append(bookmark_start)
        break_paragraph._p.append(bookmark_end)
        document.add_paragraph("Second page")

        conversion_processor._normalize_empty_next_page_section_paragraphs(document)

        self.assertIs(break_paragraph._p.getparent(), document.element.body)
        self.assertEqual(len(document.element.body.xpath(".//w:bookmarkStart")), 1)
        self.assertEqual(len(document.element.body.xpath(".//w:bookmarkEnd")), 1)
        spacing = break_paragraph._p.xpath("./w:pPr/w:spacing")
        self.assertEqual(len(spacing), 1)
        self.assertEqual(spacing[0].get(qn("w:line")), "20")
        self.assertEqual(spacing[0].get(qn("w:lineRule")), "exact")

    def test_global_section_cleanup_does_not_compact_a_real_section_end(self) -> None:
        from docx.enum.section import WD_SECTION
        from docx.oxml.ns import qn
        from docx.shared import Pt

        document = Document()
        section_end = document.add_paragraph("Visible final line of source page")
        section_end.paragraph_format.line_spacing = Pt(11)
        document.add_section(WD_SECTION.NEW_PAGE)
        carrier = document.paragraphs[-1]
        section_properties = carrier._p.get_or_add_pPr().sectPr
        section_end._p.get_or_add_pPr().append(section_properties)
        document.element.body.remove(carrier._p)
        document.add_paragraph("Next source page")

        conversion_processor._normalize_empty_next_page_section_paragraphs(document)

        spacing = section_end._p.xpath("./w:pPr/w:spacing")
        self.assertEqual(len(spacing), 1)
        self.assertEqual(spacing[0].get(qn("w:line")), "220")

    def test_global_section_cleanup_preserves_all_substantive_payload_types(
        self,
    ) -> None:
        from docx.enum.section import WD_SECTION
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt

        payload_kinds = (
            "text",
            "field",
            "hyperlink",
            "drawing",
            "object",
            "line_break",
            "page_break",
        )
        for payload_kind in payload_kinds:
            with self.subTest(payload_kind=payload_kind):
                document = Document()
                document.add_paragraph("First source page")
                document.add_section(WD_SECTION.NEW_PAGE)
                carrier = document.paragraphs[-1]
                carrier.paragraph_format.line_spacing = Pt(11)

                if payload_kind == "text":
                    carrier.add_run("Real section-ending text")
                elif payload_kind == "field":
                    field = OxmlElement("w:fldChar")
                    field.set(qn("w:fldCharType"), "begin")
                    carrier.add_run()._r.append(field)
                elif payload_kind == "hyperlink":
                    carrier._p.append(OxmlElement("w:hyperlink"))
                elif payload_kind == "drawing":
                    carrier.add_run()._r.append(OxmlElement("w:drawing"))
                elif payload_kind == "object":
                    carrier.add_run()._r.append(OxmlElement("w:object"))
                elif payload_kind == "line_break":
                    carrier.add_run()._r.append(OxmlElement("w:br"))
                elif payload_kind == "page_break":
                    page_break = OxmlElement("w:br")
                    page_break.set(qn("w:type"), "page")
                    carrier.add_run()._r.append(page_break)
                document.add_paragraph("Next source page")

                conversion_processor._normalize_empty_next_page_section_paragraphs(
                    document
                )

                spacing = carrier._p.xpath("./w:pPr/w:spacing")
                self.assertEqual(len(spacing), 1)
                self.assertEqual(spacing[0].get(qn("w:line")), "220")
                self.assertTrue(carrier._p.xpath("./w:pPr/w:sectPr"))

    def test_global_section_cleanup_compacts_empty_page_break_before_carrier(
        self,
    ) -> None:
        from docx.enum.section import WD_SECTION
        from docx.oxml.ns import qn
        from docx.shared import Pt

        document = Document()
        document.add_paragraph("First source page")
        document.add_section(WD_SECTION.NEW_PAGE)
        carrier = document.paragraphs[-1]
        carrier.paragraph_format.line_spacing = Pt(11)
        carrier.paragraph_format.page_break_before = True
        document.add_paragraph("Next source page")

        conversion_processor._normalize_empty_next_page_section_paragraphs(document)

        spacing = carrier._p.xpath("./w:pPr/w:spacing")
        self.assertEqual(len(spacing), 1)
        self.assertEqual(spacing[0].get(qn("w:line")), "20")
        self.assertEqual(spacing[0].get(qn("w:lineRule")), "exact")
        page_break_before = carrier._p.xpath("./w:pPr/w:pageBreakBefore")
        self.assertEqual(len(page_break_before), 1)
        self.assertEqual(page_break_before[0].get(qn("w:val")), "0")

    def test_global_section_cleanup_preserves_nonempty_page_break_before_paragraph(
        self,
    ) -> None:
        from docx.enum.section import WD_SECTION
        from docx.oxml.ns import qn
        from docx.shared import Pt

        document = Document()
        document.add_paragraph("First source page")
        document.add_section(WD_SECTION.NEW_PAGE)
        carrier = document.paragraphs[-1]
        carrier.add_run("Visible section-ending text")
        carrier.paragraph_format.line_spacing = Pt(11)
        carrier.paragraph_format.page_break_before = True
        document.add_paragraph("Next source page")

        conversion_processor._normalize_empty_next_page_section_paragraphs(document)

        spacing = carrier._p.xpath("./w:pPr/w:spacing")
        self.assertEqual(len(spacing), 1)
        self.assertEqual(spacing[0].get(qn("w:line")), "220")
        page_break_before = carrier._p.xpath("./w:pPr/w:pageBreakBefore")
        self.assertEqual(len(page_break_before), 1)
        self.assertNotEqual(page_break_before[0].get(qn("w:val")), "0")

    def test_footer_detector_ignores_unique_figure_link_and_bottom_drawing(
        self,
    ) -> None:
        from docx.enum.section import WD_SECTION
        from docx.oxml import OxmlElement
        from docx.shared import Pt

        document = Document()
        reference = document.add_paragraph(
            "Fig 9 data: https://doi.org/10.1000/unique-figure-link"
        )
        reference.runs[0].font.size = Pt(8)
        reference.paragraph_format.space_before = Pt(12)
        drawing = document.add_paragraph("Figure 9")
        drawing.runs[0].font.size = Pt(8)
        drawing.paragraph_format.space_before = Pt(8)
        drawing.runs[0]._r.append(OxmlElement("w:drawing"))
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_paragraph("Second source page")

        self.assertEqual(
            conversion_processor._float_pdf2docx_running_footers(document),
            0,
        )
        self.assertFalse(reference._p.xpath("./w:pPr/w:framePr"))
        self.assertFalse(drawing._p.xpath("./w:pPr/w:framePr"))

    def test_global_section_cleanup_keeps_internal_bootstrap_break_independent(
        self,
    ) -> None:
        from docx.enum.section import WD_SECTION
        from docx.oxml.ns import qn

        document = Document()
        title = document.add_paragraph("Paper title and authors")
        document.add_section(WD_SECTION.CONTINUOUS)
        bootstrap_break = document.paragraphs[-1]
        document.add_paragraph("Editable abstract in the same source page")

        self.assertTrue(bootstrap_break._p.xpath("./w:pPr/w:sectPr"))
        self.assertFalse(bootstrap_break._p.xpath("./w:pPr/w:sectPr/w:type"))

        conversion_processor._normalize_empty_next_page_section_paragraphs(document)

        self.assertIs(bootstrap_break._p.getparent(), document.element.body)
        self.assertFalse(title._p.xpath("./w:pPr/w:sectPr"))
        self.assertTrue(bootstrap_break._p.xpath("./w:pPr/w:sectPr"))
        spacing = bootstrap_break._p.xpath("./w:pPr/w:spacing")
        self.assertEqual(spacing[0].get(qn("w:line")), "20")
        self.assertEqual(spacing[0].get(qn("w:lineRule")), "exact")

    def test_formula_detector_ignores_inline_math_prose_and_plain_cm_fonts(
        self,
    ) -> None:
        page_rect = SimpleNamespace(
            x0=0.0,
            y0=0.0,
            x1=600.0,
            y1=800.0,
            width=600.0,
            height=800.0,
        )

        def span(text: str, font: str, bbox: tuple[float, float, float, float]):
            return {"text": text, "font": font, "bbox": bbox}

        page_dict = {
            "blocks": [
                {
                    "type": 0,
                    "lines": [
                        {
                            "bbox": (300, 50, 555, 66),
                            "spans": [
                                span("Here, ", "LMRoman10-Regular", (300, 50, 335, 66)),
                                span("W", "CMBX10", (335, 50, 345, 66)),
                                span("Q", "CMMI10", (345, 50, 355, 66)),
                                span(
                                    " = X, where model represents sequence length",
                                    "CMR10",
                                    (355, 50, 555, 66),
                                ),
                            ],
                        }
                    ],
                },
                {
                    "type": 0,
                    "lines": [
                        {
                            "bbox": (300, 85, 565, 101),
                            "spans": [
                                span(
                                    "acquirecontextualinformationacrossseveralpaths.Let",
                                    "LMRoman10-Regular",
                                    (300, 85, 490, 101),
                                ),
                                span("V", "CMMI10", (490, 85, 500, 101)),
                                span(" = ", "CMR10", (500, 85, 515, 101)),
                                span("X", "CMMI10", (515, 85, 525, 101)),
                                span(
                                    ",beasliceofthevalue",
                                    "LMRoman10-Regular",
                                    (525, 85, 565, 101),
                                ),
                            ],
                        }
                    ],
                },
                {
                    "type": 0,
                    "lines": [
                        {
                            "bbox": (120, 200, 350, 216),
                            "spans": [
                                span("Y", "CMMI10", (120, 200, 130, 216)),
                                span(" = ", "CMR10", (130, 200, 150, 216)),
                                span("∑", "CMEX10", (150, 198, 168, 218)),
                                span("x", "CMMI10", (168, 200, 180, 216)),
                            ],
                        },
                        {
                            "bbox": (354, 200, 374, 216),
                            "spans": [span("(4)", "CMR10", (354, 200, 374, 216))],
                        },
                    ],
                },
            ]
        }

        regions = conversion_processor._detect_pdf_formula_regions(
            0,
            page_dict,
            page_rect,
            dpi=300,
        )

        self.assertFalse(conversion_processor._math_font("CMR10"))
        self.assertFalse(conversion_processor._math_font("CMBX10"))
        self.assertEqual(len(regions), 1)
        self.assertGreater(regions[0].rect[1], 190)
        self.assertLess(regions[0].rect[3], 225)

    def test_numbered_display_formula_is_one_high_dpi_noneditable_region(
        self,
    ) -> None:
        page_rect = SimpleNamespace(
            x0=0.0,
            y0=0.0,
            x1=600.0,
            y1=800.0,
            width=600.0,
            height=800.0,
        )

        def span(text: str, font: str, bbox: tuple[float, float, float, float]):
            return {"text": text, "font": font, "bbox": bbox}

        inline_line = {
            "bbox": (70, 100, 530, 116),
            "spans": [
                span(
                    "The model uses x = y to describe the normal editable output.",
                    "TimesNewRomanPSMT",
                    (70, 100, 530, 116),
                )
            ],
        }
        formula_lines = [
            {
                "bbox": (170, 300, 390, 318),
                "spans": [
                    span("LCIoU", "NimbusRomNo9L-ReguItal", (170, 300, 205, 318)),
                    span(" = ", "CMR10", (205, 300, 225, 318)),
                    span(
                        "arctan width height",
                        "NimbusRomNo9L-Regu",
                        (225, 300, 390, 318),
                    ),
                ],
            },
            {
                "bbox": (350, 293, 363, 304),
                "spans": [span("2", "CMMI10", (350, 293, 363, 304))],
            },
            {
                "bbox": (258, 316, 270, 328),
                "spans": [span("i", "CMMI10", (258, 316, 270, 328))],
            },
            {
                "bbox": (520, 301, 545, 317),
                "spans": [span("(12)", "TimesNewRomanPSMT", (520, 301, 545, 317))],
            },
        ]
        page_dict = {
            "blocks": [
                {"type": 0, "lines": [inline_line]},
                {"type": 0, "lines": formula_lines},
            ]
        }

        regions = conversion_processor._detect_pdf_formula_regions(
            0,
            page_dict,
            page_rect,
            dpi=300,
        )

        self.assertEqual(len(regions), 1)
        self.assertGreaterEqual(regions[0].dpi, 450)
        self.assertEqual(
            conversion_processor._expand_region_to_nearby_labels(
                regions[0],
                page_dict,
                page_rect,
            ).rect,
            regions[0].rect,
        )
        for line in formula_lines:
            self.assertTrue(
                conversion_processor._rect_fully_contains(
                    regions[0].rect,
                    tuple(float(value) for value in line["bbox"]),
                )
            )
        self.assertGreater(regions[0].rect[1], 285)

        editable = "\n".join(
            conversion_processor._extract_editable_pdf_text_blocks(
                page_dict,
                page_rect,
                regions,
            )
        )
        self.assertIn("normal editable output", editable)
        self.assertNotIn("LCIoU", editable)
        self.assertNotIn("arctan", editable)
        self.assertNotIn("(12)", editable)

    def test_symbol_private_use_equation_number_merges_with_formula_body(
        self,
    ) -> None:
        page_rect = SimpleNamespace(
            x0=0.0,
            y0=0.0,
            x1=600.0,
            y1=800.0,
            width=600.0,
            height=800.0,
        )

        formula_bbox = (170.0, 300.0, 405.0, 318.0)
        number_bbox = (520.0, 301.0, 545.0, 317.0)
        page_dict = {
            "blocks": [
                {
                    "type": 0,
                    "lines": [
                        {
                            "bbox": formula_bbox,
                            "spans": [
                                {
                                    "text": "L = x + y",
                                    "font": "CMMI10",
                                    "bbox": formula_bbox,
                                }
                            ],
                        },
                        {
                            "bbox": number_bbox,
                            "spans": [
                                {
                                    "text": "\uf028",
                                    "font": "Symbol",
                                    "bbox": (520.0, 301.0, 528.0, 317.0),
                                },
                                {
                                    "text": "6",
                                    "font": "TimesNewRomanPSMT",
                                    "bbox": (528.0, 301.0, 537.0, 317.0),
                                },
                                {
                                    "text": "\uf029",
                                    "font": "Symbol",
                                    "bbox": (537.0, 301.0, 545.0, 317.0),
                                },
                            ],
                        },
                    ],
                }
            ]
        }

        regions = conversion_processor._detect_pdf_formula_regions(
            0,
            page_dict,
            page_rect,
            dpi=300,
        )

        self.assertEqual(len(regions), 1)
        self.assertIn("编号显示公式", regions[0].reasons)
        self.assertTrue(
            conversion_processor._rect_fully_contains(
                regions[0].rect,
                formula_bbox,
            )
        )
        self.assertTrue(
            conversion_processor._rect_fully_contains(
                regions[0].rect,
                number_bbox,
            )
        )

    def test_numbered_formulas_stay_separate_across_intervening_prose(self) -> None:
        page_rect = SimpleNamespace(
            x0=0.0,
            y0=0.0,
            x1=600.0,
            y1=800.0,
            width=600.0,
            height=800.0,
        )

        def line(
            text: str,
            bbox: tuple[float, float, float, float],
            font: str,
        ) -> dict[str, object]:
            return {
                "bbox": bbox,
                "spans": [{"text": text, "font": font, "bbox": bbox}],
            }

        formula_lines = []
        for number, center_y in ((8, 120.0), (9, 210.0), (10, 300.0)):
            formula_lines.extend(
                [
                    line(
                        f"L{number} = x + y",
                        (170, center_y - 9, 390, center_y + 9),
                        "CMMI10",
                    ),
                    line(
                        f"({number})",
                        (520, center_y - 8, 545, center_y + 8),
                        "TimesNewRomanPSMT",
                    ),
                ]
            )
            if number != 10:
                formula_lines.extend(
                    [
                        line(
                            "The surrounding explanation remains editable and should not bridge equations.",
                            (70, center_y + 32, 500, center_y + 46),
                            "TimesNewRomanPSMT",
                        ),
                        line(
                            "x",
                            (300, center_y + 32, 308, center_y + 46),
                            "CMMI10",
                        ),
                    ]
                )
        formula_lines.extend(
            [
                line(
                    "LCIoU = arctan width height",
                    (170, 360, 390, 378),
                    "NimbusRomNo9L-Regu",
                ),
                line("42", (520, 361, 545, 377), "TimesNewRomanPSMT"),
            ]
        )
        page_dict = {"blocks": [{"type": 0, "lines": formula_lines}]}

        regions = conversion_processor._detect_pdf_formula_regions(
            0,
            page_dict,
            page_rect,
            dpi=300,
        )

        self.assertEqual(len(regions), 3)
        self.assertTrue(all(region.dpi >= 450 for region in regions))
        self.assertLess(regions[0].rect[3], regions[1].rect[1])
        self.assertLess(regions[1].rect[3], regions[2].rect[1])
        self.assertTrue(all(region.rect[3] - region.rect[1] < 30 for region in regions))

    def test_column_hint_keeps_left_column_numbered_formulas_separate_from_right_prose(
        self,
    ) -> None:
        page_rect = SimpleNamespace(
            x0=0.0,
            y0=0.0,
            x1=600.0,
            y1=800.0,
            width=600.0,
            height=800.0,
        )

        def line(
            text: str,
            bbox: tuple[float, float, float, float],
            font: str,
        ) -> dict[str, object]:
            return {
                "bbox": bbox,
                "spans": [{"text": text, "font": font, "bbox": bbox}],
            }

        formula_groups: list[
            tuple[
                list[dict[str, object]],
                dict[str, object],
                dict[str, object],
            ]
        ] = []
        all_lines: list[dict[str, object]] = []
        for number, center_y in ((11, 180.0), (12, 250.0), (13, 320.0)):
            formula_lines = [
                line(
                    f"F{number} = x + y",
                    (70, center_y - 9, 230, center_y + 9),
                    "CMMI10",
                ),
                line("2", (165, center_y - 15, 175, center_y - 4), "CMMI10"),
                line("n", (165, center_y + 5, 175, center_y + 16), "CMMI10"),
            ]
            number_line = line(
                f"({number})",
                (270, center_y - 8, 296, center_y + 8),
                "TimesNewRomanPSMT",
            )
            prose_line = line(
                f"Right column prose beside equation {number} remains fully editable text.",
                (330, center_y - 8, 570, center_y + 8),
                "TimesNewRomanPSMT",
            )
            formula_groups.append((formula_lines, number_line, prose_line))
            all_lines.extend([*formula_lines, number_line, prose_line])

        page_dict = {"blocks": [{"type": 0, "lines": all_lines}]}
        for column_layout in ("double", "mixed"):
            with self.subTest(column_layout=column_layout):
                regions = conversion_processor._detect_pdf_formula_regions(
                    0,
                    page_dict,
                    page_rect,
                    dpi=300,
                    column_layout=column_layout,
                )

                self.assertEqual(len(regions), 3)
                self.assertTrue(all(region.dpi >= 450 for region in regions))
                self.assertLess(regions[0].rect[3], regions[1].rect[1])
                self.assertLess(regions[1].rect[3], regions[2].rect[1])
                for region, (formula_lines, number_line, prose_line) in zip(
                    regions,
                    formula_groups,
                    strict=True,
                ):
                    for formula_line in [*formula_lines, number_line]:
                        self.assertTrue(
                            conversion_processor._rect_fully_contains(
                                region.rect,
                                tuple(float(value) for value in formula_line["bbox"]),
                            )
                        )
                    self.assertEqual(
                        conversion_processor._rect_intersection_ratio(
                            tuple(float(value) for value in prose_line["bbox"]),
                            region.rect,
                        ),
                        0.0,
                    )

                editable = "\n".join(
                    conversion_processor._extract_editable_pdf_text_blocks(
                        page_dict,
                        page_rect,
                        regions,
                    )
                )
                self.assertIn("Right column prose", editable)
                self.assertNotIn("F11 = x + y", editable)
                self.assertNotIn("(13)", editable)

        single_regions = conversion_processor._detect_pdf_formula_regions(
            0,
            page_dict,
            page_rect,
            dpi=300,
            column_layout="single",
        )
        self.assertEqual(len(single_regions), 3)
        for _formula_lines, number_line, _prose_line in formula_groups:
            number_rect = tuple(float(value) for value in number_line["bbox"])
            self.assertFalse(
                any(
                    conversion_processor._rect_fully_contains(region.rect, number_rect)
                    for region in single_regions
                )
            )
        single_editable = "\n".join(
            conversion_processor._extract_editable_pdf_text_blocks(
                page_dict,
                page_rect,
                single_regions,
            )
        )
        self.assertIn("(11)", single_editable)
        self.assertIn("(12)", single_editable)
        self.assertIn("(13)", single_editable)

    def test_hybrid_assessment_passes_column_hint_to_formula_detector(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            source = self._make_pdf(Path(folder) / "columns.pdf")
            with patch.object(
                conversion_processor,
                "_detect_pdf_formula_regions",
                return_value=[],
            ) as detector:
                assessments = conversion_processor._assess_pdf_pages_for_hybrid(
                    source,
                    None,
                    dpi=300,
                    column_layout="double",
                )

        self.assertEqual(len(assessments), 1)
        detector.assert_called_once()
        self.assertEqual(detector.call_args.kwargs["column_layout"], "double")

    def test_numbered_formula_bbox_trims_overlapping_adjacent_body_prose(
        self,
    ) -> None:
        page_rect = SimpleNamespace(
            x0=0.0,
            y0=0.0,
            x1=600.0,
            y1=800.0,
            width=600.0,
            height=800.0,
        )

        def line(
            text: str,
            bbox: tuple[float, float, float, float],
            font: str,
        ) -> dict[str, object]:
            return {
                "bbox": bbox,
                "spans": [{"text": text, "font": font, "bbox": bbox}],
            }

        previous_body = line(
            "Previous editable prose overlaps a malformed formula bbox.",
            (60, 90, 540, 110),
            "TimesNewRomanPSMT",
        )
        formula = line("E = mc", (180, 108, 390, 136), "CMMI10")
        number = line("(8)", (520, 114, 545, 130), "TimesNewRomanPSMT")
        following_body = line(
            "Following editable prose also touches the malformed bbox.",
            (60, 134, 540, 150),
            "TimesNewRomanPSMT",
        )
        page_dict = {
            "blocks": [
                {
                    "type": 0,
                    "lines": [previous_body, formula, number, following_body],
                }
            ]
        }

        detected = conversion_processor._detect_pdf_formula_regions(
            0,
            page_dict,
            page_rect,
            dpi=300,
        )

        self.assertEqual(len(detected), 1)
        self.assertAlmostEqual(detected[0].rect[1], 110.0)
        self.assertAlmostEqual(detected[0].rect[3], 134.0)
        self.assertGreaterEqual(detected[0].dpi, 450)

        closed = conversion_processor._close_hybrid_regions_over_text(
            detected,
            page_dict,
            page_rect,
        )
        constrained = (
            conversion_processor._constrain_numbered_formula_regions_to_body_prose(
                closed,
                page_dict,
                page_rect,
            )
        )
        self.assertAlmostEqual(constrained[0].rect[1], 110.0)
        self.assertAlmostEqual(constrained[0].rect[3], 134.0)

        editable = "\n".join(
            conversion_processor._extract_editable_pdf_text_blocks(
                page_dict,
                page_rect,
                constrained,
            )
        )
        self.assertIn("Previous editable prose", editable)
        self.assertIn("Following editable prose", editable)
        self.assertNotIn("E = mc", editable)
        self.assertNotIn("(8)", editable)

    def test_final_region_closure_contains_suppressed_text_but_skips_labels(
        self,
    ) -> None:
        page_rect = SimpleNamespace(
            x0=0.0,
            y0=0.0,
            x1=600.0,
            y1=800.0,
            width=600.0,
            height=800.0,
        )
        body_line = {
            "bbox": (148, 110, 220, 124),
            "spans": [
                {"text": "partial", "bbox": (148, 110, 180, 124)},
                {"text": " label", "bbox": (170, 110, 190, 124)},
            ],
        }
        caption_line = {
            "bbox": (195, 112, 245, 126),
            "spans": [{"text": "Figure 2. Result", "bbox": (195, 112, 245, 126)}],
        }
        header_line = {
            "bbox": (150, 20, 260, 32),
            "spans": [{"text": "Running header", "bbox": (150, 20, 260, 32)}],
        }
        footer_line = {
            "bbox": (290, 780, 310, 792),
            "spans": [{"text": "17", "bbox": (290, 780, 310, 792)}],
        }
        page_dict = {
            "blocks": [
                {
                    "type": 0,
                    "lines": [header_line, body_line, caption_line, footer_line],
                }
            ]
        }
        initial = conversion_processor._HybridRegion(
            page_index=0,
            rect=(100.0, 100.0, 160.0, 140.0),
            kind="figure",
            dpi=300,
        )

        closed = conversion_processor._close_hybrid_regions_over_text(
            [initial],
            page_dict,
            page_rect,
        )

        self.assertEqual(len(closed), 1)
        self.assertTrue(
            conversion_processor._rect_fully_contains(
                closed[0].rect,
                tuple(float(value) for value in body_line["bbox"]),
            )
        )
        self.assertFalse(
            conversion_processor._rect_fully_contains(
                closed[0].rect,
                tuple(float(value) for value in caption_line["bbox"]),
            )
        )
        editable = "\n".join(
            conversion_processor._extract_editable_pdf_text_blocks(
                page_dict,
                page_rect,
                closed,
            )
        )
        self.assertNotIn("partial label", editable)
        self.assertIn("Figure 2. Result", editable)
        self.assertIn("Running header", editable)
        self.assertIn("17", editable)

    def test_quality_baseline_excludes_a_caption_fully_inside_visual_region(
        self,
    ) -> None:
        page_rect = SimpleNamespace(
            x0=0.0,
            y0=0.0,
            x1=600.0,
            y1=800.0,
            width=600.0,
            height=800.0,
        )
        caption = {
            "bbox": (120.0, 120.0, 480.0, 136.0),
            "spans": [
                {
                    "text": "Table 2. Results already preserved in the image",
                    "bbox": (120.0, 120.0, 480.0, 136.0),
                }
            ],
        }
        body = {
            "bbox": (60.0, 180.0, 540.0, 196.0),
            "spans": [
                {
                    "text": "Editable body remains outside the visual region.",
                    "bbox": (60.0, 180.0, 540.0, 196.0),
                }
            ],
        }
        region = conversion_processor._HybridRegion(
            page_index=0,
            rect=(90.0, 90.0, 510.0, 160.0),
            kind="table",
            dpi=300,
        )

        editable = "\n".join(
            conversion_processor._extract_editable_pdf_text_blocks(
                {"blocks": [{"type": 0, "lines": [caption, body]}]},
                page_rect,
                [region],
            )
        )

        self.assertNotIn("Table 2", editable)
        self.assertIn("Editable body", editable)

    def test_image_detector_does_not_merge_stacked_figures_across_caption(
        self,
    ) -> None:
        page_rect = SimpleNamespace(
            x0=0.0,
            y0=0.0,
            x1=600.0,
            y1=800.0,
            width=600.0,
            height=800.0,
        )
        page_dict = {
            "blocks": [
                {"type": 1, "bbox": (50, 50, 250, 200)},
                {
                    "type": 0,
                    "lines": [
                        {
                            "bbox": (70, 205, 230, 218),
                            "spans": [
                                {
                                    "text": "Figure 1. First result.",
                                    "font": "Times New Roman",
                                    "bbox": (70, 205, 230, 218),
                                }
                            ],
                        }
                    ],
                },
                {"type": 1, "bbox": (50, 230, 250, 380)},
            ]
        }

        regions = conversion_processor._detect_pdf_image_regions(
            0,
            page_dict,
            page_rect,
            dpi=300,
        )

        self.assertEqual(len(regions), 2)
        caption_rect = (70.0, 205.0, 230.0, 218.0)
        self.assertFalse(
            any(
                conversion_processor._rect_center_inside(caption_rect, region.rect)
                for region in regions
            )
        )

    def test_table_detector_merges_segmented_rules_but_keeps_title_editable(
        self,
    ) -> None:
        page_rect = SimpleNamespace(
            x0=0.0,
            y0=0.0,
            x1=600.0,
            y1=800.0,
            width=600.0,
            height=800.0,
        )
        page = SimpleNamespace(
            rect=page_rect,
            find_tables=lambda: SimpleNamespace(tables=[]),
        )
        drawings = [
            {"rect": (x0, y, x1, y), "items": [("l",)]}
            for y in (120.0, 145.0, 210.0)
            for x0, x1 in ((50.0, 210.0), (210.0, 380.0), (380.0, 550.0))
        ]
        page_dict = {
            "blocks": [
                {
                    "type": 0,
                    "bbox": (150, 98, 450, 110),
                    "lines": [
                        {
                            "bbox": (150, 98, 450, 110),
                            "spans": [
                                {
                                    "text": "Table 4 Results of ablation studies",
                                    "font": "Times New Roman",
                                    "bbox": (150, 98, 450, 110),
                                }
                            ],
                        }
                    ],
                },
                {
                    "type": 0,
                    "bbox": (60, 126, 540, 202),
                    "lines": [
                        {
                            "bbox": (60, 126, 540, 140),
                            "spans": [
                                {
                                    "text": "Baseline  Module  mAP",
                                    "font": "Times New Roman",
                                    "bbox": (60, 126, 540, 140),
                                }
                            ],
                        }
                    ],
                },
            ]
        }

        regions = conversion_processor._detect_pdf_table_regions(
            page,
            0,
            page_dict,
            drawings,
            dpi=300,
        )

        self.assertEqual(len(regions), 1)
        title_rect = (150.0, 98.0, 450.0, 110.0)
        row_rect = (60.0, 126.0, 540.0, 140.0)
        self.assertFalse(
            conversion_processor._rect_center_inside(title_rect, regions[0].rect)
        )
        self.assertTrue(
            conversion_processor._rect_center_inside(row_rect, regions[0].rect)
        )

    def test_table_detector_finds_column_table_with_caption_below(self) -> None:
        page_rect = SimpleNamespace(
            x0=0.0,
            y0=0.0,
            x1=612.0,
            y1=792.0,
            width=612.0,
            height=792.0,
        )
        page = SimpleNamespace(
            rect=page_rect,
            find_tables=lambda: SimpleNamespace(tables=[]),
        )
        drawings = [
            {"rect": (319.5, y, 558.0, y), "items": [("l",)]}
            for y in (54.5, 74.2, 146.1, 191.9)
        ]
        page_dict = {
            "blocks": [
                {
                    "type": 0,
                    "bbox": (319.5, 60.7, 558.0, 69.9),
                    "lines": [
                        {
                            "bbox": (319.5, 60.7, 558.0, 69.9),
                            "spans": [
                                {
                                    "text": "Model AUC mIoU Dice Hit",
                                    "font": "Times New Roman",
                                    "bbox": (319.5, 60.7, 558.0, 69.9),
                                }
                            ],
                        }
                    ],
                },
                {
                    "type": 0,
                    "bbox": (319.5, 152.4, 558.0, 187.5),
                    "lines": [
                        {
                            "bbox": (319.5, 152.4, 558.0, 163.0),
                            "spans": [
                                {
                                    "text": "LoGIC-CXR 0.7066 0.4573",
                                    "font": "Times New Roman",
                                    "bbox": (319.5, 152.4, 558.0, 163.0),
                                }
                            ],
                        }
                    ],
                },
                {
                    "type": 0,
                    "bbox": (319.5, 202.5, 558.0, 267.3),
                    "lines": [
                        {
                            "bbox": (319.5, 202.5, 558.0, 214.0),
                            "spans": [
                                {
                                    "text": "Table 1: Classification results",
                                    "font": "Times New Roman",
                                    "bbox": (319.5, 202.5, 558.0, 214.0),
                                }
                            ],
                        }
                    ],
                },
            ]
        }

        regions = conversion_processor._detect_pdf_table_regions(
            page,
            0,
            page_dict,
            drawings,
            dpi=450,
        )

        self.assertEqual(len(regions), 1)
        self.assertLessEqual(regions[0].rect[0], 319.5)
        self.assertGreaterEqual(regions[0].rect[2], 558.0)
        self.assertLess(regions[0].rect[3], 202.5)

    def test_table_detector_does_not_swallow_journal_footer_separator(self) -> None:
        page_rect = SimpleNamespace(
            x0=0.0,
            y0=0.0,
            x1=612.0,
            y1=792.0,
            width=612.0,
            height=792.0,
        )
        page = SimpleNamespace(
            rect=page_rect,
            find_tables=lambda: SimpleNamespace(tables=[]),
        )
        drawings = [
            {"rect": (36.0, y, 576.0, y), "items": [("l",)]}
            for y in (522.0, 550.0, 675.0, 744.0)
        ]

        def line(text: str, bbox: tuple[float, float, float, float]):
            return {"bbox": bbox, "spans": [{"text": text, "bbox": bbox}]}

        page_dict = {
            "blocks": [
                {
                    "type": 0,
                    "lines": [
                        line("Model Precision Recall", (60.0, 528.0, 300.0, 540.0)),
                        line("ST-YOLO 96.6 93.2", (60.0, 660.0, 300.0, 672.0)),
                        line(
                            "PLOS ONE | https://doi.org/10.1371/journal.pone.test",
                            (36.0, 750.0, 260.0, 760.0),
                        ),
                        line("December 12, 2024", (270.0, 750.0, 390.0, 760.0)),
                        line("14 / 19", (545.0, 750.0, 576.0, 760.0)),
                    ],
                }
            ]
        }

        regions = conversion_processor._detect_pdf_table_regions(
            page,
            0,
            page_dict,
            drawings,
            dpi=300,
        )

        self.assertTrue(regions)
        self.assertTrue(all(region.rect[3] < 720.0 for region in regions))

    def test_auto_column_detector_resolves_a_two_column_paper_to_mixed(
        self,
    ) -> None:
        page_rect = SimpleNamespace(
            x0=0.0,
            y0=0.0,
            x1=600.0,
            y1=800.0,
            width=600.0,
            height=800.0,
        )
        lines = []
        for index in range(10):
            y0 = 120.0 + index * 24.0
            lines.extend(
                [
                    {
                        "bbox": (50.0, y0, 285.0, y0 + 10.0),
                        "spans": [
                            {
                                "text": f"Left column scientific body line {index}",
                                "bbox": (50.0, y0, 285.0, y0 + 10.0),
                            }
                        ],
                    },
                    {
                        "bbox": (325.0, y0, 560.0, y0 + 10.0),
                        "spans": [
                            {
                                "text": f"Right column scientific body line {index}",
                                "bbox": (325.0, y0, 560.0, y0 + 10.0),
                            }
                        ],
                    },
                ]
            )
        page_dict = {"blocks": [{"type": 0, "lines": lines}]}
        self.assertTrue(
            conversion_processor._pdf_page_looks_two_column(
                page_dict,
                page_rect,
            )
        )
        assessments = [
            conversion_processor._HybridPageAssessment(
                page_index=0,
                source_text="editable scientific body " * 20,
                editable_source_text="editable scientific body " * 20,
                editable_text_blocks=("editable scientific body",),
                draw_items=0,
                draw_bbox_max_ratio=0.0,
                reasons=[],
                visual_regions=[],
                detected_two_columns=True,
            )
        ]
        self.assertEqual(
            conversion_processor._resolve_hybrid_auto_column_layout(
                "auto",
                assessments,
            ),
            "mixed",
        )

    def test_auto_column_detector_accepts_narrow_central_gutter(self) -> None:
        page_rect = SimpleNamespace(
            x0=0.0,
            y0=0.0,
            x1=612.2,
            y1=792.0,
            width=612.2,
            height=792.0,
        )
        lines = []
        for index in range(10):
            y0 = 120.0 + index * 24.0
            lines.extend(
                [
                    {
                        "bbox": (48.0, y0, 301.6, y0 + 11.0),
                        "spans": [
                            {
                                "text": f"Left column scientific sentence {index}",
                                "bbox": (48.0, y0, 301.6, y0 + 11.0),
                            }
                        ],
                    },
                    {
                        "bbox": (313.2, y0, 565.0, y0 + 11.0),
                        "spans": [
                            {
                                "text": f"Right column scientific sentence {index}",
                                "bbox": (313.2, y0, 565.0, y0 + 11.0),
                            }
                        ],
                    },
                ]
            )

        self.assertTrue(
            conversion_processor._pdf_page_looks_two_column(
                {"blocks": [{"type": 0, "lines": lines}]},
                page_rect,
            )
        )

    def test_flatten_pdf_regions_preserves_outside_text_and_inserts_one_image(
        self,
    ) -> None:
        pymupdf = conversion_processor._require_pymupdf()

        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "region-source.pdf"
            target = root / "region-flattened.pdf"
            region_rect = (70.0, 60.0, 230.0, 150.0)

            source_document = pymupdf.open()
            source_page = source_document.new_page(width=300, height=200)
            source_page.insert_text((20, 30), "OUTSIDE EDITABLE TEXT")
            source_page.insert_text((100, 110), "INSIDE VISUAL LABEL")
            source_page.draw_rect(
                pymupdf.Rect(80, 70, 220, 140),
                color=(0, 0, 0),
                width=1,
            )
            source_page.draw_line(
                pymupdf.Point(20, 105),
                pymupdf.Point(280, 105),
                color=(0, 0, 0),
                width=1,
            )
            source_document.save(str(source))
            source_document.close()

            region = conversion_processor._HybridRegion(
                page_index=0,
                rect=region_rect,
                kind="figure",
                reasons=("unit-test complex figure",),
                dpi=144,
            )
            conversion_processor._flatten_pdf_regions(
                source,
                target,
                [region],
                password=None,
                dpi=144,
            )

            flattened_document = pymupdf.open(str(target))
            try:
                self.assertEqual(flattened_document.page_count, 1)
                flattened_page = flattened_document[0]
                text_layer = flattened_page.get_text("text")
                self.assertIn("OUTSIDE EDITABLE TEXT", text_layer)
                self.assertNotIn("INSIDE VISUAL LABEL", text_layer)

                image_blocks = [
                    block
                    for block in flattened_page.get_text("dict").get("blocks", ())
                    if block.get("type") == 1
                ]
                self.assertEqual(len(image_blocks), 1)
                for actual, expected in zip(image_blocks[0]["bbox"], region_rect):
                    self.assertAlmostEqual(actual, expected, delta=0.2)
                self.assertGreaterEqual(image_blocks[0]["width"], 300)
                self.assertGreaterEqual(image_blocks[0]["height"], 160)

                drawing_items = [
                    item
                    for drawing in flattened_page.get_drawings()
                    for item in drawing.get("items", ())
                ]
                self.assertTrue(
                    any(
                        item[0] == "l"
                        and min(float(item[1].x), float(item[2].x)) <= 20.5
                        and max(float(item[1].x), float(item[2].x)) >= 279.5
                        and abs(float(item[1].y) - 105.0) <= 0.5
                        and abs(float(item[2].y) - 105.0) <= 0.5
                        for item in drawing_items
                    ),
                    "a vector line that only crosses the rasterized region must remain "
                    "intact outside that region",
                )
            finally:
                flattened_document.close()

    def test_flatten_pdf_regions_removes_a_fully_covered_source_image(self) -> None:
        pymupdf = conversion_processor._require_pymupdf()

        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            image_path = root / "source-image.png"
            Image.new("RGB", (360, 240), (30, 120, 220)).save(image_path)
            source = root / "fully-covered-source.pdf"
            target = root / "fully-covered-target.pdf"

            source_document = pymupdf.open()
            page = source_document.new_page(width=300, height=200)
            source_image_rect = pymupdf.Rect(60, 40, 240, 160)
            page.insert_image(source_image_rect, filename=str(image_path))
            source_document.save(str(source))
            source_document.close()

            region_rect = (50.0, 30.0, 250.0, 170.0)
            conversion_processor._flatten_pdf_regions(
                source,
                target,
                [
                    conversion_processor._HybridRegion(
                        page_index=0,
                        rect=region_rect,
                        kind="figure",
                        dpi=144,
                    )
                ],
                dpi=144,
            )

            flattened = pymupdf.open(str(target))
            try:
                image_info = flattened[0].get_image_info(xrefs=True)
                self.assertEqual(len(image_info), 1)
                for actual, expected in zip(image_info[0]["bbox"], region_rect):
                    self.assertAlmostEqual(actual, expected, delta=0.2)
            finally:
                flattened.close()

    def test_flatten_pdf_regions_keeps_partially_covered_image_outside_region(
        self,
    ) -> None:
        pymupdf = conversion_processor._require_pymupdf()

        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            image_path = root / "source-image.png"
            source_image = Image.new("RGB", (360, 240), (30, 120, 220))
            ImageDraw.Draw(source_image).rectangle(
                (0, 0, 179, 239),
                fill=(220, 60, 40),
            )
            source_image.save(image_path)
            source = root / "partially-covered-source.pdf"
            target = root / "partially-covered-target.pdf"

            source_document = pymupdf.open()
            page = source_document.new_page(width=300, height=200)
            source_image_rect = (60.0, 40.0, 240.0, 160.0)
            page.insert_image(
                pymupdf.Rect(source_image_rect),
                filename=str(image_path),
            )
            source_document.save(str(source))
            source_document.close()

            region_rect = (150.0, 30.0, 250.0, 170.0)
            conversion_processor._flatten_pdf_regions(
                source,
                target,
                [
                    conversion_processor._HybridRegion(
                        page_index=0,
                        rect=region_rect,
                        kind="figure",
                        dpi=144,
                    )
                ],
                dpi=144,
            )

            flattened = pymupdf.open(str(target))
            try:
                image_rects = [
                    tuple(float(value) for value in info["bbox"])
                    for info in flattened[0].get_image_info(xrefs=True)
                ]
                self.assertEqual(len(image_rects), 2)
                self.assertTrue(
                    any(
                        all(
                            abs(actual - expected) <= 0.2
                            for actual, expected in zip(rect, source_image_rect)
                        )
                        for rect in image_rects
                    )
                )
                self.assertTrue(
                    any(
                        all(
                            abs(actual - expected) <= 0.2
                            for actual, expected in zip(rect, region_rect)
                        )
                        for rect in image_rects
                    )
                )
                pixmap = flattened[0].get_pixmap(alpha=False)
                outside_pixel = pixmap.pixel(90, 100)
                self.assertGreater(outside_pixel[0], 180)
                self.assertLess(outside_pixel[1], 100)
            finally:
                flattened.close()

    def test_flatten_pdf_regions_mixes_full_and_partial_image_redactions_per_page(
        self,
    ) -> None:
        pymupdf = conversion_processor._require_pymupdf()

        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            fully_covered_path = root / "fully-covered.png"
            partially_covered_path = root / "partially-covered.png"
            Image.new("RGB", (240, 240), (30, 120, 220)).save(fully_covered_path)
            partially_covered = Image.new("RGB", (300, 240), (20, 170, 70))
            ImageDraw.Draw(partially_covered).rectangle(
                (0, 0, 149, 239),
                fill=(220, 60, 40),
            )
            partially_covered.save(partially_covered_path)
            source = root / "mixed-image-redactions-source.pdf"
            target = root / "mixed-image-redactions-target.pdf"

            source_document = pymupdf.open()
            page = source_document.new_page(width=420, height=220)
            fully_covered_image_rect = (30.0, 40.0, 150.0, 160.0)
            partially_covered_image_rect = (220.0, 40.0, 370.0, 160.0)
            page.insert_image(
                pymupdf.Rect(fully_covered_image_rect),
                filename=str(fully_covered_path),
            )
            page.insert_image(
                pymupdf.Rect(partially_covered_image_rect),
                filename=str(partially_covered_path),
            )
            source_document.save(str(source))
            source_document.close()

            # Leave a sub-threshold edge sliver: the image is still considered
            # fully covered by the 98% rule, but PIXELS mode alone would retain
            # the old image placement behind the replacement.
            fully_covered_region = (20.0, 30.0, 149.0, 170.0)
            partially_covered_region = (295.0, 30.0, 380.0, 170.0)
            conversion_processor._flatten_pdf_regions(
                source,
                target,
                [
                    conversion_processor._HybridRegion(
                        page_index=0,
                        rect=fully_covered_region,
                        kind="figure",
                        dpi=144,
                    ),
                    conversion_processor._HybridRegion(
                        page_index=0,
                        rect=partially_covered_region,
                        kind="figure",
                        dpi=144,
                    ),
                ],
                dpi=144,
            )

            flattened = pymupdf.open(str(target))
            try:
                image_rects = [
                    tuple(float(value) for value in info["bbox"])
                    for info in flattened[0].get_image_info(xrefs=True)
                ]
                self.assertEqual(len(image_rects), 3)
                self.assertFalse(
                    any(
                        all(
                            abs(actual - expected) <= 0.2
                            for actual, expected in zip(
                                rect,
                                fully_covered_image_rect,
                            )
                        )
                        for rect in image_rects
                    ),
                    "the fully covered source image must not remain behind its replacement",
                )
                for expected_rect in (
                    partially_covered_image_rect,
                    fully_covered_region,
                    partially_covered_region,
                ):
                    self.assertTrue(
                        any(
                            all(
                                abs(actual - expected) <= 0.2
                                for actual, expected in zip(rect, expected_rect)
                            )
                            for rect in image_rects
                        )
                    )

                pixmap = flattened[0].get_pixmap(alpha=False)
                outside_pixel = pixmap.pixel(250, 100)
                self.assertGreater(outside_pixel[0], 180)
                self.assertLess(outside_pixel[1], 100)
            finally:
                flattened.close()

    def test_redaction_partition_uses_the_visible_clipped_image_bbox(self) -> None:
        page_rect = SimpleNamespace(
            x0=0.0,
            y0=0.0,
            x1=200.0,
            y1=200.0,
            width=200.0,
            height=200.0,
        )

        class ClippedImagePage:
            @staticmethod
            def get_text(mode: str) -> dict[str, object]:
                self.assertEqual(mode, "dict")
                return {
                    "blocks": [
                        {
                            "type": 1,
                            # The source placement is clipped to this visible box.
                            "bbox": (40.0, 50.0, 120.0, 130.0),
                        }
                    ]
                }

            @staticmethod
            def get_image_info(*, xrefs: bool) -> list[dict[str, object]]:
                self.assertTrue(xrefs)
                return [
                    {
                        # The placement matrix extends far beyond the clip path.
                        "bbox": (10.0, 20.0, 180.0, 160.0),
                    }
                ]

        region = conversion_processor._HybridRegion(
            page_index=0,
            rect=(38.0, 48.0, 122.0, 132.0),
            kind="figure",
            dpi=450,
        )

        remove_regions, pixel_regions = (
            conversion_processor._partition_pdf_redaction_regions(
                ClippedImagePage(),
                [region],
                page_rect,
            )
        )

        self.assertEqual(remove_regions, [region])
        self.assertEqual(pixel_regions, [])

    def test_hybrid_layout_detector_flags_full_width_content_before_columns(
        self,
    ) -> None:
        class Section(list):
            pass

        def section(column_count: int, y0: float, y1: float) -> Section:
            value = Section(object() for _ in range(column_count))
            value.bbox = SimpleNamespace(y0=y0, y1=y1)
            return value

        risky_page = SimpleNamespace(
            height=842,
            sections=[
                section(1, 6, 105),
                section(2, 143, 786),
                section(1, 818, 835),
            ],
        )
        regular_page = SimpleNamespace(
            height=842,
            sections=[
                section(1, 6, 30),
                section(2, 57, 774),
                section(1, 782, 835),
            ],
        )

        self.assertEqual(
            conversion_processor._pdf2docx_page_layout_risk_reason(risky_page),
            "全宽内容与多栏正文混排",
        )
        self.assertIsNone(
            conversion_processor._pdf2docx_page_layout_risk_reason(regular_page)
        )

    def test_layout_detector_flags_unequal_short_sidebar_and_main_body(
        self,
    ) -> None:
        class Section(list):
            pass

        sidebar = SimpleNamespace(
            bbox=SimpleNamespace(
                x0=36.0,
                y0=120.0,
                x1=187.0,
                y1=300.0,
                width=151.0,
                height=180.0,
            )
        )
        main_body = SimpleNamespace(
            bbox=SimpleNamespace(
                x0=200.0,
                y0=120.0,
                x1=576.0,
                y1=775.0,
                width=376.0,
                height=655.0,
            )
        )
        body_section = Section([sidebar, main_body])
        body_section.bbox = SimpleNamespace(
            x0=36.0,
            y0=120.0,
            x1=576.0,
            y1=775.0,
            width=540.0,
            height=655.0,
        )
        page = SimpleNamespace(
            width=612.0,
            height=842.0,
            sections=[body_section],
        )

        self.assertIsNotNone(
            conversion_processor._pdf2docx_page_layout_risk_reason(page)
        )

    def test_pdf_asymmetric_sidebar_region_prefers_local_card_background(
        self,
    ) -> None:
        page_rect = SimpleNamespace(
            x0=0.0,
            y0=0.0,
            x1=612.0,
            y1=840.0,
            width=612.0,
            height=840.0,
        )

        def line(bbox: tuple[float, float, float, float], text: str) -> dict:
            return {
                "bbox": bbox,
                "spans": [{"bbox": bbox, "text": text}],
            }

        lines = [
            line((60.0, 60.0, 145.0, 72.0), "Running publication logo"),
        ]
        for index in range(9):
            y0 = 205.0 + index * 17.0
            lines.append(
                line(
                    (62.0, y0, 188.0, y0 + 11.0),
                    f"Sidebar author metadata affiliation item {index}",
                )
            )
        for index in range(18):
            y0 = 190.0 + index * 17.0
            lines.append(
                line(
                    (235.0, y0, 570.0, y0 + 11.0),
                    f"Editable main scientific body sentence number {index} with details",
                )
            )
        page_dict = {"blocks": [{"type": 0, "lines": lines}]}
        region = conversion_processor._pdf_asymmetric_sidebar_region(
            0,
            page_dict,
            [
                {
                    "rect": (54.0, 188.0, 201.0, 520.0),
                    "items": [("re", (54.0, 188.0, 201.0, 520.0), 1)],
                    "fill": (0.9, 0.9, 0.9),
                }
            ],
            page_rect,
            dpi=360,
        )

        self.assertIsNotNone(region)
        assert region is not None
        self.assertEqual(region.kind, "sidebar")
        self.assertAlmostEqual(region.rect[0], 54.0, delta=0.5)
        self.assertAlmostEqual(region.rect[1], 188.0, delta=0.5)
        self.assertAlmostEqual(region.rect[2], 201.0, delta=0.5)
        self.assertAlmostEqual(region.rect[3], 520.0, delta=0.5)
        self.assertGreater(region.rect[1], 100.0)

    def test_sidebar_mask_removes_only_the_parser_copy_region(self) -> None:
        import pymupdf

        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "sidebar.pdf"
            masked = root / "sidebar-masked.pdf"
            canvas = Canvas(str(source), pagesize=(320, 220))
            canvas.drawString(25, 170, "SIDEBAR SECRET")
            canvas.drawString(180, 170, "EDITABLE BODY")
            canvas.save()

            conversion_processor._mask_pdf_regions_for_editable_conversion(
                source,
                masked,
                [
                    conversion_processor._HybridRegion(
                        page_index=0,
                        rect=(10.0, 20.0, 155.0, 90.0),
                        kind="sidebar",
                    )
                ],
                password=None,
            )

            document = pymupdf.open(masked)
            try:
                text = document[0].get_text("text")
            finally:
                document.close()
            self.assertNotIn("SIDEBAR SECRET", text)
            self.assertIn("EDITABLE BODY", text)

    def test_pdf_region_overlay_uses_page_relative_anchor_without_inline_height(
        self,
    ) -> None:
        import pymupdf
        from docx.shared import Pt

        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "anchor-source.pdf")
            output = root / "anchor.docx"
            document = Document()
            document.add_paragraph("Editable body remains in normal flow")
            pdf_document = pymupdf.open(source)
            try:
                restored = conversion_processor._append_page_relative_pdf_regions(
                    document,
                    pdf_document[0],
                    [
                        conversion_processor._HybridRegion(
                            page_index=0,
                            rect=(20.0, 30.0, 120.0, 95.0),
                            kind="sidebar",
                            dpi=240,
                        )
                    ],
                    root,
                    existing_body_ids=set(),
                )
            finally:
                pdf_document.close()
            document.save(output)

            namespaces = {
                "wp": (
                    "http://schemas.openxmlformats.org/drawingml/2006/"
                    "wordprocessingDrawing"
                )
            }
            with ZipFile(output) as archive, archive.open(
                "word/document.xml"
            ) as stream:
                root_element = ElementTree.parse(stream).getroot()
            anchors = root_element.findall(".//wp:anchor", namespaces)
            self.assertEqual(restored, 1)
            self.assertEqual(len(anchors), 1)
            self.assertEqual(
                root_element.findall(".//wp:inline", namespaces),
                [],
            )
            self.assertEqual(
                anchors[0].find("wp:positionH", namespaces).attrib["relativeFrom"],
                "page",
            )
            self.assertEqual(
                anchors[0].find("wp:positionV", namespaces).attrib["relativeFrom"],
                "page",
            )
            self.assertEqual(
                anchors[0].find("wp:positionH/wp:posOffset", namespaces).text,
                str(int(Pt(20.0))),
            )
            self.assertEqual(
                anchors[0].find("wp:positionV/wp:posOffset", namespaces).text,
                str(int(Pt(30.0))),
            )

    def test_dense_two_column_figure_page_uses_visual_compatibility_fallback(
        self,
    ) -> None:
        page_rect = SimpleNamespace(
            x0=0.0,
            y0=0.0,
            x1=600.0,
            y1=840.0,
            width=600.0,
            height=840.0,
        )
        regions = [
            conversion_processor._HybridRegion(
                page_index=0,
                rect=rect,
                kind="figure",
                dpi=450,
            )
            for rect in (
                (45.0, 55.0, 285.0, 305.0),
                (45.0, 330.0, 285.0, 580.0),
                (350.0, 55.0, 555.0, 250.0),
                (350.0, 285.0, 555.0, 480.0),
            )
        ]

        self.assertTrue(
            conversion_processor._hybrid_dense_two_column_figure_risk(
                regions,
                page_rect,
                column_layout="mixed",
            )
        )
        self.assertTrue(
            conversion_processor._hybrid_dense_two_column_figure_risk(
                regions,
                page_rect,
                column_layout="auto",
            )
        )
        self.assertFalse(
            conversion_processor._hybrid_dense_two_column_figure_risk(
                regions,
                page_rect,
                column_layout="single",
            )
        )

    def test_bottom_table_stays_local_unless_it_is_exceptionally_large(self) -> None:
        page_rect = SimpleNamespace(
            x0=0.0,
            y0=0.0,
            x1=612.0,
            y1=792.0,
            width=612.0,
            height=792.0,
        )
        ordinary = conversion_processor._HybridRegion(
            page_index=0,
            rect=(32.0, 518.0, 580.0, 693.0),
            kind="complex",
            dpi=300,
        )
        exceptional = conversion_processor._HybridRegion(
            page_index=0,
            rect=(30.0, 480.0, 582.0, 790.0),
            kind="complex",
            dpi=300,
        )

        self.assertFalse(
            conversion_processor._hybrid_bottom_full_width_visual_overflow_risk(
                [ordinary],
                page_rect,
            )
        )
        self.assertTrue(
            conversion_processor._hybrid_bottom_full_width_visual_overflow_risk(
                [exceptional],
                page_rect,
            )
        )

    def test_intercolumn_merge_detector_returns_the_narrow_local_rect(self) -> None:
        merged = SimpleNamespace(
            lines=[object()],
            text="Left section heading accidentally joined with right body text.",
            bbox=SimpleNamespace(x0=65.0, y0=108.0, x1=548.0, y1=128.0),
        )
        page = SimpleNamespace(
            width=612.0,
            sections=[[SimpleNamespace(blocks=[merged])]],
        )

        rects = conversion_processor._pdf2docx_intercolumn_text_merge_rects(
            page,
            split_x=306.0,
            paired_bands=[(107.0, 250.0)],
        )

        self.assertEqual(rects, [(65.0, 108.0, 548.0, 128.0)])

    def test_hybrid_keeps_full_width_intro_and_two_column_body_editable(
        self,
    ) -> None:
        from docx.oxml.ns import qn

        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "two-column-paper.pdf"
            output = root / "two-column-paper.docx"

            canvas = Canvas(str(source), pagesize=(595, 842))
            canvas.setFont("Helvetica-Bold", 14)
            canvas.drawString(50, 760, "Two Column Layout Validation")
            canvas.setFont("Helvetica", 9)
            for line_number, y_position in enumerate(
                (730, 705, 680, 655),
                start=1,
            ):
                canvas.drawString(
                    50,
                    y_position,
                    f"Abstract overview line {line_number} spans the page with "
                    "reliable scientific context and stable wording.",
                )
            y_position = 610
            for line_number in range(1, 17):
                canvas.drawString(
                    50,
                    y_position,
                    f"Left item {line_number:02d} alpha beta gamma delta.",
                )
                canvas.drawString(
                    335,
                    y_position,
                    f"Right item {line_number:02d} sigma tau omega theta.",
                )
                y_position -= 22
            canvas.save()

            self.assertEqual(
                pdf_to_docx(
                    source,
                    output,
                    mode="hybrid",
                    dpi=144,
                    low_quality_policy="discard",
                    column_layout="mixed",
                ),
                [output],
            )

            document = Document(output)
            self.assertEqual(len(document.inline_shapes), 0)
            self.assertEqual(
                len(
                    document.element.body.xpath(
                        ".//wp:anchor/wp:docPr["
                        "@descr='LayoutLoom fixed-layout background']"
                    )
                ),
                1,
            )
            self.assertGreater(
                len(document.element.body.xpath("./w:p/w:pPr/w:framePr")),
                20,
            )
            editable_text = _extract_docx_text(output)
            self.assertIn("Two Column Layout Validation", editable_text)
            self.assertIn("Left item 01", editable_text)
            self.assertIn("Right item 16", editable_text)

    def test_single_column_hint_flattens_detected_pdf_columns_before_docx_build(
        self,
    ) -> None:
        from docx.oxml.ns import qn

        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "detected-two-columns.pdf"
            output = root / "forced-single-column.docx"

            canvas = Canvas(str(source), pagesize=(595, 842))
            canvas.setFont("Helvetica", 9)
            y_position = 760
            for line_number in range(1, 19):
                canvas.drawString(
                    50,
                    y_position,
                    f"Left column sentence {line_number:02d} alpha beta gamma.",
                )
                canvas.drawString(
                    335,
                    y_position,
                    f"Right column sentence {line_number:02d} sigma tau omega.",
                )
                y_position -= 22
            canvas.save()

            with warnings.catch_warnings():
                warnings.simplefilter("ignore")
                self.assertEqual(
                    pdf_to_docx(
                        source,
                        output,
                        mode="editable",
                        low_quality_policy="keep",
                        column_layout="single",
                    ),
                    [output],
                )

            document = Document(output)
            self.assertFalse(
                any(
                    columns.get(qn("w:num")) == "2"
                    for section in document.sections
                    for columns in section._sectPr.xpath("./w:cols")
                )
            )
            self.assertFalse(
                document.element.body.xpath(".//w:sectPr[w:type/@w:val='nextColumn']")
            )
            editable_text = _extract_docx_text(output)
            self.assertIn("Left column sentence 01", editable_text)
            self.assertIn("Right column sentence 18", editable_text)

    def test_hybrid_page_quality_accepts_small_order_drift_at_high_recall(
        self,
    ) -> None:
        words = (
            "alpha beta gamma delta epsilon zeta eta theta iota kappa lambda mu "
            "nu xi omicron pi rho sigma tau upsilon"
        ).split()
        source = " ".join(words)
        reordered = " ".join([*words[:-2], words[-1], words[-2]])

        reason = conversion_processor._pdf2docx_page_quality_reason(source, reordered)

        self.assertIsNone(reason)

    def test_hybrid_page_quality_rejects_material_column_order_drift(self) -> None:
        words = (
            "alpha beta gamma delta epsilon zeta eta theta iota kappa lambda mu "
            "nu xi omicron pi rho sigma tau upsilon"
        ).split()
        source = " ".join(words)
        reordered = " ".join([*words[:-3], *reversed(words[-3:])])

        reason = conversion_processor._pdf2docx_page_quality_reason(source, reordered)

        self.assertIsNotNone(reason)
        self.assertIn("词序", reason)

    def test_hybrid_character_detector_rejects_controls_and_private_use(self) -> None:
        detector = conversion_processor._has_suspicious_pdf_characters
        self.assertFalse(detector("Normal 中文 text √"))
        self.assertTrue(detector("bad\u008a control"))
        self.assertTrue(detector("formula \ue000 glyph"))

    def test_known_pdf_encoding_repair_requires_ascii_word_context(self) -> None:
        repair = conversion_processor._repair_known_pdf_text_encoding

        self.assertEqual(repair("self\uff27worth"), "self-worth")
        self.assertEqual(repair("students\U001001b3growth"), "students'growth")
        self.assertEqual(repair("well\uff27\nmeaning"), "well-\nmeaning")
        self.assertEqual(repair("students\U001001b3 future"), "students' future")
        self.assertEqual(repair("型号\uff27 与 A \uff27 B"), "型号\uff27 与 A \uff27 B")
        self.assertEqual(repair("marker \U001001b3 value"), "marker \U001001b3 value")

    def test_pdf2docx_known_encoding_repair_crosses_span_boundaries(self) -> None:
        def span(text: str) -> SimpleNamespace:
            return SimpleNamespace(
                text=text,
                chars=[SimpleNamespace(c=character) for character in text],
                size=10,
                bbox=SimpleNamespace(x0=0, x1=20),
            )

        parts = [
            span("self"),
            span("\uff27"),
            span("worth and well"),
            span("\uff27"),
            span("meaning and students"),
            span("\U001001b3"),
            span(" growth 型号"),
            span("\uff27"),
            span(" 中文"),
        ]
        block = SimpleNamespace(
            lines=[
                SimpleNamespace(spans=parts[:4]),
                SimpleNamespace(spans=parts[4:]),
            ]
        )
        column = SimpleNamespace(blocks=[block])
        page = SimpleNamespace(finalized=True, sections=[[column]])

        repaired = conversion_processor._repair_pdf2docx_known_encoding([page])

        self.assertEqual(repaired, 3)
        self.assertEqual(parts[1].text, "-")
        self.assertEqual(parts[3].text, "-")
        self.assertEqual(parts[5].text, "'")
        self.assertEqual(parts[7].text, "\uff27")
        self.assertEqual(parts[1].chars, [])
        self.assertEqual(parts[3].chars, [])
        self.assertEqual(parts[5].chars, [])

    def test_hybrid_encoding_detector_keeps_recoverable_descenders_editable(
        self,
    ) -> None:
        page_rect = SimpleNamespace(
            x0=0.0,
            y0=0.0,
            x1=600.0,
            y1=800.0,
            width=600.0,
            height=800.0,
        )

        def line(
            bbox: tuple[float, float, float, float],
            spans: list[tuple[str, tuple[float, float, float, float]]],
        ) -> dict[str, object]:
            return {
                "bbox": bbox,
                "spans": [
                    {"text": text, "font": "Embedded", "bbox": span_bbox}
                    for text, span_bbox in spans
                ],
            }

        proven_mapping_line = line(
            (60, 70, 210, 88),
            [("editable self\uff27worth text", (60, 70, 210, 88))],
        )
        private_use_line = line(
            (60, 110, 240, 128),
            [("unknown \U001001ba placeholder", (60, 110, 240, 128))],
        )
        recoverable_fragmented_line = line(
            (60, 150, 180, 168),
            [
                ("p", (60, 150, 72, 168)),
                ("ractical ex", (66, 149, 122, 167)),
                ("p", (120, 150, 132, 168)),
                ("erience", (126, 149, 180, 167)),
            ],
        )
        malformed_fragmented_line = line(
            (60, 190, 180, 208),
            [
                ("g", (95, 190, 107, 208)),
                ("arbled", (66, 189, 100, 207)),
                ("y", (94, 190, 106, 208)),
                (" order", (100, 189, 180, 207)),
            ],
        )
        neighboring_line = line(
            (60, 207, 260, 225),
            [("following baseline must stay whole", (60, 207, 260, 225))],
        )
        ordinary_line = line(
            (60, 240, 250, 258),
            [("ordinary editable prose", (60, 240, 250, 258))],
        )
        page_dict = {
            "blocks": [
                {
                    "type": 0,
                    "lines": [
                        proven_mapping_line,
                        private_use_line,
                        recoverable_fragmented_line,
                        malformed_fragmented_line,
                        neighboring_line,
                        ordinary_line,
                    ],
                }
            ]
        }

        regions = conversion_processor._detect_pdf_encoding_regions(
            0,
            page_dict,
            page_rect,
            dpi=300,
        )

        self.assertEqual(len(regions), 2)
        self.assertTrue(all(region.kind == "complex" for region in regions))
        self.assertTrue(all(region.dpi >= 450 for region in regions))
        for region, expected_line in zip(
            regions, [private_use_line, malformed_fragmented_line], strict=True
        ):
            self.assertTrue(
                conversion_processor._rect_fully_contains(
                    region.rect,
                    expected_line["bbox"],
                )
            )
        self.assertFalse(
            conversion_processor._rect_fully_contains(
                regions[1].rect, neighboring_line["bbox"]
            )
        )

        editable = "\n".join(
            conversion_processor._extract_editable_pdf_text_blocks(
                page_dict,
                page_rect,
                regions,
            )
        )
        self.assertIn("editable self-worth text", editable)
        self.assertIn("practical experience", editable)
        self.assertIn("following baseline must stay whole", editable)
        self.assertIn("ordinary editable prose", editable)
        self.assertNotIn("unknown", editable)
        self.assertNotIn("garbled", editable)

    def test_pdf2docx_space_helpers_restore_span_and_cross_line_words(self) -> None:
        def bbox(x0: float, x1: float) -> SimpleNamespace:
            return SimpleNamespace(x0=x0, x1=x1)

        internal_span = SimpleNamespace(
            text="AB",
            chars=[
                SimpleNamespace(c="A", bbox=bbox(0, 4)),
                SimpleNamespace(c="B", bbox=bbox(7, 11)),
            ],
            size=10,
            bbox=bbox(0, 11),
        )
        self.assertEqual(_restore_pdf2docx_span_spaces(internal_span), 1)
        self.assertEqual(internal_span.text, "A B")

        punctuation_span = SimpleNamespace(
            text="A.",
            chars=[
                SimpleNamespace(c="A", bbox=bbox(0, 4)),
                SimpleNamespace(c=".", bbox=bbox(8, 10)),
            ],
            size=10,
            bbox=bbox(0, 10),
        )
        self.assertEqual(_restore_pdf2docx_span_spaces(punctuation_span), 0)
        self.assertEqual(punctuation_span.text, "A.")

        first = SimpleNamespace(text="High", chars=[], size=10, bbox=bbox(0, 20))
        second = SimpleNamespace(text="precision", chars=[], size=10, bbox=bbox(24, 66))
        third = SimpleNamespace(text="layout", chars=[], size=10, bbox=bbox(0, 28))
        first_line = SimpleNamespace(
            spans=[first, second], dir=(1.0, 0.0), line_break=0
        )
        second_line = SimpleNamespace(spans=[third], dir=(1.0, 0.0), line_break=0)
        block = SimpleNamespace(lines=[first_line, second_line])

        self.assertEqual(_restore_pdf2docx_text_block_spaces(block), 2)
        self.assertEqual(first.text, "High ")
        self.assertEqual(second.text, "precision ")
        self.assertEqual(third.text, "layout")

        wide_digits = SimpleNamespace(
            text="１９８０", chars=[], size=10, bbox=bbox(0, 28)
        )
        preceding = SimpleNamespace(text="from", chars=[], size=10, bbox=bbox(0, 20))
        fullwidth_boundary = SimpleNamespace(
            lines=[
                SimpleNamespace(spans=[preceding], dir=(1.0, 0.0), line_break=0),
                SimpleNamespace(spans=[wide_digits], dir=(1.0, 0.0), line_break=0),
            ]
        )
        self.assertEqual(
            _restore_pdf2docx_text_block_spaces(fullwidth_boundary),
            1,
        )
        self.assertEqual(preceding.text, "from ")
        self.assertFalse(
            conversion_processor._pdf2docx_coordinate_space_candidate("19.", "３９%")
        )
        self.assertFalse(
            conversion_processor._pdf2docx_coordinate_space_candidate(
                "．",
                "３",
                previous_context="９",
            )
        )

        decimal_left = SimpleNamespace(text="１９", chars=[], size=10, bbox=bbox(0, 10))
        decimal_point = SimpleNamespace(text="．", chars=[], size=10, bbox=bbox(10, 12))
        decimal_right = SimpleNamespace(
            text="３９", chars=[], size=10, bbox=bbox(0, 10)
        )
        decimal_block = SimpleNamespace(
            lines=[
                SimpleNamespace(
                    spans=[decimal_left, decimal_point],
                    dir=(1.0, 0.0),
                    line_break=0,
                ),
                SimpleNamespace(spans=[decimal_right], dir=(1.0, 0.0), line_break=0),
            ]
        )
        self.assertEqual(_restore_pdf2docx_text_block_spaces(decimal_block), 0)
        self.assertEqual(decimal_point.text, "．")

    def test_pdf2docx_space_helper_uses_visible_glyph_outlines_when_bboxes_overlap(
        self,
    ) -> None:
        def bbox(x0: float, x1: float) -> SimpleNamespace:
            return SimpleNamespace(x0=x0, x1=x1)

        first = SimpleNamespace(c="A", origin=(0.0, 10.0), bbox=bbox(0, 10))
        second = SimpleNamespace(c="B", origin=(5.0, 10.0), bbox=bbox(5, 15))
        span = SimpleNamespace(
            text="AB",
            chars=[first, second],
            font="damaged-font-name",
            size=10.0,
            bbox=bbox(0, 15),
        )
        geometry = {
            ("", 10.0, 0.0, 10.0, None): (0.0, 3.0),
            ("", 10.0, 5.0, 10.0, None): (6.0, 9.0),
        }

        self.assertEqual(
            _restore_pdf2docx_span_spaces(span, glyph_geometry=geometry),
            1,
        )
        self.assertEqual(span.text, "A B")

    def test_pdf2docx_reconstructs_split_prose_and_list_rows(self) -> None:
        from pdf2docx.layout.Column import Column
        from pdf2docx.text.TextBlock import TextBlock

        def line(text: str, x0: float, y0: float, x1: float) -> dict[str, object]:
            return {
                "bbox": (x0, y0, x1, y0 + 12),
                "dir": (1.0, 0.0),
                "spans": [
                    {
                        "bbox": (x0, y0, x1, y0 + 12),
                        "font": "Arial",
                        "size": 10.0,
                        "text": text,
                        "chars": [],
                    }
                ],
            }

        prose_blocks = [
            TextBlock(
                {
                    "lines": [
                        line(
                            "The unusual punctuation metrics split this ordinary English",
                            80,
                            100,
                            480,
                        )
                    ]
                }
            ),
            TextBlock(
                {
                    "lines": [
                        line(
                            "paragraph even though its main baselines remain perfectly",
                            60,
                            120,
                            480,
                        )
                    ]
                }
            ),
            TextBlock({"lines": [line("aligned and fully editable.", 60, 140, 260)]}),
        ]
        prose_column = Column().update_bbox((0, 0, 500, 700))
        prose_column.blocks.reset(prose_blocks)
        prose_page = SimpleNamespace(finalized=True, sections=[[prose_column]])

        self.assertEqual(
            conversion_processor._reconstruct_pdf2docx_prose_blocks([prose_page]),
            2,
        )
        self.assertEqual(len(prose_column.blocks), 1)
        reconstructed = prose_column.blocks[0]
        reconstructed.lines[0].line_break = 1
        reconstructed.lines[-1].line_break = 1
        conversion_processor._normalize_pdf2docx_prose_alignment([prose_page])
        self.assertEqual(
            len(conversion_processor._pdf2docx_block_rows(reconstructed)), 3
        )
        self.assertEqual(str(reconstructed.alignment), "TextAlignment.JUSTIFY")
        self.assertAlmostEqual(reconstructed.first_line_space, 20.0)
        self.assertEqual(
            [line.line_break for line in reconstructed.lines],
            [1, 0, 0],
        )

        option_lines = [
            line("A) First independent answer", 80, 200, 270),
            line("B) Second independent answer", 81, 220, 285),
            line("C) Third independent answer", 79, 240, 270),
            line("D) Fourth independent answer", 80, 260, 280),
        ]
        option_block = TextBlock({"lines": option_lines})
        option_column = Column().update_bbox((0, 0, 500, 700))
        option_column.blocks.reset([option_block])

        self.assertEqual(
            conversion_processor._split_pdf2docx_list_blocks(option_column),
            3,
        )
        self.assertEqual(len(option_column.blocks), 4)
        self.assertEqual(
            [block.text[:2] for block in option_column.blocks],
            ["A)", "B)", "C)", "D)"],
        )

    def test_pdf2docx_splits_short_cjk_structure_from_english_flow(self) -> None:
        from pdf2docx.common.share import TextAlignment
        from pdf2docx.layout.Column import Column
        from pdf2docx.text.TextBlock import TextBlock

        def line(
            text: str,
            x0: float,
            y0: float,
            x1: float,
            *,
            line_break: int = 0,
        ) -> dict[str, object]:
            return {
                "bbox": (x0, y0, x1, y0 + 12),
                "dir": (1.0, 0.0),
                "line_break": line_break,
                "spans": [
                    {
                        "bbox": (x0, y0, x1, y0 + 12),
                        "font": "Arial",
                        "size": 10.0,
                        "text": text,
                        "chars": [],
                    }
                ],
            }

        mixed = TextBlock(
            {
                "alignment": TextAlignment.JUSTIFY.value,
                "lines": [
                    line("２０２５年１２月", 86, 100, 158, line_break=1),
                    line(
                        "１．For this part, you are allowed 30 minutes to write",
                        86,
                        122,
                        480,
                    ),
                    line("with the sentence and complete the essay.", 62, 144, 360),
                ],
            }
        )
        mixed.alignment = TextAlignment.JUSTIFY
        mixed.left_space = 2.0
        mixed.right_space = 3.0
        mixed.first_line_space = 24.0
        mixed.tab_stops = [120.0]
        mixed.lines[0].tab_stop = 2
        mixed.lines[1].tab_stop = 1
        column = Column().update_bbox((62, 0, 482, 700))
        column.blocks.reset([mixed])
        page = SimpleNamespace(finalized=True, sections=[[column]])

        self.assertEqual(
            conversion_processor._reconstruct_pdf2docx_prose_blocks([page]),
            1,
        )
        self.assertEqual(len(column.blocks), 2)
        date, english = column.blocks
        conversion_processor._normalize_pdf2docx_list_prefixes([page])
        conversion_processor._normalize_pdf2docx_prose_alignment([page])

        self.assertEqual(date.text, "２０２５年１２月")
        self.assertEqual(date.alignment, TextAlignment.LEFT)
        self.assertAlmostEqual(date.left_space, 24.0)
        self.assertEqual(date.first_line_space, 0.0)
        self.assertEqual(date.tab_stops, [])
        self.assertTrue(all(line.tab_stop == 0 for line in date.lines))
        self.assertTrue(english.text.startswith("1. For this part"))
        self.assertEqual(english.alignment, TextAlignment.LEFT)
        self.assertAlmostEqual(english.left_space, 0.0)
        self.assertAlmostEqual(english.first_line_space, 24.0)
        self.assertEqual(english.tab_stops, [])
        self.assertTrue(all(line.tab_stop == 0 for line in english.lines))

    def test_pdf2docx_splits_one_embedded_list_marker_but_not_decimal(self) -> None:
        from pdf2docx.layout.Column import Column
        from pdf2docx.text.TextBlock import TextBlock

        def line(text: str, x0: float, y0: float, x1: float) -> dict[str, object]:
            return {
                "bbox": (x0, y0, x1, y0 + 12),
                "dir": (1.0, 0.0),
                "spans": [
                    {
                        "bbox": (x0, y0, x1, y0 + 12),
                        "font": "Arial",
                        "size": 10.0,
                        "text": text,
                        "chars": [],
                    }
                ],
            }

        embedded = TextBlock(
            {
                "lines": [
                    line(
                        "Robots can work without a stop for 24 hours",
                        62,
                        100,
                        470,
                    ),
                    line("(5)尊老是社会和谐与发展的重要基石.", 86, 122, 350),
                ]
            }
        )
        embedded.lines[0].tab_stop = 1
        embedded.lines[1].tab_stop = 2
        column = Column().update_bbox((62, 0, 482, 700))
        column.blocks.reset([embedded])

        self.assertEqual(
            conversion_processor._split_pdf2docx_list_blocks(column),
            1,
        )
        self.assertEqual(len(column.blocks), 2)
        self.assertFalse(bool(getattr(column.blocks[0], "_docuforge_list_item", False)))
        self.assertTrue(column.blocks[1]._docuforge_list_item)
        self.assertTrue(
            all(
                line_item.tab_stop == 0
                for block in column.blocks
                for line_item in block.lines
            )
        )

        decimal = TextBlock(
            {
                "lines": [
                    line("The measured value was", 62, 160, 260),
                    line("5.4 million units in total.", 86, 182, 330),
                ]
            }
        )
        decimal_column = Column().update_bbox((62, 0, 482, 700))
        decimal_column.blocks.reset([decimal])
        self.assertEqual(
            conversion_processor._split_pdf2docx_list_blocks(decimal_column),
            0,
        )
        self.assertEqual(len(decimal_column.blocks), 1)

    def test_pdf2docx_list_item_can_absorb_its_unfinished_continuation(self) -> None:
        from pdf2docx.layout.Column import Column
        from pdf2docx.text.TextBlock import TextBlock

        def block(text: str, x0: float, y0: float, x1: float) -> TextBlock:
            return TextBlock(
                {
                    "lines": [
                        {
                            "bbox": (x0, y0, x1, y0 + 12),
                            "dir": (1.0, 0.0),
                            "spans": [
                                {
                                    "bbox": (x0, y0, x1, y0 + 12),
                                    "font": "Arial",
                                    "size": 10.0,
                                    "text": text,
                                    "chars": [],
                                }
                            ],
                        }
                    ]
                }
            )

        numbered = block(
            "2.... begins with the sentence and teachers can exert a",
            86,
            100,
            480,
        )
        numbered._docuforge_list_item = True
        continuation = block(
            "profound influence on their students academic pursuit.",
            62,
            120,
            390,
        )
        column = Column().update_bbox((62, 0, 482, 700))
        column.blocks.reset([numbered, continuation])
        page = SimpleNamespace(finalized=True, sections=[[column]])

        self.assertEqual(
            conversion_processor._reconstruct_pdf2docx_prose_blocks([page]),
            1,
        )
        self.assertEqual(len(column.blocks), 1)
        self.assertIn("profound influence", column.blocks[0].text)

    def test_pdf2docx_recovers_short_single_line_geometric_center(self) -> None:
        from pdf2docx.common.share import TextAlignment
        from pdf2docx.layout.Column import Column
        from pdf2docx.text.TextBlock import TextBlock

        title = TextBlock(
            {
                "lines": [
                    {
                        "bbox": (170, 100, 330, 114),
                        "dir": (1.0, 0.0),
                        "spans": [
                            {
                                "bbox": (170, 100, 330, 114),
                                "font": "Arial",
                                "size": 12.0,
                                "text": "写作与翻译",
                                "chars": [],
                            }
                        ],
                    }
                ]
            }
        )
        title.alignment = TextAlignment.RIGHT
        title.left_space = 40.0
        title.right_space = 12.0
        title.first_line_space = 28.0
        title.tab_stops = [120.0]
        title.lines[0].tab_stop = 2
        column = Column().update_bbox((0, 0, 500, 700))
        column.blocks.reset([title])
        page = SimpleNamespace(finalized=True, sections=[[column]])

        conversion_processor._normalize_pdf2docx_prose_alignment([page])

        self.assertEqual(title.alignment, TextAlignment.CENTER)
        self.assertEqual(title.left_space, 0.0)
        self.assertEqual(title.right_space, 0.0)
        self.assertEqual(title.first_line_space, 0.0)
        self.assertEqual(title.tab_stops, [])
        self.assertEqual(title.lines[0].tab_stop, 0)

    def test_final_docx_layout_quality_detects_mixed_softbreak_and_tabs(self) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        with tempfile.TemporaryDirectory() as folder:
            path = Path(folder) / "mixed.docx"
            document = Document()
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = paragraph.add_run()
            run.add_tab()
            run.add_tab()
            run.add_text("第二段")
            run.add_break()
            run.add_tab()
            run.add_text("First and foremost, young people should remain confident.")
            document.save(path)

            initial = conversion_processor._pdf2docx_docx_layout_quality_reasons(path)
            self.assertTrue(any("兼容设置" in reason for reason in initial))
            self.assertTrue(any("两端对齐" in reason for reason in initial))
            self.assertTrue(any("前导 Tab" in reason for reason in initial))

            conversion_processor._stabilize_pdf2docx_paragraph_layout(document)
            document.save(path)
            stabilized = conversion_processor._pdf2docx_docx_layout_quality_reasons(
                path
            )
            self.assertFalse(any("兼容设置" in reason for reason in stabilized))
            self.assertTrue(any("两端对齐" in reason for reason in stabilized))
            self.assertTrue(any("前导 Tab" in reason for reason in stabilized))

            safe = Document()
            safe_paragraph = safe.add_paragraph()
            safe_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            safe_run = safe_paragraph.add_run("第二段")
            safe_run.add_break()
            safe_run.add_text(
                "First and foremost, young people should remain confident."
            )
            conversion_processor._stabilize_pdf2docx_paragraph_layout(safe)
            safe.save(path)
            self.assertEqual(
                conversion_processor._pdf2docx_docx_layout_quality_reasons(path),
                (),
            )

    def test_final_render_quality_rejects_extra_physical_page(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "source.pdf")
            rendered = root / "rendered.pdf"
            writer = PdfWriter()
            writer.add_blank_page(width=320, height=220)
            writer.add_blank_page(width=320, height=220)
            with rendered.open("wb") as stream:
                writer.write(stream)

            reason = conversion_processor._pdf2docx_rendered_pdf_quality_reason(
                source,
                rendered,
                expected_pages=1,
                password=None,
            )

            self.assertIn("疑似空白或低内容溢出页", reason)
            self.assertIn("有效字符 0/要求", reason)

    def test_final_render_quality_allows_substantive_page_reflow(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "source.pdf")
            rendered = root / "rendered.pdf"
            canvas = Canvas(str(rendered), pagesize=(320, 220))
            canvas.setFont("Helvetica", 9)
            line = "Substantive editable paragraph content remains on this Word page."
            for page_number in range(2):
                for line_number in range(6):
                    canvas.drawString(
                        18,
                        195 - line_number * 26,
                        f"{page_number}-{line_number} {line}",
                    )
                if page_number == 0:
                    canvas.showPage()
                    canvas.setFont("Helvetica", 9)
            canvas.save()

            reason, warning = (
                conversion_processor._pdf2docx_rendered_pdf_content_quality_result(
                    source,
                    rendered,
                    expected_pages=1,
                    password=None,
                )
            )

            self.assertIsNone(reason)
            self.assertIn("源文件 1 页", warning)
            self.assertIn("Word 渲染 2 页", warning)
            self.assertIn("允许保存", warning)

    def test_final_render_quality_counts_large_image_as_substantive_content(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "source.pdf")
            image_path = root / "figure.png"
            Image.new("RGB", (240, 130), "navy").save(image_path)
            rendered = root / "rendered.pdf"
            canvas = Canvas(str(rendered), pagesize=(320, 220))
            canvas.setFont("Helvetica", 9)
            for line_number in range(6):
                canvas.drawString(
                    18,
                    195 - line_number * 26,
                    "Substantive text content for the first rendered Word page.",
                )
            canvas.showPage()
            canvas.drawImage(str(image_path), 40, 42, width=240, height=130)
            canvas.drawString(40, 28, "Figure 1.")
            canvas.save()

            reason, warning = (
                conversion_processor._pdf2docx_rendered_pdf_content_quality_result(
                    source,
                    rendered,
                    expected_pages=1,
                    password=None,
                )
            )

            self.assertIsNone(reason)
            self.assertIsNotNone(warning)

    def test_final_render_quality_rejects_small_logo_with_only_one_character(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "source.pdf")
            logo_path = root / "logo.png"
            Image.new("RGB", (24, 16), "black").save(logo_path)
            rendered = root / "rendered.pdf"
            canvas = Canvas(str(rendered), pagesize=(320, 220))
            canvas.setFont("Helvetica", 9)
            for line_number in range(6):
                canvas.drawString(
                    18,
                    195 - line_number * 26,
                    "Substantive text content for the first rendered Word page.",
                )
            canvas.showPage()
            canvas.drawImage(str(logo_path), 18, 185, width=24, height=16)
            canvas.drawString(18, 160, "x")
            canvas.save()

            reason, warning = (
                conversion_processor._pdf2docx_rendered_pdf_content_quality_result(
                    source,
                    rendered,
                    expected_pages=1,
                    password=None,
                )
            )

            self.assertIn("疑似空白或低内容溢出页", reason)
            self.assertIn("2（有效字符 1/要求", reason)
            self.assertIsNone(warning)

    def test_final_render_quality_preserves_a_source_blank_page_after_reflow(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "source-with-blank.pdf"
            canvas = Canvas(str(source), pagesize=(320, 220))
            canvas.setFont("Helvetica", 9)
            for line_number in range(6):
                canvas.drawString(
                    18,
                    195 - line_number * 26,
                    "Source first page contains substantive body text.",
                )
            canvas.showPage()
            canvas.showPage()
            canvas.setFont("Helvetica", 9)
            for line_number in range(6):
                canvas.drawString(
                    18,
                    195 - line_number * 26,
                    "Source third page contains substantive body text.",
                )
            canvas.save()

            rendered = root / "rendered-with-blank.pdf"
            canvas = Canvas(str(rendered), pagesize=(320, 220))
            canvas.setFont("Helvetica", 9)
            for page_number in (1, 2):
                for line_number in range(6):
                    canvas.drawString(
                        18,
                        195 - line_number * 26,
                        f"Reflow page {page_number} contains substantive text.",
                    )
                canvas.showPage()
                canvas.setFont("Helvetica", 9)
            canvas.showPage()
            canvas.setFont("Helvetica", 9)
            for line_number in range(6):
                canvas.drawString(
                    18,
                    195 - line_number * 26,
                    "Rendered final page contains substantive body text.",
                )
            canvas.save()

            reason, warning = (
                conversion_processor._pdf2docx_rendered_pdf_content_quality_result(
                    source,
                    rendered,
                    expected_pages=3,
                    password=None,
                )
            )

            self.assertIsNone(reason)
            self.assertIn("源第 2 页→Word 第 3 页", warning)

    def test_final_render_quality_allows_a_source_backed_short_fragment(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "source.pdf"
            source_text = "This source paragraph contains a short authentic continuation fragment."
            canvas = Canvas(str(source), pagesize=(320, 220))
            canvas.setFont("Helvetica", 9)
            canvas.drawString(18, 180, source_text)
            canvas.save()

            rendered = root / "rendered.pdf"
            canvas = Canvas(str(rendered), pagesize=(320, 220))
            canvas.setFont("Helvetica", 9)
            for line_number in range(6):
                canvas.drawString(
                    18,
                    195 - line_number * 26,
                    "Substantive reconstructed Word body content.",
                )
            canvas.showPage()
            canvas.setFont("Helvetica", 9)
            canvas.drawString(18, 180, "authentic continuation fragment")
            canvas.save()

            reason, warning = (
                conversion_processor._pdf2docx_rendered_pdf_content_quality_result(
                    source,
                    rendered,
                    expected_pages=1,
                    password=None,
                )
            )

            self.assertIsNone(reason)
            self.assertIn("有效重排片段", warning)

    def test_final_render_quality_reports_an_isolated_chinese_page_counter(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "source.pdf")
            rendered = root / "rendered.pdf"
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
            canvas = Canvas(str(rendered), pagesize=(320, 220))
            canvas.setFont("Helvetica", 9)
            for line_number in range(6):
                canvas.drawString(
                    18,
                    195 - line_number * 26,
                    "Substantive reconstructed Word body content.",
                )
            canvas.showPage()
            canvas.setFont("STSong-Light", 9)
            canvas.drawString(125, 190, "第1 5页共53页")
            canvas.save()

            reason, warning = (
                conversion_processor._pdf2docx_rendered_pdf_content_quality_result(
                    source,
                    rendered,
                    expected_pages=1,
                    password=None,
                )
            )

            self.assertIn("页码或页脚独占页面", reason)
            self.assertIn("第1 5页共53页", reason)
            self.assertIsNone(warning)

    def test_final_render_quality_reports_multiple_counters_on_one_page(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "source.pdf")
            rendered = root / "rendered.pdf"
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
            canvas = Canvas(str(rendered), pagesize=(320, 220))
            canvas.setFont("Helvetica", 9)
            for line_number in range(6):
                canvas.drawString(
                    18,
                    195 - line_number * 26,
                    "Substantive reconstructed Word body content.",
                )
            canvas.setFont("STSong-Light", 9)
            canvas.drawString(35, 12, "第24页共53页")
            canvas.drawString(205, 12, "第26页共53页")
            canvas.save()

            conflicts = conversion_processor._pdf2docx_rendered_page_counter_conflicts(
                rendered
            )
            self.assertEqual(
                conflicts,
                (
                    conversion_processor._RenderedPageCounterConflict(
                        page_number=1,
                        counters=((24, 53), (26, 53)),
                    ),
                ),
            )
            reason, warning = (
                conversion_processor._pdf2docx_rendered_pdf_content_quality_result(
                    source,
                    rendered,
                    expected_pages=1,
                    password=None,
                )
            )
            self.assertIn("同一页面存在多个源页码", reason)
            self.assertIn("第24页共53页", reason)
            self.assertIn("第26页共53页", reason)
            self.assertIsNone(warning)

    def test_rendered_page_content_counts_detected_table_area(self) -> None:
        page_rect = SimpleNamespace(
            x0=0.0,
            y0=0.0,
            x1=600.0,
            y1=800.0,
            width=600.0,
            height=800.0,
        )
        page = SimpleNamespace(
            rect=page_rect,
            get_text=lambda _kind: {"blocks": []},
            get_drawings=lambda: [
                {"rect": (60.0, 120.0 + index, 540.0, 121.0 + index)}
                for index in range(4)
            ],
            find_tables=lambda: SimpleNamespace(
                tables=[SimpleNamespace(bbox=(60.0, 120.0, 540.0, 520.0))]
            ),
        )

        assessment = conversion_processor._rendered_pdf_page_content_assessment(page)

        self.assertGreaterEqual(assessment.visual_ratio, 0.28)
        self.assertTrue(assessment.substantive)

    def test_rendered_page_content_ignores_large_unfilled_page_border(self) -> None:
        page_rect = SimpleNamespace(
            x0=0.0,
            y0=0.0,
            x1=600.0,
            y1=800.0,
            width=600.0,
            height=800.0,
        )
        page = SimpleNamespace(
            rect=page_rect,
            get_text=lambda _kind: {"blocks": []},
            get_drawings=lambda: [
                {
                    "rect": (30.0, 30.0, 570.0, 770.0),
                    "fill": None,
                    "items": [("re", (30.0, 30.0, 570.0, 770.0), 1)],
                }
            ],
            find_tables=lambda: SimpleNamespace(tables=[]),
        )

        assessment = conversion_processor._rendered_pdf_page_content_assessment(page)

        self.assertEqual(assessment.visual_ratio, 0.0)
        self.assertFalse(assessment.substantive)

    def test_cross_page_short_english_fragment_is_merged_conservatively(
        self,
    ) -> None:
        from docx.enum.section import WD_SECTION
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt

        document = Document()
        previous = document.add_paragraph(
            "It is high time that the subject became a focus among"
        )
        previous.runs[0].font.name = "Times New Roman"
        previous.runs[0].font.size = Pt(10.5)
        document.add_paragraph("22")
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_paragraph()
        logo = document.add_paragraph()
        logo.add_run()._r.append(OxmlElement("w:drawing"))
        following = document.add_paragraph("youngsters.")
        following.runs[0].font.name = "Times New Roman"
        following.runs[0].font.size = Pt(10.5)

        self.assertEqual(
            conversion_processor._merge_cross_page_english_orphan_fragments(document),
            1,
        )
        self.assertEqual(
            previous.text,
            "It is high time that the subject became a focus among youngsters.",
        )
        self.assertIsNone(following._p.getparent())
        self.assertEqual(len(document.element.body.xpath(".//w:sectPr")), 2)
        self.assertEqual(
            conversion_processor._merge_cross_page_english_orphan_fragments(document),
            0,
        )
        section_breaks = document.element.body.xpath("./w:p/w:pPr/w:sectPr")
        self.assertEqual(len(section_breaks), 1)
        self.assertTrue(
            conversion_processor._section_break_is_next_page(section_breaks[0])
        )
        self.assertEqual(section_breaks[0].tag, qn("w:sectPr"))

    def test_cross_page_fragment_merge_ignores_headings_long_text_and_tables(
        self,
    ) -> None:
        from docx.enum.section import WD_SECTION

        for following_text, add_table in (
            ("Youngsters.", False),
            ("young people with a substantially longer independent paragraph.", False),
            ("youngsters.", True),
        ):
            with self.subTest(following_text=following_text, add_table=add_table):
                document = Document()
                previous = document.add_paragraph(
                    "It is high time that the subject became a focus among"
                )
                document.add_section(WD_SECTION.NEW_PAGE)
                if add_table:
                    document.add_table(rows=1, cols=1).cell(0, 0).text = "Table"
                following = document.add_paragraph(following_text)

                self.assertEqual(
                    conversion_processor._merge_cross_page_english_orphan_fragments(
                        document
                    ),
                    0,
                )
                self.assertEqual(
                    previous.text,
                    "It is high time that the subject became a focus among",
                )
                self.assertIsNotNone(following._p.getparent())

    def test_flatten_pdf_regions_preserves_external_tall_boundary_glyphs(self) -> None:
        import pymupdf

        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "source.pdf"
            target = root / "flattened.pdf"
            document = pymupdf.open()
            page = document.new_page(width=300, height=200)
            page.insert_text((30, 70), "2. PRECEDING MARKER", fontsize=12)
            page.insert_text((30, 84), "TARGET VISUAL LINE", fontsize=12)
            document.save(source)
            document.close()

            source_document = pymupdf.open(source)
            source_page = source_document[0]
            lines = [
                line
                for block in (source_page.get_text("dict") or {}).get("blocks", ())
                if int(block.get("type", 0) or 0) == 0
                for line in block.get("lines", ())
            ]
            target_line = next(
                line
                for line in lines
                if "TARGET VISUAL LINE"
                in "".join(str(span.get("text", "")) for span in line.get("spans", ()))
            )
            target_bbox = tuple(float(value) for value in target_line["bbox"])
            region_rect = (
                target_bbox[0] - 1.0,
                target_bbox[1] - 0.5,
                target_bbox[2] + 1.0,
                target_bbox[3] + 0.5,
            )
            source_document.close()

            conversion_processor._flatten_pdf_regions(
                source,
                target,
                [
                    conversion_processor._HybridRegion(
                        page_index=0,
                        rect=region_rect,
                        kind="complex",
                        reasons=(conversion_processor._OVERLAPPING_PDF_GLYPH_REASON,),
                        dpi=300,
                    )
                ],
                dpi=300,
            )

            flattened = pymupdf.open(target)
            flattened_page = flattened[0]
            text = flattened_page.get_text()
            self.assertIn("2. PRECEDING MARKER", text)
            self.assertNotIn("TARGET VISUAL LINE", text)
            image_blocks = [
                block
                for block in (flattened_page.get_text("dict") or {}).get("blocks", ())
                if int(block.get("type", 0) or 0) == 1
            ]
            self.assertEqual(len(image_blocks), 1)
            image_bbox = tuple(float(value) for value in image_blocks[0]["bbox"])
            for actual, expected in zip(image_bbox, region_rect):
                self.assertAlmostEqual(actual, expected, places=2)
            flattened.close()

    def test_pdf2docx_normalizes_list_prefix_and_preserves_short_centered_title(
        self,
    ) -> None:
        from pdf2docx.common.share import TextAlignment
        from pdf2docx.layout.Column import Column
        from pdf2docx.text.TextBlock import TextBlock

        def block(text: str, x0: float, x1: float, alignment: TextAlignment) -> object:
            value = TextBlock(
                {
                    "alignment": alignment.value,
                    "lines": [
                        {
                            "bbox": (x0, 100, x1, 112),
                            "dir": (1.0, 0.0),
                            "spans": [
                                {
                                    "bbox": (x0, 100, x1, 112),
                                    "font": "Arial",
                                    "size": 10.0,
                                    "text": text,
                                    "chars": [],
                                }
                            ],
                        }
                    ],
                }
            )
            value.alignment = alignment
            return value

        numbered = block(
            "５５．What should the reader do?", 60, 480, TextAlignment.CENTER
        )
        title = block("Introduction", 170, 330, TextAlignment.CENTER)
        title.lines[0].line_break = 1
        column = Column().update_bbox((0, 0, 500, 700))
        column.blocks.reset([numbered, title])
        page = SimpleNamespace(finalized=True, sections=[[column]])

        conversion_processor._normalize_pdf2docx_list_prefixes([page])
        conversion_processor._normalize_pdf2docx_prose_alignment([page])

        self.assertEqual(numbered.text, "55. What should the reader do?")
        self.assertEqual(numbered.alignment, TextAlignment.LEFT)
        self.assertEqual(title.alignment, TextAlignment.CENTER)
        self.assertEqual(title.lines[0].line_break, 0)

    def test_pdf2docx_uses_half_width_ascii_only_in_english_flow(self) -> None:
        from pdf2docx.layout.Column import Column
        from pdf2docx.text.TextBlock import TextBlock

        def block(text: str, y0: float) -> object:
            return TextBlock(
                {
                    "lines": [
                        {
                            "bbox": (60, y0, 480, y0 + 12),
                            "dir": (1.0, 0.0),
                            "spans": [
                                {
                                    "bbox": (60, y0, 480, y0 + 12),
                                    "font": "Arial",
                                    "size": 10.0,
                                    "text": text,
                                    "chars": [],
                                }
                            ],
                        }
                    ]
                }
            )

        english = block(
            "（２０２５年１２月）Many people value reliable friendships．",
            100,
        )
        chinese = block("２０２５年１２月课程安排", 130)
        short_english = block("Ａ）Yes．", 160)
        column = Column().update_bbox((0, 0, 500, 700))
        column.blocks.reset([english, chinese, short_english])
        page = SimpleNamespace(finalized=True, sections=[[column]])

        self.assertGreater(
            conversion_processor._normalize_pdf2docx_english_widths([page]),
            0,
        )
        self.assertEqual(
            english.text,
            "(2025年12月)Many people value reliable friendships.",
        )
        self.assertEqual(chinese.text, "２０２５年１２月课程安排")
        self.assertEqual(short_english.text, "A)Yes.")

    def test_pdf2docx_hyphen_join_requires_a_proven_source_word(self) -> None:
        def span(text: str) -> SimpleNamespace:
            return SimpleNamespace(
                text=text,
                chars=[],
                size=10,
                bbox=SimpleNamespace(x0=0, x1=20),
            )

        broken_left = span("prod-")
        broken_right = span("ucts.")
        self.assertTrue(
            _join_pdf2docx_line_break_hyphen(broken_left, broken_right, {"products"})
        )
        self.assertEqual(broken_left.text, "prod")
        self.assertEqual(broken_right.text, "ucts.")

        compound_left = span("scale-")
        compound_right = span("aware")
        self.assertFalse(
            _join_pdf2docx_line_break_hyphen(
                compound_left, compound_right, {"scale", "aware"}
            )
        )
        self.assertEqual(compound_left.text, "scale-")

        uppercase_left = span("fea-")
        uppercase_right = span("Tures")
        self.assertFalse(
            _join_pdf2docx_line_break_hyphen(
                uppercase_left, uppercase_right, {"features"}
            )
        )

        integrated_left = span("dimension-")
        integrated_right = span("ality")
        block = SimpleNamespace(
            lines=[
                SimpleNamespace(spans=[integrated_left], dir=(1.0, 0.0), line_break=0),
                SimpleNamespace(spans=[integrated_right], dir=(1.0, 0.0), line_break=0),
            ]
        )
        self.assertEqual(
            _restore_pdf2docx_text_block_spaces(block, source_words={"dimensionality"}),
            1,
        )
        self.assertEqual(integrated_left.text, "dimension")
        self.assertEqual(integrated_right.text, "ality")

    def test_editable_quality_helpers_detect_joined_english_words(self) -> None:
        source = "High precision editable Word document"
        damaged = "Highprecision editable Word document"
        self.assertEqual(_english_word_multiset_recall(source, source), 1.0)
        self.assertEqual(_adjacent_english_word_coverage(source, source), 1.0)
        self.assertLess(_english_word_multiset_recall(source, damaged), 1.0)
        self.assertLess(_adjacent_english_word_coverage(source, damaged), 1.0)
        self.assertFalse(
            conversion_processor._pdf_english_boundary_recovery_pass(
                source,
                damaged,
            )
        )

        lost_boundaries = (
            "Manyseefriendshipsasacomfortblanketashouldertocryon "
            "areliablesoultoconfidein"
        )
        restored = (
            "Many see friendships as a comfort blanket a shoulder to cry on "
            "a reliable soul to confide in"
        )
        self.assertTrue(
            conversion_processor._pdf_english_boundary_recovery_pass(
                lost_boundaries,
                restored,
            )
        )
        self.assertTrue(
            conversion_processor._pdf_english_boundary_recovery_pass(
                lost_boundaries,
                restored,
                character_coverage=0.94,
            )
        )
        self.assertIsNone(
            conversion_processor._pdf2docx_page_quality_reason(
                lost_boundaries,
                restored,
            )
        )

    def test_quality_tokenization_normalizes_only_proven_line_break_words(self) -> None:
        source = "products improve industrial prod-\nucts"
        output = "products improve industrial products"
        normalized = _normalize_english_line_break_hyphens(source)
        self.assertNotIn("prod-\nucts", normalized)
        self.assertEqual(normalized.count("products"), 2)
        self.assertEqual(_english_word_multiset_recall(source, output), 1.0)
        self.assertEqual(_adjacent_english_word_coverage(source, output), 1.0)

        compound = "real work continues in real-\ntime"
        normalized_compound = _normalize_english_line_break_hyphens(compound)
        self.assertIn("real-\ntime", normalized_compound)
        self.assertNotIn("realtime", normalized_compound)

    def test_extract_docx_text_preserves_paragraphs_and_table_cells(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            path = Path(folder) / "structured.docx"
            document = Document()
            document.add_paragraph("Paragraph one")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Cell A"
            table.cell(0, 1).text = "Cell B"
            document.add_paragraph("Paragraph two")
            document.save(path)

            text = _extract_docx_text(path)
            self.assertIn("Paragraph one\n", text)
            self.assertIn("Cell A\tCell B", text)
            self.assertIn("\nParagraph two", text)

    def test_pdf2docx_table_grid_normalizer_restores_asymmetric_two_columns(
        self,
    ) -> None:
        from docx.oxml.ns import qn

        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.autofit = False
        self._set_merged_row_widths(table, 0, [(0, 2, 9360)])
        self._set_merged_row_widths(
            table,
            1,
            [(0, 1, 720), (1, 1, 8640)],
        )
        self._set_table_grid_widths(table, [4680, 4680])

        self.assertEqual(
            conversion_processor._normalize_pdf2docx_table_grids(document),
            1,
        )
        self.assertEqual(
            [
                int(column.get(qn("w:w")))
                for column in table._tbl.tblGrid.findall(qn("w:gridCol"))
            ],
            [720, 8640],
        )
        table_width = table._tbl.tblPr.find(qn("w:tblW"))
        self.assertIsNotNone(table_width)
        assert table_width is not None
        self.assertEqual(table_width.get(qn("w:type")), "dxa")
        self.assertEqual(table_width.get(qn("w:w")), "9360")
        self.assertEqual(
            [
                int(cell.tcPr.find(qn("w:tcW")).get(qn("w:w")))
                for cell in table.rows[1]._tr.findall(qn("w:tc"))
            ],
            [720, 8640],
        )

    def test_pdf2docx_table_grid_normalizer_solves_nine_column_merged_form(
        self,
    ) -> None:
        from docx.oxml.ns import qn

        document = Document()
        table = document.add_table(rows=4, cols=9)
        table.autofit = False
        self._set_merged_row_widths(
            table,
            0,
            [
                (0, 1, 900),
                (1, 2, 1764),
                (3, 1, 950),
                (4, 1, 1248),
                (5, 2, 1440),
                (7, 2, 3058),
            ],
        )
        self._set_merged_row_widths(
            table,
            1,
            [
                (0, 1, 900),
                (1, 2, 1764),
                (3, 1, 950),
                (4, 2, 2328),
                (6, 2, 1440),
                (8, 1, 1978),
            ],
        )
        self._set_merged_row_widths(
            table,
            2,
            [
                (0, 2, 1442),
                (2, 3, 3420),
                (5, 2, 1440),
                (7, 2, 3058),
            ],
        )
        self._set_merged_row_widths(table, 3, [(0, 9, 9360)])
        self._set_table_grid_widths(table, [1040] * 9)

        self.assertEqual(
            conversion_processor._normalize_pdf2docx_table_grids(document),
            1,
        )
        self.assertEqual(
            [
                int(column.get(qn("w:w")))
                for column in table._tbl.tblGrid.findall(qn("w:gridCol"))
            ],
            [900, 542, 1222, 950, 1248, 1080, 360, 1080, 1978],
        )

    def test_pdf2docx_table_grid_normalizer_leaves_correct_grid_unchanged(
        self,
    ) -> None:
        from docx.oxml.ns import qn

        document = Document()
        table = document.add_table(rows=1, cols=2)
        table.autofit = False
        self._set_merged_row_widths(
            table,
            0,
            [(0, 1, 720), (1, 1, 8640)],
        )
        self._set_table_grid_widths(table, [720, 8640])
        table_width = table._tbl.tblPr.find(qn("w:tblW"))
        assert table_width is not None
        table_width.set(qn("w:type"), "dxa")
        table_width.set(qn("w:w"), "9360")
        xml_before = table._tbl.xml

        self.assertEqual(
            conversion_processor._normalize_pdf2docx_table_grids(document),
            0,
        )
        self.assertEqual(table._tbl.xml, xml_before)

    def test_pdf2docx_table_grid_normalizer_leaves_underdetermined_grid_unchanged(
        self,
    ) -> None:
        document = Document()
        table = document.add_table(rows=1, cols=3)
        table.autofit = False
        self._set_merged_row_widths(table, 0, [(0, 3, 9000)])
        self._set_table_grid_widths(table, [1000, 2000, 3000])
        xml_before = table._tbl.xml

        self.assertEqual(
            conversion_processor._normalize_pdf2docx_table_grids(document),
            0,
        )
        self.assertEqual(table._tbl.xml, xml_before)

    def test_pdf2docx_table_grid_normalizer_leaves_autofit_grid_unchanged(
        self,
    ) -> None:
        document = Document()
        table = document.add_table(rows=1, cols=2)
        table.autofit = True
        self._set_merged_row_widths(
            table,
            0,
            [(0, 1, 720), (1, 1, 8640)],
        )
        self._set_table_grid_widths(table, [4680, 4680])
        xml_before = table._tbl.xml

        self.assertEqual(
            conversion_processor._normalize_pdf2docx_table_grids(document),
            0,
        )
        self.assertEqual(table._tbl.xml, xml_before)

    def test_pdf2docx_table_grid_normalizer_leaves_inconsistent_constraints_unchanged(
        self,
    ) -> None:
        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.autofit = False
        self._set_merged_row_widths(
            table,
            0,
            [(0, 1, 720), (1, 1, 8640)],
        )
        self._set_merged_row_widths(
            table,
            1,
            [(0, 1, 1200), (1, 1, 8160)],
        )
        self._set_table_grid_widths(table, [4680, 4680])
        xml_before = table._tbl.xml

        self.assertEqual(
            conversion_processor._normalize_pdf2docx_table_grids(document),
            0,
        )
        self.assertEqual(table._tbl.xml, xml_before)

    def test_editable_pdf_to_docx_preserves_fonts_tables_images_and_page_size(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "复杂排版.pdf"
            image_path = root / "插图.png"
            image = Image.new("RGB", (120, 60), "white")
            draw = ImageDraw.Draw(image)
            draw.rectangle((2, 2, 117, 57), outline="blue", width=4)
            draw.text((20, 20), "IMAGE", fill="red")
            image.save(image_path)

            canvas = Canvas(str(source), pagesize=(595, 842))
            canvas.setFont("Helvetica-Bold", 18)
            canvas.drawString(50, 790, "Editable Layout Test")
            canvas.setFont("Helvetica", 11)
            canvas.drawString(50, 760, "Left column paragraph one.")
            canvas.drawString(320, 760, "Right column paragraph two.")
            for x in (50, 200, 350):
                canvas.line(x, 700, x, 640)
            for y in (700, 670, 640):
                canvas.line(50, y, 350, y)
            canvas.drawString(60, 680, "Cell A")
            canvas.drawString(210, 680, "Cell B")
            canvas.drawString(60, 650, "Cell C")
            canvas.drawString(210, 650, "Cell D")
            canvas.drawImage(str(image_path), 50, 540, width=120, height=60)
            canvas.save()

            output = root / "复杂排版.docx"
            self.assertEqual(pdf_to_docx(source, output, mode="editable"), [output])
            document = Document(output)
            all_text = "".join(
                node.text or "" for node in document.element.body.xpath(".//w:t")
            )
            for expected in (
                "Editable Layout Test",
                "Left column paragraph one.",
                "Right column paragraph two.",
                "Cell A",
                "Cell D",
            ):
                self.assertIn(expected, all_text)
            self.assertGreaterEqual(len(document.tables), 1)
            self.assertGreaterEqual(len(document.inline_shapes), 1)
            title_runs = [
                run
                for paragraph in document.paragraphs
                for run in paragraph.runs
                if "Editable Layout Test" in run.text
            ]
            self.assertEqual(len(title_runs), 1)
            self.assertTrue(title_runs[0].bold)
            self.assertAlmostEqual(title_runs[0].font.size.pt, 18.0, places=1)
            self.assertAlmostEqual(document.sections[0].page_width.pt, 595, delta=1)
            self.assertAlmostEqual(document.sections[0].page_height.pt, 842, delta=1)

    def test_editable_pdf_to_docx_preserves_chinese_content(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "contract.pdf"
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
            canvas = Canvas(str(source), pagesize=(595, 842))
            canvas.setFont("STSong-Light", 16)
            canvas.drawString(50, 790, "中文合同标题")
            canvas.setFont("STSong-Light", 11)
            canvas.drawString(50, 755, "甲方：测试公司")
            canvas.drawString(50, 730, "乙方：示例客户")
            canvas.save()

            output = root / "contract.docx"
            self.assertEqual(pdf_to_docx(source, output, mode="editable"), [output])
            document = Document(output)
            all_text = "".join(
                node.text or "" for node in document.element.body.xpath(".//w:t")
            )
            self.assertIn("中文合同标题", all_text)
            self.assertIn("甲方：测试公司", all_text)
            self.assertIn("乙方：示例客户", all_text)

    def test_pdf_word_layout_profile_detects_dense_two_column_paper(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "paper.pdf"
            canvas = Canvas(str(source), pagesize=(595, 842))
            for page_number in range(2):
                canvas.setFont("Helvetica", 9)
                for line_number in range(24):
                    y = 790 - line_number * 29
                    canvas.drawString(
                        45,
                        y,
                        f"Left scientific paragraph {page_number}-{line_number} text.",
                    )
                    canvas.drawString(
                        315,
                        y,
                        f"Right scientific paragraph {page_number}-{line_number} text.",
                    )
                canvas.showPage()
            canvas.save()

            profile = conversion_processor._pdf_word_layout_profile(
                source,
                None,
                column_layout="auto",
            )

            self.assertEqual(profile.page_count, 2)
            self.assertEqual(profile.two_column_pages, 2)
            self.assertTrue(profile.fixed_layout_recommended)
            self.assertIn("双栏/混合分栏论文", profile.reasons)

    def test_fixed_layout_editable_path_preserves_artwork_and_editable_frames(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "designed.pdf"
            canvas = Canvas(str(source), pagesize=(420, 595))
            canvas.setFillColorRGB(0.18, 0.42, 0.58)
            canvas.rect(20, 515, 380, 55, stroke=0, fill=1)
            canvas.setFillColorRGB(1, 1, 1)
            canvas.setFont("Helvetica-Bold", 22)
            canvas.drawCentredString(210, 535, "Designed Resume")
            canvas.setFillColorRGB(0, 0, 0)
            canvas.setFont("Helvetica", 11)
            canvas.drawString(35, 470, "Editable left profile information")
            canvas.drawString(230, 470, "Editable right contact information")
            canvas.drawString(35, 430, "Editable experience and education content")
            canvas.save()

            output = root / "designed.docx"
            with self.assertWarnsRegex(UserWarning, "固定坐标可编辑布局"):
                self.assertEqual(
                    pdf_to_docx(
                        source,
                        output,
                        mode="editable",
                        column_layout="double",
                    ),
                    [output],
                )

            document = Document(output)
            backgrounds = document.element.body.xpath(
                ".//wp:anchor/wp:docPr[@descr='LayoutLoom fixed-layout background']"
            )
            self.assertEqual(len(backgrounds), 1)
            self.assertGreaterEqual(
                len(document.element.body.xpath("./w:p/w:pPr/w:framePr")),
                4,
            )
            output_text = _extract_docx_text(output)
            self.assertIn("Designed Resume", output_text)
            self.assertIn("Editable left profile information", output_text)
            self.assertIn("Editable right contact information", output_text)

    def test_editable_pdf_to_docx_stops_when_content_validation_fails(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "合同.pdf")
            output = root / "合同.docx"
            with patch(
                "docuforge.processors.conversion._extract_docx_text", return_value=""
            ), self.assertRaisesRegex(ValidationError, "文本完整度自动校验未通过"):
                pdf_to_docx(source, output, mode="editable")
            self.assertFalse(output.exists())

    def test_editable_pdf_to_docx_can_keep_a_low_quality_result(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "合同.pdf")
            output = root / "合同.docx"
            with patch(
                "docuforge.processors.conversion._extract_docx_text", return_value=""
            ), self.assertWarnsRegex(UserWarning, "已按用户选择保留"):
                self.assertEqual(
                    pdf_to_docx(
                        source,
                        output,
                        mode="editable",
                        low_quality_policy="keep",
                    ),
                    [output],
                )
            self.assertTrue(output.is_file())

    def test_editable_keep_policy_preserves_a_wps_pagination_failure(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "合同.pdf")
            output = root / "合同.docx"
            with patch.object(
                conversion_processor,
                "_pdf2docx_wps_render_quality_result",
                return_value=("检测到页码或页脚独占页面：2", None),
            ), self.assertWarnsRegex(UserWarning, "已按用户选择保留"):
                self.assertEqual(
                    pdf_to_docx(
                        source,
                        output,
                        mode="editable",
                        low_quality_policy="keep",
                    ),
                    [output],
                )
            self.assertTrue(output.is_file())

    def test_editable_discard_policy_rejects_a_wps_pagination_failure(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "合同.pdf")
            output = root / "合同.docx"
            with patch.object(
                conversion_processor,
                "_pdf2docx_wps_render_quality_result",
                return_value=("检测到页码或页脚独占页面：2", None),
            ), self.assertRaisesRegex(ValidationError, "分页复检未通过"):
                pdf_to_docx(
                    source,
                    output,
                    mode="editable",
                    low_quality_policy="discard",
                )
            self.assertFalse(output.exists())

    def test_hybrid_keep_policy_preserves_a_wps_pagination_failure(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "合同.pdf")
            output = root / "合同.docx"
            with patch.object(
                conversion_processor,
                "_pdf2docx_wps_render_quality_result",
                return_value=("检测到页码或页脚独占页面：2", None),
            ), self.assertWarnsRegex(UserWarning, "已按用户选择保留"):
                self.assertEqual(
                    pdf_to_docx(
                        source,
                        output,
                        mode="hybrid",
                        low_quality_policy="keep",
                        dpi=110,
                    ),
                    [output],
                )
            self.assertTrue(output.is_file())

    def test_hybrid_discard_policy_rejects_a_wps_pagination_failure(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "合同.pdf")
            output = root / "合同.docx"
            with patch.object(
                conversion_processor,
                "_pdf2docx_wps_render_quality_result",
                return_value=("检测到页码或页脚独占页面：2", None),
            ), self.assertRaisesRegex(ValidationError, "二次检测未通过"):
                pdf_to_docx(
                    source,
                    output,
                    mode="hybrid",
                    low_quality_policy="discard",
                    dpi=110,
                )
            self.assertFalse(output.exists())

    def test_hybrid_keep_policy_preserves_a_final_layout_failure(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "合同.pdf")
            output = root / "合同.docx"

            def final_only_layout_reason(path: Path) -> tuple[str, ...]:
                return () if "probe-page-" in str(path) else ("最终结构异常",)

            with patch.object(
                conversion_processor,
                "_pdf2docx_docx_layout_quality_reasons",
                side_effect=final_only_layout_reason,
            ), self.assertWarnsRegex(UserWarning, "已按用户选择保留"):
                self.assertEqual(
                    pdf_to_docx(
                        source,
                        output,
                        mode="hybrid",
                        low_quality_policy="keep",
                        dpi=110,
                    ),
                    [output],
                )
            self.assertTrue(output.is_file())

    def test_hybrid_discard_policy_rejects_a_final_layout_failure(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "合同.pdf")
            output = root / "合同.docx"

            def final_only_layout_reason(path: Path) -> tuple[str, ...]:
                return () if "probe-page-" in str(path) else ("最终结构异常",)

            with patch.object(
                conversion_processor,
                "_pdf2docx_docx_layout_quality_reasons",
                side_effect=final_only_layout_reason,
            ), self.assertRaisesRegex(ValidationError, "排版结构二次校验未通过"):
                pdf_to_docx(
                    source,
                    output,
                    mode="hybrid",
                    low_quality_policy="discard",
                    dpi=110,
                )
            self.assertFalse(output.exists())

    def test_editable_warning_as_error_does_not_delete_committed_output(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "合同.pdf")
            output = root / "合同.docx"

            with patch(
                "docuforge.processors.conversion._extract_docx_text", return_value=""
            ), warnings.catch_warnings():
                warnings.simplefilter("error", UserWarning)
                with self.assertRaisesRegex(UserWarning, "已按用户选择保留"):
                    pdf_to_docx(
                        source,
                        output,
                        mode="editable",
                        low_quality_policy="keep",
                    )

            self.assertTrue(output.is_file())
            self.assertGreater(output.stat().st_size, 0)

    def test_editable_pdf_to_docx_warns_for_mixed_image_only_pages(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            image_path = root / "scan.png"
            Image.new("RGB", (300, 200), "white").save(image_path)
            source = root / "mixed.pdf"
            canvas = Canvas(str(source), pagesize=(300, 200))
            canvas.drawString(30, 150, "Editable first page")
            canvas.showPage()
            canvas.drawImage(str(image_path), 0, 0, width=300, height=200)
            canvas.save()

            output = root / "mixed.docx"
            with self.assertWarnsRegex(UserWarning, "第 2 页没有文字层"):
                self.assertEqual(pdf_to_docx(source, output, mode="editable"), [output])
            self.assertTrue(output.is_file())

    def test_editable_pdf_to_docx_rejects_a_scan_without_a_text_layer(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            image_path = root / "scan.png"
            Image.new("RGB", (500, 300), "white").save(image_path)
            source = root / "扫描件.pdf"
            canvas = Canvas(str(source), pagesize=(500, 300))
            canvas.drawImage(str(image_path), 0, 0, width=500, height=300)
            canvas.save()

            with self.assertRaisesRegex(ValidationError, "没有可编辑文字层"):
                pdf_to_docx(
                    source,
                    root / "扫描件.docx",
                    mode="editable",
                    low_quality_policy="keep",
                )
            visual = root / "扫描件_原样.docx"
            self.assertEqual(
                pdf_to_docx(source, visual, mode="visual", dpi=110), [visual]
            )
            self.assertEqual(len(Document(visual).inline_shapes), 1)

    def test_pdf_to_docx_rejects_an_unknown_content_mode(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "合同.pdf")
            with self.assertRaisesRegex(
                ValidationError, "editable.*hybrid.*office_native.*visual"
            ):
                pdf_to_docx(source, root / "合同.docx", mode="unknown")

    def test_pdf_to_docx_rejects_an_unknown_engine(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "合同.pdf")
            output = root / "合同.docx"
            with self.assertRaisesRegex(
                ValidationError, "auto.*layoutloom.*microsoft_office"
            ):
                pdf_to_docx(source, output, engine="imaginary")
            self.assertFalse(output.exists())

    def test_pdf_to_docx_rejects_forced_office_in_legacy_hybrid_mode(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "合同.pdf")
            output = root / "hybrid.docx"
            with self.assertRaisesRegex(ValidationError, "Microsoft Word 原生转换"):
                pdf_to_docx(
                    source,
                    output,
                    mode="hybrid",
                    engine="microsoft_office",
                )
            self.assertFalse(output.exists())

    def test_pdf_to_docx_office_native_rejects_layoutloom_engine(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "论文.pdf")
            output = root / "Office原生.docx"
            with self.assertRaisesRegex(ValidationError, "必须调用桌面版 Microsoft Word"):
                pdf_to_docx(
                    source,
                    output,
                    mode="office_native",
                    engine="layoutloom",
                )
            self.assertFalse(output.exists())

    def test_office_native_reports_missing_microsoft_word_without_fallback(self) -> None:
        with tempfile.TemporaryDirectory() as folder, patch(
            "docuforge.processors.office.detect_office_engines",
            return_value={
                "microsoft_word": SimpleNamespace(
                    available=False,
                    reason="未检测到桌面版 Microsoft Word",
                )
            },
        ):
            root = Path(folder)
            source = root / "论文.pdf"
            source.touch()
            with self.assertRaisesRegex(MissingEngineError, "未检测到桌面版 Microsoft Word"):
                conversion_processor._convert_pdf_with_microsoft_reflow_candidate(
                    source,
                    root / "论文.docx",
                    password=None,
                    expected_pages=1,
                    low_quality_policy="discard",
                    progress_start=0.0,
                    progress_span=1.0,
                )

    def test_forced_office_keep_policy_does_not_claim_failed_quality_passed(
        self,
    ) -> None:
        failed_quality = conversion_processor._PdfWordCandidateQuality(
            0.40, 0.35, 0.20, (), None, None, 0, 10, 0, False
        )

        def build_office(
            _source: Path, candidate: Path, **_kwargs: object
        ) -> object:
            candidate.write_bytes(b"office")
            return failed_quality

        with tempfile.TemporaryDirectory() as folder, patch.object(
            conversion_processor,
            "_inspect_pdf_text_layers",
            return_value=(["source text"], []),
        ), patch.object(
            conversion_processor,
            "_convert_pdf_with_microsoft_reflow_candidate",
            side_effect=build_office,
        ):
            root = Path(folder)
            source = root / "source.pdf"
            source.write_bytes(b"pdf")
            target = root / "target.docx"
            with self.assertWarnsRegex(UserWarning, "未完全通过"):
                conversion_processor._execute_pdf_to_docx(
                    source,
                    target,
                    password=None,
                    mode="office_native",
                    dpi=300,
                    low_quality_policy="keep",
                    hybrid_force_visual_pages="",
                    column_layout="auto",
                    engine="microsoft_office",
                )
            self.assertEqual(target.read_bytes(), b"office")

    def test_pdf_to_docx_rejects_an_unknown_column_layout(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "columns.pdf")
            output = root / "columns.docx"
            with self.assertRaisesRegex(
                ValidationError,
                "auto.*single.*double.*mixed",
            ):
                pdf_to_docx(source, output, column_layout="three")
            self.assertFalse(output.exists())

    def test_pdf_to_docx_rejects_an_unknown_low_quality_policy(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "合同.pdf")
            output = root / "合同.docx"
            with self.assertRaisesRegex(ValidationError, "discard 或 keep"):
                pdf_to_docx(source, output, low_quality_policy="ask")
            self.assertFalse(output.exists())

    def test_visual_pdf_to_docx_preserves_rotated_pages_with_compact_boundaries(
        self,
    ) -> None:
        from docx.oxml.ns import qn

        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            original = root / "two-pages.pdf"
            canvas = Canvas(str(original), pagesize=(320, 220))
            canvas.drawString(30, 170, "Page one")
            canvas.showPage()
            canvas.drawString(30, 170, "Page two")
            canvas.save()

            rotated = root / "rotated.pdf"
            reader = PdfReader(original)
            writer = PdfWriter()
            writer.add_page(reader.pages[0])
            writer.add_page(reader.pages[1].rotate(90))
            with rotated.open("wb") as stream:
                writer.write(stream)
            reader.close()

            output = root / "visual.docx"
            self.assertEqual(
                pdf_to_docx(rotated, output, mode="visual", dpi=110), [output]
            )
            document = Document(output)
            self.assertEqual(len(document.inline_shapes), 2)
            self.assertEqual(len(document.paragraphs), 3)
            boundary = document.paragraphs[1]._p
            self.assertFalse(boundary.xpath(".//w:drawing|.//w:t[normalize-space(.)]"))
            self.assertTrue(boundary.xpath("./w:pPr/w:sectPr"))
            spacing = boundary.xpath("./w:pPr/w:spacing")
            self.assertEqual(len(spacing), 1)
            self.assertEqual(spacing[0].get(qn("w:line")), "20")
            self.assertEqual(spacing[0].get(qn("w:lineRule")), "exact")
            self.assertEqual(len(document.sections), 2)
            self.assertAlmostEqual(
                document.sections[0].page_width.inches, 320 / 72, places=2
            )
            self.assertAlmostEqual(
                document.sections[0].page_height.inches, 220 / 72, places=2
            )
            self.assertAlmostEqual(
                document.sections[1].page_width.inches, 220 / 72, places=2
            )
            self.assertAlmostEqual(
                document.sections[1].page_height.inches, 320 / 72, places=2
            )

    def test_unicode_note_and_visual_signature(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_pdf(root / "待签署.pdf")
            note_path = root / "批注.pdf"
            add_pdf_note(source, note_path, page=1, text="请复核中文条款")
            annotations = PdfReader(note_path).pages[0].get("/Annots")
            self.assertIsNotNone(annotations)
            self.assertEqual(len(annotations), 1)

            markup_path = root / "高亮.pdf"
            add_pdf_markup(
                source,
                markup_path,
                page=1,
                kind="highlight",
                x=25,
                y=155,
                width=180,
                height=24,
                comment="重点条款",
            )
            markup = PdfReader(markup_path).pages[0]["/Annots"][0].get_object()
            self.assertEqual(markup["/Subtype"], "/Highlight")
            self.assertEqual(markup["/Contents"], "重点条款")

            signature = root / "签名.png"
            image = Image.new("RGBA", (240, 90), (255, 255, 255, 0))
            draw = ImageDraw.Draw(image)
            draw.line((10, 70, 225, 20), fill=(20, 40, 140, 255), width=5)
            image.save(signature)
            signed_path = root / "已签署.pdf"
            self.assertEqual(
                add_visual_signature(
                    source, signature, signed_path, pages="1", x=30, y=30
                ),
                [signed_path],
            )
            self.assertEqual(len(PdfReader(signed_path).pages), 1)

    def test_fill_standard_acroform(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "表单.pdf"
            canvas = Canvas(str(source), pagesize=(320, 220))
            canvas.drawString(30, 170, "Name")
            canvas.acroForm.textfield(name="name", x=80, y=155, width=150, height=24)
            canvas.save()
            output = root / "已填写.pdf"
            self.assertEqual(
                fill_pdf_form(source, output, {"name": "Zhang San"}), [output]
            )
            fields = PdfReader(output).get_fields()
            self.assertIsNotNone(fields)
            self.assertEqual(fields["name"].get("/V"), "Zhang San")

    def test_remove_background_uses_rembg_session_and_writes_png(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "portrait.jpg"
            Image.new("RGB", (4, 3), "blue").save(source)
            session = object()
            calls: list[tuple[str, tuple[int, int], object]] = []

            def fake_remove(image: Image.Image, *, session: object) -> Image.Image:
                calls.append((image.mode, image.size, session))
                return Image.new("RGBA", image.size, (0, 0, 255, 128))

            rembg = ModuleType("rembg")
            rembg.new_session = Mock(return_value=session)
            rembg.remove = Mock(side_effect=fake_remove)
            output_dir = root / "output"
            with patch.dict(sys.modules, {"rembg": rembg}):
                outputs = remove_background([source], output_dir)

            expected = output_dir / "portrait_透明背景.png"
            self.assertEqual(outputs, [expected])
            self.assertTrue(expected.is_file())
            self.assertEqual(calls, [("RGBA", (4, 3), session)])
            rembg.new_session.assert_called_once_with()
            rembg.remove.assert_called_once()
            with Image.open(expected) as result:
                self.assertEqual(result.mode, "RGBA")
                self.assertEqual(result.getpixel((0, 0)), (0, 0, 255, 128))

    def test_svg_to_images_validates_parameters_and_calls_cairosvg(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            output_dir = root / "output"
            source = root / "icon.svg"
            svg_bytes = (
                b'<svg xmlns="http://www.w3.org/2000/svg" width="8" height="6">'
                b'<rect width="8" height="6" fill="red"/></svg>'
            )
            source.write_bytes(svg_bytes)

            def fake_svg2png(**kwargs: object) -> None:
                Image.new("RGBA", (8, 6), "red").save(str(kwargs["write_to"]), "PNG")

            cairosvg = ModuleType("cairosvg")
            cairosvg.svg2png = Mock(side_effect=fake_svg2png)
            with patch.dict(sys.modules, {"cairosvg": cairosvg}):
                with self.assertRaisesRegex(ValidationError, "目标格式"):
                    svg_to_images([], output_dir, target_format="gif")
                for invalid_scale in (0, -1, 21, float("inf"), float("nan")):
                    with self.subTest(scale=invalid_scale):
                        with self.assertRaisesRegex(ValidationError, "渲染倍率"):
                            svg_to_images([], output_dir, scale=invalid_scale)

                outputs = svg_to_images(
                    [source], output_dir, target_format=".PNG", scale=2.5
                )

            expected = output_dir / "icon.png"
            self.assertEqual(outputs, [expected])
            self.assertTrue(expected.is_file())
            cairosvg.svg2png.assert_called_once()
            call = cairosvg.svg2png.call_args.kwargs
            self.assertEqual(call["bytestring"], svg_bytes)
            self.assertEqual(call["scale"], 2.5)
            self.assertFalse(call["unsafe"])
            with Image.open(expected) as result:
                self.assertEqual(result.size, (8, 6))


if __name__ == "__main__":
    unittest.main()
