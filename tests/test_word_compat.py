from __future__ import annotations

import io
import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from docuforge.models import ValidationError
from docuforge.processors import word_compat, word_region


class WordRegionCompatibilityTests(unittest.TestCase):
    @staticmethod
    def _png_bytes(width: int = 80, height: int = 24) -> bytes:
        from PIL import Image, ImageDraw

        image = Image.new("RGB", (width, height), "white")
        ImageDraw.Draw(image).rectangle((8, 7, width - 8, height - 7), fill="black")
        payload = io.BytesIO()
        image.save(payload, format="PNG")
        return payload.getvalue()

    def test_region_markers_accept_current_and_legacy_brand_names(self) -> None:
        for brand in ("LayoutLoom", "DocuForge"):
            page_match = word_region._PAGE_NAME.match(
                f"{brand} region page 7 text 3"
            )
            visual_match = word_region._VISUAL_NAME.match(
                f"{brand} region page 7 visual 2 formula"
            )
            self.assertIsNotNone(page_match)
            self.assertIsNotNone(visual_match)
            self.assertEqual(page_match.group(1), "7")
            self.assertEqual(visual_match.group(1), "7")

    def test_integrity_check_allows_tab_and_coordinate_token_reassembly(self) -> None:
        source = (
            "The editable table values remain complete.\n"
            + "\n".join("0" for _index in range(80))
        )
        output = "The editable table values remain complete.\t" + ("0" * 80)

        metrics = word_region._validate_text_integrity(source, output)

        self.assertLess(metrics.token_recall, 0.985)
        self.assertEqual(metrics.character_recall, 1.0)
        self.assertEqual(metrics.character_precision, 1.0)
        self.assertEqual(metrics.sequence_coverage, 1.0)

    def test_integrity_check_still_rejects_real_character_loss(self) -> None:
        with self.assertRaisesRegex(ValidationError, "文字完整度检查未通过"):
            word_region._validate_text_integrity(
                "Critical editable content 1234567890",
                "Critical editable content 12345",
            )

    def test_page_hint_generators_are_materialized_once(self) -> None:
        mapping = {
            2: (
                hint
                for hint in (
                    ("inline_math", 10, 20, 30, 40),
                    ("text_visual", 50, 60, 70, 80),
                )
            )
        }

        materialized = word_region._materialize_page_hints(mapping)

        self.assertEqual(
            materialized,
            {
                2: (
                    ("inline_math", 10, 20, 30, 40),
                    ("text_visual", 50, 60, 70, 80),
                )
            },
        )

    def test_empty_formula_page_entry_still_runs_conservative_formula_pass(self) -> None:
        source_page = SimpleNamespace(
            index=0,
            width=2000,
            height=3000,
            frames=[],
            background=None,
        )
        with (
            patch.object(word_region, "_source_pages", return_value=[source_page]),
            patch.object(word_region, "_merge_frame_rows", return_value=[]),
            patch.object(word_region, "_refine_visual_regions", return_value=[]),
            patch.object(
                word_region,
                "_formula_line_replacements",
                return_value=([], []),
            ) as formula_pass,
        ):
            word_region._plan_pages(
                Document(),
                [self._png_bytes()],
                preserve_editable_text=False,
                formula_hints_by_page={0: ()},
            )

        formula_pass.assert_called_once()

    def test_typed_visual_page_still_runs_residual_formula_pass(self) -> None:
        source_page = SimpleNamespace(
            index=0,
            width=2000,
            height=3000,
            frames=[],
            background=None,
        )
        with (
            patch.object(word_region, "_source_pages", return_value=[source_page]),
            patch.object(word_region, "_merge_frame_rows", return_value=[]),
            patch.object(word_region, "_refine_visual_regions", return_value=[]),
            patch.object(
                word_region,
                "_formula_line_replacements",
                return_value=([], []),
            ) as formula_pass,
        ):
            word_region._plan_pages(
                Document(),
                [self._png_bytes(200, 300)],
                preserve_editable_text=False,
                formula_hints_by_page={0: ()},
                visual_hints_by_page={
                    0: (("text_visual", 100, 200, 300, 400),)
                },
            )

        formula_pass.assert_called_once()

    def test_standalone_bullet_symbol_is_not_discarded_as_formula_debris(self) -> None:
        document = Document()
        paragraph = document.add_paragraph("•")
        frame = word_region._FrameLine(
            paragraph, 200, 200, 120, 220, "•", 10.0, False
        )
        line = word_region._MergedLine(
            [frame], 200, 200, 120, 220, "•", 10.0, False
        )
        source_page = SimpleNamespace(
            index=0,
            width=2000,
            height=3000,
            frames=[frame],
            background=None,
        )

        def keep_residual_lines(lines, visuals, *_args):
            return lines, visuals

        with (
            patch.object(word_region, "_source_pages", return_value=[source_page]),
            patch.object(word_region, "_merge_frame_rows", return_value=[line]),
            patch.object(word_region, "_refine_visual_regions", return_value=[]),
            patch.object(
                word_region,
                "_formula_line_replacements",
                side_effect=keep_residual_lines,
            ),
        ):
            pages = word_region._plan_pages(
                document,
                [self._png_bytes()],
                preserve_editable_text=False,
                formula_hints_by_page={0: ()},
            )

        self.assertIn(
            "•",
            word_region._planned_editable_text(pages, normalize_text=False),
        )

    def test_unconfirmed_formula_requires_positive_math_evidence(self) -> None:
        document = Document()

        def line(text: str) -> word_region._MergedLine:
            frame = word_region._FrameLine(
                document.add_paragraph(text),
                200, 200, 500, 220, text, 10.0, False
            )
            return word_region._MergedLine(
                [frame], 200, 200, 500, 220, text, 10.0, False
            )

        visual = word_region._VisualRegion(
            self._png_bytes(), 200, 200, 500, 220, 80, 24
        )
        setattr(visual, "kind", "formula_unconfirmed")

        for ordinary in (
            "Results", "Conclusion", "•", "—", "✓", "3", "AP50",
            "state-of-the-art", "end-to-end", "pre-trained", "A/B testing",
            "(a)", "page-3",
        ):
            self.assertFalse(
                word_region._line_is_replaced_by_formula(line(ordinary), [visual]),
                ordinary,
            )
        for formula in ("x = y", "x²", "∑x", "12/24"):
            self.assertTrue(
                word_region._line_is_replaced_by_formula(line(formula), [visual]),
                formula,
            )

        setattr(visual, "kind", "formula")
        self.assertFalse(
            word_region._line_is_replaced_by_formula(line("Results"), [visual])
        )

    def test_residual_formula_pass_keeps_an_ordinary_short_heading_editable(self) -> None:
        document = Document()
        paragraph = document.add_paragraph("Results")
        frame = word_region._FrameLine(
            paragraph, 200, 200, 500, 220, "Results", 10.0, False
        )
        line = word_region._MergedLine(
            [frame], 200, 200, 500, 220, "Results", 10.0, False
        )
        visual = word_region._VisualRegion(
            self._png_bytes(), 200, 200, 500, 220, 80, 24
        )
        setattr(visual, "kind", "formula")

        output_lines, output_visuals = word_region._formula_line_replacements(
            [line],
            [visual],
            self._png_bytes(800, 800),
            2000,
            2000,
        )

        self.assertEqual(output_lines, [line])
        self.assertIn(visual, output_visuals)
        self.assertEqual(getattr(visual, "kind", ""), "figure")
        self.assertFalse(
            any(
                getattr(item, "kind", "") == "formula_line"
                for item in output_visuals
            )
        )

    def test_planned_page_keeps_input_hint_for_local_render_audit(self) -> None:
        source_page = SimpleNamespace(
            index=0,
            width=2000,
            height=3000,
            frames=[],
            background=None,
        )
        with (
            patch.object(word_region, "_source_pages", return_value=[source_page]),
            patch.object(word_region, "_merge_frame_rows", return_value=[]),
            patch.object(word_region, "_refine_visual_regions", return_value=[]),
        ):
            pages = word_region._plan_pages(
                Document(),
                [self._png_bytes(200, 300)],
                visual_hints_by_page={
                    0: (("text_visual", 100, 200, 300, 400),)
                },
            )

        self.assertEqual(pages[0].precise_audit_regions, [(100, 200, 300, 400)])

    def test_dense_math_page_keeps_one_redacted_background_anchor(self) -> None:
        background = self._png_bytes(400, 600)
        source_page = SimpleNamespace(
            index=0,
            width=2000,
            height=3000,
            frames=[],
            background=background,
        )
        hints = tuple(
            (
                "inline_math",
                100 + (index % 6) * 250,
                200 + (index // 6) * 300,
                180 + (index % 6) * 250,
                320 + (index // 6) * 300,
            )
            for index in range(word_region._DENSE_BACKGROUND_HINT_THRESHOLD)
        )
        with patch.object(word_region, "_source_pages", return_value=[source_page]):
            pages = word_region._plan_pages(
                Document(),
                [background],
                preserve_editable_text=True,
                visual_hints_by_page={0: hints},
            )

        self.assertEqual(len(pages[0].visual_regions), 1)
        visual = pages[0].visual_regions[0]
        self.assertEqual(
            (visual.x, visual.y, visual.x1, visual.y1),
            (0, 0, 2000, 3000),
        )
        self.assertEqual(getattr(visual, "kind", ""), "background_safe_page")
        self.assertTrue(getattr(visual, "background_safe", False))

    def test_precise_inline_visual_keeps_one_tab_aligned_editable_row(self) -> None:
        document = Document()
        left_paragraph = document.add_paragraph("left editable")
        right_paragraph = document.add_paragraph("right editable")
        left = word_region._FrameLine(
            left_paragraph, 100, 100, 340, 220, "left editable", 10.0, False
        )
        right = word_region._FrameLine(
            right_paragraph, 700, 100, 420, 220, "right editable", 10.0, False
        )
        line = word_region._MergedLine(
            [left, right],
            100,
            100,
            1020,
            220,
            "left editable right editable",
            10.0,
            False,
        )

        output_lines, visuals = word_region._precise_visual_replacements(
            [line],
            [],
            self._png_bytes(200, 200),
            2000,
            2000,
            [("inline_math", 500, 100, 650, 320)],
        )

        self.assertEqual(len(output_lines), 1)
        self.assertEqual(output_lines[0].text, "left editable right editable")
        self.assertEqual(len(output_lines[0].frames), 2)
        self.assertTrue(output_lines[0].precision)
        self.assertEqual(len(visuals), 1)
        self.assertEqual(getattr(visuals[0], "kind", ""), "inline_math")

    def test_one_sided_precise_visual_marks_neighboring_row_as_precision(self) -> None:
        document = Document()
        paragraph = document.add_paragraph("right editable")
        frame = word_region._FrameLine(
            paragraph, 700, 100, 420, 220, "right editable", 10.0, False
        )
        line = word_region._MergedLine(
            [frame], 700, 100, 420, 220, frame.text, 10.0, False
        )

        output_lines, _visuals = word_region._precise_visual_replacements(
            [line],
            [],
            self._png_bytes(200, 200),
            2000,
            2000,
            [("inline_math", 500, 100, 650, 320)],
        )

        self.assertEqual(len(output_lines), 1)
        self.assertTrue(output_lines[0].precision)

    def test_two_precise_gaps_keep_three_fragments_in_one_precision_row(self) -> None:
        document = Document()
        frames = [
            word_region._FrameLine(
                document.add_paragraph(text), x, 100, 220, 220, text, 10.0, False
            )
            for text, x in (("left", 100), ("middle", 500), ("right", 900))
        ]
        lines = [
            word_region._MergedLine(
                [frame], frame.x, frame.y, frame.width, frame.height,
                frame.text, frame.font_size, frame.bold
            )
            for frame in frames
        ]

        output_lines, visuals = word_region._precise_visual_replacements(
            lines,
            [],
            self._png_bytes(300, 200),
            1600,
            1600,
            [
                ("inline_math", 350, 100, 450, 320),
                ("inline_math", 750, 100, 850, 320),
            ],
        )

        self.assertEqual(len(output_lines), 1)
        self.assertEqual([frame.text for frame in output_lines[0].frames], [
            "left", "middle", "right"
        ])
        self.assertTrue(output_lines[0].precision)
        self.assertEqual(len(visuals), 2)

    def test_precision_table_fragments_on_one_row_coalesce(self) -> None:
        document = Document()
        left_frame = word_region._FrameLine(
            document.add_paragraph("eff = min"),
            2456, 8138, 808, 232, "eff = min", 10.0, False,
        )
        right_frame = word_region._FrameLine(
            document.add_paragraph("loc +"),
            3793, 8292, 433, 245, "loc +", 10.0, False,
        )
        left = word_region._MergedLine(
            [left_frame], 2456, 8138, 808, 232,
            left_frame.text, 10.0, False, "left",
        )
        right = word_region._MergedLine(
            [right_frame], 3793, 8292, 433, 245,
            right_frame.text, 10.0, False, "left",
        )
        setattr(left, "precision", True)
        setattr(right, "precision", True)

        output = word_region._coalesce_same_row_fragments(
            [left, right],
            12240,
        )

        self.assertEqual(len(output), 1)
        self.assertEqual(output[0].frames, [left_frame, right_frame])
        self.assertTrue(output[0].precision)

    def test_precision_table_fragments_do_not_merge_across_rows(self) -> None:
        document = Document()
        first_frame = word_region._FrameLine(
            document.add_paragraph("row one"),
            1600, 2700, 2100, 220, "row one", 10.0, False,
        )
        second_frame = word_region._FrameLine(
            document.add_paragraph("row two"),
            1600, 3060, 2100, 220, "row two", 10.0, False,
        )
        first = word_region._MergedLine(
            [first_frame], 1600, 2700, 2100, 220,
            first_frame.text, 10.0, False, "left",
        )
        second = word_region._MergedLine(
            [second_frame], 1600, 3060, 2100, 220,
            second_frame.text, 10.0, False, "left",
        )
        setattr(first, "precision", True)
        setattr(second, "precision", True)

        output = word_region._coalesce_same_row_fragments(
            [first, second],
            12240,
        )

        self.assertEqual(output, [first, second])

    def test_precision_table_rows_form_non_overlapping_regions(self) -> None:
        document = Document()

        def line(text: str, x: int, y: int, width: int) -> object:
            frame = word_region._FrameLine(
                document.add_paragraph(text),
                x, y, width, 220, text, 10.0, False,
            )
            merged = word_region._MergedLine(
                [frame], x, y, width, 220, text, 10.0, False, "right",
            )
            setattr(merged, "precision", True)
            return merged

        lines = [
            line("DenseNet + Grad-CAM", 6390, 1872, 1729),
            line("0.7550 0.2190 0.3306 0.4639", 8687, 1872, 2497),
            line("ConvNeXt-Tiny + Grad-CAM", 6390, 2134, 2200),
            line("0.7640 0.1988 0.3132 0.6836", 8687, 2134, 2497),
            line("Swin-Tiny + Grad-CAM", 6390, 2396, 1769),
        ]

        blocks = word_region._paragraph_blocks(lines, 12240)
        regions = word_region._merge_blocks_to_regions(
            blocks,
            [],
            12240,
            15840,
        )

        for index, first in enumerate(regions):
            for second in regions[index + 1 :]:
                self.assertLessEqual(
                    word_region._overlap_fraction(first, second),
                    0.12,
                )

    def test_precision_full_width_table_row_absorbs_left_label(self) -> None:
        document = Document()

        def line(text: str, x: int, y: int, width: int) -> object:
            frame = word_region._FrameLine(
                document.add_paragraph(text),
                x, y, width, 230, text, 10.0, False,
            )
            merged = word_region._MergedLine(
                [frame], x, y, width, 230, text, 10.0, False,
            )
            setattr(merged, "precision", True)
            return merged

        label = line("Swin-Tiny + Grad-CAM", 1584, 3068, 2118)
        values = line(
            "0.7208 0.1308 0.1953 0.3184 0.7941 0.1329 0.2080 0.2708",
            4365,
            3065,
            6316,
        )

        output = word_region._coalesce_same_row_fragments(
            [label, values],
            12240,
        )

        self.assertEqual(len(output), 1)
        self.assertEqual(len(output[0].frames), 2)
        self.assertEqual(output[0].column, "main")
        self.assertTrue(output[0].precision)

    def test_precision_region_overlap_uses_actual_frame_geometry(self) -> None:
        document = Document()

        def region(text: str, frame_x: int, region_x: int) -> object:
            frame = word_region._FrameLine(
                document.add_paragraph(text),
                frame_x, 100, 120, 180, text, 10.0, False,
            )
            line = word_region._MergedLine(
                [frame], frame_x, 100, 120, 180,
                text, 10.0, False, "main",
            )
            setattr(line, "precision", True)
            block = word_region._TextBlock(
                [line], frame_x, 100, 120, 180,
                "main", 10.0, False, True,
            )
            return word_region._TextRegion(
                [block], region_x, 80, 500, 240, "main", True,
            )

        first = region("numerator", 100, 80)
        disjoint = region("denominator", 700, 300)
        colliding = region("duplicate", 150, 300)

        self.assertFalse(
            word_region._text_regions_have_substantive_frame_overlap(
                first,
                disjoint,
            )
        )
        self.assertTrue(
            word_region._text_regions_have_substantive_frame_overlap(
                first,
                colliding,
            )
        )

    def test_precise_inline_visual_removes_duplicate_orphan_symbol_frame(self) -> None:
        document = Document()
        paragraph = document.add_paragraph("ˆ")
        frame = word_region._FrameLine(
            paragraph, 500, 100, 120, 220, "ˆ", 10.0, False
        )
        line = word_region._MergedLine(
            [frame], 500, 100, 120, 220, "ˆ", 10.0, False
        )

        output_lines, visuals = word_region._precise_visual_replacements(
            [line],
            [],
            self._png_bytes(200, 200),
            2000,
            2000,
            [("inline_math", 480, 90, 650, 330)],
        )

        self.assertEqual(output_lines, [])
        self.assertEqual(len(visuals), 1)

    def test_background_safe_visual_crop_keeps_neighboring_editable_text(self) -> None:
        from PIL import Image, ImageDraw

        document = Document()
        paragraph = document.add_paragraph("low")
        frame = word_region._FrameLine(
            paragraph, 110, 110, 40, 40, "low", 10.0, False
        )
        line = word_region._MergedLine(
            [frame], 110, 110, 40, 40, frame.text, 10.0, False
        )
        rendered = Image.new("RGB", (400, 400), "white")
        ImageDraw.Draw(rendered).rectangle((110, 110, 150, 150), fill="black")
        ImageDraw.Draw(rendered).rectangle((190, 180, 210, 220), fill="black")
        background = Image.new("RGB", (400, 400), "white")
        ImageDraw.Draw(background).rectangle((190, 180, 210, 220), fill="black")
        rendered_payload = io.BytesIO()
        background_payload = io.BytesIO()
        rendered.save(rendered_payload, format="PNG")
        background.save(background_payload, format="PNG")

        output_lines, visuals = word_region._precise_visual_replacements(
            [line],
            [],
            rendered_payload.getvalue(),
            400,
            400,
            [("inline_math", 100, 100, 300, 300)],
            background_page=background_payload.getvalue(),
        )

        self.assertEqual(output_lines, [line])
        self.assertEqual(len(visuals), 1)
        self.assertEqual(
            getattr(visuals[0], "kind", ""),
            "background_safe_inline_math",
        )
        crop = Image.open(io.BytesIO(visuals[0].blob)).convert("RGB")
        self.assertEqual(crop.getpixel((20, 20)), (255, 255, 255))
        self.assertEqual(crop.getpixel((100, 100)), (0, 0, 0))

    def test_small_rendered_visual_hint_cannot_remove_a_large_symbol_frame(self) -> None:
        document = Document()
        paragraph = document.add_paragraph("ˆ")
        frame = word_region._FrameLine(
            paragraph, 100, 100, 400, 220, "ˆ", 10.0, False
        )
        line = word_region._MergedLine(
            [frame], 100, 100, 400, 220, frame.text, 10.0, False
        )

        with self.assertRaisesRegex(ValidationError, "未以高覆盖"):
            word_region._precise_visual_replacements(
                [line],
                [],
                self._png_bytes(600, 400),
                600,
                400,
                [("inline_math", 250, 100, 350, 320)],
            )

    def test_precise_hint_reuses_a_larger_covering_figure(self) -> None:
        old_visual = word_region._VisualRegion(
            self._png_bytes(100, 100), 100, 100, 100, 100, 100, 100
        )
        setattr(old_visual, "kind", "figure")

        _lines, visuals = word_region._precise_visual_replacements(
            [],
            [old_visual],
            self._png_bytes(200, 200),
            400,
            400,
            [("inline_math", 135, 135, 200, 200)],
        )

        self.assertIn(old_visual, visuals)
        self.assertEqual(len(visuals), 1)

    def test_precise_hint_reuses_a_larger_matching_visual_with_one_axis_cover(self) -> None:
        old_visual = word_region._VisualRegion(
            self._png_bytes(360, 269), 100, 100, 360, 269, 360, 269
        )
        setattr(old_visual, "kind", "inline_math")

        _lines, visuals = word_region._precise_visual_replacements(
            [],
            [old_visual],
            self._png_bytes(500, 500),
            500,
            500,
            [("inline_math", 325, 282, 395, 422)],
        )

        self.assertEqual(visuals, [old_visual])

    def test_precise_hint_replaces_a_smaller_old_visual_once(self) -> None:
        old_visual = word_region._VisualRegion(
            self._png_bytes(50, 50), 120, 120, 50, 50, 50, 50
        )
        setattr(old_visual, "kind", "inline_math")

        _lines, visuals = word_region._precise_visual_replacements(
            [],
            [old_visual],
            self._png_bytes(200, 200),
            400,
            400,
            [("inline_math", 100, 100, 180, 180)],
        )

        self.assertNotIn(old_visual, visuals)
        self.assertEqual(len(visuals), 1)
        self.assertEqual((visuals[0].x, visuals[0].y), (100, 100))
        self.assertEqual((visuals[0].width, visuals[0].height), (80, 80))

    def test_precise_replacement_expands_to_cover_removed_visual_fringe(self) -> None:
        old_visual = word_region._VisualRegion(
            self._png_bytes(100, 100), 100, 90, 100, 100, 100, 100
        )
        setattr(old_visual, "kind", "inline_math")
        audit_regions: list[tuple[int, int, int, int]] = []

        _lines, visuals = word_region._precise_visual_replacements(
            [],
            [old_visual],
            self._png_bytes(300, 300),
            300,
            300,
            [("inline_math", 50, 100, 250, 200)],
            audit_regions=audit_regions,
        )

        self.assertNotIn(old_visual, visuals)
        self.assertEqual(len(visuals), 1)
        self.assertEqual(
            (visuals[0].x, visuals[0].y, visuals[0].x1, visuals[0].y1),
            (50, 90, 250, 200),
        )
        self.assertIn((50, 90, 250, 200), audit_regions)

    def test_offset_same_kind_visual_is_replaced_by_a_union_crop(self) -> None:
        old_visual = word_region._VisualRegion(
            self._png_bytes(100, 100), 100, 100, 100, 100, 100, 100
        )
        setattr(old_visual, "kind", "inline_math")

        _lines, visuals = word_region._precise_visual_replacements(
            [],
            [old_visual],
            self._png_bytes(300, 300),
            300,
            300,
            [("inline_math", 130, 120, 230, 220)],
        )

        self.assertNotIn(old_visual, visuals)
        self.assertEqual(len(visuals), 1)
        self.assertEqual(
            (visuals[0].x, visuals[0].y, visuals[0].x1, visuals[0].y1),
            (100, 100, 230, 220),
        )

    def test_background_safe_visual_remains_compatible_with_typed_hint(self) -> None:
        old_visual = word_region._VisualRegion(
            self._png_bytes(100, 100), 100, 100, 100, 100, 100, 100
        )
        setattr(old_visual, "kind", "background_safe_inline_math")
        setattr(old_visual, "background_safe", True)

        _lines, visuals = word_region._precise_visual_replacements(
            [],
            [old_visual],
            self._png_bytes(300, 300),
            300,
            300,
            [("inline_math", 120, 110, 220, 210)],
            background_page=self._png_bytes(300, 300),
        )

        self.assertEqual(len(visuals), 1)
        self.assertEqual(
            getattr(visuals[0], "kind", ""),
            "background_safe_inline_math",
        )

    def test_same_source_background_crop_drops_a_contained_duplicate(self) -> None:
        from PIL import Image, ImageChops, ImageDraw

        background = Image.new("RGB", (1000, 1000), "white")
        draw = ImageDraw.Draw(background)
        draw.rectangle((300, 410, 700, 590), fill=(20, 80, 160))
        draw.ellipse((430, 450, 540, 560), fill=(220, 30, 40))
        background_payload = io.BytesIO()
        background.save(background_payload, format="PNG")

        outer_box = (250, 350, 750, 650)
        inner_box = (420, 440, 560, 570)

        def visual(box: tuple[int, int, int, int]) -> object:
            crop = background.crop(box)
            payload = io.BytesIO()
            crop.save(payload, format="PNG")
            region = word_region._VisualRegion(
                payload.getvalue(),
                box[0],
                box[1],
                box[2] - box[0],
                box[3] - box[1],
                crop.width,
                crop.height,
            )
            setattr(region, "kind", "background_safe_inline_math")
            setattr(region, "background_safe", True)
            return region

        outer = visual(outer_box)
        inner = visual(inner_box)
        visuals = word_region._deduplicate_contained_background_visuals(
            [inner, outer],
            background_payload.getvalue(),
            1000,
            1000,
        )

        self.assertEqual(visuals, [outer])
        outer_image = Image.open(io.BytesIO(outer.blob)).convert("RGB")
        retained_inner = outer_image.crop(
            (
                inner_box[0] - outer_box[0],
                inner_box[1] - outer_box[1],
                inner_box[2] - outer_box[0],
                inner_box[3] - outer_box[1],
            )
        )
        self.assertIsNone(
            ImageChops.difference(
                retained_inner,
                background.crop(inner_box),
            ).getbbox()
        )

    def test_background_crop_keeps_duplicate_when_outer_pixels_do_not_match(self) -> None:
        from PIL import Image

        background = Image.new("RGB", (400, 400), "white")
        background_payload = io.BytesIO()
        background.save(background_payload, format="PNG")

        outer_image = Image.new("RGB", (300, 300), "black")
        outer_payload = io.BytesIO()
        outer_image.save(outer_payload, format="PNG")
        outer = word_region._VisualRegion(
            outer_payload.getvalue(), 50, 50, 300, 300, 300, 300
        )
        inner_image = background.crop((100, 100, 200, 200))
        inner_payload = io.BytesIO()
        inner_image.save(inner_payload, format="PNG")
        inner = word_region._VisualRegion(
            inner_payload.getvalue(), 100, 100, 100, 100, 100, 100
        )
        for visual in (outer, inner):
            setattr(visual, "kind", "background_safe_inline_math")
            setattr(visual, "background_safe", True)

        visuals = word_region._deduplicate_contained_background_visuals(
            [outer, inner],
            background_payload.getvalue(),
            400,
            400,
        )

        self.assertEqual(visuals, [outer, inner])

    def test_background_visual_reanchors_to_source_pixel_edges(self) -> None:
        from PIL import Image, ImageChops, ImageDraw

        background = Image.new("RGB", (401, 503), "white")
        ImageDraw.Draw(background).rectangle(
            (41, 77, 166, 201),
            fill=(12, 90, 180),
        )
        background_payload = io.BytesIO()
        background.save(background_payload, format="PNG")
        visual = word_region._VisualRegion(
            self._png_bytes(30, 30),
            103,
            207,
            811,
            677,
            30,
            30,
        )
        setattr(visual, "kind", "background_safe_inline_math")
        setattr(visual, "background_safe", True)

        output = word_region._align_background_visuals_to_source_pixels(
            [visual],
            background_payload.getvalue(),
            2000,
            2500,
        )

        self.assertEqual(output, [visual])
        pixel_box = (
            round(visual.x * background.width / 2000),
            round(visual.y * background.height / 2500),
            round(visual.x1 * background.width / 2000),
            round(visual.y1 * background.height / 2500),
        )
        crop = Image.open(io.BytesIO(visual.blob)).convert("RGB")
        self.assertEqual(
            crop.size,
            (pixel_box[2] - pixel_box[0], pixel_box[3] - pixel_box[1]),
        )
        self.assertIsNone(
            ImageChops.difference(
                crop,
                background.crop(pixel_box),
            ).getbbox()
        )

    def test_rendered_union_crop_cannot_expand_over_editable_text(self) -> None:
        document = Document()
        paragraph = document.add_paragraph("editable")
        frame = word_region._FrameLine(
            paragraph, 100, 100, 25, 100, "editable", 10.0, False
        )
        line = word_region._MergedLine(
            [frame], 100, 100, 25, 100, frame.text, 10.0, False
        )
        old_visual = word_region._VisualRegion(
            self._png_bytes(100, 100), 100, 100, 100, 100, 100, 100
        )
        setattr(old_visual, "kind", "inline_math")

        with self.assertRaisesRegex(ValidationError, "扩展后会覆盖可编辑文字"):
            word_region._precise_visual_replacements(
                [line],
                [old_visual],
                self._png_bytes(300, 300),
                300,
                300,
                [("inline_math", 130, 120, 230, 220)],
            )

    def test_precise_hint_rejects_ambiguous_partial_visual_overlap(self) -> None:
        old_visual = word_region._VisualRegion(
            self._png_bytes(100, 100), 100, 100, 100, 100, 100, 100
        )
        setattr(old_visual, "kind", "figure")

        with self.assertRaisesRegex(ValidationError, "部分重叠"):
            word_region._precise_visual_replacements(
                [],
                [old_visual],
                self._png_bytes(300, 300),
                400,
                400,
                [("text_visual", 150, 150, 250, 250)],
            )

    def test_precise_hint_ignores_a_one_axis_antialias_fringe(self) -> None:
        old_visual = word_region._VisualRegion(
            self._png_bytes(100, 60), 100, 195, 100, 60, 100, 60
        )
        setattr(old_visual, "kind", "inline_math")

        _lines, visuals = word_region._precise_visual_replacements(
            [],
            [old_visual],
            self._png_bytes(300, 300),
            400,
            400,
            [("inline_math", 100, 100, 300, 200)],
        )

        self.assertIn(old_visual, visuals)
        self.assertEqual(len(visuals), 2)

    def test_precise_hint_does_not_expand_a_sub_twenty_twip_crop(self) -> None:
        _lines, visuals = word_region._precise_visual_replacements(
            [],
            [],
            self._png_bytes(400, 400),
            400,
            400,
            [("inline_math", 100, 100, 110, 115)],
        )

        self.assertEqual(len(visuals), 1)
        self.assertEqual((visuals[0].width, visuals[0].height), (10, 15))

    def test_overlapping_precise_hints_coalesce_to_one_visual(self) -> None:
        _lines, visuals = word_region._precise_visual_replacements(
            [],
            [],
            self._png_bytes(400, 400),
            400,
            400,
            [
                ("text_visual", 100, 100, 200, 200),
                ("inline_math", 110, 105, 210, 205),
            ],
        )

        self.assertEqual(len(visuals), 1)
        self.assertEqual(getattr(visuals[0], "kind", ""), "inline_math")
        self.assertEqual(
            (visuals[0].x, visuals[0].y, visuals[0].x1, visuals[0].y1),
            (100, 100, 210, 205),
        )

    def test_overlapping_formula_hints_coalesce_to_one_visual(self) -> None:
        _lines, visuals = word_region._precise_formula_replacements(
            [],
            [],
            self._png_bytes(400, 400),
            400,
            400,
            [
                (100, 100, 200, 200),
                (110, 105, 210, 205),
            ],
        )

        self.assertEqual(len(visuals), 1)
        self.assertEqual(getattr(visuals[0], "kind", ""), "formula_exact")
        self.assertEqual(
            (visuals[0].x, visuals[0].y, visuals[0].x1, visuals[0].y1),
            (100, 100, 210, 205),
        )

    def test_confirmed_formula_crop_absorbs_a_larger_matching_component(self) -> None:
        old_visual = word_region._VisualRegion(
            self._png_bytes(100, 100), 100, 100, 100, 100, 100, 100
        )
        setattr(old_visual, "kind", "formula")

        _lines, visuals = word_region._precise_formula_replacements(
            [],
            [old_visual],
            self._png_bytes(300, 300),
            300,
            300,
            [(50, 130, 250, 230)],
        )

        self.assertNotIn(old_visual, visuals)
        self.assertEqual(len(visuals), 1)
        self.assertEqual(
            (visuals[0].x, visuals[0].y, visuals[0].x1, visuals[0].y1),
            (50, 100, 250, 230),
        )

    def test_formula_hint_center_does_not_delete_a_wide_editable_frame(self) -> None:
        document = Document()
        paragraph = document.add_paragraph("Important editable surrounding text")
        frame = word_region._FrameLine(
            paragraph,
            100, 100, 1000, 220,
            "Important editable surrounding text", 10.0, False
        )
        line = word_region._MergedLine(
            [frame], 100, 100, 1000, 220,
            frame.text, 10.0, False
        )

        with self.assertRaisesRegex(ValidationError, "未高覆盖可编辑文字"):
            word_region._precise_formula_replacements(
                [line],
                [],
                self._png_bytes(1200, 400),
                1200,
                400,
                [(550, 100, 650, 320)],
            )

    def test_formula_hint_requires_high_coverage_and_clear_math_evidence(self) -> None:
        document = Document()

        def line(text: str) -> word_region._MergedLine:
            frame = word_region._FrameLine(
                document.add_paragraph(text),
                100, 100, 500, 220, text, 10.0, False
            )
            return word_region._MergedLine(
                [frame], 100, 100, 500, 220, text, 10.0, False
            )

        with self.assertRaisesRegex(ValidationError, "非数学可编辑文字"):
            word_region._precise_formula_replacements(
                [line("Results")],
                [],
                self._png_bytes(700, 400),
                700,
                400,
                [(100, 100, 600, 320)],
            )
        with self.assertRaisesRegex(ValidationError, "未高覆盖可编辑文字"):
            word_region._precise_formula_replacements(
                [line("x = y")],
                [],
                self._png_bytes(700, 400),
                700,
                400,
                [(100, 100, 450, 320)],
            )

        output_lines, visuals = word_region._precise_formula_replacements(
            [line("x = y")],
            [],
            self._png_bytes(700, 400),
            700,
            400,
            [(100, 100, 600, 320)],
        )

        self.assertEqual(output_lines, [])
        self.assertEqual(len(visuals), 1)
        self.assertEqual(getattr(visuals[0], "kind", ""), "formula_exact")

    def test_background_safe_formula_crop_keeps_non_math_editable_frame(self) -> None:
        from PIL import Image, ImageDraw

        document = Document()
        paragraph = document.add_paragraph("cls")
        frame = word_region._FrameLine(
            paragraph, 110, 110, 40, 40, "cls", 10.0, False
        )
        line = word_region._MergedLine(
            [frame], 110, 110, 40, 40, frame.text, 10.0, False
        )
        rendered = Image.new("RGB", (400, 400), "white")
        ImageDraw.Draw(rendered).rectangle((110, 110, 150, 150), fill="black")
        ImageDraw.Draw(rendered).rectangle((180, 180, 220, 220), fill="black")
        background = Image.new("RGB", (400, 400), "white")
        ImageDraw.Draw(background).rectangle((180, 180, 220, 220), fill="black")
        rendered_payload = io.BytesIO()
        background_payload = io.BytesIO()
        rendered.save(rendered_payload, format="PNG")
        background.save(background_payload, format="PNG")

        output_lines, visuals = word_region._precise_formula_replacements(
            [line],
            [],
            rendered_payload.getvalue(),
            400,
            400,
            [(100, 100, 300, 300)],
            background_page=background_payload.getvalue(),
        )

        self.assertEqual(output_lines, [line])
        self.assertEqual(len(visuals), 1)
        self.assertEqual(
            getattr(visuals[0], "kind", ""),
            "background_safe_formula_exact",
        )
        crop = Image.open(io.BytesIO(visuals[0].blob)).convert("RGB")
        self.assertEqual(crop.getpixel((20, 20)), (255, 255, 255))
        self.assertEqual(crop.getpixel((100, 100)), (0, 0, 0))

    def _overlapping_region_page(self, visual_kind: str) -> word_region._RegionPage:
        source_document = Document()
        paragraph = source_document.add_paragraph("editable table value")
        frame = word_region._FrameLine(
            paragraph, 200, 200, 900, 240, "editable table value", 10.0, False
        )
        line = word_region._MergedLine(
            [frame], 200, 200, 900, 240, frame.text, 10.0, False, "main"
        )
        block = word_region._TextBlock(
            [line], 200, 200, 900, 240, "main", 10.0, False
        )
        region = word_region._TextRegion(
            [block], 190, 190, 930, 270, "main"
        )
        visual = word_region._VisualRegion(
            self._png_bytes(), 300, 200, 360, 240, 80, 24
        )
        setattr(visual, "kind", visual_kind)
        return word_region._RegionPage(0, 2200, 3200, [region], [visual])

    @staticmethod
    def _reading_order_region(
        document: Document,
        text: str,
        *,
        x: int,
        y: int,
        width: int,
        column: str,
        heading: bool = False,
    ) -> word_region._TextRegion:
        paragraph = document.add_paragraph(text)
        frame = word_region._FrameLine(
            paragraph, x, y, width, 220, text, 10.0, heading
        )
        line = word_region._MergedLine(
            [frame], x, y, width, 220, text, 10.0, heading, column
        )
        block = word_region._TextBlock(
            [line], x, y, width, 220, column, 10.0, heading
        )
        return word_region._TextRegion(
            [block], x, y, width, 240, column
        )

    @classmethod
    def _reading_order_page(
        cls, *, include_middle_heading: bool = False, include_visual: bool = False
    ) -> word_region._RegionPage:
        document = Document()
        regions = [
            cls._reading_order_region(
                document,
                "Page title",
                x=200,
                y=100,
                width=2000,
                column="full",
                heading=True,
            ),
            cls._reading_order_region(
                document, "Left upper", x=200, y=500, width=850, column="left"
            ),
            cls._reading_order_region(
                document, "Right upper", x=1350, y=500, width=850, column="right"
            ),
        ]
        if include_middle_heading:
            regions.extend(
                [
                    cls._reading_order_region(
                        document,
                        "Middle heading",
                        x=200,
                        y=1000,
                        width=2000,
                        column="full",
                        heading=True,
                    ),
                    cls._reading_order_region(
                        document,
                        "Left lower",
                        x=200,
                        y=1400,
                        width=850,
                        column="left",
                    ),
                    cls._reading_order_region(
                        document,
                        "Right lower",
                        x=1350,
                        y=1400,
                        width=850,
                        column="right",
                    ),
                ]
            )
        visuals: list[word_region._VisualRegion] = []
        if include_visual:
            visual = word_region._VisualRegion(
                cls._png_bytes(160, 48), 250, 940, 1900, 180, 160, 48
            )
            setattr(visual, "kind", "figure")
            visuals.append(visual)
        return word_region._RegionPage(0, 2400, 3200, regions, visuals)

    @staticmethod
    def _reorder_anchor_runs(
        path: Path,
        *,
        description: str,
        order: list[int],
    ) -> None:
        document = Document(path)
        anchors = document.element.body.xpath(
            f".//wp:anchor[wp:docPr[@descr='{description}']]"
        )
        runs = [anchor.getparent().getparent() for anchor in anchors]
        if not runs:
            raise AssertionError("test document contains no matching anchors")
        parent = runs[0].getparent()
        if any(run.getparent() is not parent for run in runs):
            raise AssertionError("test anchors are not attached to one page carrier")
        for run in runs:
            parent.remove(run)
        for index in order:
            parent.append(runs[index])
        document.save(path)

    def test_structure_gate_rejects_reversed_double_column_regions(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            output = Path(folder) / "reversed-columns.docx"
            page = self._reading_order_page()
            source_text = word_region._planned_editable_text(
                [page], normalize_text=False
            )
            word_region._build_document(
                [page], source_text, font_scale=100000, normalize_text=False
            ).save(output)
            # Expected: title, left, right.  Serialize the right column first.
            self._reorder_anchor_runs(
                output,
                description="LayoutLoom editable region",
                order=[0, 2, 1],
            )

            with self.assertRaisesRegex(
                ValidationError, "双栏/局部分栏区域整体逆序"
            ):
                word_region._inspect_and_validate_structure(
                    output, [page], source_text, source_frames=3
                )

    def test_planned_text_uses_serialized_double_column_reading_order(self) -> None:
        page = self._reading_order_page()
        page.text_regions = [
            page.text_regions[0],
            page.text_regions[2],
            page.text_regions[1],
        ]

        source_text = word_region._planned_editable_text(
            [page],
            normalize_text=False,
        )

        self.assertEqual(
            source_text.splitlines(),
            ["Page title", "Left upper", "Right upper"],
        )

    def test_full_page_background_does_not_split_text_reading_order(self) -> None:
        page = self._reading_order_page()
        background = word_region._full_page_background_visual(
            self._png_bytes(240, 320),
            page.width,
            page.height,
        )
        page.visual_regions = [background]

        source_text = word_region._planned_editable_text(
            [page],
            normalize_text=False,
        )

        self.assertEqual(
            source_text.splitlines(),
            ["Page title", "Left upper", "Right upper"],
        )

    def test_structure_gate_rejects_cross_column_heading_order(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            output = Path(folder) / "heading-order.docx"
            page = self._reading_order_page(include_middle_heading=True)
            source_text = word_region._planned_editable_text(
                [page], normalize_text=False
            )
            word_region._build_document(
                [page], source_text, font_scale=100000, normalize_text=False
            ).save(output)
            # Expected: title, left upper, right upper, middle heading, lower
            # columns.  Move lower-left ahead of the cross-column heading.
            self._reorder_anchor_runs(
                output,
                description="LayoutLoom editable region",
                order=[0, 1, 2, 4, 3, 5],
            )

            with self.assertRaisesRegex(ValidationError, "跨栏标题顺序异常"):
                word_region._inspect_and_validate_structure(
                    output, [page], source_text, source_frames=6
                )

    def test_structure_gate_ignores_ordinary_visual_anchor_order(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            output = Path(folder) / "visual-order.docx"
            page = self._reading_order_page(
                include_middle_heading=True, include_visual=True
            )
            source_text = word_region._planned_editable_text(
                [page], normalize_text=False
            )
            word_region._build_document(
                [page], source_text, font_scale=100000, normalize_text=False
            ).save(output)
            self._reorder_anchor_runs(
                output,
                description="LayoutLoom region visual",
                order=[0],
            )

            inspection = word_region._inspect_and_validate_structure(
                output, [page], source_text, source_frames=6
            )

            self.assertTrue(inspection.structural_passed)

    def test_structure_check_allows_transparent_box_across_exact_formula(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            output = Path(folder) / "overlap.docx"
            page = self._overlapping_region_page("formula_exact")
            source_text = word_region._planned_editable_text(
                [page], normalize_text=False
            )
            word_region._build_document(
                [page], source_text, font_scale=100000, normalize_text=False
            ).save(output)

            inspection = word_region._inspect_and_validate_structure(
                output, [page], source_text, source_frames=1
            )

            self.assertTrue(inspection.structural_passed)

    def test_planning_rejects_real_frame_overlap_with_exact_formula(self) -> None:
        page = self._overlapping_region_page("formula_exact")

        with self.assertRaisesRegex(ValidationError, "formula_exact 视觉层重叠"):
            word_region._validate_planned_frame_visual_overlap([page])

    def test_planning_allows_editable_text_inside_large_figure_background(self) -> None:
        page = self._overlapping_region_page("figure")

        word_region._validate_planned_frame_visual_overlap([page])

    def test_planning_allows_a_thin_visual_bbox_edge_touch(self) -> None:
        page = self._overlapping_region_page("inline_math")
        visual = page.visual_regions[0]
        visual.x = 1090
        visual.width = 140

        word_region._validate_planned_frame_visual_overlap([page])

    def test_planning_allows_background_safe_visual_over_editable_text(self) -> None:
        page = self._overlapping_region_page("background_safe_inline_math")

        word_region._validate_planned_frame_visual_overlap([page])

    def test_adjacent_precision_rows_can_share_one_editable_region(self) -> None:
        document = Document()
        first_frame = word_region._FrameLine(
            document.add_paragraph("first"),
            200, 200, 900, 220, "first", 10.0, False
        )
        second_frame = word_region._FrameLine(
            document.add_paragraph("second"),
            210, 450, 880, 220, "second", 10.0, False
        )
        first_line = word_region._MergedLine(
            [first_frame], 200, 200, 900, 220, "first", 10.0, False, "main"
        )
        second_line = word_region._MergedLine(
            [second_frame], 210, 450, 880, 220, "second", 10.0, False, "main"
        )
        first_block = word_region._TextBlock(
            [first_line], 200, 200, 900, 220, "main", 10.0, False, True
        )
        second_block = word_region._TextBlock(
            [second_line], 210, 450, 880, 220, "main", 10.0, False, True
        )

        self.assertTrue(
            word_region._may_share_region(
                [first_block], second_block, [], 2200, 5000
            )
        )

        second_block.y = 900
        self.assertFalse(
            word_region._may_share_region(
                [first_block], second_block, [], 2200, 5000
            )
        )

    def test_structure_check_allows_text_over_table_grid(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            output = Path(folder) / "table.docx"
            page = self._overlapping_region_page("table")
            source_text = word_region._planned_editable_text(
                [page], normalize_text=False
            )
            word_region._build_document(
                [page], source_text, font_scale=100000, normalize_text=False
            ).save(output)

            inspection = word_region._inspect_and_validate_structure(
                output, [page], source_text, source_frames=1
            )

            self.assertTrue(inspection.structural_passed)

    def test_structure_check_allows_text_inside_figure_background(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            output = Path(folder) / "figure-background.docx"
            page = self._overlapping_region_page("figure")
            source_text = word_region._planned_editable_text(
                [page], normalize_text=False
            )
            word_region._build_document(
                [page], source_text, font_scale=100000, normalize_text=False
            ).save(output)

            inspection = word_region._inspect_and_validate_structure(
                output, [page], source_text, source_frames=1
            )

            self.assertTrue(inspection.structural_passed)

    def test_full_scale_text_boxes_disable_costly_autofit(self) -> None:
        page = self._overlapping_region_page("table")

        full_scale = word_region._build_document(
            [page], "editable table value", font_scale=100000,
            normalize_text=False
        )
        corrected = word_region._build_document(
            [page], "editable table value", font_scale=96500,
            normalize_text=False
        )

        self.assertTrue(
            full_scale.element.body.xpath(
                ".//*[local-name()='bodyPr']/*[local-name()='noAutofit']"
            )
        )
        self.assertFalse(
            full_scale.element.body.xpath(
                ".//*[local-name()='bodyPr']/*[local-name()='normAutofit']"
            )
        )
        self.assertTrue(
            corrected.element.body.xpath(
                ".//*[local-name()='bodyPr']/*[local-name()='normAutofit']"
            )
        )

    def test_adjacent_identically_formatted_frames_share_one_run(self) -> None:
        source = Document()
        first_paragraph = source.add_paragraph()
        first_run = first_paragraph.add_run("left")
        first_run.font.name = "Times New Roman"
        first_run.font.size = Pt(10)
        second_paragraph = source.add_paragraph()
        second_run = second_paragraph.add_run("right")
        second_run.font.name = "Times New Roman"
        second_run.font.size = Pt(10)
        first = word_region._FrameLine(
            first_paragraph, 100, 100, 240, 220, "left", 10.0, False
        )
        second = word_region._FrameLine(
            second_paragraph, 700, 100, 260, 220, "right", 10.0, False
        )
        line = word_region._MergedLine(
            [first, second], 100, 100, 860, 220,
            "left right", 10.0, False, "main"
        )
        block = word_region._TextBlock(
            [line], 100, 100, 860, 220, "main", 10.0, False, True
        )
        region = word_region._TextRegion(
            [block], 100, 100, 860, 220, "main", True
        )

        paragraph = word_region._append_region_paragraph(
            Document(), block, region, previous_block=None,
            word_lexicon=set(), normalize_text=False
        )

        self.assertEqual(len(paragraph.xpath("./w:r")), 1)
        self.assertEqual(len(paragraph.xpath("./w:r/w:tab")), 1)
        self.assertEqual(
            "".join(node.text or "" for node in paragraph.xpath("./w:r/w:t")),
            "leftright",
        )

    def test_same_format_soft_line_break_is_folded_into_one_run(self) -> None:
        source = Document()

        def line(text: str, y: int) -> word_region._MergedLine:
            paragraph = source.add_paragraph()
            run = paragraph.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            frame = word_region._FrameLine(
                paragraph, 100, y, 600, 220, text, 10.0, False
            )
            return word_region._MergedLine(
                [frame], 100, y, 600, 220, text, 10.0, False, "main"
            )

        lines = [line("first line", 100), line("second line", 340)]
        block = word_region._TextBlock(
            lines, 100, 100, 600, 460, "main", 10.0, False, False
        )
        region = word_region._TextRegion(
            [block], 100, 100, 600, 460, "main", False
        )

        paragraph = word_region._append_region_paragraph(
            Document(), block, region, previous_block=None,
            word_lexicon=set(), normalize_text=False
        )

        self.assertEqual(len(paragraph.xpath("./w:r")), 1)
        self.assertEqual(len(paragraph.xpath("./w:r/w:br")), 1)
        self.assertEqual(
            "".join(node.text or "" for node in paragraph.xpath("./w:r/w:t")),
            "first linesecond line",
        )

    def test_complex_source_run_forces_safe_fixed_layout_fallback(self) -> None:
        source = Document()
        paragraph = source.add_paragraph("editable")
        paragraph.runs[0]._r.append(OxmlElement("w:sym"))
        frame = word_region._FrameLine(
            paragraph, 100, 100, 600, 220, "editable", 10.0, False
        )
        line = word_region._MergedLine(
            [frame], 100, 100, 600, 220, "editable", 10.0, False, "main"
        )
        block = word_region._TextBlock(
            [line], 100, 100, 600, 220, "main", 10.0, False, False
        )
        region = word_region._TextRegion(
            [block], 100, 100, 600, 220, "main", False
        )

        with self.assertRaisesRegex(ValidationError, "无法无损迁移的复杂文字节点"):
            word_region._append_region_paragraph(
                Document(), block, region, previous_block=None,
                word_lexicon=set(), normalize_text=False
            )

    def test_region_document_marks_proofing_clean_for_faster_initial_open(self) -> None:
        page = self._overlapping_region_page("table")

        document = word_region._build_document(
            [page], "editable table value", font_scale=100000,
            normalize_text=False
        )

        proof_state = document.settings.element.find(qn("w:proofState"))
        self.assertIsNotNone(proof_state)
        assert proof_state is not None
        self.assertEqual(proof_state.get(qn("w:spelling")), "clean")
        self.assertEqual(proof_state.get(qn("w:grammar")), "clean")

    def test_local_visual_difference_detects_extra_and_missing_ink(self) -> None:
        import pymupdf

        source = pymupdf.open()
        candidate = pymupdf.open()
        source_page = source.new_page(width=100, height=100)
        candidate_page = candidate.new_page(width=100, height=100)
        candidate_page.draw_rect(
            pymupdf.Rect(20, 20, 80, 80), color=(0, 0, 0), fill=(0, 0, 0)
        )
        try:
            changed, extra, missing = word_region._local_visual_difference(
                source_page,
                candidate_page,
                [(0, 0, 2000, 2000)],
            )
            reverse_changed, reverse_extra, reverse_missing = (
                word_region._local_visual_difference(
                    candidate_page,
                    source_page,
                    [(0, 0, 2000, 2000)],
                )
            )
        finally:
            source.close()
            candidate.close()

        self.assertGreater(changed, 0.20)
        self.assertGreater(extra, 0.20)
        self.assertLess(missing, 0.01)
        self.assertGreater(reverse_changed, 0.20)
        self.assertLess(reverse_extra, 0.01)
        self.assertGreater(reverse_missing, 0.20)

    def test_formula_candidate_rejects_long_english_prose_with_digits_and_inline_math(
        self,
    ) -> None:
        prose_samples = (
            "We use pretrained DINOv2 ViT-S/14 as the shared visual encoder.",
            "Each image is divided into non-overlapping 14 x 14 patches, "
            "yielding an editable feature sequence.",
            "The objective uses x = y only as an inline example while the "
            "surrounding sentence remains ordinary editable English prose.",
            "Theobjectiveusesx=yonlyasaninlineexamplewhilethesurrounding"
            "sentenceremainsordinaryeditableEnglishprose.",
            "We use DINOv2 (ViT-S/14) as the shared visual encoder.",
            "The model (version 2.1) achieves 93.5% accuracy.",
            "Fig. 3(a) shows...",
            "The loss L = sum_i w_i x_i is optimized end-to-end.",
            "In Eq. (3), x^2 + y^2 = z^2 defines the constraint.",
            "Algorithm 1: Set x = y, then return z.",
            "当 x = y 时，模型输出保持不变。",
            "公式（3）中的 x 表示输入特征，y 表示目标标签。",
            "图 3(a) 显示 AP50/AP75 提升了 2.1%。",
            "损失函数 L = sum_i w_i x_i 通过反向传播进行优化。",
        )
        formula_samples = (
            "L = sum_i w_i x_i",
            "x^2 + y^2 = z^2",
        )

        for text in prose_samples:
            with self.subTest(kind="prose", text=text):
                line = word_region._MergedLine(
                    frames=[],
                    x=0,
                    y=0,
                    width=4000,
                    height=260,
                    text=text,
                    font_size=10.0,
                    bold=False,
                )
                self.assertFalse(word_region._formula_line_text_candidate(line))

        for text in formula_samples:
            with self.subTest(kind="formula", text=text):
                line = word_region._MergedLine(
                    frames=[],
                    x=0,
                    y=0,
                    width=1600,
                    height=260,
                    text=text,
                    font_size=10.0,
                    bold=False,
                )
                self.assertTrue(word_region._formula_line_text_candidate(line))

    def test_formula_visual_does_not_rasterize_neighboring_prose_line(self) -> None:
        visual = word_region._VisualRegion(
            blob=b"",
            x=100,
            y=100,
            width=500,
            height=220,
            pixel_width=50,
            pixel_height=22,
        )
        setattr(visual, "kind", "formula")
        prose = word_region._MergedLine(
            frames=[],
            x=80,
            y=90,
            width=1800,
            height=260,
            text="We use pretrained visual features for the shared representation.",
            font_size=10.0,
            bold=False,
        )
        equation = word_region._MergedLine(
            frames=[],
            x=80,
            y=90,
            width=900,
            height=260,
            text="low = Conv 1 loc 1 (x)",
            font_size=10.0,
            bold=False,
        )

        self.assertFalse(word_region._line_is_replaced_by_formula(prose, [visual]))
        self.assertTrue(word_region._line_is_replaced_by_formula(equation, [visual]))

    def test_mixed_chinese_english_boundaries_are_normalized(self) -> None:
        self.assertEqual(
            word_region._normalize_mixed_script_text("高配版:Chinese"),
            "高配版：Chinese",
        )
        self.assertEqual(
            word_region._normalize_mixed_script_text("in a word简而言之"),
            "in a word 简而言之",
        )
        self.assertEqual(
            word_region._normalize_mixed_script_text("user@example.com"),
            "user@example.com",
        )

    def test_normalization_keeps_separate_run_format_owners(self) -> None:
        document = Document()
        paragraph = document.add_paragraph()
        first = paragraph.add_run("高配版")
        first.bold = True
        second = paragraph.add_run(":Chinese")
        frame = word_region._FrameLine(
            paragraph=paragraph,
            x=0,
            y=0,
            width=2000,
            height=240,
            text="高配版:Chinese",
            font_size=10.5,
            bold=False,
        )

        segments = word_region._normalized_frame_runs(frame)

        self.assertEqual("".join(text for _run, text in segments), "高配版：Chinese")
        self.assertIs(segments[0][0]._r, first._r)
        self.assertIs(segments[1][0]._r, second._r)

    @staticmethod
    def _add_positioned_text(
        document: Document,
        text: str,
        *,
        x: int,
        y: int,
        width: int = 7200,
        size: float = 10.5,
        bold: bool = False,
    ) -> None:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        properties = paragraph._p.get_or_add_pPr()
        frame = OxmlElement("w:framePr")
        for name, value in (
            ("w:hAnchor", "page"),
            ("w:vAnchor", "page"),
            ("w:x", str(x)),
            ("w:y", str(y)),
            ("w:w", str(width)),
            ("w:h", str(max(200, round(size * 24)))),
        ):
            frame.set(qn(name), value)
        properties.insert(0, frame)
        fit = OxmlElement("w:fitText")
        fit.set(qn("w:val"), str(width))
        run._r.get_or_add_rPr().append(fit)

    @classmethod
    def _make_fixed_layout_docx(cls, path: Path, pages: int = 3) -> Path:
        document = Document()
        for page_number in range(1, pages + 1):
            if page_number > 1:
                document.add_section(WD_SECTION.NEW_PAGE)
            cls._add_positioned_text(
                document,
                f"Page {page_number} heading",
                x=1500,
                y=1000,
                width=5000,
                size=16,
                bold=True,
            )
            cls._add_positioned_text(
                document,
                f"Page {page_number} editable paragraph begins with normal text",
                x=1450,
                y=1450,
            )
            cls._add_positioned_text(
                document,
                "and continues on another positioned visual line for copying.",
                x=1450,
                y=1740,
            )
            cls._add_positioned_text(
                document,
                "A second paragraph remains editable after the region rebuild.",
                x=1450,
                y=2260,
            )
        document.save(path)
        return path

    def test_region_rebuild_removes_line_frames_and_preserves_text(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_fixed_layout_docx(root / "fixed.docx", pages=3)
            output = root / "fixed_区域级全兼容Word.docx"
            source_bytes = source.read_bytes()

            results = word_compat.optimize_word_full_compatibility(
                source,
                output,
                verification_engine="none",
            )

            self.assertEqual(results, [output.resolve()])
            self.assertEqual(source.read_bytes(), source_bytes)
            document = Document(output)
            body = document.element.body
            self.assertFalse(body.xpath(".//w:pPr/w:framePr"))
            self.assertFalse(body.xpath(".//w:fitText"))
            boxes = body.xpath(
                ".//*[namespace-uri()='http://schemas.microsoft.com/office/word/2010/wordprocessingShape' and local-name()='txbx']"
            )
            self.assertGreaterEqual(len(boxes), 3)
            self.assertLess(len(boxes), 12)
            self.assertEqual(
                len(body.xpath(".//wp:docPr[@descr='LayoutLoom editable region']")),
                len(boxes),
            )
            text = "".join(node.text or "" for node in body.xpath(".//w:txbxContent//w:t"))
            self.assertIn("Page 1 editable paragraph", text)
            self.assertIn("second paragraph remains editable", text)
            self.assertEqual(
                len(body.xpath(".//w:pPr/w:pageBreakBefore")),
                2,
            )
            drawing_ids = [
                item.get("id") for item in body.xpath(".//wp:docPr")
            ]
            self.assertEqual(len(drawing_ids), len(set(drawing_ids)))
            self.assertTrue(
                document.settings.element.xpath("./w:doNotAutoCompressPictures")
            )

    def test_region_contains_normal_multi_line_paragraphs(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_fixed_layout_docx(root / "fixed.docx", pages=1)
            output = root / "result.docx"
            word_compat.optimize_word_full_compatibility(
                source,
                output,
                verification_engine="none",
            )

            document = Document(output)
            paragraph_texts = [
                "".join(node.text or "" for node in paragraph.xpath(".//w:t"))
                for paragraph in document.element.body.xpath(".//w:txbxContent/w:p")
            ]
            self.assertTrue(
                any(
                    "paragraph begins" in text and "continues on another" in text
                    for text in paragraph_texts
                )
            )

    def test_plain_word_is_rejected_instead_of_silently_reformatted(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "plain.docx"
            document = Document()
            document.add_paragraph("Already normal Word content")
            document.save(source)

            with self.assertRaisesRegex(ValidationError, "未检测到.*定位文字"):
                word_compat.optimize_word_full_compatibility(
                    source,
                    root / "out.docx",
                    verification_engine="none",
                )

    def test_render_check_rebuilds_once_with_conservative_font_fit(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_fixed_layout_docx(root / "fixed.docx", pages=2)
            output = root / "result.docx"
            failed = word_region._RenderAudit(False, "疑似文字重叠")
            passed = word_region._RenderAudit(True, page_count=2, text_recall=1.0)

            with (
                patch.object(
                    word_region,
                    "_render_source_for_regions",
                    return_value=root / "source-render.pdf",
                ),
                patch.object(
                    word_region,
                    "_rasterized_source_pages",
                    return_value=[],
                ),
                patch.object(
                    word_region,
                    "_render_audit",
                    side_effect=(failed, passed),
                ) as audit,
            ):
                results = word_compat.optimize_word_full_compatibility(
                    source,
                    output,
                    verification_engine="wps",
                    timeout=123,
                )

            self.assertEqual(results, [output.resolve()])
            self.assertEqual(audit.call_count, 2)
            self.assertEqual(audit.call_args.kwargs["engine"], "wps")
            self.assertEqual(audit.call_args.kwargs["timeout"], 123.0)

    def test_second_render_failure_rejects_output(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = self._make_fixed_layout_docx(root / "fixed.docx", pages=1)
            output = root / "result.docx"
            failure = word_region._RenderAudit(False, "仍然存在异常分页")

            with (
                patch.object(
                    word_region,
                    "_render_source_for_regions",
                    return_value=root / "source-render.pdf",
                ),
                patch.object(
                    word_region,
                    "_rasterized_source_pages",
                    return_value=[],
                ),
                patch.object(
                    word_region,
                    "_render_audit",
                    side_effect=(failure, failure),
                ),
            ):
                with self.assertRaisesRegex(ValidationError, "二重检查未通过"):
                    word_compat.optimize_word_full_compatibility(
                        source,
                        output,
                        verification_engine="wps",
                    )


if __name__ == "__main__":
    unittest.main()
