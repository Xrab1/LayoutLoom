from __future__ import annotations

from typing import Any

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree

from docuforge.models import ValidationError
from docuforge.processors import word_region


_THEME_RELATIONSHIP = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme"
)


def _single_line_region(
    paragraph: Any,
    *,
    text: str | None = None,
    font_size: float = 10.0,
) -> tuple[Any, Any]:
    visible_text = paragraph.text if text is None else text
    frame = word_region._FrameLine(
        paragraph,
        100,
        100,
        1200,
        max(220, round(font_size * 24)),
        visible_text,
        font_size,
        False,
    )
    line = word_region._MergedLine(
        [frame],
        frame.x,
        frame.y,
        frame.width,
        frame.height,
        visible_text,
        font_size,
        False,
        "main",
    )
    block = word_region._TextBlock(
        [line],
        frame.x,
        frame.y,
        frame.width,
        frame.height,
        "main",
        font_size,
        False,
        False,
    )
    region = word_region._TextRegion(
        [block],
        frame.x,
        frame.y,
        frame.width,
        frame.height,
        "main",
        False,
    )
    return block, region


def _append_single_line(paragraph: Any, *, text: str | None = None) -> tuple[Any, Any]:
    block, region = _single_line_region(paragraph, text=text)
    target = Document()
    element = word_region._append_region_paragraph(
        target,
        block,
        region,
        previous_block=None,
        word_lexicon=set(),
        normalize_text=False,
    )
    return target, element


def _append_nested_text(paragraph: Any, wrapper: str, text: str) -> None:
    run = OxmlElement("w:r")
    value = OxmlElement("w:t")
    value.text = text
    run.append(value)
    if wrapper == "hyperlink":
        container = OxmlElement("w:hyperlink")
        container.append(run)
        paragraph._p.append(container)
        return
    if wrapper == "sdt":
        container = OxmlElement("w:sdt")
        container.append(OxmlElement("w:sdtPr"))
        content = OxmlElement("w:sdtContent")
        content.append(run)
        container.append(content)
        paragraph._p.append(container)
        return
    raise AssertionError(f"unknown wrapper: {wrapper}")


@pytest.mark.parametrize("wrapper", ["hyperlink", "sdt"])
def test_nested_run_content_forces_safe_fallback(wrapper: str) -> None:
    source = Document()
    paragraph = source.add_paragraph("A")
    _append_nested_text(paragraph, wrapper, "NESTED")
    block, region = _single_line_region(paragraph, text="ANESTED")

    with pytest.raises(ValidationError):
        word_region._append_region_paragraph(
            Document(),
            block,
            region,
            previous_block=None,
            word_lexicon=set(),
            normalize_text=False,
        )


def _direct_font_matches(element: Any, name: str, size_pt: float) -> bool:
    expected_size = str(round(size_pt * 2))
    for run_properties in element.xpath("./w:r/w:rPr"):
        fonts = run_properties.find(qn("w:rFonts"))
        size = run_properties.find(qn("w:sz"))
        if fonts is None or size is None:
            continue
        names = {
            fonts.get(qn("w:ascii")),
            fonts.get(qn("w:hAnsi")),
            fonts.get(qn("w:eastAsia")),
            fonts.get(qn("w:cs")),
        }
        if name in names and size.get(qn("w:val")) == expected_size:
            return True
    return False


def _target_style_matches(
    document: Any,
    style_id: str,
    name: str,
    size_pt: float,
) -> bool:
    for style in document.styles:
        if style.style_id != style_id:
            continue
        size = style.font.size
        return bool(
            style.font.name == name
            and size is not None
            and abs(size.pt - size_pt) < 0.01
        )
    return False


def test_character_style_dependency_is_rejected_or_preserved() -> None:
    source = Document()
    style = source.styles.add_style("SafetyChar", WD_STYLE_TYPE.CHARACTER)
    style.font.name = "Courier New"
    style.font.size = Pt(18)
    paragraph = source.add_paragraph()
    run = paragraph.add_run("Styled")
    run.style = style

    try:
        target, element = _append_single_line(paragraph)
    except ValidationError:
        return

    style_ids = element.xpath("./w:r/w:rPr/w:rStyle/@w:val")
    copied_style = bool(
        style_ids
        and _target_style_matches(target, style_ids[0], "Courier New", 18)
    )
    assert _direct_font_matches(element, "Courier New", 18) or copied_style


def test_paragraph_style_dependency_is_rejected_or_preserved() -> None:
    source = Document()
    style = source.styles.add_style("SafetyParagraph", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Courier New"
    style.font.size = Pt(18)
    paragraph = source.add_paragraph("Styled", style=style)

    try:
        target, element = _append_single_line(paragraph)
    except ValidationError:
        return

    style_ids = element.xpath("./w:pPr/w:pStyle/@w:val")
    copied_style = bool(
        style_ids
        and _target_style_matches(target, style_ids[0], "Courier New", 18)
    )
    assert _direct_font_matches(element, "Courier New", 18) or copied_style


def _theme_part(document: Any) -> Any:
    for relationship in document.part.rels.values():
        if relationship.reltype == _THEME_RELATIONSHIP:
            return relationship.target_part
    raise AssertionError("document has no theme part")


def test_theme_font_dependency_is_rejected_resolved_or_copied() -> None:
    source = Document()
    source_theme = _theme_part(source)
    theme = etree.fromstring(source_theme.blob)
    namespaces = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main"
    }
    major_latin = theme.xpath(
        ".//a:themeElements/a:fontScheme/a:majorFont/a:latin",
        namespaces=namespaces,
    )[0]
    major_latin.set("typeface", "Safety Theme Font")
    source_theme._blob = etree.tostring(
        theme,
        xml_declaration=True,
        encoding="UTF-8",
        standalone=True,
    )

    paragraph = source.add_paragraph()
    run = paragraph.add_run("Theme styled")
    run.font.size = Pt(18)
    fonts = run._r.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:asciiTheme"), "majorHAnsi")
    fonts.set(qn("w:hAnsiTheme"), "majorHAnsi")

    try:
        target, element = _append_single_line(paragraph)
    except ValidationError:
        return

    explicit = _direct_font_matches(element, "Safety Theme Font", 18)
    copied_theme = _theme_part(target).blob == source_theme.blob
    assert explicit or copied_theme


def _fragment_line(
    document: Any,
    *,
    y: int,
    target_x: int,
    left_text: str,
    right_text: str,
) -> Any:
    left_paragraph = document.add_paragraph(left_text)
    right_paragraph = document.add_paragraph(right_text)
    left = word_region._FrameLine(
        left_paragraph, 100, y, 100, 220, left_text, 10.0, False
    )
    right = word_region._FrameLine(
        right_paragraph, target_x, y, 100, 220, right_text, 10.0, False
    )
    return word_region._MergedLine(
        [left, right],
        100,
        y,
        target_x,
        220,
        f"{left_text} {right_text}",
        10.0,
        False,
        "main",
    )


def _paragraph_line_segments(paragraph: Any) -> list[tuple[str, int]]:
    segments: list[list[Any]] = [["", 0]]
    for run in paragraph.xpath("./w:r"):
        for node in run:
            local_name = str(node.tag).rsplit("}", 1)[-1]
            if local_name == "rPr":
                continue
            if local_name == "t":
                segments[-1][0] += str(node.text or "")
            elif local_name == "tab":
                segments[-1][1] += 1
            elif local_name in {"br", "cr"}:
                segments.append(["", 0])
    return [(str(text), int(tabs)) for text, tabs in segments]


def test_multiline_different_tab_targets_are_unambiguous() -> None:
    source = Document()
    lines = [
        _fragment_line(
            source, y=100, target_x=750, left_text="A", right_text="B"
        ),
        _fragment_line(
            source, y=340, target_x=900, left_text="C", right_text="D"
        ),
    ]
    blocks = word_region._paragraph_blocks(lines, 12240)
    region = word_region._region_from_blocks(blocks, 12240, 15840)
    page = word_region._RegionPage(0, 12240, 15840, [region], [])
    document = word_region._build_document(
        [page], "A B C D", font_scale=100000, normalize_text=False
    )

    expected = {
        "AB": 750 - region.x,
        "CD": 900 - region.x,
    }
    observed: dict[str, int] = {}
    for paragraph in document.element.body.xpath(".//w:txbxContent/w:p"):
        stops = sorted(
            int(node.get(qn("w:pos"), "0"))
            for node in paragraph.xpath("./w:pPr/w:tabs/w:tab")
        )
        for text, tab_count in _paragraph_line_segments(paragraph):
            compact = "".join(text.split())
            for marker in expected:
                if marker not in compact:
                    continue
                assert 0 < tab_count <= len(stops)
                observed[marker] = stops[tab_count - 1]

    assert observed == expected


def test_mixed_font_row_exact_spacing_covers_largest_font() -> None:
    source = Document()
    label_paragraph = source.add_paragraph()
    label_run = label_paragraph.add_run("X")
    label_run.font.size = Pt(20)
    values_paragraph = source.add_paragraph()
    values_run = values_paragraph.add_run("12345678901234567890")
    values_run.font.size = Pt(10)

    label = word_region._FrameLine(
        label_paragraph, 100, 100, 300, 420, "X", 20.0, False
    )
    values = word_region._FrameLine(
        values_paragraph,
        900,
        100,
        1800,
        220,
        values_run.text,
        10.0,
        False,
    )
    left = word_region._MergedLine(
        [label], 100, 100, 300, 420, "X", 20.0, False, "main"
    )
    right = word_region._MergedLine(
        [values], 900, 100, 1800, 220, values_run.text, 10.0, False, "main"
    )
    merged = word_region._coalesce_same_row_fragments([left, right], 12240)
    assert len(merged) == 1
    block = word_region._block_from_lines(merged, 10.0)
    region = word_region._region_from_blocks([block], 12240, 15840)
    paragraph = word_region._append_region_paragraph(
        Document(),
        block,
        region,
        previous_block=None,
        word_lexicon=set(),
        normalize_text=False,
    )

    spacing = paragraph.xpath("./w:pPr/w:spacing")[0]
    if spacing.get(qn("w:lineRule")) == "exact":
        assert int(spacing.get(qn("w:line"), "0")) >= 20 * 20


@pytest.mark.parametrize(
    ("property_name", "property_value"),
    [
        ("w:vanish", None),
        ("w:rtl", None),
        ("w:vertAlign", "superscript"),
    ],
)
def test_dangerous_run_properties_do_not_carry_folded_soft_break(
    property_name: str,
    property_value: str | None,
) -> None:
    source = Document()

    def make_line(text: str, y: int) -> Any:
        paragraph = source.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        dangerous = OxmlElement(property_name)
        if property_value is not None:
            dangerous.set(qn("w:val"), property_value)
        run._r.get_or_add_rPr().append(dangerous)
        frame = word_region._FrameLine(
            paragraph, 100, y, 600, 220, text, 10.0, False
        )
        return word_region._MergedLine(
            [frame], 100, y, 600, 220, text, 10.0, False, "main"
        )

    lines = [make_line("first", 100), make_line("second", 340)]
    block = word_region._TextBlock(
        lines, 100, 100, 600, 460, "main", 10.0, False, False
    )
    region = word_region._TextRegion(
        [block], 100, 100, 600, 460, "main", False
    )

    try:
        paragraph = word_region._append_region_paragraph(
            Document(),
            block,
            region,
            previous_block=None,
            word_lexicon=set(),
            normalize_text=False,
        )
    except ValidationError:
        return

    assert len(paragraph.xpath("./w:r/w:br")) == 1
    assert not paragraph.xpath(
        f"./w:r[w:br and w:rPr/{property_name}]"
    )
